#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Alinea FINAL_COMPLETA.docx a docs/informedetesis.txt (secciones faltantes).

Aplica:
- Cap. 2: Estado del arte + Trabajos relacionados (titulos/contenido)
- Cap. 3: Herramientas explicitas en 3.5
- Cap. 5: subtitulos Experimentos / Metricas si faltan
- Cap. 6: regenera conclusiones preliminares (hallazgos, limitaciones,
  trabajo pendiente, plan de culminacion)
- Sincroniza copias canónicas del informe final
"""

from __future__ import annotations

import json
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "scripts"))

from generate_borrador_tesis_docx import (  # noqa: E402
    add_table,
    bullet,
    heading,
    p,
    style_base,
)
from thesis_doctoral_sections import (  # noqa: E402
    RUN_ID,
    add_chapter_6_doctoral,
    verify_doctoral_docx,
)

OUT = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx"
COPIES = [
    REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx",
    REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_skill.docx",
    REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_50EP_ANTECEDENTES.docx",
]
AUDIT_OUT = (
    REPO
    / "outputs"
    / "_drive_madrl"
    / "full_data"
    / "analysis_real_drive"
    / "informe_final_structure_completion_2026-07-14.json"
)


def _para_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def _body_children(doc: Document) -> list:
    return [c for c in doc.element.body if c.tag != qn("w:sectPr")]


def _find_para(doc: Document, predicate) -> Paragraph | None:
    for para in doc.paragraphs:
        if predicate(para.text.strip()):
            return para
    return None


def _set_heading_text(para: Paragraph, text: str) -> None:
    # Preserve heading style; replace visible text.
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _insert_after(anchor: Paragraph, new_elements: list) -> None:
    el = anchor._element
    for new_el in reversed(new_elements):
        el.addnext(new_el)


def _make_temp_paras(factory) -> list:
    tmp = Document()
    style_base(tmp)
    factory(tmp)
    return [deepcopy(c) for c in _body_children(tmp)]


def _replace_range(doc: Document, start_prefix: str, end_prefix: str | None, new_elements: list) -> None:
    body = doc.element.body
    children = [c for c in body if c.tag != qn("w:sectPr")]
    sect_pr = body.find(qn("w:sectPr"))

    start = None
    end = len(children)
    for i, child in enumerate(children):
        txt = _para_text(child)
        if start is None and txt.startswith(start_prefix):
            start = i
            continue
        if start is not None and end_prefix and txt.startswith(end_prefix):
            end = i
            break
    if start is None:
        raise RuntimeError(f"No se encontro inicio: {start_prefix!r}")

    rebuilt = children[:start] + new_elements + children[end:]
    for child in list(children):
        body.remove(child)
    for child in rebuilt:
        if sect_pr is not None:
            body.insert(body.index(sect_pr), child)
        else:
            body.append(child)


def patch_cap2(doc: Document) -> list[str]:
    actions: list[str] = []
    full = "\n".join(p.text for p in doc.paragraphs)

    # Rename Antecedentes -> Trabajos relacionados (estado del arte synthesised separately)
    ante = _find_para(
        doc,
        lambda t: t.startswith("2.4 Antecedentes")
        or t.startswith("2.5 Trabajos relacionados")
        or t.lower().startswith("2.4 trabajos relacionados")
        or t.lower().startswith("2.5 trabajos relacionados"),
    )
    if ante and "Trabajos relacionados" not in ante.text:
        # Will be renumbered below if Estado del arte is inserted.
        _set_heading_text(ante, ante.text.replace("Antecedentes", "Trabajos relacionados y antecedentes", 1))
        actions.append(f"Renombrado antecedentes -> trabajos relacionados: {ante.text}")
    elif ante:
        actions.append(f"2.x ya incluye Trabajos relacionados: {ante.text}")

    # Insert Estado del arte after Capitulo 2 if missing as heading
    if "Estado del arte actualizado" not in full:
        cap2 = _find_para(doc, lambda t: t.startswith("Capitulo 2"))
        if cap2 is None:
            raise RuntimeError("Capitulo 2 no encontrado")

        def _estado(tmp):
            heading(tmp, "2.1 Estado del arte actualizado", 2)
            p(
                tmp,
                "El estado del arte se organiza en cuatro ejes alineados con los objetivos "
                "doctorales (OE.1 flexibilidad, OE.2 emisiones de CO2, OE.3 costos energeticos) "
                "y con el marco tecnico MADRL. CityLearn estandariza el control multiedificio "
                "con KPIs comparables de pico, rampa y factor de carga (Vazquez-Canteli y Nagy, "
                "2019; Vazquez-Canteli et al., 2020). CityLearn v2 incorpora EV/V2G, intensidad "
                "de carbono dinamica y comunidades grid-interactive (Nweye et al., 2024), "
                "mientras que trabajos de escala similar a 17 edificios confirman la "
                "pertinencia del caso SEAI Iquitos (Nweye et al., 2023a; Nweye et al., 2023b).",
            )
            p(
                tmp,
                "En flexibilidad y coordinacion, el MARL cooperativo y mecanismos de atencion "
                "reportan mejoras frente a agentes independientes (Yao et al., 2023; Xie et al., "
                "2023; Hribar et al., 2025). En carbono, enfoques multiobjetivo y restricciones "
                "seguras reducen emisiones y costo en redes/microredes (Liu et al., 2022; Ye et al., "
                "2025; Ma et al., 2025; Sarkar et al., 2024). En costos, MADDPG/MASAC y control "
                "TOU-BESS muestran ahorros tarifarios bajo senales dinamicas (Yao et al., 2023; "
                "Gao et al., 2023; Xiong et al., 2024; Kim et al., 2025).",
            )
            p(
                tmp,
                "El marco tecnico se sustenta en Dec-POMDP (Oliehoek y Amato, 2016), CTDE "
                "(Lowe et al., 2017), HAPPO (Kuba et al., 2021), SAC/MASAC (Haarnoja et al., 2018; "
                "Gao et al., 2023), MATD3 (Fujimoto et al., 2018) y MAAC (Iqbal y Sha, 2019), con "
                "soporte de bibliotecas unificadas (Hu et al., 2023) y HPO (Akiba et al., 2019). "
                "En Peru, el SEAI Iquitos concentra redes aisladas diesel-PV con CI y tarifas "
                "reguladas (MINAM, 2019; OSINERGMIN, 2024); antecedentes doctorales UNI y regionales "
                "abordan generacion, PV hibrido e IA, y microredes (Chevarria Moscoso, 2024; "
                "Penalva Sanchez, 2024; Rosero Bernal, 2024; Dominguez Barbero, 2026). La brecha "
                "persistente es la ausencia de un benchmark unificado HAPPO/MASAC/MATD3/MAAC "
                "sobre dataset real de Iquitos bajo los tres ejes OE.",
            )
            p(
                tmp,
                "Las subsecciones siguientes desarrollan los fundamentos matematicos, las bases "
                "teoricas por eje y la sintesis critica de trabajos relacionados que operacionalizan "
                "esta lectura del estado del arte.",
            )

        elems = _make_temp_paras(_estado)
        _insert_after(cap2, elems)
        actions.append("Insertado 2.1 Estado del arte actualizado (contenido sustentado)")

        # Renumber former Cap.2 Heading-2 sections to keep a linear outline.
        renames = [
            ("2.1 Fundamentos", "2.2 Fundamentos teoricos y matematicos"),
            ("2.2 Variables de la investigacion", "2.3 Variables de la investigacion"),
            ("2.3 Bases teoricas por eje", "2.4 Bases teoricas por eje"),
            ("2.4 Trabajos relacionados y antecedentes", "2.5 Trabajos relacionados y antecedentes"),
            ("2.4 Antecedentes", "2.5 Trabajos relacionados y antecedentes"),
            ("2.5 Definicion de terminos", "2.6 Definicion de terminos y posicion teorica"),
            ("2.5 Definicion de terminos y posicion teorica", "2.6 Definicion de terminos y posicion teorica"),
        ]
        # Apply from high numbers first to avoid cascading collisions.
        for old, new in sorted(renames, key=lambda x: -float(x[0].split()[0])):
            para = _find_para(doc, lambda t, o=old: t.startswith(o))
            if para and para.text.strip() != new:
                _set_heading_text(para, new)
                actions.append(f"Renumberto H2: {old} -> {new}")
    else:
        actions.append("Estado del arte ya presente")

    bases = _find_para(doc, lambda t: "Bases teoricas" in t and t.startswith("2."))
    if bases:
        actions.append(f"Bases teoricas presentes: {bases.text}")

    return actions


def patch_cap3(doc: Document) -> list[str]:
    actions: list[str] = []
    tec = _find_para(
        doc,
        lambda t: t.startswith("3.5") and ("Tecnica" in t or "Herramient" in t),
    )
    if tec is None:
        actions.append("AVISO: no se encontro 3.5")
        return actions

    if "herramient" not in tec.text.lower():
        _set_heading_text(tec, "3.5 Tecnicas, herramientas e instrumentos")
        actions.append("Renombrado 3.5 para incluir Herramientas")

        def _tools(tmp):
            p(tmp, "Herramientas e instrumentos (operativos):", bold=True)
            bullet(
                tmp,
                "Simulacion CityLearn v2 / CityLearn v3 propuesto; backends HAPPO, MASAC, "
                "MATD3 y MAAC; Optuna; Python 3.9 + PyTorch/CUDA.",
            )
            bullet(
                tmp,
                "Dataset citylearn_iquitos_2023_2025 y artefactos de la corrida canonica "
                f"{RUN_ID} (KPIs, figuras Drive, pruebas no parametricas).",
            )

        _insert_after(tec, _make_temp_paras(_tools))
        actions.append("Insertado bloque explicitos de Herramientas bajo 3.5")
    else:
        actions.append("3.5 ya menciona herramientas")
    return actions


def patch_cap5(doc: Document) -> list[str]:
    actions: list[str] = []
    full_h2 = [
        para.text.strip()
        for para in doc.paragraphs
        if (para.style and para.style.name.startswith("Heading"))
        and para.text.strip().startswith("5.")
    ]
    has_exp = any("Experimentos realizados" in t for t in full_h2)
    has_met = any("Metricas utilizadas" in t or "Métricas utilizadas" in t for t in full_h2)

    cap5 = _find_para(doc, lambda t: t.startswith("Capitulo 5"))
    if cap5 is None:
        return ["AVISO: Capitulo 5 no encontrado"]

    # Rename 5.1 to make Experimentos explicit when needed
    h51 = _find_para(doc, lambda t: t.startswith("5.1 "))
    if h51 and not has_exp:
        if "cobertura experimental" in h51.text.lower() or "marco de contrast" in h51.text.lower():
            _set_heading_text(
                h51,
                "5.1 Experimentos realizados y cobertura experimental",
            )
            actions.append("Renombrado 5.1 -> Experimentos realizados y cobertura experimental")
            has_exp = True

    if not has_met and h51 is not None:
        def _met(tmp):
            heading(tmp, "5.1.1 Metricas utilizadas", 3)
            p(
                tmp,
                "Las metricas utilizadas provienen de CityLearn v2 (evaluate_v2) y del "
                "agregado doctoral: flex_composite / peak / ramping (OE.1), "
                "carbon_emissions_delta_kgco2 (OE.2), electricity_cost_delta_eur (OE.3), "
                "scores normalizados por eje, score global min-max, tasas de exito EV y "
                "pruebas no parametricas (Shapiro-Wilk, Kruskal-Wallis, Mann-Whitney U, "
                "Wilcoxon). La fuente canonica es best_madrl_report.json y los CSV de "
                f"estadistica de la corrida {RUN_ID}.",
            )

        _insert_after(h51, _make_temp_paras(_met))
        actions.append("Insertado 5.1.1 Metricas utilizadas")
    else:
        actions.append("Metricas/Experimentos Cap.5 ya cubiertos o insertados")

    return actions


def patch_cap6(doc: Document) -> list[str]:
    new_cap6 = _make_temp_paras(
        lambda tmp: add_chapter_6_doctoral(tmp, p, heading, bullet, add_table)
    )
    _replace_range(doc, "Capitulo 6", "Referencias bibliograficas", new_cap6)
    return [
        "Regenerado Capitulo 6 (Principales hallazgos, Limitaciones encontradas, "
        "Trabajo pendiente, Plan para culminar la tesis)"
    ]


def main() -> int:
    if not OUT.is_file():
        print(f"FALTA: {OUT}")
        return 1

    # Backup lightly
    bak = OUT.with_suffix(".docx.bak_structure_2026-07-14")
    if not bak.exists():
        shutil.copyfile(OUT, bak)

    doc = Document(str(OUT))
    actions = []
    actions.extend(patch_cap2(doc))
    actions.extend(patch_cap3(doc))
    actions.extend(patch_cap5(doc))
    actions.extend(patch_cap6(doc))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    for copy in COPIES:
        copy.write_bytes(OUT.read_bytes())
        actions.append(f"Sincronizado {copy.name}")

    checks = verify_doctoral_docx(OUT)
    report = {
        "docx": str(OUT),
        "run_id": RUN_ID,
        "actions": actions,
        "verification": checks,
    }
    AUDIT_OUT.parent.mkdir(parents=True, exist_ok=True)
    AUDIT_OUT.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0 if checks.get("complete") else 2


if __name__ == "__main__":
    raise SystemExit(main())
