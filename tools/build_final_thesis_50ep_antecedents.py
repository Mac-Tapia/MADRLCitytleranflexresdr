from __future__ import annotations

import csv
import json
import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
SRC = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_resultados_drive_integrados_ordenado_con_diagramas_marco_teorico_doctoral_sustentado.docx"
OUT = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_50EP_ANTECEDENTES.docx"
CSV = REPO / "outputs" / "_drive_madrl" / "full_data" / "analysis_real_drive" / "tables" / "district_summary_by_algorithm_scenario.csv"
METRICS = REPO / "outputs" / "_drive_madrl" / "full_data" / "analysis_real_drive" / "thesis_docx_version_final_50ep_antecedentes_metrics.json"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)


def text_of(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def clear_body_keep_sectpr(document: Document) -> None:
    body = document.element.body
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def append_before_sectpr(document: Document, el) -> None:
    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(el)
    else:
        body.insert(body.index(sect_pr), el)


def style_doc(document: Document) -> None:
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        if name not in [s.name for s in document.styles]:
            continue
        st = document.styles[name]
        st.font.name = "Calibri"
        if name == "Normal":
            st.font.size = Pt(11)
            st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.line_spacing = 1.15
        else:
            st.font.bold = True
            st.font.color.rgb = ACCENT
            st.font.size = Pt(16 if name == "Heading 1" else 13 if name == "Heading 2" else 11.5)


def set_bg(cell, color: str = "1F4E79") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_para(doc: Document, text: str):
    para = doc.add_paragraph()
    para.add_run(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    return para


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], font_size: float = 7.4):
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GREY
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(head)
        rr.bold = True
        rr.font.size = Pt(font_size)
        rr.font.color.rgb = RGBColor(255, 255, 255)
        set_bg(cell)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(font_size)
    doc.add_paragraph()
    return tbl


def fmt(x: str, nd: int = 3) -> str:
    try:
        return f"{float(x):,.{nd}f}"
    except Exception:
        return str(x)


def read_results_rows() -> list[list[str]]:
    with CSV.open(newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    out: list[list[str]] = []
    for r in sorted(rows, key=lambda x: (x["algorithm"], x["scenario"])):
        status = "50 episodios completos" if int(float(r["episodes"])) == 50 else f"{r['episodes']} episodios reales; no completar artificialmente"
        out.append([
            r["algorithm"],
            r["scenario"],
            r["episodes"],
            r["steps_total"],
            fmt(r["reward_mean"], 6),
            fmt(r["district_net_kwh_mean"], 2),
            fmt(r["district_cost_mean"], 2),
            fmt(r["district_emission_mean"], 2),
            status,
        ])
    return out


def add_antecedents_block(doc: Document) -> None:
    doc.add_heading("2.5 Antecedentes internacionales y nacionales usados en la version final", level=2)
    add_para(doc, "Para cerrar el marco teorico, se incorporan antecedentes internacionales y nacionales/contextuales. Los internacionales provienen de la lista APA del proyecto y sustentan CityLearn, MADRL, flexibilidad, CO2 y costos. Los nacionales sustentan el contexto peruano de emisiones, tarifas, red aislada, vehiculos electricos y evidencia local de aprendizaje automatico aplicado al sector energetico. Las fuentes nacionales no se usan para inventar resultados MADRL; se emplean para justificar pertinencia del caso SEAI Iquitos y condiciones de frontera.")
    rows = [
        ["Internacional", "Vazquez-Canteli & Nagy (2019a)", "CityLearn v1.0 para respuesta a la demanda multiedificio.", "Base historica del entorno de simulacion y KPIs de flexibilidad."],
        ["Internacional", "Nweye et al. (2024)", "CityLearn v2 integra EV, BESS, PV, carbono y comunidades grid-interactive.", "Base oficial sobre la que se construye CityLearn v3 propuesto."],
        ["Internacional", "Yao et al. (2023)", "MADRL para gestion energetica de comunidades con PV/BESS/EV.", "Antecedente directo para OE.1 y OE.3."],
        ["Internacional", "Liu et al. (2022)", "MADRL en edificios con renovables para costo y CO2.", "Antecedente directo para OE.2 y multiobjetivo."],
        ["Internacional", "Iqbal & Sha (2019)", "Actor-Attention-Critic para coordinacion multiagente.", "Fundamento teorico de MAAC y coordinacion selectiva."],
        ["Nacional", "MINAM (2019)", "Factor de emision RAGEI para energia en Peru.", "Sustenta CI base 0,790 kgCO2/kWh del SEAI Iquitos."],
        ["Nacional", "OSINERGMIN (2024)", "Tarifas de distribucion y demanda maxima para Electro Oriente.", "Sustenta costos, TOU y KPI de pico facturable."],
        ["Nacional", "Electro Oriente S.A. / dataset del proyecto", "Facturacion y perfiles de edificios reales del SEAI Iquitos.", "Sustenta demanda, edificios y caso aplicado nacional."],
        ["Nacional", "Juarez Valles (2024), registro ALICIA/CONCYTEC", "Aprendizaje por refuerzo y LSTM para eficiencia de subestacion de 215 kVA.", "Antecedente peruano de IA/RL aplicada a eficiencia electrica; requiere ficha APA completa antes de deposito final."],
        ["Nacional", "Prociencia (2024), proyecto PUCP 94714", "IA e ISO 15118 para gestion inteligente de estaciones de recarga EV con PV y baterias de segundo uso.", "Antecedente nacional de I+D en EV, PV, baterias e IA aplicada."],
    ]
    add_table(doc, "Tabla 2.4. Cinco antecedentes internacionales y cinco nacionales/contextuales.", ["Tipo", "Fuente", "Aporte", "Uso en la tesis"], rows, 7.2)


def add_50ep_results_block(doc: Document) -> None:
    doc.add_heading("5.10 Resultados reales de entrenamiento por MADRL y cobertura de 50 episodios", level=2)
    add_para(doc, "La version final reporta la cobertura real de episodios observada en los artefactos descargados desde Google Drive. MAAC, MASAC y MATD3 tienen 50 episodios por escenario E1/E2/E3. HAPPO contiene 49 episodios reales por escenario en timeseries.csv y trace.csv; no se imputa ni se declara un episodio 50 inexistente. Por tanto, HAPPO se incluye en resultados distritales y de traza, pero queda excluido de comparaciones por edificio y checkpoints por ausencia de building_kpis.csv, building_behavior_summary.csv y checkpoint_manifest.json.")
    add_table(
        doc,
        "Tabla 5.10. Cobertura y KPIs distritales reales por algoritmo y escenario.",
        ["MADRL", "Esc.", "Epis.", "Pasos", "Reward medio", "Net kWh medio", "Costo medio", "CO2 medio", "Estado"],
        read_results_rows(),
        6.8,
    )
    add_para(doc, "Interpretacion directa de la tabla: los resultados de 50 episodios completos existen para MAAC, MASAC y MATD3 en los tres escenarios. HAPPO alcanza 49 episodios; reportarlo como 50 seria una alteracion de la evidencia. En lectura descriptiva distrital, el reporte real identifica el mayor reward medio en HAPPO-E2, el menor costo distrital medio en HAPPO-E2 y las menores emisiones distritales medias en HAPPO-E3. Sin embargo, estas conclusiones para HAPPO son distritales y no deben extenderse a KPIs por edificio ni a checkpoints.")
    add_para(doc, "Para comparaciones completas por edificio y reproducibilidad de modelos, la evidencia cerrada corresponde a MAAC, MASAC y MATD3. Las tablas building_kpis_all.csv, building_behavior_summary_all.csv, checkpoint_summary.csv y controlled_uncontrolled_equipment_by_building.csv preservan 17 edificios por corrida completa y 153 filas de equipamiento controlado/no controlado. Las figuras A.1-A.9 muestran la evidencia grafica generada desde timeseries.csv, trace.csv, building_kpis.csv y checkpoint_manifest.json reales.")


def insert_before_heading(doc: Document, heading_text: str, builder) -> None:
    children = list(doc.element.body)
    idx = None
    for i, el in enumerate(children):
        if text_of(el).startswith(heading_text):
            idx = i
            break
    if idx is None:
        raise RuntimeError(f"No se encontro heading destino: {heading_text}")
    before = [deepcopy(el) for el in children[:idx]]
    after = [deepcopy(el) for el in children[idx:] if el.tag != qn("w:sectPr")]
    clear_body_keep_sectpr(doc)
    for el in before:
        append_before_sectpr(doc, el)
    builder(doc)
    for el in after:
        append_before_sectpr(doc, el)


def main() -> None:
    shutil.copyfile(SRC, OUT)
    doc = Document(OUT)
    style_doc(doc)
    insert_before_heading(doc, "Capitulo 3. Metodologia", add_antecedents_block)
    insert_before_heading(doc, "Capitulo 6. Conclusiones", add_50ep_results_block)
    doc.save(OUT)

    v = Document(OUT)
    paras = [p.text.strip() for p in v.paragraphs if p.text.strip()]
    full = "\n".join(paras)
    metrics = {
        "output": str(OUT),
        "size_bytes": OUT.stat().st_size,
        "paragraphs_non_empty": len(paras),
        "word_count_estimated": len(re.findall(r"\b[\wáéíóúÁÉÍÓÚñÑüÜ-]+\b", full, re.UNICODE)),
        "tables": len(v.tables),
        "inline_images": len(v.inline_shapes),
        "has_antecedents_5_5": "Tabla 2.4. Cinco antecedentes internacionales y cinco nacionales/contextuales." in full,
        "has_50ep_results_table": "Tabla 5.10. Cobertura y KPIs distritales reales por algoritmo y escenario." in full,
        "happo_not_imputed": "no se imputa ni se declara un episodio 50 inexistente" in full,
        "figures_a_1_a_9": all(f"Figura A.{i}" in full for i in range(1, 10)),
        "figures_b_1_a_9": all(f"Figura B.{i}" in full for i in range(1, 10)),
    }
    METRICS.parent.mkdir(parents=True, exist_ok=True)
    METRICS.write_text(json.dumps(metrics, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(metrics, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
