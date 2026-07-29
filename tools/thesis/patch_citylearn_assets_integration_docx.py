#!/usr/bin/env python3
"""Parche mínimo: integrar activos CityLearn retenidos en Cap. III del Word canónico.

Fuente:
- docs/tesis_capitulos/Capitulo_3_Metodologia.md §3.4.6
- docs/INTEGRACION_CITYLEARN_THESIS_2026-07-29.md

No inventa métricas ni corridas sobre barrios/challenges upstream.
"""
from __future__ import annotations

import json
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

import sys

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, TESIS, INFORME  # noqa: E402

REPORT = DOCS / "CITYLEARN_ASSETS_INTEGRATION_PATCH_REPORT_2026-07-29.json"

MARKER = "arbol citylearn retenido: usados vs disponibles"
ANCHOR_HINTS = (
    "compuertas de validacion",
    "compuertas de validación",
    "check_training_dataset_ready",
    "status=ready",
    "3.4.5",
    "datos utilizados",
    "citylearn_iquitos_2023_2025",
)

HEADING = "3.4.6 Árbol CityLearn retenido: usados vs disponibles (reproducibilidad)"

BODY_PARAS = [
    (
        "El submódulo CityLearn/ conserva, además del dataset Iquitos, el árbol de datos "
        "del ecosistema CityLearn v2 embebido en el fork de la tesis. Ese material no se "
        "elimina (decisión 2026-07-29): sirve de contexto metodológico, integridad de "
        "pruebas del paquete y reproducibilidad offline del simulador. El caso empírico "
        "de contraste de hipótesis (Capítulo 5) sigue siendo exclusivamente "
        "citylearn_iquitos_2023_2025 + capa MADRL v3."
    ),
    (
        "En CityLearn/data/datasets/ se retienen: (i) citylearn_iquitos_2023_2025 "
        "(empírico canónico OE.1–OE.3); (ii) citylearn_challenge_2020–2023 (contexto "
        "histórico de benchmarking comunitario; challenge 2022 y demos alimentan tests "
        "del submódulo); (iii) barrios de referencia quebec_neighborhood_*, "
        "ca_alameda_county_neighborhood, tx_travis_county_neighborhood y "
        "vt_chittenden_county_neighborhood (inventario reproducible; contraste "
        "cualitativo con el SEAI aislado). No se atribuyen KPIs ni pruebas estadísticas "
        "a barrios o challenges upstream en ausencia de artefactos en outputs/."
    ),
    (
        "Los launchers launch_citylearn_v3_iquitos_training.ps1 y "
        "monitor_citylearn_v3_iquitos_training.ps1 se retienen como vía histórica "
        "local; la vía canónica de entrenamiento permanece "
        "launch_citylearn_v3_official_training.ps1 y, para la corrida de 50 episodios, "
        "el protocolo Colab del notebook madrl_citylearn_v3_tutorial.ipynb. Detalle: "
        "docs/INTEGRACION_CITYLEARN_THESIS_2026-07-29.md."
    ),
]


def set_run_font(run, bold: bool = False, size: float = 12.0, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def replace_paragraph_text(p: Paragraph, text: str, *, bold: bool = False) -> None:
    p.clear()
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, bold=bold)
    return new_para


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def chapter_bounds(doc: Document, chapter: int) -> tuple[int | None, int | None]:
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if re.search(rf"^Cap[ií]tulo\s*{chapter}\b", t, re.I):
            start = i
        elif start is not None and re.search(rf"^Cap[ií]tulo\s*{chapter + 1}\b", t, re.I):
            end = i
            break
    return start, end


def already_integrated(doc: Document, start: int | None, end: int | None) -> bool:
    lo = start or 0
    hi = end if end is not None else len(doc.paragraphs)
    for i in range(lo, hi):
        if MARKER in norm(doc.paragraphs[i].text or ""):
            return True
        if "barrios de referencia quebec" in norm(doc.paragraphs[i].text or ""):
            return True
    return False


def find_anchor(doc: Document, start: int | None, end: int | None) -> int | None:
    lo = start or 0
    hi = end if end is not None else len(doc.paragraphs)
    best = None
    for i in range(lo, hi):
        n = norm(doc.paragraphs[i].text or "")
        for hint in ANCHOR_HINTS:
            if hint in n:
                best = i
    return best


def patch_doc(doc: Document) -> list[str]:
    actions: list[str] = []
    start, end = chapter_bounds(doc, 3)
    if start is None:
        actions.append("cap3_not_found")
        return actions

    if already_integrated(doc, start, end):
        actions.append("already_present")
        return actions

    anchor = find_anchor(doc, start, end)
    if anchor is None:
        actions.append("anchor_not_found")
        return actions

    cursor = doc.paragraphs[anchor]
    cursor = insert_paragraph_after(cursor, HEADING, bold=True)
    actions.append(f"insert_heading_after@{anchor}")
    for para in BODY_PARAS:
        cursor = insert_paragraph_after(cursor, para)
    actions.append(f"insert_body_paras={len(BODY_PARAS)}")
    return actions


def manual_insertion_text() -> dict:
    return {
        "where": "Capítulo 3, después de §3.4.5 (compuertas / gates) y antes de Técnicas e instrumentos",
        "heading": HEADING,
        "paragraphs": BODY_PARAS,
        "also_see_md": "docs/tesis_capitulos/Capitulo_3_Metodologia.md §3.4.6",
    }


def process_path(path: Path) -> dict:
    if not path.is_file():
        return {"path": str(path), "ok": False, "error": "missing"}
    doc = Document(str(path))
    actions = patch_doc(doc)
    result = {
        "path": str(path.relative_to(REPO)),
        "actions": actions,
        "saved": False,
        "manual_insertion": None,
    }
    if actions == ["already_present"]:
        result["ok"] = True
        result["note"] = "Texto de integración ya presente; sin cambios."
        return result
    if "anchor_not_found" in actions or "cap3_not_found" in actions:
        result["ok"] = False
        result["manual_insertion"] = manual_insertion_text()
        result["reason"] = "No se encontró ancla segura en Cap. III; usar texto markdown + manual_insertion."
        return result

    # Política 2-Word: backups fuera de docs/; editar solo canónicos.
    backup_dir = REPO / "outputs" / "_word_backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = backup_dir / f"{path.stem}_antes_citylearn_assets_{stamp}.docx"
    shutil.copy2(path, backup)
    doc.save(str(path))
    result["ok"] = True
    result["saved"] = True
    result["backup"] = str(backup.relative_to(REPO))
    return result


def main() -> int:
    targets = [TESIS]
    if INFORME.is_file():
        targets.append(INFORME)

    results = [process_path(p) for p in targets]
    report = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "decision": "RETAIN_AND_INTEGRATE",
        "source_md": "docs/tesis_capitulos/Capitulo_3_Metodologia.md §3.4.6",
        "integration_report": "docs/INTEGRACION_CITYLEARN_THESIS_2026-07-29.md",
        "results": results,
        "claim_boundary": (
            "No KPIs / no Cap.5 results for Quebec/Alameda/Travis/Chittenden/challenges; "
            "empirical case remains Iquitos 2023-2025 + MADRL v3."
        ),
    }
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if all(r.get("ok") for r in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
