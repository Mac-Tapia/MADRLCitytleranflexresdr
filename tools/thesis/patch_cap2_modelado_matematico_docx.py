#!/usr/bin/env python3
"""Parche Cap. II UC3M en los 2 Word canónicos (Tesis + Informe).

NO crea Word nuevos en docs/. Solo edita:
  - docs/Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx
  - docs/Inforne_tesisV4_FINAL_REVISADO_50_EPISODIOS.docx

Backups temporales van a outputs/_word_backups/ (fuera de docs/).
Idempotente: si 2.2.6 ya existe, no duplica.
"""
from __future__ import annotations

import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, INFORME, TESIS  # noqa: E402

REPORT = DOCS / "CAP2_MODELADO_MATEMATICO_PATCH_REPORT_2026-07-29.json"
BACKUP_DIR = REPO / "outputs" / "_word_backups"

SECTIONS: list[tuple[str, list[str]]] = [
    (
        "2.2.5 CityLearn v3 propuesto y distincion UC3M (7 ejes vs 3 ejes)",
        [
            (
                "CityLearn v3 propuesto es la extension experimental de esta tesis sobre "
                "CityLearn v2 (Nweye et al., 2024): conserva fisica y KPIs evaluate_v2, y "
                "agrega el contrato Dec-POMDP/CTDE, objetivos OE.1–OE.3, recompensa "
                "multiobjetivo y adaptadores a backends MADRL en external/. Modulos: "
                "v3/environment.py, v3/objectives.py, v3/config.py, v3/backends.py y "
                "CityLearnV3MADRLRewardFunction (perfil unified_comparable_v4)."
            ),
            (
                "El sustento formal agent-skills/madrl-sustento-doc-capa v3/"
                "madrl-modeladomatematico.md y el paquete uc3m/ axiomatizan un operador "
                "holistico de siete ejes (CO2, costo, flexibilidad, confort, degradacion "
                "BESS, resiliencia, ACS) con BACT y HPHI. Las 12 corridas canónicas "
                "operacionalizan solo OE.1/OE.2/OE.3. Mapeo: UC3M-1↔OE.2, UC3M-2↔OE.3, "
                "UC3M-3↔OE.1; ejes 4–7 son sustento, no evidencia Cap. 5. Parametros: "
                "N=17, gamma=0,9999, T=8 760, r_team=0,70."
            ),
        ],
    ),
    (
        "2.2.6 Formalizacion matematica del Meta-Dec-POMDP UC3M",
        [
            (
                "Fuente: madrl-modeladomatematico.md (adaptado a Cap. 2). El Meta-Dec-POMDP "
                "UC3M es la tupla 11-aria M_UC3M = <I, S, A, O, T, R, Z, gamma, H, b0, Lambda>, "
                "que extiende Oliehoek y Amato (2016) con vector R=(r^(1),...,r^(7)) y "
                "simplex Lambda. En Iquitos N=17, gamma=0,9999, H=8 760. La ejecucion "
                "restringe Lambda al subsimplex OE.1/OE.2/OE.3."
            ),
            (
                "Definicion 2.1 (BACT): B en R^{N x Ka x Kc x Kb} codifica activos, clima y "
                "constructivo (uc3m/env/bact.py). Observaciones parciales: carga, termica, "
                "PV, SoC, EV, tarifas, CI y meteorologia; acciones en [-1,1]^{d_ai} "
                "heterogeneas. Prop. 2.1: S compacto. Lema 2.2: R escalarizada Borel-medible "
                "y acotada. Nucleo T factorizado exogeno (clima/red) x endogeno por edificio."
            ),
        ],
    ),
    (
        "2.2.7 Operador de recompensa holistico, Pareto y convergencia",
        [
            (
                "Definicion 2.2: R_i = -sum_k lambda_k * r_i~(k) con normalizacion a base RBC. "
                "Teorema 2.3: V_pi existe y |V|<=M/(1-gamma). Prop. 2.4: Lipschitz con "
                "L=sum lambda_k L_k. Teorema 2.5: existencia de frontera de Pareto si Pi "
                "compacto (Roijers et al., 2013). Definicion 2.4 (HPHI): hipervolumen 7-D "
                "normalizado (uc3m/reward/hphi.py); esta tesis no ejecuta HPHI 7-D en Cap. 5. "
                "Prop. 2.6: CTDE mitiga no-estacionariedad (Lowe et al., 2017); HAPPO/HATRPO "
                "bajo Zhong et al. (2023) y Kuba et al. (2022)."
            ),
        ],
    ),
    (
        "2.2.8 Modelado fisico-matematico de los ejes operacionales",
        [
            (
                "Ejes ejecutados: UC3M-1/OE.2 (CO2 marginal + desplazamiento EV), "
                "UC3M-2/OE.3 (costo TOU Electro Oriente + potencia), UC3M-3/OE.1 (ramping y "
                "picos vs umbral DSO; KPIs peak/ramping/load_factor). Ejes de sustento "
                "(no Cap. 5): 4 confort adaptativo De Dear-Brager, 5 degradacion "
                "Arrhenius-SEI/Peukert, 6 resiliencia isla (CCI/LOLP), 7 ACS con perdidas UA. "
                "Balance Kirchhoff: P_net = P_load + P_HP + P_EH + P_BESS + P_EV - P_PV - P_wind."
            ),
        ],
    ),
    (
        "2.2.9 Arquitectura MARLlib-CTDE y universalidad algoritmica (sustento)",
        [
            (
                "MARLlib (Hu et al., 2023) y el plugin algoritmico UC3M "
                "P=<Theta, Phi, L_actor, L_critic, U_step, B> sustentan universalidad via "
                "AlgorithmFactory (uc3m/). CTDE: pi_i(a_i|o_i) descentralizado y "
                "Q_centr(s,a) centralizado. HAPPO, MASAC, MATD3 y MAAC son instancias; el "
                "launcher de 12 corridas usa wrappers CityLearn/scripts/ + external/, no "
                "UC3MEnv/MARLlib como canal de entrenamiento."
            ),
        ],
    ),
]


def set_run_font(run, bold: bool = False, size: float = 12.0) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, bold=bold)
    return new_para


def set_heading_style(paragraph: Paragraph, level: int) -> None:
    try:
        paragraph.style = f"Heading {level}"
    except KeyError:
        if paragraph.runs:
            set_run_font(paragraph.runs[0], bold=True, size=12.0)


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def marco_bounds(doc: Document) -> tuple[int | None, int | None]:
    """Cap. 2: 'Capítulo 2...' (Tesis) o 'Marco teórico' (Informe)."""
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if start is None:
            if re.search(r"^Cap[ií]tulo\s*2\b", t, re.I) or norm(t) in {
                "marco teorico",
                "marco teórico",
            }:
                start = i
            continue
        if re.search(r"^Cap[ií]tulo\s*3\b", t, re.I):
            end = i
            break
    return start, end


def already_patched(doc: Document, start: int, end: int | None) -> bool:
    end_i = end if end is not None else len(doc.paragraphs)
    for i in range(start, end_i):
        t = norm(doc.paragraphs[i].text)
        if "2.2.6 formalizacion matematica del meta-dec-pomdp" in t:
            return True
    return False


def find_insert_anchor(doc: Document, start: int, end: int | None) -> int | None:
    """Último párrafo del bloque CityLearn, antes del siguiente Heading."""
    end_i = end if end is not None else len(doc.paragraphs)
    heading = None
    for i in range(start, end_i):
        t = norm(doc.paragraphs[i].text)
        if t.startswith("2.2.4 ") or "citylearn y simulacion multiobjetivo" in t or "citylearn y simulación multiobjetivo" in t:
            heading = i
            break
    if heading is None:
        return None

    j = heading
    while j + 1 < end_i:
        nxt_p = doc.paragraphs[j + 1]
        nxt = norm(nxt_p.text)
        style = nxt_p.style.name if nxt_p.style else ""
        if nxt.startswith("2.2.5 ") or nxt.startswith("2.2.6 "):
            break
        if nxt.startswith("2.2 variables") or nxt.startswith("variables de la investig"):
            break
        if style.startswith("Heading") and j > heading:
            break
        j += 1
    return j


def insert_sections_after(anchor: Paragraph) -> int:
    cursor = anchor
    n = 0
    for title, bodies in SECTIONS:
        h = insert_paragraph_after(cursor, title, bold=True)
        set_heading_style(h, 3)
        cursor = h
        n += 1
        for body in bodies:
            cursor = insert_paragraph_after(cursor, body)
            n += 1
    return n


def patch_doc(path: Path) -> dict:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{path.stem}_before_cap2_{datetime.now().strftime('%Y%m%d_%H%M%S')}{path.suffix}"
    shutil.copy2(path, backup)
    doc = Document(str(path))
    start, end = marco_bounds(doc)
    if start is None:
        return {"path": str(path), "ok": False, "error": "Marco teorico / Cap. 2 no encontrado", "backup": str(backup)}

    if already_patched(doc, start, end):
        backup.unlink(missing_ok=True)
        return {
            "path": str(path),
            "ok": True,
            "skipped": True,
            "reason": "§§2.2.5–2.2.9 ya presentes",
        }

    anchor_i = find_insert_anchor(doc, start, end)
    if anchor_i is None:
        return {
            "path": str(path),
            "ok": False,
            "error": "Ancla CityLearn no encontrada",
            "backup": str(backup),
        }

    inserted = insert_sections_after(doc.paragraphs[anchor_i])
    doc.save(str(path))
    return {
        "path": str(path),
        "ok": True,
        "backup": str(backup),
        "inserted": inserted,
        "anchor_index": anchor_i,
        "sections": [t for t, _ in SECTIONS],
    }


def cleanup_docs_backup_docx() -> list[str]:
    """Política canónica: solo 2 Word en docs/; quitar _backup_*.docx de la raíz."""
    removed: list[str] = []
    for p in DOCS.glob("_backup_*.docx"):
        dest = BACKUP_DIR / p.name
        BACKUP_DIR.mkdir(parents=True, exist_ok=True)
        shutil.move(str(p), str(dest))
        removed.append(str(p.name))
    return removed


def main() -> int:
    targets = [TESIS, INFORME]
    results = []
    for path in targets:
        if not path.is_file():
            results.append({"path": str(path), "ok": False, "error": "missing"})
            continue
        results.append(patch_doc(path))
    moved = cleanup_docs_backup_docx()
    report = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "policy": "solo 2 Word canonicos en docs/; sin generar terceros",
        "canons": [str(TESIS), str(INFORME)],
        "backups_moved_from_docs": moved,
        "backup_dir": str(BACKUP_DIR),
        "results": results,
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0 if all(r.get("ok") for r in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
