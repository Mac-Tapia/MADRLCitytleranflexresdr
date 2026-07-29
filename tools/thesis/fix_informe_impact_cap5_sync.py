#!/usr/bin/env python3
"""Fix Informe: move Cap.1 criteria to §1.3 and resync Cap.5 from Tesis."""
from __future__ import annotations

import json
import re
import shutil
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from sync_cap5_to_canon_words import extract_cap5, replace_cap5  # noqa: E402
from thesis_word_canons import INFORME, TESIS  # noqa: E402

REPORT = REPO / "docs" / "INFORME_IMPACT_CAP5_FIX_REPORT_2026-07-29.json"

CRITERIA_ROWS = [
    ["C1", "Impacto vs baseline", "Wilcoxon KPI-gains vs cero + Holm", "Inferencial HE"],
    ["C2", "Diferencias entre algoritmos", "Kruskal-Wallis / Friedman + Holm", "Inferencial HE"],
    ["C3", "KPIs fisicos de distrito por eje", "flex_composite / delta CO2 / delta costo", "Descriptivo distrito"],
    ["C4", "KPIs desagregados por edificio por eje", "17 edificios x E1/E2/E3", "Descriptivo edificio"],
    ["C5", "Control de recursos", "BESS, EV/V2G, carga desplazable (acciones y exito EV)", "Obligatorio (atribuibilidad)"],
]


def set_run_font(run, bold: bool = False, size: float = 12.0, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except (KeyError, ValueError):
            pass
    if text:
        run = new_para.add_run(text)
        set_run_font(run, bold=("Heading" in (style or "")))
    return new_para


def remove_misplaced_criteria(doc: Document) -> int:
    """Remove 1.4.1 / Tabla 1.3 block wrongly inserted inside Cap. 5."""
    removed = 0
    # Collect paragraphs to delete (by element)
    to_delete = []
    paras = list(doc.paragraphs)
    i = 0
    while i < len(paras):
        t = (paras[i].text or "").strip()
        if t.startswith("1.4.1 Criterios") or (
            t.startswith("Tabla 1.3") and "control de recursos" in t.lower()
        ):
            # delete this and following related paras until Cap 5 heading resumes
            j = i
            while j < len(paras):
                tj = (paras[j].text or "").strip()
                to_delete.append(paras[j]._p)
                j += 1
                if j >= len(paras):
                    break
                nxt = (paras[j].text or "").strip()
                # stop before next real Cap5 heading/table that isn't criteria
                if nxt.startswith("5.") or nxt.startswith("Tabla 5.") or nxt.startswith("Capitulo 5"):
                    break
                if nxt.startswith("1.4.1") or nxt.startswith("Tabla 1.3") or nxt.startswith("Para cumplir") or nxt.startswith("Regla de cumplimiento"):
                    continue
                if not nxt:
                    continue
                # if we hit unrelated content after note, stop
                if not (
                    nxt.startswith("Para cumplir")
                    or nxt.startswith("Regla de cumplimiento")
                    or nxt.startswith("1.4.1")
                    or nxt.startswith("Tabla 1.3")
                ):
                    break
            i = j
            continue
        i += 1
    # Also remove orphan tables that are only criteria (4 cols Id/Criterio...)
    body = doc.element.body
    for el in list(to_delete):
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            removed += 1
    # Remove criteria tables near Cap5: detect by header text
    for tbl in list(doc.tables):
        try:
            header = " | ".join(c.text.strip() for c in tbl.rows[0].cells)
        except Exception:
            continue
        if "Criterio" in header and "Medida" in header and any(
            "Control de recursos" in (cell.text or "") for row in tbl.rows for cell in row.cells
        ):
            # Only remove if table is NOT in Cap1 region: check preceding para
            tbl_el = tbl._tbl
            prev = tbl_el.getprevious()
            prev_text = ""
            while prev is not None:
                if prev.tag == qn("w:p"):
                    prev_text = "".join(t.text or "" for t in prev.iter(qn("w:t")))
                    break
                prev = prev.getprevious()
            if "1.4.1" in prev_text or "Capitulo 5" in prev_text or "5.1" in prev_text or not prev_text:
                # if Cap5 context or orphan after deletion, remove
                if "1.3" not in prev_text and "Objetivo" not in prev_text:
                    parent = tbl_el.getparent()
                    if parent is not None:
                        parent.remove(tbl_el)
                        removed += 1
    return removed


def insert_criteria_after_hypotheses(doc: Document) -> str:
    """Insert Cap.1 criteria after hipótesis específicas."""
    blob = "\n".join(p.text for p in doc.paragraphs)
    if "1.4.1 Criterios de determinacion del impacto" in blob and "1.3.2" in blob:
        # check if already correctly placed near objectives
        for i, p in enumerate(doc.paragraphs):
            if "1.4.1 Criterios" in (p.text or ""):
                if i < 250:
                    return "already_in_cap1"
                break

    anchor = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("1.3.2.2") or (
            "Hipótesis específicas" in t and t.startswith("1.3")
        ):
            anchor = p
        # keep updating through HE paragraphs until Cap 2 / Justificacion section end
        if anchor is not None and (
            t.startswith("1.4 ")
            or t.startswith("Capítulo 2")
            or t.startswith("Capitulo 2")
            or t.startswith("2.")
            and "Marco" in t
        ):
            break
    if anchor is None:
        # fallback: last paragraph under 1.3.2
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t.startswith("HE3") or t.startswith("HE31") or "HE31" in t:
                anchor = p
    if anchor is None:
        return "anchor_not_found"

    # Move past following HE body paragraphs
    paras = list(doc.paragraphs)
    idx = next((i for i, p in enumerate(paras) if p._p is anchor._p), None)
    insert_at = anchor
    if idx is not None:
        for p in paras[idx + 1 : idx + 25]:
            t = (p.text or "").strip()
            if not t:
                continue
            if t.startswith("1.4") or t.startswith("Capítulo 2") or t.startswith("Capitulo 2") or (
                t.startswith("2.") and "Marco" in t
            ):
                break
            if (
                t.startswith("1.3")
                or t.startswith("HE")
                or "hipótesis" in t.lower()
                or "hipotesis" in t.lower()
            ):
                insert_at = p
                continue
            if len(t) > 40:
                insert_at = p

    h = insert_paragraph_after(
        insert_at,
        "1.3.1.3 Criterios de determinacion del impacto",
        style="Heading 3",
    )
    intro = insert_paragraph_after(
        h,
        "Para cumplir el OG y los OE.1–OE.3, y para demostrar las hipotesis, la tesis exige "
        "el conjunto completo de criterios de determinacion del impacto (C1–C5). "
        "Uno de ellos —C5— es el control de recursos energeticos (BESS, EV/V2G y carga "
        "desplazable). Los tres ejes se reportan a nivel de distrito y a nivel de edificio "
        "(evidencia Cap. 5).",
    )
    table = doc.add_table(rows=1 + len(CRITERIA_ROWS), cols=4)
    table.style = "Table Grid"
    headers = ["Id", "Criterio", "Medida / prueba", "Rol"]
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, bold=True, size=10)
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(CRITERIA_ROWS):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=10)
    intro._p.addnext(table._tbl)
    caption_el = OxmlElement("w:p")
    table._tbl.addnext(caption_el)
    caption = Paragraph(caption_el, insert_at._parent)
    run = caption.add_run(
        "Tabla 1.3. Criterios completos de determinacion del impacto "
        "(C5 = control de recursos; evidencia Cap. 5)."
    )
    set_run_font(run, italic=True, size=10)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    insert_paragraph_after(
        caption,
        "Regla de cumplimiento sin parciales: OG/OE.1–OE.3 y la demostracion de hipotesis "
        "exigen C1–C5. Cada eje se reporta a nivel distrito y a nivel edificio. "
        "C5 (control de recursos) es obligatorio para atribuir el impacto a las acciones "
        "MADRL sobre BESS, EV/V2G y carga desplazable.",
    )
    return f"inserted_after:{insert_at.text[:60]}"


def save_safe(doc: Document, path: Path) -> Path:
    """Save via temp file to avoid Windows lock Invalid argument."""
    backup_dir = REPO / "outputs" / "_word_backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    bak = backup_dir / f"{path.name}.pre_informe_fix_{datetime.now().strftime('%Y%m%d_%H%M%S')}.bak"
    if path.is_file():
        shutil.copy2(path, bak)
    tmp = path.with_suffix(".docx.tmp_write")
    if tmp.exists():
        tmp.unlink()
    doc.save(str(tmp))
    # replace
    if path.exists():
        path.unlink()
    tmp.rename(path)
    return bak


def main() -> int:
    report: dict = {"generated_at": datetime.now().isoformat(timespec="seconds")}
    if not INFORME.is_file():
        print(json.dumps({"ok": False, "error": "informe_missing"}))
        return 1

    # 1) Resync Cap5 from Tesis first (removes misplaced Cap.1 block inside Cap.5)
    elements, meta = extract_cap5(TESIS)
    report["cap5_extract"] = meta
    try:
        info = replace_cap5(INFORME, elements, toc_safe=False, dry_run=False)
        report["cap5_sync"] = info
    except OSError as exc:
        tmp_path = INFORME.with_suffix(".docx.resync_tmp")
        shutil.copy2(INFORME, tmp_path)
        info = replace_cap5(tmp_path, elements, toc_safe=False, dry_run=False)
        if tmp_path.is_file():
            if INFORME.exists():
                try:
                    INFORME.unlink()
                except OSError:
                    pass
            tmp_path.replace(INFORME)
        report["cap5_sync"] = info
        report["cap5_sync_note"] = f"retried_via_tmp after {exc}"

    # 2) Cap.1 criteria placement
    doc = Document(str(INFORME))
    removed = remove_misplaced_criteria(doc)
    report["removed_misplaced"] = removed
    placed = insert_criteria_after_hypotheses(doc)
    report["criteria_placement"] = placed
    bak = save_safe(doc, INFORME)
    report["backup"] = str(bak)

    # Verify
    d2 = Document(str(INFORME))
    blob = "\n".join(p.text for p in d2.paragraphs)
    checks = {
        "control_de_recursos": "control de recursos" in blob.lower(),
        "tabla_1_3": "Tabla 1.3" in blob,
        "criterios_cap1": "1.3.1.3 Criterios" in blob or "1.4.1 Criterios" in blob,
        "5_1_1_criterios": "5.1.1 Criterios de determinacion del impacto" in blob,
        "5_4_5_control": "5.4.5 Control de recursos" in blob,
        "inventario_ev": "Inventario EV" in blob or "inventario EV" in blob.lower(),
        "c5": "C5" in blob,
        "misplaced_in_cap5": False,
    }
    seen_cap5 = False
    for p in d2.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Capitulo 5") or t.startswith("Capítulo 5"):
            seen_cap5 = True
        if seen_cap5 and (t.startswith("1.4.1 Criterios") or t.startswith("1.3.1.3 Criterios")):
            checks["misplaced_in_cap5"] = True
            break
    # Cap1 criteria should appear before Cap5
    cap1_idx = None
    cap5_idx = None
    for i, p in enumerate(d2.paragraphs):
        t = (p.text or "").strip()
        if cap1_idx is None and ("1.3.1.3 Criterios" in t or "1.4.1 Criterios" in t):
            cap1_idx = i
        if cap5_idx is None and (t.startswith("Capitulo 5") or t.startswith("Capítulo 5")):
            cap5_idx = i
    checks["criteria_before_cap5"] = (
        cap1_idx is not None and cap5_idx is not None and cap1_idx < cap5_idx
    )
    report["checks"] = checks
    report["ok"] = all(
        [
            checks["control_de_recursos"],
            checks["tabla_1_3"],
            checks["criterios_cap1"],
            checks["5_1_1_criterios"],
            checks["5_4_5_control"],
            checks["c5"],
            not checks["misplaced_in_cap5"],
            checks["criteria_before_cap5"],
        ]
    )
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
