#!/usr/bin/env python3
"""Remisión Dec-POMDP dims en Cap. VI de los 2 Word canónicos.

Añade frase con d_s / d_oi / d_ai / r_team + Caps. 2 y 4 (Tabla 2.A).
Idempotente.
"""
from __future__ import annotations

import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, INFORME, TESIS  # noqa: E402

REPORT = DOCS / "CAP6_DECPOMDP_XREF_PATCH_REPORT_2026-07-29.json"
BACKUP_DIR = REPO / "outputs" / "_word_backups"
MARKER = "Caps. 2 y 4 (Tabla 2.A"
ADD = (
    " La formulacion Dec-POMDP ejecutada (d_s = 1 856; d_oi en [54, 327]; "
    "d_ai en [5, 44]; r_team = 0,70) queda fijada en Caps. 2 y 4 (Tabla 2.A / §4.2)."
)


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def cap6_bounds(doc: Document) -> tuple[int, int]:
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if start is None and (
            re.search(r"^Cap[ií]tulo\s*6\b", t, re.I)
            or t.lower().startswith("6.1 principales")
        ):
            start = i
        if start is not None and (
            re.search(r"^Referencias\b", t, re.I)
            or t.lower().startswith("referencias bibliograficas")
            or t.lower().startswith("referencias bibliográficas")
            or re.search(r"^Anexo\b", t, re.I)
            or re.search(r"^Ap[eé]ndice\b", t, re.I)
        ):
            end = i
            break
    return start or 0, end or len(doc.paragraphs)


def patch_doc(path: Path) -> dict:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / (
        f"{path.stem}_before_cap6_xref_{datetime.now().strftime('%Y%m%d_%H%M%S')}{path.suffix}"
    )
    shutil.copy2(path, backup)
    doc = Document(str(path))
    start, end = cap6_bounds(doc)
    changes: list[str] = []

    # Already done?
    for i in range(start, end):
        if MARKER in (doc.paragraphs[i].text or ""):
            # still patch diagram if needed outside? check full doc for fig
            break
    else:
        pass

    already = any(MARKER in (doc.paragraphs[i].text or "") for i in range(start, end))
    if not already:
        for i in range(start, end):
            p = doc.paragraphs[i]
            t = p.text or ""
            tl = t.lower()
            if ("integracion oe" in tl or "integración oe" in tl) and "matd3" in tl:
                set_paragraph_text(p, t.rstrip().rstrip(".") + "." + ADD)
                changes.append(f"integracion_oe_p{i}")
                break
            if tl.startswith("respuesta a pe.1") and "flex_composite" in tl:
                set_paragraph_text(p, t.rstrip().rstrip(".") + "." + ADD)
                changes.append(f"pe1_p{i}")
                break
            if tl.startswith("og:") and "madrl" in tl:
                set_paragraph_text(p, t.rstrip().rstrip(".") + "." + ADD)
                changes.append(f"og_p{i}")
                break
    else:
        # Cap.6 already has xref; only fig interpretations may still need it
        pass

    # Diagram interpretations anywhere (appendix often after refs)
    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        tl = t.lower()
        if ("interpretacion de la figura" in tl or "interpretación de la figura" in tl) and (
            "54" in t and "327" in t
        ):
            if "tabla 2.a" not in tl and "cap. 2" not in tl and "caps. 2" not in tl:
                set_paragraph_text(
                    p,
                    t.rstrip().rstrip(".")
                    + " (valores alineados a Cap. 2 Tabla 2.A y Cap. 4 §4.2).",
                )
                changes.append(f"fig_p{i}")

    if already and not changes:
        backup.unlink(missing_ok=True)
        return {"path": str(path), "ok": True, "skipped": True, "reason": "already xref"}

    if not changes:
        return {
            "path": str(path),
            "ok": False,
            "error": "no Cap.6 anchor found",
            "backup": str(backup),
            "bounds": [start, end],
        }

    doc.save(str(path))
    return {"path": str(path), "ok": True, "backup": str(backup), "changes": changes}


def main() -> int:
    results = []
    for path in (TESIS, INFORME):
        if path.is_file():
            results.append(patch_doc(path))
    report = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "marker": MARKER,
        "results": results,
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0 if results and all(r.get("ok") for r in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
