#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Sync Capítulo 5 body from Tesis canon into Informe (2 canons).

TOC-safe: skips TOC entries that look like 'Capítulo 5 …89' when needed.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

_THESIS_DIR = Path(__file__).resolve().parent
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, INFORME, TESIS, mirrors_of_tesis, require_tesis  # noqa: E402

CAP5_RE = re.compile(r"(?i)^\s*cap[ií]tulo\s*5\b")
CAP6_RE = re.compile(r"(?i)^\s*cap[ií]tulo\s*6\b")
END_RE = re.compile(r"(?i)^\s*(cap[ií]tulo\s*6\b|referencias\b|bibliograf)")
TOC_TAIL_RE = re.compile(r"\d+\s*$")


def is_toc_heading(text: str) -> bool:
    t = text.strip()
    if "\t" in t:
        return True
    if TOC_TAIL_RE.search(t) and not re.search(r"(?i)(2026|episod|kpi|tabla)", t):
        if re.search(r"(?i)cap[ií]tulo\s*\d+.*[a-záéíóúñ]{3,}\d+\s*$", t):
            return True
    return False


def iter_block_items(doc: Document):
    for child in list(doc.element.body.iterchildren()):
        tag = child.tag
        if tag == qn("w:p"):
            texts = [t.text or "" for t in child.iter(qn("w:t"))]
            yield child, "p", "".join(texts)
        elif tag == qn("w:tbl"):
            texts = [t.text or "" for t in child.iter(qn("w:t"))]
            yield child, "tbl", "".join(texts)
        else:
            yield child, "other", ""


def find_body_cap5_range(doc: Document, *, toc_safe: bool) -> tuple[int | None, int | None, list]:
    items = list(iter_block_items(doc))
    start = None
    end = None
    for i, (_el, _kind, text) in enumerate(items):
        t = text.strip()
        if start is None and CAP5_RE.search(t):
            if toc_safe and is_toc_heading(t):
                continue
            start = i
            continue
        if start is not None:
            if toc_safe:
                if CAP6_RE.search(t) and not is_toc_heading(t):
                    end = i
                    break
            elif END_RE.search(t) and not CAP5_RE.search(t):
                end = i
                break
    return start, end, items


def cap5_checks(text: str) -> dict[str, int]:
    return {
        "Shapiro": text.count("Shapiro"),
        "Kruskal": text.count("Kruskal"),
        "Wilcoxon": text.count("Wilcoxon"),
        "0,6667": text.count("0,6667") + text.count("0.6667"),
        "normalidad": text.lower().count("normalidad"),
        "Mann": len(re.findall(r"Mann.?Whitney", text, re.I)),
        "no_param": len(re.findall(r"no\s*param", text, re.I)),
    }


def extract_cap5(src_path: Path) -> tuple[list, dict]:
    doc = Document(str(src_path))
    # Tesis usually has no TOC page-number glue; still use toc_safe=False for simple range
    start, end, items = find_body_cap5_range(doc, toc_safe=False)
    if start is None:
        # fallback TOC-safe (if Tesis ever gets TOC)
        start, end, items = find_body_cap5_range(doc, toc_safe=True)
    if start is None:
        raise RuntimeError(f"Capítulo 5 no encontrado en {src_path.name}")
    if end is None:
        end = len(items)
        while end > start and items[end - 1][1] == "other":
            end -= 1
    elements = [deepcopy(items[i][0]) for i in range(start, end)]
    blob = "\n".join("".join(t.text or "" for t in el.iter(qn("w:t"))) for el in elements)
    meta = {
        "source": src_path.name,
        "range": [start, end],
        "blocks": len(elements),
        "checks": cap5_checks(blob),
    }
    return elements, meta


def replace_cap5(
    target_path: Path,
    cap5_elements: list,
    *,
    toc_safe: bool,
    dry_run: bool,
) -> dict:
    if not target_path.is_file():
        return {"file": target_path.name, "ok": False, "error": "missing"}

    doc = Document(str(target_path))
    start, end, items = find_body_cap5_range(doc, toc_safe=toc_safe)
    if start is None or end is None:
        return {
            "file": target_path.name,
            "ok": False,
            "error": f"Cap5 body not found (start={start}, end={end})",
            "toc_safe": toc_safe,
        }

    info: dict = {
        "file": target_path.name,
        "ok": True,
        "toc_safe": toc_safe,
        "old_range": [start, end],
        "dry_run": dry_run,
    }
    if dry_run:
        info["would_replace_blocks"] = end - start
        info["incoming_blocks"] = len(cap5_elements)
        return info

    to_remove = [items[i][0] for i in range(start, end)]
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    anchor = items[end][0]
    for el in cap5_elements:
        anchor.addprevious(el)

    # Política 2-Word: editar solo el canónico; no crear .bak/.docx nuevos en docs/.
    doc.save(str(target_path))
    info["backup"] = None

    doc2 = Document(str(target_path))
    s2, e2, items2 = find_body_cap5_range(doc2, toc_safe=toc_safe)
    texts = "\n".join(t for _, _, t in items2[s2:e2]) if s2 is not None and e2 is not None else ""
    info["new_range"] = [s2, e2]
    info["checks"] = cap5_checks(texts)
    return info


def sync(
    *,
    dry_run: bool = False,
    targets: list[Path] | None = None,
) -> dict:
    src = require_tesis()
    elements, src_meta = extract_cap5(src)
    dests = targets or list(mirrors_of_tesis())
    results = []
    for tgt in dests:
        # Informe has TOC fields; use toc_safe to avoid matching TOC lines.
        toc_safe = tgt.resolve() == INFORME.resolve()
        results.append(
            replace_cap5(
                tgt,
                [deepcopy(el) for el in elements],
                toc_safe=toc_safe,
                dry_run=dry_run,
            )
        )
    report = {"source": src_meta, "targets": results, "dry_run": dry_run}
    out = DOCS / "_cap5_sync_report.json"
    if not dry_run:
        out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        report["report_path"] = str(out)
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Sync Cap.5 from Tesis to Informe")
    parser.add_argument("--dry-run", action="store_true", help="Locate ranges; do not write")
    parser.add_argument(
        "--informe-only",
        action="store_true",
        help="Only sync Informe (default behaviour; kept for compatibility)",
    )
    args = parser.parse_args()
    targets: list[Path] | None = [INFORME] if args.informe_only else None
    report = sync(dry_run=args.dry_run, targets=targets)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    ok = all(t.get("ok") for t in report["targets"])
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
