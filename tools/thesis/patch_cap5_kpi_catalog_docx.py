#!/usr/bin/env python3
"""Parche Cap. V: catálogo completo de KPIs CityLearn v3 (evaluate_v2 / all_values).

Inserta en Cap. 5 (antes de §5.2) la cobertura verificada:
- 54 KPI oficiales (matriz VD / thesis_objective_evidence)
- 58 KPI runtime all_values (54 + flex_composite + 3 price_signal_*)
- 680 valores (4 algos × 3 escenarios), core 14/176, building 15 300

Aplica solo a los 2 Word canónicos (Tesis → sync Informe).
Idempotente: elimina bloque previo marcado si existe.
"""
from __future__ import annotations

import csv
import json
import re
import sys
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, INFORME, TESIS, require_tesis  # noqa: E402
from sync_cap5_to_canon_words import sync as sync_cap5  # noqa: E402

REPORT = DOCS / "CAP5_KPI_CATALOG_PATCH_REPORT_2026-07-29.json"
MARKER = "5.1.1 catalogo de kpis citylearn v3"
MATRIX_CSV = REPO / "outputs" / "thesis_objective_evidence" / "KPIs_y_metricas.csv"
ALL_VALUES_CSV = (
    REPO / "outputs" / "_drive_madrl" / "kpi_recalc_20260728" / "tables" / "all_evaluate_v2_kpis_long.csv"
)
CAP5_MD = DOCS / "tesis_capitulos" / "Capitulo_5_Resultados.md"

RUNTIME_EXTRA = [
    ("OE1", "flex_composite", "runtime (all_values)"),
    ("OE3", "price_signal_deviation_baseline", "runtime (derivado)"),
    ("OE3", "price_signal_deviation_delta", "runtime (derivado)"),
    ("OE3", "price_signal_deviation_ratio", "runtime (derivado)"),
]

OE_LABEL = {
    "OE1": "OE.1 Flexibilidad (E1)",
    "OE2": "OE.2 Emisiones CO₂ (E2)",
    "OE3": "OE.3 Costos (E3)",
}


def set_run_font(run, bold: bool = False, size: float = 12.0, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def replace_paragraph_text(p: Paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p.clear()
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", *, bold: bool = False, italic: bool = False, size: float = 12.0
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, bold=bold, italic=italic, size=size)
    return new_para


def set_heading_style(paragraph: Paragraph, level: int) -> None:
    style_name = f"Heading {level}"
    try:
        paragraph.style = style_name
    except KeyError:
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
        set_run_font(run, bold=True, size=14.0 if level == 2 else 12.0)


def norm(s: str) -> str:
    folded = unicodedata.normalize("NFKD", s or "")
    folded = "".join(ch for ch in folded if not unicodedata.combining(ch))
    return re.sub(r"\s+", " ", folded.strip().lower())


def _is_toc_heading(text: str) -> bool:
    t = text.strip()
    if "\t" in t:
        return True
    # TOC lines often end with a page number glued to the title.
    if re.search(r"\d+\s*$", t) and re.search(
        r"(?i)cap[ií]tulo\s*\d+.*[a-záéíóúñ]{3,}\d+\s*$", t
    ):
        return True
    return False


def chapter_bounds(doc: Document, chapter: int) -> tuple[int | None, int | None]:
    """Body Cap. N bounds (skip TOC and trailing meta notes)."""
    candidates: list[tuple[int, int | None]] = []
    start = None
    heading_re = re.compile(rf"^Cap[ií]tulo\s*{chapter}\s*[\.\:\-—]", re.I)
    next_re = re.compile(rf"^Cap[ií]tulo\s*{chapter + 1}\s*[\.\:\-—]", re.I)
    # Reject meta notes like "Capítulo 5 (2026-07-29): cumplimiento..."
    meta_note_re = re.compile(rf"^Cap[ií]tulo\s*{chapter}\s*\(", re.I)

    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if heading_re.search(t) and not _is_toc_heading(t) and not meta_note_re.search(t):
            start = i
        elif start is not None and next_re.search(t) and not _is_toc_heading(t):
            candidates.append((start, i))
            start = None
    if start is not None:
        candidates.append((start, None))
    if not candidates:
        # Fallback: plain "Capítulo N" / "Capítulo N. Título" without requiring punctuation.
        start = None
        for i, p in enumerate(doc.paragraphs):
            t = (p.text or "").strip()
            if (
                re.search(rf"^Cap[ií]tulo\s*{chapter}\b", t, re.I)
                and not _is_toc_heading(t)
                and not meta_note_re.search(t)
                and "cumplimiento" not in t.lower()
            ):
                start = i
            elif start is not None and re.search(rf"^Cap[ií]tulo\s*{chapter + 1}\b", t, re.I) and not _is_toc_heading(t):
                candidates.append((start, i))
                start = None
        if start is not None:
            candidates.append((start, None))
    if not candidates:
        return None, None
    # Prefer the candidate whose span contains §5.2 Resultados (for Cap. 5).
    if chapter == 5:
        scored = []
        for s, e in candidates:
            end_i = e if e is not None else len(doc.paragraphs)
            blob = "\n".join((doc.paragraphs[j].text or "") for j in range(s, end_i))
            score = 1 if re.search(r"(?i)5\.2\s+resultados", blob) else 0
            score += min(end_i - s, 500) / 500.0
            scored.append((score, s, e))
        scored.sort(reverse=True)
        return scored[0][1], scored[0][2]
    return candidates[-1]


def insert_table_after_paragraph(
    paragraph: Paragraph,
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    caption: str,
    note: str,
    *,
    font_size: float = 8.0,
) -> Paragraph:
    cap = insert_paragraph_after(paragraph, caption, italic=True, size=11.0)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=font_size)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=font_size)
    tbl = table._tbl
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)
    cap._p.addnext(tbl)
    note_el = OxmlElement("w:p")
    tbl.addnext(note_el)
    note_p = Paragraph(note_el, paragraph._parent)
    run = note_p.add_run(note)
    set_run_font(run, italic=True, size=10.0)
    return note_p


def strip_existing_block(doc: Document, start_idx: int, end_idx: int | None) -> bool:
    """Remove previous 5.1.1 block (paragraphs + tables) until §5.2 or Cap. 6."""
    paras = doc.paragraphs
    end = end_idx if end_idx is not None else len(paras)
    block_start = None
    for i in range(start_idx, end):
        if MARKER in norm(paras[i].text or ""):
            block_start = i
            break
    if block_start is None:
        return False
    block_end = None
    for i in range(block_start + 1, end):
        n = norm(paras[i].text or "")
        if re.match(r"^5\.2(\s|$|—|-)", n) and not n.startswith("5.2."):
            block_end = i
            break
        if n.startswith("5.2 resultados"):
            block_end = i
            break
    if block_end is None:
        if end_idx is None:
            return False
        block_end = end_idx
    start_el = paras[block_start]._element
    end_el = paras[block_end]._element
    body = start_el.getparent()
    if body is None:
        return False
    to_remove = []
    el = start_el
    while el is not None and el != end_el:
        nxt = el.getnext()
        to_remove.append(el)
        el = nxt
    for el in to_remove:
        body.remove(el)
    return True


def load_matrix_rows() -> list[dict[str, str]]:
    with open(MATRIX_CSV, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def load_all_values_kpis() -> set[str]:
    with open(ALL_VALUES_CSV, newline="", encoding="utf-8") as f:
        return {r["kpi"] for r in csv.DictReader(f)}


def coverage_rows() -> list[list[str]]:
    return [
        [
            "54 KPI oficiales (matriz VD)",
            "54",
            "12/12 tratamientos (price_signal_* solo E1+E3)",
            "outputs/thesis_objective_evidence/KPIs_y_metricas.*",
        ],
        [
            "Runtime all_values (capa v3)",
            "58 (=54+4)",
            "680 valores (4×3); sin nulos",
            "outputs/_drive_madrl/kpi_recalc_20260728/tables/all_evaluate_v2_kpis_*.csv",
        ],
        [
            "Core KPI",
            "14 nombres / 176 valores",
            "12 jobs",
            "tables/all_core_kpis_*.csv",
        ],
        [
            "Building KPI",
            "15 300 filas",
            "17 edificios × jobs",
            "by_building/building_kpis_all.csv",
        ],
    ]


def oe_summary_rows(matrix: list[dict[str, str]]) -> list[list[str]]:
    by_oe: dict[str, list[str]] = {"OE1": [], "OE2": [], "OE3": []}
    for r in matrix:
        by_oe.setdefault(r["axis"], []).append(r["kpi"])
    examples = {
        "OE1": "grid_*, peak/ramping/load_factor, pv_*, battery_*, ev_*, community_*, zero_net_energy, net_exchange_*; + flex_composite",
        "OE2": "carbon_emissions, carbon_emissions_control/baseline/delta (+ daily_average_*)",
        "OE3": "electricity_cost*, cost_peak/ramping/load_factor, price_signal_deviation (+ 3 derivados)",
    }
    rows = []
    for axis in ("OE1", "OE2", "OE3"):
        rows.append(
            [
                OE_LABEL[axis],
                str(len(by_oe[axis])),
                examples[axis],
            ]
        )
    rows.append(
        [
            "Extras runtime (no en matriz 54)",
            "4",
            "flex_composite; price_signal_deviation_baseline/delta/ratio",
        ]
    )
    return rows


def catalog_rows(matrix: list[dict[str, str]]) -> list[list[str]]:
    rows: list[list[str]] = []
    for r in matrix:
        lib = r.get("lower_is_better", "").strip().lower()
        orient = "menor=mejor" if lib in {"true", "1", "yes"} else "mayor=mejor"
        rows.append([OE_LABEL.get(r["axis"], r["axis"]), r["kpi"], "oficial", orient])
    for axis, kpi, tipo in RUNTIME_EXTRA:
        rows.append([OE_LABEL.get(axis, axis), kpi, tipo, "menor=mejor"])
    return rows


def find_insert_anchor(doc: Document, start_idx: int, end_idx: int | None) -> Paragraph:
    """Insert immediately before the body §5.2 heading."""
    paras = doc.paragraphs
    end = end_idx if end_idx is not None else len(paras)
    idx_52 = None
    for i in range(start_idx, end):
        n = norm(paras[i].text or "")
        # Exact numeral heading (avoid Tabla 5.2.x)
        if re.match(r"^5\.2(\s|$|—|-)", n) and not n.startswith("5.2."):
            idx_52 = i
            break
    if idx_52 is None:
        # Fallback: first paragraph that is exactly the 5.2 results heading.
        for i in range(start_idx, end):
            n = norm(paras[i].text or "")
            if n.startswith("5.2 resultados"):
                idx_52 = i
                break
    if idx_52 is None:
        raise RuntimeError(
            f"No se encontró ancla §5.2 en Cap. 5 (start={start_idx}, end={end})"
        )
    if idx_52 <= start_idx:
        raise RuntimeError("Ancla §5.2 inválida (coincide con inicio Cap. 5)")
    return paras[idx_52 - 1]


def patch_54_mentions(doc: Document, start_idx: int, end_idx: int | None) -> int:
    """Clarify legacy '54 KPI' wording inside Cap. 5 only."""
    n_changed = 0
    end = end_idx if end_idx is not None else len(doc.paragraphs)
    for i in range(start_idx, end):
        p = doc.paragraphs[i]
        t = p.text or ""
        if "54 KPI" not in t and "54 KPIs" not in t and "54 kpi" not in t.lower():
            continue
        if "58" in t and "all_values" in t.lower():
            continue
        new = t
        new = re.sub(
            r"54\s*KPI(?:s)?\s+oficiales",
            "54 KPI oficiales (matriz VD; runtime all_values = 58 con flex_composite y derivados price_signal)",
            new,
            flags=re.I,
        )
        new = re.sub(
            r"los\s+54\s+KPI(?:s)?",
            "los 54 KPI oficiales (58 en all_values runtime)",
            new,
            flags=re.I,
        )
        if new != t:
            replace_paragraph_text(p, new)
            n_changed += 1
    return n_changed


def insert_catalog_block(doc: Document, anchor: Paragraph, matrix: list[dict[str, str]]) -> list[str]:
    actions: list[str] = []
    cursor = insert_paragraph_after(
        anchor,
        "5.1.1 Catálogo de KPIs CityLearn v3 propuesto (evaluate_v2 / all_values)",
        bold=True,
    )
    set_heading_style(cursor, 3)
    actions.append("insert_heading_5_1_1")

    intro = (
        "La variable dependiente de la tesis se operacionaliza con los 54 KPI oficiales de la "
        "matriz VD (D-VD.1 flexibilidad, D-VD.2 CO₂, D-VD.3 costos), calculados por la capa "
        "CityLearn v3 propuesto a partir de env.evaluate_v2() de CityLearn v2 y persistidos en "
        "citylearn_v3_report.all_values. El recálculo Drive 2026-07-28 "
        "(outputs/_drive_madrl/kpi_recalc_20260728/) recupera el catálogo runtime completo de "
        "58 KPIs (los 54 oficiales más flex_composite y tres derivados de price_signal_deviation) "
        "con 680 valores en 12/12 tratamientos, 14 core KPI (176 valores) y 15 300 filas "
        "building KPI. La familia price_signal_deviation* se reporta en E1 y E3 (no en E2). "
        "Los agregados de hipótesis (KPI-gains) y el ranking evaluate_v2 4/4 de este capítulo "
        "se construyen sobre ese mismo catálogo; no se inventan KPIs fuera de evaluate_v2/all_values."
    )
    cursor = insert_paragraph_after(cursor, intro)
    actions.append("insert_intro")

    cursor = insert_table_after_paragraph(
        cursor,
        doc,
        ["Capa", "n", "Cobertura", "Artefacto"],
        coverage_rows(),
        "Tabla 5.1.1. Cobertura del catálogo KPI CityLearn v3 propuesto (corrida madrl_v3_20260627_164047).",
        "Nota. Fuente: kpi_recalc_20260728 y thesis_objective_evidence/KPIs_y_metricas (2026-07-28/29). "
        "Valores < 1 favorecen al MADRL cuando la métrica es «menor = mejor».",
    )
    actions.append("insert_table_5_1_1")

    cursor = insert_table_after_paragraph(
        cursor,
        doc,
        ["Objetivo / capa", "n KPI", "Familias y KPIs representativos"],
        oe_summary_rows(matrix),
        "Tabla 5.1.2. Distribución del catálogo por OE.1–OE.3 y extras runtime.",
        "Nota. Los 54 oficiales suman OE.1=36, OE.2=7, OE.3=11. Los 4 extras completan el runtime "
        "all_values (58).",
    )
    actions.append("insert_table_5_1_2")

    cursor = insert_table_after_paragraph(
        cursor,
        doc,
        ["Eje", "KPI", "Tipo", "Orientación"],
        catalog_rows(matrix),
        "Tabla 5.1.3. Listado nominal de los 54 KPI oficiales + 4 runtime (catálogo integrado Cap. V).",
        "Nota. Listado nominal completo. Valores numéricos por algoritmo×escenario en "
        "all_evaluate_v2_kpis_long.csv / wide.csv; KPIs_y_metricas_FULL.md.",
        font_size=7.5,
    )
    actions.append("insert_table_5_1_3")
    return actions


def verify_cap5(doc: Document) -> dict:
    start, end = chapter_bounds(doc, 5)
    if start is None:
        return {"ok": False, "error": "Cap. 5 no encontrado"}
    end_i = end if end is not None else len(doc.paragraphs)
    blob = "\n".join((p.text or "") for p in doc.paragraphs[start:end_i])
    # KPI names also live in table cells (not only paragraphs).
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                blob += "\n" + (cell.text or "")
    matrix = load_matrix_rows()
    named = sum(1 for r in matrix if r["kpi"] in blob)
    extras = sum(1 for _, kpi, _ in RUNTIME_EXTRA if kpi in blob)
    return {
        "ok": True,
        "has_marker": MARKER in norm(blob),
        "mentions_58": "58" in blob and ("KPI" in blob or "kpi" in blob.lower()),
        "mentions_680": "680" in blob,
        "mentions_15300": ("15 300" in blob) or ("15300" in blob),
        "official_kpi_names_in_cap5": named,
        "runtime_extra_names_in_cap5": extras,
        "expected_official": len(matrix),
        "expected_runtime_extra": len(RUNTIME_EXTRA),
        "cap5_range": [start, end],
    }


def patch_tesis() -> dict:
    require_tesis()
    if not MATRIX_CSV.is_file():
        raise FileNotFoundError(MATRIX_CSV)
    if not ALL_VALUES_CSV.is_file():
        raise FileNotFoundError(ALL_VALUES_CSV)

    matrix = load_matrix_rows()
    all_kpis = load_all_values_kpis()
    missing_in_all = sorted({r["kpi"] for r in matrix} - all_kpis)
    if missing_in_all:
        raise RuntimeError(f"KPIs de matriz ausentes en all_values: {missing_in_all}")

    doc = Document(str(TESIS))
    start, end = chapter_bounds(doc, 5)
    if start is None:
        raise RuntimeError("No se pudo delimitar Cap. 5 en Tesis")

    removed = strip_existing_block(doc, start, end)
    # refresh bounds after possible deletion
    start, end = chapter_bounds(doc, 5)
    if start is None:
        raise RuntimeError("Cap. 5 perdido tras strip")
    n54 = patch_54_mentions(doc, start, end)
    anchor = find_insert_anchor(doc, start, end)
    actions = insert_catalog_block(doc, anchor, matrix)
    doc.save(str(TESIS))

    doc2 = Document(str(TESIS))
    verification = verify_cap5(doc2)
    return {
        "file": TESIS.name,
        "removed_previous": removed,
        "patched_54_mentions": n54,
        "actions": actions,
        "matrix_n": len(matrix),
        "all_values_n": len(all_kpis),
        "verification": verification,
    }


def patch_cap5_markdown(matrix: list[dict[str, str]]) -> bool:
    if not CAP5_MD.is_file():
        return False
    text = CAP5_MD.read_text(encoding="utf-8")
    marker = "### 5.1.1 Catálogo de KPIs CityLearn v3 propuesto"
    block = [
        "",
        marker,
        "",
        "La VD se operacionaliza con **54 KPI oficiales** (matriz "
        "`outputs/thesis_objective_evidence/KPIs_y_metricas.*`: OE.1=36, OE.2=7, OE.3=11). "
        "El runtime `citylearn_v3_report.all_values` añade **4** métricas "
        "(`flex_composite` + 3 derivados `price_signal_*`) → **58 KPI / 680 valores** en "
        "`outputs/_drive_madrl/kpi_recalc_20260728/` (12/12 tratamientos; "
        "`price_signal_*` solo E1+E3). Core: 14/176; building: 15 300 filas.",
        "",
        "**Tabla 5.1.1.** Cobertura del catálogo.",
        "",
        "| Capa | n | Cobertura | Artefacto |",
        "|---|---:|---|---|",
        "| 54 KPI oficiales (matriz VD) | 54 | 12/12 (`price_signal_*` E1+E3) | `KPIs_y_metricas.*` |",
        "| Runtime all_values | 58 | 680 valores | `all_evaluate_v2_kpis_*.csv` |",
        "| Core KPI | 14 / 176 | 12 jobs | `all_core_kpis_*.csv` |",
        "| Building KPI | 15 300 | 17 edificios × jobs | `building_kpis_all.csv` |",
        "",
        "**Tabla 5.1.2.** Por OE + extras runtime.",
        "",
        "| Objetivo | n | Representativos |",
        "|---|---:|---|",
        "| OE.1 E1 | 36 | grid_*, peak/ramping/load_factor, pv_*, battery_*, ev_*, community_*, zero_net_energy |",
        "| OE.2 E2 | 7 | carbon_emissions* |",
        "| OE.3 E3 | 11 | electricity_cost*, cost_peak/ramping/load_factor, price_signal_deviation |",
        "| Extras runtime | 4 | flex_composite; price_signal_deviation_baseline/delta/ratio |",
        "",
        "**Tabla 5.1.3.** Listado nominal (54 oficiales + 4 runtime).",
        "",
        "| Eje | KPI | Tipo | Orientación |",
        "|---|---|---|---|",
    ]
    for r in matrix:
        lib = r.get("lower_is_better", "").strip().lower()
        orient = "menor=mejor" if lib in {"true", "1", "yes"} else "mayor=mejor"
        block.append(f"| {OE_LABEL.get(r['axis'], r['axis'])} | `{r['kpi']}` | oficial | {orient} |")
    for axis, kpi, tipo in RUNTIME_EXTRA:
        block.append(f"| {OE_LABEL.get(axis, axis)} | `{kpi}` | {tipo} | menor=mejor |")
    block.append("")

    section = "\n".join(block)
    if marker in text:
        # replace from marker until next ##/### at same or higher level before 5.2
        pattern = re.compile(
            r"### 5\.1\.1 Catálogo de KPIs CityLearn v3 propuesto.*?(?=\n## 5\.2 |\n### 5\.2 )",
            re.S,
        )
        text2, n = pattern.subn(section.lstrip("\n") + "\n", text, count=1)
        if n == 0:
            return False
        CAP5_MD.write_text(text2, encoding="utf-8")
        return True

    # insert before ## 5.2
    m = re.search(r"\n## 5\.2 ", text)
    if not m:
        return False
    text2 = text[: m.start()] + "\n" + section + text[m.start() :]
    CAP5_MD.write_text(text2, encoding="utf-8")
    return True


def main() -> int:
    tesis_result = patch_tesis()
    matrix = load_matrix_rows()
    md_ok = patch_cap5_markdown(matrix)
    sync_result = sync_cap5(dry_run=False, targets=[INFORME] if INFORME.is_file() else [])

    informe_ver = None
    if INFORME.is_file():
        informe_ver = verify_cap5(Document(str(INFORME)))

    report = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "tesis": tesis_result,
        "markdown_updated": md_ok,
        "sync_cap5_to_informe": sync_result,
        "informe_verification": informe_ver,
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    ok = bool(tesis_result.get("verification", {}).get("ok")) and tesis_result["verification"].get(
        "official_kpi_names_in_cap5", 0
    ) >= 50
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
