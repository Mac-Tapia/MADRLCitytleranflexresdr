#!/usr/bin/env python3
"""Parche quirúrgico Cap. 6 (Informe) + APA de alto impacto en canons Word.

- Corrige Tabla 6.1 del Informe (veredicto KPI-gains Cap. 5).
- Alinea hallazgos/limitaciones clave Informe ↔ Tesis.
- Añade referencias faltantes (Shadish, Creswell, Nweye 2023a/b).
- Normaliza Vazquez-Canteli → Vázquez-Canteli en Tesis.
- Actualiza entradas MINAM/OSINERGMIN con acrónimo APA [MINAM]/[OSINERGMIN].

Backups: outputs/_word_backups/ (nunca docs/).
"""
from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import INFORME, TESIS  # noqa: E402

BACKUP_DIR = REPO / "outputs" / "_word_backups"
STAMP = datetime.now().strftime("%Y%m%d_%H%M%S")

CAP6_TABLE_ROWS = [
    ["Hipótesis", "Decisión", "Fundamento cuantitativo (KPI-gains Cap. 5)"],
    [
        "H0G / H1G",
        "Rechazar H0G de forma exploratoria; respaldar H1G exploratoria (sin ganador único)",
        "Friedman integración p = 0,0096; KW ALL p = 0,1554 n.s.; no implica HE11∧HE21∧HE31",
    ],
    [
        "HE10 / HE11",
        "No rechazar HE10; no respaldar HE11",
        "Kruskal–Wallis E1 p = 0,4685 (KPI-gains)",
    ],
    [
        "HE20 / HE21",
        "No rechazar HE20; no respaldar HE21",
        "Kruskal–Wallis E2 p = 0,7648 (KPI-gains)",
    ],
    [
        "HE30 / HE31",
        "No rechazar HE30; no respaldar HE31",
        "Kruskal–Wallis E3 p = 0,7357 (KPI-gains); TOPSIS/4/4 no sustituyen HE31",
    ],
]

REFS_TO_ENSURE = [
    "Creswell, J. W., & Creswell, J. D. (2023). Research design: Qualitative, quantitative, and mixed methods approaches (6th ed.). SAGE Publications.",
    "Nweye, K., Sankaranarayanan, S., & Nagy, Z. (2023a). MERLIN: Multi-agent offline and transfer learning for occupant-centric operation of grid-interactive communities. Applied Energy. https://arxiv.org/abs/2301.01148",
    "Nweye, K., et al. (2023b). Heterogeneous multi-agent reinforcement learning for grid-interactive communities. En Proceedings of the 10th ACM International Conference on Systems for Energy-Efficient Buildings, Cities, and Transportation. ACM. https://doi.org/10.1145/3600100.3626276",
    "Shadish, W. R., Cook, T. D., & Campbell, D. T. (2002). Experimental and quasi-experimental designs for generalized causal inference. Houghton Mifflin.",
]

MINAM_OLD = "Ministerio del Ambiente del Perú. (2019). INFOCARBONO — RAGEI 2019 Energía. MINAM. https://infocarbono.minam.gob.pe/"
MINAM_NEW = "Ministerio del Ambiente del Perú [MINAM]. (2019). INFOCARBONO — RAGEI 2019 Energía. https://infocarbono.minam.gob.pe/"
OSIN_OLD = "Organismo Supervisor de la Inversión en Energía y Minería. (2024). Resolución de Consejo Directivo N.° 0024-2024-OS/CD — Tarifas de Distribución Eléctrica MT-3/MT-4, Electro Oriente S.A. OSINERGMIN."
OSIN_NEW = "Organismo Supervisor de la Inversión en Energía y Minería [OSINERGMIN]. (2024). Resolución de Consejo Directivo N.° 0024-2024-OS/CD — Tarifas de Distribución Eléctrica MT-3/MT-4, Electro Oriente S.A."

HALLAZGOS_HE = (
    "Respecto de las hipótesis (Cap. 5 §5.5, KPI-gains): H0G se rechaza de forma exploratoria "
    "(Friedman p = 0,0096) mientras KW ALL no es significativo (p = 0,1554). "
    "HE10 (KW p = 0,4685), HE20 (KW p = 0,7648) y HE30 (KW p = 0,7357) no se rechazan; "
    "HE11/HE21/HE31 no se respaldan. Descriptivamente MATD3 lidera OE.1/OE.2 y MAAC lidera OE.3. "
    "Wilcoxon/Mann–Whitney exploratorios no sustituyen el contraste omnibus ni la replicación multi-semilla."
)

LIMITACIONES_ALIGN = (
    "Las limitaciones principales son: (i) semilla única (seed 0) en la campaña entrenada —el protocolo "
    "n_seeds=12 y el runner están implementados y validados con smoke ilustrativo, pero no se entrenaron "
    "12 semillas reales—; (ii) HAPPO con 49/50 episodios y KPI-gains evaluate_v2 peores que el trío "
    "(ranking 4/4 score 0); las HE sobre KPI-gains de entrenamiento usan MASAC/MATD3/MAAC; "
    "(iii) simulación sin validación en red física; (iv) trade-off multiobjetivo sin ganador Pareto "
    "universal; (v) MADRL por debajo del baseline RBC en score global HPHI. Las decisiones de HE "
    "usan la capa KPI-gains (Cap. 5); TOPSIS/4/4 y medias episódicas son descriptivas."
)


def set_run_font(run, bold: bool = False, size: float = 12.0) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def replace_paragraph_text(p: Paragraph, text: str, *, bold: bool = False) -> None:
    p.clear()
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def backup(path: Path) -> Path:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    dest = BACKUP_DIR / f"{path.stem}_{STAMP}{path.suffix}"
    shutil.copy2(path, dest)
    return dest


def fill_table(table, rows: list[list[str]]) -> None:
    # Resize rows if needed
    while len(table.rows) < len(rows):
        table.add_row()
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        while len(row.cells) < len(row_data):
            # cannot easily add cells; write into existing
            break
        for j, val in enumerate(row_data):
            if j < len(row.cells):
                # clear cell paragraphs
                cell = row.cells[j]
                for pi, para in enumerate(cell.paragraphs):
                    if pi == 0:
                        replace_paragraph_text(para, val, bold=(i == 0))
                    else:
                        para.clear()


def find_cap6_hypothesis_table(doc: Document):
    body = doc.element.body
    capture = False
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            texts = [n.text or "" for n in child.iter(qn("w:t"))]
            t = "".join(texts).strip()
            if "Capítulo 6" in t or "Capitulo 6" in t:
                capture = True
            elif capture and ("Referencias" in t or t.startswith("Anexo")):
                break
        elif capture and child.tag == qn("w:tbl"):
            # map to python-docx table
            for table in doc.tables:
                if table._tbl is child:
                    flat = " | ".join(
                        c.text.strip() for r in table.rows for c in r.cells
                    )
                    if "H0G" in flat and "HE10" in flat:
                        return table
    return None


def ensure_refs(doc: Document) -> list[str]:
    added: list[str] = []
    paras = list(doc.paragraphs)
    ref_idx = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip().lower()
        # "bibliogr" evita fallo por á/í en «bibliográficas»
        if "referencias" in t and "bibliogr" in t:
            ref_idx = i
            break
    if ref_idx is None:
        return ["ERROR: no se encontró sección Referencias"]

    existing = "\n".join((p.text or "") for p in paras[ref_idx:])
    # Update MINAM/OSINERGMIN entries in place
    for p in paras[ref_idx:]:
        t = (p.text or "").strip()
        if t.startswith("Ministerio del Ambiente del Perú.") and "INFOCARBONO" in t:
            replace_paragraph_text(p, MINAM_NEW)
            added.append("updated MINAM entry")
        elif t.startswith("Organismo Supervisor") and "0024-2024" in t:
            replace_paragraph_text(p, OSIN_NEW)
            added.append("updated OSINERGMIN entry")

    # Refresh existing after updates
    existing = "\n".join((p.text or "") for p in doc.paragraphs[ref_idx:])

    # Find insertion anchor: alphabetically near neighbors
    insert_after_map = {
        "Creswell": ("Campbell", "Demsar", "Colas", "Chen"),
        "Nweye, K., Sankaranarayanan": ("Nweye, K., Liu", "Nweye, K., Kaspar, K., Buscemi, G., Pinto"),
        "Nweye, K., et al. (2023b)": ("Nweye, K., Sankaranarayanan", "Nweye, K., Liu", "Nweye, K., Kaspar"),
        "Shadish": ("Sarkar", "Shojaeighadikolaei", "Sutton", "Rajagopalan"),
    }

    for ref in REFS_TO_ENSURE:
        key = ref.split("(")[0].strip()
        # presence checks
        if "Creswell, J. W" in existing and ref.startswith("Creswell"):
            continue
        if "2023a" in ref and "Sankaranarayanan" in existing and "(2023a)" in existing:
            continue
        if "2023b" in ref and "3626276" in existing:
            continue
        if ref.startswith("Shadish") and "Shadish, W" in existing:
            continue

        # find anchor paragraph
        anchors = []
        for prefix_key, candidates in insert_after_map.items():
            if ref.startswith(prefix_key.split(",")[0]) or prefix_key in ref[:40]:
                anchors = list(candidates)
                if ref.startswith("Creswell"):
                    anchors = insert_after_map["Creswell"]
                elif "2023a" in ref:
                    anchors = insert_after_map["Nweye, K., Sankaranarayanan"]
                elif "2023b" in ref:
                    anchors = insert_after_map["Nweye, K., et al. (2023b)"]
                elif ref.startswith("Shadish"):
                    anchors = insert_after_map["Shadish"]
                break

        anchor_para = None
        for p in doc.paragraphs[ref_idx:]:
            t = (p.text or "").strip()
            for a in anchors:
                if t.startswith(a):
                    anchor_para = p
                    break
            if anchor_para:
                break
        if anchor_para is None:
            # append before first Vázquez or at end of refs block
            for p in doc.paragraphs[ref_idx:]:
                t = (p.text or "").strip()
                if t.startswith("Vázquez-Canteli") or t.startswith("Vazquez-Canteli") or t.startswith("Zhong"):
                    anchor_para = p
                    # insert before: use previous
                    # we'll insert after previous sibling by inserting before via previous para
                    break
            if anchor_para is not None:
                # insert before anchor: get previous paragraph
                prev = anchor_para._p.getprevious()
                from docx.text.paragraph import Paragraph as P

                # create after previous element
                if prev is not None and prev.tag == qn("w:p"):
                    prev_para = P(prev, anchor_para._parent)
                    insert_paragraph_after(prev_para, ref)
                else:
                    insert_paragraph_after(anchor_para, ref)
                added.append(f"added {ref[:40]}...")
                existing += "\n" + ref
                continue
            # fallback: after ref heading
            insert_paragraph_after(doc.paragraphs[ref_idx], ref)
            added.append(f"added-at-head {ref[:40]}...")
            existing += "\n" + ref
            continue

        insert_paragraph_after(anchor_para, ref)
        added.append(f"added {ref[:40]}...")
        existing += "\n" + ref

    return added


def patch_informe(doc: Document) -> list[str]:
    actions: list[str] = []
    table = find_cap6_hypothesis_table(doc)
    if table is None:
        actions.append("WARN: no Cap.6 H0G/HE10 table found")
    else:
        fill_table(table, CAP6_TABLE_ROWS)
        actions.append("fixed Cap.6 hypothesis table to KPI-gains verdict")

    # Insert hallazgos HE paragraph after OG paragraph if missing
    body_text = "\n".join(p.text for p in doc.paragraphs if p.text)
    if "HE11/HE21/HE31 no se respaldan" not in body_text and "HE11/HE21/HE31 no se respaldan" not in body_text.replace(" ", ""):
        pass
    if "HE11/HE21/HE31 no se respaldan" not in body_text:
        for i, p in enumerate(doc.paragraphs):
            t = (p.text or "").strip()
            style = p.style.name if p.style else ""
            if style.startswith("Heading") and ("6.2 Veredicto" in t or "6.2 " in t and "hipótesis" in t.lower()):
                # insert before this heading: find previous paragraph
                if i > 0:
                    insert_paragraph_after(doc.paragraphs[i - 1], HALLAZGOS_HE)
                    actions.append("inserted Cap.6 hallazgos HE paragraph")
                break

    # Align limitaciones section first substantive paragraph
    in_lim = False
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        style = p.style.name if p.style else ""
        if style.startswith("Heading") and "6.3" in t and ("Limitacion" in t or "Limitación" in t or "alcance" in t.lower()):
            in_lim = True
            continue
        if in_lim and style.startswith("Heading"):
            break
        if in_lim and t and len(t) > 80:
            if "semilla única (seed 0)" not in t and "n_seeds=12" not in t:
                replace_paragraph_text(p, LIMITACIONES_ALIGN)
                actions.append("aligned Cap.6 limitaciones with Tesis KPI-gains framing")
            break

    actions.extend(f"refs:{a}" for a in ensure_refs(doc))
    return actions


def patch_tesis(doc: Document) -> list[str]:
    actions: list[str] = []
    n_fix = 0
    for p in doc.paragraphs:
        t = p.text or ""
        if "Vazquez-Canteli" in t:
            new = t.replace("Vazquez-Canteli", "Vázquez-Canteli")
            if new != t:
                # preserve bold of first run roughly
                bold = bool(p.runs and p.runs[0].bold)
                replace_paragraph_text(p, new, bold=bool(bold))
                n_fix += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text or ""
                    if "Vazquez-Canteli" in t:
                        replace_paragraph_text(p, t.replace("Vazquez-Canteli", "Vázquez-Canteli"))
                        n_fix += 1
    if n_fix:
        actions.append(f"normalized Vazquez→Vázquez in {n_fix} paragraphs/cells")

    actions.extend(f"refs:{a}" for a in ensure_refs(doc))
    return actions


def main() -> int:
    report: dict = {"stamp": STAMP, "files": {}}
    for path, patcher in ((INFORME, patch_informe), (TESIS, patch_tesis)):
        if not path.is_file():
            report["files"][path.name] = {"error": "missing"}
            continue
        b = backup(path)
        doc = Document(str(path))
        actions = patcher(doc)
        doc.save(str(path))
        report["files"][path.name] = {"backup": str(b), "actions": actions}
        print(path.name, "→", actions)

    out = REPO / "outputs" / f"cap6_apa_patch_report_{STAMP}.json"
    import json

    out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print("report", out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
