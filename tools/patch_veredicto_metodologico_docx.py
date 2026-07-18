#!/usr/bin/env python3
"""Portar veredicto metodologico (Caps. 1, 3, 6) a los Word PATCHED/SYNCED."""
from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
DOCS = REPO / "docs"

TARGETS = [
    DOCS / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS_PATCHED.docx",
    DOCS / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA_PATCHED.docx",
    DOCS / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA_SYNCED.docx",
]

FALLBACK_FROM = {
    DOCS / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS_PATCHED.docx": DOCS
    / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx",
    DOCS / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA_PATCHED.docx": DOCS
    / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx",
    DOCS / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA_SYNCED.docx": DOCS
    / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA_PATCHED.docx",
}


def set_run_font(run, bold: bool = False, size: float = 12.0) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def clear_paragraph(p: Paragraph) -> None:
    p.clear()


def replace_paragraph_text(p: Paragraph, text: str, *, bold: bool = False) -> None:
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, bold=bold)
    return new_para


def insert_table_after(paragraph: Paragraph, headers: list[str], rows: list[list[str]]) -> Paragraph:
    """Insert a simple table after paragraph; return last paragraph after table (marker)."""
    # Create table via a temporary document body append then move XML — simpler: use doc.add_table
    # and relocate. We use parent document.
    doc = paragraph.part.document
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=9)
    tbl = table._tbl
    # Move table XML to after paragraph
    paragraph._p.addnext(tbl)
    # Remove orphan from end of body (python-docx appends at end)
    body = paragraph.part.element.body
    # table already moved via addnext which also leaves a copy? addnext moves the element.
    # Actually add_table appends to body; addnext relocates the same element. Good.
    return paragraph


def find_para_idx(doc: Document, pattern: str, start: int = 0, end: int | None = None) -> int | None:
    rx = re.compile(pattern, re.IGNORECASE)
    last = end if end is not None else len(doc.paragraphs)
    for i in range(start, last):
        if rx.search(doc.paragraphs[i].text or ""):
            return i
    return None


def find_heading_range(doc: Document, start_pat: str, end_pat: str) -> tuple[int, int] | None:
    s = find_para_idx(doc, start_pat)
    if s is None:
        return None
    e = find_para_idx(doc, end_pat, start=s + 1)
    if e is None:
        e = min(s + 80, len(doc.paragraphs))
    return s, e


def delete_paragraphs_between(doc: Document, start: int, end: int, keep_start: bool = True) -> None:
    """Delete paragraphs in (start, end) exclusive of end; optionally keep start."""
    # Delete from end-1 down to start+1 (or start)
    first = start if not keep_start else start + 1
    for i in range(end - 1, first - 1, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)


def patch_cap1(doc: Document) -> dict:
    info: dict = {"ok": False, "actions": []}

    # Prefer replacing the formulation / objectives / hypotheses blocks by markers.
    # Strategy: find "Formulación del problema" or "Problema general" and rewrite contiguous block until "1.4" or "Justificación".

    start = find_para_idx(doc, r"Formulaci[oó]n del problema|Problema general\s*\(PG\)")
    end = find_para_idx(doc, r"^1\.4\b|Justificaci[oó]n", start=start or 0)
    if start is None:
        info["actions"].append("cap1_block_not_found")
        return info

    if end is None:
        end = min(start + 60, len(doc.paragraphs))

    # Keep the heading at start if it is a heading; replace content after
    anchor = doc.paragraphs[start]
    # Clear from start to end-1 and rewrite
    delete_paragraphs_between(doc, start, end, keep_start=False)

    # Re-find anchor (indices shifted) — we deleted including start, so use previous
    # After deleting start..end-1, the old end becomes the next paragraph after removal.
    # Safer: insert before the paragraph that was `end` by finding Justificación again.
    just = find_para_idx(doc, r"^1\.4\b|Justificaci[oó]n")
    if just is None:
        info["actions"].append("cap1_justificacion_lost")
        return info
    # Insert before justificación by using paragraph just-1 as anchor for addnext chain... 
    # We need insert BEFORE just. Get paragraph before just.
    before = doc.paragraphs[just - 1] if just > 0 else doc.paragraphs[0]

    blocks = [
        ("### 1.1.3 Formulación del problema", True),
        ("Problema general (PG):", True),
        (
            "¿Qué algoritmo MADRL ofrece el mejor compromiso (ranking / frontera de Pareto) "
            "de gestión coordinada de la flexibilidad energética, las emisiones de CO₂ y los "
            "costos energéticos en comunidades inteligentes simuladas bajo CityLearn v3 en el SEAI Iquitos?",
            False,
        ),
        ("Problemas específicos:", True),
        (
            "PE.1: ¿Qué MADRL lidera la flexibilidad energética (escenario E1) y con qué evidencia descriptiva e inferencial?",
            False,
        ),
        (
            "PE.2: ¿Qué MADRL lidera la reducción de emisiones de CO₂ (escenario E2) y con qué evidencia descriptiva e inferencial?",
            False,
        ),
        (
            "PE.3: ¿Qué MADRL lidera la optimización de costos energéticos (escenario E3) y con qué evidencia descriptiva e inferencial?",
            False,
        ),
        ("1.2 Objetivos", True),
        ("1.2.1 Objetivo general", True),
        (
            "OG. Identificar el(los) MADRL recomendable(s) por eje y el ranking integrado de gestión "
            "coordinada de flexibilidad, CO₂ y costos en el SEAI Iquitos, sin asumir dominancia Pareto universal.",
            False,
        ),
        ("1.2.2 Objetivos específicos", True),
        (
            "OE.1. Identificar el MADRL líder en flexibilidad energética (E1) y contrastar si la diferencia entre algoritmos es estadísticamente sustentable.",
            False,
        ),
        (
            "OE.2. Identificar el MADRL líder en reducción de emisiones de CO₂ (E2) y contrastar si la diferencia entre algoritmos es estadísticamente sustentable.",
            False,
        ),
        (
            "OE.3. Identificar el MADRL líder en costos energéticos (E3) y contrastar si la diferencia entre algoritmos es estadísticamente sustentable.",
            False,
        ),
        (
            "Coherencia vertical: cada OE responde a su PE, se operacionaliza con E1/E2/E3 y se evalúa con KPIs CityLearn v2 más recompensa episódica. "
            "Las métricas primarias son KPIs energéticos y recompensa MADRL; accuracy/precision/recall/F1 no se usan como métricas centrales.",
            False,
        ),
        ("1.3 Hipótesis", True),
        (
            "Nota metodológica: el estudio es cuantitativo, aplicado y cuasiexperimental factorial 4×3 (algoritmo × escenario), basado en simulación. "
            "Se formula H₀/H₁ por eje y se contrastan dos capas de evidencia: (A) series episódicas alineadas a OE; (B) KPI-gains de entrenamiento.",
            False,
        ),
        ("Hipótesis general (HG) — ranking multiobjetivo:", True),
        (
            "H₁(G): no existe un único MADRL que domine simultáneamente los tres ejes; el ranking integrado y los líderes por eje pueden diferir (trade-off Pareto). "
            "H₀(G): las distribuciones de desempeño entre algoritmos son idénticas en el agregado de ejes (omnibus).",
            False,
        ),
        ("Hipótesis específicas (contraste de efecto del algoritmo):", True),
        ("HE.1: H₁₁ = las distribuciones de desempeño de flexibilidad difieren entre algoritmos; H₀₁ = son idénticas.", False),
        ("HE.2: H₁₂ = las distribuciones de emisiones de CO₂ difieren entre algoritmos; H₀₂ = son idénticas.", False),
        ("HE.3: H₁₃ = las distribuciones de costo energético difieren entre algoritmos; H₀₃ = son idénticas.", False),
        (
            "Contrastación (α = 0,05): Shapiro–Wilk → Kruskal–Wallis → Mann–Whitney U (Holm) y Wilcoxon signed-rank (exploratorio). "
            "Corrida canónica madrl_v3_20260627_164047 (seed = 0; ≈50 episodios; HAPPO n = 49).",
            False,
        ),
        (
            "Veredicto: HG ranking/Pareto aceptada (MATD3 score 0,6667; sin dominador universal); superioridad omnibus KPI-gains no confirmada (p = 0,155). "
            "HE.1: rechazar H₀ capa A (p = 1,305×10⁻⁸), no rechazar capa B (p = 0,281). "
            "HE.2: rechazar H₀ capa A (p = 0,0439), no rechazar capa B (p = 0,546). "
            "HE.3: no rechazar H₀ capas A/B (p = 0,251 / 0,388); liderazgo MAAC descriptivo (Δcosto 9 515 EUR en E3). "
            "Fuentes: gdrive_objective_aligned_statistics.csv; hipotesis_estadisticas_madrl.csv.",
            False,
        ),
    ]

    cursor = before
    # Insert in reverse so order is preserved when using addnext on same anchor... 
    # Better insert sequentially with addnext on the newly created para.
    for text, bold in blocks:
        cursor = insert_paragraph_after(cursor, text, bold=bold)

    # Fix "no experimental" in alcances if present
    for p in doc.paragraphs:
        t = p.text or ""
        if "no experimental" in t.lower() and ("metodol" in t.lower() or "alcance" in t.lower() or "cuantitativo" in t.lower()):
            nt = re.sub(
                r"no experimental",
                "cuasiexperimental factorial 4×3 (algoritmo × escenario)",
                t,
                flags=re.IGNORECASE,
            )
            replace_paragraph_text(p, nt)
            info["actions"].append("cap1_alcance_cuasiexperimental")

    info["ok"] = True
    info["actions"].append(f"cap1_rewrote_block_before_idx_{just}")
    return info


def patch_cap3(doc: Document) -> dict:
    info: dict = {"ok": False, "actions": []}
    start = find_para_idx(doc, r"3\.1\b.*[Tt]ipo|[Tt]ipo y nivel de investigaci")
    if start is None:
        # broader
        start = find_para_idx(doc, r"Dise[nñ]o:\s*no experimental|Enfoque:\s*cuantitativo")
    if start is None:
        info["actions"].append("cap3_block_not_found")
        return info

    end = find_para_idx(doc, r"^3\.2\b|Variable independiente|3\.3\b|Unidad de an[aá]lisis", start=start + 1)
    if end is None:
        end = min(start + 25, len(doc.paragraphs))

    delete_paragraphs_between(doc, start, end, keep_start=False)
    # Find next section after deletion
    nxt = find_para_idx(doc, r"^3\.2\b|Variable independiente|3\.3\b|Unidad de an[aá]lisis|Datos utilizados")
    if nxt is None:
        info["actions"].append("cap3_next_lost")
        return info
    before = doc.paragraphs[nxt - 1] if nxt > 0 else doc.paragraphs[0]

    blocks = [
        ("3.1 Tipo y nivel de investigación", True),
        ("Enfoque: cuantitativo (Hernández-Sampieri et al.).", False),
        ("Tipo: aplicada (Tamayo y Tamayo; Arias).", False),
        ("Nivel: comparativo y propositivo (con componente descriptivo de KPIs).", False),
        (
            "Diseño: cuasiexperimental, factorial 4×3 (algoritmo MADRL × escenario E1/E2/E3), basado en simulación computacional. "
            "Se manipula deliberadamente la variable independiente (algoritmo y pesos de recompensa por escenario) bajo protocolo fijo; "
            "no hay aleatorización de unidades naturales ni sujetos humanos, por lo que no constituye experimento puro "
            "(Campbell & Stanley vía Hernández-Sampieri; Bunge sobre experimentación/simulación controlada).",
            False,
        ),
        (
            "Método: modelamiento computacional Dec-POMDP/CTDE, simulación CityLearn v2/v3, comparación de algoritmos MADRL "
            "y análisis no paramétrico de KPIs y recompensa. Se descarta el rótulo “no experimental” porque la VI se manipula sistemáticamente.",
            False,
        ),
        (
            "Variable independiente (tratamiento): algoritmo MADRL — HAPPO, MASAC, MATD3 y MAAC — bajo Dec-POMDP y CTDE, con escenarios E1/E2/E3. "
            "La capa CityLearn v3 es el entorno común del cuasiexperimento, no la VI primaria.",
            False,
        ),
        (
            "Variable dependiente: desempeño coordinado en flexibilidad, CO₂ y costos (KPIs CityLearn v2 y recompensa episódica). "
            "Métricas primarias: KPIs y reward; accuracy/precision/recall/F1 no son métricas centrales de este diseño.",
            False,
        ),
        (
            "Pruebas: Shapiro–Wilk → Kruskal–Wallis → Mann–Whitney U (Holm) → Wilcoxon (exploratorio). "
            "Friedman solo si hubiera ≥3 semillas (no aplicable con seed = 0). "
            "Dos capas de evidencia (no fusionar): (A) series episódicas OE-alineadas; (B) KPI-gains de entrenamiento.",
            False,
        ),
    ]
    cursor = before
    for text, bold in blocks:
        cursor = insert_paragraph_after(cursor, text, bold=bold)

    # Fix remaining "no experimental" in Cap 3 region
    for p in doc.paragraphs:
        t = p.text or ""
        if re.search(r"no experimental", t, re.I) and re.search(r"dise[nñ]o|metodol|simulaci", t, re.I):
            nt = re.sub(r"no experimental", "cuasiexperimental", t, flags=re.I)
            replace_paragraph_text(p, nt)
            info["actions"].append("cap3_replaced_no_experimental")

    info["ok"] = True
    info["actions"].append("cap3_rewrote_31")
    return info


def patch_cap6(doc: Document) -> dict:
    info: dict = {"ok": False, "actions": []}
    # Insert after principales hallazgos / before 6.2 Limitaciones
    start = find_para_idx(doc, r"6\.1\b|Principales hallazgos")
    lim = find_para_idx(doc, r"6\.2\b|Limitaciones encontradas", start=start or 0)
    if lim is None:
        info["actions"].append("cap6_limitaciones_not_found")
        return info

    # If already patched, skip duplicate
    existing = find_para_idx(doc, r"Veredicto de hip[oó]tesis")
    if existing is not None and (start is None or existing > (start or 0)) and existing < lim:
        # Replace existing block between existing and lim
        delete_paragraphs_between(doc, existing, lim, keep_start=False)
        lim = find_para_idx(doc, r"6\.2\b|Limitaciones encontradas")
        if lim is None:
            info["actions"].append("cap6_limitaciones_lost_after_delete")
            return info

    before = doc.paragraphs[lim - 1] if lim > 0 else doc.paragraphs[0]
    cursor = before
    blocks = [
        ("6.1.1 Veredicto de hipótesis (aceptación / rechazo)", True),
        (
            "Diseño adoptado: cuasiexperimental factorial 4×3; formulación PG/OG tipo ranking–Pareto; "
            "contraste H₀/H₁ por eje con dos capas (A = episódica OE-alineada; B = KPI-gains). α = 0,05.",
            False,
        ),
        (
            "HG (ranking/Pareto): ACEPTADA como ranking multiobjetivo sin dominador universal (MATD3 score 0,6667; MAAC lidera costos). "
            "Superioridad omnibus en KPI-gains: H₀ no rechazada (p = 0,155).",
            False,
        ),
        (
            "HE.1 (flexibilidad): H₀ rechazada en capa A (p = 1,305×10⁻⁸); H₀ no rechazada en capa B (p = 0,281). Cumplimiento de OE.1: sí (comparativo).",
            False,
        ),
        (
            "HE.2 (CO₂): H₀ rechazada en capa A (p = 0,0439, ε² ≈ 0,029); H₀ no rechazada en capa B (p = 0,546). Cumplimiento de OE.2: sí (descriptivo + inferencia episódica débil).",
            False,
        ),
        (
            "HE.3 (costos): H₀ no rechazada en capas A (p = 0,251) ni B (p = 0,388). Liderazgo MAAC: descriptivo. "
            "Cumplimiento de OE.3: sí a nivel identificación comparativa; no a nivel superioridad omnibus.",
            False,
        ),
        (
            "Nota. No se fusionan capas A y B. Accuracy/precision/recall/F1 no intervienen en este veredicto "
            "(métricas no primarias del control continuo MADRL). Veredicto documental 2026-07-18.",
            False,
        ),
    ]
    for text, bold in blocks:
        cursor = insert_paragraph_after(cursor, text, bold=bold)

    # Optional table after note
    headers = ["Hipótesis", "Decisión", "Fundamento"]
    rows = [
        [
            "HG",
            "Ranking/Pareto aceptado; omnibus KPI-gains no confirma superioridad",
            "p=0,155 capa B; score MATD3 0,6667",
        ],
        ["HE.1", "Rechazar H₀ capa A; no rechazar capa B", "p=1,305e-8 / p=0,281"],
        ["HE.2", "Rechazar H₀ capa A; no rechazar capa B", "p=0,0439 / p=0,546"],
        ["HE.3", "No rechazar H₀ (A y B)", "p=0,251 / p=0,388; MAAC descriptivo"],
    ]
    insert_table_after(cursor, headers, rows)
    info["ok"] = True
    info["actions"].append("cap6_inserted_veredicto")
    return info


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(str(path))
        return path
    except PermissionError:
        alt = path.with_name(path.stem + f"_VEREDICTO_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(str(alt))
        return alt


def ensure_source(path: Path) -> Path | None:
    if path.exists():
        return path
    fb = FALLBACK_FROM.get(path)
    if fb and fb.exists():
        shutil.copy2(fb, path)
        return path
    return None


def patch_one(path: Path) -> dict:
    src = ensure_source(path)
    if src is None:
        return {"file": str(path), "ok": False, "error": "missing"}
    doc = Document(str(src))
    r1 = patch_cap1(doc)
    r3 = patch_cap3(doc)
    r6 = patch_cap6(doc)
    out = save_doc(doc, path)
    return {
        "file": str(path),
        "saved": str(out),
        "ok": bool(r1.get("ok") or r3.get("ok") or r6.get("ok")),
        "cap1": r1,
        "cap3": r3,
        "cap6": r6,
    }


def verify(path: Path) -> dict:
    if not path.exists():
        return {"file": str(path), "checks": {}}
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    checks = {
        "has_cuasiexperimental": "cuasiexperimental" in text.lower(),
        "has_veredicto": "veredicto de hipótesis" in text.lower() or "veredicto de hipotesis" in text.lower(),
        "has_ranking_pareto": "ranking" in text.lower() and "pareto" in text.lower(),
        "has_two_layers": "capa a" in text.lower() or "dos capas" in text.lower(),
        "no_experimental_remaining_cap1_style": len(
            re.findall(r"dise[nñ]o[:\s].{0,40}no experimental", text, flags=re.I)
        ),
    }
    return {"file": path.name, "checks": checks}


def main() -> int:
    results = []
    for t in TARGETS:
        results.append(patch_one(t))
    verifications = [verify(Path(r["saved"])) for r in results if r.get("saved")]
    import json

    report = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "results": results,
        "verifications": verifications,
    }
    out_json = DOCS / "VEREDICTO_WORD_PATCH_REPORT_2026-07-18.json"
    out_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    # Success if at least one file saved with some ok
    ok = any(r.get("ok") for r in results)
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
