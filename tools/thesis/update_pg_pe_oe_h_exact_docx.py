#!/usr/bin/env python
"""Actualiza PG/PE/OG/OE/H0G–HE31 en Word tesis con formulaciones exactas del autor."""

from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

import sys

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import CANONS, existing_canons  # noqa: E402

BACKUP_DIR = REPO / "outputs" / "_word_backups"
LOG = REPO / "tmp" / "update_pg_pe_oe_h_exact_docx.log"

# Exact locked wording (author-validated). Do not paraphrase.
PG = (
    "¿En qué medida el algoritmo MADRL (aprendizaje por refuerzo profundo multiagente) "
    "impacta en la gestión coordinada de la flexibilidad energética, las emisiones de CO₂ "
    "y los costos energéticos en comunidades inteligentes de la ciudad de Iquitos, y cuál "
    "de los algoritmos presenta el mejor desempeño a nivel global?"
)
PE1 = (
    "PE.1: ¿En qué medida el algoritmo MADRL impacta en la flexibilidad energética en "
    "comunidades inteligentes de la ciudad de Iquitos, y cuál de los algoritmos presenta "
    "el mejor desempeño en el escenario E1?"
)
PE2 = (
    "PE.2: ¿En qué medida el algoritmo MADRL impacta en las emisiones de CO₂ en "
    "comunidades inteligentes de la ciudad de Iquitos, y cuál de los algoritmos presenta "
    "el mejor desempeño en el escenario E2?"
)
PE3 = (
    "PE.3: ¿En qué medida el algoritmo MADRL impacta en los costos energéticos en "
    "comunidades inteligentes de la ciudad de Iquitos, y cuál de los algoritmos presenta "
    "el mejor desempeño en el escenario E3?"
)
OG = (
    "OG. - Determinar el impacto de los algoritmos aprendizaje por refuerzo profundo "
    "multiagente (MADRLs) en la gestión coordinada de la flexibilidad energética, las "
    "emisiones de CO₂ y los costos energéticos en comunidades inteligentes de la ciudad "
    "de Iquitos, e identificar cuál de los algoritmos presenta el mejor desempeño a nivel global."
)
OE1 = (
    "OE.1: Determinar el impacto de los algoritmos MADRLs en la flexibilidad energética "
    "en comunidades inteligentes de la ciudad de Iquitos e identificar cuál de los "
    "algoritmos presenta el mejor desempeño en el escenario E1."
)
OE2 = (
    "OE.2: Determinar el impacto de los algoritmos MADRLs en las emisiones de CO₂ en "
    "comunidades inteligentes de la ciudad de Iquitos e identificar cuál de los "
    "algoritmos presenta el mejor desempeño en el escenario E2."
)
OE3 = (
    "OE.3: Determinar el impacto de los algoritmos MADRLs en los costos energéticos "
    "en comunidades inteligentes de la ciudad de Iquitos e identificar cuál de los "
    "algoritmos presenta el mejor desempeño en el escenario E3."
)
H0G = (
    "H0G.-El algoritmo MADRL no impacta de manera estadísticamente significativa y "
    "diferenciada en la gestión coordinada de la flexibilidad energética, las emisiones "
    "de CO₂ y los costos energéticos en comunidades inteligentes de la ciudad de Iquitos, "
    "y no existen diferencias significativas en el desempeño global de los algoritmos."
)
H1G = (
    "H1G.- El algoritmo MADRL impacta de manera estadísticamente significativa y "
    "diferenciada en la gestión coordinada de la flexibilidad energética, las emisiones "
    "de CO₂ y los costos energéticos en comunidades inteligentes de la ciudad de Iquitos, "
    "y el desempeño global difiere entre los algoritmos."
)
HE10 = (
    "HE10.- El algoritmo MADRL no impacta de manera estadísticamente significativa en "
    "la flexibilidad energética en comunidades inteligentes de la ciudad de Iquitos, y "
    "no existen diferencias significativas entre los algoritmos evaluados en el escenario E1."
)
HE11 = (
    "HE11.- El algoritmo MADRL impacta de manera estadísticamente significativa en la "
    "flexibilidad energética en comunidades inteligentes de la ciudad de Iquitos, y "
    "existen diferencias significativas entre los algoritmos evaluados en el escenario E1."
)
HE20 = (
    "HE20.- El algoritmo MADRL no impacta de manera estadísticamente significativa en "
    "las emisiones de CO₂ en comunidades inteligentes de la ciudad de Iquitos, y no "
    "existen diferencias significativas entre los algoritmos evaluados en el escenario E2."
)
HE21 = (
    "HE21.- El algoritmo MADRL impacta de manera estadísticamente significativa en las "
    "emisiones de CO₂ en comunidades inteligentes de la ciudad de Iquitos, y existen "
    "diferencias significativas entre los algoritmos evaluados en el escenario E2."
)
HE30 = (
    "HE30.-El algoritmo MADRL no impacta de manera estadísticamente significativa en "
    "los costos energéticos en comunidades inteligentes de la ciudad de Iquitos, y no "
    "existen diferencias significativas entre los algoritmos evaluados en el escenario E3."
)
HE31 = (
    "HE31.-El algoritmo MADRL impacta de manera estadísticamente significativa en los "
    "costos energéticos en comunidades inteligentes de la ciudad de Iquitos, y existen "
    "diferencias significativas entre los algoritmos evaluados en el escenario E3."
)

HEADINGS = [
    (("1.2.1.1", "Formulación del problema general", "Formulacion del problema general"), "1.2.1.1 Formulación del problema general"),
    (("1.2.1.2", "Formulación de los problemas específicos", "Formulacion de los problemas especificos"), "1.2.1.2 Formulación de los problemas específicos"),
    (("1.3.1.1", "Objetivo general"), "1.3.1.1 Objetivo general"),
    (("1.3.1.2", "Objetivos específicos", "Objetivos especificos"), "1.3.1.2 Objetivos específicos"),
    (("1.3.2 Hipótesis", "1.3.2 Hipotesis"), "1.3.2 Hipótesis"),
    (("1.3.2.1", "Hipótesis general", "Hipotesis general"), "1.3.2.1 Hipótesis general"),
    (("1.3.2.2", "Hipótesis específicas", "Hipotesis especificas"), "1.3.2.2 Hipótesis específicas"),
]

TARGET_DOCS = list(CANONS)


def clean(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _starts_with_any(text: str, labels: tuple[str, ...]) -> bool:
    t = clean(text)
    tl = t.lower()
    for label in labels:
        lab = clean(label)
        if t.startswith(lab) or tl.startswith(lab.lower()):
            return True
        bare = lab.rstrip(".:- ").lower()
        for sep in (" ", ".", ":", ".-", "-", ":"):
            if tl.startswith(bare + sep):
                return True
    return False


def is_statement_paragraph(text: str, labels: tuple[str, ...], *, markers: tuple[str, ...]) -> bool:
    t = clean(text)
    if not _starts_with_any(t, labels):
        return False
    if len(t) > 900:
        return False
    tl = t.lower()
    return any(m in tl for m in markers)


def is_pg_paragraph(text: str) -> bool:
    t = clean(text)
    tl = t.lower()
    if not t.startswith("¿") and "problema general" not in tl and "en qué medida el algoritmo madrl" not in tl:
        # allow standalone question
        if not t.startswith("¿En qué medida el algoritmo MADRL"):
            return False
    if "gestión coordinada" not in tl and "gestion coordinada" not in tl:
        return False
    if "iquitos" not in tl:
        return False
    if len(t) > 900:
        return False
    return "flexibilidad" in tl and ("co₂" in tl or "co2" in tl) and "costos" in tl


STATEMENT_MAP: list[tuple[tuple[str, ...], str, tuple[str, ...]]] = [
    (("PE.1:", "PE.1.", "PE1:", "**PE.1:**"), PE1, ("flexibilidad", "e1")),
    (("PE.2:", "PE.2.", "PE2:", "**PE.2:**"), PE2, ("emisiones", "e2")),
    (("PE.3:", "PE.3.", "PE3:", "**PE.3:**"), PE3, ("costos", "e3")),
    (("OG. -", "OG.-", "OG.", "OG:"), OG, ("determinar el impacto", "madrls")),
    (("OE.1:", "OE.1.", "OE1:"), OE1, ("determinar el impacto", "e1")),
    (("OE.2:", "OE.2.", "OE2:"), OE2, ("determinar el impacto", "e2")),
    (("OE.3:", "OE.3.", "OE3:"), OE3, ("determinar el impacto", "e3")),
    (("H0G.-", "H0G."), H0G, ("no impacta", "estadísticamente")),
    (("H1G.-", "H1G."), H1G, ("impacta de manera", "estadísticamente")),
    (("HE10.-", "HE10."), HE10, ("no impacta", "e1")),
    (("HE11.-", "HE11."), HE11, ("impacta de manera", "e1")),
    (("HE20.-", "HE20."), HE20, ("no impacta", "e2")),
    (("HE21.-", "HE21."), HE21, ("impacta de manera", "e2")),
    (("HE30.-", "HE30."), HE30, ("no impacta", "e3")),
    (("HE31.-", "HE31."), HE31, ("impacta de manera", "e3")),
]


def replace_statements(doc: Document) -> dict[str, int]:
    counts = {labels[0]: 0 for labels, _, _ in STATEMENT_MAP}
    counts["PG"] = 0
    for para in doc.paragraphs:
        raw = para.text or ""
        if not raw.strip():
            continue
        if is_pg_paragraph(raw):
            # strip leading labels like "Problema general (PG):"
            set_paragraph_text(para, PG)
            counts["PG"] += 1
            continue
        for labels, new_text, markers in STATEMENT_MAP:
            if is_statement_paragraph(raw, labels, markers=markers):
                if clean(raw) != clean(new_text):
                    set_paragraph_text(para, new_text)
                counts[labels[0]] = counts.get(labels[0], 0) + 1
                break
    return counts


def replace_in_tables(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    raw = para.text or ""
                    if not raw.strip():
                        continue
                    if is_pg_paragraph(raw):
                        set_paragraph_text(para, PG)
                        n += 1
                        continue
                    for labels, new_text, markers in STATEMENT_MAP:
                        if is_statement_paragraph(raw, labels, markers=markers):
                            if clean(raw) != clean(new_text):
                                set_paragraph_text(para, new_text)
                                n += 1
                            break
    return n


def ensure_headings(doc: Document) -> list[str]:
    notes: list[str] = []
    for para in doc.paragraphs:
        t = clean(para.text)
        if not t or len(t) > 140:
            continue
        style = (para.style.name if para.style is not None else "") or ""
        is_heading = style.lower().startswith("heading") or len(t) < 100
        if not is_heading:
            continue
        for keys, exact in HEADINGS:
            if any(k.lower() in t.lower() for k in keys):
                if clean(t) != clean(exact):
                    set_paragraph_text(para, exact)
                    notes.append(f"heading -> {exact}")
                break
    return notes


def process_doc(path: Path, log: list[str]) -> None:
    if not path.is_file():
        log.append(f"SKIP missing: {path.name}")
        return
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{path.stem}_pre_pg_pe_oe_h_{stamp}{path.suffix}"
    try:
        shutil.copy2(path, backup)
        log.append(f"BACKUP {path.name} -> {backup.name}")
    except OSError as exc:
        log.append(f"BACKUP_FAIL {path.name}: {exc}")

    doc = Document(str(path))
    notes = ensure_headings(doc)
    counts = replace_statements(doc)
    table_n = replace_in_tables(doc)

    # Special wrappers
    for para in doc.paragraphs:
        t = clean(para.text)
        tl = t.lower()
        if tl.startswith("objetivo general (og)"):
            set_paragraph_text(para, OG)
            counts["OG. -"] = counts.get("OG. -", 0) + 1
        elif tl.startswith("problema general (pg)"):
            set_paragraph_text(para, PG)
            counts["PG"] = counts.get("PG", 0) + 1

    try:
        doc.save(str(path))
        saved = path
    except OSError as exc:
        raise SystemExit(
            f"No se pudo guardar el Word canónico {path.name} (¿abierto en Word?): {exc}"
        ) from exc

    log.append(f"UPDATED {saved.name}")
    for k, v in counts.items():
        log.append(f"  {k}: {v}")
    log.append(f"  table_cell_replacements: {table_n}")
    for note in notes:
        log.append(f"  {note}")

    blob = "\n".join(p.text for p in Document(str(saved)).paragraphs)
    for label, exact in [
        ("PG", PG),
        ("PE.1", PE1),
        ("PE.2", PE2),
        ("PE.3", PE3),
        ("OG", OG),
        ("OE.1", OE1),
        ("OE.2", OE2),
        ("OE.3", OE3),
        ("H0G", H0G),
        ("H1G", H1G),
        ("HE10", HE10),
        ("HE11", HE11),
        ("HE20", HE20),
        ("HE21", HE21),
        ("HE30", HE30),
        ("HE31", HE31),
    ]:
        ok = exact in blob or clean(exact) in clean(blob)
        log.append(f"  VERIFY {label}: {'OK' if ok else 'MISSING'}")


def main() -> int:
    docs = existing_canons()
    if not docs:
        docs = [p for p in TARGET_DOCS if p.is_file()]

    log: list[str] = [f"run={datetime.now().isoformat()}", f"targets={[p.name for p in docs]}"]
    if not docs:
        log.append("ERROR: ningún Word canónico encontrado")
        LOG.parent.mkdir(parents=True, exist_ok=True)
        LOG.write_text("\n".join(log) + "\n", encoding="utf-8")
        print("\n".join(log))
        return 1

    for path in docs:
        try:
            process_doc(path, log)
        except Exception as exc:  # noqa: BLE001
            log.append(f"ERROR {path.name}: {exc}")
    LOG.parent.mkdir(parents=True, exist_ok=True)
    LOG.write_text("\n".join(log) + "\n", encoding="utf-8")
    print("\n".join(log))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
