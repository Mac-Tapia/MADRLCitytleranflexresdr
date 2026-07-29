"""Genera resumen multiobjetivo Colab/Drive en Word (.docx) para la tesis."""

from __future__ import annotations

import csv
import datetime as dt
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[2]
RUN_ID = "madrl_v3_20260627_164047"
MO_DIR = REPO / "outputs" / RUN_ID / "resumen_comparativo" / "multiobjetivo"
OUT_DOCX = MO_DIR / "RESUMEN_MULTIOBJETIVO_TESIS.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, caption: str | None = None) -> None:
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = GREY
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 16.0) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)
    doc.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    doc.add_paragraph()


def fmt_num(val: str | float | None, decimals: int = 0) -> str:
    if val is None or val == "":
        return "-"
    try:
        x = float(val)
        if decimals == 0:
            return f"{x:,.0f}"
        return f"{x:,.{decimals}f}"
    except (TypeError, ValueError):
        return str(val)


def build(out_docx: Path = OUT_DOCX) -> Path:
    district_rows = read_csv(MO_DIR / "district_objectives_by_algorithm.csv")
    inventory_rows = read_csv(MO_DIR / "building_inventory_multiobjective.csv")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Portada
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Resultados multiobjetivo MADRL — Iquitos")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Flexibilidad (OE1) · Emisiones CO₂ (OE2) · Costo energético (OE3)\n"
        f"Corrida canónica: {RUN_ID}"
    ).font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "17 edificios · 185 cargadores EV · MASAC / MATD3 / MAAC\n"
        f"Generado: {dt.date.today().strftime('%d/%m/%Y')}"
    ).font.size = Pt(10)

    doc.add_page_break()

    doc.add_heading("1. Alcance del análisis", level=1)
    doc.add_paragraph(
        "El análisis desagrega KPIs en nivel distrito (citylearn_v3_report.all_values) "
        "y nivel edificio (building_behavior_summary.csv, 17 agentes). "
        "Fuente: artefactos Colab/Drive en outputs/_drive_madrl/."
    )
    doc.add_paragraph(
        "Elementos controlados: BESS (electrical_storage), cargadores EV, carga desplazable (washing_machine). "
        "No controlados: non_shiftable_load, refrigeración/ACS modeladas, generación FV fija."
    )

    doc.add_heading("2. Mejor algoritmo por objetivo (distrito)", level=1)
    add_table(
        doc,
        ["Objetivo", "Escenario", "Mejor", "Valor principal"],
        [
            ["OE1 Flexibilidad", "E1", "MATD3", "flex_composite = 1.001"],
            ["OE2 Emisiones CO₂", "E2", "MATD3", "ΔCO₂ = 23,070 kg"],
            ["OE3 Costo energético", "E3", "MAAC", "Δcosto = 9,515 EUR"],
        ],
        caption="Tabla 1. Ranking por objetivo multiobjetivo.",
    )

    doc.add_heading("3. KPIs de distrito — todos los algoritmos", level=1)
    dist_table: list[list[str]] = []
    for row in district_rows:
        dist_table.append(
            [
                row["algorithm"],
                row["scenario"],
                fmt_num(row["flex_composite"], 3),
                fmt_num(row["carbon_emissions_delta_kg"]),
                fmt_num(row["electricity_cost_delta_eur"]),
                f"{float(row['ev_departure_success_rate']) * 100:.1f}%",
            ]
        )
    add_table(
        doc,
        ["Algoritmo", "Escenario", "Flex", "ΔCO₂ (kg)", "ΔCosto (EUR)", "EV éxito"],
        dist_table,
        caption="Tabla 2. KPIs agregados de distrito por algoritmo y escenario.",
    )

    doc.add_heading("4. Figuras comparativas", level=1)
    main_figs = [
        (MO_DIR / "drive_district_objectives.png", "Figura 1. KPIs multiobjetivo — distrito."),
        (MO_DIR / "drive_building_E1_flex_composite_proxy.png", "Figura 2. OE1 Flexibilidad por edificio."),
        (MO_DIR / "drive_building_E2_carbon_emissions_delta_kgco2.png", "Figura 3. OE2 Δ CO₂ por edificio (kg)."),
        (MO_DIR / "drive_building_E3_electricity_cost_delta_eur.png", "Figura 4. OE3 Δ costo por edificio (EUR)."),
        (MO_DIR / "drive_building_ev_inventory.png", "Figura 5. Inventario EV por edificio (185 tomas)."),
        (MO_DIR / "drive_building_ev_success_matd3_e2.png", "Figura 6. Desempeño EV — MATD3 / E2."),
    ]
    for path, caption in main_figs:
        doc.add_page_break()
        add_figure(doc, path, caption)

    doc.add_page_break()
    doc.add_heading("5. Inventario de edificios", level=1)
    inv_table = [
        [
            f"B{int(row['building_id']):02d}",
            row["nombre"][:40],
            row["tipo_uso"],
            row["ev_total"],
            row["elementos_controlados"][:50] + ("…" if len(row["elementos_controlados"]) > 50 else ""),
        ]
        for row in inventory_rows
    ]
    add_table(
        doc,
        ["ID", "Edificio", "Tipo", "EV", "Controlados"],
        inv_table,
        caption="Tabla 3. Inventario multiobjetivo — 17 edificios Iquitos.",
    )

    doc.add_heading("6. Detalle por edificio (B01–B17)", level=1)
    doc.add_paragraph(
        "Cada figura siguiente compara MASAC, MATD3 y MAAC en los tres escenarios "
        "multiobjetivo para un edificio."
    )
    per_dir = MO_DIR / "por_edificio"
    for png in sorted(per_dir.glob("drive_building_B*_objectives.png")):
        bid = png.stem.replace("drive_building_", "").replace("_objectives", "")
        inv = next(
            (r for r in inventory_rows if f"B{int(r['building_id']):02d}" == bid),
            None,
        )
        name = inv["nombre"] if inv else bid
        doc.add_page_break()
        add_figure(doc, png, f"Figura — {bid} {name}: tres objetivos × algoritmo.", width_cm=15.5)

    doc.add_heading("7. Referencias de artefactos", level=1)
    for label, rel in (
        ("CSV distrito", "district_objectives_by_algorithm.csv"),
        ("CSV edificio", "building_objectives_by_algorithm.csv"),
        ("Inventario", "building_inventory_multiobjective.csv"),
        ("Detalle MASAC", "building_detail_masac_by_scenario.md"),
        ("Detalle MATD3", "building_detail_matd3_by_scenario.md"),
        ("Detalle MAAC", "building_detail_maac_by_scenario.md"),
    ):
        doc.add_paragraph(f"{label}: {MO_DIR / rel}", style="List Bullet")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    size_mb = out_docx.stat().st_size / 1024 / 1024
    print(f"Word generado: {out_docx}")
    print(f"Tamaño: {size_mb:.1f} MB")
    return out_docx


if __name__ == "__main__":
    build()
