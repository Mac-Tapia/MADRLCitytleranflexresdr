#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Regenerate Figura A.9: tamaño total en checkpoint_manifest.json por tratamiento.

Corrige la gráfica antigua (bytes crudos en notación científica, MASAC invisible,
etiquetas en inglés, HAPPO omitido) usando GB en escala logarítmica y los 12
tratamientos MADRL de la corrida canónica madrl_v3_20260627_164047.
"""
from __future__ import annotations

import json
import math
import re
import zipfile
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document

REPO = Path(__file__).resolve().parents[1]
FULL = REPO / "outputs" / "_drive_madrl" / "full_data"
FIG_DIR = REPO / "outputs" / "_drive_madrl" / "gdrive_20260627_164047_objective_analysis" / "figures"
TABLE_DIR = FIG_DIR.parent / "tables"
LEGACY_FIG = FULL / "analysis_real_drive" / "figures" / "checkpoint_manifest_bytes.png"
PNG = FIG_DIR / "checkpoint_manifest_total_size_by_treatment.png"
REPORT = FIG_DIR.parent / "figura_a9_checkpoint_size_report.json"

ALGOS = ["HAPPO", "MAAC", "MASAC", "MATD3"]
SCENARIOS = ["E1", "E2", "E3"]

DOC_TARGETS = [
    REPO / "docs" / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx",
    REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx",
]


def fmt_num(x: float, nd: int = 2) -> str:
    if x is None or (isinstance(x, float) and (math.isnan(x) or math.isinf(x))):
        return "N/D"
    return f"{float(x):,.{nd}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def scan_checkpoint_bytes() -> pd.DataFrame:
    rows = []
    for algo in ALGOS:
        for scen in SCENARIOS:
            man = FULL / algo / scen / "data" / "checkpoint_manifest.json"
            total_b = 0
            n = 0
            source = "missing"
            max_file_mb = 0.0
            max_rel = ""
            if man.exists():
                obj = json.loads(man.read_text(encoding="utf-8"))
                cps = obj.get("checkpoints") or []
                n = len(cps)
                source = str(man.relative_to(REPO)).replace("\\", "/")
                for ck in cps:
                    b = ck.get("bytes")
                    if b is None:
                        continue
                    b = int(b)
                    total_b += b
                    mb = b / (1024 * 1024)
                    if mb > max_file_mb:
                        max_file_mb = mb
                        max_rel = str(ck.get("relative_path") or "")
            rows.append(
                {
                    "algorithm": algo,
                    "scenario": scen,
                    "treatment": f"{algo}-{scen}",
                    "n_files_listed": n,
                    "total_bytes": total_b,
                    "total_mb": total_b / (1024 * 1024),
                    "total_gb": total_b / (1024**3),
                    "max_file_mb": max_file_mb,
                    "max_relative_path": max_rel,
                    "source": source,
                }
            )
    return pd.DataFrame(rows)


def _label_for_bar(gb: float, n: int) -> str:
    if n <= 0 or gb <= 0:
        return "0\n(n=0)"
    if gb < 0.01:
        return f"{gb * 1024:.2f} MB\n(n={n})"
    if gb < 1:
        return f"{gb:.3f}\n(n={n})"
    return f"{gb:.2f}\n(n={n})"


def plot_a9(df: pd.DataFrame) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    d = df.copy()
    d["algorithm"] = pd.Categorical(d["algorithm"], ALGOS, ordered=True)
    d["scenario"] = pd.Categorical(d["scenario"], SCENARIOS, ordered=True)
    d = d.sort_values(["algorithm", "scenario"])

    labels = d["treatment"].tolist()
    values = d["total_gb"].astype(float).tolist()
    ns = d["n_files_listed"].astype(int).tolist()

    # Escala log: HAPPO=0 se representa en el piso visual sin inventar bytes.
    floor_gb = 1e-4
    plot_vals = [max(v, floor_gb) if n > 0 else floor_gb for v, n in zip(values, ns)]

    fig, ax = plt.subplots(figsize=(10.2, 4.8))
    palette = {"HAPPO": "#7A7A7A", "MAAC": "#406A9F", "MASAC": "#2E8B57", "MATD3": "#C45C26"}
    colors = [palette.get(lab.split("-")[0], "#406A9F") for lab in labels]
    bars = ax.bar(labels, plot_vals, color=colors, edgecolor="white", linewidth=0.4)
    ax.set_yscale("log")
    ax.set_title("Tamaño total listado en manifiestos de checkpoint por algoritmo y escenario")
    ax.set_ylabel("GB listados en checkpoint_manifest.json (escala log)")
    ax.set_xlabel("Tratamiento (algoritmo × escenario)")
    ax.tick_params(axis="x", rotation=55)
    ax.grid(axis="y", alpha=0.28, which="both")
    ax.set_ylim(floor_gb, max(plot_vals) * 3.5)

    ymax = max(plot_vals)
    for bar, gb, n in zip(bars, values, ns):
        y_text = max(bar.get_height(), floor_gb)
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            y_text * 1.35,
            _label_for_bar(gb, n),
            ha="center",
            va="bottom",
            fontsize=6.8,
        )
        if n <= 0 or gb <= 0:
            bar.set_hatch("///")
            bar.set_alpha(0.55)

    ax.text(
        0.01,
        0.98,
        "HAPPO: sin checkpoint_manifest.json en madrl_v3_20260627_164047 (tamaño listado = 0). "
        "Escala logarítmica para hacer visibles MASAC (~MB) frente a MAAC (~38 GB).",
        transform=ax.transAxes,
        va="top",
        fontsize=7.2,
        color="#333333",
        wrap=True,
    )
    fig.tight_layout()
    fig.savefig(PNG, dpi=180, bbox_inches="tight")
    plt.close(fig)

    # Espejo legible en la carpeta analysis_real_drive (nombre legado del Word/informe).
    LEGACY_FIG.parent.mkdir(parents=True, exist_ok=True)
    LEGACY_FIG.write_bytes(PNG.read_bytes())

    d.to_csv(TABLE_DIR / "figura_a9_checkpoint_manifest_sizes.csv", index=False)
    (FULL / "analysis_real_drive" / "tables").mkdir(parents=True, exist_ok=True)
    d.to_csv(FULL / "analysis_real_drive" / "tables" / "figura_a9_checkpoint_manifest_sizes.csv", index=False)
    return PNG


def replace_paragraph_text(para, new_text: str) -> None:
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def patch_interpretation(doc: Document, new_interp: str, new_note: str) -> int:
    n = 0
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if "Figura A.9" in t:
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                tj = (doc.paragraphs[j].text or "").strip()
                if tj.startswith("Nota."):
                    replace_paragraph_text(doc.paragraphs[j], new_note)
                    n += 1
                if tj.startswith("Interpretacion de la figura") or tj.startswith("Interpretación de la figura"):
                    replace_paragraph_text(doc.paragraphs[j], new_interp)
                    n += 1
                    return n
            break
    return n


def replace_image_after_caption(docx_path: Path, png_path: Path) -> dict:
    info = {"file": docx_path.name, "replaced": False}
    if not docx_path.exists():
        info["error"] = "missing"
        return info
    png_bytes = png_path.read_bytes()
    caption_re = r"Figura A\.9|Tamano total listado en manifiestos|Tamaño total listado en manifiestos"
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            xml = zin.read("word/document.xml").decode("utf-8")
            rels = zin.read("word/_rels/document.xml.rels").decode("utf-8")
    except Exception as exc:
        info["error"] = f"{type(exc).__name__}: {exc}"
        return info
    m = re.search(caption_re, xml, flags=re.I)
    if not m:
        info["error"] = "caption not found"
        return info
    after = xml[m.end() :]
    blips = re.findall(r'a:blip[^>]*r:embed="(rId\d+)"', after)
    if not blips:
        info["error"] = "no blip after caption"
        return info
    rid = blips[0]
    rm = re.search(rf'Relationship[^>]*Id="{rid}"[^>]*Target="([^"]+)"', rels)
    if not rm:
        info["error"] = f"no relationship for {rid}"
        return info
    media = "word/" + rm.group(1).lstrip("/")
    tmp = docx_path.with_suffix(".docx.tmp_A9")
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            replaced = False
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.replace("\\", "/") == media.replace("\\", "/"):
                    data = png_bytes
                    replaced = True
                zout.writestr(item, data)
        if not replaced:
            tmp.unlink(missing_ok=True)
            info["error"] = f"media missing: {media}"
            return info
        tmp.replace(docx_path)
        info.update({"replaced": True, "rId": rid, "media": media})
    except PermissionError:
        alt = docx_path.with_name(docx_path.stem + "_A9_PATCHED.docx")
        if tmp.exists():
            tmp.replace(alt)
        info.update({"replaced": False, "saved_alt": alt.name, "note": "Word bloqueado"})
    return info


def build_texts(a9: pd.DataFrame) -> tuple[str, str]:
    total_gb = float(a9["total_gb"].sum())
    total_mb = float(a9["total_mb"].sum())
    n_files = int(a9["n_files_listed"].sum())
    max_row = a9.sort_values("max_file_mb", ascending=False).iloc[0]
    maac_mean = float(a9.loc[a9["algorithm"] == "MAAC", "total_gb"].mean())
    masac_mb = float(a9.loc[a9["algorithm"] == "MASAC", "total_mb"].mean())
    matd3_mb = float(a9.loc[a9["algorithm"] == "MATD3", "total_mb"].mean())
    interp = (
        "Interpretación de la figura. La figura resume el tamaño total listado en checkpoint_manifest.json "
        f"para los 12 tratamientos MADRL del proyecto (HAPPO, MAAC, MASAC, MATD3 × E1–E3). "
        f"Se listan {n_files} archivos con un total de {fmt_num(total_gb, 2)} GB ({fmt_num(total_mb, 2)} MB). "
        "HAPPO aparece en 0 GB porque no hay checkpoint_manifest.json en la corrida canónica. "
        f"MAAC concentra el mayor volumen (~{fmt_num(maac_mean, 2)} GB por escenario); "
        f"MATD3 ≈ {fmt_num(matd3_mb, 1)} MB y MASAC ≈ {fmt_num(masac_mb, 2)} MB por escenario "
        "(por eso la figura usa escala logarítmica). "
        f"El archivo individual más grande listado corresponde a {max_row['algorithm']}-{max_row['scenario']} "
        f"({fmt_num(max_row['max_file_mb'], 2)} MB). No se inventan bytes: solo se agregan los campos bytes del manifiesto."
    )
    note = (
        "Nota. La figura representa el tamaño total listado en los manifiestos de checkpoint por algoritmo y escenario. "
        "Fuente: elaboración propia a partir de outputs/_drive_madrl/full_data/{ALGO}/{E}/data/checkpoint_manifest.json "
        "de la corrida canónica madrl_v3_20260627_164047."
    )
    return interp, note


def main() -> int:
    a9 = scan_checkpoint_bytes()
    png = plot_a9(a9)
    interp, note = build_texts(a9)

    embeds = []
    patches = []
    for doc_path in DOC_TARGETS:
        if not doc_path.exists():
            patches.append({"file": doc_path.name, "error": "missing"})
            continue
        try:
            doc = Document(str(doc_path))
            n_patch = patch_interpretation(doc, interp, note)
            doc.save(str(doc_path))
            patches.append({"file": doc_path.name, "interp_patched": n_patch, "ok": True})
        except PermissionError:
            alt = doc_path.with_name(doc_path.stem + "_A9_PATCHED.docx")
            doc = Document(str(doc_path))
            n_patch = patch_interpretation(doc, interp, note)
            doc.save(str(alt))
            patches.append({"file": alt.name, "interp_patched": n_patch, "ok": False, "note": "Word bloqueado"})
            doc_path = alt
        embeds.append(replace_image_after_caption(doc_path, png))

    report = {
        "png": str(png.relative_to(REPO)).replace("\\", "/"),
        "legacy_png": str(LEGACY_FIG.relative_to(REPO)).replace("\\", "/"),
        "total_files": int(a9["n_files_listed"].sum()),
        "total_gb": float(a9["total_gb"].sum()),
        "total_mb": float(a9["total_mb"].sum()),
        "per_treatment": a9[["treatment", "n_files_listed", "total_gb", "total_mb", "source"]].to_dict(orient="records"),
        "problems_fixed": [
            "Eje Y en bytes crudos (4e10) sustituido por GB con escala log",
            "MASAC/MATD3 visibles frente a MAAC (~38 GB)",
            "HAPPO incluido con tamaño 0 (manifiesto ausente)",
            "Etiquetas y título en español",
            "Interpretación alineada al agregado real del manifiesto canónico",
        ],
        "doc_patches": patches,
        "embeds": embeds,
        "interpretation": interp,
        "note": note,
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({
        "png": report["png"],
        "total_gb": report["total_gb"],
        "total_files": report["total_files"],
        "embeds": embeds,
        "patches": patches,
        "report": str(REPORT.relative_to(REPO)).replace("\\", "/"),
    }, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
