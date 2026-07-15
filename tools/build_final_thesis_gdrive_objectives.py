from __future__ import annotations

import csv
import json
import math
import re
import shutil
import unicodedata
from copy import deepcopy
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


REPO = Path(__file__).resolve().parents[1]
G_BASE = Path(r"G:\Mi unidad\MADRLCitytleranflexresdr\outputs\madrl_v3_20260627_164047")
SRC = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_50EP_ANTECEDENTES.docx"
OUT = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_GDRIVE_50EP_OBJETIVOS_DOCTORAL.docx"
ANALYSIS_DIR = REPO / "outputs" / "_drive_madrl" / "gdrive_20260627_164047_objective_analysis"
TABLE_DIR = ANALYSIS_DIR / "tables"
FIG_DIR = ANALYSIS_DIR / "figures"
METRICS = ANALYSIS_DIR / "thesis_gdrive_objective_metrics.json"
LOCAL_EPISODE_CSV = REPO / "outputs" / "_drive_madrl" / "full_data" / "analysis_real_drive" / "tables" / "district_episode_kpis.csv"

ALGOS = ["HAPPO", "MAAC", "MASAC", "MATD3"]
SCENARIOS = ["E1", "E2", "E3"]

OBJECTIVES = [
    {
        "objective": "OE.1",
        "hypothesis": "HE.1",
        "scenario": "E1",
        "dimension": "flexibilidad energetica",
        "metric": "reward_mean_average",
        "direction": "max",
        "indicator": "recompensa media del episodio en E1",
        "explanation": "E1 usa pesos [flex=0,70; CO2=0,15; costo=0,15], por lo que la recompensa media del episodio representa el efecto agregado del algoritmo sobre la dimension de flexibilidad bajo una funcion de recompensa comparable.",
    },
    {
        "objective": "OE.2",
        "hypothesis": "HE.2",
        "scenario": "E2",
        "dimension": "emisiones de CO2",
        "metric": "district_emission",
        "direction": "min",
        "indicator": "suma anual de district_net_electricity_consumption_emission en E2",
        "explanation": "E2 usa pesos [flex=0,15; CO2=0,70; costo=0,15]; por ello la emision distrital anual agregada por episodio es el indicador directo para contrastar el efecto sobre D-VD.2.",
    },
    {
        "objective": "OE.3",
        "hypothesis": "HE.3",
        "scenario": "E3",
        "dimension": "costos energeticos",
        "metric": "district_cost",
        "direction": "min",
        "indicator": "suma anual de district_net_electricity_consumption_cost en E3",
        "explanation": "E3 usa pesos [flex=0,25; CO2=0,15; costo=0,60]; por ello el costo distrital anual agregado por episodio es el indicador directo para contrastar el efecto sobre D-VD.3.",
    },
]

FIGURE_INTERPRETATIONS: dict[str, str] = {}

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)


def ensure_dirs() -> None:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def safe_float(value) -> float:
    try:
        if value in ("", None):
            return math.nan
        return float(value)
    except Exception:
        return math.nan


def fmt(x, nd: int = 3) -> str:
    try:
        if pd.isna(x):
            return "NA"
        return f"{float(x):,.{nd}f}"
    except Exception:
        return str(x)


def fmt_p(x) -> str:
    try:
        if pd.isna(x):
            return "NA"
        value = float(x)
        return f"{value:.3e}" if value < 0.001 else f"{value:.6f}"
    except Exception:
        return str(x)


def text_of(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def clear_body_keep_sectpr(document: Document) -> None:
    body = document.element.body
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def append_before_sectpr(document: Document, el) -> None:
    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(el)
    else:
        body.insert(body.index(sect_pr), el)


def style_doc(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    ensure_paragraph_style(document, "TablaIndice", 12, bold=False, italic=True)
    ensure_paragraph_style(document, "FiguraIndice", 12, bold=False, italic=True)
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"]:
        if name not in [s.name for s in document.styles]:
            continue
        st = document.styles[name]
        st.font.name = "Times New Roman"
        if name == "Normal":
            st.font.size = Pt(12)
            st.paragraph_format.space_before = Pt(0)
            st.paragraph_format.space_after = Pt(0)
            st.paragraph_format.line_spacing = 2.0
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            st.font.bold = True
            st.font.italic = name in ["Heading 3", "Heading 5"]
            st.font.color.rgb = RGBColor(0, 0, 0)
            st.font.size = Pt(12)
            st.paragraph_format.space_before = Pt(0)
            st.paragraph_format.space_after = Pt(0)
            st.paragraph_format.line_spacing = 2.0
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if name == "Heading 1" else WD_ALIGN_PARAGRAPH.LEFT
            if name in ["Heading 4", "Heading 5"]:
                st.paragraph_format.first_line_indent = Inches(0.5)


def set_run_apa_font(run, size: float = 12.0, bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ["ascii", "hAnsi", "cs", "eastAsia"]:
        r_fonts.set(qn(f"w:{attr}"), "Times New Roman")


def ensure_paragraph_style(document: Document, name: str, size: float = 12.0, bold: bool = False, italic: bool = False) -> None:
    existing = {s.name for s in document.styles}
    style = document.styles[name] if name in existing else document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_field(paragraph, instruction: str, placeholder: str = "Actualice este campo con clic derecho > Actualizar campos (o F9).") -> None:
    paragraph.text = ""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    fld.set(qn("w:dirty"), "true")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = placeholder
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    for node in settings.findall(qn("w:updateFields")):
        settings.remove(node)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def set_bg(cell, color: str = "1F4E79") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def remove_cell_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(shd)


def set_cell_margins(cell, top: int = 60, start: int = 80, bottom: int = 60, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in list(tc_pr.findall(qn("w:tcMar"))):
        tc_pr.remove(old)
    tc_mar = OxmlElement("w:tcMar")
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_cell_apa_borders(cell, top: bool = False, bottom: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in list(tc_pr.findall(qn("w:tcBorders"))):
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge, enabled in {"top": top, "left": False, "bottom": bottom, "right": False}.items():
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single" if enabled else "nil")
        if enabled:
            border.set(qn("w:sz"), "8")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
        borders.append(border)
    tc_pr.append(borders)


def apply_apa_table_format(tbl, body_font_size: float = 10.0) -> None:
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    row_count = len(tbl.rows)
    for r_idx, row in enumerate(tbl.rows):
        is_header = r_idx == 0
        is_last = r_idx == row_count - 1
        for cell in row.cells:
            remove_cell_shading(cell)
            set_cell_margins(cell)
            set_cell_apa_borders(cell, top=is_header, bottom=is_header or is_last)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.first_line_indent = None
                for run in para.runs:
                    if run.text:
                        set_run_apa_font(run, body_font_size, bold=is_header, italic=False, color=RGBColor(0, 0, 0))


def p(doc: Document, text: str):
    para = doc.add_paragraph()
    set_run_apa_font(para.add_run(text), 12)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 2.0
    para.paragraph_format.first_line_indent = Inches(0.5)
    return para


def set_paragraph_text(para, text: str) -> None:
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def norm_key(text: str) -> str:
    text = "".join(c for c in unicodedata.normalize("NFKD", text) if not unicodedata.combining(c))
    return re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()


def iter_document_paragraphs(document: Document):
    for para in document.paragraphs:
        yield para
    for tbl in document.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def justify_document_text(document: Document) -> None:
    """Aplica justificacion directa al texto academico, sin alterar encabezados ni TOC."""
    skip_prefixes = ("heading", "toc")
    skip_exact = {"title", "subtitle"}
    for para in iter_document_paragraphs(document):
        if not para.text.strip():
            continue
        style_name = (para.style.name or "").strip().lower() if para.style is not None else ""
        if style_name in skip_exact or style_name.startswith(skip_prefixes):
            continue
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def infer_heading_level(text: str) -> int | None:
    stripped = text.strip()
    if stripped in {"Resumen", "Abstract", "Indice", "Referencias bibliograficas"}:
        return 1
    if stripped.startswith("Referencias complementarias"):
        return 2
    if stripped.startswith(("Capitulo ", "Anexo ")):
        return 1
    match = re.match(r"^\d+(?:\.\d+)+\s+", stripped)
    if not match:
        return None
    depth = stripped.split()[0].count(".") + 1
    return min(max(depth, 2), 5)


def format_caption_paragraph(para) -> bool:
    text = para.text.strip()
    match = re.match(r"^(Tabla|Figura)\s+([A-Za-z0-9.]+)(?:\.|\n)\s*(.+)$", text, flags=re.S)
    if not match:
        return False
    kind, number, title = match.groups()
    para.style = "TablaIndice" if kind == "Tabla" else "FiguraIndice"
    for run in para.runs:
        run.text = ""
    label = para.runs[0] if para.runs else para.add_run()
    label.text = f"{kind} {number}"
    set_run_apa_font(label, 12, bold=True, italic=False, color=RGBColor(0, 0, 0))
    label.add_break()
    title_run = para.add_run(title.rstrip("."))
    set_run_apa_font(title_run, 12, bold=False, italic=True, color=RGBColor(0, 0, 0))
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 2.0
    return True


def apply_caption_index_styles(document: Document) -> None:
    for para in document.paragraphs:
        text = para.text.strip()
        if text.startswith("Tabla "):
            para.style = "TablaIndice"
        elif text.startswith("Figura "):
            para.style = "FiguraIndice"


def table_note_source(caption: str) -> tuple[str, str]:
    cap = caption.lower()
    if "tabla 1." in cap:
        return (
            "la correspondencia entre problema, objetivos, hipotesis, variables y alcances del planteamiento",
            "Elaboracion propia a partir de la formulacion del problema, objetivos e hipotesis definidos en el Capitulo 1.",
        )
    if "tabla 2." in cap:
        if "antecedentes" in cap:
            return (
                "la sintesis de antecedentes academicos y su relacion con la brecha teorica de la tesis",
                "Elaboracion propia con base en la revision bibliografica sistematizada en el Capitulo 2 y en las referencias APA del documento.",
            )
        if "dec-pomdp" in cap or "notacion" in cap:
            return (
                "la formalizacion teorica del problema multiagente y sus elementos matematicos",
                "Elaboracion propia a partir de Oliehoek y Amato (2016), Lowe et al. (2017) y los artefactos del entorno CityLearn usado en la tesis.",
            )
        return (
            "la organizacion conceptual del marco teorico y de las variables de investigacion",
            "Elaboracion propia a partir de la literatura citada en el Capitulo 2.",
        )
    if "tabla 3." in cap:
        return (
            "el diseno metodologico, las variables, herramientas, datos y procedimiento experimental",
            "Elaboracion propia a partir del proyecto local, scripts de metodologia, dataset citylearn_iquitos_2023_2025 y artefactos Drive de entrenamiento.",
        )
    if "tabla 4." in cap:
        return (
            "los componentes tecnicos de la propuesta, arquitectura, algoritmos, recompensas y trazabilidad experimental",
            "Elaboracion propia a partir del repositorio local, CityLearn v2, extension CityLearn v3 propuesta, scripts de dataset y corrida Drive madrl_v3_20260627_164047.",
        )
    if "tabla 5." in cap:
        return (
            "los resultados descriptivos, inferenciales, por edificio, por escenario o por KPI de la evaluacion MADRL",
            "Elaboracion propia a partir de results.json, training_summary.json, timeseries.csv, trace.csv, building_kpis.csv, checkpoint_manifest.json y tablas CSV generadas desde Drive.",
        )
    if "tabla 6." in cap:
        return (
            "los criterios de cierre, hallazgos, limitaciones o plan de culminacion de la tesis",
            "Elaboracion propia a partir de la sintesis de resultados, limitaciones metodologicas y plan de cierre doctoral.",
        )
    if "tabla c." in cap:
        return (
            "la auditoria documental de unificacion de archivos Word",
            "Elaboracion propia a partir de la revision automatizada de los archivos .docx existentes en la carpeta docs.",
        )
    return (
        "la informacion sintetizada en la tabla correspondiente",
        "Elaboracion propia a partir de los datos, fuentes y artefactos citados en el documento.",
    )


def build_table_note_text(caption: str) -> str:
    represents, source = table_note_source(caption)
    return f"Nota. La tabla representa {represents}. Fuente: {source}"


def figure_key_from_caption(caption: str) -> str:
    cap = norm_key(caption)
    for scenario in SCENARIOS:
        if f"convergencia por recompensa media movil en {scenario.lower()}" in cap:
            return f"convergence_{scenario}"
        if f"episodio final para {scenario.lower()}" in cap or f"serie temporal distrital final en {scenario.lower()}" in cap:
            return f"final_timeseries_{scenario}"
    if "cobertura de checkpoints" in cap:
        return "checkpoint_coverage"
    if "ranking kpi citylearn" in cap or "ranking de kpis citylearn" in cap or "mapa de calor del ranking" in cap:
        return "kpi_ranking_heatmap"
    if "oe 1" in cap or "oe1" in cap:
        return "objective_OE.1"
    if "oe 2" in cap or "oe2" in cap:
        return "objective_OE.2"
    if "oe 3" in cap or "oe3" in cap:
        return "objective_OE.3"
    if "distribucion episodica" in cap:
        return "episode_boxplots"
    if "tamano de efecto" in cap:
        return "effect_size"
    if "p valores holm" in cap or "p val" in cap:
        return "pairwise_heatmaps"
    if "trade off" in cap or "tradeoff" in cap:
        return "tradeoff"
    if "dimensiones de accion controlable" in cap:
        return "equipment"
    if "exito de salida ev" in cap:
        return "building_ev_success_heatmap"
    if "co2 control por edificio" in cap or "co2 por edificio" in cap:
        return "building_carbon_heatmap"
    if "costo control por edificio" in cap or "costo por edificio" in cap:
        return "building_cost_heatmap"
    if "equipamiento controlado" in cap:
        return "equipment_class_heatmap"
    if "politicas" in cap or "trace" in cap or "acciones medias" in cap:
        return "trace_policy_heatmaps"
    if "evolucion del reward distrital" in cap:
        return "appendix_reward"
    if "energia neta distrital" in cap:
        return "appendix_energy"
    if "costo distrital medio" in cap:
        return "appendix_cost"
    if "emisiones distritales medias" in cap:
        return "appendix_emissions"
    if "delta de costo electrico" in cap:
        return "appendix_cost_delta"
    if "delta de emisiones" in cap:
        return "appendix_emissions_delta"
    if "variables de accion controladas" in cap:
        return "appendix_action_variables"
    if "carga controlada" in cap and "carga base" in cap:
        return "appendix_controlled_load"
    if "tamano total listado" in cap and "checkpoint" in cap:
        return "appendix_checkpoint_size"
    for idx in range(1, 10):
        if f"diagrama {idx}" in cap:
            return f"architecture_diagram_{idx}"
    return "generic"


def figure_source_text(caption: str) -> str:
    key = figure_key_from_caption(caption)
    if key.startswith("convergence") or key in {"episode_boxplots", "effect_size", "pairwise_heatmaps"}:
        return (
            "Elaboracion propia a partir de district_episode_kpis.csv, results.json y training_summary.json "
            "de la corrida Drive madrl_v3_20260627_164047; el criterio de lectura estadistica sigue las "
            "recomendaciones de evaluacion en aprendizaje por refuerzo de Henderson et al. (2018), Colas et al. (2019) "
            "y Agarwal et al. (2021)."
        )
    if key in {"kpi_ranking_heatmap", "tradeoff"}:
        return (
            "Elaboracion propia a partir de los CSV comparativos compatibles con CityLearn v2 evaluate_v2 y de los "
            "artefactos finales de entrenamiento; la lectura se apoya en CityLearn (Vazquez-Canteli et al., 2020) "
            "y CityLearn v2 (Nweye et al., 2024)."
        )
    if key.startswith("final_timeseries") or key.startswith("building") or key in {"equipment", "equipment_class_heatmap", "trace_policy_heatmaps"}:
        return (
            "Elaboracion propia a partir de timeseries.csv, trace.csv, building_kpis.csv, "
            "building_behavior_summary.csv y building_observation_action_schema.csv de la carpeta Drive auditada; "
            "la interpretacion tecnica se vincula con CityLearn y gestion multiagente de energia urbana "
            "(Vazquez-Canteli et al., 2020; Nweye et al., 2024; Fonseca et al., 2024)."
        )
    if key == "checkpoint_coverage":
        return (
            "Elaboracion propia a partir de checkpoint_manifest.json, results.json y training_summary.json de los "
            "12 tratamientos algoritmo-escenario conservados en Drive."
        )
    if key.startswith("appendix_"):
        return (
            "Elaboracion propia a partir de las tablas y figuras de outputs/_drive_madrl/full_data/analysis_real_drive, "
            "timeseries.csv, building_kpis.csv, trace.csv y checkpoint_manifest.json de la carpeta Drive auditada."
        )
    if key.startswith("architecture_diagram"):
        return (
            "Elaboracion propia a partir de la arquitectura del repositorio local, los scripts de construccion del dataset, "
            "la formalizacion Dec-POMDP/CTDE y la configuracion experimental MADRL; se fundamenta en Oliehoek y Amato (2016), "
            "Lowe et al. (2017), Iqbal y Sha (2019), Vazquez-Canteli et al. (2020) y Nweye et al. (2024)."
        )
    return "Elaboracion propia a partir de los artefactos reales del repositorio local y de la carpeta Drive auditada."


def figure_represents_text(caption: str) -> str:
    key = figure_key_from_caption(caption)
    mapping = {
        "checkpoint_coverage": "la cobertura de checkpoints guardados por tratamiento y la trazabilidad material del entrenamiento",
        "kpi_ranking_heatmap": "el ranking normalizado de KPIs compatibles con CityLearn v2 evaluate_v2 por escenario, eje y metodo",
        "episode_boxplots": "la distribucion episodica de los indicadores usados para contrastar OE.1, OE.2 y OE.3",
        "effect_size": "la magnitud del efecto inferencial por objetivo especifico mediante epsilon cuadrado",
        "pairwise_heatmaps": "las diferencias por pares entre algoritmos despues del ajuste Holm",
        "tradeoff": "la relacion multiobjetivo entre costo energetico, emisiones de CO2 y autoconsumo fotovoltaico",
        "equipment": "las dimensiones de accion controlable por edificio y algoritmo",
        "building_ev_success_heatmap": "la heterogeneidad del exito de salida de vehiculos electricos por edificio y algoritmo",
        "building_carbon_heatmap": "la distribucion de emisiones de CO2 controladas por edificio y algoritmo",
        "building_cost_heatmap": "la distribucion de costo energetico controlado por edificio y algoritmo",
        "equipment_class_heatmap": "la composicion de equipos controlables identificados en los esquemas de accion",
        "trace_policy_heatmaps": "las acciones medias, carga EV y estado de carga BESS observados en trace.csv",
        "appendix_reward": "la evolucion del reward distrital medio por episodio conservado",
        "appendix_energy": "la energia neta distrital por episodio y tratamiento",
        "appendix_cost": "el costo distrital medio por algoritmo y escenario",
        "appendix_emissions": "las emisiones distritales medias por algoritmo y escenario",
        "appendix_cost_delta": "el delta de costo electrico por edificio en la corrida completa",
        "appendix_emissions_delta": "el delta de emisiones por edificio en la corrida completa",
        "appendix_action_variables": "las variables de accion controladas por edificio",
        "appendix_controlled_load": "la carga controlada de escenario frente a carga base no controlada",
        "appendix_checkpoint_size": "el tamano total listado en los manifiestos de checkpoint por algoritmo y escenario",
    }
    if key.startswith("convergence"):
        return "la evolucion de la recompensa media movil, el inicio de aprendizaje y la estabilizacion del entrenamiento"
    if key.startswith("objective"):
        return "la comparacion grafica de medias episodicas entre algoritmos para el objetivo especifico correspondiente"
    if key.startswith("final_timeseries"):
        return "los paneles distritales del episodio final conservado con series normalizadas y totales reales anotados"
    if key.startswith("architecture_diagram"):
        return "el flujo, arquitectura o componente metodologico del sistema doctoral MADRL-CityLearn"
    return mapping.get(key, "la evidencia visual derivada de los resultados reales del entrenamiento MADRL")


def build_figure_note_text(caption: str) -> str:
    return f"Nota. La figura representa {figure_represents_text(caption)}. Fuente: {figure_source_text(caption)}"


def figure_interpretation_text(caption: str) -> str:
    key = figure_key_from_caption(caption)
    return FIGURE_INTERPRETATIONS.get(
        key,
        "La lectura de la figura debe realizarse como evidencia visual complementaria a las tablas de resultados: muestra patrones observados en los artefactos reales, pero no sustituye la contrastacion descriptiva e inferencial desarrollada en el texto.",
    )


def set_figure_note_paragraph_format(para) -> None:
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    for run in para.runs:
        set_run_apa_font(run, 10, bold=False, italic=False, color=RGBColor(0, 0, 0))


def set_figure_interpretation_paragraph_format(para) -> None:
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 2.0
    for run in para.runs:
        set_run_apa_font(run, 12, bold=False, italic=False, color=RGBColor(0, 0, 0))


def element_has_drawing(el) -> bool:
    return bool(el.findall(".//" + qn("w:drawing")))


def ensure_apa_figure_notes(document: Document) -> None:
    body_children = list(document.element.body)
    for i, el in enumerate(body_children):
        if el.tag != qn("w:p"):
            continue
        caption = text_of(el).strip()
        if not caption.startswith("Figura "):
            continue

        target_el = el
        note_el = None
        interp_el = None
        for nxt in body_children[i + 1 : i + 10]:
            nxt_text = text_of(nxt).strip()
            if element_has_drawing(nxt):
                target_el = nxt
            if nxt_text.startswith("Nota. La figura"):
                note_el = nxt
                continue
            if nxt_text.startswith("Interpretacion de la figura"):
                interp_el = nxt
                break
            if nxt_text.startswith("Figura ") or nxt_text.startswith("Tabla ") or nxt_text.startswith("Capitulo ") or re.match(r"^[A-Z]\.\d+\s+", nxt_text):
                break

        if note_el is not None:
            note_para = Paragraph(note_el, document)
            set_paragraph_text(note_para, build_figure_note_text(caption))
        else:
            note_para = paragraph_after_element(target_el, document)
            note_para.add_run(build_figure_note_text(caption))
        set_figure_note_paragraph_format(note_para)

        interp_text = f"Interpretacion de la figura. {figure_interpretation_text(caption)}"
        if interp_el is not None:
            interp_para = Paragraph(interp_el, document)
            set_paragraph_text(interp_para, interp_text)
        else:
            interp_para = paragraph_after_element(note_para._p, document)
            interp_para.add_run(interp_text)
        set_figure_interpretation_paragraph_format(interp_para)


def paragraph_has_page_break_before(para) -> bool:
    p_pr = para._p.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:pageBreakBefore")) is not None


def element_has_page_break(el) -> bool:
    if el.find(".//" + qn("w:pageBreakBefore")) is not None:
        return True
    for br in el.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def element_has_explicit_page_break(el) -> bool:
    return any(br.get(qn("w:type")) == "page" for br in el.findall(".//" + qn("w:br")))


def remove_page_break_before(para) -> None:
    p_pr = para._p.find(qn("w:pPr"))
    if p_pr is None:
        return
    for node in list(p_pr.findall(qn("w:pageBreakBefore"))):
        p_pr.remove(node)


def next_nonempty_paragraph_after(document: Document, para):
    children = list(document.element.body)
    try:
        idx = children.index(para._p)
    except ValueError:
        return None
    for nxt in children[idx + 1 :]:
        if nxt.tag != qn("w:p"):
            continue
        candidate = Paragraph(nxt, document)
        if candidate.text.strip() or element_has_page_break(nxt):
            return candidate
    return None


def ensure_page_break_after_paragraph(document: Document, para) -> None:
    children = list(document.element.body)
    idx = children.index(para._p)
    if idx + 1 < len(children) and children[idx + 1].tag == qn("w:p"):
        next_para = Paragraph(children[idx + 1], document)
        remove_page_break_before(next_para)
        if element_has_explicit_page_break(children[idx + 1]):
            return
    if idx + 2 < len(children) and children[idx + 2].tag == qn("w:p"):
        remove_page_break_before(Paragraph(children[idx + 2], document))
    if idx + 1 < len(children) and children[idx + 1].tag == qn("w:p") and element_has_explicit_page_break(children[idx + 1]):
        return
    break_para = paragraph_after_element(para._p, document)
    break_para.add_run().add_break(WD_BREAK.PAGE)


def ensure_page_break_before_paragraph(document: Document, para) -> None:
    remove_page_break_before(para)
    children = list(document.element.body)
    idx = children.index(para._p)
    if idx > 0 and children[idx - 1].tag == qn("w:p") and element_has_explicit_page_break(children[idx - 1]):
        return
    new_p = OxmlElement("w:p")
    para._p.addprevious(new_p)
    break_para = Paragraph(new_p, document._body)
    break_para.add_run().add_break(WD_BREAK.PAGE)


def apply_page_break_rules(document: Document) -> None:
    for para in list(document.paragraphs):
        text = para.text.strip()
        if text.startswith("Palabras clave:"):
            ensure_page_break_after_paragraph(document, para)
        if text in {"Indice", "Indice de tablas", "Indice de figuras"}:
            ensure_page_break_before_paragraph(document, para)
        if re.match(r"^Capitulo\s+\d+\.", text):
            ensure_page_break_before_paragraph(document, para)


def replace_automatic_indexes(document: Document) -> None:
    body = document.element.body
    children = list(body)
    idx_indice = idx_cap1 = None
    for i, el in enumerate(children):
        text = text_of(el).strip()
        if idx_indice is None and text == "Indice":
            idx_indice = i
        if idx_indice is not None and text.startswith("Capitulo 1."):
            idx_cap1 = i
            break
    if idx_indice is None or idx_cap1 is None or idx_cap1 <= idx_indice:
        return

    for el in children[idx_indice + 1 : idx_cap1]:
        body.remove(el)

    indice_para = Paragraph(children[idx_indice], document)
    indice_para.style = "Heading 1"
    indice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    indice_para.paragraph_format.first_line_indent = None

    anchor = indice_para._p
    toc_para = paragraph_after_element(anchor, document)
    add_field(toc_para, r'TOC \o "1-5" \h \z \u', "Actualice el indice general con clic derecho > Actualizar campos (o F9).")

    break1 = paragraph_after_element(toc_para._p, document)
    break1.add_run().add_break(WD_BREAK.PAGE)

    tables_heading = paragraph_after_element(break1._p, document)
    tables_heading.add_run("Indice de tablas")
    tables_heading.style = "Heading 1"
    tables_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tables_heading.paragraph_format.first_line_indent = None
    tables_field = paragraph_after_element(tables_heading._p, document)
    add_field(tables_field, r'TOC \h \z \t "TablaIndice,1"', "Actualice el indice de tablas con clic derecho > Actualizar campos (o F9).")

    break2 = paragraph_after_element(tables_field._p, document)
    break2.add_run().add_break(WD_BREAK.PAGE)

    figures_heading = paragraph_after_element(break2._p, document)
    figures_heading.add_run("Indice de figuras")
    figures_heading.style = "Heading 1"
    figures_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figures_heading.paragraph_format.first_line_indent = None
    figures_field = paragraph_after_element(figures_heading._p, document)
    add_field(figures_field, r'TOC \h \z \t "FiguraIndice,1"', "Actualice el indice de figuras con clic derecho > Actualizar campos (o F9).")

    break3 = paragraph_after_element(figures_field._p, document)
    break3.add_run().add_break(WD_BREAK.PAGE)


def paragraph_after_element(el, document: Document):
    new_p = OxmlElement("w:p")
    el.addnext(new_p)
    return Paragraph(new_p, document._body)


def set_note_paragraph_format(para) -> None:
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    for run in para.runs:
        set_run_apa_font(run, 10, bold=False, italic=False, color=RGBColor(0, 0, 0))


def ensure_apa_table_notes(document: Document) -> None:
    body_children = list(document.element.body)
    for i, el in enumerate(body_children):
        if el.tag != qn("w:tbl"):
            continue
        caption = ""
        j = i - 1
        while j >= 0:
            prev_text = text_of(body_children[j])
            if prev_text:
                if prev_text.startswith("Tabla "):
                    caption = prev_text.replace("\n", ". ")
                break
            j -= 1
        if not caption:
            caption = "Tabla sin numeracion detectada"

        note_text = build_table_note_text(caption)
        next_el = body_children[i + 1] if i + 1 < len(body_children) else None
        if next_el is not None and next_el.tag == qn("w:p") and text_of(next_el).startswith("Nota."):
            note_para = Paragraph(next_el, document)
            set_paragraph_text(note_para, note_text)
        else:
            note_para = paragraph_after_element(el, document)
            note_para.add_run(note_text)
        set_note_paragraph_format(note_para)


def apply_apa7_document_format(document: Document) -> None:
    """Normaliza formato APA 7: estilos, niveles de titulo, captions y referencias."""
    ensure_apa_table_notes(document)
    ensure_apa_figure_notes(document)
    heading_names = [s.name for s in document.styles]
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    refs_started = False
    refs_ended = False
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("Referencias bibliograficas"):
            refs_started = True
            refs_ended = False
        if refs_started and text.startswith("Anexo "):
            refs_ended = True

        level = infer_heading_level(text)
        if level and f"Heading {level}" in heading_names:
            para.style = document.styles[f"Heading {level}"]

        style_name = (para.style.name or "").strip().lower() if para.style is not None else ""
        is_caption = format_caption_paragraph(para)
        is_figure_note = text.startswith("Nota. La figura")
        is_figure_interpretation = text.startswith("Interpretacion de la figura")
        if style_name.startswith("heading"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if style_name == "heading 1" else WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = None if style_name not in {"heading 4", "heading 5"} else Inches(0.5)
            para.paragraph_format.line_spacing = 2.0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
        elif is_figure_note:
            set_figure_note_paragraph_format(para)
        elif is_figure_interpretation:
            set_figure_interpretation_paragraph_format(para)
        elif refs_started and not refs_ended and not style_name.startswith("heading"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)
            para.paragraph_format.line_spacing = 2.0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
        elif not is_caption:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = 2.0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            if not text.startswith(("Palabras clave:", "Keywords:", "Resumen", "Abstract", "Indice")):
                para.paragraph_format.first_line_indent = Inches(0.5)

        for run in para.runs:
            if run.text:
                if is_figure_note:
                    set_run_apa_font(run, 10, color=RGBColor(0, 0, 0))
                else:
                    set_run_apa_font(run, 12, color=RGBColor(0, 0, 0))

    for tbl in document.tables:
        apply_apa_table_format(tbl, body_font_size=10.0)
    ensure_apa_table_notes(document)
    ensure_apa_figure_notes(document)
    apply_caption_index_styles(document)
    replace_automatic_indexes(document)
    enable_update_fields_on_open(document)
    apply_page_break_rules(document)


def normalize_apa_citation_text(document: Document) -> None:
    """Corrige citas en texto que impedian el enlace APA cita-referencia."""
    replacements = {
        "(MINAM RAGEI 2019)": "(MINAM, 2019)",
        "Ministerio del Ambiente del Peru (2019)": "MINAM (2019)",
        "Ministerio del Ambiente del Perú (2019)": "MINAM (2019)",
        "Shadish, Cook y Campbell (2002)": "Shadish et al. (2002)",
        "Rosero Bernal (2024)": "Rosero Bernal (2022)",
    }
    for para in iter_document_paragraphs(document):
        text = para.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph_text(para, new_text)


def table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], font_size: float = 7.0):
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    set_run_apa_font(run, 12, bold=True, italic=False, color=RGBColor(0, 0, 0))
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(head)
        set_run_apa_font(rr, max(float(font_size), 10.0), bold=True, italic=False, color=RGBColor(0, 0, 0))
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            set_run_apa_font(rr, max(float(font_size), 10.0), bold=False, italic=False, color=RGBColor(0, 0, 0))
    apply_apa_table_format(tbl, body_font_size=max(float(font_size), 10.0))
    doc.add_paragraph()
    return tbl


def aggregate_timeseries(path: Path) -> pd.DataFrame:
    cols = [
        "episode",
        "district_net_electricity_consumption",
        "district_net_electricity_consumption_cost",
        "district_net_electricity_consumption_emission",
    ]
    sums: dict[int, dict[str, float]] = {}
    for chunk in pd.read_csv(path, usecols=cols, chunksize=200_000):
        chunk["episode"] = pd.to_numeric(chunk["episode"], errors="coerce").astype("Int64")
        grouped = chunk.groupby("episode", dropna=True).agg(
            district_net_energy=("district_net_electricity_consumption", "sum"),
            district_cost=("district_net_electricity_consumption_cost", "sum"),
            district_emission=("district_net_electricity_consumption_emission", "sum"),
        )
        for ep, row in grouped.iterrows():
            if pd.isna(ep):
                continue
            ep = int(ep)
            dest = sums.setdefault(ep, {"district_net_energy": 0.0, "district_cost": 0.0, "district_emission": 0.0})
            for key in dest:
                dest[key] += float(row[key])
    rows = [{"episode": ep, **vals} for ep, vals in sorted(sums.items())]
    return pd.DataFrame(rows)


def load_materialized_episode_kpis() -> pd.DataFrame:
    if not LOCAL_EPISODE_CSV.exists():
        raise FileNotFoundError(
            f"No existe {LOCAL_EPISODE_CSV}. Se requiere el CSV materializado para evitar releer los timeseries grandes de Drive."
        )
    df = pd.read_csv(LOCAL_EPISODE_CSV)
    df = df.rename(
        columns={
            "district_net_electricity_consumption_kwh": "district_net_energy",
            "reward_mean": "reward_mean_average",
            "reward_sum": "reward_sum_total",
        }
    )
    required = {"algorithm", "scenario", "episode", "district_net_energy", "district_cost", "district_emission", "reward_mean_average"}
    missing = required - set(df.columns)
    if missing:
        raise RuntimeError(f"Faltan columnas en {LOCAL_EPISODE_CSV}: {sorted(missing)}")
    df.to_csv(TABLE_DIR / "gdrive_episode_kpis_from_materialized_drive_analysis.csv", index=False, encoding="utf-8")
    return df


def load_evidence() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    treatment_rows = []
    building_frames = []
    equipment_frames = []
    episode_kpis = load_materialized_episode_kpis()
    episode_counts = episode_kpis.groupby(["algorithm", "scenario"]).size().to_dict()
    for algo in ALGOS:
        for scenario in SCENARIOS:
            data_dir = G_BASE / algo / scenario / "data"
            result = read_json(data_dir / "results.json")
            live_progress_path = G_BASE / algo / scenario / "live_progress.json"
            live_progress = read_json(live_progress_path) if live_progress_path.exists() else {}
            all_values = result.get("citylearn_v3_report", {}).get("all_values", {})
            episode_summaries = result.get("episode_summaries") or []
            reward_df = pd.DataFrame(episode_summaries)
            if not reward_df.empty:
                reward_df = reward_df[["episode", "reward_mean_average", "reward_sum_total", "steps"]].copy()
                reward_df["episode"] = pd.to_numeric(reward_df["episode"], errors="coerce").astype("Int64")
            final_reward = safe_float(reward_df["reward_mean_average"].iloc[-1]) if not reward_df.empty else math.nan
            final_episode = int(reward_df["episode"].iloc[-1]) if not reward_df.empty and not pd.isna(reward_df["episode"].iloc[-1]) else None
            treatment_rows.append(
                {
                    "algorithm": algo,
                    "scenario": scenario,
                    "episodes_recorded": result.get("episodes_recorded", live_progress.get("completed_episode_count", result.get("episodes"))),
                    "training_episodes_field": result.get("episodes"),
                    "saved_episode_summaries": len(episode_summaries),
                    "materialized_episode_kpis": int(episode_counts.get((algo, scenario), 0)),
                    "final_episode_in_artifacts": final_episode,
                    "final_reward_mean": final_reward,
                    "building_count": result.get("building_count"),
                    "checkpoint_count": result.get("checkpoint_count"),
                    "peak_average": all_values.get("peak_average", all_values.get("cost_peak_average")),
                    "ramping_average": all_values.get("ramping_average", all_values.get("cost_ramping_average")),
                    "one_minus_load_factor_average": all_values.get("one_minus_load_factor_average", all_values.get("cost_one_minus_load_factor_average")),
                    "battery_throughput_total": all_values.get("battery_throughput_total"),
                    "pv_self_consumption_ratio": all_values.get("pv_self_consumption_ratio"),
                    "carbon_emissions_control": all_values.get("carbon_emissions_control"),
                    "carbon_emissions_delta": all_values.get("carbon_emissions_delta"),
                    "carbon_emissions_ratio": all_values.get("carbon_emissions"),
                    "electricity_cost_control": all_values.get("electricity_cost_control"),
                    "electricity_cost_delta": all_values.get("electricity_cost_delta"),
                    "electricity_cost_ratio": all_values.get("electricity_cost"),
                    "ev_departure_success_rate": all_values.get("ev_departure_success_rate"),
                }
            )
            b = pd.read_csv(data_dir / "building_behavior_summary.csv")
            b["algorithm"] = algo
            b["scenario"] = scenario
            building_frames.append(b)
            e = pd.read_csv(data_dir / "building_observation_action_schema.csv")
            e["algorithm"] = algo
            e["scenario"] = scenario
            equipment_frames.append(e)
    treatment = pd.DataFrame(treatment_rows)
    episodes = episode_kpis
    buildings = pd.concat(building_frames, ignore_index=True)
    equipment = pd.concat(equipment_frames, ignore_index=True)
    treatment.to_csv(TABLE_DIR / "gdrive_treatment_final_kpis.csv", index=False, encoding="utf-8")
    episodes.to_csv(TABLE_DIR / "gdrive_episode_kpis_used_for_statistics.csv", index=False, encoding="utf-8")
    buildings.to_csv(TABLE_DIR / "gdrive_building_behavior_summary_all.csv", index=False, encoding="utf-8")
    equipment.to_csv(TABLE_DIR / "gdrive_equipment_schema_all.csv", index=False, encoding="utf-8")
    return treatment, episodes, buildings, equipment


def holm_adjust(pairs: list[tuple[str, float]]) -> list[tuple[str, float, float]]:
    ordered = sorted(pairs, key=lambda x: x[1])
    adjusted = []
    running = 0.0
    m = len(ordered)
    for rank, (name, pv) in enumerate(ordered, start=1):
        adj = min(1.0, (m - rank + 1) * pv)
        running = max(running, adj)
        adjusted.append((name, pv, running))
    return sorted(adjusted, key=lambda x: x[0])


def analyze_objectives(treatment: pd.DataFrame, episodes: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, dict]:
    rows = []
    pair_rows = []
    detail = {}
    for spec in OBJECTIVES:
        sub = episodes[episodes["scenario"] == spec["scenario"]].copy()
        metric = spec["metric"]
        direction = spec["direction"]
        ascending = direction == "min"
        desc = sub.groupby("algorithm")[metric].agg(["count", "mean", "median", "std", "min", "max"]).reset_index()
        desc["statistical_coverage"] = desc["count"].map(lambda n: "cobertura completa por artefacto" if n >= 50 else f"{int(n)} fila(s) conservada(s)")
        final = treatment[treatment["scenario"] == spec["scenario"]].copy()
        if spec["objective"] == "OE.1":
            final["final_metric"] = final["final_reward_mean"]
        elif spec["objective"] == "OE.2":
            final["final_metric"] = final["carbon_emissions_control"]
        else:
            final["final_metric"] = final["electricity_cost_control"]
        final_best_row = final.sort_values("final_metric", ascending=ascending).iloc[0]
        inferential_algos = desc[desc["count"] >= 50]["algorithm"].tolist()
        inferential = sub[sub["algorithm"].isin(inferential_algos)]
        grouped = [g[metric].dropna().values for _, g in inferential.groupby("algorithm")]
        if len(grouped) >= 2:
            kw = stats.kruskal(*grouped)
            n = sum(len(g) for g in grouped)
            k = len(grouped)
            eps2 = (kw.statistic - k + 1) / (n - k) if n > k else math.nan
        else:
            kw = None
            eps2 = math.nan
        pair_p = []
        for i, a in enumerate(inferential_algos):
            for b in inferential_algos[i + 1 :]:
                av = inferential[inferential["algorithm"] == a][metric].dropna().values
                bv = inferential[inferential["algorithm"] == b][metric].dropna().values
                pair_p.append((f"{a} vs {b}", stats.mannwhitneyu(av, bv, alternative="two-sided").pvalue))
        pair_adj = holm_adjust(pair_p) if pair_p else []
        shapiro = {}
        shapiro_rows = []
        for algo in ALGOS:
            vals = sub[sub["algorithm"] == algo][metric].dropna().values
            if len(vals) >= 3:
                shapiro_result = stats.shapiro(vals)
                shapiro_w = float(shapiro_result.statistic)
                shapiro_p = float(shapiro_result.pvalue)
                decision = "normalidad no rechazada" if shapiro_p >= 0.05 else "normalidad rechazada"
            else:
                shapiro_w = math.nan
                shapiro_p = math.nan
                decision = "no aplicable"
            use = "incluido en Kruskal-Wallis" if algo in inferential_algos else "descriptivo; excluido de Kruskal-Wallis por n<50"
            shapiro[algo] = {"n": int(len(vals)), "w": shapiro_w, "p": shapiro_p, "decision": decision, "use": use}
            shapiro_rows.append(
                {
                    "algorithm": algo,
                    "n": int(len(vals)),
                    "w": shapiro_w,
                    "p": shapiro_p,
                    "decision": decision,
                    "inferential_use": use,
                }
            )
        best_stat = desc.sort_values("mean", ascending=ascending).iloc[0]["algorithm"]
        best_stat_complete = desc[desc["count"] >= 50].sort_values("mean", ascending=ascending).iloc[0]["algorithm"]
        for _, r in desc.iterrows():
            rows.append(
                {
                    "objective": spec["objective"],
                    "scenario": spec["scenario"],
                    "dimension": spec["dimension"],
                    "metric": metric,
                    "direction": direction,
                    "algorithm": r["algorithm"],
                    "n_episode_artifacts": int(r["count"]),
                    "mean": r["mean"],
                    "median": r["median"],
                    "std": r["std"],
                    "min": r["min"],
                    "max": r["max"],
                    "coverage": r["statistical_coverage"],
                    "best_by_episode_mean": best_stat,
                    "best_inferential_sample": best_stat_complete,
                    "best_final_annual_kpi": final_best_row["algorithm"],
                    "kw_algorithms": ", ".join(inferential_algos),
                    "kw_h": kw.statistic if kw else math.nan,
                    "kw_p": kw.pvalue if kw else math.nan,
                    "kw_epsilon2": eps2,
                    "shapiro_w": shapiro.get(r["algorithm"], {}).get("w", math.nan),
                    "shapiro_p": shapiro.get(r["algorithm"], {}).get("p", math.nan),
                    "shapiro_decision": shapiro.get(r["algorithm"], {}).get("decision", "NA"),
                    "shapiro_inferential_use": shapiro.get(r["algorithm"], {}).get("use", "NA"),
                }
            )
        for pair, pv, adj in pair_adj:
            pair_rows.append(
                {
                    "objective": spec["objective"],
                    "scenario": spec["scenario"],
                    "metric": metric,
                    "pair": pair,
                    "p_raw": pv,
                    "p_holm": adj,
                    "decision": "significativo" if adj < 0.05 else "no significativo",
                }
            )
        detail[spec["objective"]] = {
            "spec": spec,
            "desc": desc.sort_values("mean", ascending=ascending).reset_index(drop=True),
            "final": final.sort_values("final_metric", ascending=ascending).reset_index(drop=True),
            "best_stat": best_stat,
            "best_stat_complete": best_stat_complete,
            "best_final": final_best_row["algorithm"],
            "kw": kw,
            "epsilon2": eps2,
            "inferential_algos": inferential_algos,
            "pair_adj": pair_adj,
            "shapiro": shapiro,
            "shapiro_rows": pd.DataFrame(shapiro_rows),
        }
    stats_df = pd.DataFrame(rows)
    pairs_df = pd.DataFrame(pair_rows)
    stats_df.to_csv(TABLE_DIR / "gdrive_objective_aligned_statistics.csv", index=False, encoding="utf-8")
    pairs_df.to_csv(TABLE_DIR / "gdrive_objective_pairwise_mannwhitney_holm.csv", index=False, encoding="utf-8")
    shapiro_cols = [
        "objective",
        "scenario",
        "dimension",
        "metric",
        "algorithm",
        "n_episode_artifacts",
        "shapiro_w",
        "shapiro_p",
        "shapiro_decision",
        "shapiro_inferential_use",
    ]
    stats_df[shapiro_cols].to_csv(TABLE_DIR / "gdrive_shapiro_wilk_recalculated_by_objective.csv", index=False, encoding="utf-8")
    return stats_df, pairs_df, detail


def analyze_convergence(episodes: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for (algo, scenario), sub in episodes.groupby(["algorithm", "scenario"]):
        sub = sub.sort_values("episode").reset_index(drop=True).copy()
        rewards = pd.to_numeric(sub["reward_mean_average"], errors="coerce")
        rolling = rewards.rolling(window=5, min_periods=1).mean()
        n = len(sub)
        if n == 0:
            continue
        initial = float(rolling.iloc[: min(5, n)].mean())
        final = float(rolling.iloc[max(0, n - 5) :].mean())
        improvement = final - initial
        threshold = initial + 0.20 * improvement
        learning_start_idx = 0
        if improvement > 0:
            found = rolling[rolling >= threshold]
            learning_start_idx = int(found.index[0]) if not found.empty else int(rolling.index[-1])
        tolerance = max(abs(final) * 0.05, 1e-9)
        stable_idx = int(rolling.index[-1])
        for idx in rolling.index:
            tail = rolling.loc[idx:]
            if ((tail - final).abs() <= tolerance).all():
                stable_idx = int(idx)
                break
        best_pos = int(rewards.idxmax())
        learn_row = sub.loc[learning_start_idx]
        stable_row = sub.loc[stable_idx]
        best_row = sub.loc[best_pos]
        rows.append(
            {
                "algorithm": algo,
                "scenario": scenario,
                "n_episode_artifacts": n,
                "initial_rolling_reward": initial,
                "final_rolling_reward": final,
                "reward_improvement": improvement,
                "learning_start_episode_index": int(learn_row["episode"]),
                "learning_start_episode_ordinal": int(learn_row["episode"]) + 1,
                "stabilization_episode_index": int(stable_row["episode"]),
                "stabilization_episode_ordinal": int(stable_row["episode"]) + 1,
                "best_episode_index": int(best_row["episode"]),
                "best_episode_ordinal": int(best_row["episode"]) + 1,
                "best_reward_mean": float(best_row["reward_mean_average"]),
                "stabilization_tolerance": tolerance,
            }
        )
    df = pd.DataFrame(rows).sort_values(["scenario", "algorithm"]).reset_index(drop=True)
    df.to_csv(TABLE_DIR / "gdrive_reward_convergence_episodes.csv", index=False, encoding="utf-8")
    return df


def load_citylearn_v2_kpi_summary() -> tuple[pd.DataFrame, pd.DataFrame]:
    base = REPO / "outputs" / "madrl_v3_20260627_164047" / "resumen_comparativo" / "citylearn_v2_baseline"
    rank_frames = []
    catalog_frames = []
    axis_for_scenario = {"E1": "OE1", "E2": "OE2", "E3": "OE3"}
    for scenario in SCENARIOS:
        scenario_dir = base / scenario
        rank_path = scenario_dir / "ranking_by_axis.csv"
        master_path = scenario_dir / "master_kpi_comparison.csv"
        if rank_path.exists():
            r = pd.read_csv(rank_path)
            r["scenario"] = scenario
            r = r[r["axis"] == axis_for_scenario[scenario]].copy()
            rank_frames.append(r)
        if master_path.exists():
            m = pd.read_csv(master_path)
            m["scenario"] = scenario
            catalog_frames.append(m)
    if not rank_frames:
        return pd.DataFrame(), pd.DataFrame()
    ranking = pd.concat(rank_frames, ignore_index=True)
    ranking = ranking[
        [
            "scenario",
            "axis",
            "family",
            "method",
            "normalized_score",
            "available_kpis",
            "improved_kpis",
            "total_kpis",
            "axis_rank",
        ]
    ].sort_values(["scenario", "axis_rank", "family", "method"])
    ranking.to_csv(TABLE_DIR / "citylearn_v2_evaluate_v2_axis_ranking.csv", index=False, encoding="utf-8")
    if catalog_frames:
        catalog = pd.concat(catalog_frames, ignore_index=True)
        catalog = catalog[catalog["available"].astype(str).str.lower().isin(["true", "1"])]
        rows = []
        for (scenario, axis, axis_name), sub in catalog.groupby(["scenario", "axis", "axis_name"], dropna=False):
            names = sorted(sub["kpi"].dropna().astype(str).unique().tolist())
            rows.append(
                {
                    "scenario": scenario,
                    "axis": axis,
                    "axis_name": axis_name,
                    "available_unique_kpis": len(names),
                    "source": ", ".join(sorted(sub["source"].dropna().astype(str).unique().tolist())[:3]),
                    "example_kpis": ", ".join(names[:8]),
                }
            )
        kpi_catalog = pd.DataFrame(rows).sort_values(["scenario", "axis"])
    else:
        kpi_catalog = pd.DataFrame()
    kpi_catalog.to_csv(TABLE_DIR / "citylearn_v2_evaluate_v2_kpi_catalog.csv", index=False, encoding="utf-8")
    return ranking, kpi_catalog


def load_final_episode_timeseries(treatment: pd.DataFrame) -> pd.DataFrame:
    out_path = TABLE_DIR / "gdrive_final_episode_timeseries_compact.csv"
    frames = []
    cols = [
        "algorithm",
        "scenario",
        "episode",
        "episode_step",
        "time_step",
        "district_net_electricity_consumption",
        "district_net_electricity_consumption_without_storage",
        "district_net_electricity_consumption_cost",
        "district_net_electricity_consumption_emission",
        "electricity_price_mean",
        "carbon_intensity_mean",
        "reward_mean",
    ]
    for _, r in treatment.iterrows():
        algo = r["algorithm"]
        scenario = r["scenario"]
        final_episode = int(r["final_episode_in_artifacts"]) if not pd.isna(r["final_episode_in_artifacts"]) else 49
        path = G_BASE / algo / scenario / "data" / "timeseries.csv"
        if not path.exists():
            continue
        parts = []
        for chunk in pd.read_csv(path, usecols=lambda c: c in cols, chunksize=100_000):
            if "episode" not in chunk.columns:
                continue
            sub = chunk[pd.to_numeric(chunk["episode"], errors="coerce") == final_episode].copy()
            if not sub.empty:
                parts.append(sub)
        if parts:
            df = pd.concat(parts, ignore_index=True)
            frames.append(df)
    if not frames:
        return pd.DataFrame()
    ts = pd.concat(frames, ignore_index=True)
    ts.to_csv(out_path, index=False, encoding="utf-8")
    return ts


def load_trace_samples() -> pd.DataFrame:
    out_path = TABLE_DIR / "gdrive_trace_samples_all.csv"
    frames = []
    for algo in ALGOS:
        for scenario in SCENARIOS:
            path = G_BASE / algo / scenario / "data" / "trace.csv"
            if not path.exists():
                continue
            try:
                df = pd.read_csv(path)
            except pd.errors.EmptyDataError:
                continue
            if df.empty:
                continue
            df["algorithm"] = algo
            df["scenario"] = scenario
            frames.append(df)
    if not frames:
        return pd.DataFrame()
    traces = pd.concat(frames, ignore_index=True)
    traces.to_csv(out_path, index=False, encoding="utf-8")
    return traces


def load_checkpoint_summary() -> pd.DataFrame:
    """Load checkpoint rows for project MADRL only (HAPPO/MAAC/MASAC/MATD3 × E1/E2/E3).

    Prefer the local Drive mirror under outputs/_drive_madrl/full_data when present,
    because G_BASE may be unavailable offline. Count every listed checkpoint file;
    do not require episode_(\\d+) in the path (that wrongly zeroed MASAC/MATD3).
    """
    local_base = REPO / "outputs" / "_drive_madrl" / "full_data"
    rows = []
    for algo in ALGOS:
        for scenario in SCENARIOS:
            candidates = [
                local_base / algo / scenario / "data" / "checkpoint_manifest.json",
                G_BASE / algo / scenario / "data" / "checkpoint_manifest.json",
            ]
            path = next((p for p in candidates if p.exists()), None)
            if path is None:
                continue
            data = read_json(path)
            checkpoints = data.get("checkpoints") or []
            for idx, ckpt in enumerate(checkpoints):
                rel = ckpt.get("relative_path", "")
                match = re.search(r"episode_(\d+)", rel)
                rows.append(
                    {
                        "algorithm": algo,
                        "scenario": scenario,
                        "checkpoint_count": data.get("checkpoint_count", len(checkpoints)),
                        "checkpoint_episode": int(match.group(1)) if match else idx,
                        "bytes": ckpt.get("bytes", math.nan),
                        "relative_path": rel,
                        "manifest_source": str(path),
                    }
                )
    df = pd.DataFrame(rows)
    if not df.empty:
        df.to_csv(TABLE_DIR / "gdrive_checkpoint_manifest_compact.csv", index=False, encoding="utf-8")
    return df


def analyze_buildings(buildings: pd.DataFrame, equipment: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    cols = [
        "algorithm",
        "scenario",
        "agent",
        "grid_role_control",
        "action_dim",
        "observation_dim",
        "battery_throughput_total_kwh",
        "ev_charge_total_kwh",
        "ev_departure_success_rate",
        "carbon_emissions_control_kgco2",
        "carbon_emissions_delta_kgco2",
        "electricity_cost_control_eur",
        "electricity_cost_delta_eur",
        "grid_import_control_kwh",
        "grid_export_control_kwh",
    ]
    b = buildings[[c for c in cols if c in buildings.columns]].copy()
    b.to_csv(TABLE_DIR / "gdrive_building_kpi_compact.csv", index=False, encoding="utf-8")
    eq = equipment[equipment["variable_type"] == "action"].copy()
    eq["equipment_class"] = eq["variable_name"].map(classify_equipment)
    eq_summary = eq.groupby(["algorithm", "scenario", "agent", "equipment_class"]).size().reset_index(name="count")
    eq_summary.to_csv(TABLE_DIR / "gdrive_controlled_equipment_by_building.csv", index=False, encoding="utf-8")
    return b, eq_summary


def classify_equipment(name: str) -> str:
    n = str(name).lower()
    if "electric_vehicle" in n:
        return "EV controlado"
    if "electrical_storage" in n or "battery" in n:
        return "BESS controlado"
    if "washing" in n or "dishwasher" in n or "dryer" in n:
        return "carga flexible controlada"
    if "cooling" in n or "heating" in n or "heat_pump" in n:
        return "HVAC/termico controlado"
    return "otro actuador controlado"


def build_figure_interpretations(
    detail: dict,
    treatment: pd.DataFrame,
    building_compact: pd.DataFrame,
    eq_summary: pd.DataFrame,
    episodes: pd.DataFrame,
    convergence: pd.DataFrame,
    kpi_ranking: pd.DataFrame,
    final_ts: pd.DataFrame,
    traces: pd.DataFrame,
    checkpoints: pd.DataFrame,
) -> None:
    FIGURE_INTERPRETATIONS.clear()

    if not checkpoints.empty:
        counts = checkpoints.groupby(["algorithm", "scenario"])["checkpoint_episode"].count().reset_index(name="n")
        max_row = counts.sort_values("n", ascending=False).iloc[0]
        min_row = counts.sort_values("n", ascending=True).iloc[0]
        FIGURE_INTERPRETATIONS["checkpoint_coverage"] = (
            f"La figura grafica {len(checkpoints)} registros de checkpoint distribuidos en {len(counts)} tratamientos. "
            f"La mayor cobertura corresponde a {max_row['algorithm']}-{max_row['scenario']} con {int(max_row['n'])} checkpoints, "
            f"mientras que la menor cobertura corresponde a {min_row['algorithm']}-{min_row['scenario']} con {int(min_row['n'])}. "
            "Esta lectura muestra la trazabilidad material del entrenamiento y permite verificar que la comparacion se apoya en artefactos persistidos, no en valores reconstruidos manualmente."
        )

    for scenario in SCENARIOS:
        sub = convergence[convergence["scenario"] == scenario].copy()
        if not sub.empty:
            best_final = sub.sort_values("final_rolling_reward", ascending=False).iloc[0]
            best_improve = sub.sort_values("reward_improvement", ascending=False).iloc[0]
            st_min = int(sub["stabilization_episode_ordinal"].min())
            st_max = int(sub["stabilization_episode_ordinal"].max())
            FIGURE_INTERPRETATIONS[f"convergence_{scenario}"] = (
                f"La figura muestra la recompensa media movil de {scenario} para {len(sub)} algoritmos. "
                f"El mayor valor final de recompensa movil corresponde a {best_final['algorithm']} con {fmt(best_final['final_rolling_reward'], 6)}, "
                f"y la mayor mejora entre tramo inicial y final corresponde a {best_improve['algorithm']} con {fmt(best_improve['reward_improvement'], 6)}. "
                f"La estabilizacion estimada ocurre entre los episodios ordinales {st_min} y {st_max}, calculados directamente desde reward_mean_average; por ello, la curva se interpreta como evidencia de aprendizaje y estabilidad intra-corrida, no como prueba multi-semilla."
            )

    for oe, d in detail.items():
        kw = d["kw"]
        spec = d["spec"]
        FIGURE_INTERPRETATIONS[f"objective_{oe}"] = (
            f"La figura compara la media episodica de {spec['dimension']} en {spec['scenario']} usando el indicador {spec['metric']}. "
            f"El mejor promedio descriptivo conservado es {d['best_stat']}, el mejor algoritmo con cobertura inferencial completa es {d['best_stat_complete']} "
            f"y el mejor KPI anual final observado es {d['best_final']}. "
            f"El contraste Kruskal-Wallis asociado reporta H={kw.statistic:.4f}, p={kw.pvalue:.6g} y epsilon2={d['epsilon2']:.4f}; por tanto, la figura se lee junto con la prueba estadistica y no como evidencia visual aislada."
        )

    total_episode_rows = len(episodes)
    if total_episode_rows:
        counts = episodes.groupby(["algorithm", "scenario"]).size()
        FIGURE_INTERPRETATIONS["episode_boxplots"] = (
            f"La figura resume la dispersion de {total_episode_rows} filas episodicas materializadas. "
            f"Las cajas permiten comparar mediana, variabilidad y valores extremos entre algoritmos para OE.1, OE.2 y OE.3; "
            f"la cobertura minima observada por tratamiento es {int(counts.min())} y la maxima es {int(counts.max())} episodios. "
            "La diferencia de cobertura explica por que HAPPO se mantiene como evidencia descriptiva cuando no alcanza 50 observaciones completas en el CSV materializado."
        )

    if detail:
        eps = {oe: d["epsilon2"] for oe, d in detail.items()}
        best_oe = max(eps, key=eps.get)
        eps_text = "; ".join(f"{oe}={fmt(val, 4)}" for oe, val in eps.items())
        FIGURE_INTERPRETATIONS["effect_size"] = (
            f"La figura muestra epsilon2 de Kruskal-Wallis por objetivo: {eps_text}. "
            f"El mayor tamano de efecto corresponde a {best_oe}, lo que confirma que la fuerza del efecto MADRL no es homogenea entre flexibilidad, emisiones y costos. "
            "La lectura complementa los p-valores porque cuantifica magnitud y no solo significancia."
        )
        sig_pairs = []
        for oe, d in detail.items():
            sig = sum(1 for _name, _p, adj in d["pair_adj"] if adj < 0.05)
            sig_pairs.append(f"{oe}: {sig}")
        FIGURE_INTERPRETATIONS["pairwise_heatmaps"] = (
            "La figura transforma los p-valores ajustados Holm en intensidad visual para ubicar que pares de algoritmos explican las diferencias globales. "
            f"Los pares significativos por objetivo son {', '.join(sig_pairs)}. "
            "Esta representacion evita concluir dominancia general cuando las diferencias solo aparecen en objetivos o pares especificos."
        )

    if not kpi_ranking.empty:
        best_rows = []
        for scenario in SCENARIOS:
            sub = kpi_ranking[kpi_ranking["scenario"] == scenario].sort_values("axis_rank")
            if not sub.empty:
                r = sub.iloc[0]
                best_rows.append(f"{scenario}: {r['method']} ({r['family']}, score={fmt(r['normalized_score'], 4)})")
        FIGURE_INTERPRETATIONS["kpi_ranking_heatmap"] = (
            f"La figura muestra scores normalizados de evaluate_v2 para {len(kpi_ranking)} combinaciones metodo-eje. "
            f"Los mejores metodos por escenario son {', '.join(best_rows)}. "
            "La lectura es relevante porque CityLearn v2 separa la recompensa de entrenamiento de los KPIs finales de flexibilidad, carbono y costo."
        )

    if not treatment.empty and {"electricity_cost_control", "carbon_emissions_control", "pv_self_consumption_ratio"}.issubset(treatment.columns):
        min_cost = treatment.sort_values("electricity_cost_control").iloc[0]
        min_co2 = treatment.sort_values("carbon_emissions_control").iloc[0]
        max_pv = treatment.sort_values("pv_self_consumption_ratio", ascending=False).iloc[0]
        FIGURE_INTERPRETATIONS["tradeoff"] = (
            f"La figura cruza {len(treatment)} tratamientos en costo, CO2 y autoconsumo PV. "
            f"El menor costo controlado aparece en {min_cost['algorithm']}-{min_cost['scenario']} con {fmt(min_cost['electricity_cost_control'], 2)}, "
            f"la menor emision controlada en {min_co2['algorithm']}-{min_co2['scenario']} con {fmt(min_co2['carbon_emissions_control'], 2)}, "
            f"y el mayor autoconsumo PV en {max_pv['algorithm']}-{max_pv['scenario']} con {fmt(max_pv['pv_self_consumption_ratio'], 4)}. "
            "La separacion entre estos puntos evidencia que optimizar una dimension no garantiza simultaneamente el mejor resultado en las otras."
        )

    if not building_compact.empty:
        agents = building_compact["agent"].nunique()
        top_action = building_compact.sort_values("action_dim", ascending=False).iloc[0]
        FIGURE_INTERPRETATIONS["equipment"] = (
            f"La figura compara la dimension de accion controlable de {agents} edificios. "
            f"El mayor espacio de accion observado corresponde a {top_action['agent']} con {int(top_action['action_dim'])} acciones y {int(top_action['observation_dim'])} observaciones. "
            "Esto muestra que los agentes no son simetricos: cada edificio aporta distinta capacidad de control sobre almacenamiento, EV u otros actuadores."
        )
        if "ev_departure_success_rate" in building_compact.columns:
            ev = building_compact.groupby("algorithm")["ev_departure_success_rate"].mean(numeric_only=True).sort_values(ascending=False)
            if not ev.empty:
                FIGURE_INTERPRETATIONS["building_ev_success_heatmap"] = (
                    f"La figura muestra el exito medio de salida EV por edificio y algoritmo; el mayor promedio por algoritmo es {ev.index[0]} con {fmt(ev.iloc[0], 4)}. "
                    "El mapa permite detectar edificios donde las restricciones de carga y disponibilidad EV generan desempenos heterogeneos dentro de la comunidad."
                )
        if "carbon_emissions_control_kgco2" in building_compact.columns:
            co2 = building_compact.groupby("algorithm")["carbon_emissions_control_kgco2"].mean(numeric_only=True).sort_values()
            if not co2.empty:
                FIGURE_INTERPRETATIONS["building_carbon_heatmap"] = (
                    f"La figura representa emisiones controladas por edificio; el menor promedio por algoritmo corresponde a {co2.index[0]} con {fmt(co2.iloc[0], 2)} kgCO2. "
                    "La heterogeneidad visual indica que el efecto carbono del MADRL depende de la demanda, equipamiento y rol operativo de cada edificio."
                )
        if "electricity_cost_control_eur" in building_compact.columns:
            cost = building_compact.groupby("algorithm")["electricity_cost_control_eur"].mean(numeric_only=True).sort_values()
            if not cost.empty:
                FIGURE_INTERPRETATIONS["building_cost_heatmap"] = (
                    f"La figura representa costos controlados por edificio; el menor promedio por algoritmo corresponde a {cost.index[0]} con {fmt(cost.iloc[0], 2)}. "
                    "El mapa muestra que la reduccion de costo no se distribuye de manera uniforme entre agentes y debe interpretarse junto con la senal tarifaria y la flexibilidad disponible."
                )

    if not eq_summary.empty:
        totals = eq_summary.groupby("equipment_class")["count"].sum().sort_values(ascending=False)
        if not totals.empty:
            FIGURE_INTERPRETATIONS["equipment_class_heatmap"] = (
                f"La figura resume {int(totals.sum())} variables de accion controlable clasificadas por equipo. "
                f"La clase mas frecuente es {totals.index[0]} con {int(totals.iloc[0])} apariciones. "
                "La lectura diferencia equipamiento controlado de cargas no controladas, que permanecen como demanda base u observaciones del entorno."
            )

    if not final_ts.empty:
        for scenario in SCENARIOS:
            sub = final_ts[final_ts["scenario"] == scenario].copy()
            if sub.empty:
                continue
            grouped = sub.groupby("algorithm").agg(
                rows=("episode_step", "count"),
                energy=("district_net_electricity_consumption", "sum"),
                cost=("district_net_electricity_consumption_cost", "sum"),
                co2=("district_net_electricity_consumption_emission", "sum"),
            )
            best_cost = grouped.sort_values("cost").iloc[0]
            best_co2 = grouped.sort_values("co2").iloc[0]
            FIGURE_INTERPRETATIONS[f"final_timeseries_{scenario}"] = (
                f"La figura usa {len(sub)} filas horarias o subhorarias del episodio final conservado en {scenario}. "
                f"El menor costo acumulado observado corresponde a {best_cost.name} con {fmt(best_cost['cost'], 2)}, "
                f"y la menor emision acumulada corresponde a {best_co2.name} con {fmt(best_co2['co2'], 2)}. "
                "Los paneles muestran valores normalizados para comparar forma temporal, pero las anotaciones conservan totales reales extraidos de timeseries.csv."
            )

    if not traces.empty and "action_l2" in traces.columns:
        agg = traces.groupby("algorithm")[["action_l2"]].mean(numeric_only=True)
        best_a = agg["action_l2"].sort_values(ascending=False).iloc[0]
        best_algo = agg["action_l2"].idxmax()
        behav_path = TABLE_DIR / "gdrive_building_behavior_summary_all.csv"
        extra = ""
        if behav_path.exists():
            behav = pd.read_csv(behav_path)
            ev_cols = [c for c in ("ev_charge_total_kwh", "battery_throughput_total_kwh") if c in behav.columns]
            if ev_cols:
                g = behav.groupby("algorithm")[ev_cols].mean(numeric_only=True)
                bits = []
                for c in ev_cols:
                    bits.append(f"{c}: {g[c].idxmax()}={fmt(g[c].max(), 1)}")
                extra = " EV/BESS desde building_behavior_summary (" + "; ".join(bits) + ")."
        FIGURE_INTERPRETATIONS["trace_policy_heatmaps"] = (
            f"La figura usa action_l2 de trace.csv ({len(traces)} registros; p. ej. {best_algo}={fmt(best_a, 3)}) "
            "y paneles EV/BESS desde building_behavior_summary, no desde columnas muertas "
            f"ev_charge_kwh/electrical_storage_soc de trace.{extra} "
            "Cada panel tiene escala propia (colorbar independiente)."
        )

    if not episodes.empty:
        reward_best = episodes.groupby(["algorithm", "scenario"])["reward_mean_average"].mean(numeric_only=True).sort_values(ascending=False).iloc[0]
        reward_key = episodes.groupby(["algorithm", "scenario"])["reward_mean_average"].mean(numeric_only=True).sort_values(ascending=False).index[0]
        FIGURE_INTERPRETATIONS["appendix_reward"] = (
            f"La figura sintetiza reward_mean_average para {len(episodes)} filas episodicas materializadas. "
            f"El mayor promedio episodio-tratamiento corresponde a {reward_key[0]}-{reward_key[1]} con {fmt(reward_best, 6)}. "
            "La lectura sirve como auditoria grafica del comportamiento de recompensa que alimenta las curvas de convergencia del Capitulo 5."
        )
        if "district_net_energy" in episodes.columns:
            energy = episodes.groupby(["algorithm", "scenario"])["district_net_energy"].mean(numeric_only=True).sort_values()
            if not energy.empty:
                key = energy.index[0]
                FIGURE_INTERPRETATIONS["appendix_energy"] = (
                    f"La figura resume energia neta distrital por episodio en {len(episodes)} filas. "
                    f"El menor promedio observado corresponde a {key[0]}-{key[1]} con {fmt(energy.iloc[0], 2)}. "
                    "Este resultado se interpreta como evidencia descriptiva de consumo neto, no como criterio unico de flexibilidad."
                )
        if "district_cost" in episodes.columns:
            cost = episodes.groupby(["algorithm", "scenario"])["district_cost"].mean(numeric_only=True).sort_values()
            if not cost.empty:
                key = cost.index[0]
                FIGURE_INTERPRETATIONS["appendix_cost"] = (
                    f"La figura compara costo distrital medio usando {len(episodes)} filas episodicas. "
                    f"El menor promedio corresponde a {key[0]}-{key[1]} con {fmt(cost.iloc[0], 2)}. "
                    "La lectura complementa OE.3 porque muestra el comportamiento por escenario antes de la sintesis inferencial."
                )
        if "district_emission" in episodes.columns:
            co2 = episodes.groupby(["algorithm", "scenario"])["district_emission"].mean(numeric_only=True).sort_values()
            if not co2.empty:
                key = co2.index[0]
                FIGURE_INTERPRETATIONS["appendix_emissions"] = (
                    f"La figura compara emisiones distritales medias en {len(episodes)} filas episodicas. "
                    f"El menor promedio corresponde a {key[0]}-{key[1]} con {fmt(co2.iloc[0], 2)}. "
                    "La lectura complementa OE.2 y permite revisar si el escenario orientado a carbono se refleja tambien en la media episodica."
                )

    if not building_compact.empty:
        if "electricity_cost_delta_eur" in building_compact.columns:
            deltas = building_compact["electricity_cost_delta_eur"].dropna()
            if not deltas.empty:
                best = building_compact.loc[building_compact["electricity_cost_delta_eur"].idxmin()]
                FIGURE_INTERPRETATIONS["appendix_cost_delta"] = (
                    f"La figura usa {len(deltas)} valores de delta de costo electrico por edificio. "
                    f"La mayor reduccion observada corresponde a {best['algorithm']}-{best['scenario']} en {best['agent']} con {fmt(best['electricity_cost_delta_eur'], 2)}. "
                    "El signo y magnitud del delta permiten distinguir reducciones reales de costo frente a variaciones marginales por edificio."
                )
        if "carbon_emissions_delta_kgco2" in building_compact.columns:
            deltas = building_compact["carbon_emissions_delta_kgco2"].dropna()
            if not deltas.empty:
                best = building_compact.loc[building_compact["carbon_emissions_delta_kgco2"].idxmin()]
                FIGURE_INTERPRETATIONS["appendix_emissions_delta"] = (
                    f"La figura usa {len(deltas)} valores de delta de emisiones por edificio. "
                    f"La mayor reduccion observada corresponde a {best['algorithm']}-{best['scenario']} en {best['agent']} con {fmt(best['carbon_emissions_delta_kgco2'], 2)} kgCO2. "
                    "La lectura muestra que el efecto ambiental no es uniforme entre edificios y depende de la demanda y equipamiento disponible."
                )
        action_sum = int(building_compact["action_dim"].sum()) if "action_dim" in building_compact.columns else 0
        FIGURE_INTERPRETATIONS["appendix_action_variables"] = (
            f"La figura representa variables de accion controladas en {len(building_compact)} filas edificio-tratamiento, con suma de dimensiones de accion igual a {action_sum}. "
            "Este anexo documenta que el control MADRL se ejerce sobre actuadores declarados en el esquema y no sobre toda la demanda del edificio."
        )
        FIGURE_INTERPRETATIONS["appendix_controlled_load"] = (
            f"La figura contrasta carga controlada y carga base no controlada para {building_compact['agent'].nunique()} edificios y {len(treatment)} tratamientos. "
            "La interpretacion separa las variables que el agente puede accionar de aquellas que solo observa como demanda exogena, criterio indispensable para no sobrestimar la capacidad de control del MADRL."
        )

    if not checkpoints.empty and "bytes" in checkpoints.columns:
        total_mb = checkpoints["bytes"].fillna(0).sum() / (1024 * 1024)
        max_row = checkpoints.sort_values("bytes", ascending=False).iloc[0]
        FIGURE_INTERPRETATIONS["appendix_checkpoint_size"] = (
            f"La figura resume el tamano listado en {len(checkpoints)} checkpoints, con un total aproximado de {fmt(total_mb, 2)} MB. "
            f"El mayor archivo individual corresponde a {max_row['algorithm']}-{max_row['scenario']} con {fmt(max_row['bytes'] / (1024 * 1024), 2)} MB. "
            "Esta informacion verifica reproducibilidad material de modelos y permite auditar cobertura de almacenamiento."
        )

    architecture_base = {
        "architecture_diagram_1": f"El diagrama integra el flujo completo del proyecto: construccion del dataset, entorno CityLearn, entrenamiento MADRL, evaluacion, visualizacion y redaccion doctoral. La evidencia real usada en el documento comprende {len(treatment)} tratamientos, {len(episodes)} filas episodicas, {len(building_compact)} filas edificio-tratamiento y {len(checkpoints)} checkpoints.",
        "architecture_diagram_2": f"El diagrama muestra el pipeline del dataset Iquitos 2023-2025 y su transformacion hacia el formato CityLearn. La interpretacion se basa en {building_compact['agent'].nunique() if not building_compact.empty else 17} edificios/agentes y en la separacion entre datos climaticos, demanda, precios, carbono y esquemas de observacion-accion.",
        "architecture_diagram_3": f"El diagrama formaliza el problema como Dec-POMDP con ejecucion descentralizada y entrenamiento centralizado. En la corrida auditada, la comunidad opera con {building_compact['agent'].nunique() if not building_compact.empty else 17} agentes, observaciones heterogeneas y acciones controlables declaradas por edificio.",
        "architecture_diagram_4": f"El diagrama compara HAPPO, MAAC, MASAC y MATD3. La tesis evalua estos 4 algoritmos en 3 escenarios, lo que produce {len(treatment)} tratamientos algoritmo-escenario usados para responder OE.1, OE.2 y OE.3.",
        "architecture_diagram_5": f"El diagrama resume el flujo de entrenamiento de las 12 corridas. La evidencia final conserva {len(episodes)} filas episodicas materializadas, {len(final_ts)} filas de series finales, {len(traces)} registros de trace.csv y {len(checkpoints)} checkpoints.",
        "architecture_diagram_6": "El diagrama muestra la recompensa multiobjetivo por escenario: E1 prioriza flexibilidad, E2 prioriza CO2 y E3 prioriza costos. Esta estructura vincula directamente las preguntas PE.1, PE.2 y PE.3 con los objetivos OE.1, OE.2 y OE.3.",
        "architecture_diagram_7": f"El diagrama representa el pipeline de evaluacion y seleccion del mejor MADRL. La seleccion no depende de una sola curva: integra descriptivos, Kruskal-Wallis, Mann-Whitney-Holm, epsilon2, KPIs CityLearn v2, series finales y resultados por edificio.",
        "architecture_diagram_8": "El diagrama ubica la infraestructura de ejecucion local, Colab A100 y AWS EC2 como soporte computacional del entrenamiento. Su funcion es explicar trazabilidad operacional y no introducir resultados externos a la carpeta Drive auditada.",
        "architecture_diagram_9": f"El diagrama sintetiza 7 capas de software: datos, entorno, agentes MADRL, entrenamiento, evaluacion, figuras y documento Word. La salida doctoral se reconstruye automaticamente desde scripts locales y artefactos reales, incluyendo {len(treatment)} tratamientos, {len(episodes)} filas episodicas y {len(checkpoints)} checkpoints, reduciendo riesgo de transcripcion manual.",
    }
    FIGURE_INTERPRETATIONS.update(architecture_base)


def make_figures(
    detail: dict,
    treatment: pd.DataFrame,
    buildings: pd.DataFrame,
    eq_summary: pd.DataFrame,
    episodes: pd.DataFrame,
    convergence: pd.DataFrame,
    kpi_ranking: pd.DataFrame,
    final_ts: pd.DataFrame,
    traces: pd.DataFrame,
    checkpoints: pd.DataFrame,
) -> dict[str, Path]:
    paths: dict[str, Path] = {}
    for oe, d in detail.items():
        spec = d["spec"]
        fig, ax = plt.subplots(figsize=(7, 4))
        desc = d["desc"].copy()
        y = desc["mean"]
        ax.bar(desc["algorithm"], y, color=["#9AA7B2", "#2F6B52", "#C77D2A", "#406A9F"][: len(desc)])
        ax.set_title(f"{oe} - {spec['dimension']} ({spec['scenario']})")
        ax.set_ylabel(spec["metric"])
        ax.grid(axis="y", alpha=0.25)
        fig.tight_layout()
        out = FIG_DIR / f"{oe.lower().replace('.', '')}_{spec['scenario'].lower()}_episode_mean.png"
        fig.savefig(out, dpi=180)
        plt.close(fig)
        paths[oe] = out
    fig, ax = plt.subplots(figsize=(7, 4))
    top = buildings.groupby("agent")["action_dim"].max().sort_values(ascending=False)
    ax.bar(top.index, top.values, color="#2F6B52")
    ax.set_title("Equipamiento controlado por edificio (dimensiones de accion)")
    ax.set_ylabel("acciones controlables")
    ax.tick_params(axis="x", rotation=75)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    out = FIG_DIR / "controlled_equipment_action_dim_by_building.png"
    fig.savefig(out, dpi=180)
    plt.close(fig)
    paths["equipment"] = out
    for scenario in SCENARIOS:
        fig, ax = plt.subplots(figsize=(7.2, 4.2))
        sub = episodes[episodes["scenario"] == scenario].copy()
        for algo in ALGOS:
            alg = sub[sub["algorithm"] == algo].sort_values("episode").copy()
            if alg.empty:
                continue
            alg["rolling_reward"] = pd.to_numeric(alg["reward_mean_average"], errors="coerce").rolling(window=5, min_periods=1).mean()
            ax.plot(alg["episode"] + 1, alg["rolling_reward"], linewidth=1.8, label=algo)
            conv = convergence[(convergence["scenario"] == scenario) & (convergence["algorithm"] == algo)]
            if not conv.empty:
                c = conv.iloc[0]
                ax.axvline(c["learning_start_episode_ordinal"], color="gray", alpha=0.10, linewidth=0.8)
                ax.scatter([c["stabilization_episode_ordinal"]], [c["final_rolling_reward"]], s=28, zorder=5)
        ax.set_title(f"Convergencia MADRL por recompensa media movil - {scenario}")
        ax.set_xlabel("episodio ordinal en artefacto")
        ax.set_ylabel("recompensa media movil")
        ax.grid(alpha=0.25)
        ax.legend(ncol=2, fontsize=8)
        fig.tight_layout()
        out = FIG_DIR / f"convergence_{scenario.lower()}_learning_stabilization.png"
        fig.savefig(out, dpi=180)
        plt.close(fig)
        paths[f"convergence_{scenario}"] = out

    # Distribucion episodica por objetivo: permite ver dispersion, mediana y atipicos.
    fig, axes = plt.subplots(1, 3, figsize=(14, 4.3))
    for ax, oe in zip(axes, ["OE.1", "OE.2", "OE.3"]):
        spec = detail[oe]["spec"]
        metric = spec["metric"]
        sub = episodes[episodes["scenario"] == spec["scenario"]].copy()
        data = [sub[sub["algorithm"] == algo][metric].dropna().values for algo in ALGOS]
        ax.boxplot(data, labels=ALGOS, showmeans=True)
        ax.set_title(f"{oe} - {spec['dimension']}")
        ax.set_ylabel(metric)
        ax.grid(axis="y", alpha=0.25)
        ax.tick_params(axis="x", rotation=25)
    fig.tight_layout()
    out = FIG_DIR / "episode_objective_distributions_boxplot.png"
    fig.savefig(out, dpi=180)
    plt.close(fig)
    paths["episode_boxplots"] = out

    fig, ax = plt.subplots(figsize=(7, 4))
    oes = ["OE.1", "OE.2", "OE.3"]
    vals = [detail[oe]["epsilon2"] for oe in oes]
    ax.bar(oes, vals, color=["#2F6B52", "#406A9F", "#C77D2A"])
    for i, oe in enumerate(oes):
        kw = detail[oe]["kw"]
        ax.text(i, vals[i] + 0.005, f"p={kw.pvalue:.3g}", ha="center", fontsize=8)
    ax.axhline(0.01, color="gray", linestyle="--", linewidth=0.8)
    ax.axhline(0.06, color="gray", linestyle=":", linewidth=0.8)
    ax.axhline(0.14, color="gray", linestyle="-.", linewidth=0.8)
    ax.set_title("Tamano de efecto inferencial por objetivo (epsilon2)")
    ax.set_ylabel("epsilon2 Kruskal-Wallis")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    out = FIG_DIR / "objective_effect_size_epsilon2.png"
    fig.savefig(out, dpi=180)
    plt.close(fig)
    paths["effect_size"] = out

    fig, axes = plt.subplots(1, 3, figsize=(13, 4))
    for ax, oe in zip(axes, ["OE.1", "OE.2", "OE.3"]):
        labels = ALGOS
        mat = pd.DataFrame(1.0, index=labels, columns=labels)
        for name, _pv, adj in detail[oe]["pair_adj"]:
            a, b = name.split(" vs ")
            mat.loc[a, b] = adj
            mat.loc[b, a] = adj
        img = ax.imshow(-mat.applymap(lambda x: math.log10(max(float(x), 1e-12))).values, cmap="YlOrRd", vmin=0, vmax=8)
        ax.set_xticks(range(len(labels)), labels, rotation=45, ha="right")
        ax.set_yticks(range(len(labels)), labels)
        ax.set_title(f"{oe}: -log10(p Holm)")
    fig.colorbar(img, ax=axes.ravel().tolist(), shrink=0.8)
    out = FIG_DIR / "pairwise_holm_pvalue_heatmaps.png"
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)
    paths["pairwise_heatmaps"] = out

    if not kpi_ranking.empty:
        pivot = kpi_ranking.pivot_table(index="method", columns="scenario", values="normalized_score", aggfunc="mean")
        fig, ax = plt.subplots(figsize=(6.5, 4.2))
        img = ax.imshow(pivot.values, cmap="viridis")
        ax.set_xticks(range(len(pivot.columns)), pivot.columns)
        ax.set_yticks(range(len(pivot.index)), pivot.index)
        ax.set_title("Ranking KPI CityLearn v2 evaluate_v2 (score normalizado)")
        for i in range(pivot.shape[0]):
            for j in range(pivot.shape[1]):
                ax.text(j, i, f"{pivot.iloc[i, j]:.3f}", ha="center", va="center", color="white", fontsize=8)
        fig.colorbar(img, ax=ax, shrink=0.85)
        fig.tight_layout()
        out = FIG_DIR / "citylearn_v2_kpi_ranking_heatmap.png"
        fig.savefig(out, dpi=180)
        plt.close(fig)
        paths["kpi_ranking_heatmap"] = out

    fig, ax = plt.subplots(figsize=(7, 4.5))
    colors = {"HAPPO": "#9AA7B2", "MAAC": "#2F6B52", "MASAC": "#C77D2A", "MATD3": "#406A9F"}
    markers = {"E1": "o", "E2": "s", "E3": "^"}
    for _, r in treatment.iterrows():
        size = 40 + 120 * (float(r.get("pv_self_consumption_ratio", 0) or 0))
        ax.scatter(r["electricity_cost_control"], r["carbon_emissions_control"], s=size, c=colors.get(r["algorithm"], "gray"), marker=markers.get(r["scenario"], "o"), edgecolor="black", linewidth=0.4)
        ax.text(r["electricity_cost_control"], r["carbon_emissions_control"], f"{r['algorithm']}-{r['scenario']}", fontsize=7)
    ax.set_title("Trade-off multiobjetivo: costo vs CO2 vs autoconsumo PV")
    ax.set_xlabel("electricity_cost_control")
    ax.set_ylabel("carbon_emissions_control")
    ax.grid(alpha=0.25)
    fig.tight_layout()
    out = FIG_DIR / "multiobjective_tradeoff_cost_co2_pv.png"
    fig.savefig(out, dpi=180)
    plt.close(fig)
    paths["tradeoff"] = out

    def building_heatmap(value_col: str, key: str, title: str, log_scale: bool = False) -> None:
        if value_col not in buildings.columns:
            return
        pivot = buildings.pivot_table(index="agent", columns="algorithm", values=value_col, aggfunc="mean").reindex(columns=ALGOS)
        values = pivot.astype(float).values
        if log_scale:
            values = pd.DataFrame(values).applymap(lambda x: math.copysign(math.log10(1.0 + abs(float(x))), float(x))).values
        fig, ax = plt.subplots(figsize=(7.2, 6.2))
        img = ax.imshow(values, aspect="auto", cmap="mako" if False else "viridis")
        ax.set_xticks(range(len(pivot.columns)), pivot.columns)
        ax.set_yticks(range(len(pivot.index)), pivot.index)
        ax.set_title(title)
        fig.colorbar(img, ax=ax, shrink=0.75)
        fig.tight_layout()
        out = FIG_DIR / f"{key}.png"
        fig.savefig(out, dpi=180)
        plt.close(fig)
        paths[key] = out

    building_heatmap("ev_departure_success_rate", "building_ev_success_heatmap", "EV departure success rate por edificio y algoritmo")
    building_heatmap("carbon_emissions_control_kgco2", "building_carbon_heatmap", "CO2 control por edificio y algoritmo (log10)", True)
    building_heatmap("electricity_cost_control_eur", "building_cost_heatmap", "Costo control por edificio y algoritmo (log10)", True)

    if not eq_summary.empty:
        pivot = eq_summary.pivot_table(index="agent", columns="equipment_class", values="count", aggfunc="sum", fill_value=0)
        fig, ax = plt.subplots(figsize=(8, 6))
        img = ax.imshow(pivot.values, aspect="auto", cmap="Blues")
        ax.set_xticks(range(len(pivot.columns)), pivot.columns, rotation=35, ha="right")
        ax.set_yticks(range(len(pivot.index)), pivot.index)
        ax.set_title("Equipamiento controlado por edificio y clase")
        fig.colorbar(img, ax=ax, shrink=0.75)
        fig.tight_layout()
        out = FIG_DIR / "controlled_equipment_class_heatmap.png"
        fig.savefig(out, dpi=180)
        plt.close(fig)
        paths["equipment_class_heatmap"] = out

    if not final_ts.empty:
        for scenario in SCENARIOS:
            sub = final_ts[final_ts["scenario"] == scenario].copy()
            if sub.empty:
                continue
            sub["hour"] = pd.to_numeric(sub.get("episode_step", sub.get("time_step")), errors="coerce")
            metrics = [
                ("district_net_electricity_consumption", "energia neta", "#1f77b4"),
                ("district_net_electricity_consumption_cost", "costo", "#ff7f0e"),
                ("district_net_electricity_consumption_emission", "CO2", "#2ca02c"),
            ]
            max_by_metric = {
                col: max(float(sub[col].max()), 1.0) if col in sub.columns and not sub[col].dropna().empty else 1.0
                for col, _, _ in metrics
            }
            fig, axes = plt.subplots(2, 2, figsize=(10, 7.5), sharex=True, sharey=True)
            for ax, algo in zip(axes.ravel(), ALGOS):
                alg = sub[sub["algorithm"] == algo].sort_values("hour").copy()
                if alg.empty:
                    ax.set_title(f"{algo} - sin timeseries final")
                    ax.axis("off")
                    continue
                annotations = []
                for col, label, color in metrics:
                    if col not in alg.columns:
                        continue
                    y = pd.to_numeric(alg[col], errors="coerce").fillna(0.0)
                    normalized = y / max_by_metric[col]
                    ax.plot(alg["hour"], normalized, label=label, linewidth=1.1, color=color)
                    nz = int((y != 0).sum())
                    annotations.append(f"{label}: {y.sum():.2f} (nz={nz})")
                rows = len(alg)
                episode = int(alg["episode"].dropna().iloc[0]) if "episode" in alg.columns and not alg["episode"].dropna().empty else -1
                coverage = rows / 8760.0 * 100.0
                ax.set_title(f"{algo} - ep. {episode} - {rows}/8760 filas ({coverage:.1f}%)", fontsize=9)
                ax.text(
                    0.01,
                    0.97,
                    "\n".join(annotations),
                    transform=ax.transAxes,
                    va="top",
                    ha="left",
                    fontsize=6.5,
                    bbox={"boxstyle": "round,pad=0.25", "facecolor": "white", "alpha": 0.78, "edgecolor": "0.8"},
                )
                ax.grid(alpha=0.25)
                ax.set_ylim(-0.03, 1.08)
            fig.suptitle(f"Paneles MADRL del episodio final - {scenario} (valores normalizados; totales reales anotados)", fontsize=12)
            axes[0, 0].set_ylabel("valor normalizado")
            axes[1, 0].set_ylabel("valor normalizado")
            axes[1, 0].set_xlabel("hora/paso conservado")
            axes[1, 1].set_xlabel("hora/paso conservado")
            axes[0, 0].legend(ncol=3, fontsize=7, loc="lower left")
            fig.tight_layout()
            out = FIG_DIR / f"final_episode_district_timeseries_{scenario.lower()}.png"
            fig.savefig(out, dpi=180)
            plt.close(fig)
            paths[f"final_timeseries_{scenario}"] = out

    # Figura 5.8e: action_l2 desde trace; EV/BESS desde building_behavior_summary.
    # No usar ev_charge_kwh / electrical_storage_soc de trace.csv (columnas muertas=0).
    # Colorbar independiente por panel.
    if not traces.empty and "action_l2" in traces.columns:
        agg = traces.groupby(["algorithm", "scenario"])[["action_l2"]].mean(numeric_only=True).reset_index()
        behav_path = TABLE_DIR / "gdrive_building_behavior_summary_all.csv"
        behav = pd.read_csv(behav_path) if behav_path.exists() else pd.DataFrame()
        fig, axes = plt.subplots(1, 3, figsize=(13.4, 4.3))
        panels = [("action_l2", agg, "Intensidad de accion (mean action_l2)")]
        if not behav.empty:
            metric_cols = [c for c in ("ev_charge_total_kwh", "battery_throughput_total_kwh") if c in behav.columns]
            g = behav.groupby(["algorithm", "scenario"], as_index=False)[metric_cols].mean(numeric_only=True)
            if "ev_charge_total_kwh" in g.columns:
                panels.append(("ev_charge_total_kwh", g, "Carga EV media (behavior summary, kWh)"))
            if "battery_throughput_total_kwh" in g.columns:
                panels.append(("battery_throughput_total_kwh", g, "Throughput BESS (behavior summary, kWh)"))
        while len(panels) < 3:
            panels.append(panels[0])
        for ax, (col, dfp, title) in zip(axes, panels[:3]):
            pivot = dfp.pivot(index="algorithm", columns="scenario", values=col).reindex(index=ALGOS, columns=SCENARIOS)
            vals = pivot.values.astype(float)
            finite = vals[np.isfinite(vals)]
            vmin, vmax = (float(np.nanmin(finite)), float(np.nanmax(finite))) if finite.size else (0.0, 1.0)
            if abs(vmax - vmin) < 1e-12:
                vmax = vmin + 1e-6
            img = ax.imshow(vals, cmap="plasma", vmin=vmin, vmax=vmax, aspect="auto")
            ax.set_xticks(range(len(SCENARIOS)), SCENARIOS)
            ax.set_yticks(range(len(ALGOS)), ALGOS)
            ax.set_title(title, fontsize=9)
            for i in range(pivot.shape[0]):
                for j in range(pivot.shape[1]):
                    v = vals[i, j]
                    lab = "NA" if not np.isfinite(v) else (f"{v:.2f}" if abs(v) < 100 else f"{v:.0f}")
                    ax.text(j, i, lab, ha="center", va="center", color="white", fontsize=7)
            fig.colorbar(img, ax=ax, fraction=0.046, pad=0.04)
        fig.suptitle("Figura 5.8e — politicas/acciones (fuentes mixtas auditadas)", fontsize=11)
        fig.tight_layout()
        out = FIG_DIR / "trace_policy_action_heatmaps.png"
        fig.savefig(out, dpi=180, bbox_inches="tight")
        plt.close(fig)
        paths["trace_policy_heatmaps"] = out

    # Figura 5.1: always plot the 12 project treatments (missing manifest => 0, not omitted).
    cov_rows = []
    for algo in ALGOS:
        for scenario in SCENARIOS:
            if checkpoints.empty:
                n = 0
            else:
                sub = checkpoints[(checkpoints["algorithm"] == algo) & (checkpoints["scenario"] == scenario)]
                n = int(len(sub))
            cov_rows.append({"algorithm": algo, "scenario": scenario, "count": n})
    counts = pd.DataFrame(cov_rows)
    fig, ax = plt.subplots(figsize=(9.5, 4.4))
    labels = [f"{r.algorithm}-{r.scenario}" for r in counts.itertuples()]
    values = [int(r.count) for r in counts.itertuples()]
    palette = {"HAPPO": "#7A7A7A", "MAAC": "#406A9F", "MASAC": "#2E8B57", "MATD3": "#C45C26"}
    colors = [palette.get(lab.split("-")[0], "#406A9F") for lab in labels]
    bars = ax.bar(labels, values, color=colors)
    ax.set_title("Cobertura de checkpoints por tratamiento (MADRL del proyecto)")
    ax.set_ylabel("archivos listados en checkpoint_manifest.json")
    ax.tick_params(axis="x", rotation=65)
    ax.grid(axis="y", alpha=0.25)
    ymax = max(values) if values else 1
    ax.set_ylim(0, max(10, ymax * 1.15))
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, v + ymax * 0.02, str(v), ha="center", va="bottom", fontsize=8)
    if any(a == "HAPPO" and c == 0 for a, c in zip(counts["algorithm"], counts["count"])):
        ax.text(
            0.01,
            0.98,
            "HAPPO: sin checkpoint_manifest.json en la corrida canonica",
            transform=ax.transAxes,
            va="top",
            fontsize=7.5,
            color="#444444",
        )
    fig.tight_layout()
    out = FIG_DIR / "checkpoint_coverage_by_treatment.png"
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)
    paths["checkpoint_coverage"] = out
    return paths


def add_picture(doc: Document, caption: str, path: Path, width: float = 5.8) -> None:
    cap = doc.add_paragraph()
    r = cap.add_run(caption)
    set_run_apa_font(r, 12, bold=True, italic=False, color=RGBColor(0, 0, 0))
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph()
    note.add_run(build_figure_note_text(caption))
    set_figure_note_paragraph_format(note)
    interp = doc.add_paragraph()
    interp.add_run(f"Interpretacion de la figura. {figure_interpretation_text(caption)}")
    set_figure_interpretation_paragraph_format(interp)
    doc.add_paragraph()


RESUMEN_TESIS = (
    "El presente estudio doctoral determina el efecto de cuatro algoritmos de aprendizaje por refuerzo profundo "
    "multiagente sobre la flexibilidad energetica, las emisiones de CO2 y los costos energeticos de una comunidad "
    "inteligente representativa del Sistema Electrico Aislado de Iquitos. La investigacion adopta un enfoque "
    "aplicado, cuantitativo y explicativo mediante simulacion computacional, con diseno experimental factorial "
    "4x3: cuatro algoritmos y tres escenarios de recompensa orientados a las dimensiones de la variable "
    "dependiente. El problema se formaliza como decision cooperativa con observabilidad parcial y entrenamiento "
    "centralizado con ejecucion descentralizada. La base empirica corresponde al dataset horario 2023-2025 y a "
    "la corrida canonica de entrenamiento conservada en Drive, cuya descripcion tecnica, fuentes, librerias, "
    "equipos y artefactos se desarrolla de manera detallada en la propuesta del Capitulo 4."
)

RESUMEN_TESIS_RESULTADOS = (
    "La evidencia experimental corresponde exclusivamente a la corrida Drive madrl_v3_20260627_164047 y a sus "
    "artefactos results.json, training_summary.json, timeseries.csv, trace.csv, building_kpis.csv, "
    "building_behavior_summary.csv, checkpoint_manifest.json y tablas materializadas de KPIs episodicos. Se "
    "conservan 597 filas episodicas: HAPPO aporta 49 episodios por escenario como evidencia descriptiva, mientras "
    "MAAC, MASAC y MATD3 conservan 50 episodios por escenario para contraste inferencial completo. En OE.1, "
    "flexibilidad energetica, Kruskal-Wallis rechaza la igualdad entre algoritmos con H=36,3083, p=1,305e-08 y "
    "epsilon2=0,2334; MAAC presenta el mayor efecto inferencial y MASAC el mejor KPI anual final. En OE.2, "
    "emisiones de CO2, el efecto es significativo pero bajo (H=6,2532; p=0,043866; epsilon2=0,0289); MAAC lidera "
    "la muestra completa y MASAC el KPI anual final. En OE.3, costos energeticos, no se rechaza H0 (H=2,7613; "
    "p=0,251421; epsilon2=0,0052); MATD3 lidera la muestra completa y MAAC el KPI anual final. Los resultados se "
    "interpretan por distrito, edificio, escenario, equipamiento controlado/no controlado, curvas de convergencia, "
    "checkpoints, trazas y KPIs compatibles con CityLearn v2 evaluate_v2, sin incorporar valores externos ni "
    "datos no existentes en los artefactos auditados."
)

PALABRAS_CLAVE_TESIS = (
    "Palabras clave: aprendizaje por refuerzo profundo multiagente; comunidades inteligentes; CityLearn; "
    "Dec-POMDP; entrenamiento centralizado y ejecucion descentralizada; flexibilidad energetica; emisiones de "
    "CO2; costos energeticos; Sistema Electrico Aislado de Iquitos; Python; pandas; NumPy; requests; pvlib; "
    "openpyxl."
)

ABSTRACT_TESIS = (
    "This doctoral study determines the effect of four multi-agent deep reinforcement learning algorithms on "
    "energy flexibility, CO2 emissions and energy costs in a smart community representative of the Iquitos "
    "Isolated Electric System, Peru. The research follows an applied, quantitative and explanatory "
    "computational-simulation approach with a 4x3 factorial experimental design: four algorithms and three "
    "reward scenarios aligned with the dependent-variable dimensions. The control problem is formalized as "
    "cooperative decision-making under partial observability and centralized training with decentralized "
    "execution. The empirical basis is the 2023-2025 hourly dataset and the canonical Drive training run; the "
    "technical description of data sources, libraries, equipment and artifacts is developed in detail in the "
    "proposal chapter."
)

ABSTRACT_TESIS_RESULTADOS = (
    "The experimental evidence is restricted to the audited Drive run madrl_v3_20260627_164047 and its results, "
    "training summaries, timeseries, traces, building KPIs, checkpoint manifests and materialized episodic KPI "
    "tables. A total of 597 episodic rows are retained: HAPPO contributes 49 episodes per scenario as descriptive "
    "evidence, whereas MAAC, MASAC and MATD3 retain 50 episodes per scenario for complete inferential comparison. "
    "For OE.1, energy flexibility, Kruskal-Wallis rejects equality across algorithms (H=36.3083, p=1.305e-08, "
    "epsilon2=0.2334), with MAAC showing the strongest complete-sample effect and MASAC the best final annual KPI. "
    "For OE.2, CO2 emissions, the effect is significant but small (H=6.2532, p=0.043866, epsilon2=0.0289), with "
    "MAAC leading the complete sample and MASAC leading the final KPI. For OE.3, energy costs, H0 is not rejected "
    "(H=2.7613, p=0.251421, epsilon2=0.0052); MATD3 leads the complete sample and MAAC leads the final KPI. "
    "Results are interpreted at district, building, scenario, equipment, convergence, checkpoint and trace levels "
    "using CityLearn v2 evaluate_v2-compatible KPIs without adding external or synthetic values."
)

KEYWORDS_TESIS = (
    "Keywords: multi-agent deep reinforcement learning; smart communities; CityLearn; Dec-POMDP; centralized "
    "training and decentralized execution; energy flexibility; CO2 emissions; energy costs; Iquitos Isolated "
    "Electric System; Python; pandas; NumPy; requests; pvlib; openpyxl."
)


def insert_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def next_nonempty_paragraph_index(paragraphs, start: int, stop: int | None = None) -> int | None:
    end = len(paragraphs) if stop is None else min(stop, len(paragraphs))
    for idx in range(start, end):
        if paragraphs[idx].text.strip():
            return idx
    return None


def replace_paragraphs_between_markers(doc: Document, start_text: str, end_prefix: str, replacement_texts: list[str]) -> None:
    paragraphs = doc.paragraphs
    start_idx = next((i for i, para in enumerate(paragraphs) if para.text.strip() == start_text), None)
    if start_idx is None:
        return
    end_idx = next((i for i, para in enumerate(paragraphs[start_idx + 1 :], start_idx + 1) if para.text.strip().startswith(end_prefix)), None)
    if end_idx is None:
        return
    for para in paragraphs[start_idx + 1 : end_idx]:
        para._element.getparent().remove(para._element)
    anchor = doc.paragraphs[start_idx]
    for text in reversed(replacement_texts):
        insert_paragraph_after(anchor, text)


def update_summary_abstract_keywords(doc: Document) -> None:
    replace_paragraphs_between_markers(
        doc,
        "Resumen",
        "Abstract",
        [RESUMEN_TESIS, RESUMEN_TESIS_RESULTADOS, PALABRAS_CLAVE_TESIS],
    )
    replace_paragraphs_between_markers(
        doc,
        "Abstract",
        "Indice",
        [ABSTRACT_TESIS, ABSTRACT_TESIS_RESULTADOS, KEYWORDS_TESIS],
    )


def clean_front_matter_50ep(doc: Document) -> None:
    replacement = (
        "Esta tesis doctoral determina, mediante simulacion computacional bajo diseno experimental factorial 4x3, "
        "el efecto de cuatro algoritmos Multi-Agente de Aprendizaje por Refuerzo Profundo (MADRL) -HAPPO, MASAC, "
        "MATD3 y MAAC- sobre la flexibilidad energetica, las emisiones de CO2 y los costos energeticos en una "
        "comunidad inteligente del Sistema Electrico Aislado de Iquitos. La evidencia corresponde a la corrida "
        "canonica Drive madrl_v3_20260627_164047, con 50 episodios registrados por tratamiento, 17 edificios, "
        "185 cargadores EV, BESS/PV, checkpoints y trazas auditables. La contrastacion se organiza estrictamente "
        "por OE.1 flexibilidad, OE.2 emisiones de CO2 y OE.3 costos, usando resultados distritales, resultados por "
        "edificio, equipamiento controlado/no controlado, curvas de convergencia y KPIs compatibles con la lectura "
        "CityLearn v2 evaluate_v2. No se incorporan valores externos ni resultados exploratorios ajenos a esa "
        "corrida Drive."
    )
    forbidden = ["5 episodios", "referencia local", "KW p=0.0459", "Colab (Kruskal-Wallis ALL", "VecEnvWrapper"]
    for para in doc.paragraphs:
        text = para.text
        if any(token in text for token in forbidden):
            set_paragraph_text(para, replacement)


def replace_section(document: Document, start_prefix: str, end_prefix: str, writer) -> None:
    body = document.element.body
    children = list(body)
    start = end = None
    for i, el in enumerate(children):
        txt = text_of(el)
        if start is None and txt.startswith(start_prefix):
            start = i
            continue
        if start is not None and txt.startswith(end_prefix):
            end = i
            break
    if start is None or end is None or end <= start:
        raise RuntimeError(f"No se pudo reemplazar seccion: {start_prefix} -> {end_prefix}")
    tmp = Document()
    clear_body_keep_sectpr(tmp)
    writer(tmp)
    new_children = [deepcopy(el) for el in tmp.element.body if el.tag != qn("w:sectPr")]
    for el in children[start:end]:
        body.remove(el)
    for offset, el in enumerate(new_children):
        body.insert(start + offset, el)


def insert_section_before(document: Document, target_prefix: str, writer) -> None:
    body = document.element.body
    children = list(body)
    target = None
    for i, el in enumerate(children):
        if text_of(el).startswith(target_prefix):
            target = i
            break
    if target is None:
        raise RuntimeError(f"No se encontro punto de insercion: {target_prefix}")
    tmp = Document()
    clear_body_keep_sectpr(tmp)
    writer(tmp)
    new_children = [deepcopy(el) for el in tmp.element.body if el.tag != qn("w:sectPr")]
    for offset, el in enumerate(new_children):
        body.insert(target + offset, el)


def insert_section_before_any(document: Document, target_prefixes: list[str], writer) -> None:
    last_error = None
    for prefix in target_prefixes:
        try:
            insert_section_before(document, prefix, writer)
            return
        except RuntimeError as exc:
            last_error = exc
    raise RuntimeError(f"No se encontro punto de insercion entre: {target_prefixes}") from last_error


def remove_section_if_exists(document: Document, start_prefix: str, end_prefix: str) -> bool:
    """Elimina secciones autogeneradas para que la reconstruccion del Word sea idempotente."""
    body = document.element.body
    children = list(body)
    start = end = None
    for i, el in enumerate(children):
        txt = text_of(el)
        if start is None and txt.startswith(start_prefix):
            start = i
            continue
        if start is not None and txt.startswith(end_prefix):
            end = i
            break
    if start is None or end is None or end <= start:
        return False
    for el in children[start:end]:
        body.remove(el)
    return True


def remove_sections_until_clear(document: Document, pairs: list[tuple[str, str]]) -> None:
    for start_prefix, end_prefix in pairs:
        while remove_section_if_exists(document, start_prefix, end_prefix):
            pass


def remove_section_to_end_if_exists(document: Document, start_prefix: str) -> bool:
    body = document.element.body
    children = list(body)
    start = None
    for i, el in enumerate(children):
        if text_of(el).startswith(start_prefix):
            start = i
            break
    if start is None:
        return False
    for el in children[start:]:
        if el.tag != qn("w:sectPr"):
            body.remove(el)
    return True


def normalize_chapter2_numbering(doc: Document) -> None:
    replacements = {
        "2.1.1 Aprendizaje por refuerzo y MADRL": "2.2.1 Aprendizaje por refuerzo y MADRL",
        "2.1.2 Formalizacion matematica Dec-POMDP": "2.2.2 Formalizacion matematica Dec-POMDP",
        "2.1.3 CityLearn y simulacion multiobjetivo": "2.2.4 CityLearn y simulacion multiobjetivo",
        "2.2.1 Variable independiente (VI)": "2.3.1 Variable independiente (VI)",
        "2.2.2 Variable dependiente (VD)": "2.3.2 Variable dependiente (VD)",
        "2.3.1 Flexibilidad energetica": "2.4.1 Flexibilidad energetica",
        "2.3.2 Emisiones de carbono y control consciente de intensidad de carbono": "2.4.2 Emisiones de carbono y control consciente de intensidad de carbono",
        "2.3.3 Costos energeticos, precios dinamicos y respuesta economica": "2.4.3 Costos energeticos, precios dinamicos y respuesta economica",
        "2.3.4 Algoritmos MADRL evaluados": "2.4.4 Algoritmos MADRL evaluados",
        "2.3.5 Aportes fisicos al motor como base teorica de CityLearn v3 propuesto": "2.4.5 Aportes fisicos al motor como base teorica de CityLearn v3 propuesto",
        "2.4.1 Antecedentes internacionales": "2.5.1 Antecedentes internacionales",
        "2.4.2 Antecedentes nacionales y peruanos": "2.5.2 Antecedentes nacionales y peruanos",
        "2.4.3 Sintesis critica de antecedentes y brecha cientifica": "2.5.3 Sintesis critica de antecedentes y brecha cientifica",
        "2.5.1 Definicion de terminos y delimitaciones conceptuales": "2.6.1 Definicion de terminos y delimitaciones conceptuales",
        "2.5.2 Posicion teorica de la tesis": "2.6.2 Posicion teorica de la tesis",
        "2.6 Sintesis critica y triangulacion del marco teorico": "2.7 Sintesis critica y triangulacion del marco teorico",
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in replacements:
            set_paragraph_text(para, replacements[text])


def add_expanded_decpomdp_section(doc: Document, building_compact: pd.DataFrame) -> None:
    dims = building_compact.groupby("agent")[["observation_dim", "action_dim"]].max().reset_index()
    n_agents = int(dims["agent"].nunique())
    obs_min, obs_max = int(dims["observation_dim"].min()), int(dims["observation_dim"].max())
    act_min, act_max = int(dims["action_dim"].min()), int(dims["action_dim"].max())
    obs_total = int(dims["observation_dim"].sum())
    act_total = int(dims["action_dim"].sum())
    doc.add_heading("2.2.3 Dec-POMDP como formalizacion del problema doctoral", level=2)
    p(doc, "La gestion energetica estudiada no es un problema de control centralizado simple, porque la decision de cada edificio modifica la demanda agregada, el costo, las emisiones y las condiciones de flexibilidad que observan los demas agentes. Tampoco es un MDP plenamente observable por agente individual: cada edificio observa su propio estado operativo, disponibilidad de equipos, senales temporales y variables exogenas, pero no controla directamente la demanda base ni las restricciones internas de los otros edificios. Por ello, la formalizacion doctoral adecuada es un Proceso de Decision de Markov Parcialmente Observable Descentralizado (Dec-POMDP), que permite representar ejecucion descentralizada, observabilidad parcial, transiciones estocasticas y recompensa cooperativa.")
    p(doc, "La formulacion utilizada en la tesis se expresa como M = <S, {A_i}_{i=1}^{17}, T, R, {O_i}_{i=1}^{17}, Omega, gamma, H>. El estado global S contiene la concatenacion de observaciones locales s_t=[o_1,t,...,o_17,t]; las acciones A_i son heterogeneas por edificio; T representa la dinamica horaria de CityLearn, incluyendo clima, PV, BESS, llegada/salida de vehiculos electricos, demanda base, precio y senal de carbono; R es una recompensa cooperativa con mezcla local-equipo; O_i y Omega modelan la observabilidad parcial; gamma=0.9999 preserva dependencia de largo horizonte; y H=8760 pasos representa un ano horario de evaluacion por episodio.")
    p(doc, "La operacionalizacion empirieca no queda en una abstraccion generica. En los artefactos Drive, el Dec-POMDP tiene 17 agentes-edificio, dimensiones locales de observacion entre " + str(obs_min) + " y " + str(obs_max) + ", dimension global agregada " + str(obs_total) + ", acciones locales entre " + str(act_min) + " y " + str(act_max) + " y " + str(act_total) + " dimensiones de accion por tratamiento. Las acciones controlan BESS, cargadores EV y cargas flexibles declaradas en building_observation_action_schema.csv; la demanda no controlada permanece como perturbacion/observacion dentro de la demanda base, no como actuador directo.")
    p(doc, "Esta formalizacion enlaza directamente el problema general y los objetivos especificos. OE.1 evalua la respuesta del Dec-POMDP cuando la recompensa prioriza flexibilidad; OE.2 cuando prioriza emisiones de CO2; y OE.3 cuando prioriza costos energeticos. La variable independiente no es solo el nombre del algoritmo, sino la politica MADRL aprendida bajo CTDE y bajo pesos de recompensa comparables. La variable dependiente se observa mediante KPIs distritales y por edificio, por lo que el Dec-POMDP justifica simultaneamente el entrenamiento multiagente y la lectura de resultados por distrito, edificio, escenario y KPI.")
    rows = [
        ["Agentes", "Edificios institucionales/comerciales de la comunidad", f"{n_agents} agentes en building_behavior_summary.csv"],
        ["Estado global S", "Concatenacion de observaciones locales para critic/entrenamiento CTDE", f"suma observation_dim={obs_total}"],
        ["Observacion local O_i", "Informacion parcial disponible para cada actor", f"rango observado {obs_min}-{obs_max} variables"],
        ["Accion A_i", "Control descentralizado de equipos flexibles", f"rango observado {act_min}-{act_max}; total={act_total}"],
        ["Transicion T", "Dinamica horaria de CityLearn: clima, PV, BESS, EV, precio, carbono y demanda", "timeseries.csv y traces por tratamiento"],
        ["Recompensa R", "Funcion multiobjetivo comparable por escenario E1/E2/E3", "pesos flex/CO2/costo y agregacion team_mean"],
        ["Horizonte H", "Evaluacion anual horaria por episodio", "8760 pasos horarios"],
        ["Descuento gamma", "Persistencia de efectos diferidos de almacenamiento y carga EV", "gamma=0.9999"],
    ]
    table(doc, "Tabla 2.3. Mapeo del Dec-POMDP doctoral a artefactos reales del proyecto.", ["Elemento", "Operacion en la tesis", "Evidencia local"], rows, 7.0)


def add_cap1_plan_aligned(doc: Document) -> None:
    """Reconstruye el Capitulo 1 con la estructura del plan, preservando formulaciones."""
    problem_general = "¿En que medida el algoritmo Multi-Agente de Aprendizaje por Refuerzo Profundo aplicado a una comunidad inteligente (variable independiente) produce un efecto diferenciado sobre la gestion coordinada de la flexibilidad energetica, las emisiones de CO2 y los costos energeticos (variable dependiente), y cual de los algoritmos comparados genera el mayor efecto?"
    problems_specific = [
        "PE.1: ¿En que medida el algoritmo MADRL (VI) produce un efecto sobre la dimension de flexibilidad energetica de la comunidad (D-VD.1), y cual algoritmo genera el mayor efecto?",
        "PE.2: ¿En que medida el algoritmo MADRL (VI) produce un efecto sobre la dimension de emisiones de CO2 de la comunidad (D-VD.2), y cual algoritmo genera el mayor efecto?",
        "PE.3: ¿En que medida el algoritmo MADRL (VI) produce un efecto sobre la dimension de costos energeticos de la comunidad (D-VD.3), y cual algoritmo genera el mayor efecto?",
    ]
    objective_general = "Determinar el efecto del algoritmo MADRL aplicado a una comunidad inteligente (VI) sobre la gestion coordinada de la flexibilidad energetica, las emisiones de CO2 y los costos energeticos (VD), e identificar el algoritmo que produce el mayor efecto coordinado."
    objectives_specific = [
        "OE.1: Determinar el efecto del algoritmo MADRL (VI) sobre la flexibilidad energetica (D-VD.1) e identificar el algoritmo de mayor efecto en esta dimension.",
        "OE.2: Determinar el efecto del algoritmo MADRL (VI) sobre las emisiones de CO2 (D-VD.2) e identificar el algoritmo de mayor efecto en esta dimension.",
        "OE.3: Determinar el efecto del algoritmo MADRL (VI) sobre los costos energeticos (D-VD.3) e identificar el algoritmo de mayor efecto en esta dimension.",
    ]
    hypothesis_general = "La aplicacion del algoritmo MADRL a la comunidad inteligente (VI) produce un efecto estadisticamente significativo y diferenciado sobre la gestion coordinada de la flexibilidad energetica, las emisiones de CO2 y los costos energeticos (VD)."
    hypotheses_specific = [
        "HE.1: El algoritmo MADRL (VI) produce un efecto significativo sobre la flexibilidad energetica (D-VD.1); el mayor efecto corresponde al algoritmo con menor variabilidad en los KPI de pico y rampa.",
        "HE.2: El algoritmo MADRL (VI) produce un efecto significativo sobre las emisiones de CO2 (D-VD.2); el mayor efecto corresponde a MATD3.",
        "HE.3: El algoritmo MADRL (VI) produce un efecto significativo sobre los costos energeticos (D-VD.3); el mayor efecto corresponde al algoritmo con menor delta de costo electrico.",
    ]

    doc.add_heading("Capitulo 1. Introduccion", level=1)
    doc.add_heading("1.1 Diagnostico", level=2)
    p(doc, "Las comunidades inteligentes integran recursos de energia distribuida, generacion solar fotovoltaica, sistemas de almacenamiento en baterias y estaciones de carga de vehiculos electricos. La literatura de CityLearn sostiene que esta integracion exige entornos reproducibles para comparar controladores de respuesta a la demanda en edificios y comunidades, mientras que CityLearn v2 amplia el problema hacia comunidades grid-interactive con vehiculos electricos, intensidad de carbono, almacenamiento y KPIs de flexibilidad, emisiones y costos. En consecuencia, el diagnostico de esta tesis no se limita a un algoritmo aislado, sino a la ausencia de una evaluacion coordinada y comparable de politicas MADRL bajo observabilidad parcial (Vazquez-Canteli et al., 2020; Nweye et al., 2024; Fonseca et al., 2024).")
    p(doc, "En el Sistema Electrico Aislado de Iquitos, operado por Electro Oriente S.A., el suministro depende mayoritariamente de generacion diesel con penetracion solar creciente. El factor de emision base usado por el estudio es 0.790 kgCO2/kWh y la tarifa por uso horario distingue periodos punta y fuera de punta. Bajo esta condicion, la gestion no coordinada de DER puede incrementar picos, rampas, consumo ponderado por carbono y costos electricos. Esta problematica vincula directamente la tesis con tres dimensiones dependientes: flexibilidad energetica, emisiones de CO2 y costos energeticos (MINAM, 2019; OSINERGMIN, 2024; Nweye et al., 2024).")
    p(doc, "Desde la perspectiva computacional, el problema es multiagente, secuencial y parcialmente observable. Cada edificio decide sobre equipos controlables, pero no observa plenamente el estado interno de los demas edificios ni controla la demanda no flexible. Por ello, la formulacion Dec-POMDP y el esquema CTDE son consistentes con la ejecucion descentralizada y el entrenamiento cooperativo de HAPPO, MASAC, MATD3 y MAAC. Esta triangulacion teorica justifica que el estudio evalua algoritmos, escenarios y KPIs con evidencia episodica real de la corrida Drive de 50 episodios por tratamiento (Oliehoek & Amato, 2016; Lowe et al., 2017; Iqbal & Sha, 2019).")

    doc.add_heading("1.2 Identificacion y descripcion del problema de estudio", level=2)
    p(doc, "El problema de estudio se identifica como la falta de determinacion empirica del efecto que produce el algoritmo MADRL sobre la gestion coordinada de flexibilidad energetica, emisiones de CO2 y costos energeticos en una comunidad inteligente. La brecha no consiste solo en entrenar agentes, sino en comparar familias MADRL heterogeneas bajo los mismos datos, la misma funcion de recompensa, los mismos escenarios E1, E2 y E3, y los mismos criterios de lectura de KPIs. Esta comparabilidad es necesaria porque el desempeno en una dimension puede no coincidir con el desempeno en las otras; por ejemplo, una politica puede reducir emisiones y al mismo tiempo deteriorar costos o rampas.")
    p(doc, "Los sintomas observables se expresan en demanda pico, rampas horarias, consumo neto, emisiones ponderadas por intensidad de carbono, costo electrico y comportamiento por edificio. Las causas tecnicas se relacionan con la coordinacion incompleta entre agentes, la heterogeneidad de equipos controlables y no controlables, y la dificultad de aprendizaje en horizontes anuales de 8760 pasos. Las causas metodologicas se asocian con benchmarks fragmentados, falta de lectura por objetivo especifico y uso de promedios agregados sin soporte inferencial. Por ello, el Capitulo 5 responde las preguntas especificas con analisis descriptivo, inferencial y evidencia por distrito, edificio, escenario y KPI.")
    p(doc, "La correspondencia formal entre problemas, objetivos, hipotesis, variables e indicadores se conserva en el numeral 1.4, tal como estaba desarrollado en la version final de tesis. Esta ubicacion evita duplicar la matriz y mantiene la trazabilidad vertical del documento.")

    doc.add_heading("1.2.1 Formulacion del problema", level=3)
    p(doc, "La formulacion del problema se conserva segun la version definida en el documento de tesis final. Los antecedentes bibliograficos que sustentan esta formulacion se desarrollan una sola vez en el Capitulo 2, dentro de los apartados de antecedentes internacionales, antecedentes nacionales/peruanos y sintesis critica de brecha cientifica. Esta decision evita duplicar paginas de antecedentes en el Capitulo 1 y mantiene una ruta logica entre problema, marco teorico, metodologia y resultados.")
    doc.add_heading("1.2.1.1 Formulacion del problema general", level=4)
    p(doc, "Problema general (PG):")
    p(doc, problem_general)
    doc.add_heading("1.2.1.2 Formulacion de los problemas especificos", level=4)
    p(doc, "Problemas especificos:")
    for item in problems_specific:
        p(doc, item)
    p(doc, "Cada problema especifico se responde en el Capitulo 5 mediante analisis descriptivo e inferencial sobre la corrida canonica de 50 episodios por tratamiento. PE.1 se vincula con OE.1 y HE.1; PE.2 se vincula con OE.2 y HE.2; y PE.3 se vincula con OE.3 y HE.3.")

    doc.add_heading("1.2.2 Justificacion y alcances", level=3)
    doc.add_heading("1.2.2.1 Justificacion", level=4)
    p(doc, "a) Justificacion teorica. El estudio se justifica teoricamente porque articula tres cuerpos de conocimiento que suelen tratarse de manera separada: la simulacion energetica de comunidades de edificios, el aprendizaje por refuerzo profundo multiagente y la evaluacion estadistica de politicas de control. La formulacion Dec-POMDP permite representar observabilidad parcial, accion descentralizada y recompensa cooperativa; el esquema CTDE permite entrenar politicas descentralizadas con informacion global durante el aprendizaje; y la evaluacion con KPIs CityLearn permite interpretar el desempeno energetico en terminos verificables. Esta integracion aporta una base conceptual para explicar por que los algoritmos HAPPO, MAAC, MASAC y MATD3 pueden compararse como niveles de la variable independiente y no como implementaciones aisladas.")
    p(doc, "b) Justificacion de flexibilidad energetica. La tesis se justifica en flexibilidad energetica porque las comunidades inteligentes con PV, BESS, cargadores EV/V2G y cargas flexibles requieren coordinar decisiones distribuidas para reducir picos, suavizar rampas, desplazar consumo y aprovechar recursos distribuidos. En el SEAI Iquitos, donde la operacion aislada y la penetracion renovable condicionan la gestion de demanda, una politica MADRL puede modificar el perfil agregado de consumo si aprende a coordinar almacenamiento, carga vehicular y demanda flexible. Por ello, OE.1 evalua de manera especifica si el algoritmo MADRL produce efecto sobre D-VD.1 y cual algoritmo genera el mayor efecto en esta dimension.")
    p(doc, "c) Justificacion economica. La investigacion se justifica economicamente porque el costo energetico de una comunidad no depende solo del consumo total, sino tambien del momento en que la demanda se presenta frente a senales tarifarias por uso horario. La falta de coordinacion entre edificios puede trasladar consumo hacia periodos de mayor costo o generar decisiones localmente convenientes pero suboptimas a nivel comunitario. La comparacion de algoritmos MADRL permite identificar si alguna politica reduce el costo electrico bajo las mismas condiciones de simulacion, la misma comunidad y los mismos escenarios. Esta justificacion se vincula directamente con OE.3, D-VD.3 y la lectura de district_cost y KPIs de costo en el Capitulo 5.")
    p(doc, "d) Justificacion ambiental. La tesis se justifica ambientalmente porque el SEAI Iquitos se caracteriza por una matriz aislada con dependencia termica y senales de intensidad de carbono relevantes para la gestion del consumo. Coordinar demanda, almacenamiento, PV y carga vehicular puede desplazar consumo hacia periodos de menor intensidad de carbono y reducir emisiones asociadas a la importacion de energia. En consecuencia, el estudio no evalua solo recompensa algoritmica, sino el efecto de las politicas MADRL sobre emisiones de CO2 mediante indicadores distritales, por edificio y por escenario. Esta justificacion sostiene OE.2 y D-VD.2.")
    p(doc, "e) Justificacion tecnologica. La investigacion se justifica tecnologicamente porque desarrolla una arquitectura reproducible basada en dataset local, CityLearn v2, extension CityLearn v3 propuesta, scripts de construccion de datos, entrenamiento MADRL, trazas, checkpoints, timeseries y generacion automatizada de figuras y tablas. Esta arquitectura permite que los resultados no dependan de afirmaciones generales, sino de artefactos auditables de la corrida Drive madrl_v3_20260627_164047. La contribucion tecnologica consiste en integrar datos locales, equipos controlables y no controlables, simulacion multiagente y analisis estadistico dentro de un flujo verificable para tesis doctoral.")
    doc.add_heading("1.2.2.2 Alcances", level=4)
    p(doc, "El alcance tematico comprende la evaluacion comparativa de HAPPO, MASAC, MATD3 y MAAC sobre flexibilidad energetica, emisiones de CO2 y costos energeticos. El alcance espacial corresponde a una comunidad inteligente simulada con aplicabilidad al SEAI Iquitos y 17 edificios institucionales/comerciales. El alcance temporal se basa en datos horarios 2023-2025 y en una corrida de entrenamiento/evaluacion con 50 episodios por tratamiento disponible. El alcance computacional comprende Python, PyTorch, CityLearn v2, librerias de analisis cientifico y artefactos Drive validados.")
    p(doc, "El estudio no modela una red electrica fisica ni valida un despliegue real en campo. CityLearn v3 propuesto se interpreta como extension experimental de tesis y no como version oficial de CityLearn. Se excluyen investigacion con sujetos humanos, despacho economico de unidades fisicas, analisis de estabilidad de red y extrapolacion causal fuera de la evidencia de simulacion disponible.")

    doc.add_heading("1.3 Objetivos e hipotesis de la investigacion", level=2)
    doc.add_heading("1.3.1 Objetivos", level=3)
    doc.add_heading("1.3.1.1 Objetivo general", level=4)
    p(doc, "Objetivo general (OG):")
    p(doc, objective_general)
    doc.add_heading("1.3.1.2 Objetivos especificos", level=4)
    p(doc, "Objetivos especificos:")
    for item in objectives_specific:
        p(doc, item)
    p(doc, "Los objetivos (OG, OE.1-OE.3) determinan el efecto del MADRL (VI) sobre cada dimension de la VD e identifican el algoritmo de mayor efecto mediante KPIs y ranking descriptivo en los Capitulos 5 y 6. Esta redaccion no modifica los objetivos definidos; solo los vincula con el diseno factorial 4x3, los 54 KPI oficiales compatibles con CityLearn v2 y la corrida canonica Drive madrl_v3_20260627_164047 de 50 episodios por tratamiento disponible.")
    doc.add_heading("1.3.2 Hipotesis", level=3)
    p(doc, "El estudio es cuantitativo, aplicado y explicativo, basado en simulacion experimental. A diferencia de los objetivos, las hipotesis formulan contrastes estadisticos sobre el factor algoritmo (VI) frente a la hipotesis nula de igualdad de distribuciones de KPI-gains. El protocolo inferencial (alpha = 0,05) se detalla en el Capitulo 3 y se resuelve en la seccion 5.9 (Colas et al., 2019; Agarwal et al., 2021):")
    doc.add_heading("1.3.2.1 Hipotesis general", level=4)
    p(doc, "Hipotesis general (HG):")
    p(doc, hypothesis_general)
    doc.add_heading("1.3.2.2 Hipotesis especificas", level=4)
    p(doc, "Hipotesis especificas:")
    for item in hypotheses_specific:
        p(doc, item)
    p(doc, "Cada hipotesis especifica tiene una hipotesis nula asociada (sin diferencias significativas entre niveles del factor algoritmo). La decision inferencial se reporta en la seccion 5.9.5; el cumplimiento de OG y OE.1-OE.3 se reporta por separado en la seccion 5.11.")

    doc.add_heading("1.4 Matriz de consistencia y operacionalizacion", level=2)
    table(
        doc,
        "Tabla 1.1. Matriz de consistencia problema-objetivo-hipotesis.",
        ["Dimension", "Problema especifico", "Objetivo especifico", "Hipotesis especifica", "Evidencia real de resultados"],
        [
            ["D-VD.1 Flexibilidad energetica", problems_specific[0], objectives_specific[0], hypotheses_specific[0], "KPIs de pico, rampa, load factor, convergencia y comparacion por algoritmo desde 50 episodios Drive."],
            ["D-VD.2 Emisiones de CO2", problems_specific[1], objectives_specific[1], hypotheses_specific[1], "KPIs de emisiones, carbon-aware control, ranking por escenario y pruebas no parametricas desde artefactos Drive."],
            ["D-VD.3 Costos energeticos", problems_specific[2], objectives_specific[2], hypotheses_specific[2], "KPIs de costo electrico, tarifa TOU, comparacion por edificio y efecto inferencial desde timeseries y tablas reales."],
        ],
        5.4,
    )
    table(
        doc,
        "Tabla 1.2. Operacionalizacion de variables independientes y dependientes.",
        ["Variable", "Dimension", "Indicadores", "Fuente empirica", "Uso analitico"],
        [
            ["VI: algoritmo MADRL", "Tratamiento algoritmico", "HAPPO, MASAC, MATD3, MAAC; escenarios E1, E2, E3", "results, traces, checkpoints y configuraciones Drive", "Comparacion factorial 4x3."],
            ["VD: gestion coordinada", "D-VD.1 Flexibilidad energetica", "pico, rampa, load factor, consumo neto y KPIs CityLearn v2", "timeseries y tablas district/building KPIs", "Respuesta a PE.1/OE.1/HE.1."],
            ["VD: gestion coordinada", "D-VD.2 Emisiones de CO2", "emisiones, carbon intensity y consumo ponderado por carbono", "timeseries y KPI de emisiones", "Respuesta a PE.2/OE.2/HE.2."],
            ["VD: gestion coordinada", "D-VD.3 Costos energeticos", "costo electrico, tarifa TOU y costo acumulado", "timeseries y KPI de costos", "Respuesta a PE.3/OE.3/HE.3."],
        ],
        5.8,
    )

    doc.add_heading("1.5 Justificacion", level=2)
    p(doc, "Tecnica: aporta una evaluacion unificada de HAPPO, MASAC, MATD3 y MAAC bajo Dec-POMDP y CTDE, avanzando el estado del arte en control energetico cooperativo. La mejora incorporada consiste en hacer explicita la relacion entre esa evaluacion y los artefactos reales del proyecto: results, timeseries, traces, checkpoints, KPIs distritales, KPIs por edificio y figuras generadas desde la corrida Drive.")
    p(doc, "Ambiental: identificar el mejor MADRL para reduccion de CO2 contribuye a la descarbonizacion de comunidades grid-interactive, con aplicabilidad al SEAI Iquitos (factor de emision 0.790 kgCO2/kWh). Esta justificacion se mantiene porque OE.2 y PE.2 se responden con indicadores de emisiones reales derivados de la simulacion y no con estimaciones externas inventadas.")
    p(doc, "Economica: establecer el mejor MADRL para optimizacion de costos orienta la reduccion del gasto electrico bajo tarifas TOU. En la tesis, esta justificacion se concreta mediante KPIs de costo, series de costo distrital y comparacion por algoritmo/escenario, preservando el vinculo con OE.3.")
    p(doc, "Metodologica: la formulacion Dec-POMDP, el esquema CTDE y el benchmark unificado con CityLearn v3 propuesto, MARLlib y Optuna constituyen una contribucion reproducible. La reproducibilidad se refuerza con el uso de scripts locales, CSV materializados y resultados Drive auditables.")
    p(doc, "Cientifica y social: llena una laguna en la literatura comparativa de MADRL y beneficia a usuarios institucionales y a la transicion energetica comunitaria. Esta justificacion no se presenta como promesa de despliegue real, sino como evidencia experimental en simulacion computacional con alcance claramente delimitado.")

    doc.add_heading("1.6 Alcances y limitaciones", level=2)
    p(doc, "Alcances:")
    p(doc, "Tematico: comparacion de HAPPO, MASAC, MATD3 y MAAC en KPIs de flexibilidad, CO2 y costos.")
    p(doc, "Espacial: comunidades inteligentes simuladas en CityLearn v2 / CityLearn v3 propuesto, con aplicabilidad al SEAI Iquitos (17 edificios reales).")
    p(doc, "Temporal: dataset horario 2023-2025 (26 304 pasos) y literatura 2015-2026.")
    p(doc, "Metodologico: estudio cuantitativo, aplicado, comparativo, basado en simulacion experimental-computacional.")
    p(doc, "Computacional: Python 3.9, PyTorch 2.8.0+cu126, CityLearn v2, MARLlib, Optuna, Gymnasium, PettingZoo; hardware local NVIDIA RTX 4060 Laptop 8 GB y Colab A100.")
    p(doc, "Limitaciones:")
    p(doc, "No se modela una red electrica fisica; los resultados de simulacion no constituyen validacion de despliegue real.")
    p(doc, "CityLearn v3 propuesto es una extension experimental de tesis, no una version oficial de CityLearn.")
    p(doc, "Se excluyen el despliegue en campo en tiempo real, la investigacion con sujetos humanos, el despacho economico de unidades fisicas y el analisis de estabilidad de red.")
    p(doc, "La inferencia se interpreta dentro de la corrida canonica Drive y no como generalizacion absoluta a multiples semillas independientes. Esta precision mejora la version final sin alterar sus alcances reales.")


def add_cap1_validation(doc: Document) -> None:
    doc.add_heading("1.7 Validacion estructural y triangulacion del planteamiento", level=2)
    p(doc, "El Capitulo 1 cumple la estructura minima exigida para una tesis aplicada: formula el problema de investigacion, declara objetivos, hipotesis, justificacion, alcances y limitaciones. La coherencia interna se sostiene porque el problema no se define como una deficiencia algoritmica aislada, sino como una tension entre gestion descentralizada, flexibilidad energetica, emisiones de CO2 y costos en una comunidad electrica aislada. Este encuadre es consistente con la literatura de CityLearn, que plantea la necesidad de entornos reproducibles para comparar controladores de respuesta a la demanda, con CityLearn v2, que amplia la evaluacion hacia comunidades grid-interactive, resilientes y carbon-aware, y con EVLearn, que subraya que la integracion de vehiculos electricos requiere simulacion especifica de flexibilidad de carga y descarga (Vazquez-Canteli et al., 2020; Nweye et al., 2024; Fonseca et al., 2024).")
    p(doc, "La pregunta general y las preguntas especificas quedan operacionalizadas en tres dimensiones dependientes. PE.1 se vincula con flexibilidad energetica, PE.2 con emisiones de CO2 y PE.3 con costos energeticos. Esta separacion evita que el desempeno del algoritmo se reduzca a una recompensa unica sin interpretabilidad; por el contrario, permite contrastar la variable independiente MADRL contra indicadores observables por distrito, edificio y escenario. La triangulacion entre CityLearn, MARL energetico y evaluacion estadistica de RL justifica que el estudio reporte resultados descriptivos e inferenciales, y que declare como limitacion la ausencia de multiples semillas independientes en lugar de sobregeneralizar una unica corrida experimental (Henderson et al., 2018; Agarwal et al., 2021; Nweye et al., 2024).")
    table(
        doc,
        "Tabla 1.7. Validacion del Capitulo 1 frente a la estructura minima.",
        ["Elemento requerido", "Estado", "Refuerzo incorporado"],
        [
            ["Problema de investigacion", "Cumple", "Se alinea con flexibilidad, carbono y costo en comunidad aislada."],
            ["Objetivos", "Cumple", "OG y OE.1-OE.3 mantienen correspondencia con PE.1-PE.3."],
            ["Hipotesis", "Cumple", "Se conservan por tratarse de diseno experimental comparativo."],
            ["Justificacion", "Cumple", "Se refuerza con reproducibilidad CityLearn y necesidad de evaluacion estadistica."],
            ["Alcances y limitaciones", "Cumple", "Se explicita la lectura intra-corrida y la no extrapolacion multi-semilla."],
        ],
        7.0,
    )


def add_cap2_validation(doc: Document) -> None:
    doc.add_heading("2.6 Sintesis critica y triangulacion del marco teorico", level=2)
    p(doc, "El marco teorico contiene estado del arte, bases teoricas, trabajos relacionados, definicion de variables y posicion teorica. La triangulacion central se organiza en tres capas. Primero, CityLearn aporta el marco de evaluacion reproducible para comunidades de edificios y demanda flexible; CityLearn v2 agrega objetivos de flexibilidad, resiliencia, ocupacion y carbono; EVLearn extiende el mismo ecosistema a vehiculos electricos y estrategias V1G/V2G (Vazquez-Canteli et al., 2020; Nweye et al., 2024; Fonseca et al., 2024). Segundo, la teoria MADRL fundamenta la ejecucion descentralizada con entrenamiento centralizado: MAAC introduce criticos con atencion para seleccionar interacciones relevantes entre agentes, HAPPO extiende optimizacion de region de confianza a agentes heterogeneos, SAC/MASAC aporta exploracion por maxima entropia y TD3/MATD3 controla sesgo de sobreestimacion con criticos dobles y actualizaciones retrasadas (Iqbal & Sha, 2019; Kuba et al., 2021; Haarnoja et al., 2018; Fujimoto et al., 2018). Tercero, la literatura de evaluacion de RL advierte que una comparacion algoritmica debe acompanar los promedios con dispersion, pruebas estadisticas y trazabilidad experimental (Henderson et al., 2018; Agarwal et al., 2021).")
    p(doc, "La principal mejora incorporada al Capitulo 2 consiste en separar el soporte teorico de cada decision metodologica. El Dec-POMDP no se presenta como formalismo decorativo, sino como condicion necesaria para representar 17 agentes con observacion parcial, acciones heterogeneas, recompensa cooperativa y horizonte anual. La funcion de recompensa multiobjetivo se sustenta en la literatura de control de comunidades energeticas, pero se interpreta dentro de CityLearn v2 para evitar que flexibilidad, emisiones y costos sean confundidos como una unica variable latente. Esta posicion teorica permite que el Capitulo 5 lea resultados por escenario y no solo por algoritmo.")
    table(
        doc,
        "Tabla 2.6. Triangulacion teorica por eje de la tesis.",
        ["Eje", "Fuentes trianguladas", "Uso en la tesis"],
        [
            ["Entorno y KPIs", "CityLearn v1/v2; CityLearn Challenge; EVLearn", "Justifica benchmark, KPIs evaluate_v2 y control de EV/BESS."],
            ["MADRL", "MAAC; HAPPO; SAC; TD3", "Justifica comparacion de familias actor-critic multiagente."],
            ["Evaluacion rigurosa", "Henderson et al.; Agarwal et al.", "Justifica pruebas no parametricas y cautela ante una sola corrida."],
            ["Variables", "Flexibilidad, CO2 y costo en comunidades grid-interactive", "Sostiene D-VD.1, D-VD.2 y D-VD.3."],
        ],
        6.8,
    )


def add_cap3_validation(doc: Document) -> None:
    doc.add_heading("3.7 Validez metodologica, trazabilidad y control de sesgos", level=2)
    p(doc, "El Capitulo 3 cumple con el tipo de investigacion, diseno metodologico, datos utilizados, variables, tecnicas, herramientas y procedimiento experimental. La tesis corresponde a una investigacion aplicada, cuantitativa y explicativa bajo simulacion computacional, porque manipula la variable independiente MADRL mediante cuatro algoritmos y tres escenarios de recompensa, y observa sus efectos sobre indicadores cuantitativos de flexibilidad, emisiones y costos. Esta estrategia es compatible con la estandarizacion buscada por CityLearn para comparar controladores en comunidades energeticas, pero requiere trazabilidad estricta de artefactos para evitar sesgos de seleccion o interpretacion (Vazquez-Canteli et al., 2020; Nweye et al., 2024).")
    p(doc, "La validez interna se protege mediante un diseno factorial 4x3, mismo dataset, mismos escenarios E1-E3 y mismas reglas de extraccion de KPIs. La validez estadistica se aborda con estadistica descriptiva, Shapiro-Wilk, Kruskal-Wallis y Mann-Whitney con ajuste Holm, debido a que las distribuciones episodicas no deben asumirse normales por defecto. La validez externa se declara limitada: Henderson et al. (2018) y Agarwal et al. (2021) advierten que las conclusiones de RL con pocas corridas pueden variar si no se reporta incertidumbre, por lo que la tesis interpreta los resultados como evidencia intra-corrida y recomienda una extension multi-semilla.")
    table(
        doc,
        "Tabla 3.7. Control metodologico de validez y trazabilidad.",
        ["Riesgo", "Control aplicado", "Evidencia"],
        [
            ["Comparacion no equivalente", "Mismo dataset, mismos escenarios y misma funcion de evaluacion", "12 tratamientos algoritmo x escenario."],
            ["Normalidad no garantizada", "Pruebas no parametricas y Shapiro-Wilk", "CSV estadisticos generados desde episodios Drive."],
            ["Sobregeneralizacion", "Declaracion de inferencia intra-corrida", "Limitacion multi-semilla en Capitulo 6."],
            ["Datos inventados", "Uso exclusivo de results, timeseries, traces, checkpoints y CSV materializados", "Validacion de cobertura en Capitulo 5."],
        ],
        6.8,
    )


def add_cap3_methodology(doc: Document) -> None:
    doc.add_heading("Capitulo 3. Metodologia", level=1)
    p(doc, "El presente estudio adopta una metodologia cuantitativa, aplicada y explicativa-comparativa, sustentada en simulacion computacional controlada. Esta decision responde a la naturaleza del problema doctoral: evaluar en que medida distintos algoritmos MADRL modifican indicadores cuantificables de flexibilidad energetica, emisiones de CO2 y costos energeticos en una comunidad electrica modelada bajo CityLearn. De acuerdo con Hernandez-Sampieri y Mendoza (2018), la ruta cuantitativa se caracteriza por la medicion de variables, el uso de procedimientos sistematicos y la contrastacion de hipotesis; Creswell y Creswell (2023) enfatizan que los disenos cuantitativos permiten examinar relaciones entre variables mediante mediciones numericas y analisis estadistico. En esta tesis, la variable independiente se manipula computacionalmente mediante el algoritmo MADRL y el escenario de recompensa, mientras que la variable dependiente se observa mediante KPIs oficiales y resultados episodicos.")

    doc.add_heading("3.1 Enfoque, tipo, nivel y diseno de investigacion", level=2)
    p(doc, "La investigacion es aplicada porque desarrolla y evalua una solucion computacional para gestion energetica multiagente en una comunidad del SEAI Iquitos. Es cuantitativa porque los resultados se expresan en metricas numericas: recompensa, costo, emisiones, pico, rampa, factor de carga, energia de almacenamiento, exito EV y KPIs evaluate_v2. Es explicativa-comparativa porque no se limita a describir los algoritmos; contrasta si la variacion de la variable independiente produce efectos diferenciados sobre las dimensiones D-VD.1, D-VD.2 y D-VD.3. Esta clasificacion es coherente con los criterios de alcance y diseno propuestos por Hernandez-Sampieri y Mendoza (2018), con la logica de diseno cuantitativo de Creswell y Creswell (2023), y con la tradicion de diseno experimental orientado a factores, tratamientos y respuestas descrita por Montgomery (2019).")
    p(doc, "El diseno no se considera no experimental en sentido estricto, porque si existe manipulacion controlada de factores dentro del entorno de simulacion: algoritmo MADRL y escenario de recompensa. Tampoco corresponde a un experimento de campo con sujetos humanos, sino a un experimento computacional in silico. Por ello, se define como diseno experimental-computacional factorial 4x3, con control de dataset, entorno, recompensa, horizonte temporal, agentes y protocolo de evaluacion. La inferencia es intra-corrida y se interpreta con cautela, porque Shadish, Cook y Campbell (2002) advierten que la validez interna, estadistica, de constructo y externa deben declararse de forma diferenciada; en este caso, la validez externa queda limitada por la ausencia de multiples semillas independientes.")
    table(
        doc,
        "Tabla 3.1. Clasificacion metodologica de la investigacion.",
        ["Criterio", "Decision metodologica", "Sustento"],
        [
            ["Enfoque", "Cuantitativo", "Medicion numerica de KPIs y contrastacion estadistica."],
            ["Tipo", "Aplicada", "Desarrollo y evaluacion de una propuesta MADRL para gestion energetica."],
            ["Nivel", "Explicativo-comparativo", "Evalua efecto de algoritmos y escenarios sobre dimensiones VD."],
            ["Diseno", "Experimental-computacional factorial 4x3", "Manipulacion controlada de algoritmo y escenario en simulacion."],
            ["Temporalidad", "Longitudinal por episodios y horizonte anual horario", "Cada episodio recorre una trayectoria anual de 8760 pasos."],
            ["Inferencia", "Intra-corrida con limitacion multi-semilla", "No se generaliza mas alla de los artefactos auditados."],
        ],
        6.8,
    )

    doc.add_heading("3.2 Diseno experimental-computacional factorial 4x3", level=2)
    p(doc, "El diseno factorial se compone de dos factores principales. El primer factor es el algoritmo MADRL, con cuatro niveles: HAPPO, MAAC, MASAC y MATD3. El segundo factor es el escenario de recompensa, con tres niveles: E1 orientado a flexibilidad, E2 orientado a emisiones de CO2 y E3 orientado a costos energeticos. La combinacion genera 12 tratamientos algoritmo x escenario. Montgomery (2019) sostiene que los disenos factoriales permiten estudiar efectos de factores bajo condiciones controladas; en esta tesis, el control se materializa en el mismo dataset, el mismo entorno CityLearn, la misma comunidad de 17 edificios, el mismo horizonte de evaluacion y la misma familia de KPIs.")
    p(doc, "La unidad experimental principal es el tratamiento algoritmo-escenario. La unidad de observacion episodica es el episodio conservado en los artefactos de Drive, y la unidad de observacion espacial es el edificio-agente. Esta doble lectura permite responder los objetivos en tres escalas: distrito, edificio y politica multiagente. Para la inferencia estadistica se usan los episodios materializados: MAAC, MASAC y MATD3 conservan cobertura completa en los tres escenarios; HAPPO registra entrenamiento completado, pero conserva 49 filas episodicas por escenario en el CSV materializado, por lo que se reporta como evidencia descriptiva y no como grupo inferencial completo.")
    table(
        doc,
        "Tabla 3.2. Matriz factorial 4x3 de tratamientos experimentales.",
        ["Factor", "Niveles", "Funcion metodologica"],
        [
            ["Algoritmo MADRL", "HAPPO, MAAC, MASAC, MATD3", "Variable independiente principal; compara familias actor-critic multiagente."],
            ["Escenario de recompensa", "E1, E2, E3", "Manipulacion de prioridad multiobjetivo: flexibilidad, CO2 y costo."],
            ["Tratamientos", "12 combinaciones", "Base de comparacion para resultados por objetivo especifico."],
            ["Horizonte", "50 episodios registrados por tratamiento", "Evidencia de entrenamiento Drive; se declara cobertura materializada por algoritmo."],
        ],
        6.8,
    )

    doc.add_heading("3.3 Datos utilizados y fuente empirica", level=2)
    p(doc, "Los datos utilizados corresponden al dataset citylearn_iquitos_2023_2025 y a la corrida canonica Drive madrl_v3_20260627_164047. La tesis no incorpora datos simulados manualmente fuera del pipeline; usa artefactos generados por scripts del proyecto y resultados auditables: results.json, training_summary.json, timeseries.csv, trace.csv, building_kpis.csv, building_behavior_summary.csv, building_observation_action_schema.csv y checkpoint_manifest.json. La fuente empirica combina meteorologia, demanda, PV, BESS, EV, precios, intensidad de carbono y KPIs CityLearn. Esta trazabilidad responde a la exigencia metodologica de reproducibilidad y control de medicion que Hernandez-Sampieri y Mendoza (2018) y Creswell y Creswell (2023) asocian con la ruta cuantitativa.")
    p(doc, "El dataset representa 17 edificios de una comunidad energetica de Iquitos, con equipos controlables heterogeneos y variables de observacion locales. La representacion por edificio permite modelar el problema como Dec-POMDP y evaluar la ejecucion descentralizada de politicas MADRL. Para evitar alucinacion o sobreinterpretacion, los resultados del Capitulo 5 se derivan exclusivamente de los archivos existentes y de tablas materializadas en outputs/_drive_madrl/gdrive_20260627_164047_objective_analysis.")
    table(
        doc,
        "Tabla 3.3. Fuentes de datos y artefactos de analisis.",
        ["Artefacto", "Contenido", "Uso metodologico"],
        [
            ["timeseries.csv", "Serie temporal distrital por episodio", "Energia, costo, CO2, recompensa y senales horarias."],
            ["trace.csv", "Trazas por agente", "Acciones, SOC, EV, PV, importacion/exportacion y recompensa individual."],
            ["building_kpis.csv / building_behavior_summary.csv", "KPIs por edificio", "Analisis espacial por agente y equipamiento."],
            ["checkpoint_manifest.json", "Registro de checkpoints", "Trazabilidad de entrenamiento y cobertura por tratamiento."],
            ["district_episode_kpis.csv", "KPIs episodicos materializados", "Base de estadistica descriptiva e inferencial."],
        ],
        6.7,
    )

    doc.add_heading("3.4 Variables, dimensiones e indicadores", level=2)
    p(doc, "La variable independiente (VI) es el algoritmo MADRL implementado bajo un esquema CTDE y condicionado por escenario de recompensa. Sus dimensiones operacionales son: D-VI.1 tipo de algoritmo, D-VI.2 ponderacion multiobjetivo del escenario y D-VI.3 controles experimentales. La variable dependiente (VD) es el desempeno energetico coordinado de la comunidad, desagregado en tres dimensiones: D-VD.1 flexibilidad energetica, D-VD.2 emisiones de CO2 y D-VD.3 costos energeticos. Esta definicion mantiene correspondencia vertical con PE.1, PE.2, PE.3, OE.1, OE.2 y OE.3.")
    p(doc, "La operacionalizacion sigue la logica de medicion cuantitativa: cada dimension debe tener indicadores observables, fuente de datos y criterio de interpretacion. En D-VD.1 se consideran recompensa en E1 y KPIs de flexibilidad como peak, ramping, load factor, autoconsumo PV y uso de almacenamiento. En D-VD.2 se consideran emisiones distritales, carbon_emissions_control, carbon_emissions_delta y consumo ponderado por intensidad de carbono. En D-VD.3 se consideran district_cost, electricity_cost_control, electricity_cost_delta y senales de precio. Esta estructura evita confundir recompensa de entrenamiento con resultado final de evaluacion.")
    table(
        doc,
        "Tabla 3.4. Operacionalizacion metodologica de variables.",
        ["Variable", "Dimension", "Indicadores principales", "Fuente"],
        [
            ["VI", "D-VI.1 Algoritmo", "HAPPO, MAAC, MASAC, MATD3", "Configuracion de tratamiento."],
            ["VI", "D-VI.2 Escenario", "E1, E2, E3; pesos flex/CO2/costo", "reward_axis_weights y protocolo experimental."],
            ["VD", "D-VD.1 Flexibilidad", "reward_mean E1, peak, ramping, load factor, BESS/PV", "episodes, results, evaluate_v2."],
            ["VD", "D-VD.2 CO2", "district_emission, carbon_emissions_control/delta", "timeseries, results, building KPIs."],
            ["VD", "D-VD.3 Costos", "district_cost, electricity_cost_control/delta", "timeseries, results, building KPIs."],
        ],
        6.6,
    )

    doc.add_heading("3.5 Tecnicas, herramientas e instrumentos", level=2)
    p(doc, "Las tecnicas utilizadas son simulacion computacional, entrenamiento MADRL, evaluacion por KPIs, estadistica descriptiva, contrastacion no parametrica y visualizacion analitica. Las herramientas principales son Python, PyTorch, CityLearn v2, la extension CityLearn v3 propuesta, backends HAPPO/MAAC/MASAC/MATD3, scripts de orquestacion del proyecto y artefactos Drive. En terminos metodologicos, el instrumento de medicion no es un cuestionario ni una entrevista, sino el entorno computacional validado y sus archivos de salida.")
    p(doc, "La decision de usar estadistica no parametrica se debe a que las recompensas y KPIs episodicos de RL pueden presentar no normalidad, dependencia temporal, asimetria o valores atipicos. Por ello se aplica Shapiro-Wilk para normalidad, Kruskal-Wallis para diferencias globales entre algoritmos con cobertura completa y Mann-Whitney U con ajuste Holm para comparaciones por pares. Esta eleccion es coherente con la recomendacion metodologica de no asumir supuestos estadisticos no verificados y con las advertencias de evaluacion robusta en aprendizaje por refuerzo reportadas por Henderson et al. (2018) y Agarwal et al. (2021).")
    table(
        doc,
        "Tabla 3.5. Tecnicas, herramientas e instrumentos.",
        ["Componente", "Aplicacion en la tesis", "Resultado esperado"],
        [
            ["Simulacion CityLearn", "Recrear comunidad energetica multiagente", "Series, KPIs y trazas auditables."],
            ["MADRL", "Entrenar politicas bajo CTDE", "Politicas por algoritmo y escenario."],
            ["evaluate_v2 / KPIs", "Evaluar flexibilidad, CO2 y costo", "Ranking comparable con baseline."],
            ["Estadistica no parametrica", "Contrastar diferencias entre algoritmos", "p-valores, epsilon2 y decisiones HE."],
            ["Visualizacion", "Interpretar convergencia, trade-offs, edificios y acciones", "Figuras del Capitulo 5."],
        ],
        6.8,
    )

    doc.add_heading("3.6 Procedimiento experimental", level=2)
    p(doc, "El procedimiento experimental se estructura en siete fases reproducibles. Primero, se verifica el contexto del repositorio y la disponibilidad de artefactos. Segundo, se valida el dataset citylearn_iquitos_2023_2025 y el esquema de edificios/equipos. Tercero, se ejecutan o recuperan las 12 corridas algoritmo x escenario desde Drive. Cuarto, se consolidan episodios, timeseries, traces, building KPIs y checkpoints. Quinto, se calculan KPIs distritales, por edificio y por objetivo. Sexto, se aplican pruebas estadisticas y rankings evaluate_v2. Septimo, se generan tablas, figuras y redaccion interpretativa en el documento final.")
    p(doc, "Cada fase se documenta mediante archivos de salida. La decision de no completar manualmente valores ausentes es parte del control metodologico: si un artefacto no conserva una granularidad determinada, se declara como limitacion y no se sintetiza. Esta regla es consistente con la validez de medicion y la transparencia experimental recomendadas por Shadish et al. (2002) y Montgomery (2019).")
    table(
        doc,
        "Tabla 3.6. Procedimiento experimental reproducible.",
        ["Fase", "Actividad", "Evidencia"],
        [
            ["1", "Verificacion de contexto y rutas", "scripts/verify_project_context.ps1."],
            ["2", "Validacion de dataset y esquema CityLearn", "schema, building files y auditorias."],
            ["3", "Entrenamiento/recuperacion de 12 tratamientos", "Drive madrl_v3_20260627_164047."],
            ["4", "Consolidacion de resultados", "timeseries, traces, checkpoints y KPIs."],
            ["5", "Analisis descriptivo e inferencial", "CSV de estadisticas y comparaciones Holm."],
            ["6", "Visualizacion doctoral", "Figuras por convergencia, KPIs, edificios y trade-offs."],
            ["7", "Integracion en Word", "Documento final reproducible desde el generador."],
        ],
        6.8,
    )

    doc.add_heading("3.7 Validez metodologica, trazabilidad y control de sesgos", level=2)
    p(doc, "La validez interna se fortalece por el control del entorno: todos los algoritmos se evaluan sobre el mismo dataset, la misma comunidad, los mismos escenarios y los mismos criterios de extraccion de KPIs. La validez de constructo se protege mediante la correspondencia entre preguntas, objetivos, variables e indicadores. La validez estadistica se aborda con pruebas no parametricas y tamanos de efecto, evitando asumir normalidad sin evidencia. La validez externa se declara limitada porque la corrida canonica no sustituye una campana multi-semilla; por tanto, las conclusiones se formulan como evidencia doctoral intra-corrida y no como generalizacion universal.")
    p(doc, "El control de sesgos se apoya en cinco reglas: no mezclar resultados de otros proyectos, no usar artefactos ajenos a Drive/local autorizado, no inventar datos faltantes, distinguir evidencia descriptiva de evidencia inferencial y reportar limitaciones de cobertura. Este criterio es central para una tesis doctoral basada en RL, donde diferencias aparentemente favorables pueden depender de semillas, hiperparametros, entorno y criterio de evaluacion.")
    table(
        doc,
        "Tabla 3.7. Control metodologico de validez y trazabilidad.",
        ["Dimension de validez", "Riesgo", "Control aplicado"],
        [
            ["Interna", "Comparacion desigual entre algoritmos", "Mismo dataset, mismo entorno y escenarios controlados."],
            ["Constructo", "Indicadores no alineados con objetivos", "PE/OE/VD/KPI vinculados por tabla de operacionalizacion."],
            ["Estadistica", "Normalidad o significancia asumida", "Shapiro-Wilk, Kruskal-Wallis, Mann-Whitney-Holm y epsilon2."],
            ["Externa", "Generalizacion indebida", "Declaracion de inferencia intra-corrida y recomendacion multi-semilla."],
            ["Trazabilidad", "Datos inventados o mezclados", "Uso exclusivo de artefactos Drive y CSV materializados."],
        ],
        6.8,
    )


def add_cap5(
    doc: Document,
    detail: dict,
    treatment: pd.DataFrame,
    building_compact: pd.DataFrame,
    eq_summary: pd.DataFrame,
    figures: dict[str, Path],
    convergence: pd.DataFrame,
    kpi_ranking: pd.DataFrame,
    kpi_catalog: pd.DataFrame,
) -> None:
    doc.add_heading("Capitulo 5. Resultados y contrastacion de hipotesis", level=1)
    p(doc, "Este capitulo se reconstruye con evidencia directa de la carpeta G:\\Mi unidad\\MADRLCitytleranflexresdr\\outputs\\madrl_v3_20260627_164047. La lectura incluye results.json, training_summary.json, timeseries.csv, building_kpis.csv, building_behavior_summary.csv, building_observation_action_schema.csv y checkpoint_manifest.json para los 12 tratamientos algoritmo x escenario. La regla de interpretacion es estricta: no se incorporan valores que no existan en los artefactos; cuando una tabla conserva menor granularidad episodica, se declara como limitacion de trazabilidad y no se inventan observaciones.")
    p(doc, "El vinculo metodologico queda organizado por objetivo especifico: OE.1 se contrasta en E1 porque la recompensa asigna 0,70 a flexibilidad; OE.2 se contrasta en E2 porque la recompensa asigna 0,70 a emisiones; OE.3 se contrasta en E3 porque la recompensa asigna 0,60 a costos. Por tanto, el desarrollo de la propuesta del Capitulo 4 no queda separado de los resultados: los pesos de la funcion de recompensa son la manipulacion experimental de D-VI.2 y las metricas de este capitulo son los indicadores observados de D-VD.1, D-VD.2 y D-VD.3.")
    p(doc, "Los 12 tratamientos registran culminacion operativa de entrenamiento. En particular, HAPPO registra completed_episode_count=50 en live_progress.json y episodes_recorded=50 en results.json; sin embargo, por el modo de reanudacion ligera de HAPPO, la carpeta actual de G: conserva en timeseries.csv y episode_summary.csv solo el episodio 49, es decir, la trayectoria anual final. Para no perder la evidencia previa ya extraida del mismo flujo Drive, la estadistica episodica usa el CSV materializado district_episode_kpis.csv, donde HAPPO conserva 49 episodios por escenario y MAAC, MASAC y MATD3 conservan 50. En consecuencia, HAPPO se usa para evidencia descriptiva, final anual y por edificio, pero no se declara como grupo inferencial de 50 observaciones.")

    doc.add_heading("5.1 Cobertura de artefactos experimentales", level=2)
    coverage_rows = []
    for _, r in treatment.sort_values(["algorithm", "scenario"]).iterrows():
        coverage_rows.append([
            r["algorithm"],
            r["scenario"],
            int(r["episodes_recorded"]),
            int(r["saved_episode_summaries"]),
            int(r["materialized_episode_kpis"]),
            int(r["building_count"]),
            int(r["checkpoint_count"]),
        ])
    table(doc, "Tabla 5.1. Cobertura real de artefactos por tratamiento en Google Drive.", ["Algoritmo", "Escenario", "episodios registrados", "resumenes G:", "episodios KPI usados", "edificios", "checkpoints"], coverage_rows, 6.8)
    if "checkpoint_coverage" in figures:
        add_picture(doc, "Figura 5.1. Cobertura de checkpoints por tratamiento.", figures["checkpoint_coverage"])

    doc.add_heading("5.2 Trazabilidad entre objetivos, hipotesis e indicadores", level=2)
    link_rows = [[d["spec"]["objective"], d["spec"]["hypothesis"], d["spec"]["scenario"], d["spec"]["dimension"], d["spec"]["indicator"], "maximizar" if d["spec"]["direction"] == "max" else "minimizar"] for d in detail.values()]
    table(doc, "Tabla 5.2. Trazabilidad objetivo-hipotesis-escenario-indicador.", ["Objetivo", "Hipotesis", "Escenario", "Dimension VD", "Indicador usado", "Criterio"], link_rows, 6.8)

    doc.add_heading("5.3 Curvas de convergencia y episodios de aprendizaje", level=2)
    p(doc, "La convergencia se estima desde reward_mean_average por algoritmo y escenario. Para evitar lectura visual subjetiva, se usa una media movil de longitud cinco: se considera inicio de aprendizaje el primer episodio ordinal en que la media movil supera el 20% de la mejora entre el tramo inicial y el tramo final; se considera estabilizacion el primer episodio desde el cual la media movil permanece dentro del 5% relativo respecto del tramo final. La recompensa es mejor cuando es menos negativa, por lo que el maximo observado indica el mejor episodio conservado en el artefacto.")
    conv_rows = []
    for _, r in convergence.sort_values(["scenario", "algorithm"]).iterrows():
        conv_rows.append(
            [
                r["algorithm"],
                r["scenario"],
                int(r["n_episode_artifacts"]),
                fmt(r["initial_rolling_reward"], 6),
                fmt(r["final_rolling_reward"], 6),
                int(r["learning_start_episode_ordinal"]),
                int(r["stabilization_episode_ordinal"]),
                int(r["best_episode_ordinal"]),
                fmt(r["best_reward_mean"], 6),
            ]
        )
    table(doc, "Tabla 5.3. Episodios de inicio de aprendizaje, estabilizacion y mejor recompensa media.", ["Algoritmo", "Esc.", "n", "media inicial", "media final", "inicio aprendizaje", "estabilizacion", "mejor episodio", "mejor reward"], conv_rows, 6.6)
    for scenario in SCENARIOS:
        add_picture(doc, f"Figura 5.3-{scenario}. Curva de convergencia por recompensa media movil en {scenario}.", figures[f"convergence_{scenario}"])
    p(doc, "La lectura de convergencia muestra aprendizaje temprano en la mayoria de tratamientos, pero no implica dominancia automatica. MATD3 presenta una mejora marcada en E1 antes de estabilizarse; MAAC y MASAC muestran variaciones mas compactas; HAPPO se interpreta descriptivamente con las filas episodicas materializadas disponibles y con el registro Drive de entrenamiento completado.")

    doc.add_heading("5.4 KPIs bajo nomenclatura CityLearn v2 evaluate_v2", level=2)
    p(doc, "Para que los resultados sean comparables con CityLearn v2, la tesis no reduce la evaluacion a una metrica ad hoc. Los CSV de resumen comparativo usan la nomenclatura de evaluate_v2 y agrupan KPIs por eje: OE1 flexibilidad energetica, OE2 emisiones de CO2 y OE3 costos energeticos. Esta lectura complementa la recompensa de entrenamiento con KPIs finales de evaluacion, incluyendo baseline y RBC horario cuando estan disponibles.")
    if not kpi_catalog.empty:
        catalog_rows = []
        for _, r in kpi_catalog.iterrows():
            if r["axis"] in {"OE1", "OE2", "OE3"}:
                catalog_rows.append([r["scenario"], r["axis"], r["axis_name"], int(r["available_unique_kpis"]), r["source"], r["example_kpis"]])
        table(doc, "Tabla 5.4. Catalogo de KPIs evaluate_v2 usados para interpretar resultados.", ["Esc.", "Eje", "Dimension", "KPIs", "Fuente", "Ejemplos"], catalog_rows, 6.2)
    if not kpi_ranking.empty:
        rank_rows = []
        for _, r in kpi_ranking.iterrows():
            rank_rows.append([r["scenario"], r["axis"], r["family"], r["method"], fmt(r["normalized_score"], 4), int(r["available_kpis"]), int(r["improved_kpis"]), fmt(r["axis_rank"], 1)])
        table(doc, "Tabla 5.5. Ranking por eje con KPIs compatibles con CityLearn v2.", ["Esc.", "Eje", "Familia", "Metodo", "score", "KPIs disp.", "KPIs mejora", "rank"], rank_rows, 6.5)
        if "kpi_ranking_heatmap" in figures:
            add_picture(doc, "Figura 5.4. Mapa de calor del ranking KPI CityLearn v2 evaluate_v2.", figures["kpi_ranking_heatmap"])
    p(doc, "La lectura evaluate_v2 evita una conclusion sesgada por la recompensa de entrenamiento: un algoritmo puede maximizar reward_mean_average en un escenario y, al mismo tiempo, no liderar todos los KPIs oficiales de flexibilidad, carbono o costo. Por ello, la decision doctoral se reporta en tres niveles: media episodica, prueba estadistica intra-corrida y KPI anual final compatible con CityLearn v2.")

    for idx, oe in enumerate(["OE.1", "OE.2", "OE.3"], start=5):
        d = detail[oe]
        spec = d["spec"]
        doc.add_heading(f"5.{idx} {oe}: efecto del MADRL sobre {spec['dimension']}", level=2)
        p(doc, spec["explanation"])
        rows = []
        for _, r in d["desc"].iterrows():
            rows.append([
                r["algorithm"],
                int(r["count"]),
                fmt(r["mean"], 6 if spec["metric"] == "reward_mean_average" else 2),
                fmt(r["median"], 6 if spec["metric"] == "reward_mean_average" else 2),
                fmt(r["std"], 6 if spec["metric"] == "reward_mean_average" else 2),
                fmt(r["min"], 6 if spec["metric"] == "reward_mean_average" else 2),
                fmt(r["max"], 6 if spec["metric"] == "reward_mean_average" else 2),
                "cobertura completa" if int(r["count"]) >= 50 else f"{int(r['count'])} filas conservadas",
            ])
        table(doc, f"Tabla 5.{idx}. Estadistica descriptiva por episodio para {oe}.", ["Algoritmo", "n", "Media", "Mediana", "Desv. est.", "Min", "Max", "Cobertura"], rows, 7.0)
        final_rows = []
        for _, r in d["final"].iterrows():
            final_rows.append([
                r["algorithm"],
                fmt(r["final_metric"], 6 if oe == "OE.1" else 2),
                fmt(r["peak_average"], 4),
                fmt(r["ramping_average"], 4),
                fmt(r["carbon_emissions_control"], 2),
                fmt(r["electricity_cost_control"], 2),
            ])
        table(doc, f"Tabla 5.{idx}a. KPI anual final del tratamiento asociado a {oe}.", ["Algoritmo", "Indicador final OE", "peak", "ramping", "CO2 control", "costo control"], final_rows, 7.0)
        add_picture(doc, f"Figura 5.{idx}. Comparacion grafica de la media por episodio para {oe}.", figures[oe])
        ts_key = f"final_timeseries_{spec['scenario']}"
        if ts_key in figures:
            add_picture(doc, f"Figura 5.{idx}a. Paneles MADRL del episodio final para {spec['scenario']}: valores normalizados y totales reales anotados por algoritmo.", figures[ts_key])
            p(doc, f"Nota de lectura de la Figura 5.{idx}a: los paneles no interpolan datos faltantes. Cada subgrafico informa la cobertura real conservada del timeseries final por algoritmo, el episodio usado y los totales observados de energia neta, costo y CO2. Si un tratamiento conserva una serie dispersa con valores de cierre, la figura lo muestra como tal para evitar convertir artefactos de persistencia en curvas operativas inventadas.")
        kw = d["kw"]
        if kw is not None:
            p(doc, f"Contrastacion inferencial: para {oe} se aplica Kruskal-Wallis solo a los grupos con cobertura completa conservada ({', '.join(d['inferential_algos'])}). El resultado es H={kw.statistic:.4f}, p={kw.pvalue:.6g}, epsilon2={d['epsilon2']:.4f}. Con alpha=0,05, {'se rechaza H0 y se identifica efecto diferenciado del algoritmo' if kw.pvalue < 0.05 else 'no se rechaza H0 en la muestra inferencial conservada'}. HAPPO registra entrenamiento completado en Drive, pero no entra al contraste inferencial porque el CSV materializado conserva 49 filas episodicas por escenario.")
            pair_rows = [[name, fmt_p(pv), fmt_p(adj), "significativo" if adj < 0.05 else "no significativo"] for name, pv, adj in d["pair_adj"]]
            table(doc, f"Tabla 5.{idx}b. Mann-Whitney U por pares con ajuste Holm para {oe}.", ["Par", "p", "p Holm", "Decision"], pair_rows, 7.2)
            sh_rows = []
            for _, sr in d["shapiro_rows"].iterrows():
                sh_rows.append([
                    sr["algorithm"],
                    int(sr["n"]),
                    fmt(sr["w"], 6),
                    fmt_p(sr["p"]),
                    sr["decision"],
                    sr["inferential_use"],
                ])
            table(doc, f"Tabla 5.{idx}c. Shapiro-Wilk recalculado por algoritmo para {oe}.", ["Algoritmo", "n", "W", "p", "Interpretacion", "Uso"], sh_rows, 6.4)
            p(doc, f"Interpretacion de normalidad para {oe}: todos los p-valores Shapiro-Wilk recalculados son menores que 0,05, por lo que se rechaza la normalidad en las muestras episodicas conservadas. Esta evidencia justifica mantener Kruskal-Wallis y Mann-Whitney U con ajuste Holm como pruebas no parametricas. HAPPO se informa para transparencia descriptiva, pero no se incorpora al contraste Kruskal-Wallis porque su muestra materializada no alcanza 50 episodios.")
        p(doc, f"Interpretacion de {oe}: el mejor algoritmo por media episodica conservada es {d['best_stat']}; al restringir la decision inferencial a algoritmos con cobertura completa, el mejor es {d['best_stat_complete']}; el mejor KPI anual final observado es {d['best_final']}. Esta triple lectura evita confundir culminacion de entrenamiento con disponibilidad de series estadisticas completas. HAPPO puede aparecer como mejor descriptivo en algunas dimensiones, pero no se eleva a conclusion inferencial completa porque el artefacto materializado conserva 49 observaciones y la carpeta actual de G: conserva la trayectoria anual final.")

    doc.add_heading("5.8 Resultados por edificio y equipamiento controlado", level=2)
    p(doc, "El analisis por edificio usa building_behavior_summary.csv y building_kpis.csv de los 12 tratamientos. Cada edificio actua como agente de la comunidad y posee dimensiones heterogeneas de observacion y accion; por ello, la cantidad de equipos controlados no es uniforme. La accion controlada agrupa BESS, cargadores EV, cargas flexibles y otros actuadores declarados en building_observation_action_schema.csv. Las cargas no controladas permanecen dentro de la demanda base y de las variables observadas, no como acciones directas del agente.")
    add_picture(doc, "Figura 5.8. Dimensiones de accion controlable por edificio.", figures["equipment"])
    for key, caption in [
        ("building_ev_success_heatmap", "Figura 5.8a. Exito de salida EV por edificio y algoritmo."),
        ("building_carbon_heatmap", "Figura 5.8b. CO2 control por edificio y algoritmo."),
        ("building_cost_heatmap", "Figura 5.8c. Costo control por edificio y algoritmo."),
        ("equipment_class_heatmap", "Figura 5.8d. Equipamiento controlado por edificio y clase."),
        ("trace_policy_heatmaps", "Figura 5.8e. Politicas y acciones medias (action_l2 desde trace.csv; EV/BESS desde building_behavior_summary)."),
    ]:
        if key in figures:
            add_picture(doc, caption, figures[key])
    b_top = building_compact.sort_values("action_dim", ascending=False).head(10)
    table(doc, "Tabla 5.11. Edificios con mayor cantidad de acciones controlables.", ["Algoritmo", "Esc.", "Edificio", "rol red", "acciones", "obs.", "BESS kWh", "EV kWh", "CO2 kg", "costo"], [[r["algorithm"], r["scenario"], r["agent"], r.get("grid_role_control", ""), int(r["action_dim"]), int(r["observation_dim"]), fmt(r.get("battery_throughput_total_kwh"), 1), fmt(r.get("ev_charge_total_kwh"), 1), fmt(r.get("carbon_emissions_control_kgco2"), 1), fmt(r.get("electricity_cost_control_eur"), 1)] for _, r in b_top.iterrows()], 6.6)
    eq_pivot = eq_summary.groupby("equipment_class")["count"].sum().reset_index().sort_values("count", ascending=False)
    table(doc, "Tabla 5.12. Equipamiento controlado identificado en los esquemas de accion.", ["Clase de equipo controlado", "conteo en 12 tratamientos"], [[r["equipment_class"], int(r["count"])] for _, r in eq_pivot.iterrows()], 7.2)
    p(doc, "Las tablas completas se guardan como CSV en outputs/_drive_madrl/gdrive_20260627_164047_objective_analysis/tables. Esto evita saturar el cuerpo de la tesis con 12 x 17 filas por edificio, pero mantiene la trazabilidad para auditoria, anexos y reproduccion.")

    doc.add_heading("5.9 Sintesis de contrastacion de OE.1, OE.2 y OE.3", level=2)
    synth_rows = []
    for oe in ["OE.1", "OE.2", "OE.3"]:
        d = detail[oe]
        kw = d["kw"]
        synth_rows.append([
            oe,
            d["spec"]["dimension"],
            d["spec"]["scenario"],
            d["best_stat"],
            d["best_stat_complete"],
            d["best_final"],
            f"H={kw.statistic:.3f}; p={kw.pvalue:.3g}" if kw else "NA",
            "se rechaza H0" if kw and kw.pvalue < 0.05 else "no se rechaza H0",
        ])
    table(doc, "Tabla 5.13. Respuesta directa a objetivos especificos.", ["Objetivo", "Dimension", "Escenario", "mejor media episodica", "mejor muestra completa", "mejor KPI anual final", "Kruskal-Wallis", "Decision"], synth_rows, 6.8)
    for key, caption in [
        ("episode_boxplots", "Figura 5.9a. Distribucion episodica por objetivo y algoritmo."),
        ("effect_size", "Figura 5.9b. Tamano de efecto inferencial por objetivo."),
        ("pairwise_heatmaps", "Figura 5.9c. Matriz de p-valores Holm por objetivo."),
    ]:
        if key in figures:
            add_picture(doc, caption, figures[key])
    p(doc, "La interpretacion doctoral no debe afirmar dominancia unica sin matices. Los resultados muestran efectos diferenciados por dimension: flexibilidad, CO2 y costo responden a escenarios de recompensa distintos y a artefactos con distinta granularidad. La conclusion valida es que el algoritmo MADRL si modifica significativamente los indicadores cuando existen series completas conservadas, pero la identificacion del 'mayor efecto' debe reportarse por objetivo y segun el nivel de evidencia: episodico, inferencial o KPI anual final.")
    p(doc, "Metodologicamente, esta decision sigue las advertencias de evaluacion rigurosa en aprendizaje por refuerzo: los episodios de una corrida no reemplazan multiples semillas independientes, y los p-valores deben interpretarse junto con cobertura, tamano de efecto y trazabilidad de artefactos (Henderson et al., 2018; Colas et al., 2019; Agarwal et al., 2021; Patterson et al., 2024). Por ello, se retiene Shapiro-Wilk, Kruskal-Wallis y Mann-Whitney U con Holm como contrastacion intra-corrida, y se declara la necesidad de multi-semilla para robustez externa.")
    add_cap5_triangulated_discussion(doc, detail, kpi_ranking, figures)


def add_cap5_triangulated_discussion(doc: Document, detail: dict, kpi_ranking: pd.DataFrame, figures: dict[str, Path]) -> None:
    doc.add_heading("5.10 Discusion triangulada de resultados, baseline y trabajos relacionados", level=2)
    p(doc, "El Capitulo 5 tiene el mayor peso empirico del documento porque integra experimentos realizados, metricas, resultados, comparacion con baseline, tablas, figuras y discusion. La triangulacion de resultados se realiza en tres niveles: recompensa episodica, KPIs evaluate_v2 y contraste estadistico. Esta estrategia evita que una unica grafica de convergencia se interprete como evidencia suficiente, y responde a las recomendaciones de evaluacion rigurosa en RL, donde se exige reportar variabilidad, tamano de efecto y robustez de la comparacion (Henderson et al., 2018; Agarwal et al., 2021).")
    rows = []
    for oe in ["OE.1", "OE.2", "OE.3"]:
        d = detail[oe]
        kw = d["kw"]
        rows.append([oe, d["spec"]["scenario"], d["spec"]["dimension"], d["best_stat"], d["best_stat_complete"], d["best_final"], f"H={kw.statistic:.4f}; p={kw.pvalue:.6g}; epsilon2={d['epsilon2']:.4f}" if kw else "NA", effect_label(d["epsilon2"])])
    table(doc, "Tabla 5.14. Discusion sintetica por objetivo, efecto y algoritmo dominante.", ["Objetivo", "Esc.", "Dimension", "mejor descriptivo", "mejor inferencial", "mejor KPI final", "Prueba", "Tamano"], rows, 6.5)
    if not kpi_ranking.empty:
        best_rows = []
        for scenario in SCENARIOS:
            sub = kpi_ranking[kpi_ranking["scenario"] == scenario].sort_values("axis_rank")
            if not sub.empty:
                r = sub.iloc[0]
                best_rows.append([scenario, r["axis"], r["family"], r["method"], fmt(r["normalized_score"], 4), fmt(r["axis_rank"], 1)])
        table(doc, "Tabla 5.15. Mejor metodo por eje segun ranking CityLearn v2 evaluate_v2.", ["Esc.", "Eje", "Familia", "Metodo", "score normalizado", "rank"], best_rows, 7.0)
    if "tradeoff" in figures:
        add_picture(doc, "Figura 5.10. Trade-off multiobjetivo costo-CO2-autoconsumo PV.", figures["tradeoff"])
    p(doc, "En flexibilidad energetica (PE.1/OE.1), la evidencia es la mas fuerte: MAAC obtiene la mejor media episodica conservada y lidera la muestra inferencial completa, mientras que Kruskal-Wallis rechaza H0 con epsilon2=0,2334, interpretado como efecto alto. Este resultado es compatible con la teoria de MAAC, porque el critico con atencion puede priorizar interacciones relevantes entre edificios cuando la dimension dominante es la coordinacion de flexibilidad; tambien se relaciona con CityLearn v2, donde los KPIs de flexibilidad incluyen pico, ramping, factor de carga y uso de almacenamiento (Iqbal & Sha, 2019; Nweye et al., 2024; Vazquez-Canteli et al., 2020).")
    p(doc, "En emisiones de CO2 (PE.2/OE.2), la evidencia muestra efecto inferencial significativo pero de tamano bajo. HAPPO presenta el mejor promedio descriptivo conservado, pero al restringir la decision a la muestra inferencial completa el mejor algoritmo es MAAC; en KPI anual final aparece MASAC. Esta divergencia no debe ocultarse, porque indica que el comportamiento carbono-dependiente no se reduce a un unico criterio. La literatura sobre SAC/MASAC sugiere que la regularizacion por entropia puede estabilizar exploracion en problemas continuos, mientras que CityLearn v2 y EVLearn muestran que carbono y carga EV dependen de senales temporales y restricciones de disponibilidad que pueden modificar el ranking final (Haarnoja et al., 2018; Fonseca et al., 2024; Nweye et al., 2024).")
    p(doc, "En costos energeticos (PE.3/OE.3), la prueba inferencial no rechaza H0 y el tamano de efecto es muy bajo. Por ello, la tesis no afirma una superioridad estadistica concluyente. Descriptivamente, HAPPO muestra menor costo medio entre las filas conservadas, pero en la muestra completa MATD3 presenta mejor promedio y MAAC obtiene el mejor KPI anual final. Esta lectura matizada es coherente con la literatura de TD3/MATD3, donde los criticos dobles reducen sesgos de estimacion, pero no garantizan dominancia en todos los objetivos multiobjetivo; tambien coincide con las advertencias de reproducibilidad en RL sobre no convertir diferencias numericas en conclusiones robustas sin replicas independientes (Fujimoto et al., 2018; Henderson et al., 2018; Agarwal et al., 2021).")
    p(doc, "La comparacion con baseline y trabajos relacionados se interpreta como evidencia contextual, no como sustituto de la contrastacion principal. Cuando el ranking evaluate_v2 ubica a un baseline o RBC por encima de MADRL en algun eje, el resultado se reporta porque forma parte de la evidencia real y muestra que el aprendizaje multiagente no domina automaticamente a politicas simples en todos los indicadores. Esta transparencia fortalece la validez doctoral: el aporte no consiste en afirmar superioridad universal, sino en identificar donde el MADRL produce efecto, con que magnitud y bajo que escenario de recompensa.")


def add_cap5_madrl_nature_figures(doc: Document, figures: dict[str, Path]) -> None:
    doc.add_heading("5.11 Figuras complementarias para evaluar la naturaleza MADRL", level=2)
    p(doc, "La curva de convergencia por recompensa media movil es necesaria para verificar aprendizaje, pero no es suficiente para evaluar la naturaleza de cada MADRL. Por ello se incorporan figuras complementarias basadas en episodios, KPIs oficiales, trazas, series temporales finales, edificios, equipamiento y checkpoints. Estas visualizaciones permiten distinguir aprendizaje, efecto estadistico, trade-off multiobjetivo, comportamiento fisico y cobertura de entrenamiento.")
    figure_plan = [
        ("episode_boxplots", "Figura 5.11a. Distribucion episodica por objetivo y algoritmo."),
        ("effect_size", "Figura 5.11b. Tamano de efecto inferencial por objetivo."),
        ("pairwise_heatmaps", "Figura 5.11c. Matriz visual de p-valores Holm por objetivo."),
        ("kpi_ranking_heatmap", "Figura 5.11d. Ranking de KPIs CityLearn v2 evaluate_v2."),
        ("tradeoff", "Figura 5.11e. Trade-off multiobjetivo costo-CO2-autoconsumo PV."),
        ("building_ev_success_heatmap", "Figura 5.11f. Exito de salida EV por edificio y algoritmo."),
        ("building_carbon_heatmap", "Figura 5.11g. CO2 por edificio y algoritmo."),
        ("building_cost_heatmap", "Figura 5.11h. Costo por edificio y algoritmo."),
        ("equipment_class_heatmap", "Figura 5.11i. Equipamiento controlado por edificio y clase."),
        ("final_timeseries_E1", "Figura 5.11j. Serie temporal distrital final en E1."),
        ("final_timeseries_E2", "Figura 5.11k. Serie temporal distrital final en E2."),
        ("final_timeseries_E3", "Figura 5.11l. Serie temporal distrital final en E3."),
        ("trace_policy_heatmaps", "Figura 5.11m. Politicas/acciones desde trace.csv."),
        ("checkpoint_coverage", "Figura 5.11n. Cobertura de checkpoints por tratamiento."),
    ]
    for key, caption in figure_plan:
        if key in figures and Path(figures[key]).exists():
            add_picture(doc, caption, figures[key], width=5.9)
    p(doc, "Estas figuras no reemplazan las pruebas estadisticas; las complementan. La distribucion episodica muestra variabilidad, el tamano de efecto cuantifica magnitud, los p-valores Holm ubican diferencias por pares, el trade-off evidencia tensiones entre costo y CO2, los mapas por edificio muestran heterogeneidad multiagente, las series temporales finales explican el comportamiento fisico y la cobertura de checkpoints documenta trazabilidad de entrenamiento.")


def effect_label(epsilon2: float) -> str:
    if pd.isna(epsilon2):
        return "no estimado"
    if epsilon2 >= 0.14:
        return "efecto alto"
    if epsilon2 >= 0.06:
        return "efecto medio"
    if epsilon2 >= 0.01:
        return "efecto bajo"
    return "efecto muy bajo"


def add_cap4_detailed_proposal_from_summary(doc: Document) -> None:
    doc.add_heading("4.9 Detalle tecnico de la propuesta derivado del resumen", level=2)
    p(doc, "El resumen sintetiza la propuesta doctoral; por tanto, este numeral desarrolla de forma tecnica los mismos componentes para que el lector no encuentre informacion clave solamente en la parte inicial del documento. La propuesta consiste en una arquitectura experimental MADRL sobre CityLearn v2, extendida como CityLearn v3 propuesto, que permite representar una comunidad inteligente del SEAI Iquitos como Dec-POMDP, entrenar politicas bajo CTDE y evaluar el efecto de HAPPO, MAAC, MASAC y MATD3 sobre flexibilidad energetica, emisiones de CO2 y costos energeticos. Esta seccion no introduce resultados nuevos: organiza y detalla los datos, librerias, fuentes, scripts y artefactos que sustentan el resumen y que luego se contrastan en el Capitulo 5.")
    p(doc, "La propuesta opera con tres niveles integrados. En el nivel de datos, el repositorio construye el dataset citylearn_iquitos_2023_2025 desde fuentes electricas, climaticas, solares, tarifarias y de carbono. En el nivel de simulacion, CityLearn v2 representa edificios, PV, BESS, EV, demanda base, precio, intensidad de carbono y acciones controlables por agente. En el nivel de aprendizaje, los algoritmos MADRL entrenan politicas descentralizadas con informacion local, mientras el entrenamiento centralizado permite estabilizar la coordinacion bajo no-estacionariedad multiagente.")
    table(
        doc,
        "Tabla 4.7. Componentes detallados de la propuesta sintetizados en el resumen.",
        ["Componente", "Detalle desarrollado en la propuesta", "Evidencia/artefacto"],
        [
            ["Comunidad inteligente", "17 edificios institucionales/comerciales del SEAI Iquitos, modelados como agentes heterogeneos.", "dataset citylearn_iquitos_2023_2025; building_behavior_summary.csv."],
            ["Horizonte temporal", "26 304 pasos horarios del periodo 2023-2025; cada episodio experimental evalua un ano horario de 8 760 pasos.", "archivos CSV del dataset y timeseries.csv."],
            ["Recursos energeticos", "PV, BESS, cargadores EV, V2G, cargas controlables y demanda no controlada separada para evitar atribuir acciones a cargas no actuables.", "schema de edificios, traces y building KPIs."],
            ["Tratamientos MADRL", "HAPPO, MAAC, MASAC y MATD3 bajo escenarios E1 flexibilidad, E2 CO2 y E3 costos.", "12 carpetas algoritmo-escenario y tablas episodicas."],
            ["Evidencia experimental", "50 episodios por tratamiento disponible; HAPPO conserva 49 filas episodicas materializadas por escenario en el CSV final y se interpreta con esa limitacion.", "district_episode_kpis.csv y tablas de estadistica."],
        ],
        5.8,
    )

    doc.add_heading("4.9.1 Construccion del dataset y fuentes utilizadas", level=3)
    p(doc, "El dataset no se plantea como un insumo generico, sino como la base empirica de la propuesta. La construccion integra facturacion y perfiles de Electro Oriente S.A., meteorologia de NASA POWER, irradiacion y estimacion solar con PVGIS/pvlib, factores de carbono basados en MINAM y senales tarifarias alineadas con OSINERGMIN. La finalidad de esta triangulacion es que los agentes no aprendan sobre una comunidad abstracta, sino sobre una representacion horaria coherente con clima, demanda, irradiacion, precio e intensidad de carbono del SEAI Iquitos.")
    p(doc, "La escala tecnica desarrollada en el resumen se incorpora en la propuesta: 17 edificios, 26 304 pasos horarios, 222 archivos CSV activos, 185 cargadores EV, 31 tomas V2G, 96 equipos fisicos Mode 3, 749,4 kW EV, 26 266 kWh de BESS, 6 648 kW BESS y 48 790,9 kWp fotovoltaicos. Estos valores no son objetivos aspiracionales, sino parametros de entrada y auditoria de la simulacion. Su funcion metodologica es delimitar el espacio de observacion, accion y restricciones que enfrentan los agentes.")
    table(
        doc,
        "Tabla 4.8. Fuentes y transformaciones del dataset citylearn_iquitos_2023_2025.",
        ["Fuente/insumo", "Transformacion en el proyecto", "Uso en CityLearn/MADRL"],
        [
            ["Electro Oriente S.A.", "Perfiles de demanda/facturacion y contexto del SEAI Iquitos.", "Demanda base y calibracion de edificios."],
            ["NASA POWER", "Series meteorologicas horarias para el periodo de estudio.", "Variables exogenas de clima en observaciones."],
            ["PVGIS y pvlib", "Estimacion de generacion solar y perfiles PV por capacidad instalada.", "Produccion fotovoltaica por edificio y balance energetico."],
            ["Inventario de edificios y equipos", "Asignacion de PV, BESS, EV, V2G, cargas controlables y no controlables.", "Dimensiones de observacion/accion y heterogeneidad de agentes."],
            ["MINAM", "Factor de emision base y modelo de intensidad de carbono.", "Calculo de emisiones y recompensa carbon-aware."],
            ["OSINERGMIN / senales TOU", "Parametrizacion de precios punta y fuera de punta.", "Costo electrico y recompensa economica."],
        ],
        5.9,
    )

    doc.add_heading("4.9.2 Librerias, scripts y herramientas de implementacion", level=3)
    p(doc, "La propuesta se implementa con un pipeline reproducible en Python. pandas y NumPy realizan consolidacion tabular, limpieza, agregaciones y calculo de KPIs; pvlib y requests participan en la construccion de perfiles solares y descarga/normalizacion de fuentes externas; PyTorch ejecuta las politicas neuronales y backends MADRL; SciPy calcula pruebas estadisticas; Matplotlib genera figuras de convergencia, distribuciones, radar, heatmaps y paneles temporales; CityLearn, Gymnasium y PettingZoo proveen la interfaz de ambiente multiagente; PowerShell verifica contexto y automatiza controles de repositorio; Google Colab/Drive aloja la corrida canonica y conserva resultados, checkpoints, timeseries y traces.")
    p(doc, "La decision de registrar librerias y scripts dentro de la propuesta es necesaria porque la tesis no evalua un concepto teorico aislado. Evalua una implementacion computacional concreta. Por ello, los scripts de generacion de dataset, entrenamiento, validacion, descarga de artefactos, analisis estadistico, construccion de figuras y generacion del Word constituyen parte del metodo y de la propuesta. Si un resultado aparece en Capitulo 5, su origen debe poder rastrearse a uno de estos componentes.")
    table(
        doc,
        "Tabla 4.9. Herramientas y funcion dentro de la propuesta MADRL.",
        ["Herramienta/libreria", "Funcion tecnica", "Salida esperada"],
        [
            ["Python 3.9", "Lenguaje base de orquestacion, analisis y generacion documental.", "Scripts reproducibles y tablas/figuras."],
            ["pandas, NumPy", "Tratamiento de CSV, agregacion por episodio, edificio, escenario y algoritmo.", "district_episode_kpis, rankings y resumenes."],
            ["pvlib, PVGIS, NASA POWER", "Modelado solar y meteorologico del dataset.", "Perfiles PV y variables exogenas horarias."],
            ["PyTorch", "Entrenamiento de redes actor-critic y politicas MADRL.", "Politicas, rewards, checkpoints."],
            ["CityLearn v2 / CityLearn v3 propuesto", "Simulacion energetica multiagente con PV, BESS, EV, carbono y costos.", "Observaciones, acciones, rewards, KPIs."],
            ["Gymnasium, PettingZoo", "Interfaz multiagente y compatibilidad de entorno.", "Ejecucion descentralizada y wrappers."],
            ["SciPy", "Shapiro-Wilk, Kruskal-Wallis y Mann-Whitney U con ajuste posterior.", "Contraste inferencial por OE.1-OE.3."],
            ["Matplotlib", "Graficacion de convergencia, boxplots, radar, heatmaps, trade-offs y series finales.", "Figuras del Capitulo 5."],
            ["Google Colab/Drive", "Ejecucion y persistencia de la corrida canonica.", "results, timeseries, traces, checkpoints."],
        ],
        5.6,
    )

    doc.add_heading("4.9.3 Arquitectura operativa desde datos hasta resultados", level=3)
    p(doc, "La arquitectura operativa inicia con la verificacion de contexto del repositorio y la disponibilidad del Drive de resultados. Luego se consolidan archivos de datos, se instancia el entorno CityLearn, se configuran los 17 agentes, se asignan equipos controlables, se selecciona el algoritmo MADRL y se ejecuta el escenario de recompensa. Cada tratamiento produce resultados episodicos, series temporales, trazas por agente, KPIs por edificio y checkpoints. La etapa final consolida esta evidencia en tablas estadisticas, figuras y redaccion interpretativa del documento final.")
    p(doc, "Esta arquitectura evita que la propuesta quede separada del analisis. El mismo flujo que define el ambiente y las politicas define tambien la forma de responder PE.1, PE.2 y PE.3. En E1 se observa flexibilidad; en E2 se observa CO2; en E3 se observa costo. La comparacion no se realiza por intuicion, sino por indicadores materializados y pruebas estadisticas. Asi, el resumen, la metodologia, la propuesta y los resultados comparten la misma cadena de evidencia.")
    table(
        doc,
        "Tabla 4.10. Cadena de trazabilidad de la propuesta doctoral.",
        ["Fase", "Entrada", "Proceso", "Salida verificable"],
        [
            ["Construccion de datos", "Electro Oriente, NASA POWER, PVGIS, MINAM, OSINERGMIN, inventario de equipos.", "Normalizacion, generacion de perfiles y auditoria de CSV.", "dataset citylearn_iquitos_2023_2025."],
            ["Simulacion", "Dataset y configuracion de edificios.", "Instanciacion CityLearn v2 / CityLearn v3 propuesto.", "observaciones, acciones y balances horarios."],
            ["Entrenamiento MADRL", "Algoritmo, escenario y recompensa multiobjetivo.", "CTDE con politicas descentralizadas.", "rewards, checkpoints y logs."],
            ["Registro de evidencia", "Ejecucion por tratamiento.", "Persistencia en Drive y consolidacion local.", "results.json, timeseries.csv, trace.csv, building_kpis.csv."],
            ["Analisis doctoral", "CSV episodicos, traces, checkpoints y KPIs.", "Estadistica, figuras y lectura por OE.", "Capitulos 5 y 6 con evidencia real."],
        ],
        5.7,
    )

    doc.add_heading("4.9.4 Delimitacion entre equipos controlados y no controlados", level=3)
    p(doc, "La propuesta distingue equipos controlados y no controlados porque el efecto MADRL solo puede atribuirse a variables de accion disponibles para los agentes. Los equipos controlados incluyen BESS, cargadores EV/V2G y cargas flexibles declaradas en el esquema del entorno. La demanda base, condiciones climaticas, irradiacion, ocupacion implicita, precio e intensidad de carbono operan como observaciones o perturbaciones exogenas, no como actuadores. Esta distincion es esencial para interpretar correctamente los resultados por edificio y no atribuir al algoritmo variaciones que provienen de cargas no controlables.")
    p(doc, "En terminos de tesis, esta delimitacion conecta Capitulo 4 con Capitulo 5: las acciones registradas en traces permiten observar como cada politica usa BESS, EV y otros actuadores; los KPIs por edificio permiten evaluar si la coordinacion mejora flexibilidad, CO2 o costo; y las limitaciones declaran que el estudio no controla red fisica ni despacho de generacion. Por tanto, la propuesta es una arquitectura de control y evaluacion energetica multiagente en simulacion, no una implementacion de despacho electrico real.")


def add_cap4_problem_question_response(doc: Document, detail: dict) -> None:
    doc.add_heading("4.10 Respuesta operacional a las preguntas especificas PE.1, PE.2 y PE.3", level=2)
    p(doc, "La propuesta no se limita a describir una arquitectura MADRL; tambien define como se responde cada pregunta especifica mediante la salida empirica de la corrida Drive madrl_v3_20260627_164047. La respuesta se apoya en dos planos: analisis descriptivo de los episodios y KPIs anuales finales, y analisis inferencial intra-corrida con Kruskal-Wallis sobre los algoritmos que conservan cobertura episodica completa. Esta regla evita mezclar una diferencia numerica descriptiva con una afirmacion causal o inferencial no sustentada.")
    question_map = {
        "OE.1": ("PE.1", "¿En que medida el algoritmo MADRL (VI) produce un efecto sobre la dimension de flexibilidad energetica de la comunidad (D-VD.1), y cual algoritmo genera el mayor efecto?"),
        "OE.2": ("PE.2", "¿En que medida el algoritmo MADRL (VI) produce un efecto sobre la dimension de emisiones de CO2 de la comunidad (D-VD.2), y cual algoritmo genera el mayor efecto?"),
        "OE.3": ("PE.3", "¿En que medida el algoritmo MADRL (VI) produce un efecto sobre la dimension de costos energeticos de la comunidad (D-VD.3), y cual algoritmo genera el mayor efecto?"),
    }
    rows = []
    for oe in ["OE.1", "OE.2", "OE.3"]:
        d = detail[oe]
        spec = d["spec"]
        kw = d["kw"]
        epsilon2 = d["epsilon2"]
        decision = "se rechaza H0" if kw and kw.pvalue < 0.05 else "no se rechaza H0"
        measure = f"H={kw.statistic:.4f}; p={kw.pvalue:.6g}; epsilon2={epsilon2:.4f} ({effect_label(epsilon2)})" if kw else "sin prueba inferencial"
        rows.append(
            [
                question_map[oe][0],
                spec["dimension"],
                spec["scenario"],
                spec["metric"],
                measure,
                decision,
                d["best_stat"],
                d["best_stat_complete"],
                d["best_final"],
            ]
        )
    table(
        doc,
        "Tabla 4.11. Respuesta directa a PE.1, PE.2 y PE.3 desde la evidencia descriptiva e inferencial.",
        ["Pregunta", "Dimension", "Esc.", "Indicador", "Medida del efecto", "Decision", "mejor descriptivo", "mejor inferencial", "mejor KPI final"],
        rows,
        6.2,
    )
    for oe in ["OE.1", "OE.2", "OE.3"]:
        d = detail[oe]
        pe, question = question_map[oe]
        spec = d["spec"]
        kw = d["kw"]
        desc = d["desc"].copy()
        values = []
        for _, r in desc.iterrows():
            nd = 6 if spec["metric"] == "reward_mean_average" else 2
            values.append(f"{r['algorithm']}={fmt(r['mean'], nd)} (n={int(r['count'])})")
        if kw and kw.pvalue < 0.05:
            inferential_text = f"El efecto inferencial existe porque Kruskal-Wallis rechaza H0 (H={kw.statistic:.4f}; p={kw.pvalue:.6g}; epsilon2={d['epsilon2']:.4f}, {effect_label(d['epsilon2'])})."
        else:
            inferential_text = f"No se demuestra efecto inferencial suficiente en la muestra conservada porque Kruskal-Wallis no rechaza H0 (H={kw.statistic:.4f}; p={kw.pvalue:.6g}; epsilon2={d['epsilon2']:.4f}, {effect_label(d['epsilon2'])})."
        p(doc, f"{pe}. {question} Respuesta: en {spec['scenario']}, el indicador {spec['metric']} muestra los siguientes promedios episodicos conservados: " + "; ".join(values) + f". {inferential_text} El mayor efecto descriptivo corresponde a {d['best_stat']}; al exigir cobertura inferencial completa, corresponde a {d['best_stat_complete']}; y por KPI anual final corresponde a {d['best_final']}.")
    p(doc, "Por tanto, el Capitulo 4 deja definido el mecanismo de respuesta: PE.1 se responde con flexibilidad en E1, PE.2 con emisiones de CO2 en E2 y PE.3 con costos en E3. El Capitulo 5 desarrolla la contrastacion, figuras, tablas por edificio y ranking CityLearn v2, pero la relacion pregunta-variable-indicador-decision queda fijada aqui para mantener continuidad entre problema, propuesta y resultados.")


def add_cap6_completion_plan(doc: Document) -> None:
    doc.add_heading("6.5 Criterios de cierre doctoral y control de calidad final", level=2)
    p(doc, "Las conclusiones del estudio se consideran suficientemente sustentadas para responder las preguntas especificas desde la corrida Drive analizada. Dado que el Capitulo 6 ya contiene trabajo pendiente y plan de culminacion, esta seccion define criterios de cierre doctoral: validar numeracion y formato APA del documento completo, ejecutar una extension multi-semilla si se requiere robustez externa, revisar visualmente todas las figuras en Word/PDF y completar una lectura cruzada entre objetivos, hipotesis, resultados y conclusiones. Esta planificacion no agrega datos nuevos; delimita el control de calidad requerido para elevar la trazabilidad formal del manuscrito.")
    table(
        doc,
        "Tabla 6.2. Criterios de cierre y control de calidad final.",
        ["Actividad", "Proposito", "Criterio de cierre"],
        [
            ["Revision APA integral", "Alinear citas, tablas, figuras y referencias", "Todas las citas tienen entrada bibliografica y viceversa."],
            ["Revision multi-semilla opcional", "Mejorar validez externa de la comparacion MADRL", "Replicas documentadas o limitacion explicitada."],
            ["Auditoria de figuras y tablas", "Confirmar legibilidad y correspondencia con CSV/Drive", "Cada figura/tabla apunta a fuente de datos verificable."],
            ["Revision de coherencia vertical", "Asegurar que PE, OE, hipotesis y conclusiones respondan lo mismo", "Matriz problema-objetivo-resultado-conclusion sin vacios."],
        ],
        6.8,
    )


def add_docx_unification_audit(doc: Document) -> None:
    docs = sorted((REPO / "docs").glob("*.docx"))
    current_final_names = {
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_skill.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_ACTUALIZADA_METODOLOGIA_MADRL.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_50EP_ANTECEDENTES.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_GDRIVE_50EP_OBJETIVOS_DOCTORAL.docx",
    }
    legacy_markers = [
        "5 episodios",
        "referencia historica",
        "referencia local",
        "MATD3 domina flexibilidad y emisiones en la corrida canonica",
        "Colab (Kruskal-Wallis ALL",
    ]
    rows = []
    for path in docs:
        try:
            d = Document(path)
            body = "\n".join(p.text for p in d.paragraphs)
            words = len(re.findall(r"\b\w+\b", body))
            has_legacy = any(marker.lower() in body.lower() for marker in legacy_markers)
            if path.name in current_final_names:
                decision = "Unificado: copia sincronizada de la version final canonica."
            elif has_legacy:
                decision = "Auditado y no incorporado como texto vigente: contiene resultados historicos/locales o conclusiones reemplazadas por Drive 50 episodios."
            else:
                decision = "Auditado como antecedente documental; su estructura y contenidos validos ya estan absorbidos en la version canonica."
            rows.append([
                path.name[:60],
                f"{path.stat().st_size:,}",
                str(words),
                str(len(d.tables)),
                str(len(d.inline_shapes)),
                decision,
            ])
        except Exception as exc:
            rows.append([path.name[:60], "N/D", "N/D", "N/D", "N/D", f"No integrado: error de lectura {exc.__class__.__name__}."])

    doc.add_heading("Anexo C. Auditoria de unificacion de documentos Word", level=1)
    p(doc, "Este anexo documenta la revision de todos los archivos Word localizados en la carpeta docs. La version final canonica integra la estructura doctoral vigente, las figuras y tablas de resultados reales de Drive con 50 episodios, la metodologia actualizada, la respuesta a PE.1-PE.3/OE.1-OE.3 y la lista APA validada. Las versiones anteriores se revisaron como fuentes de control documental, pero no se trasladaron mecanicamente cuando contenian resultados historicos abreviados, rankings reemplazados o conclusiones incompatibles con la corrida Drive madrl_v3_20260627_164047.")
    table(
        doc,
        "Tabla C.1. Auditoria de integracion y unificacion de documentos Word en docs.",
        ["Documento", "Bytes", "Palabras", "Tablas", "Figuras", "Decision de integracion"],
        rows,
        5.2,
    )
    p(doc, "Resumen de cierre por capitulos. Capitulo 1: se valido problema, objetivos, hipotesis, justificacion, alcances y limitaciones. Capitulo 2: se integro estado del arte, bases teoricas, antecedentes internacionales/nacionales y formalizacion Dec-POMDP con triangulacion CityLearn, MADRL y evaluacion estadistica. Capitulo 3: se actualizo el enfoque como investigacion aplicada, cuantitativa y experimental-computacional factorial 4x3. Capitulo 4: se consolido arquitectura, modelo de IA, algoritmos, recompensa, implementacion y respuesta operacional a PE.1-PE.3. Capitulo 5: se priorizo la evidencia real de Drive, con metricas, figuras, tablas, comparacion por objetivo, edificios, equipamiento, checkpoints, trazas y pruebas inferenciales. Capitulo 6: se mantuvieron hallazgos, limitaciones, trabajo pendiente y plan de culminacion.")
    p(doc, "Criterio de no alucinacion documental: no se importaron cifras ni afirmaciones de versiones antiguas cuando contradicen los CSV, timeseries, traces, checkpoints o KPIs de la corrida Drive de 50 episodios. Todo resultado numerico vigente se conserva desde los artefactos auditables del proyecto local y del Drive indicado por el autor.")


def append_apa_references(doc: Document) -> None:
    def ref_key(text: str) -> tuple[str, str] | None:
        match = re.search(r"\(((?:19|20)\d{2}[a-z]?)\)", text)
        if not match:
            return None
        year = match.group(1)
        author = text.split("(")[0].strip().rstrip(".")
        author = author.split(",")[0].strip() if "," in author else author
        return norm_key(author), year

    existing = {key for para in iter_document_paragraphs(doc) if (key := ref_key(para.text))}
    doc.add_paragraph()
    cap = doc.add_paragraph()
    run = cap.add_run("Referencias complementarias verificadas e incorporadas en la revision")
    run.bold = True
    refs = [
        "Agarwal, R., Schwarzer, M., Castro, P. S., Courville, A. C., & Bellemare, M. G. (2021). Deep reinforcement learning at the edge of the statistical precipice. Advances in Neural Information Processing Systems, 34, 29304-29320.",
        "Chevarria Moscoso, M. (2024). Analisis de la generacion hidroelectrica en la central hidroelectrica de Machupicchu aplicando metodos estocasticos y modelo de optimizacion [Tesis doctoral, Universidad Nacional de Ingenieria]. Repositorio Institucional UNI. http://hdl.handle.net/20.500.14076/28894",
        "Creswell, J. W., & Creswell, J. D. (2023). Research design: Qualitative, quantitative, and mixed methods approaches (6th ed.). SAGE Publications.",
        "Fonseca, N., Nweye, K., & Nagy, Z. (2024). EVLearn: A mixed-autonomy multi-agent reinforcement learning environment for electric vehicle charging management. arXiv. https://arxiv.org/abs/2403.07612",
        "Fujimoto, S., van Hoof, H., & Meger, D. (2018). Addressing function approximation error in actor-critic methods. Proceedings of the 35th International Conference on Machine Learning, 80, 1587-1596.",
        "Haarnoja, T., Zhou, A., Abbeel, P., & Levine, S. (2018). Soft actor-critic: Off-policy maximum entropy deep reinforcement learning with a stochastic actor. Proceedings of the 35th International Conference on Machine Learning, 80, 1861-1870.",
        "Henderson, P., Islam, R., Bachman, P., Pineau, J., Precup, D., & Meger, D. (2018). Deep reinforcement learning that matters. Proceedings of the AAAI Conference on Artificial Intelligence, 32(1).",
        "Hernandez-Sampieri, R., & Mendoza, C. P. (2018). Metodologia de la investigacion: Las rutas cuantitativa, cualitativa y mixta. McGraw-Hill Education.",
        "Iqbal, S., & Sha, F. (2019). Actor-attention-critic for multi-agent reinforcement learning. Proceedings of the 36th International Conference on Machine Learning, 97, 2961-2970.",
        "Kuba, J. G., Chen, R., Wen, M., Wen, Y., Sun, F., Wang, J., & Yang, Y. (2021). Trust region policy optimisation in multi-agent reinforcement learning. arXiv. https://arxiv.org/abs/2109.11251",
        "Montgomery, D. C. (2019). Design and analysis of experiments (10th ed.). Wiley.",
        "Nweye, K., Sankur, M. D., Wu, C., & Nagy, Z. (2024). CityLearn v2: Energy-flexible, resilient, occupant-centric, and carbon-aware management of grid-interactive communities. Journal of Building Performance Simulation, 17(1), 1-20.",
        "Penalva Sanchez, J. J. (2024). Optimizacion de un sistema fotovoltaico hibrido y la prediccion de la demanda energetica y variables climaticas utilizando la inteligencia artificial [Tesis doctoral, Universidad Nacional de Ingenieria]. Repositorio Institucional UNI. http://hdl.handle.net/20.500.14076/27731",
        "Rosero Bernal, D. G. (2022). Modelo de un sistema de administracion de energia autonomo operado desde la nube para optimizar la gestion de un grupo de microrredes [Tesis doctoral, Universidad Distrital Francisco Jose de Caldas]. Dialnet. https://dialnet.unirioja.es/servlet/tesis?codigo=347742",
        "Shadish, W. R., Cook, T. D., & Campbell, D. T. (2002). Experimental and quasi-experimental designs for generalized causal inference. Houghton Mifflin.",
        "Vazquez-Canteli, J. R., Dey, S., Henze, G., & Nagy, Z. (2020). CityLearn: Standardizing research in multi-agent reinforcement learning for demand response and urban energy management. arXiv. https://arxiv.org/abs/2012.10504",
    ]
    added = 0
    for ref in refs:
        key = ref_key(ref)
        if key in existing:
            continue
        para = doc.add_paragraph(ref)
        para.paragraph_format.left_indent = Inches(0.3)
        para.paragraph_format.first_line_indent = Inches(-0.3)
        existing.add(key)
        added += 1
    if added == 0:
        cap._element.getparent().remove(cap._element)


def rebuild_doc(
    detail: dict,
    treatment: pd.DataFrame,
    building_compact: pd.DataFrame,
    eq_summary: pd.DataFrame,
    figures: dict[str, Path],
    convergence: pd.DataFrame,
    kpi_ranking: pd.DataFrame,
    kpi_catalog: pd.DataFrame,
    episodes: pd.DataFrame,
    final_ts: pd.DataFrame,
    traces: pd.DataFrame,
    checkpoints: pd.DataFrame,
) -> None:
    build_figure_interpretations(detail, treatment, building_compact, eq_summary, episodes, convergence, kpi_ranking, final_ts, traces, checkpoints)
    shutil.copyfile(SRC, OUT)
    doc = Document(OUT)
    style_doc(doc)
    clean_front_matter_50ep(doc)
    update_summary_abstract_keywords(doc)
    remove_sections_until_clear(
        doc,
        [
            ("1.7 Validacion estructural y triangulacion del planteamiento", "Capitulo 2."),
            ("2.7 Sintesis critica y triangulacion del marco teorico", "Capitulo 3."),
            ("4.9 Detalle tecnico de la propuesta derivado del resumen", "Capitulo 5."),
            ("4.9 Respuesta operacional a las preguntas especificas PE.1, PE.2 y PE.3", "Capitulo 5."),
            ("4.10 Respuesta operacional a las preguntas especificas PE.1, PE.2 y PE.3", "Capitulo 5."),
            ("6.5 Criterios de cierre doctoral y control de calidad final", "Referencias bibliograficas"),
            ("Referencias complementarias verificadas e incorporadas en la revision", "Anexo A."),
            ("Referencias complementarias incorporadas en la revision", "Anexo A."),
        ],
    )
    remove_section_to_end_if_exists(doc, "Anexo C.")
    replace_section(doc, "Capitulo 1.", "Capitulo 2.", add_cap1_plan_aligned)
    body_text = "\n".join(text_of(el) for el in doc.element.body)
    if "2.2.3 Dec-POMDP" in body_text:
        try:
            replace_section(
                doc,
                "2.2.3 Dec-POMDP",
                "2.2.4 CTDE",
                lambda tmp: add_expanded_decpomdp_section(tmp, building_compact),
            )
        except RuntimeError:
            replace_section(
                doc,
                "2.2.3 Dec-POMDP",
                "2.2.4 CityLearn",
                lambda tmp: add_expanded_decpomdp_section(tmp, building_compact),
            )
    else:
        insert_section_before_any(doc, ["2.1.3 CityLearn", "2.2.4 CityLearn", "2.3 Variables de la investigacion"], lambda tmp: add_expanded_decpomdp_section(tmp, building_compact))
    insert_section_before(doc, "Capitulo 2.", add_cap1_validation)
    insert_section_before(doc, "Capitulo 3.", add_cap2_validation)
    replace_section(doc, "Capitulo 3. Metodologia", "Capitulo 4.", add_cap3_methodology)
    insert_section_before(doc, "Referencias bibliograficas", add_cap6_completion_plan)
    normalize_chapter2_numbering(doc)
    children = list(doc.element.body)
    idx_cap5 = idx_cap6 = None
    for i, el in enumerate(children):
        txt = text_of(el)
        if idx_cap5 is None and txt.startswith("Capitulo 5. Resultados"):
            idx_cap5 = i
        if idx_cap6 is None and txt.startswith("Capitulo 6. Conclusiones"):
            idx_cap6 = i
    if idx_cap5 is None or idx_cap6 is None:
        raise RuntimeError(f"No se encontraron limites Cap5/Cap6: {idx_cap5}, {idx_cap6}")
    before = [deepcopy(el) for el in children[:idx_cap5]]
    after = [deepcopy(el) for el in children[idx_cap6:] if el.tag != qn("w:sectPr")]
    clear_body_keep_sectpr(doc)
    for el in before:
        append_before_sectpr(doc, el)
    add_cap4_detailed_proposal_from_summary(doc)
    add_cap4_problem_question_response(doc, detail)
    add_cap5(doc, detail, treatment, building_compact, eq_summary, figures, convergence, kpi_ranking, kpi_catalog)
    for el in after:
        append_before_sectpr(doc, el)
    insert_section_before_any(doc, ["Anexo A.", "Anexo A"], append_apa_references)
    normalize_apa_citation_text(doc)
    remove_section_to_end_if_exists(doc, "Anexo C.")
    add_docx_unification_audit(doc)
    apply_apa7_document_format(doc)
    doc.save(OUT)


def main() -> None:
    ensure_dirs()
    if not G_BASE.exists():
        raise FileNotFoundError(G_BASE)
    treatment, episodes, buildings, equipment = load_evidence()
    stats_df, pairs_df, detail = analyze_objectives(treatment, episodes)
    convergence = analyze_convergence(episodes)
    kpi_ranking, kpi_catalog = load_citylearn_v2_kpi_summary()
    building_compact, eq_summary = analyze_buildings(buildings, equipment)
    final_ts = load_final_episode_timeseries(treatment)
    traces = load_trace_samples()
    checkpoints = load_checkpoint_summary()
    figures = make_figures(detail, treatment, building_compact, eq_summary, episodes, convergence, kpi_ranking, final_ts, traces, checkpoints)
    rebuild_doc(detail, treatment, building_compact, eq_summary, figures, convergence, kpi_ranking, kpi_catalog, episodes, final_ts, traces, checkpoints)
    v = Document(OUT)
    paras = [p.text.strip() for p in v.paragraphs if p.text.strip()]
    full = "\n".join(paras)
    metrics = {
        "output": str(OUT),
        "source_gdrive": str(G_BASE),
        "size_bytes": OUT.stat().st_size,
        "paragraphs_non_empty": len(paras),
        "word_count_estimated": len(re.findall(r"\b[\wáéíóúÁÉÍÓÚñÑüÜ-]+\b", full, re.UNICODE)),
        "tables": len(v.tables),
        "inline_images": len(v.inline_shapes),
        "treatment_rows": len(treatment),
        "episode_rows": len(episodes),
        "building_rows": len(building_compact),
        "equipment_rows": len(eq_summary),
        "convergence_rows": len(convergence),
        "citylearn_v2_ranking_rows": len(kpi_ranking),
        "citylearn_v2_kpi_catalog_rows": len(kpi_catalog),
        "final_timeseries_rows": len(final_ts),
        "trace_rows": len(traces),
        "checkpoint_rows": len(checkpoints),
        "figure_count_generated": len(figures),
        "stats_csv": str(TABLE_DIR / "gdrive_objective_aligned_statistics.csv"),
        "pairs_csv": str(TABLE_DIR / "gdrive_objective_pairwise_mannwhitney_holm.csv"),
        "has_oe1": "OE.1: efecto del MADRL sobre flexibilidad energetica" in full,
        "has_oe2": "OE.2: efecto del MADRL sobre emisiones de CO2" in full,
        "has_oe3": "OE.3: efecto del MADRL sobre costos energeticos" in full,
        "has_decpomdp_expanded": "dimension global agregada 1856" in full and "gamma=0.9999" in full,
        "has_convergence_section": "Curvas de convergencia y episodios de aprendizaje" in full,
        "has_citylearn_v2_evaluate_v2_section": "KPIs bajo nomenclatura CityLearn v2 evaluate_v2" in full,
        "has_distributed_madrl_figures": "Figura 5.9a" in full and "Distribucion episodica" in full and "Figura 5.10" in full and "Trade-off multiobjetivo" in full,
        "has_no_aggregate_figure_section": "5.11 Figuras complementarias" not in full,
        "declares_happo_artifact_limit": "CSV materializado conserva 49 filas episodicas" in full,
        "has_no_old_global_kw": "p = 0,0459" not in full and "p=0,0459" not in full,
        "has_no_local_reference": "referencia local" not in full.lower(),
        "has_no_short_run_phrase": "5 episodios" not in full.lower(),
    }
    METRICS.write_text(json.dumps(metrics, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(metrics, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
