#!/usr/bin/env python3
"""Alinea Resumen/Abstract del Informe al veredicto KPI-gains (Friedman p=0,0096)."""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

REPO = Path(__file__).resolve().parents[2]
INFORME = REPO / "docs" / "Inforne_tesisV4_FINAL_REVISADO_50_EPISODIOS.docx"
BACKUP_DIR = REPO / "outputs" / "_word_backups"

SUBS = [
    (
        "Friedman global no alcanzó significancia (p=0,085801)",
        "H0G se rechaza de forma exploratoria (Friedman integración p = 0,0096) "
        "mientras KW ALL no es significativo (p = 0,1554); HE11/HE21/HE31 no se respaldan",
    ),
    (
        "while the global Friedman test was not significant (p=0.085801)",
        "while H0G is rejected exploratorily (integration Friedman p = 0.0096) "
        "and KW ALL is non-significant (p = 0.1554); HE11/HE21/HE31 are not supported",
    ),
]


def set_run_font(run) -> None:
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def main() -> int:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(INFORME, BACKUP_DIR / f"{INFORME.stem}_{stamp}{INFORME.suffix}")
    doc = Document(str(INFORME))
    n = 0
    for p in doc.paragraphs:
        t = p.text or ""
        new = t
        for old, repl in SUBS:
            if old in new:
                new = new.replace(old, repl)
        if new != t:
            p.clear()
            run = p.add_run(new)
            set_run_font(run)
            n += 1
    doc.save(str(INFORME))
    text = "\n".join(x.text for x in Document(str(INFORME)).paragraphs)
    print(
        {
            "patched": n,
            "leftover_085801": ("0,085801" in text or "0.085801" in text),
            "has_0096": ("0,0096" in text or "0.0096" in text),
        }
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
