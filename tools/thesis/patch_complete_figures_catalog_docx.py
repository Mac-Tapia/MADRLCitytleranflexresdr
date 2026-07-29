#!/usr/bin/env python3
"""Integración COMPLETA (sin parciales) de figuras en los 2 Word canónicos.

Incluye:
  1) Multicriterio (Pareto / learning / degradación) con captions.
  2) Resúmenes performance_comparison por MADRL (4).
  3) Catálogo completo por job: performance_comparison + KPIs/entrenamiento
     (150 PNG: MATD3/MAAC/MASAC 14×9 + HAPPO 8×3).

Idempotente vía marcador CATALOGO_COMPLETO_FIGURAS_PERFORMANCE_50EP.
Backups: outputs/_word_backups/
"""
from __future__ import annotations

import json
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import CANONS, RUN_ID  # noqa: E402

BACKUP_DIR = REPO / "outputs" / "_word_backups"
RUN = REPO / "outputs" / RUN_ID
MC_DIR = REPO / "outputs" / "madrl_multicriteria_selection" / "figures"
PERF_DIR = RUN / "resumen_comparativo" / "performance_comparison"
REPORT = REPO / "docs" / "COMPLETE_FIGURES_CATALOG_WORD_PATCH_REPORT_2026-07-29.json"
STAMP = datetime.now().strftime("%Y%m%d_%H%M%S")
GREY = RGBColor(0x55, 0x55, 0x55)

MARKER = "CATALOGO_COMPLETO_FIGURAS_PERFORMANCE_50EP"
ALGOS = ("MATD3", "MAAC", "MASAC", "HAPPO")
SCENS = ("E1", "E2", "E3")

FULL_FIGS = (
    "performance_comparison.png",
    "convergence_returns.png",
    "reward_timeseries.png",
    "episode_reward_summary.png",
    "learning_efficiency.png",
    "agent_reward_contribution.png",
    "exploration_action_l2.png",
    "citylearn_v2_district_timeseries.png",
    "core_kpis.png",
    "OE1_flexibility_kpis.png",
    "OE2_co2_kpis.png",
    "OE3_cost_kpis.png",
    "axis_baseline_comparison.png",
    "baseline_gain_by_kpi.png",
)
HAPPO_FIGS = (
    "performance_comparison.png",
    "convergence_returns.png",
    "reward_timeseries.png",
    "episode_reward_summary.png",
    "learning_efficiency.png",
    "agent_reward_contribution.png",
    "exploration_action_l2.png",
    "citylearn_v2_district_timeseries.png",
)

FIG_EXPLAIN = {
    "performance_comparison.png": "Distrito (efecto % vs baseline de 4 MADRL) + edificio (17 edificios).",
    "convergence_returns.png": "Convergencia de retornos episódicos.",
    "reward_timeseries.png": "Serie temporal de recompensa.",
    "episode_reward_summary.png": "Resumen de recompensa por episodio.",
    "learning_efficiency.png": "Eficiencia de aprendizaje.",
    "agent_reward_contribution.png": "Contribución de recompensa por agente/edificio.",
    "exploration_action_l2.png": "Exploración (norma L2 de acciones).",
    "citylearn_v2_district_timeseries.png": "Series temporales distritales CityLearn.",
    "core_kpis.png": "KPIs núcleo del job.",
    "OE1_flexibility_kpis.png": "KPIs de flexibilidad (OE.1).",
    "OE2_co2_kpis.png": "KPIs de emisiones (OE.2).",
    "OE3_cost_kpis.png": "KPIs de costo (OE.3).",
    "axis_baseline_comparison.png": "Comparación por eje vs baseline.",
    "baseline_gain_by_kpi.png": "Ganancia/pérdida por KPI vs baseline.",
}


def set_run_font(run, *, size: float = 12.0, italic: bool = False, bold: bool = False, grey: bool = False) -> None:
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.italic = italic
    run.bold = bold
    if grey:
        run.font.color.rgb = GREY
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(anchor: Paragraph, text: str = "", *, bold: bool = False, italic: bool = False, size: float = 12.0) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        run = para.add_run(text)
        set_run_font(run, bold=bold, italic=italic, size=size, grey=italic and size <= 9.5)
    return para


def set_heading(paragraph: Paragraph, level: int) -> None:
    try:
        paragraph.style = f"Heading {level}"
    except KeyError:
        if paragraph.runs:
            set_run_font(paragraph.runs[0], bold=True, size=14.0 if level <= 2 else 12.0)


def insert_figure_after(anchor: Paragraph, path: Path, caption: str, width_cm: float = 14.0) -> Paragraph:
    pic = insert_paragraph_after(anchor)
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = insert_paragraph_after(anchor=pic, text=caption, italic=True, size=9.0)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        set_run_font(cap.runs[0], size=9.0, italic=True, grey=True)
    return cap


def catalog_for(algo: str) -> tuple[str, ...]:
    return HAPPO_FIGS if algo == "HAPPO" else FULL_FIGS


def all_required_paths() -> list[Path]:
    paths: list[Path] = [
        MC_DIR / "pareto_cost_co2_flex.png",
        MC_DIR / "learning_curves.png",
        MC_DIR / "degradation_bars.png",
    ]
    for algo in ALGOS:
        paths.append(PERF_DIR / f"{algo}_performance_comparison.png")
        for scen in SCENS:
            for name in catalog_for(algo):
                paths.append(RUN / algo / scen / "figures" / name)
    return paths


def find_anchor(doc: Document) -> Paragraph | None:
    """Insert before 5.4.2 if present; else after Tabla 5.4.1 / performance block."""
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if t.startswith("5.4.2"):
            return paras[i - 1] if i else p
    for needle in (
        "Figura 5.3h. Performance comparison",
        "Tabla 5.4.1. TOPSIS descriptivo",
        "5.4.1 OG",
    ):
        for i, p in enumerate(paras):
            if needle in (p.text or ""):
                return p
    return None


def already_complete(doc: Document) -> bool:
    return MARKER in "\n".join(p.text or "" for p in doc.paragraphs)


def patch_doc(doc_path: Path) -> dict:
    required = all_required_paths()
    missing = [str(p.relative_to(REPO)).replace("\\", "/") for p in required if not p.is_file()]
    if missing:
        return {"path": doc_path.name, "ok": False, "error": "missing_source_pngs", "missing_count": len(missing), "missing_sample": missing[:10]}

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{doc_path.stem}_antes_catalogo_completo_{STAMP}{doc_path.suffix}"
    shutil.copy2(doc_path, backup)

    doc = Document(str(doc_path))
    if already_complete(doc):
        return {
            "path": str(doc_path.relative_to(REPO)).replace("\\", "/"),
            "ok": True,
            "skipped": True,
            "reason": "marker_present",
            "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        }

    anchor = find_anchor(doc)
    if anchor is None:
        return {
            "path": str(doc_path.relative_to(REPO)).replace("\\", "/"),
            "ok": False,
            "error": "anchor_not_found",
            "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        }

    counts = {"multicriteria": 0, "perf_summaries": 0, "job_figures": 0}
    cursor = insert_paragraph_after(
        anchor,
        "5.4.1-bis Catálogo completo de figuras (multicriterio + performance distrito/edificio)",
        bold=True,
        size=14.0,
    )
    set_heading(cursor, 3)
    cursor = insert_paragraph_after(
        cursor,
        "Integración completa — sin parciales — de artefactos reales de la corrida "
        f"{RUN_ID} (50 episodios). Incluye multicriterio TOPSIS/AHP, resúmenes "
        "performance_comparison por MADRL y el catálogo íntegro por job "
        f"({len(required) - 7} PNG en {{ALGO}}/{{E}}/figures/ + 3 multicriterio + 4 resúmenes). "
        "Naturaleza descriptiva; no sustituye §5.3 ni §5.5 (HE/H0G).",
    )
    cursor = insert_paragraph_after(cursor, MARKER, italic=True, size=8.0)

    # --- Multicriterio ---
    cursor = insert_paragraph_after(cursor, "A. Selección multicriterio (TOPSIS/AHP) — figuras", bold=True, size=12.0)
    cursor = insert_paragraph_after(
        cursor,
        "Complemento visual de la Tabla 5.4.1 TOPSIS. Pareto, curvas de aprendizaje y "
        "barras de degradación desde outputs/madrl_multicriteria_selection/figures/.",
    )
    for fname, caption in (
        ("pareto_cost_co2_flex.png", "Figura MC.1. Multicriterio — frente de Pareto costo–CO₂–flexibilidad."),
        ("learning_curves.png", "Figura MC.2. Multicriterio — curvas de aprendizaje (50 episodios)."),
        ("degradation_bars.png", "Figura MC.3. Multicriterio — barras de degradación."),
    ):
        cursor = insert_figure_after(cursor, MC_DIR / fname, caption, width_cm=14.5)
        counts["multicriteria"] += 1

    # --- Resúmenes por algoritmo ---
    cursor = insert_paragraph_after(cursor, "B. Performance comparison por MADRL (distrito + edificio, E1–E3)", bold=True)
    cursor = insert_paragraph_after(
        cursor,
        "Una figura por algoritmo: fila superior = efecto distrital % vs baseline de los "
        "cuatro MADRL (focal resaltado); fila inferior = heterogeneidad de 17 edificios.",
    )
    for algo in ALGOS:
        cursor = insert_figure_after(
            cursor,
            PERF_DIR / f"{algo}_performance_comparison.png",
            f"Figura PC-{algo}. Performance comparison — {algo} (distrito + edificio, E1–E3).",
            width_cm=15.5,
        )
        counts["perf_summaries"] += 1

    # --- Catálogo completo por job ---
    cursor = insert_paragraph_after(cursor, "C. Catálogo completo por job (algoritmo × escenario)", bold=True)
    cursor = insert_paragraph_after(
        cursor,
        "Se embebe el inventario completo de figuras de performance/KPI por tratamiento. "
        "MATD3/MAAC/MASAC: 14 figuras × 9 jobs; HAPPO: 8 figuras × 3 jobs. "
        "Cada caption indica archivo, métrica y alcance (distrito/edificio/entrenamiento).",
    )

    fig_idx = 0
    for algo in ALGOS:
        cursor = insert_paragraph_after(cursor, f"C.{algo} — figuras de performance", bold=True, size=12.0)
        for scen in SCENS:
            cursor = insert_paragraph_after(cursor, f"{algo} / {scen}", bold=True, size=11.0)
            for name in catalog_for(algo):
                fig_idx += 1
                fig_path = RUN / algo / scen / "figures" / name
                explain = FIG_EXPLAIN.get(name, name)
                caption = (
                    f"Figura JOB.{fig_idx:03d}. {algo}/{scen} — {name} — {explain}"
                )
                width = 15.2 if name.endswith("performance_comparison.png") else 13.8
                cursor = insert_figure_after(cursor, fig_path, caption, width_cm=width)
                counts["job_figures"] += 1

    cursor = insert_paragraph_after(
        cursor,
        "Nota. Catálogo completo integrado. Fuentes: "
        "madrl_multicriteria_selection/figures/, "
        f"resumen_comparativo/performance_comparison/, "
        f"y {{ALGO}}/{{E}}/figures/ bajo outputs/{RUN_ID}/. "
        "Mapping: performance_comparison_mapping.md.",
        italic=True,
        size=9.0,
    )

    doc.save(str(doc_path))

    verify = Document(str(doc_path))
    text = "\n".join(p.text or "" for p in verify.paragraphs)
    return {
        "path": str(doc_path.relative_to(REPO)).replace("\\", "/"),
        "ok": True,
        "skipped": False,
        "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        "counts": counts,
        "required_pngs": len(required),
        "checks": {
            "marker": MARKER in text,
            "mc1": "Figura MC.1." in text,
            "mc2": "Figura MC.2." in text,
            "mc3": "Figura MC.3." in text,
            "pc_matd3": "Figura PC-MATD3." in text,
            "pc_happo": "Figura PC-HAPPO." in text,
            "job_001": "Figura JOB.001." in text,
            "job_last": f"Figura JOB.{counts['job_figures']:03d}." in text,
            "n_job_captions": text.count("Figura JOB."),
        },
    }


def main() -> int:
    report: dict = {"stamp": STAMP, "run_id": RUN_ID, "files": {}}
    ok_all = True
    for path in CANONS:
        print(f"PATCH {path.name} ...", flush=True)
        result = patch_doc(path)
        report["files"][path.name] = result
        ok_all = ok_all and bool(result.get("ok"))
        print(json.dumps(result, ensure_ascii=False), flush=True)
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print("WROTE", REPORT, flush=True)
    return 0 if ok_all else 1


if __name__ == "__main__":
    raise SystemExit(main())
