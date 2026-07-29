#!/usr/bin/env python3
"""Embebe resúmenes performance_comparison (4 MADRL) en Word canónicos Cap. 5.

Inserta tras el bloque 5.4.1 (tras Figuras 5.3d multicriterio o Tabla 5.4.1)
un subbloque explicativo + 4 PNG distrito/edificio.

Backups: outputs/_word_backups/
"""
from __future__ import annotations

import json
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import CANONS, RUN_ID  # noqa: E402

BACKUP_DIR = REPO / "outputs" / "_word_backups"
FIG_DIR = REPO / "outputs" / RUN_ID / "resumen_comparativo" / "performance_comparison"
REPORT = REPO / "docs" / "PERFORMANCE_COMPARISON_WORD_PATCH_REPORT_2026-07-29.json"
STAMP = datetime.now().strftime("%Y%m%d_%H%M%S")
GREY = RGBColor(0x55, 0x55, 0x55)
MARKER = "Figura 5.3e. Performance comparison — MATD3"

FIGURES = (
    (FIG_DIR / "MATD3_performance_comparison.png", "Figura 5.3e. Performance comparison — MATD3 (distrito + edificio, E1–E3)."),
    (FIG_DIR / "MAAC_performance_comparison.png", "Figura 5.3f. Performance comparison — MAAC (distrito + edificio, E1–E3)."),
    (FIG_DIR / "MASAC_performance_comparison.png", "Figura 5.3g. Performance comparison — MASAC (distrito + edificio, E1–E3)."),
    (FIG_DIR / "HAPPO_performance_comparison.png", "Figura 5.3h. Performance comparison — HAPPO (distrito + edificio, E1–E3)."),
)

INTRO = (
    "Performance comparison por MADRL (distrito y edificio). Cada figura resume, "
    "sobre los 50 episodios Drive, (i) el efecto primario distrital vs baseline de los "
    "cuatro algoritmos en E1/E2/E3, con el algoritmo focal resaltado, y (ii) la "
    "heterogeneidad de los 17 edificios Iquitos en el KPI del eje. Lectura: en distrito, "
    "% negativo indica empeoramiento vs baseline; en edificio, Δ negativo indica "
    "reducción local. Son resultados descriptivos; no deciden HE10–HE31."
)

NOTE = (
    "Nota. Archivos en outputs/.../performance_comparison/{ALGO}_performance_comparison.png; "
    "también existe performance_comparison.png por job en {ALGO}/{E}/figures/. "
    "Mapping: performance_comparison_mapping.md."
)


def set_run_font(run, *, size: float = 12.0, italic: bool = False, grey: bool = False) -> None:
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.italic = italic
    if grey:
        run.font.color.rgb = GREY
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(anchor: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        run = para.add_run(text)
        set_run_font(run)
    return para


def insert_figure_after(anchor: Paragraph, path: Path, caption: str, width_cm: float = 15.5) -> Paragraph:
    pic = insert_paragraph_after(anchor)
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = insert_paragraph_after(pic, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        set_run_font(cap.runs[0], size=9.0, italic=True, grey=True)
    return cap


def find_insert_anchor(doc: Document) -> Paragraph | None:
    # Prefer after last multicriteria figure note; else after Tabla 5.4.1 block.
    prefer = (
        "Figura 5.3d. Multicriterio",
        "Nota. Figuras descriptivas de outputs/madrl_multicriteria_selection",
        "Tabla 5.4.1. TOPSIS descriptivo",
    )
    paras = list(doc.paragraphs)
    for needle in prefer:
        for i, p in enumerate(paras):
            if needle in (p.text or ""):
                # if caption of 5.3d, skip following empty; if table caption, skip tbl
                el = p._p
                nxt = el.getnext()
                if needle.startswith("Tabla") and nxt is not None and nxt.tag.endswith("}tbl"):
                    after = nxt.getnext()
                    if after is not None and after.tag.endswith("}p"):
                        return Paragraph(after, p._parent)
                # walk forward past any immediately following empty/note until before 5.4.2
                cur = p
                for j in range(i, min(i + 12, len(paras))):
                    t = (paras[j].text or "").strip()
                    if t.startswith("5.4.2"):
                        return paras[j - 1] if j > 0 else p
                    cur = paras[j]
                return cur
    return None


def already_patched(doc: Document) -> bool:
    return MARKER in "\n".join(p.text or "" for p in doc.paragraphs)


def patch_doc(path: Path) -> dict:
    for fig, _ in FIGURES:
        if not fig.is_file():
            raise FileNotFoundError(fig)

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{path.stem}_antes_perf_comparison_{STAMP}{path.suffix}"
    shutil.copy2(path, backup)

    doc = Document(str(path))
    if already_patched(doc):
        return {
            "path": str(path.relative_to(REPO)).replace("\\", "/"),
            "ok": True,
            "skipped": True,
            "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        }

    anchor = find_insert_anchor(doc)
    if anchor is None:
        return {
            "path": str(path.relative_to(REPO)).replace("\\", "/"),
            "ok": False,
            "error": "insert anchor not found",
            "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        }

    cursor = insert_paragraph_after(anchor, INTRO)
    inserted = []
    for fig, caption in FIGURES:
        cursor = insert_figure_after(cursor, fig, caption)
        inserted.append(caption)
    note = insert_paragraph_after(cursor, NOTE)
    if note.runs:
        set_run_font(note.runs[0], size=9.0, italic=True)

    doc.save(str(path))
    text = "\n".join(p.text or "" for p in Document(str(path)).paragraphs)
    return {
        "path": str(path.relative_to(REPO)).replace("\\", "/"),
        "ok": True,
        "skipped": False,
        "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        "inserted": inserted,
        "checks": {cap.split("—")[0].strip(): (cap in text) for _, cap in FIGURES},
    }


def main() -> int:
    report: dict = {"stamp": STAMP, "files": {}}
    ok_all = True
    for path in CANONS:
        result = patch_doc(path)
        report["files"][path.name] = result
        ok_all = ok_all and bool(result.get("ok"))
        print(json.dumps(result, ensure_ascii=False))
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print("WROTE", REPORT)
    return 0 if ok_all else 1


if __name__ == "__main__":
    raise SystemExit(main())
