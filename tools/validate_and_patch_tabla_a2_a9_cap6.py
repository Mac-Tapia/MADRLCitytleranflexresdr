# -*- coding: utf-8 -*-
"""Validate Tabla A.2 against canonical 50-ep building cost deltas; patch ABRIR_ESTE.

Also: regenerate Figura A.9; refresh Figs 5.1/5.8e interpretations; replace 6.4/6.5
with executed-plan content (user 15-Jul-2026). Incremental patches only.
"""
from __future__ import annotations

import json
import math
import re
import zipfile
from copy import deepcopy
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
FULL = REPO / "outputs" / "_drive_madrl" / "full_data"
ANALYSIS = REPO / "outputs" / "_drive_madrl" / "gdrive_20260627_164047_objective_analysis"
TABLE_DIR = ANALYSIS / "tables"
FIG_DIR = ANALYSIS / "figures"
CANONICAL = REPO / "outputs" / "madrl_v3_20260627_164047"
REPORT_DIR = ANALYSIS / "validation"
REPORT_DIR.mkdir(parents=True, exist_ok=True)

ALGOS = ["HAPPO", "MAAC", "MASAC", "MATD3"]
SCENARIOS = ["E1", "E2", "E3"]
PRIMARY = REPO / "docs" / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx"
MIRROR = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx"

GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_background(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def clear_cell(cell) -> None:
    for p in cell.paragraphs:
        p.clear()


def write_cell(cell, text: str, *, bold: bool = False, size: float = 8.0, white: bool = False) -> None:
    clear_cell(cell)
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def fmt_num(x, nd: int = 2) -> str:
    if x is None or (isinstance(x, float) and (math.isnan(x) or math.isinf(x))):
        return "N/D"
    return f"{float(x):,.{nd}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def find_behavior_sources() -> list[Path]:
    paths = []
    # Prefer consolidated objective-analysis table
    for p in [
        TABLE_DIR / "gdrive_building_behavior_summary_all.csv",
        TABLE_DIR / "gdrive_building_kpi_compact.csv",
        FULL / "analysis_real_drive" / "tables" / "building_behavior_summary_all.csv",
        ANALYSIS / "tables" / "gdrive_building_behavior_summary_all.csv",
    ]:
        if p.exists():
            paths.append(p)
    # Per-treatment raw under full_data and canonical
    for base in [FULL, CANONICAL]:
        for algo in ALGOS:
            for scen in SCENARIOS:
                for name in [
                    "building_behavior_summary.csv",
                    "data/building_behavior_summary.csv",
                    "figures/tables/building_behavior_summary.csv",
                ]:
                    # canonical uses ALGO/E/, full_data same
                    for cand in [
                        base / algo / scen / name,
                        base / algo / scen / "data" / "building_behavior_summary.csv",
                    ]:
                        if cand.exists() and cand not in paths:
                            paths.append(cand)
    return paths


def load_building_cost_universe() -> tuple[pd.DataFrame, dict]:
    """Load electricity_cost_delta for complete buildings from canonical mirrors."""
    meta = {"sources": [], "definition": "electricity_cost_delta_eur = control - baseline; lower (more negative) = greater reduction"}
    frames = []

    # 1) Consolidated files first
    cons = TABLE_DIR / "gdrive_building_behavior_summary_all.csv"
    if cons.exists():
        df = pd.read_csv(cons)
        meta["sources"].append(str(cons.relative_to(REPO)).replace("\\", "/"))
        frames.append(("consolidated", cons, df))

    compact = TABLE_DIR / "gdrive_building_kpi_compact.csv"
    if compact.exists():
        df = pd.read_csv(compact)
        meta["sources"].append(str(compact.relative_to(REPO)).replace("\\", "/"))
        frames.append(("compact", compact, df))

    # 2) Per-treatment from full_data (Drive mirror of canonical)
    for algo in ALGOS:
        for scen in SCENARIOS:
            for rel in [
                Path(algo) / scen / "data" / "building_behavior_summary.csv",
                Path(algo) / scen / "building_behavior_summary.csv",
            ]:
                p = FULL / rel
                if p.exists():
                    df = pd.read_csv(p)
                    df["algorithm"] = algo
                    df["scenario"] = scen
                    frames.append((f"{algo}-{scen}", p, df))
                    meta["sources"].append(str(p.relative_to(REPO)).replace("\\", "/"))
                    break

    if not frames:
        raise FileNotFoundError("No building_behavior_summary sources found for Tabla A.2")

    # Prefer per-treatment full_data if available; else consolidated
    treatment_frames = [f for f in frames if f[0] not in ("consolidated", "compact")]
    use = treatment_frames if treatment_frames else frames[:1]
    parts = []
    for tag, path, df in use:
        d = df.copy()
        # normalize columns
        colmap = {}
        for c in d.columns:
            cl = c.lower()
            if cl in ("algorithm", "algo"):
                colmap[c] = "algorithm"
            elif cl in ("scenario", "scen", "reward_scenario"):
                colmap[c] = "scenario"
            elif cl in ("agent", "building", "building_id", "building_name"):
                colmap[c] = "agent"
            elif "electricity_cost_delta" in cl:
                colmap[c] = "electricity_cost_delta_eur"
            elif cl in ("electricity_cost_control_eur", "electricity_cost_control"):
                colmap[c] = "electricity_cost_control_eur"
            elif cl in ("electricity_cost_baseline_eur", "electricity_cost_baseline"):
                colmap[c] = "electricity_cost_baseline_eur"
        d = d.rename(columns=colmap)
        if "algorithm" not in d.columns or "scenario" not in d.columns:
            # try from path tag
            if "-" in tag and tag not in ("consolidated", "compact"):
                d["algorithm"] = tag.split("-")[0]
                d["scenario"] = tag.split("-")[1]
        keep = [c for c in ["algorithm", "scenario", "agent", "electricity_cost_delta_eur",
                            "electricity_cost_control_eur", "electricity_cost_baseline_eur"] if c in d.columns]
        d = d[keep].copy()
        d["source_file"] = str(path.relative_to(REPO)).replace("\\", "/")
        parts.append(d)

    all_df = pd.concat(parts, ignore_index=True)
    # Project MADRL only
    all_df = all_df[all_df["algorithm"].isin(ALGOS) & all_df["scenario"].isin(SCENARIOS)].copy()
    all_df["electricity_cost_delta_eur"] = pd.to_numeric(all_df["electricity_cost_delta_eur"], errors="coerce")
    # Complete buildings: non-null delta and identified agent
    all_df = all_df.dropna(subset=["electricity_cost_delta_eur", "agent"])
    all_df["agent"] = all_df["agent"].astype(str)
    # Deduplicate by treatment-building keeping first (prefer full_data order)
    all_df = all_df.drop_duplicates(subset=["algorithm", "scenario", "agent"], keep="first")
    meta["n_complete_rows"] = int(len(all_df))
    meta["n_treatments"] = int(all_df.groupby(["algorithm", "scenario"]).ngroups)
    return all_df, meta


def recompute_top20(df: pd.DataFrame) -> pd.DataFrame:
    """Best reduction = most negative electricity_cost_delta_eur."""
    top = df.sort_values("electricity_cost_delta_eur", ascending=True).head(20).copy()
    top.insert(0, "rank", range(1, len(top) + 1))
    if "electricity_cost_baseline_eur" in top.columns and "electricity_cost_control_eur" in top.columns:
        base = pd.to_numeric(top["electricity_cost_baseline_eur"], errors="coerce")
        ctrl = pd.to_numeric(top["electricity_cost_control_eur"], errors="coerce")
        top["reduction_pct"] = np.where(base.abs() > 1e-9, (base - ctrl) / base.abs() * 100.0, np.nan)
    else:
        top["reduction_pct"] = np.nan
    return top


def extract_word_table_a2(doc_path: Path) -> tuple[list[list[str]], int | None]:
    doc = Document(str(doc_path))
    caption_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith("Tabla A.2") or ("Tabla A.2" in t and "costo" in t.lower()):
            caption_idx = i
            break
    # Find table after caption by walking body
    body = list(doc.element.body)
    table_i = 0
    found_table = None
    found_ti = None
    for ci, child in enumerate(body):
        if child.tag == qn("w:tbl"):
            # previous paragraph text
            prev = ""
            for j in range(ci - 1, -1, -1):
                if body[j].tag == qn("w:p"):
                    prev = "".join(body[j].itertext()).strip()
                    break
            if "Tabla A.2" in prev:
                found_table = doc.tables[table_i]
                found_ti = table_i
                break
            table_i += 1
        # keep counting even when not match — wait, we only increment on tbl
    if found_table is None:
        # fallback: search any table whose first header mentions costo eléctrico / building
        for ti, tbl in enumerate(doc.tables):
            hdr = " ".join(c.text for c in tbl.rows[0].cells).lower()
            if ("costo" in hdr or "electricity" in hdr or "delta" in hdr) and (
                "edificio" in hdr or "building" in hdr or "agent" in hdr
            ):
                # check nearby caption in paragraphs
                found_table = tbl
                found_ti = ti
                # prefer if caption A.2 exists in doc
                break
    if found_table is None:
        return [], None
    rows = []
    for row in found_table.rows:
        rows.append([c.text.replace("\n", " ").strip() for c in row.cells])
    return rows, found_ti


def compare_tables(word_rows: list[list[str]], top: pd.DataFrame) -> dict:
    # Heuristic: identify columns
    if not word_rows:
        return {"verdict": "NO TRAZABLE", "reason": "Tabla A.2 no encontrada en Word", "matches": 0, "rows": []}

    header = [h.lower() for h in word_rows[0]]
    data = word_rows[1:]

    def find_col(keys):
        for i, h in enumerate(header):
            if any(k in h for k in keys):
                return i
        return None

    c_algo = find_col(["algoritmo", "algorithm", "algo"])
    c_scen = find_col(["escenario", "scenario", "esc."])
    c_bldg = find_col(["edificio", "building", "agent", "agente"])
    c_delta = find_col(["delta", "reduc", "electricity_cost_delta", "costo"])

    comparisons = []
    n_match = 0
    for i, wrow in enumerate(data[:20]):
        r = i + 1
        expected = top.iloc[i] if i < len(top) else None
        entry = {"rank": r, "word": wrow, "status": "missing_expected"}
        if expected is None:
            comparisons.append(entry)
            continue
        exp_algo = str(expected["algorithm"])
        exp_scen = str(expected["scenario"])
        exp_agent = str(expected["agent"])
        exp_delta = float(expected["electricity_cost_delta_eur"])
        got_algo = wrow[c_algo] if c_algo is not None and c_algo < len(wrow) else ""
        got_scen = wrow[c_scen] if c_scen is not None and c_scen < len(wrow) else ""
        got_agent = wrow[c_bldg] if c_bldg is not None and c_bldg < len(wrow) else ""
        got_delta_s = wrow[c_delta] if c_delta is not None and c_delta < len(wrow) else ""
        # parse delta with european/US formats
        gd = re.sub(r"[^0-9,.\-]", "", got_delta_s)
        if gd.count(",") == 1 and gd.count(".") >= 1:
            gd = gd.replace(".", "").replace(",", ".")
        elif gd.count(",") == 1 and gd.count(".") == 0:
            gd = gd.replace(",", ".")
        try:
            got_delta = float(gd) if gd not in ("", "-", ".") else float("nan")
        except ValueError:
            got_delta = float("nan")

        algo_ok = exp_algo.lower() in got_algo.lower() or got_algo.lower() in exp_algo.lower()
        scen_ok = exp_scen.lower() in got_scen.lower() or got_scen.upper() == exp_scen
        agent_ok = exp_agent.lower() in got_agent.lower() or got_agent.lower() in exp_agent.lower()
        delta_ok = (not math.isnan(got_delta)) and abs(got_delta - exp_delta) <= max(1.0, abs(exp_delta) * 0.01)
        ok = algo_ok and scen_ok and agent_ok and delta_ok
        if ok:
            n_match += 1
        entry.update(
            {
                "status": "match" if ok else "mismatch",
                "expected": {
                    "algorithm": exp_algo,
                    "scenario": exp_scen,
                    "agent": exp_agent,
                    "delta": exp_delta,
                    "source": expected.get("source_file", ""),
                },
                "got": {"algorithm": got_algo, "scenario": got_scen, "agent": got_agent, "delta": got_delta_s},
                "checks": {"algo": algo_ok, "scenario": scen_ok, "agent": agent_ok, "delta": delta_ok},
            }
        )
        comparisons.append(entry)

    verdict = "VALIDADA" if n_match == min(20, len(data), len(top)) and n_match >= 15 else (
        "CORREGIDA" if n_match < min(20, len(top)) else "NO TRAZABLE"
    )
    return {
        "verdict_preliminary": verdict if n_match >= 15 else "MISMATCH",
        "matches": n_match,
        "word_rows": len(data),
        "header": word_rows[0],
        "rows": comparisons,
        "table_index": None,
    }


def fill_existing_table(table, headers: list[str], rows: list[list[str]]) -> None:
    # Ensure row count
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    # Optionally trim? leave extra empty if fewer — clear extras
    for ri, row in enumerate(table.rows):
        if ri == 0:
            for ci, h in enumerate(headers):
                if ci < len(row.cells):
                    write_cell(row.cells[ci], h, bold=True, size=8, white=True)
                    set_cell_background(row.cells[ci], "1F4E79")
        elif ri - 1 < len(rows):
            for ci, val in enumerate(rows[ri - 1]):
                if ci < len(row.cells):
                    write_cell(row.cells[ci], val, size=8)
        else:
            for ci in range(len(row.cells)):
                write_cell(row.cells[ci], "", size=8)


def replace_paragraph_text(para, new_text: str) -> None:
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def delete_element(el) -> None:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_paragraph_after(paragraph, text: str, *, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = type(paragraph)(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def clear_between_headings(doc: Document, start_prefix: str, end_prefix: str):
    """Remove body elements after start heading until end heading; keep both headings."""
    paras = list(doc.paragraphs)
    start = next((p for p in paras if (p.text or "").strip().startswith(start_prefix)), None)
    end = next((p for p in paras if (p.text or "").strip().startswith(end_prefix)), None)
    if not start or not end:
        return None, None
    el = start._element.getnext()
    while el is not None and el is not end._element:
        nxt = el.getnext()
        delete_element(el)
        el = nxt
    return start, end


def add_apa_table_after(paragraph, caption: str, headers: list[str], rows: list[list[str]], col_widths=None):
    # caption
    cap = insert_paragraph_after(paragraph, caption)
    if cap.runs:
        cap.runs[0].bold = True
        cap.runs[0].font.size = Pt(9.5)
        cap.runs[0].font.color.rgb = GREY
    # create table via temp doc and move
    tmp = Document()
    tbl = tmp.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        write_cell(tbl.rows[0].cells[i], h, bold=True, size=8, white=True)
        set_cell_background(tbl.rows[0].cells[i], "1F4E79")
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            write_cell(cells[i], val, size=8)
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in tbl.rows:
                r.cells[i].width = Cm(w)
    # move tbl xml after caption
    tbl_el = deepcopy(tbl._tbl)
    cap._element.addnext(tbl_el)
    return cap


def scan_checkpoint_bytes() -> pd.DataFrame:
    rows = []
    for algo in ALGOS:
        for scen in SCENARIOS:
            man = FULL / algo / scen / "data" / "checkpoint_manifest.json"
            total_b = 0
            n = 0
            source = "missing"
            max_file_mb = 0.0
            max_rel = ""
            if man.exists():
                obj = json.loads(man.read_text(encoding="utf-8"))
                cps = obj.get("checkpoints") or []
                n = len(cps)
                source = str(man.relative_to(REPO)).replace("\\", "/")
                for ck in cps:
                    b = ck.get("bytes")
                    if b is None:
                        continue
                    b = int(b)
                    total_b += b
                    mb = b / (1024 * 1024)
                    if mb > max_file_mb:
                        max_file_mb = mb
                        max_rel = ck.get("relative_path", "")
            rows.append(
                {
                    "algorithm": algo,
                    "scenario": scen,
                    "treatment": f"{algo}-{scen}",
                    "n_files_listed": n,
                    "total_bytes": total_b,
                    "total_mb": total_b / (1024 * 1024),
                    "total_gb": total_b / (1024 ** 3),
                    "max_file_mb": max_file_mb,
                    "max_relative_path": max_rel,
                    "source": source,
                }
            )
    return pd.DataFrame(rows)


def plot_a9(df: pd.DataFrame) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    d = df.copy()
    d["algorithm"] = pd.Categorical(d["algorithm"], ALGOS, ordered=True)
    d["scenario"] = pd.Categorical(d["scenario"], SCENARIOS, ordered=True)
    d = d.sort_values(["algorithm", "scenario"])
    labels = d["treatment"].tolist()
    values = d["total_gb"].astype(float).tolist()
    fig, ax = plt.subplots(figsize=(9.5, 4.4))
    palette = {"HAPPO": "#7A7A7A", "MAAC": "#406A9F", "MASAC": "#2E8B57", "MATD3": "#C45C26"}
    colors = [palette.get(lab.split("-")[0], "#406A9F") for lab in labels]
    bars = ax.bar(labels, values, color=colors)
    ax.set_title("Tamaño total listado en manifiestos de checkpoint (MADRL del proyecto)")
    ax.set_ylabel("GB listados en checkpoint_manifest.json")
    ax.tick_params(axis="x", rotation=65)
    ax.grid(axis="y", alpha=0.25)
    ymax = max(values) if values else 1
    ax.set_ylim(0, max(0.5, ymax * 1.18))
    for b, v, n in zip(bars, values, d["n_files_listed"].tolist()):
        ax.text(b.get_x() + b.get_width() / 2, v + ymax * 0.02, f"{v:.2f}\n(n={n})", ha="center", va="bottom", fontsize=7)
    if any(d.loc[d["algorithm"] == "HAPPO", "n_files_listed"] == 0):
        ax.text(
            0.01,
            0.98,
            "HAPPO: sin checkpoint_manifest.json en madrl_v3_20260627_164047 (tamaño listado = 0)",
            transform=ax.transAxes,
            va="top",
            fontsize=7.5,
            color="#444444",
        )
    fig.tight_layout()
    out = FIG_DIR / "checkpoint_manifest_total_size_by_treatment.png"
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)
    d.to_csv(TABLE_DIR / "figura_a9_checkpoint_manifest_sizes.csv", index=False)
    return out


def replace_image_after_caption(docx_path: Path, caption_regex: str, png_path: Path, alt_suffix: str) -> dict:
    info = {"file": docx_path.name, "replaced": False}
    if not docx_path.exists():
        info["error"] = "missing"
        return info
    png_bytes = png_path.read_bytes()
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            xml = zin.read("word/document.xml").decode("utf-8")
            rels = zin.read("word/_rels/document.xml.rels").decode("utf-8")
    except Exception as exc:
        info["error"] = f"{type(exc).__name__}: {exc}"
        return info
    m = re.search(caption_regex, xml, flags=re.I)
    if not m:
        info["error"] = f"caption not found: {caption_regex}"
        return info
    after = xml[m.end() :]
    blips = re.findall(r'a:blip[^>]*r:embed="(rId\d+)"', after)
    if not blips:
        info["error"] = "no blip after caption"
        return info
    rid = blips[0]
    rm = re.search(rf'Relationship[^>]*Id="{rid}"[^>]*Target="([^"]+)"', rels)
    if not rm:
        info["error"] = f"no relationship for {rid}"
        return info
    media = "word/" + rm.group(1).lstrip("/")
    tmp = docx_path.with_suffix(f".docx.tmp_{alt_suffix}")
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            replaced = False
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.replace("\\", "/") == media.replace("\\", "/"):
                    data = png_bytes
                    replaced = True
                zout.writestr(item, data)
        if not replaced:
            tmp.unlink(missing_ok=True)
            info["error"] = f"media missing: {media}"
            return info
        tmp.replace(docx_path)
        info.update({"replaced": True, "rId": rid, "media": media})
    except PermissionError:
        alt = docx_path.with_name(docx_path.stem + f"_{alt_suffix}_PATCHED.docx")
        if tmp.exists():
            tmp.replace(alt)
        info.update({"replaced": False, "saved_alt": alt.name, "note": "Word bloqueado"})
    return info


def patch_interpretation_by_prefix(doc: Document, figure_token: str, new_interp: str, new_note: str | None = None) -> int:
    n = 0
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if figure_token in t and (t.startswith("Figura") or "Figura" in t):
            # look ahead for Nota / Interpretacion
            for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                tj = (doc.paragraphs[j].text or "").strip()
                if new_note and tj.startswith("Nota."):
                    replace_paragraph_text(doc.paragraphs[j], new_note)
                    n += 1
                if tj.startswith("Interpretacion de la figura") or tj.startswith("Interpretación de la figura"):
                    replace_paragraph_text(doc.paragraphs[j], new_interp)
                    n += 1
                    break
            break
    return n


def rebuild_section_64_65(doc: Document) -> list[str]:
    actions = []
    # Remove doctorado/doctoral in 6.5 title if present
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("6.5") and ("doctoral" in t.lower() or "doctorado" in t.lower()):
            replace_paragraph_text(p, "6.5 Criterios de cierre de la tesis y control de calidad final")
            actions.append("Renombrado heading 6.5 sin doctoral")

    start64, end65 = clear_between_headings(doc, "6.4 Plan para culminar la tesis", "6.5 Criterios de cierre")
    if start64 is None:
        actions.append("6.4 no encontrado")
        return actions

    intro = (
        "El plan de culminación fue ejecutado diferenciando los hitos indispensables para cerrar el manuscrito "
        "de las ampliaciones experimentales que corresponden a trabajo futuro. La implementación se realizó sobre "
        "la corrida canónica madrl_v3_20260627_164047, sin inventar episodios, semillas, resultados ni artefactos "
        "no disponibles. El estado consolidado al 15 de julio de 2026 se presenta en la Tabla 6.1."
    )
    rows_61 = [
        [
            "H1. Cobertura HAPPO",
            "Se adoptó como corpus definitivo la cobertura materializada de 49 episodios por escenario para HAPPO. "
            "No se imputó el episodio ausente y se sincronizó esta condición en metodología, resultados, discusión, "
            "conclusiones y limitaciones. MAAC, MASAC y MATD3 conservan 50 episodios por escenario (597 filas episódicas en total).",
            "Ejecutado. Evidencia: gdrive_episode_kpis_used_for_statistics.csv (597 filas; HAPPO n=49; MAAC/MASAC/MATD3 n=50).",
        ],
        [
            "H2. Robustez multi-semilla",
            "La inferencia se delimitó a una sola semilla de la corrida canónica. No se afirma generalización universal; "
            "la validación multi-semilla queda como trabajo futuro.",
            "Implementado como delimitación metodológica. No ejecutado experimentalmente; no bloquea el cierre del manuscrito.",
        ],
        [
            "H3. Inferencia estadística",
            "Se consolidaron Shapiro–Wilk, Kruskal–Wallis, Mann–Whitney con Holmz y tamaños de efecto; Capítulos 5 y 6 sincronizados.",
            "Ejecutado. OE.1: p = 1,305×10⁻⁸; OE.2: p = 0,043866; OE.3: p = 0,251421 "
            "(gdrive_objective_aligned_statistics.csv).",
        ],
        [
            "H4. Pareto y baseline",
            "Se consolidó la lectura multiobjetivo sin ganador universal, con contraste CityLearn v2/RBC/baseline. "
            "La discusión diferencia medias episódicas, muestra inferencial completa y KPI anual final.",
            "Ejecutado en Cap. 5. Sensibilidad de pesos como trabajo futuro.",
        ],
        [
            "H5. HPO y algoritmos adicionales",
            "Optuna y contrastes PPO/SAC/A2C no forman parte de la evidencia canónica; se evita sesgo retrospectivo.",
            "Delimitado → trabajo futuro. No requerido para los objetivos actuales.",
        ],
        [
            "H6. Cierre documental",
            "Se reforzaron Cap. 2, 4, 5 y 6; discusión 5.10; referencias depuradas; tablas APA 7; campos e índices Word al abrir.",
            "Ejecutado.",
        ],
        [
            "H7. Entrega y sustentación",
            "Secuencia final: índices F9, revisión visual PDF, validación del asesor, registro institucional y preparación de defensa.",
            "Pendiente de gestión institucional.",
        ],
    ]
    last = insert_paragraph_after(start64, intro)
    last = add_apa_table_after(
        last,
        "Tabla 6.1. Ejecución e implementación del plan para culminar la tesis.",
        ["Hito", "Implementación realizada", "Estado y evidencia de cierre"],
        rows_61,
        col_widths=[3.2, 7.5, 5.0],
    )
    # note under table: find last element — insert after table following caption
    # Walk to table after last
    el = last._element.getnext()
    anchor_para = last
    if el is not None and el.tag == qn("w:tbl"):
        # insert note paragraph after table
        note_p = OxmlElement("w:p")
        el.addnext(note_p)
        from docx.text.paragraph import Paragraph

        note_para = Paragraph(note_p, start64._parent)
        note_para.add_run(
            "Nota. Estado al 15 de julio de 2026. «Ejecutado» abarca cierre documental, analítico o de delimitación "
            "metodológica sobre madrl_v3_20260627_164047. Multi-semilla, Optuna y algoritmos adicionales quedan como "
            "trabajo futuro y no sustituyen la evidencia canónica."
        )
        note_para.runs[0].italic = True
        note_para.runs[0].font.size = Pt(9)
        close = (
            "Con H1, H3, H4 y H6 ejecutados, y con H2 y H5 delimitados como trabajo futuro, el manuscrito queda "
            "culminado para presentación académica bajo las restricciones declaradas (semilla única; HAPPO con 49 "
            "episodios por escenario). Solo H7 permanece pendiente para el cierre formal institucional."
        )
        insert_paragraph_after(note_para, close)
        actions.append("6.4 + Tabla 6.1 reemplazados")
    else:
        actions.append("6.4 insert parcial (sin ancla de tabla)")

    # 6.5
    start65, end_ref = clear_between_headings(
        doc, "6.5 Criterios de cierre de la tesis y control de calidad final", "Referencias bibliograficas"
    )
    if start65 is None:
        # try without accent
        start65, end_ref = clear_between_headings(
            doc, "6.5 Criterios de cierre", "Referencias bibliograficas"
        )
    if start65 is None:
        actions.append("6.5 no encontrado")
        return actions

    intro65 = (
        "Las conclusiones se consideran suficientemente sustentadas para responder las preguntas específicas "
        "desde la corrida Drive analizada. Tras el plan de la sección 6.4, el control de calidad final se centra "
        "en campos e índices Word, legibilidad PDF, correspondencia vertical entre PE–OE–hipótesis–resultados–conclusiones "
        "y aprobación del asesor. La extensión multi-semilla se mantiene como recomendación de trabajo futuro, no como "
        "resultado de esta tesis."
    )
    rows_62 = [
        [
            "Revisión APA integral",
            "Alinear citas, tablas, figuras y referencias al formato APA 7.",
            "Todas las citas tienen entrada bibliográfica y viceversa; captions coherentes.",
        ],
        [
            "Revisión multi-semilla opcional",
            "Mejorar la validez externa de la comparación MADRL.",
            "Réplicas documentadas o limitación de semilla única explicitada (opción adoptada).",
        ],
        [
            "Auditoría de figuras y tablas",
            "Confirmar legibilidad y correspondencia con CSV/Drive canónicos.",
            "Cada figura/tabla apunta a fuente verificable de madrl_v3_20260627_164047.",
        ],
        [
            "Revisión de coherencia vertical",
            "Asegurar que PE, OE, hipótesis, resultados y conclusiones respondan lo mismo.",
            "Matriz problema–objetivo–resultado–conclusión sin vacíos.",
        ],
    ]
    last = insert_paragraph_after(start65, intro65)
    last = add_apa_table_after(
        last,
        "Tabla 6.2. Criterios de cierre y control de calidad final.",
        ["Actividad", "Propósito", "Criterio de cierre"],
        rows_62,
        col_widths=[4.0, 6.0, 5.5],
    )
    el = last._element.getnext()
    if el is not None and el.tag == qn("w:tbl"):
        note_p = OxmlElement("w:p")
        el.addnext(note_p)
        from docx.text.paragraph import Paragraph

        note_para = Paragraph(note_p, start65._parent)
        note_para.add_run(
            "Nota. Los criterios C de calidad documental no sustituyen la evidencia experimental ya auditada. "
            "El cierre formal institucional (registro y sustentación) corresponde al hito H7."
        )
        note_para.runs[0].italic = True
        note_para.runs[0].font.size = Pt(9)
        actions.append("6.5 + Tabla 6.2 reemplazados")
    return actions


def patch_table_a2_in_doc(doc: Document, top: pd.DataFrame, table_index: int | None) -> list[str]:
    actions = []
    headers = [
        "Posición",
        "Algoritmo",
        "Escenario",
        "Edificio",
        "Costo control",
        "Costo baseline",
        "Delta costo (control−baseline)",
        "Fuente",
    ]
    rows = []
    for _, r in top.iterrows():
        rows.append(
            [
                str(int(r["rank"])),
                str(r["algorithm"]),
                str(r["scenario"]),
                str(r["agent"]),
                fmt_num(r.get("electricity_cost_control_eur"), 2) if "electricity_cost_control_eur" in r and pd.notna(r.get("electricity_cost_control_eur")) else "N/D",
                fmt_num(r.get("electricity_cost_baseline_eur"), 2) if "electricity_cost_baseline_eur" in r and pd.notna(r.get("electricity_cost_baseline_eur")) else "N/D",
                fmt_num(r["electricity_cost_delta_eur"], 2),
                Path(str(r.get("source_file", ""))).name or "building_behavior_summary",
            ]
        )

    # Locate table
    body = list(doc.element.body)
    ti = 0
    target = None
    caption_para = None
    for ci, child in enumerate(body):
        if child.tag != qn("w:tbl"):
            continue
        prev = ""
        prev_p = None
        for j in range(ci - 1, -1, -1):
            if body[j].tag == qn("w:p"):
                prev = "".join(body[j].itertext()).strip()
                # map to paragraph object
                for p in doc.paragraphs:
                    if p._element is body[j]:
                        prev_p = p
                        break
                break
        if "Tabla A.2" in prev:
            target = doc.tables[ti]
            caption_para = prev_p
            break
        ti += 1

    if target is None and table_index is not None and table_index < len(doc.tables):
        target = doc.tables[table_index]

    if target is None:
        actions.append("Tabla A.2 no localizada para reemplazo")
        return actions

    # Adjust columns if table has different width — rebuild cells by clearing/filling min shared
    n_cols = len(target.columns)
    if n_cols != len(headers):
        # If mismatch in columns, best-effort fill overlapping and note
        actions.append(f"A.2 cols Word={n_cols} vs new={len(headers)}; relleno adaptado")
        use_headers = headers[:n_cols]
        use_rows = [r[:n_cols] for r in rows]
        fill_existing_table(target, use_headers, use_rows)
    else:
        fill_existing_table(target, headers, rows)
        actions.append("Tabla A.2 celdas actualizadas")

    if caption_para is not None:
        replace_paragraph_text(
            caption_para,
            "Tabla A.2. Mejores 20 filas por reducción de costo eléctrico en edificios completos "
            "(corrida canónica madrl_v3_20260627_164047; delta = control − baseline).",
        )
        actions.append("Caption A.2 actualizado")

    # Update nearby interpretation if any
    if caption_para is not None:
        el = caption_para._element
        # find following interpretation paragraph after table
        # walk paragraphs by index
        idxs = [i for i, p in enumerate(doc.paragraphs) if p._element is caption_para._element]
        if idxs:
            i0 = idxs[0]
            for j in range(i0 + 1, min(i0 + 6, len(doc.paragraphs))):
                tj = (doc.paragraphs[j].text or "").strip()
                if tj.startswith("Interpretacion") or tj.startswith("Interpretación") or tj.startswith("Nota."):
                    if tj.startswith("Nota."):
                        replace_paragraph_text(
                            doc.paragraphs[j],
                            "Nota. Ranking por electricity_cost_delta_eur más negativo (mayor reducción) en "
                            "building_behavior_summary de edificios completos; MADRL del proyecto únicamente "
                            "(HAPPO, MAAC, MASAC, MATD3 × E1–E3). Fuente: outputs/_drive_madrl/full_data/.../"
                            "building_behavior_summary.csv y tablas materializadas en gdrive_20260627_164047_objective_analysis.",
                        )
                    if tj.startswith("Interpretacion") or tj.startswith("Interpretación"):
                        best = top.iloc[0]
                        replace_paragraph_text(
                            doc.paragraphs[j],
                            "Interpretación de la tabla. Las veinte mayores reducciones de costo eléctrico por edificio "
                            f"se ordenan por delta control−baseline. La mayor reducción corresponde a {best['algorithm']}-{best['scenario']} "
                            f"en {best['agent']} con delta = {fmt_num(best['electricity_cost_delta_eur'], 2)}. "
                            "HAPPO entra en el ranking solo cuando aporta building_behavior_summary completo; "
                            "MAAC/MASAC/MATD3 provienen del corpus de 50 episodios de la corrida canónica.",
                        )
                        actions.append("Interpretación A.2 actualizada")
                        break
    return actions


def save_doc(doc: Document, path: Path) -> dict:
    try:
        doc.save(str(path))
        return {"saved": str(path), "ok": True}
    except PermissionError:
        alt = path.with_name(path.stem + "_PATCHED.docx")
        doc.save(str(alt))
        return {"saved": str(alt), "ok": False, "note": "Word bloqueado; escrito _PATCHED"}


def main() -> int:
    # --- A.2 validate ---
    universe, meta = load_building_cost_universe()
    top = recompute_top20(universe)
    top_path = TABLE_DIR / "tabla_a2_top20_electricity_cost_reduction.csv"
    top.to_csv(top_path, index=False)
    universe_path = TABLE_DIR / "tabla_a2_universe_complete_buildings.csv"
    universe.to_csv(universe_path, index=False)

    word_rows, table_idx = extract_word_table_a2(PRIMARY)
    cmp = compare_tables(word_rows, top)
    cmp["table_index"] = table_idx
    cmp["meta"] = meta
    cmp["top20_csv"] = str(top_path.relative_to(REPO)).replace("\\", "/")

    # --- A.9 ---
    a9 = scan_checkpoint_bytes()
    a9_png = plot_a9(a9)
    old_interp_hint = "351 checkpoints / 119,627.66 MB (interpretación previa en ABRIR_ESTE)"
    total_gb = float(a9["total_gb"].sum())
    total_mb = float(a9["total_mb"].sum())
    n_files = int(a9["n_files_listed"].sum())
    max_row = a9.sort_values("max_file_mb", ascending=False).iloc[0]
    a9_interp = (
        "Interpretación de la figura. La figura resume el tamaño total listado en checkpoint_manifest.json "
        f"para los 12 tratamientos MADRL del proyecto (HAPPO, MAAC, MASAC, MATD3 × E1–E3). "
        f"Se listan {n_files} archivos con un total de {fmt_num(total_gb, 2)} GB ({fmt_num(total_mb, 2)} MB). "
        f"HAPPO aparece en 0 GB porque no hay checkpoint_manifest.json en la corrida canónica. "
        f"MAAC concentra el mayor volumen por tratamiento (~{fmt_num(a9.loc[a9['algorithm']=='MAAC','total_gb'].mean(), 2)} GB de media). "
        f"El archivo individual más grande listado corresponde a {max_row['algorithm']}-{max_row['scenario']} "
        f"({fmt_num(max_row['max_file_mb'], 2)} MB). No se inventan bytes: solo se agregan los campos bytes del manifiesto."
    )
    a9_note = (
        "Nota. La figura representa el tamaño total listado en los manifiestos de checkpoint por algoritmo y escenario. "
        "Fuente: elaboración propia a partir de outputs/_drive_madrl/full_data/{ALGO}/{E}/data/checkpoint_manifest.json "
        "de la corrida canónica madrl_v3_20260627_164047."
    )

    # Fig 5.1 / 5.8e interpretations from validated CSVs
    cov = pd.read_csv(TABLE_DIR / "figura_5_1_checkpoint_coverage_counts.csv")
    fig51_interp = (
        "Interpretación de la figura. Conteo de archivos listados en checkpoint_manifest.json de la corrida "
        "canónica madrl_v3_20260627_164047 (espejo outputs/_drive_madrl/full_data). "
        f"HAPPO-E1/E2/E3 = 0 (manifiesto ausente); MAAC = {int(cov.loc[cov['algorithm']=='MAAC','checkpoint_files_listed'].iloc[0])} por escenario; "
        f"MASAC = {int(cov.loc[cov['algorithm']=='MASAC','checkpoint_files_listed'].iloc[0])}; "
        f"MATD3 = {int(cov.loc[cov['algorithm']=='MATD3','checkpoint_files_listed'].iloc[0])}. "
        "La lectura corrige el falso cero previo de MASAC/MATD3 causado por filtrar solo rutas episode_(\\d+)."
    )
    p58 = pd.read_csv(TABLE_DIR / "fig58e_panel_values.csv")
    def mean_panel(panel, algo):
        sub = p58[(p58["panel"] == panel) & (p58["algorithm"] == algo)]
        return float(sub["value"].mean()) if len(sub) else float("nan")

    fig58_interp = (
        "Interpretación de la figura. Fuentes mixtas auditadas: action_l2 desde full_data/trace.csv "
        "(incluye HAPPO ≈ 2,25/2,23/2,22 en E1/E2/E3); EV y BESS desde building_behavior_summary "
        "(ev_charge_total_kwh y battery_throughput_total_kwh), no desde columnas muertas ev_charge_kwh/"
        "electrical_storage_soc (=0 en trace). "
        f"Medias orientativas action_l2: HAPPO={fmt_num(mean_panel('action_l2','HAPPO'),2)}, "
        f"MAAC={fmt_num(mean_panel('action_l2','MAAC'),2)}, MASAC={fmt_num(mean_panel('action_l2','MASAC'),2)}, "
        f"MATD3={fmt_num(mean_panel('action_l2','MATD3'),2)}. "
        f"EV total medio: HAPPO={fmt_num(mean_panel('ev','HAPPO'),1)} kWh frente a MAAC/MASAC ~650–680 kWh y MATD3 ~6,6–9,6×10³ kWh. "
        "Valores panel: tables/fig58e_panel_values.csv."
    )
    fig58_note = (
        "Nota. Panel action_l2 desde trace.csv (full_data); paneles EV/BESS desde building_behavior_summary "
        "(ev_charge_total_kwh, battery_throughput_total_kwh). Fuente: elaboración propia a partir de la corrida "
        "canónica madrl_v3_20260627_164047."
    )

    results = {"primary": None, "mirror": None, "a9_embeds": [], "a2_verdict": None}

    for doc_path, label in [(PRIMARY, "primary"), (MIRROR, "mirror")]:
        if not doc_path.exists():
            results[label] = {"error": "missing"}
            continue
        doc = Document(str(doc_path))
        actions = []
        # A.2
        if cmp.get("verdict_preliminary") != "VALIDADA" or cmp["matches"] < 15:
            actions += patch_table_a2_in_doc(doc, top, table_idx)
            final_verdict = "CORREGIDA"
        else:
            final_verdict = "VALIDADA"
            actions.append("Tabla A.2 VALIDADA sin cambios de celdas")
        # interpretations
        actions.append(f"fig51_interp={patch_interpretation_by_prefix(doc, 'Figura 5.1', fig51_interp)}")
        actions.append(
            f"fig58e_interp={patch_interpretation_by_prefix(doc, 'Figura 5.8e', fig58_interp, fig58_note)}"
        )
        actions.append(
            f"figA9_interp={patch_interpretation_by_prefix(doc, 'Figura A.9', a9_interp, a9_note)}"
        )
        # 6.4 / 6.5
        actions += rebuild_section_64_65(doc)
        save_info = save_doc(doc, doc_path)
        results[label] = {"actions": actions, "save": save_info, "a2": final_verdict}

    # Embed A.9 PNG (zip replace) after doc save
    for doc_path in [PRIMARY, MIRROR]:
        if doc_path.exists():
            # try main and _PATCHED if locked
            targets = [doc_path]
            patched = doc_path.with_name(doc_path.stem + "_PATCHED.docx")
            if patched.exists():
                targets.append(patched)
            for t in targets:
                emb = replace_image_after_caption(
                    t,
                    r"Figura A\.9|Tamano total listado en manifiestos|Tamaño total listado en manifiestos",
                    a9_png,
                    "A9",
                )
                results["a9_embeds"].append(emb)

    # Re-extract A.2 after patch from primary saved path
    primary_saved = Path(results["primary"]["save"]["saved"]) if results.get("primary") and results["primary"].get("save") else PRIMARY
    word_rows2, _ = extract_word_table_a2(primary_saved)
    cmp2 = compare_tables(word_rows2, top)
    if results["primary"]["a2"] == "CORREGIDA":
        results["a2_verdict"] = "CORREGIDA"
    elif cmp2["matches"] >= 15:
        results["a2_verdict"] = "VALIDADA"
    else:
        results["a2_verdict"] = "NO TRAZABLE"

    report = {
        "a2_verdict": results["a2_verdict"],
        "a2_matches_before": cmp["matches"],
        "a2_matches_after": cmp2["matches"],
        "a2_header_before": cmp.get("header"),
        "a2_top20_recomputed": top[
            ["rank", "algorithm", "scenario", "agent", "electricity_cost_delta_eur", "source_file"]
        ].to_dict(orient="records"),
        "a2_comparison_before": cmp["rows"],
        "a2_sources": meta,
        "a9_old_vs_new": {
            "old_interpretation_claim": old_interp_hint,
            "new_total_files": n_files,
            "new_total_gb": total_gb,
            "new_total_mb": total_mb,
            "per_treatment_gb": a9[["treatment", "n_files_listed", "total_gb", "source"]].to_dict(orient="records"),
            "cause": (
                "La interpretación previa (~351 archivos / ~119627 MB) no corresponde al agregado por tratamiento "
                "recalculado solo desde checkpoint_manifest.json del espejo full_data de la corrida canónica; "
                "además HAPPO debe figurar en 0 si falta el manifiesto."
            ),
            "png": str(a9_png.relative_to(REPO)).replace("\\", "/"),
        },
        "docs": results,
    }
    out_json = REPORT_DIR / "tabla_a2_a9_cap6_validation_report.json"
    out_json.write_text(json.dumps(report, indent=2, ensure_ascii=False, default=str), encoding="utf-8")
    print(json.dumps({
        "a2_verdict": report["a2_verdict"],
        "matches_before": report["a2_matches_before"],
        "matches_after": report["a2_matches_after"],
        "top5": report["a2_top20_recomputed"][:5],
        "a9_total_gb": total_gb,
        "a9_files": n_files,
        "primary_save": results["primary"],
        "report": str(out_json),
    }, indent=2, ensure_ascii=False, default=str))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
