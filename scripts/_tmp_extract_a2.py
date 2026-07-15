# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

p = Path(r"D:\MADRLCitytleranflexresdr\docs\ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx")
doc = Document(str(p))
lines = []
for i, para in enumerate(doc.paragraphs):
    t = (para.text or "").strip()
    if "Tabla A.2" in t or ("A.2" in t and "costo" in t.lower() and "Mejores" in t):
        lines.append(f"P{i}|{t[:300]}")
        for j in range(i, min(i + 4, len(doc.paragraphs))):
            tj = (doc.paragraphs[j].text or "").strip()
            if j != i and tj:
                lines.append(f"  P{j}|{tj[:300]}")

body = list(doc.element.body)
ti = 0
for ci, child in enumerate(body):
    if child.tag != qn("w:tbl"):
        continue
    prev = ""
    for j in range(ci - 1, -1, -1):
        if body[j].tag == qn("w:p"):
            prev = "".join(body[j].itertext()).strip()
            break
    if "Tabla A.2" in prev:
        tbl = doc.tables[ti]
        lines.append(f"TABLE idx={ti} rows={len(tbl.rows)} cols={len(tbl.columns)} caption={prev[:160]}")
        for r, row in enumerate(tbl.rows):
            cells = [c.text.replace("\n", " ")[:50] for c in row.cells]
            lines.append("R" + str(r) + "|" + " || ".join(cells))
    ti += 1

out = Path(r"D:\MADRLCitytleranflexresdr\outputs\_tmp_a2_extract.txt")
out.write_text("\n".join(lines), encoding="utf-8")
print("wrote", out, "n=", len(lines))
