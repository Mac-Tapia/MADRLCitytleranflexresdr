#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Completa huecos estructurales del informe FINAL_COMPLETA (una pasada).

No inventa metricas numericas nuevas: usa artefactos Drive ya auditados.
"""
from __future__ import annotations

import json
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "scripts"))

from generate_borrador_tesis_docx import bullet, heading, p, style_base  # noqa: E402
from thesis_doctoral_sections import RUN_ID, verify_doctoral_docx  # noqa: E402

OUT = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx"
AUDIT = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA_audit_estructura.json"
DISTRICT = (
    REPO
    / "outputs"
    / "_drive_madrl"
    / "full_data"
    / "analysis_real_drive"
    / "tables"
    / "district_summary_by_algorithm_scenario.csv"
)


def _para_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def _body_children(doc: Document) -> list:
    return [c for c in doc.element.body if c.tag != qn("w:sectPr")]


def _find_para(doc: Document, predicate) -> Paragraph | None:
    for para in doc.paragraphs:
        if predicate((para.text or "").strip()):
            return para
    return None


def _set_heading_text(para: Paragraph, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _insert_after(anchor: Paragraph, new_elements: list) -> None:
    el = anchor._element
    for new_el in reversed(new_elements):
        el.addnext(new_el)


def _make_temp_paras(factory) -> list:
    tmp = Document()
    style_base(tmp)
    factory(tmp)
    return [deepcopy(c) for c in _body_children(tmp)]


def _next_sibling_is_heading(para: Paragraph, prefixes: tuple[str, ...]) -> bool:
    el = para._element.getnext()
    while el is not None:
        if el.tag == qn("w:p"):
            txt = _para_text(el)
            if txt:
                return any(txt.startswith(px) for px in prefixes)
            el = el.getnext()
            continue
        break
    return False


def _section_already_has_body(para: Paragraph, stop_prefixes: tuple[str, ...]) -> bool:
    el = para._element.getnext()
    while el is not None:
        if el.tag != qn("w:p"):
            return True  # table or other content
        txt = _para_text(el)
        if not txt:
            el = el.getnext()
            continue
        if any(txt.startswith(px) for px in stop_prefixes):
            return False
        return True
    return False


def load_district_facts() -> dict:
    facts = {
        "run_id": RUN_ID,
        "algorithms": [],
        "best_reward": None,
        "lowest_cost": None,
        "lowest_emission": None,
    }
    if not DISTRICT.is_file():
        return facts
    import csv

    rows = list(csv.DictReader(DISTRICT.open(encoding="utf-8")))
    facts["algorithms"] = sorted({r["algorithm"] for r in rows})
    if not rows:
        return facts

    def pick(key, reverse=False):
        r = sorted(rows, key=lambda x: float(x[key]), reverse=reverse)[0]
        return {
            "algorithm": r["algorithm"],
            "scenario": r["scenario"],
            "value": float(r[key]),
            "episodes": int(float(r["episodes"])),
        }

    facts["best_reward"] = pick("reward_mean", reverse=True)
    facts["lowest_cost"] = pick("district_cost_mean")
    facts["lowest_emission"] = pick("district_emission_mean")
    return facts


def fix_cap2_numbering(doc: Document, actions: list[str]) -> None:
    renames = [
        # Rename deepest / higher numbers first to avoid collisions.
        (
            "2.5.1 Definicion de terminos y delimitaciones conceptuales",
            "2.6.1 Definicion de terminos y delimitaciones conceptuales",
        ),
        ("2.5.2 Posicion teorica de la tesis", "2.6.2 Posicion teorica de la tesis"),
        (
            "2.4.3 Sintesis critica de antecedentes y brecha cientifica",
            "2.5.3 Sintesis critica de antecedentes y brecha cientifica",
        ),
        ("2.4.2 Antecedentes nacionales y peruanos", "2.5.2 Antecedentes nacionales y peruanos"),
        ("2.4.1 Antecedentes internacionales", "2.5.1 Antecedentes internacionales"),
        (
            "2.3.5 Aportes fisicos al motor como base teorica de CityLearn v3 propuesto",
            "2.4.5 Aportes fisicos al motor como base teorica de CityLearn v3 propuesto",
        ),
        ("2.3.4 Algoritmos MADRL evaluados", "2.4.4 Algoritmos MADRL evaluados"),
        (
            "2.3.3 Costos energeticos, precios dinamicos y respuesta economica",
            "2.4.3 Costos energeticos, precios dinamicos y respuesta economica",
        ),
        (
            "2.3.2 Emisiones de carbono y control consciente de intensidad de carbono",
            "2.4.2 Emisiones de carbono y control consciente de intensidad de carbono",
        ),
        ("2.3.1 Flexibilidad energetica", "2.4.1 Flexibilidad energetica"),
        ("2.2.2 Variable dependiente (VD)", "2.3.2 Variable dependiente (VD)"),
        ("2.2.1 Variable independiente (VI)", "2.3.1 Variable independiente (VI)"),
        ("2.1.3 CityLearn y simulacion multiobjetivo", "2.2.3 CityLearn y simulacion multiobjetivo"),
        ("2.1.2 Formalizacion matematica Dec-POMDP", "2.2.2 Formalizacion matematica Dec-POMDP"),
        ("2.1.1 Aprendizaje por refuerzo y MADRL", "2.2.1 Aprendizaje por refuerzo y MADRL"),
    ]
    for old, new in renames:
        para = _find_para(doc, lambda t, o=old: t.startswith(o))
        if para and para.text.strip() != new:
            _set_heading_text(para, new)
            actions.append(f"Renumbered H3: {old} -> {new}")


def strengthen_sections(doc: Document, facts: dict, actions: list[str]) -> None:
    # ---- 1.4 Matriz ----
    h14 = _find_para(doc, lambda t: t.startswith("1.4 Matriz"))
    if h14:
        words_14 = 0
        el = h14._element.getnext()
        while el is not None:
            if el.tag == qn("w:p"):
                txt = _para_text(el)
                if txt.startswith("1.5"):
                    break
                words_14 += len(txt.split())
            el = el.getnext()
        if words_14 < 40:
            def _m(tmp):
                p(
                    tmp,
                    "La matriz de consistencia alinea el problema de investigacion con el objetivo "
                    "general (OG), los objetivos especificos OE.1–OE.3, las hipotesis HG/HE.1–HE.3 "
                    "y las dimensiones de la variable dependiente (D-VD.1 flexibilidad, D-VD.2 "
                    "emisiones de CO2, D-VD.3 costos energeticos). La operacionalizacion se traduce "
                    "en KPIs CityLearn v2 (peak, ramping, carbon_emissions, electricity_cost y "
                    "agregados doctorales) bajo el diseno factorial 4 algoritmos x 3 escenarios "
                    f"(corrida canonica {facts['run_id']}).",
                )
                p(
                    tmp,
                    "La tabla de consistencia y la matriz de operacionalizacion (si aparecen "
                    "inmediatamente despues de este paragrafo o en anexos) constituyen el puente "
                    "formal entre Capitulo 1 y las secciones de resultados 5.3–5.5 y 5.9.",
                )

            _insert_after(h14, _make_temp_paras(_m))
            actions.append("Completado cuerpo narrativo 1.4 Matriz de consistencia")

    # ---- 3.1 Tipo ----
    h31 = _find_para(doc, lambda t: t.startswith("3.1 Tipo"))
    if h31:
        body = []
        el = h31._element.getnext()
        while el is not None and el.tag == qn("w:p"):
            txt = _para_text(el)
            if txt.startswith("3.2"):
                break
            if txt:
                body.append(txt)
            el = el.getnext()
        if sum(len(x.split()) for x in body) < 80:
            def _t(tmp):
                p(
                    tmp,
                    "La investigacion es de tipo aplicado, con enfoque cuantitativo y alcance "
                    "explicativo-comparativo: se manipula la variable independiente (algoritmo "
                    "MADRL y escenario de ponderacion) y se observan efectos sobre KPIs energeticos "
                    "en un entorno de simulacion controlado (CityLearn). El estudio no constituye "
                    "un experimento de campo sobre la red electrica fisica de Iquitos; su validez "
                    "externa queda acotada al SEAI modelado y al protocolo reproducible de la "
                    f"corrida {facts['run_id']}.",
                )

            # Insert only if first body is very short: append after heading when thin
            if not body or sum(len(x.split()) for x in body) < 70:
                # insert after existing body end, or after heading
                anchor = h31
                el2 = h31._element.getnext()
                last = h31
                while el2 is not None and el2.tag == qn("w:p"):
                    txt = _para_text(el2)
                    if txt.startswith("3.2"):
                        break
                    last = Paragraph(el2, doc)
                    el2 = el2.getnext()
                if sum(len(x.split()) for x in body) < 70:
                    _insert_after(last if body else h31, _make_temp_paras(_t))
                    actions.append("Ampliado 3.1 Tipo de investigacion")

    # ---- 3.6 Procedimiento ----
    h36 = _find_para(doc, lambda t: t.startswith("3.6 Procedimiento"))
    if h36 and sum(
        len(_para_text(el).split())
        for el in [h36._element.getnext()]
        if el is not None and el.tag == qn("w:p") and _para_text(el) and not _para_text(el).startswith("Capitulo")
    ) < 100:
        # Measure full body words
        words = 0
        el = h36._element.getnext()
        while el is not None:
            if el.tag == qn("w:p"):
                txt = _para_text(el)
                if txt.startswith("Capitulo 4") or txt.startswith("4."):
                    if txt.startswith("Capitulo 4"):
                        break
                if txt.startswith("Capitulo"):
                    break
                words += len(txt.split())
            el = el.getnext()
            # stop at next H1 only — Cap4 is H1
        # recompute cleaner
        words = 0
        el = h36._element.getnext()
        while el is not None:
            if el.tag == qn("w:p"):
                txt = _para_text(el)
                if txt.startswith("Capitulo 4"):
                    break
                words += len(txt.split())
            el = el.getnext()
        if words < 120:
            def _proc(tmp):
                p(
                    tmp,
                    "Procedimiento experimental (resumen operativo): (1) preparar el dataset "
                    "citylearn_iquitos_2023_2025 y el schema de 17 edificios; (2) fijar semilla 0 "
                    "y la funcion de recompensa multiobjetivo unified_comparable_v4 por escenario "
                    "E1/E2/E3; (3) entrenar las 12 corridas (HAPPO, MASAC, MATD3, MAAC x E1–E3) "
                    "mediante el launcher oficial two_phase_happo_masac_v3; (4) exportar "
                    "timeseries.csv, trace.csv, KPIs y checkpoints; (5) integrar artefactos "
                    f"Drive de la corrida {facts['run_id']}; (6) calcular rankings "
                    "multiobjetivo, baseline CityLearn v2 y pruebas no parametricas; "
                    "(7) redactar contrastacion PE.1–PE.3 en el Capitulo 5.",
                )

            # append at end of section
            last = h36
            el = h36._element.getnext()
            while el is not None:
                if el.tag == qn("w:p") and _para_text(el).startswith("Capitulo 4"):
                    break
                if el.tag == qn("w:p"):
                    last = Paragraph(el, doc)
                el = el.getnext()
            _insert_after(last, _make_temp_paras(_proc))
            actions.append("Ampliado 3.6 Procedimiento experimental")

    # ---- 4.1 Sistema / Arquitectura as H3 ----
    h41 = _find_para(doc, lambda t: t.startswith("4.1 Desarrollo del sistema"))
    if h41:
        already = any(
            (para.text or "").strip().startswith("4.1.1") for para in doc.paragraphs
        )
        if not already:
            def _sys(tmp):
                heading(tmp, "4.1.1 Sistema propuesto", 3)
                p(
                    tmp,
                    "El sistema doctoral articula: (a) el dataset real de Iquitos 2023–2025 "
                    "con 17 edificios institucionales/comerciales; (b) el simulador CityLearn "
                    "v2 con extensiones propias (CityLearn v3 propuesto); (c) cuatro backends "
                    "MADRL (HAPPO, MASAC, MATD3, MAAC) bajo esquema CTDE; (d) un pipeline de "
                    "evaluacion multiobjetivo alineado a OE.1–OE.3; y (e) la generacion de "
                    "evidencia reproducible (KPIs, figuras Drive, checkpoints).",
                )
                heading(tmp, "4.1.2 Arquitectura de software", 3)
                p(
                    tmp,
                    "La arquitectura sigue capas desacopladas: simulacion y recompensa; "
                    "adaptador de entrenamiento comun (citylearn_v3_training_common.py); "
                    "backends externos (HARL/HAPPO, MADRL-MASAC, off-policy/MATD3, MAAC); "
                    "orquestacion Colab A100 (two_phase_happo_masac_v3); y postproceso de "
                    "analisis (resumen_comparativo, analysis_real_drive). Esta separacion "
                    "permite repetir las 12 corridas sin alterar el schema de observaciones "
                    "y acciones de los 17 agentes.",
                )

            _insert_after(h41, _make_temp_paras(_sys))
            actions.append("Insertado 4.1.1 Sistema y 4.1.2 Arquitectura")

    # ---- 4.4 Algoritmos prose ----
    h44 = _find_para(doc, lambda t: t.startswith("4.4 Algoritmos"))
    if h44:
        words = 0
        el = h44._element.getnext()
        while el is not None:
            if el.tag == qn("w:p"):
                txt = _para_text(el)
                if txt.startswith("4.5"):
                    break
                words += len(txt.split())
            el = el.getnext()
        if words < 120:
            def _alg(tmp):
                p(
                    tmp,
                    "Los cuatro algoritmos implementan politicas cooperativas sobre el mismo "
                    "espacio Dec-POMDP de 17 agentes, con critica centralizada en entrenamiento "
                    "y actor descentralizado en ejecucion. HAPPO (on-policy, ventajas "
                    "heterogeneas) y MASAC (off-policy, entropia maxima) se entrenan en la "
                    "Fase 1 del protocolo two_phase; MATD3 (actor-critico twin delayed) y MAAC "
                    "(atencion multiagente) en la Fase 2. La comparacion experimental fija "
                    "dataset, semilla, horizontes y schema de recompensa; solo cambia el "
                    "backend y sus hiperparametros canonicos reportados en las tablas "
                    "siguientes.",
                )

            _insert_after(h44, _make_temp_paras(_alg))
            actions.append("Ampliado texto introductorio 4.4 Algoritmos")

    # ---- 4.6 Aportes ----
    h46 = _find_para(doc, lambda t: t.startswith("4.6 Aportes"))
    if h46:
        words = 0
        el = h46._element.getnext()
        while el is not None:
            if el.tag == qn("w:p"):
                txt = _para_text(el)
                if txt.startswith("4.7"):
                    break
                words += len(txt.split())
            el = el.getnext()
        if words < 80:
            def _ap(tmp):
                p(
                    tmp,
                    "Estas extensiones habilitan el control multiobjetivo del caso Iquitos "
                    "(flexibilidad, carbono y costo) sin romper la API de evaluacion CityLearn "
                    "v2, y son la base metodologica que diferencia el CityLearn v3 propuesto "
                    "respecto del motor stock usado en antecedentes internacionales.",
                )

            last = h46
            el = h46._element.getnext()
            while el is not None:
                if el.tag == qn("w:p") and _para_text(el).startswith("4.7"):
                    break
                if el.tag in (qn("w:p"), qn("w:tbl")):
                    last = Paragraph(el, doc) if el.tag == qn("w:p") else last
                    if el.tag == qn("w:tbl"):
                        # insert after table by finding next p after tbl via last p before 4.7
                        pass
                el = el.getnext()
            # find last element before 4.7
            prev = h46._element
            el = h46._element.getnext()
            while el is not None:
                if el.tag == qn("w:p") and _para_text(el).startswith("4.7"):
                    break
                prev = el
                el = el.getnext()
            from docx.oxml import OxmlElement

            # If prev is table, insert after it
            if prev.tag == qn("w:tbl"):
                elems = _make_temp_paras(_ap)
                for new_el in reversed(elems):
                    prev.addnext(new_el)
            else:
                _insert_after(Paragraph(prev, doc) if prev.tag == qn("w:p") else h46, _make_temp_paras(_ap))
            actions.append("Ampliado 4.6 Aportes al motor")

    # ---- 5.1 Experimentos body (critical gap) ----
    h51 = _find_para(doc, lambda t: t.startswith("5.1 Experimentos"))
    if h51 and _next_sibling_is_heading(h51, ("5.1.1",)):
        br = facts.get("best_reward") or {}
        lc = facts.get("lowest_cost") or {}
        le = facts.get("lowest_emission") or {}

        def _e51(tmp):
            p(
                tmp,
                "Los experimentos reportados en este capitulo corresponden a la matriz "
                "factorial 4 algoritmos (HAPPO, MASAC, MATD3, MAAC) x 3 escenarios (E1 "
                "flexibilidad, E2 CO2, E3 costo), con horizonte objetivo de 50 episodios "
                f"anuales (8 760 pasos/episodio) y semilla 0. La corrida canonica {facts['run_id']} "
                "aporta timeseries y traces para las 12 celdas; los KPIs de edificio y "
                "checkpoint_manifest.json estan completos para MATD3, MAAC y MASAC, mientras "
                "que HAPPO permanece parcial en esos artefactos (ver Anexo A y seccion 5.2.4).",
            )
            # Only cite numbers present in district_summary
            parts = []
            if br:
                parts.append(
                    f"segun district_summary_by_algorithm_scenario.csv, el mayor reward_mean "
                    f"distrital es {br['algorithm']}-{br['scenario']} "
                    f"(reward_mean={br['value']:.6f}, episodios registrados={br['episodes']})"
                )
            if lc:
                parts.append(
                    f"el menor costo distrital medio es {lc['algorithm']}-{lc['scenario']} "
                    f"(district_cost_mean={lc['value']:.2f})"
                )
            if le:
                parts.append(
                    f"la menor emision distrital media es {le['algorithm']}-{le['scenario']} "
                    f"(district_emission_mean={le['value']:.3f})"
                )
            if parts:
                p(
                    tmp,
                    "A nivel descriptivo de agregados distritales Drive, "
                    + "; ".join(parts)
                    + ". Estos valores contextualizan la cobertura experimental; el ranking "
                    "multiobjetivo canonicamente usado para OG/OE proviene de "
                    "best_madrl_report.json y se detalla en las secciones 5.3–5.6.",
                )
            p(
                tmp,
                "La subseccion 5.1.1 precisa las metricas; las secciones 5.2–5.5 desarrollan "
                "convergencia y resultados por objetivo; 5.7 compara con baseline CityLearn "
                "v2; 5.9–5.11 cierran la contrastacion inferencial y el veredicto de cumplimiento.",
            )

        _insert_after(h51, _make_temp_paras(_e51))
        actions.append("Insertado cuerpo 5.1 Experimentos (antes de 5.1.1 Metricas)")

    # ---- 6.2 Limitaciones expand ----
    h62 = _find_para(doc, lambda t: t.startswith("6.2 Limitaciones"))
    if h62:
        words = 0
        el = h62._element.getnext()
        while el is not None:
            if el.tag == qn("w:p"):
                txt = _para_text(el)
                if txt.startswith("6.3"):
                    break
                words += len(txt.split())
            el = el.getnext()
        if words < 150:
            def _lim(tmp):
                p(
                    tmp,
                    "Adicionalmente, la evidencia de edificio y checkpoints no es homogenea: "
                    "MAAC/MASAC/MATD3 disponen de building_kpis.csv y manifiestos de "
                    "checkpoint, mientras HAPPO carece de esos archivos en el Mirror Drive "
                    "analizado. Por ello, cualquier comparacion de comportamiento por edificio "
                    "o peso de politicas (.pt/.pkl) se restringe a los tres algoritmos "
                    "completos. Finalmente, el indice Word y la normalizacion APA residual "
                    "siguen como cierre editorial (no invalidan la evidencia cuantitativa "
                    "ya integrada).",
                )

            last = h62
            el = h62._element.getnext()
            while el is not None:
                if el.tag == qn("w:p") and _para_text(el).startswith("6.3"):
                    break
                if el.tag == qn("w:p"):
                    last = Paragraph(el, doc)
                el = el.getnext()
            _insert_after(last, _make_temp_paras(_lim))
            actions.append("Ampliado 6.2 Limitaciones")

    # ---- 6.4 Plan expand ----
    h64 = _find_para(doc, lambda t: t.startswith("6.4 Plan"))
    if h64:
        words = 0
        el = h64._element.getnext()
        while el is not None:
            if el.tag == qn("w:p"):
                txt = _para_text(el)
                if txt.startswith("Referencias"):
                    break
                words += len(txt.split())
            el = el.getnext()
        if words < 200:
            def _plan(tmp):
                p(
                    tmp,
                    "Orden de ejecucion recomendado: primero H1 (cerrar HAPPO y homogeneizar "
                    "artefactos), luego H2 (multi-semilla), en paralelo H3/H4 (inferencia y "
                    "porcentajes vs baseline), despues H5 (Optuna) y finalmente H6–H7 "
                    "(redaccion APA, indice actualizado y sustentacion). Cada hito debe "
                    "dejar evidencia en outputs/ con el mismo schema de CSV/JSON usado en "
                    f"analysis_real_drive para la corrida {facts['run_id']}.",
                )

            last = h64
            el = h64._element.getnext()
            while el is not None:
                if el.tag == qn("w:p") and _para_text(el).startswith("Referencias"):
                    break
                if el.tag == qn("w:p"):
                    last = Paragraph(el, doc)
                elif el.tag == qn("w:tbl"):
                    pass
                el = el.getnext()
            # find last p before Referencias
            prev = h64._element
            el = h64._element.getnext()
            while el is not None:
                if el.tag == qn("w:p") and _para_text(el).startswith("Referencias"):
                    break
                prev = el
                el = el.getnext()
            if prev.tag == qn("w:p"):
                _insert_after(Paragraph(prev, doc), _make_temp_paras(_plan))
            else:
                elems = _make_temp_paras(_plan)
                for new_el in reversed(elems):
                    prev.addnext(new_el)
            actions.append("Ampliado 6.4 Plan de culminacion")


def checklist(doc: Document) -> dict:
    required = [
        ("1.1 Problema", "1.1"),
        ("1.2 Objetivos", "1.2 Objetivos"),
        ("1.3 Hipotesis", "1.3"),
        ("1.5 Justificacion", "1.5"),
        ("1.6 Alcances y limitaciones", "1.6"),
        ("2.1 Estado del arte", "2.1 Estado"),
        ("2.4 Bases teoricas", "2.4 Bases"),
        ("2.5 Trabajos relacionados", "2.5 Trabajos"),
        ("3.1 Tipo de investigacion", "3.1"),
        ("3.2 Diseno", "3.2"),
        ("3.3 Datos", "3.3"),
        ("3.4 Variables", "3.4"),
        ("3.5 Tecnicas/Herramientas", "3.5"),
        ("3.6 Procedimiento", "3.6"),
        ("4.1 Sistema", "4.1.1 Sistema"),
        ("4.1 Arquitectura", "4.1.2 Arquitectura"),
        ("4.2 Modelo IA", "4.2"),
        ("4.4 Algoritmos", "4.4"),
        ("4.7 Diseno experimental", "4.7"),
        ("4.8 Implementacion", "4.8"),
        ("5.1 Experimentos", "5.1 Experimentos"),
        ("5.1.1 Metricas", "5.1.1"),
        ("5.7 Comparacion baseline", "5.7"),
        ("5.10 Discusion", "5.10"),
        ("6.1 Hallazgos", "6.1"),
        ("6.2 Limitaciones", "6.2"),
        ("6.3 Trabajo pendiente", "6.3"),
        ("6.4 Plan culminacion", "6.4"),
        ("Referencias APA", "Referencias bibliograficas"),
    ]
    texts = [(p.style.name if p.style else "", (p.text or "").strip()) for p in doc.paragraphs]
    headings = [(s, t) for s, t in texts if t and s.startswith("Heading")]
    out = {}
    for label, prefix in required:
        hit = next((t for s, t in headings if t.startswith(prefix)), None)
        if not hit:
            out[label] = "missing"
            continue
        # crude word count until next same/higher
        idx = next(i for i, (s, t) in enumerate(headings) if t == hit)
        start_level = int("".join(ch for ch in headings[idx][0] if ch.isdigit()) or "2")
        # find paragraph index
        pidx = next(i for i, para in enumerate(doc.paragraphs) if (para.text or "").strip() == hit)
        words = 0
        for para in doc.paragraphs[pidx + 1 :]:
            st = para.style.name if para.style else ""
            tx = (para.text or "").strip()
            if st.startswith("Heading"):
                lvl = int("".join(ch for ch in st if ch.isdigit()) or "9")
                if lvl <= start_level:
                    break
            words += len(tx.split())
        thr = 30
        out[label] = "OK" if words >= thr else "incomplete"
        out[label + "_words"] = words
        out[label + "_heading"] = hit
    return out


def main() -> int:
    if not OUT.is_file():
        print(f"FALTA {OUT}")
        return 1

    bak = OUT.with_suffix(".docx.bak_gaps_2026-07-14")
    if not bak.exists():
        shutil.copyfile(OUT, bak)

    facts = load_district_facts()
    doc = Document(str(OUT))
    actions: list[str] = []

    fix_cap2_numbering(doc, actions)
    strengthen_sections(doc, facts, actions)

    doc.save(str(OUT))
    status = checklist(doc)
    verification = verify_doctoral_docx(OUT)

    remaining = [k for k, v in status.items() if v in ("missing", "incomplete") and not k.endswith(("_words", "_heading"))]

    report = {
        "docx": str(OUT),
        "run_id": RUN_ID,
        "actions": actions,
        "checklist_status": {k: v for k, v in status.items() if not k.endswith(("_words", "_heading"))},
        "checklist_detail": status,
        "remaining_gaps": remaining,
        "district_facts_used": facts,
        "verification": verification,
    }
    AUDIT.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    # also mirror under analysis_real_drive
    mirror = (
        REPO
        / "outputs"
        / "_drive_madrl"
        / "full_data"
        / "analysis_real_drive"
        / "informe_final_gaps_completion_2026-07-14.json"
    )
    mirror.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({"actions": actions, "remaining_gaps": remaining, "docx": str(OUT)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
