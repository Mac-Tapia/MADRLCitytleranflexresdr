#!/usr/bin/env python
"""Audit reference sections in thesis Word documents."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
DEFAULTS = [
    REPO / "docs" / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx",
    REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx",
]


def audit(path: Path) -> dict:
    doc = Document(str(path))
    ref_headings: list[dict] = []
    ref_paras: list[tuple[int, str]] = []

    for i, para in enumerate(doc.paragraphs):
        t = (para.text or "").strip()
        style = para.style.name if para.style else ""
        if "referencia" in t.lower() and (
            style.startswith("Heading") or t.lower().startswith("referencias")
        ):
            ref_headings.append({"idx": i, "style": style, "text": t})
        if re.match(r"^[A-Za-z]", t) and re.search(r"\(\d{4}", t):
            ref_paras.append((i, t[:120]))

    # assign refs to sections
    sections: list[dict] = []
    for j, h in enumerate(ref_headings):
        start = h["idx"]
        end = ref_headings[j + 1]["idx"] if j + 1 < len(ref_headings) else len(doc.paragraphs)
        count = sum(1 for idx, _ in ref_paras if start < idx < end)
        sections.append({**h, "count": count, "end": end})

    anexo = [
        (i, (p.text or "").strip()[:100])
        for i, p in enumerate(doc.paragraphs)
        if "anexo" in (p.text or "").lower() and "referencia" in (p.text or "").lower()
    ]

    return {
        "file": str(path),
        "n_ref_headings": len(ref_headings),
        "sections": sections,
        "total_ref_paras": len(ref_paras),
        "anexo_ref_mentions": anexo,
        "sample_refs": [t for _, t in ref_paras[:3]] + [t for _, t in ref_paras[-3:]],
    }


def main() -> int:
    paths = [Path(sys.argv[1])] if len(sys.argv) > 1 else DEFAULTS
    for p in paths:
        if not p.is_file():
            print(f"MISSING: {p}")
            return 1
        r = audit(p)
        print(f"\n=== {p.name} ===")
        print(f"ref headings: {r['n_ref_headings']}")
        for s in r["sections"]:
            print(f"  [{s['idx']}] {s['style']}: {s['text'][:90]} -> {s['count']} refs")
        print(f"total APA-like ref paras: {r['total_ref_paras']}")
        if r["anexo_ref_mentions"]:
            print("anexo ref mentions:", r["anexo_ref_mentions"])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
