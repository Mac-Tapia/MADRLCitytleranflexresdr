"""Create a DOCX skeleton for the Guide N. 01 thesis plan."""

from __future__ import annotations

import argparse
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from docx import Document
except ModuleNotFoundError:  # pragma: no cover
    Document = None


TITLE = (
    "Diseño y validación de un sistema eléctrico inteligente con control "
    "multiagente basado en aprendizaje por refuerzo profundo para el despacho "
    "óptimo bajo restricciones eléctricas y operación segura en el sistema "
    "eléctrico aislado de Iquitos, Loreto, Perú - 2026"
)

SECTIONS = [
    ("CARÁTULA", 0),
    ("DATOS GENERALES", 0),
    ("1. Título propuesto", 1),
    ("2. Nombre del graduando", 1),
    ("3. Nombre del asesor", 1),
    ("4. Área involucrada", 1),
    ("5. Lugar o institución donde se desarrolla el proyecto", 1),
    ("6. Duración estimada del proyecto", 1),
    ("CAPÍTULO I. PLANTEAMIENTO DEL PROBLEMA", 0),
    ("1.1 Diagnóstico", 1),
    ("1.2 Identificación y descripción del problema de estudio", 1),
    ("1.2.1 Antecedentes bibliográficos", 2),
    ("1.2.2 Formulación del problema", 2),
    ("1.2.2.1 Formulación del problema general", 3),
    ("1.2.2.2 Formulación de los problemas específicos", 3),
    ("1.2.3 Justificación y alcances", 2),
    ("1.2.3.1 Justificación", 3),
    ("1.2.3.2 Alcances", 3),
    ("CAPÍTULO II. OBJETIVOS", 0),
    ("2.1 Objetivo general", 1),
    ("2.2 Objetivos específicos", 1),
    ("CAPÍTULO III. MARCO TEÓRICO", 0),
    ("3.1 Bases teóricas", 1),
    ("3.2 Definición de términos", 1),
    ("CAPÍTULO IV. DISEÑO METODOLÓGICO", 0),
    ("4.1 Tipo y nivel de investigación", 1),
    ("4.2 Unidad de análisis", 1),
    ("4.3 Población de estudio", 1),
    ("4.4 Tamaño de muestra", 1),
    ("4.5 Selección de muestra", 1),
    ("4.6 Técnicas de recolección de datos", 1),
    ("4.7 Técnicas e instrumentos de análisis y procesamiento de datos", 1),
    ("4.8 Etapas de intervención del estudio", 1),
    ("CAPÍTULO V. ADMINISTRACIÓN DEL PLAN DE TESIS", 0),
    ("5.1 Cronograma", 1),
    ("5.2 Presupuesto", 1),
    ("5.3 Financiamiento", 1),
    ("REFERENCIAS", 0),
    ("ANEXOS", 0),
    ("Anexo 1. Matriz de consistencia", 1),
    ("Anexo 2. Matriz de operacionalización de variables", 1),
    ("Anexo 3. Matriz bibliográfica de 50 investigaciones", 1),
    ("Anexo 4. Matriz de KPIs", 1),
    ("Anexo 5. Arquitectura CityLearn v3 propuesta", 1),
    ("Anexo 6. Comparación de backends MADRL", 1),
    ("Anexo 7. Datasets y fuentes", 1),
    ("Anexo 8. Configuración preliminar de hiperparámetros", 1),
    ("Anexo 9. Función de recompensa multiobjetivo", 1),
    ("Anexo 10. Cadenas de búsqueda", 1),
    ("Anexo 11. Evidencias de GitHub", 1),
    ("Anexo 12. Glosario MADRL", 1),
]


def _xml_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _paragraph_xml(text: str, heading: bool = False) -> str:
    style = '<w:pStyle w:val="Heading1"/>' if heading else ""
    return f"<w:p><w:pPr>{style}</w:pPr><w:r><w:t>{_xml_escape(text)}</w:t></w:r></w:p>"


def _minimal_docx(output: Path) -> None:
    body = [
        _paragraph_xml(TITLE, True),
        _paragraph_xml("Plan de Tesis de Maestría de Especialización o Profesionalizante - Guía N. 01, estructura 5.1."),
        _paragraph_xml("Completar con evidencia del Módulo A, citas APA vigentes y sin inventar resultados."),
    ]
    for heading, _level in SECTIONS:
        body.append(_paragraph_xml(heading, True))
        if heading == "CARÁTULA":
            body.extend([
                _paragraph_xml("Universidad: [pendiente]"),
                _paragraph_xml("Unidad de posgrado o escuela: [pendiente]"),
                _paragraph_xml("Autor: [pendiente]"),
                _paragraph_xml("Asesor: [pendiente]"),
                _paragraph_xml("Lugar y fecha: [pendiente]"),
            ])
        elif heading == "REFERENCIAS":
            body.append(_paragraph_xml("Insertar únicamente referencias citadas en el plan, en APA vigente."))
        else:
            body.append(_paragraph_xml("[Redactar con evidencia verificable, citas APA y trazabilidad de fuente.]"))

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + "".join(body)
        + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>'
        + "</w:body></w:document>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        "</Relationships>"
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(output, "w", ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)


def build(output: Path) -> None:
    if Document is None:
        _minimal_docx(output)
        return

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph("Plan de Tesis de Maestría de Especialización o Profesionalizante - Guía N. 01, estructura 5.1.")
    doc.add_paragraph("Completar con evidencia del Módulo A, citas APA vigentes y sin inventar resultados.")
    for heading, level in SECTIONS:
        doc.add_heading(heading, level=level if level else 1)
        doc.add_paragraph("[Redactar con evidencia verificable, citas APA y trazabilidad de fuente.]")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    build(Path(args.output))
    print(Path(args.output).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

