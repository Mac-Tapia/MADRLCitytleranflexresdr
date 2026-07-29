#!/usr/bin/env python3
"""Cierra Cap. II Dec-POMDP: dims reales + jerarquía distrito/edificio + Tabla 2.A.

Edita solo los 2 Word canónicos. Idempotente.
"""
from __future__ import annotations

import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, INFORME, TESIS  # noqa: E402

REPORT = DOCS / "CAP2_DECPOMDP_DISTRICT_BUILDING_PATCH_REPORT_2026-07-29.json"
BACKUP_DIR = REPO / "outputs" / "_word_backups"
MARKER = "Tabla 2.A — Dimensiones Dec-POMDP reales por edificio"

BUILDINGS: list[tuple[str, str, int, int, int]] = [
    ("B01", "Electro Oriente S.A.", 4, 61, 6),
    ("B02", "Munic. San Juan Bautista", 6, 75, 8),
    ("B03", "Aeropuerto Internacional", 8, 89, 10),
    ("B04", "Hipermercados Tottus", 6, 75, 8),
    ("B05", "Hotel Plaza S.A.", 3, 54, 5),
    ("B06", "Mall Aventura", 32, 257, 34),
    ("B07", "UNAP Biologia", 42, 327, 44),
    ("B08", "PNP Escuela Tecnica", 17, 152, 19),
    ("B09", "GORE Loreto COER", 10, 103, 12),
    ("B10", "Gobierno Regional Loreto", 6, 75, 8),
    ("B11", "Hospital Regional", 3, 54, 5),
    ("B12", "EsSalud", 3, 54, 5),
    ("B13", "UNAP Cs. Economicas", 11, 110, 13),
    ("B14", "Autoridad Portuaria", 4, 61, 6),
    ("B15", "DREL Colegio Nacional", 8, 89, 10),
    ("B16", "SIMA Iquitos", 11, 110, 13),
    ("B17", "Asoc. Civil Selva Amazonica", 11, 110, 13),
]

HIERARCHY_PARAS = [
    (
        "Jerarquia distrito / edificio del Dec-POMDP (valores medidos en CityLearnEnv "
        "sobre citylearn_iquitos_2023_2025). No existe un agente-distrito que emita "
        "setpoints: el control distrital es emergente de 17 politicas locales."
    ),
    (
        "Nivel edificio (agente i): observacion local o_i en R^{d_oi} con "
        "d_oi en [54, 327] (calendario, meteo, carga, termica, PV, SoC BESS, "
        "precio/CI y 7 canales EV por cargador); accion a_i con "
        "d_ai = 2 + n_i^ch en [5, 44] (electrical_storage + electric_vehicle_storage "
        "x n_ch + washing_machine); politica pi_i(a_i | o_i) sin ver o_j."
    ),
    (
        "Nivel distrito (comunidad SEAI / critico CTDE): estado global "
        "s = [o_1, ..., o_17] en R^{1856} (= suma exacta de d_oi), visible solo en "
        "entrenamiento; agregados P^com(t) = sum_i P_i^net(t), pico/rampa distritales "
        "y KPIs evaluate_v2; team_reward = mean(reward_i) y "
        "mixed_reward_i = (1 - r_team)*reward_i + r_team*team_reward con r_team = 0,70 "
        "(perfil unified_comparable_v4). gamma = 0,9999; T = 8 760."
    ),
]


def set_run_font(run, bold: bool = False, size: float = 12.0) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, bold=bold)
    return new_para


def norm(s: str) -> str:
    t = (s or "").strip().lower()
    for a, b in (
        ("á", "a"),
        ("é", "e"),
        ("í", "i"),
        ("ó", "o"),
        ("ú", "u"),
        ("ñ", "n"),
    ):
        t = t.replace(a, b)
    return re.sub(r"\s+", " ", t)


def already_done(doc: Document) -> bool:
    for p in doc.paragraphs:
        if "tabla 2.a" in norm(p.text) and "dimensiones dec-pomdp" in norm(p.text):
            return True
    return False


def fix_notation_tables(doc: Document) -> list[str]:
    """Corrige O_i 57-330 -> 54-327 y aclara A_i / S en tablas de notacion Dec-POMDP."""
    changes: list[str] = []
    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        header = norm(" | ".join(c.text for c in table.rows[0].cells))
        if "simbolo" not in header and "definicion teorica" not in header:
            continue
        body = "\n".join(
            " | ".join((c.text or "") for c in row.cells) for row in table.rows[1:]
        )
        if "0.9999" not in body and "0,9999" not in body:
            continue
        if "8 760" not in body and "8760" not in body.replace(" ", ""):
            continue
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) < 3:
                continue
            sym = norm(cells[0].text).replace(" ", "")
            val = cells[2].text or ""
            if sym.startswith("a_i") or sym in {"ai", "ai,"}:
                if "2 + n" not in val and "2+n" not in val.replace(" ", ""):
                    cells[2].text = (
                        "Heterogeneo: 5-44 dims; d_ai = 2 + n_i^ch "
                        "(BESS + EV por toma + lavadora); suma distrital = 219"
                    )
                    changes.append(f"table{ti}:A_i->formula")
            if sym == "s":
                if (
                    "1…" in val
                    or "1..." in val
                    or re.search(r"\[\s*1", val)
                    or "o_1" not in val
                ):
                    cells[2].text = (
                        "Concatenacion s = [o_1, ..., o_17]; dim global = 1 856"
                    )
                    changes.append(f"table{ti}:S->o_concat")
            if sym.startswith("o_i") or sym in {"oi", "oi,"}:
                if "57" in val or "330" in val or "54-327" not in val.replace("–", "-"):
                    cells[2].text = (
                        "Heterogeneo: 54-327 dimensiones (medido CityLearnEnv; "
                        "flota EV por edificio)"
                    )
                    changes.append(f"table{ti}:O_i->54-327")
    return changes


def fix_body_typos_and_dims(doc: Document) -> list[str]:
    changes: list[str] = []
    for p in doc.paragraphs:
        t = p.text or ""
        new = t
        if "empirieca" in new:
            new = new.replace("empirieca", "empirica")
        if "57-330" in new or "57–330" in new:
            new = new.replace("57-330", "54-327").replace("57–330", "54-327")
        if "57 a 330" in new.lower():
            new = re.sub(r"57\s*a\s*330", "54 a 327", new, flags=re.I)
        if new != t and p.runs:
            # replace keeping simple single-run rewrite when possible
            if len(p.runs) == 1:
                p.runs[0].text = new
            else:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            changes.append("body_dim_or_typo")
    return changes


def find_insert_anchor(doc: Document) -> Paragraph | None:
    """Tras el parrafo de operacionalizacion empirica Dec-POMDP (54-327 / 1856)."""
    for p in doc.paragraphs:
        t = norm(p.text)
        if "17 agentes" in t and ("54" in t or "327" in t) and "1856" in t.replace(" ", ""):
            return p
        if "operacionalizacion" in t.replace("ó", "o") and "dec-pomdp" in t and "1856" in t.replace(" ", ""):
            return p
    # fallback: heading 2.2.3 block end before 2.2.4
    prev = None
    for p in doc.paragraphs:
        t = norm(p.text)
        if t.startswith("2.2.4 ") or "citylearn y simulacion" in t.replace("ó", "o"):
            return prev
        if t.strip():
            prev = p
    return None


def insert_table_after(paragraph: Paragraph, rows: list[list[str]]) -> Table:
    tbl = OxmlElement("w:tbl")
    tbl_pr = OxmlElement("w:tblPr")
    tbl.append(tbl_pr)
    tbl_grid = OxmlElement("w:tblGrid")
    for _ in rows[0]:
        tbl_grid.append(OxmlElement("w:gridCol"))
    tbl.append(tbl_grid)
    for row_data in rows:
        tr = OxmlElement("w:tr")
        for cell_text in row_data:
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            tc.append(tc_pr)
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                r_fonts.set(qn(attr), "Times New Roman")
            r_pr.append(r_fonts)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "20")
            r_pr.append(sz)
            r.append(r_pr)
            t = OxmlElement("w:t")
            t.text = cell_text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    paragraph._p.addnext(tbl)
    return Table(tbl, paragraph._parent)


def insert_hierarchy_and_table(anchor: Paragraph) -> int:
    cursor = anchor
    n = 0
    for body in HIERARCHY_PARAS:
        cursor = insert_paragraph_after(cursor, body)
        n += 1
    title = insert_paragraph_after(cursor, MARKER, bold=True)
    n += 1
    rows = [["ID", "Edificio", "Carg. EV", "d_oi", "d_ai"]]
    for bid, name, nch, do, da in BUILDINGS:
        rows.append([bid, name, str(nch), str(do), str(da)])
    rows.append(["Distrito", "suma d_oi = d_s; suma acciones", "185", "1856", "219"])
    insert_table_after(title, rows)
    n += 1
    tbl = title._p.getnext()
    note_p = OxmlElement("w:p")
    if tbl is not None and tbl.tag == qn("w:tbl"):
        tbl.addnext(note_p)
    else:
        title._p.addnext(note_p)
    note = Paragraph(note_p, title._parent)
    run = note.add_run(
        "Nota. Dimensiones medidas con CityLearnEnv(schema Iquitos, central_agent=False). "
        "Fuente: elaboracion propia a partir del dataset citylearn_iquitos_2023_2025."
    )
    set_run_font(run, size=10.0)
    n += 1
    return n


def patch_doc(path: Path) -> dict:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{path.stem}_before_cap2_decpomdp_{datetime.now().strftime('%Y%m%d_%H%M%S')}{path.suffix}"
    shutil.copy2(path, backup)
    doc = Document(str(path))

    if already_done(doc):
        # still fix notation if stale
        ch = fix_notation_tables(doc)
        ch2 = fix_body_typos_and_dims(doc)
        if ch or ch2:
            doc.save(str(path))
            return {
                "path": str(path),
                "ok": True,
                "skipped_insert": True,
                "notation_fixes": ch,
                "body_fixes": ch2,
                "backup": str(backup),
            }
        backup.unlink(missing_ok=True)
        return {"path": str(path), "ok": True, "skipped": True, "reason": "already patched"}

    notation = fix_notation_tables(doc)
    body = fix_body_typos_and_dims(doc)
    anchor = find_insert_anchor(doc)
    if anchor is None:
        return {
            "path": str(path),
            "ok": False,
            "error": "anchor operacionalizacion Dec-POMDP no encontrado",
            "backup": str(backup),
            "notation_fixes": notation,
        }
    inserted = insert_hierarchy_and_table(anchor)
    doc.save(str(path))
    return {
        "path": str(path),
        "ok": True,
        "backup": str(backup),
        "inserted_blocks": inserted,
        "notation_fixes": notation,
        "body_fixes": body,
    }


def main() -> int:
    results = []
    for path in (TESIS, INFORME):
        if not path.is_file():
            results.append({"path": str(path), "ok": False, "error": "missing"})
            continue
        results.append(patch_doc(path))
    report = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "policy": "solo 2 Word canonicos; dims CityLearnEnv 54-327 / 5-44 / ds=1856",
        "marker": MARKER,
        "results": results,
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0 if all(r.get("ok") for r in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
