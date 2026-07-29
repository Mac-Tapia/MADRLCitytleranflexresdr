#!/usr/bin/env python3
"""Inserta en Cap. II (marco teorico) la figura CityLearn v3 control edificio
Actor-Critico / Dec-POMDP, con titulo APA y nota explicativa.

Solo edita los 2 Word canonicos:
  - docs/Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx
  - docs/Inforne_tesisV4_FINAL_REVISADO_50_EPISODIOS.docx

Idempotente via MARKER en el caption.
"""
from __future__ import annotations

import json
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import CANONS, DOCS  # noqa: E402

FIG = DOCS / "architecture" / "CITYLEARN_V3_CONTROL_EDIFICIO_ACTOR_CRITIC.png"
REPORT = DOCS / "CAP2_CONTROL_EDIFICIO_ACTOR_CRITIC_FIGURE_PATCH_REPORT_2026-07-29.json"
BACKUP_DIR = REPO / "outputs" / "_word_backups"

MARKER = "Figura 2.1. Control por edificio CityLearn v3 propuesto"

CAPTION = (
    "Figura 2.1. Control por edificio CityLearn v3 propuesto (simbología CityLearn v2): "
    "Electro Oriente S.A. (B01), comunidad inteligente Iquitos (B02–B17), red pública/SEAI, "
    "Actor–Crítico CTDE y formulación Dec-POMDP."
)

NOTE = (
    "Nota. Adaptación de la Fig. 1 de CityLearn v2 (Nweye et al., 2024) a la capa "
    "CityLearn v3 propuesto de esta tesis. Se conserva la simbología oficial "
    "(energía eléctrica = línea verde punteada; energía térmica = azul discontinua; "
    "señal de control = rojo dash-dot con P/C/D; interacción del ocupante = morada "
    "punteada; observaciones = naranja long-dash). Building_1/B01 muestra datos reales "
    "de Electro Oriente S.A. (PV 3 360,2 kWp; BESS 6 747 kWh / 1 609 kW; 4 tomas EV; "
    "d_o = 61; d_a = 6). La comunidad lista B02–B17 con etiquetas del dataset "
    "citylearn_iquitos_2023_2025. En CONTROL, el Actor π_i(a_i|o_i) ejecuta de forma "
    "descentralizada y el Crítico V/Q opera solo en entrenamiento CTDE con "
    "s = [o_1,…,o_17] ∈ ℝ^1856. Acciones MADRL activas: electrical_storage, "
    "electric_vehicle_storage y washing_machine; el bloque térmico v2 se mantiene "
    "como física (acciones cooling/heating/DHW inactivas en el schema Iquitos). "
    "Fuente: elaboración propia a partir de Nweye et al. (2024), Oliehoek y Amato "
    "(2016), Lowe et al. (2017) y artefactos del proyecto "
    "(docs/architecture/CITYLEARN_V3_CONTROL_EDIFICIO_ACTOR_CRITIC.png)."
)

INTRO = (
    "La Figura 2.1 sintetiza la formalización teórica precedente: el entorno físico "
    "CityLearn v2 (edificio, DER, red y comunidad), la interfaz de control Actor–Crítico "
    "bajo CTDE y la tupla Dec-POMDP cooperativa instanciada en el SEAI Iquitos."
)


def set_run_font(run, *, bold: bool = False, italic: bool = False, size: float = 12.0, grey: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if grey:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(anchor: Paragraph, text: str = "", *, bold: bool = False, italic: bool = False, size: float = 12.0) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        run = para.add_run(text)
        set_run_font(run, bold=bold, italic=italic, size=size, grey=italic and size <= 10.0)
    return para


def already_present(doc: Document) -> bool:
    return MARKER in "\n".join(p.text or "" for p in doc.paragraphs)


def find_anchor(doc: Document) -> Paragraph | None:
    """Insertar tras la formalizacion Dec-POMDP / Tabla 2.A y antes de 2.2.4."""
    paras = list(doc.paragraphs)
    # Prefer paragraph that links Dec-POMDP to OE (after Tabla 2.A).
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if "Esta formalizacion enlaza directamente" in t or "Esta formalización enlaza directamente" in t:
            return p
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if t.startswith("2.2.4"):
            return paras[i - 1] if i else p
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if "Tabla 2.A" in t and "Dimensiones" in t:
            # after note following table if possible
            for j in range(i + 1, min(i + 6, len(paras))):
                tj = (paras[j].text or "").strip()
                if tj.startswith("Nota"):
                    return paras[j]
            return p
    return None


def patch_doc(doc_path: Path) -> dict:
    result: dict = {"doc": str(doc_path), "status": "ok"}
    if not FIG.is_file() or FIG.stat().st_size < 5000:
        result["status"] = "figure_missing"
        result["figure"] = str(FIG)
        return result

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{doc_path.name}.pre_cap2_control_fig_{stamp}.bak"
    shutil.copy2(doc_path, backup)
    result["backup"] = str(backup)

    doc = Document(str(doc_path))
    if already_present(doc):
        result["status"] = "already_present"
        return result

    anchor = find_anchor(doc)
    if anchor is None:
        result["status"] = "anchor_not_found"
        return result

    cursor = insert_paragraph_after(anchor, INTRO, size=12.0)
    if cursor.runs:
        set_run_font(cursor.runs[0], size=12.0)

    pic = insert_paragraph_after(cursor)
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic.add_run()
    run.add_picture(str(FIG), width=Cm(15.5))

    cap = insert_paragraph_after(pic, CAPTION, italic=True, size=10.0)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        set_run_font(cap.runs[0], italic=True, size=10.0)

    note = insert_paragraph_after(cap, NOTE, italic=True, size=9.0)
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if note.runs:
        set_run_font(note.runs[0], italic=True, size=9.0, grey=True)

    doc.save(str(doc_path))
    result["inserted"] = {
        "caption": CAPTION,
        "figure": str(FIG),
        "anchor_preview": (anchor.text or "")[:120],
    }
    return result


def main() -> None:
    if not FIG.is_file():
        raise SystemExit(f"Falta figura: {FIG}")

    results = [patch_doc(p) for p in CANONS if p.is_file()]
    payload = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "figure": str(FIG),
        "marker": MARKER,
        "results": results,
    }
    REPORT.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(payload, indent=2, ensure_ascii=False))
    if any(r.get("status") not in {"ok", "already_present"} for r in results):
        raise SystemExit(1)


if __name__ == "__main__":
    main()
