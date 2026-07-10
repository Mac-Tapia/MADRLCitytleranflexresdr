#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Integra la tesis doctoral completa: diagramas/anexos Drive + resultados 50 episodios."""

from __future__ import annotations

import json
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "scripts"))

from generate_borrador_tesis_docx import (  # noqa: E402
    add_table,
    bullet,
    heading,
    p,
    status_note,
    style_base,
)
from thesis_doctoral_sections import (  # noqa: E402
    RUN_ID,
    add_chapter_5_doctoral,
    add_chapter_6_doctoral,
    add_resumen_doctoral,
    verify_doctoral_docx,
)

BASE_INTEGRATED = REPO / "docs" / (
    "Tesis_Doctoral_MADRL_CityLearn_Iquitos_resultados_drive_integrados_ordenado_con_diagramas_"
    "marco_teorico_doctoral_sustentado.docx"
)
OUT_COMPLETE = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx"
OUT_CANONICAL = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx"
OUT_SKILL = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_skill.docx"
OUT_ANTECEDENTES = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_50EP_ANTECEDENTES.docx"
METRICS = (
    REPO
    / "outputs"
    / "_drive_madrl"
    / "full_data"
    / "analysis_real_drive"
    / "thesis_docx_final_completa_metrics.json"
)


def _element_text(el) -> str:
    from docx.oxml.ns import qn

    tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
    if tag != "p":
        return ""
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def _body_children(doc) -> list:
    from docx.oxml.ns import qn

    return [c for c in doc.element.body if c.tag != qn("w:sectPr")]


def _find_marker_index(children: list, prefix: str) -> int:
    for i, child in enumerate(children):
        if _element_text(child).startswith(prefix):
            return i
    raise RuntimeError(f"No se encontro marcador de seccion: {prefix!r}")


def _extract_doc_elements(doc, start_prefix: str, end_prefix: str | None = None) -> list:
    children = _body_children(doc)
    start = _find_marker_index(children, start_prefix)
    end = len(children)
    if end_prefix:
        end = _find_marker_index(children, end_prefix)
    return [deepcopy(c) for c in children[start:end]]


def _replace_body_range(doc, start_prefix: str, end_prefix: str, new_elements: list) -> None:
    from docx.oxml.ns import qn

    body = doc.element.body
    children = [c for c in body if c.tag != qn("w:sectPr")]
    sect_pr = body.find(qn("w:sectPr"))
    start = _find_marker_index(children, start_prefix)
    end = _find_marker_index(children, end_prefix)
    rebuilt = children[:start] + new_elements + children[end:]

    for child in children:
        body.remove(child)
    for child in rebuilt:
        if sect_pr is not None:
            body.insert(body.index(sect_pr), child)
        else:
            body.append(child)


def _build_results_chapters() -> list:
    temp = Document()
    style_base(temp)
    add_chapter_5_doctoral(temp, p, heading, add_table, status_note)
    temp.add_page_break()
    add_chapter_6_doctoral(temp, p, heading, bullet, add_table)
    return _body_children(temp)


def _build_resumen_abstract() -> list:
    temp = Document()
    style_base(temp)
    add_resumen_doctoral(temp, p, heading)
    return _body_children(temp)


def _doc_metrics(path: Path) -> dict:
    from docx import Document

    doc = Document(str(path))
    paras = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    full = "\n".join(paras)
    chapters = [t for t in paras if t.startswith("Capitulo ")]
    return {
        "output": str(path.relative_to(REPO)),
        "size_bytes": path.stat().st_size,
        "paragraphs_non_empty": len(paras),
        "word_count_estimated": len(re.findall(r"\b[\wáéíóúÁÉÍÓÚñÑüÜ-]+\b", full, re.UNICODE)),
        "tables": len(doc.tables),
        "inline_images": sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref),
        "chapter_order": chapters,
        "has_cap5_baseline_v2": "5.4 Comparacion con linea base CityLearn v2" in full,
        "has_cap5_stats_colab": "p=0,155" in full or "p=0.155" in full,
        "has_anexo_a": "Anexo A" in full,
        "has_anexo_b": "Anexo B" in full,
        "figures_a_1_a_9": all(f"Figura A.{i}" in full for i in range(1, 10)),
        "figures_b_1_a_9": all(f"Figura B.{i}" in full for i in range(1, 10)),
        "run_id": RUN_ID,
    }


def build_complete() -> Path:
    if not BASE_INTEGRATED.is_file():
        raise FileNotFoundError(BASE_INTEGRATED)

    shutil.copyfile(BASE_INTEGRATED, OUT_COMPLETE)
    doc = Document(str(OUT_COMPLETE))

    _replace_body_range(
        doc,
        "Capitulo 5. Resultados y contrastacion de hipotesis",
        "Referencias bibliograficas",
        _build_results_chapters(),
    )
    _replace_body_range(doc, "Resumen", "Indice", _build_resumen_abstract())

    OUT_COMPLETE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_COMPLETE))

    OUT_CANONICAL.write_bytes(OUT_COMPLETE.read_bytes())
    OUT_SKILL.write_bytes(OUT_COMPLETE.read_bytes())

    if OUT_ANTECEDENTES.is_file():
        doc_ant = Document(str(OUT_ANTECEDENTES))
        _replace_body_range(
            doc_ant,
            "Capitulo 5. Resultados y contrastacion de hipotesis",
            "Referencias bibliograficas",
            _build_results_chapters(),
        )
        _replace_body_range(doc_ant, "Resumen", "Indice", _build_resumen_abstract())
        doc_ant.save(str(OUT_ANTECEDENTES))
        print(f"OK -> {OUT_ANTECEDENTES} (Cap. 5-6 y Resumen actualizados)")

    checks = verify_doctoral_docx(OUT_COMPLETE)
    metrics = _doc_metrics(OUT_COMPLETE)
    metrics["verification"] = checks
    METRICS.parent.mkdir(parents=True, exist_ok=True)
    METRICS.write_text(json.dumps(metrics, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"OK -> {OUT_COMPLETE}")
    print(f"OK -> {OUT_CANONICAL}")
    print(f"OK -> {OUT_SKILL}")
    print(f"OK -> {METRICS}")
    print(
        "Verificacion: "
        f"complete={checks['complete']} tables={checks['tables_count']} images={checks['images_count']}"
    )
    for sec, ok in checks["sections"].items():
        print(f"  seccion {sec}: {'OK' if ok else 'FALTA'}")
    if not checks["complete"]:
        print("AVISO: revisar secciones faltantes antes de sustentacion.")
    return OUT_COMPLETE


if __name__ == "__main__":
    build_complete()
