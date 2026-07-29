#!/usr/bin/env python3
"""Inserta en Cap. 1 de ambos Word canónicos los criterios C1–C5 de impacto.

Incluye C5 = control de recursos. No regenera Caps. 2–6.
Cap. 5 se actualiza aparte con replace_cap5_structured_in_tesis.py.
"""
from __future__ import annotations

import json
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

import sys

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import CANONS, DOCS  # noqa: E402

REPORT = DOCS / "IMPACT_CRITERIA_CONTROL_RECURSOS_PATCH_REPORT_2026-07-29.json"

CRITERIA_ROWS = [
    ["C1", "Impacto vs baseline", "Wilcoxon KPI-gains vs cero + Holm", "Inferencial HE"],
    ["C2", "Diferencias entre algoritmos", "Kruskal-Wallis / Friedman + Holm", "Inferencial HE"],
    [
        "C3",
        "KPIs fisicos de distrito por eje",
        "flex_composite / delta CO2 / delta costo (OE.1-OE.3)",
        "Descriptivo distrito",
    ],
    [
        "C4",
        "KPIs desagregados por edificio por eje",
        "17 edificios x E1/E2/E3",
        "Descriptivo edificio",
    ],
    [
        "C5",
        "Control de recursos",
        "BESS, EV/V2G, carga desplazable (acciones y exito EV)",
        "Obligatorio (atribuibilidad)",
    ],
]


def set_run_font(run, bold: bool = False, size: float = 12.0, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except (KeyError, ValueError):
            pass
    if text:
        run = new_para.add_run(text)
        set_run_font(run, bold=("Heading" in (style or "")))
    return new_para


def _set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def insert_criteria_block(after: Paragraph) -> None:
    h = insert_paragraph_after(
        after,
        "1.4.1 Criterios de determinacion del impacto",
        style="Heading 3",
    )
    intro = insert_paragraph_after(
        h,
        "Para cumplir el OG y los OE.1–OE.3, y para demostrar las hipotesis, la tesis exige "
        "el conjunto completo de criterios de determinacion del impacto (C1–C5). "
        "Uno de ellos —C5— es el control de recursos energeticos (BESS, EV/V2G y carga "
        "desplazable). Los tres ejes se reportan a nivel de distrito y a nivel de edificio "
        "(evidencia Cap. 5).",
    )
    doc = after.part.document
    table = doc.add_table(rows=1 + len(CRITERIA_ROWS), cols=4)
    table.style = "Table Grid"
    headers = ["Id", "Criterio", "Medida / prueba", "Rol"]
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, bold=True, size=10)
        _set_cell_shading(cell, "1F4E79")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(CRITERIA_ROWS):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=10)
    # Relocate table from end of body to after intro
    intro._p.addnext(table._tbl)
    caption_el = OxmlElement("w:p")
    table._tbl.addnext(caption_el)
    caption = Paragraph(caption_el, after._parent)
    run = caption.add_run(
        "Tabla 1.3. Criterios completos de determinacion del impacto "
        "(C5 = control de recursos; evidencia Cap. 5)."
    )
    set_run_font(run, italic=True, size=10)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    insert_paragraph_after(
        caption,
        "Regla de cumplimiento sin parciales: OG/OE.1–OE.3 y la demostracion de hipotesis "
        "exigen C1–C5. Cada eje se reporta a nivel distrito y a nivel edificio. "
        "C5 (control de recursos) es obligatorio para atribuir el impacto a las acciones "
        "MADRL sobre BESS, EV/V2G y carga desplazable.",
    )


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def already_patched(doc: Document) -> bool:
    blob = "\n".join(p.text for p in doc.paragraphs)
    return (
        "C5" in blob
        and "control de recursos" in blob.lower()
        and ("Tabla 1.3" in blob or "1.4.1 Criterios" in blob)
    )


def find_insert_after(doc: Document) -> Paragraph | None:
    """Insert after last §1.4 content, before 1.5 Justificacion."""
    paras = list(doc.paragraphs)
    just_idx = None
    for i, p in enumerate(paras):
        n = norm(p.text)
        if n.startswith("1.5") and "justific" in n:
            just_idx = i
            break
    if just_idx is not None and just_idx > 0:
        return paras[just_idx - 1]
    # Fallback: last paragraph mentioning operacionalizacion / 1.4
    last = None
    for p in paras:
        n = norm(p.text)
        if n.startswith("1.4") or "operacionalizaci" in n or "tabla 1.2" in n:
            last = p
    return last


def patch_doc(path: Path) -> dict:
    backup_dir = REPO / "outputs" / "_word_backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    bak = backup_dir / f"{path.name}.pre_impact_criteria_{datetime.now().strftime('%Y%m%d_%H%M%S')}.bak"
    shutil.copy2(path, bak)
    doc = Document(str(path))
    if already_patched(doc):
        return {"path": str(path), "skipped": True, "reason": "already_has_criteria", "backup": str(bak)}
    anchor = find_insert_after(doc)
    if anchor is None:
        return {"path": str(path), "ok": False, "error": "anchor_not_found", "backup": str(bak)}
    insert_criteria_block(anchor)
    doc.save(str(path))
    # verify
    doc2 = Document(str(path))
    ok = already_patched(doc2)
    return {
        "path": str(path),
        "ok": ok,
        "backup": str(bak),
        "anchor": (anchor.text or "")[:100],
    }


def main() -> int:
    results = []
    for path in CANONS:
        if not (path.is_file() and path.stat().st_size > 0):
            results.append({"path": str(path), "ok": False, "error": "missing"})
            continue
        results.append(patch_doc(path))
    report = {"generated_at": datetime.now().isoformat(timespec="seconds"), "results": results}
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if all(r.get("ok") or r.get("skipped") for r in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
