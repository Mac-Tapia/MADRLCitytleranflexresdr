#!/usr/bin/env python3
"""Embebe figuras multicriterio (Pareto / learning / degradación) tras Tabla 5.4.1.

Canons: Tesis + Informe. Backups en outputs/_word_backups/ (nunca docs/).
Idempotente: no reinserta si ya existen los captions 5.3b–5.3d.
"""
from __future__ import annotations

import json
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import CANONS  # noqa: E402

BACKUP_DIR = REPO / "outputs" / "_word_backups"
FIG_DIR = REPO / "outputs" / "madrl_multicriteria_selection" / "figures"
REPORT = REPO / "docs" / "MULTICRITERIA_FIGURES_PATCH_REPORT_2026-07-29.json"
STAMP = datetime.now().strftime("%Y%m%d_%H%M%S")
GREY = RGBColor(0x55, 0x55, 0x55)

FIGURES = (
    (
        FIG_DIR / "pareto_cost_co2_flex.png",
        "Figura 5.3b. Multicriterio — frente de Pareto costo–CO₂–flexibilidad.",
    ),
    (
        FIG_DIR / "learning_curves.png",
        "Figura 5.3c. Multicriterio — curvas de aprendizaje (50 episodios).",
    ),
    (
        FIG_DIR / "degradation_bars.png",
        "Figura 5.3d. Multicriterio — barras de degradación.",
    ),
)

MARKER = "Figura 5.3b. Multicriterio"
ANCHOR = "Tabla 5.4.1. TOPSIS descriptivo"
NOTE = (
    "Nota. Figuras descriptivas de outputs/madrl_multicriteria_selection/figures/; "
    "no constituyen evidencia de HE10–HE31 (ver §5.3 y §5.5)."
)


def set_run_font(run, *, size: float = 12.0, italic: bool = False, grey: bool = False) -> None:
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.italic = italic
    if grey:
        run.font.color.rgb = GREY
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(anchor: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        run = para.add_run(text)
        set_run_font(run)
    return para


def insert_figure_after(anchor: Paragraph, path: Path, caption: str, width_cm: float = 14.0) -> Paragraph:
    pic_para = insert_paragraph_after(anchor)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = insert_paragraph_after(pic_para, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        set_run_font(cap.runs[0], size=9.0, italic=True, grey=True)
    return cap


def find_anchor_after_topsis_table(doc: Document) -> Paragraph | None:
    for p in doc.paragraphs:
        if ANCHOR not in (p.text or ""):
            continue
        el = p._p
        nxt = el.getnext()
        # caption → tbl → (optional empty p)
        if nxt is not None and nxt.tag.endswith("}tbl"):
            after = nxt.getnext()
            if after is not None and after.tag.endswith("}p"):
                return Paragraph(after, p._parent)
            # create empty paragraph after table if missing
            new_p = OxmlElement("w:p")
            nxt.addnext(new_p)
            return Paragraph(new_p, p._parent)
        return p
    return None


def already_patched(doc: Document) -> bool:
    blob = "\n".join(p.text or "" for p in doc.paragraphs)
    return MARKER in blob


def patch_doc(path: Path) -> dict:
    for fig_path, _ in FIGURES:
        if not fig_path.is_file():
            raise FileNotFoundError(fig_path)

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{path.stem}_antes_multicriteria_figs_{STAMP}{path.suffix}"
    shutil.copy2(path, backup)

    doc = Document(str(path))
    if already_patched(doc):
        return {
            "path": str(path.relative_to(REPO)).replace("\\", "/"),
            "ok": True,
            "skipped": True,
            "reason": "already_has_figura_5.3b",
            "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        }

    anchor = find_anchor_after_topsis_table(doc)
    if anchor is None:
        return {
            "path": str(path.relative_to(REPO)).replace("\\", "/"),
            "ok": False,
            "error": f"anchor not found: {ANCHOR}",
            "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        }

    cursor = insert_paragraph_after(
        anchor,
        "Figuras complementarias de la selección multicriterio TOPSIS/AHP "
        "(descriptivo; no evidencia de hipótesis específicas):",
    )
    inserted = []
    for fig_path, caption in FIGURES:
        cursor = insert_figure_after(cursor, fig_path, caption)
        inserted.append(caption)
    note = insert_paragraph_after(cursor, NOTE)
    if note.runs:
        set_run_font(note.runs[0], size=9.0, italic=True)

    doc.save(str(path))
    verify = Document(str(path))
    text = "\n".join(p.text or "" for p in verify.paragraphs)
    return {
        "path": str(path.relative_to(REPO)).replace("\\", "/"),
        "ok": True,
        "skipped": False,
        "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        "inserted": inserted,
        "checks": {
            "has_5_3b": MARKER in text,
            "has_5_3c": "Figura 5.3c. Multicriterio" in text,
            "has_5_3d": "Figura 5.3d. Multicriterio" in text,
            "media_count": len(
                [n for n in Document(str(path)).part.related_parts if "image" in str(n).lower()]
            ),
        },
    }


def main() -> int:
    report: dict = {"stamp": STAMP, "files": {}}
    ok_all = True
    for path in CANONS:
        if not (path.is_file() and path.stat().st_size > 0):
            report["files"][path.name] = {"ok": False, "error": "missing"}
            ok_all = False
            continue
        result = patch_doc(path)
        report["files"][path.name] = result
        ok_all = ok_all and bool(result.get("ok"))
        print(json.dumps(result, ensure_ascii=False))
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print("WROTE", REPORT)
    return 0 if ok_all else 1


if __name__ == "__main__":
    raise SystemExit(main())
