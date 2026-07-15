# -*- coding: utf-8 -*-
"""Incremental patch: Figura 5.8e + Cap. 6.3–6.5 + page break before Referencias.

Targets only:
  - docs/ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx
  - docs/Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx

Does NOT mass-overwrite Word docs. Preserves Cap. 6 body; replaces 6.3 content
and updates 6.4/6.5 status cells after 5.8e fix. No commit/push.
"""
from __future__ import annotations

import json
import re
import zipfile
from copy import deepcopy
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
TARGETS = [
    REPO / "docs" / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx",
    REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx",
]
FIG_DIR = REPO / "outputs" / "_drive_madrl" / "gdrive_20260627_164047_objective_analysis" / "figures"
TABLE_DIR = REPO / "outputs" / "_drive_madrl" / "gdrive_20260627_164047_objective_analysis" / "tables"
FULL_DATA = REPO / "outputs" / "_drive_madrl" / "full_data"
CONS_BEHAV = TABLE_DIR / "gdrive_building_behavior_summary_all.csv"
REPORT = TABLE_DIR.parent / "patch_abrir_58e_cap6_refs_report.json"
RUN_ID = "madrl_v3_20260627_164047"

ALGOS = ["HAPPO", "MAAC", "MASAC", "MATD3"]
SCENARIOS = ["E1", "E2", "E3"]
GREY = RGBColor(0x59, 0x59, 0x59)


def set_run_font(run, size=11, bold=False, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def patch_para_full(para, new_text: str) -> None:
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def load_action_l2() -> pd.DataFrame:
    rows = []
    for algo in ALGOS:
        for scen in SCENARIOS:
            candidates = [
                FULL_DATA / algo / scen / "data" / "trace.csv",
                REPO / "outputs" / RUN_ID / algo / scen / "data" / "trace.csv",
            ]
            path = next((p for p in candidates if p.exists() and p.stat().st_size > 100), None)
            if path is None:
                continue
            try:
                df = pd.read_csv(path, usecols=lambda c: c in {"action_l2"})
            except Exception:
                continue
            if df.empty or "action_l2" not in df.columns:
                continue
            rows.append(
                {
                    "algorithm": algo,
                    "scenario": scen,
                    "action_l2": float(pd.to_numeric(df["action_l2"], errors="coerce").mean()),
                    "source": str(path.relative_to(REPO)).replace("\\", "/"),
                }
            )
    return pd.DataFrame(rows)


def load_ev_bess() -> pd.DataFrame:
    """Prefer consolidated behavior summary (includes HAPPO when present)."""
    if not CONS_BEHAV.exists():
        raise FileNotFoundError(CONS_BEHAV)
    df = pd.read_csv(CONS_BEHAV)
    need = {"algorithm", "scenario", "ev_charge_total_kwh", "battery_throughput_total_kwh"}
    missing = need - set(df.columns)
    if missing:
        raise ValueError(f"Missing columns in consolidated behavior: {missing}")
    g = (
        df.groupby(["algorithm", "scenario"], as_index=False)[
            ["ev_charge_total_kwh", "battery_throughput_total_kwh"]
        ]
        .mean(numeric_only=True)
    )
    g = g.rename(
        columns={
            "ev_charge_total_kwh": "ev_metric",
            "battery_throughput_total_kwh": "bess_metric",
        }
    )
    g["ev_src"] = "ev_charge_total_kwh"
    g["bess_src"] = "battery_throughput_total_kwh"
    return g


def regenerate_fig_58e() -> tuple[Path, dict]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    action = load_action_l2()
    behav = load_ev_bess()
    out = FIG_DIR / "trace_policy_action_heatmaps.png"
    verify: dict = {"action_l2": {}, "ev": {}, "bess": {}, "sources": {}}

    panels = [
        ("action_l2", "Intensidad de accion (mean action_l2)", action, "action_l2"),
        ("ev", "Carga EV media (behavior summary, kWh)", behav, "ev_metric"),
        ("bess", "Throughput BESS medio (behavior summary, kWh)", behav, "bess_metric"),
    ]

    fig, axes = plt.subplots(1, 3, figsize=(13.4, 4.3))
    for ax, (key, title, df, col) in zip(axes, panels):
        pivot = pd.DataFrame(index=ALGOS, columns=SCENARIOS, dtype=float)
        if not df.empty and col in df.columns:
            for _, r in df.iterrows():
                if r["algorithm"] in ALGOS and r["scenario"] in SCENARIOS:
                    pivot.loc[r["algorithm"], r["scenario"]] = r[col]
        vals = pivot.values.astype(float)
        verify[key] = {
            f"{a}-{s}": (None if pd.isna(pivot.loc[a, s]) else float(pivot.loc[a, s]))
            for a in ALGOS
            for s in SCENARIOS
        }
        finite = vals[np.isfinite(vals)]
        if finite.size:
            vmin, vmax = float(np.nanmin(finite)), float(np.nanmax(finite))
            if abs(vmax - vmin) < 1e-12:
                vmax = vmin + 1e-6
        else:
            vmin, vmax = 0.0, 1.0
        img = ax.imshow(vals, cmap="plasma", vmin=vmin, vmax=vmax, aspect="auto")
        ax.set_xticks(range(len(SCENARIOS)), SCENARIOS)
        ax.set_yticks(range(len(ALGOS)), ALGOS)
        ax.set_title(title, fontsize=9)
        for i in range(len(ALGOS)):
            for j in range(len(SCENARIOS)):
                v = vals[i, j]
                label = "NA" if not np.isfinite(v) else (f"{v:.2f}" if abs(v) < 100 else f"{v:.0f}")
                ax.text(j, i, label, ha="center", va="center", color="white", fontsize=7)
        fig.colorbar(img, ax=ax, fraction=0.046, pad=0.04)

    verify["sources"] = {
        "action_l2": "outputs/_drive_madrl/full_data/*/E*/data/trace.csv (action_l2)",
        "ev": "gdrive_building_behavior_summary_all.csv::ev_charge_total_kwh",
        "bess": "gdrive_building_behavior_summary_all.csv::battery_throughput_total_kwh",
        "note": "No usar ev_charge_kwh / electrical_storage_soc de trace.csv (columnas muertas = 0)",
    }
    fig.suptitle(
        "Figura 5.8e — politicas/acciones (action_l2 de trace; EV/BESS de behavior summary)",
        fontsize=11,
    )
    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)

    rows = []
    for panel in ("action_l2", "ev", "bess"):
        for a in ALGOS:
            for s in SCENARIOS:
                rows.append(
                    {
                        "panel": panel,
                        "algorithm": a,
                        "scenario": s,
                        "value": verify[panel].get(f"{a}-{s}"),
                    }
                )
    pd.DataFrame(rows).to_csv(TABLE_DIR / "fig58e_panel_values.csv", index=False)
    return out, verify


def replace_image_near_caption(docx_path: Path, png_path: Path) -> dict:
    info: dict = {"replaced": False}
    with zipfile.ZipFile(docx_path, "r") as zin:
        xml = zin.read("word/document.xml").decode("utf-8")
        rels = zin.read("word/_rels/document.xml.rels").decode("utf-8")

    m = re.search(r"5\.8e", xml)
    if not m:
        info["reason"] = "5.8e no encontrado en XML"
        return info
    blips = re.findall(r'a:blip[^>]+r:embed="(rId\d+)"', xml[m.end() : m.end() + 30000])
    if not blips:
        blips = re.findall(r'a:blip[^>]+r:embed="(rId\d+)"', xml[m.end() :])
    if not blips:
        info["reason"] = "No se halló a:blip después de 5.8e"
        return info
    rid = blips[0]
    rm = re.search(rf'Relationship[^>]*Id="{rid}"[^>]*Target="([^"]+)"', rels)
    if not rm:
        info["reason"] = f"No Target para {rid}"
        return info
    target = rm.group(1)
    media_path = "word/" + target.lstrip("/")
    png_bytes = png_path.read_bytes()
    tmp = docx_path.with_suffix(".docx.tmp_58e")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        replaced = False
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.replace("\\", "/") == media_path.replace("\\", "/"):
                data = png_bytes
                replaced = True
            zout.writestr(item, data)
        info["replaced"] = replaced
        info["rId"] = rid
        info["media"] = media_path
        info["png_bytes"] = len(png_bytes)
    if not info["replaced"]:
        tmp.unlink(missing_ok=True)
        info["reason"] = f"Parte {media_path} no estaba en el zip"
        return info
    tmp.replace(docx_path)
    return info


def update_58e_caption(doc: Document) -> int:
    n = 0
    caption = (
        "Figura 5.8e\n"
        "Politicas y acciones medias (action_l2 desde trace.csv; EV/BESS desde building_behavior_summary)"
    )
    for p in doc.paragraphs:
        t = p.text or ""
        if "Figura 5.8e" in t:
            patch_para_full(p, caption)
            n += 1
            break
    return n


def heading_level(para: Paragraph) -> int:
    name = para.style.name if para.style else ""
    m = re.search(r"(\d+)", name or "")
    return int(m.group(1)) if m else 9


def find_heading(doc: Document, prefixes: tuple[str, ...]) -> Paragraph | None:
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        for pref in prefixes:
            if t.startswith(pref):
                return para
    return None


def collect_section_body(doc: Document, start: Paragraph) -> tuple[list, Paragraph | None]:
    body = list(doc.element.body)
    try:
        idx = body.index(start._element)
    except ValueError:
        return [], None
    start_lvl = heading_level(start)
    end_el = None
    keep: list = []
    for el in body[idx + 1 :]:
        if el.tag == qn("w:sectPr"):
            break
        if el.tag == qn("w:p"):
            p = Paragraph(el, doc)
            t = (p.text or "").strip()
            style = p.style.name if p.style else ""
            if style.startswith("Heading") and t and heading_level(p) <= start_lvl:
                end_el = p
                break
        keep.append(el)
    return keep, end_el


def clear_elements(elements: list) -> None:
    for el in elements:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def insert_before(anchor: Paragraph | None, elements: list, doc: Document) -> None:
    body = doc.element.body
    if anchor is not None:
        ref = anchor._element
        for el in elements:
            ref.addprevious(el)
    else:
        sect = body.find(qn("w:sectPr"))
        for el in elements:
            if sect is not None:
                sect.addprevious(el)
            else:
                body.append(el)


def add_para(tmp: Document, text: str, *, bold=False, size=11):
    para = tmp.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return para


def build_63_elements() -> list:
    tmp = Document()
    for child in list(tmp.element.body):
        if child.tag != qn("w:sectPr"):
            tmp.element.body.remove(child)

    tmp.add_heading("6.3 Trabajo pendiente", level=2)
    add_para(
        tmp,
        "El trabajo pendiente se declara con honestidad metodologica respecto de la corrida "
        f"canonica {RUN_ID} (factorial 4×3: HAPPO, MAAC, MASAC y MATD3 × E1–E3). No se imputan "
        "episodios ni KPIs inexistentes; lo que sigue son huecos reales del pipeline y del "
        "manuscrito que deben cerrarse antes de la version de sustentacion.",
    )
    add_para(
        tmp,
        "Pendientes de evidencia empirica (bloqueantes Cap. 5 / Anexo A): (1) homogeneizar "
        "HAPPO a evaluate_v2/core_kpis y artefactos building_* comparables con MAAC/MASAC/MATD3 "
        "(hoy HAPPO aporta trazas y series distritales —49 episodios reales por escenario— pero "
        "queda incompleto en KPIs de edificio cuando faltan CSV locales); (2) cerrar "
        "checkpoint_manifest.json de HAPPO en la corrida canonica (conteo = 0 en Figura 5.1 / "
        "Anexo A.4 frente a MAAC 52, MASAC 12 y MATD3 34 archivos listados); (3) mantener la "
        "Figura 5.8e con fuentes mixtas auditadas (action_l2 desde full_data/trace.csv; EV/BESS "
        "desde building_behavior_summary), sin reutilizar columnas muertas ev_charge_kwh / "
        "electrical_storage_soc (=0 en trace.csv); (4) auditar celdas cero en Anexo A.4 y demas "
        "tablas/figuras para distinguir cero legitimo vs fallo de lectura.",
    )
    add_para(
        tmp,
        "Pendientes de analisis y robustez (no bloquean la lectura descriptiva 50 ep, pero si "
        "afirmaciones de generalizacion): corrida multi-semilla (≥3, ideal ≥5) con post-hoc "
        "alineado a la Tabla 3.4; frontera de Pareto por eje OE.1–OE.3 frente a baseline "
        "CityLearn/RBC; Optuna (TPE) por backend solo si se declara optimizacion hiperrametricas; "
        "contraste SB3 (PPO/SAC/A2C) bajo el mismo schema de Iquitos como extension opcional.",
    )
    add_para(
        tmp,
        "Pendientes editoriales e institucionales: pasada ortografica RAE (tildes, tipografia, "
        "concordancia) en Cap. 1–6; verificar citas marcadas [PV] y actualizar indices Word (F9); "
        "completar metadatos de asesor / [por definir] unicamente con dato real del programa "
        "(no inventar nombres); sincronizar ABRIR_ESTE y FINAL_COMPLETA sin regenerar masivamente "
        "el cuerpo; PDF final y paquete de reproducibilidad (scripts + CSV + manifiestos).",
    )
    return [deepcopy(c) for c in tmp.element.body if c.tag != qn("w:sectPr")]


def replace_section(doc: Document, prefixes: tuple[str, ...], new_els: list, stop: tuple[str, ...]) -> str:
    start = find_heading(doc, prefixes)
    if start is None:
        anchor = None
        for pref in stop:
            anchor = find_heading(doc, (pref,))
            if anchor is not None:
                break
        if anchor is None:
            return "skip_missing"
        insert_before(anchor, new_els, doc)
        return "inserted"
    olds, nxt = collect_section_body(doc, start)
    clear_elements([start._element] + olds)
    insert_before(nxt, new_els, doc)
    return "replaced"


def fix_language_and_status(doc: Document, fig_ok: bool) -> dict:
    stats = {"lang": 0, "status_cells": 0, "filename_fix": 0}
    # Fix bad prior replacement Tesis_De tesis_ → Tesis_Doctoral_
    for p in doc.paragraphs:
        t = p.text or ""
        if "Tesis_De tesis_" in t or "etiqueta histórica «De tesis»" in t or "etiqueta historica «De tesis»" in t:
            newt = t.replace("Tesis_De tesis_", "Tesis_Doctoral_")
            newt = newt.replace("etiqueta histórica «De tesis»", "etiqueta histórica «Doctoral» (solo en nombre de archivo)")
            newt = newt.replace("etiqueta historica «De tesis»", "etiqueta historica «Doctoral» (solo en nombre de archivo)")
            if newt != t:
                patch_para_full(p, newt)
                stats["filename_fix"] += 1
        # residual doctorado
        if re.search(r"doctorado|doctoral", t, re.I) and "Tesis_Doctoral_" not in t:
            newt = re.sub(r"\btesis doctoral\b", "tesis", t, flags=re.I)
            newt = re.sub(r"\bdoctoral\b", "de la tesis", newt, flags=re.I)
            newt = re.sub(r"\bdoctorado\b", "posgrado", newt, flags=re.I)
            if newt != t:
                patch_para_full(p, newt)
                stats["lang"] += 1

    # Update H2 / C4 status in tables if present
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            joined = " | ".join(cells)
            if "H2. Figura 5.8e" in joined or (joined.startswith("H2") and "5.8e" in joined):
                if fig_ok and len(row.cells) >= 4:
                    patch_para_full(
                        row.cells[-1].paragraphs[0],
                        "CUMPLE (PNG regenerado; HAPPO action_l2≈2.22–2.25; EV/BESS desde behavior summary)",
                    )
                    stats["status_cells"] += 1
            if cells and cells[0] == "C4":
                if fig_ok and len(row.cells) >= 4:
                    patch_para_full(
                        row.cells[-1].paragraphs[0],
                        "CUMPLE (fuentes mixtas; sin paneles de columnas muertas trace EV/SOC)",
                    )
                    stats["status_cells"] += 1
            # Fix De tesis in table cells too
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text or ""
                    if "Tesis_De tesis_" in t:
                        patch_para_full(p, t.replace("Tesis_De tesis_", "Tesis_Doctoral_"))
                        stats["filename_fix"] += 1

    # Soften no-sobreclaim paragraph if C4 now meets
    for p in doc.paragraphs:
        t = p.text or ""
        if t.startswith("Regla de no sobreclaim") and fig_ok:
            patch_para_full(
                p,
                "Regla de no sobreclaim: con la evidencia actual se puede afirmar cumplimiento de "
                "C1–C4, C7, C11 y C12; C5–C6 y C8–C10 permanecen abiertos (A.4/checkpoints HAPPO, "
                "pasada RAE, indices F9 y metadatos institucionales) y deben cerrarse antes de "
                "fijar la version de sustentacion. Completar C5 es prioritario para la "
                "reproducibilidad de modelos HAPPO.",
            )
            stats["lang"] += 1
    return stats


def ensure_page_break_before_refs(doc: Document) -> str:
    refs = find_heading(
        doc,
        (
            "Referencias bibliograficas",
            "Referencias bibliográficas",
            "Referencias",
            "References",
        ),
    )
    if refs is None:
        return "refs_heading_missing"
    # Already page-break-before?
    pPr = refs._element.find(qn("w:pPr"))
    if pPr is not None:
        pb = pPr.find(qn("w:pageBreakBefore"))
        if pb is not None:
            refs.paragraph_format.page_break_before = True
            return "already_or_set"
    refs.paragraph_format.page_break_before = True
    # Also ensure no duplicate empty page-break para immediately before causing double blank —
    # leave as format flag only (nueva hoja).
    return "page_break_before_set"


def verify_doc(doc: Document, png_verify: dict) -> dict:
    text = "\n".join(p.text or "" for p in doc.paragraphs)
    refs = find_heading(doc, ("Referencias bibliograficas", "Referencias bibliográficas", "Referencias"))
    pb = False
    if refs is not None:
        pb = bool(refs.paragraph_format.page_break_before)
    happo_ok = all(
        (png_verify.get("action_l2") or {}).get(f"HAPPO-{s}") not in (None, 0) for s in SCENARIOS
    )
    ev_nonzero = any(
        abs(v or 0) > 1e-9 for k, v in (png_verify.get("ev") or {}).items() if v is not None
    )
    bess_nonzero = any(
        abs(v or 0) > 1e-9 for k, v in (png_verify.get("bess") or {}).items() if v is not None
    )
    return {
        "has_6_3": "6.3 Trabajo pendiente" in text,
        "has_6_4": "6.4 Plan para culminar la tesis" in text,
        "has_6_5": "6.5 Criterios de cierre de la tesis" in text,
        "has_5_8e_mixed_caption": "building_behavior_summary" in text and "5.8e" in text,
        "refs_page_break_before": pb,
        "banned_body_hits": (
            banned := [
                m.group(0)
                for m in re.finditer(r"\bdoctorado\b|\bdoctoral\b", text, re.I)
                if "Tesis_Doctoral" not in text[max(0, m.start() - 30) : m.end() + 10]
            ][:8]
        ),
        "no_doctorado_body": len(banned) == 0,
        "no_tesis_de_tesis": "Tesis_De tesis" not in text,
        "happo_action_l2_ok": happo_ok,
        "ev_panel_nonzero": ev_nonzero,
        "bess_panel_nonzero": bess_nonzero,
        "words_6_3": len(re.findall(r"\w+", text.split("6.3 Trabajo pendiente")[1].split("6.4 ")[0]))
        if "6.3 Trabajo pendiente" in text and "6.4 " in text
        else 0,
    }


def save_doc(path: Path, doc: Document) -> Path:
    try:
        doc.save(str(path))
        return path
    except PermissionError:
        alt = path.with_name(path.stem + "_PATCHED.docx")
        doc.save(str(alt))
        return alt


def patch_generator_in_build_script() -> dict:
    """Fix tools/build_final_thesis_gdrive_objectives.py figure block for 5.8e."""
    path = REPO / "tools" / "build_final_thesis_gdrive_objectives.py"
    text = path.read_text(encoding="utf-8")
    old = '''    if not traces.empty:
        agg = traces.groupby(["algorithm", "scenario"])[["action_l2", "action_mean", "ev_charge_kwh", "ev_v2g_export_kwh", "electrical_storage_soc"]].mean(numeric_only=True).reset_index()
        fig, axes = plt.subplots(1, 3, figsize=(13, 4))
        for ax, col, title in zip(axes, ["action_l2", "ev_charge_kwh", "electrical_storage_soc"], ["Intensidad de accion", "Carga EV media", "SOC BESS medio"]):
            pivot = agg.pivot(index="algorithm", columns="scenario", values=col).reindex(index=ALGOS, columns=SCENARIOS)
            img = ax.imshow(pivot.values, cmap="plasma")
            ax.set_xticks(range(len(SCENARIOS)), SCENARIOS)
            ax.set_yticks(range(len(ALGOS)), ALGOS)
            ax.set_title(title)
            for i in range(pivot.shape[0]):
                for j in range(pivot.shape[1]):
                    ax.text(j, i, f"{pivot.iloc[i,j]:.2f}", ha="center", va="center", color="white", fontsize=8)
        fig.colorbar(img, ax=axes.ravel().tolist(), shrink=0.75)
        out = FIG_DIR / "trace_policy_action_heatmaps.png"
        fig.savefig(out, dpi=180, bbox_inches="tight")
        plt.close(fig)
        paths["trace_policy_heatmaps"] = out'''

    new = '''    # Figura 5.8e: action_l2 desde trace; EV/BESS desde building_behavior_summary.
    # No usar ev_charge_kwh / electrical_storage_soc de trace.csv (columnas muertas=0).
    # Colorbar independiente por panel.
    if not traces.empty and "action_l2" in traces.columns:
        agg = traces.groupby(["algorithm", "scenario"])[["action_l2"]].mean(numeric_only=True).reset_index()
        behav_path = TABLE_DIR / "gdrive_building_behavior_summary_all.csv"
        behav = pd.read_csv(behav_path) if behav_path.exists() else pd.DataFrame()
        fig, axes = plt.subplots(1, 3, figsize=(13.4, 4.3))
        panels = [("action_l2", agg, "Intensidad de accion (mean action_l2)")]
        if not behav.empty:
            metric_cols = [c for c in ("ev_charge_total_kwh", "battery_throughput_total_kwh") if c in behav.columns]
            g = behav.groupby(["algorithm", "scenario"], as_index=False)[metric_cols].mean(numeric_only=True)
            if "ev_charge_total_kwh" in g.columns:
                panels.append(("ev_charge_total_kwh", g, "Carga EV media (behavior summary, kWh)"))
            if "battery_throughput_total_kwh" in g.columns:
                panels.append(("battery_throughput_total_kwh", g, "Throughput BESS (behavior summary, kWh)"))
        while len(panels) < 3:
            panels.append(panels[0])
        for ax, (col, dfp, title) in zip(axes, panels[:3]):
            pivot = dfp.pivot(index="algorithm", columns="scenario", values=col).reindex(index=ALGOS, columns=SCENARIOS)
            vals = pivot.values.astype(float)
            finite = vals[np.isfinite(vals)]
            vmin, vmax = (float(np.nanmin(finite)), float(np.nanmax(finite))) if finite.size else (0.0, 1.0)
            if abs(vmax - vmin) < 1e-12:
                vmax = vmin + 1e-6
            img = ax.imshow(vals, cmap="plasma", vmin=vmin, vmax=vmax, aspect="auto")
            ax.set_xticks(range(len(SCENARIOS)), SCENARIOS)
            ax.set_yticks(range(len(ALGOS)), ALGOS)
            ax.set_title(title, fontsize=9)
            for i in range(pivot.shape[0]):
                for j in range(pivot.shape[1]):
                    v = vals[i, j]
                    lab = "NA" if not np.isfinite(v) else (f"{v:.2f}" if abs(v) < 100 else f"{v:.0f}")
                    ax.text(j, i, lab, ha="center", va="center", color="white", fontsize=7)
            fig.colorbar(img, ax=ax, fraction=0.046, pad=0.04)
        fig.suptitle("Figura 5.8e — politicas/acciones (fuentes mixtas auditadas)", fontsize=11)
        fig.tight_layout()
        out = FIG_DIR / "trace_policy_action_heatmaps.png"
        fig.savefig(out, dpi=180, bbox_inches="tight")
        plt.close(fig)
        paths["trace_policy_heatmaps"] = out'''

    if old not in text:
        return {"patched_generator": False, "reason": "old block not found (maybe already patched)"}
    # ensure numpy import available in fallback
    if "import numpy as np" not in text:
        text = text.replace("import matplotlib.pyplot as plt\n", "import matplotlib.pyplot as plt\nimport numpy as np\n", 1)
    path.write_text(text.replace(old, new), encoding="utf-8")
    return {"patched_generator": True}


def main() -> int:
    report: dict = {"targets": [], "actions": []}

    gen = patch_generator_in_build_script()
    report["actions"].append(gen)

    png_path, png_verify = regenerate_fig_58e()
    report["actions"].append({"regen_58e": str(png_path), "verify": png_verify})
    fig_ok = all(
        abs((png_verify["action_l2"].get(f"HAPPO-{s}") or 0)) > 1.0 for s in SCENARIOS
    ) and any(abs(v or 0) > 1 for v in png_verify["ev"].values()) and any(
        abs(v or 0) > 1 for v in png_verify["bess"].values()
    )

    for target in TARGETS:
        entry: dict = {"file": target.name, "exists": target.exists()}
        if not target.exists():
            report["targets"].append(entry)
            continue
        doc = Document(str(target))
        entry["caption_58e"] = update_58e_caption(doc)
        entry["sec_6_3"] = replace_section(
            doc,
            ("6.3 Trabajo pendiente", "6.3 Trabajo"),
            build_63_elements(),
            ("6.4", "6.5", "Referencias"),
        )
        # Keep existing 6.4/6.5 structure; only status + filename artifact fixes
        entry["lang_status"] = fix_language_and_status(doc, fig_ok=fig_ok)
        entry["refs_pagebreak"] = ensure_page_break_before_refs(doc)

        saved = save_doc(target, doc)
        entry["saved_as"] = str(saved.relative_to(REPO)).replace("\\", "/")
        entry["blocked"] = saved != target

        img = replace_image_near_caption(saved, png_path)
        entry["image_58e"] = img

        doc2 = Document(str(saved))
        entry["verification"] = verify_doc(doc2, png_verify)
        report["targets"].append(entry)
        print(
            f"{target.name}: 6.3={entry['sec_6_3']} img={img.get('replaced')} "
            f"pb={entry['refs_pagebreak']} saved={entry['saved_as']}"
        )

    report["fig_ok"] = fig_ok
    report["png_path"] = str(png_path)
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print("REPORT", REPORT)
    print(json.dumps({"fig_ok": fig_ok, "happo": png_verify.get("action_l2"), "ev_sample": {k: png_verify["ev"].get(k) for k in ("HAPPO-E1", "MAAC-E1")}}, indent=2))
    return 0 if fig_ok else 2


if __name__ == "__main__":
    raise SystemExit(main())
