"""Create a Guide N.02 thesis draft from current project evidence only."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[4]
DATASET_AUDIT = ROOT / "outputs" / "dataset_audit"
TRAINING_ROOT = ROOT / "outputs" / "citylearn_v3_madrl_oficial_v4"
DATASET_DIR = ROOT / "CityLearn" / "data" / "datasets" / "citylearn_iquitos_2023_2025"

TITLE = (
    "MULTI-AGENTE DE APRENDIZAJE POR REFUERZO PROFUNDO PARA LA GESTION "
    "COORDINADA DE FLEXIBILIDAD ENERGETICA, EMISIONES DE CARBONO Y COSTOS "
    "ENERGETICOS EN COMUNIDADES INTELIGENTES"
)


def read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8-sig"))


def read_csv_rows(path: Path) -> list[dict]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def fnum(value: object, decimals: int = 1) -> str:
    try:
        return f"{float(value):,.{decimals}f}"
    except Exception:
        return str(value)


def der_totals(rows: list[dict]) -> dict[str, float]:
    fields = [
        "PV_schema_kWp",
        "BESS_schema_kWh",
        "BESS_schema_kW",
        "EV_dimensionador_v3_kW",
        "EV_dimensionador_v3_count",
        "Maquina_controlada_MWh",
        "PV_total_MWh",
    ]
    totals = {field: 0.0 for field in fields}
    for row in rows:
        for field in fields:
            try:
                totals[field] += float(row.get(field) or 0.0)
            except ValueError:
                pass
    return totals


def final_artifact_summary() -> tuple[int, int, list[str]]:
    algos = ["happo", "masac", "matd3", "maac"]
    scenarios = ["E1", "E2", "E3"]
    complete = 0
    missing: list[str] = []
    for algo in algos:
        for scenario in scenarios:
            base = TRAINING_ROOT / algo / f"{scenario}_seed_0" / "data"
            required = ["results.json", "timeseries.csv", "trace.csv"]
            if all((base / name).exists() for name in required):
                complete += 1
            else:
                missing.append(f"{algo.upper()}/{scenario}")
    return complete, len(algos) * len(scenarios), missing


def latest_live_progress() -> dict:
    candidates = list(TRAINING_ROOT.glob("*/*_seed_0/live_progress.json"))
    if not candidates:
        return {}
    newest = max(candidates, key=lambda p: p.stat().st_mtime)
    data = read_json(newest)
    data["_path"] = str(newest.relative_to(ROOT))
    return data


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build_docx(output: Path) -> None:
    csv_manifest = read_json(DATASET_AUDIT / "csv_integrity_manifest.json")
    ready = read_json(DATASET_AUDIT / "training_dataset_ready_manifest.json")
    status = read_json(TRAINING_ROOT / "official_full_status.json")
    der_rows = read_csv_rows(DATASET_AUDIT / "der_sizing_audit.csv")
    ev_rows = read_csv_rows(DATASET_AUDIT / "ev_charger_sizing_audit.csv")
    totals = der_totals(der_rows)
    complete_runs, total_runs, missing_runs = final_artifact_summary()
    live = latest_live_progress()

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph("Borrador generado desde evidencia local vigente del proyecto.")
    doc.add_paragraph(
        "Regla de uso: no se incorporan resultados historicos ni valores no observados. "
        "Los KPIs finales se mantienen como pendientes mientras no existan results.json, "
        "timeseries.csv y trace.csv para cada algoritmo y escenario."
    )

    doc.add_heading("CARATULA", level=1)
    doc.add_paragraph("Universidad: [pendiente]")
    doc.add_paragraph("Escuela de posgrado: [pendiente]")
    doc.add_paragraph("Tesista: [pendiente]")
    doc.add_paragraph("Asesor: [pendiente]")
    doc.add_paragraph("Lima, Peru - 2026")

    doc.add_heading("DATOS GENERALES", level=1)
    doc.add_paragraph("Dedicatoria: [pendiente]")
    doc.add_paragraph("Agradecimientos: [pendiente]")
    doc.add_paragraph("Resumen - Abstract: pendiente de cerrar con resultados finales validados.")
    doc.add_paragraph(
        "Introduccion: el trabajo evalua MADRL cooperativo para gestion coordinada de "
        "flexibilidad energetica, emisiones de CO2 y costos energeticos en comunidades inteligentes."
    )

    doc.add_heading("CAPITULO I. PLANTEAMIENTO DEL PROBLEMA", level=1)
    doc.add_heading("1.1 Diagnostico", level=2)
    doc.add_paragraph(
        "El SEAI Iquitos opera como sistema electrico aislado con generacion diesel dominante, "
        "alta intensidad de carbono y necesidad de coordinacion entre PV, BESS, EV y cargas controladas."
    )
    doc.add_heading("1.2 Identificacion y descripcion del problema de estudio", level=2)
    doc.add_paragraph(
        "La gestion no coordinada de edificios, almacenamiento, vehiculos electricos y cargas flexibles "
        "limita la reduccion simultanea de picos, emisiones y costos."
    )
    doc.add_heading("1.3 Formulacion del problema", level=2)
    doc.add_paragraph(
        "Problema general: determinar que MADRL gestiona mejor la flexibilidad, CO2 y costos en una comunidad inteligente."
    )
    doc.add_heading("1.4 Objetivos", level=2)
    doc.add_paragraph(
        "Objetivo general: determinar el mejor Multi-Agente de Aprendizaje por Refuerzo Profundo "
        "para la gestion coordinada de flexibilidad energetica, emisiones de CO2 y costos."
    )
    doc.add_paragraph("Objetivos especificos: OE.1 flexibilidad; OE.2 CO2; OE.3 costos.")
    doc.add_heading("1.5 Justificacion del estudio", level=2)
    doc.add_paragraph(
        "La tesis aporta un entorno CityLearn v3 propuesto, formulado como Dec-POMDP con CTDE, "
        "para comparar HAPPO, MASAC, MATD3 y MAAC sobre un dataset local de Iquitos."
    )
    doc.add_heading("1.6 Alcance del estudio", level=2)
    doc.add_paragraph(
        "Alcance actual: dataset real/sintetico auditado, entrenamiento v4 en curso y resultados finales pendientes."
    )

    doc.add_heading("CAPITULO II. MARCO TEORICO", level=1)
    doc.add_heading("2.1 Antecedentes", level=2)
    doc.add_paragraph(
        "Los antecedentes se organizan en MADRL para energia, CityLearn, Dec-POMDP/CTDE, "
        "optimizacion multiobjetivo y gestion de PV/BESS/EV."
    )
    doc.add_heading("2.2 Bases teoricas", level=2)
    doc.add_paragraph(
        "Bases: aprendizaje por refuerzo profundo multiagente, entrenamiento centralizado con ejecucion "
        "descentralizada, recompensas cooperativas y gestion energetica de comunidades inteligentes."
    )
    doc.add_heading("2.3 Definicion de terminos", level=2)
    doc.add_paragraph("MADRL, CTDE, Dec-POMDP, CityLearn v2, CityLearn v3 propuesto, PV, BESS, EV, KPI.")

    doc.add_heading("CAPITULO III. DESARROLLO DEL TRABAJO DE TESIS", level=1)
    doc.add_heading("3.1 Presentacion de la propuesta de solucion", level=2)
    doc.add_paragraph(
        "La propuesta integra CityLearn v2 como simulador base y CityLearn v3 propuesto como capa experimental "
        "MADRL con cuatro backends: HAPPO, MASAC, MATD3 y MAAC."
    )
    doc.add_heading("3.2 Desarrollo de la propuesta de solucion", level=2)
    add_table(
        doc,
        ["Componente", "Valor vigente"],
        [
            ["Dataset", "citylearn_iquitos_2023_2025"],
            ["Edificios", str(csv_manifest.get("schema_references", {}).get("buildings", 17))],
            ["CSV auditados", str(csv_manifest.get("csv_files_checked", "pendiente"))],
            ["NaN / Inf", f"{csv_manifest.get('nan_cells', 'pendiente')} / {csv_manifest.get('inf_cells', 'pendiente')}"],
            ["Cargadores EV", fnum(totals["EV_dimensionador_v3_count"], 0)],
            ["Potencia EV", f"{fnum(totals['EV_dimensionador_v3_KW'] if 'EV_dimensionador_v3_KW' in totals else totals['EV_dimensionador_v3_kW'])} kW"],
            ["PV", f"{fnum(totals['PV_schema_kWp'])} kWp"],
            ["BESS", f"{fnum(totals['BESS_schema_kWh'])} kWh / {fnum(totals['BESS_schema_kW'])} kW"],
            ["Maquinas controladas", str(csv_manifest.get("schema_references", {}).get("controlled_machines", 17))],
            ["Normalizacion permitida", str(ready.get("normalization_allowed_for_training", "pendiente"))],
        ],
    )
    doc.add_paragraph(
        "La regla operacional vigente prioriza generacion solar hacia recarga EV y carga del edificio. "
        "El BESS se dimensiona por edificio con balance PV-EV-red publica-cargas controladas/no controladas "
        "y prioriza EV dentro de la ventana operativa del edificio."
    )

    doc.add_heading("3.3 Analisis de los datos y resultados", level=2)
    add_table(
        doc,
        ["Entrenamiento", "Estado"],
        [
            ["Output root", status.get("output_root", "outputs/citylearn_v3_madrl_oficial_v4")],
            ["Estado global", status.get("status", "pendiente")],
            ["Inicio", status.get("started_at", "pendiente")],
            ["CUDA", str(status.get("cuda", "pendiente"))],
            ["Perfil GPU", status.get("gpu_optimization", {}).get("profile", "pendiente")],
            ["Episodios", str(status.get("episodes", "pendiente"))],
            ["Pasos por episodio", str(status.get("episode_time_steps", "pendiente"))],
            ["Artefactos finales completos", f"{complete_runs}/{total_runs}"],
            ["Corridas pendientes", ", ".join(missing_runs) if missing_runs else "ninguna"],
        ],
    )
    if live:
        doc.add_paragraph(
            "Ultimo progreso vivo disponible: "
            f"{live.get('algorithm', '?')}/{live.get('scenario', '?')} en episodio {live.get('episode', '?')}, "
            f"paso global {live.get('global_step', '?')}, fuente {live.get('_path', '?')}."
        )
    doc.add_paragraph(
        "Los KPIs finales por algoritmo y escenario no se reportan todavia porque faltan artefactos finales completos."
    )

    doc.add_heading("3.4 Discusion e interpretacion de los resultados", level=2)
    doc.add_paragraph(
        "Discusion cuantitativa pendiente hasta completar las 12 corridas y la comparacion estadistica. "
        "La evidencia disponible permite discutir robustez del dataset, configuracion experimental y trazabilidad."
    )
    doc.add_heading("3.5 Estimacion del impacto de la solucion", level=2)
    doc.add_paragraph(
        "Impacto esperado a validar: mejora de flexibilidad, reduccion de emisiones y optimizacion de costos mediante "
        "coordinacion MADRL. No se cuantifica hasta disponer de resultados finales."
    )

    doc.add_heading("CAPITULO IV. CONCLUSIONES Y RECOMENDACIONES", level=1)
    doc.add_heading("4.1 Conclusiones", level=2)
    doc.add_paragraph(
        "Conclusion provisional: el dataset y la arquitectura experimental estan listos y auditados; "
        "las conclusiones de desempeno MADRL quedan pendientes."
    )
    doc.add_heading("4.2 Recomendaciones", level=2)
    doc.add_paragraph(
        "Completar entrenamiento v4, verificar artefactos finales, ejecutar comparacion estadistica y recien cerrar KPIs."
    )

    doc.add_heading("REFERENCIAS", level=1)
    doc.add_paragraph("Insertar referencias APA verificadas desde la matriz bibliografica del skill.")

    doc.add_heading("ANEXOS", level=1)
    doc.add_paragraph("Anexo 1. Matriz de consistencia: pendiente de completar.")
    doc.add_paragraph("Anexo 2. Matriz de operacionalizacion: pendiente de completar.")
    doc.add_paragraph("Anexo 3. Matriz de antecedentes: usar workbook del skill.")
    doc.add_paragraph("Anexo 4. Matriz de KPIs: pendiente de resultados finales.")
    doc.add_paragraph("Anexo 5. Arquitectura CityLearn v3 propuesta: disponible en documentos del proyecto.")
    doc.add_paragraph("Anexo 6. Comparacion de backends MADRL: pendiente de completar entrenamiento.")
    doc.add_paragraph("Anexo 7. Dataset y fuentes: auditorias disponibles en outputs/dataset_audit.")
    doc.add_paragraph("Anexo 8. Hiperparametros: official_full_status.json y configs CityLearn.")
    doc.add_paragraph("Anexo 9. Recompensa multiobjetivo: CityLearnV3MADRLRewardFunction.")
    doc.add_paragraph("Anexo 10. Resultados de simulacion vigentes o pendientes.")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_markdown(output: Path) -> None:
    status = read_json(TRAINING_ROOT / "official_full_status.json")
    csv_manifest = read_json(DATASET_AUDIT / "csv_integrity_manifest.json")
    der_rows = read_csv_rows(DATASET_AUDIT / "der_sizing_audit.csv")
    totals = der_totals(der_rows)
    complete_runs, total_runs, missing_runs = final_artifact_summary()
    lines = [
        f"# {TITLE}",
        "",
        "Borrador Guia N.02 generado con informacion disponible.",
        "",
        "## Evidencia disponible",
        f"- Dataset: `citylearn_iquitos_2023_2025`.",
        f"- CSV auditados: {csv_manifest.get('csv_files_checked', 'pendiente')}.",
        f"- NaN/Inf: {csv_manifest.get('nan_cells', 'pendiente')}/{csv_manifest.get('inf_cells', 'pendiente')}.",
        f"- PV: {fnum(totals['PV_schema_kWp'])} kWp.",
        f"- BESS: {fnum(totals['BESS_schema_kWh'])} kWh / {fnum(totals['BESS_schema_kW'])} kW.",
        f"- EV: {fnum(totals['EV_dimensionador_v3_count'], 0)} cargadores / {fnum(totals['EV_dimensionador_v3_kW'])} kW.",
        f"- Entrenamiento: {status.get('status', 'pendiente')} en `{status.get('output_root', 'outputs/citylearn_v3_madrl_oficial_v4')}`.",
        f"- Artefactos finales completos: {complete_runs}/{total_runs}.",
        f"- Corridas pendientes: {', '.join(missing_runs) if missing_runs else 'ninguna'}.",
        "",
        "## Regla",
        "Los KPIs y conclusiones finales deben quedar pendientes hasta que existan `results.json`, `timeseries.csv` y `trace.csv` por algoritmo/escenario.",
    ]
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", required=True, help="Output DOCX path.")
    parser.add_argument("--markdown", required=True, help="Output Markdown summary path.")
    args = parser.parse_args()
    build_docx(Path(args.docx))
    build_markdown(Path(args.markdown))
    print(Path(args.docx).resolve())
    print(Path(args.markdown).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
