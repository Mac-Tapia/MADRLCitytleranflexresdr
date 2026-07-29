from __future__ import annotations

import csv
import hashlib
import json
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path
import sys

_THESIS_DIR = Path(__file__).resolve().parent
ROOT = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import INFORME, TESIS  # noqa: E402

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def resolve_quantitative_source() -> Path:
    """Prefer locked snapshot; else live Informe; else Tesis canon."""
    snap = ROOT / "docs" / "_working" / "Inforne_tesisV4_master_snapshot.docx"
    for candidate in (snap, INFORME, TESIS):
        if candidate.is_file() and candidate.stat().st_size > 0:
            return candidate
    raise FileNotFoundError(
        "No hay fuente para update_word_quantitative: falta snapshot/_working, Informe o Tesis."
    )


SOURCE_DOCX = resolve_quantitative_source()
OUTPUT_DOCX = INFORME
AUDIT_JSON = ROOT / "docs" / "Inforne_tesisV4_FINAL_REVISADO_50_EPISODIOS_auditoria.json"
ANALYSIS = (
    ROOT
    / "outputs"
    / "madrl_v3_20260627_164047"
    / "resumen_comparativo"
    / "estadistica"
    / "analisis_cuantitativo_completo_50_episodios"
)
DRIVE_URL = (
    "https://drive.google.com/drive/folders/"
    "1ihH6RqL2KpevfCQEUXj7PP1aS2QYssAX"
)
RUN_ID = "madrl_v3_20260627_164047"


def read_csv(name: str) -> list[dict[str, str]]:
    with (ANALYSIS / name).open("r", encoding="utf-8-sig", newline="") as fh:
        return list(csv.DictReader(fh))


def num(value: str | float | int | None) -> float:
    if value in (None, ""):
        return float("nan")
    return float(value)


def fmt(value: str | float | int | None, digits: int = 4) -> str:
    if value in (None, ""):
        return "—"
    val = float(value)
    if abs(val) != 0 and (abs(val) < 0.0001 or abs(val) >= 1_000_000_000):
        return f"{val:.3e}".replace(".", ",")
    return f"{val:,.{digits}f}".replace(",", "§").replace(".", ",").replace("§", ".")


def fmt_p(value: str | float | int | None) -> str:
    if value in (None, ""):
        return "—"
    val = float(value)
    if val < 0.0001:
        return f"{val:.3e}".replace(".", ",")
    return f"{val:.6f}".replace(".", ",")


def find_paragraph(document: Document, startswith: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"No se encontró el párrafo: {startswith}")


def replace_paragraph(document: Document, startswith: str, text: str) -> None:
    paragraph = find_paragraph(document, startswith)
    paragraph.text = text


def exact_approved_texts(document: Document) -> list[str]:
    """Capture the approved problem, objective and hypothesis statements verbatim."""
    prefixes = (
        "¿En qué medida el algoritmo MADRL (aprendizaje",
        "PE.1:",
        "PE.2:",
        "PE.3:",
        "OG. -",
        "OE.1:",
        "OE.2:",
        "OE.3:",
        "H0G.-",
        "H1G.-",
        "HE10.-",
        "HE11.-",
        "HE20.-",
        "HE21.-",
        "HE30.-",
        "HE31.-",
    )
    captured: list[str] = []
    for prefix in prefixes:
        matches = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip().startswith(prefix)
        ]
        if len(matches) != 1:
            raise RuntimeError(
                f"El texto aprobado con prefijo {prefix!r} aparece {len(matches)} veces."
            )
        captured.append(matches[0])
    return captured


def approved_hash(texts: list[str]) -> str:
    return hashlib.sha256("\n".join(texts).encode("utf-8")).hexdigest()


def disable_automatic_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id = num_pr.find(qn("w:numId"))
    if num_id is None:
        num_id = OxmlElement("w:numId")
        num_pr.append(num_id)
    num_id.set(qn("w:val"), "0")


def normalize_explicit_numbering(document: Document) -> list[dict[str, str]]:
    """Make Chapters 4–6 explicit and internally consistent without touching Chapter 1."""
    replacements = {
        "Desarrollo del sistema y arquitectura de software": "4.1 Desarrollo del sistema y arquitectura de software",
        "Algoritmos": "4.4 Algoritmos",
        "Aportes originales al motor de simulación": "4.6 Aportes originales al motor de simulación",
        "Diseño experimental: matriz de 12 corridas": "4.7 Diseño experimental: matriz de 12 corridas",
        "Implementación": "4.8 Implementación",
        "Detalle técnico de la propuesta derivado del resumen": "4.9 Detalle técnico de la propuesta derivado del resumen",
        "Construcción del dataset y fuentes utilizadas": "4.9.1 Construcción del dataset y fuentes utilizadas",
        "Librerias, scripts y herramientas de implementación": "4.9.2 Librerías, scripts y herramientas de implementación",
        "Arquitectura operativa desde datos hasta resultados": "4.9.3 Arquitectura operativa desde datos hasta resultados",
        "Delimitación entre equipos controlados y no controlados": "4.9.4 Delimitación entre equipos controlados y no controlados",
    }
    changes: list[dict[str, str]] = []
    chapter_4_seen = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "Capítulo 4. Desarrollo de la propuesta":
            chapter_4_seen = True
        if not chapter_4_seen or not paragraph.style.name.startswith("Heading"):
            continue
        if text in replacements:
            new_text = replacements[text]
            paragraph.text = new_text
            changes.append({"before": text, "after": new_text})
        disable_automatic_numbering(paragraph)
    return changes


def replace_table_cell_text(document: Document, old: str, new: str) -> int:
    replacements = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == old:
                    cell.text = new
                    replacements += 1
    return replacements


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)


def request_field_update_on_open(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def capture_removed_figure_blocks(document: Document) -> list[dict[str, object]]:
    """Preserve figures from ranges that are rewritten, retaining their image relationships."""
    chapter_5_start = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip().startswith("Capítulo 5.")
    )
    references_start = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip().startswith("Referencias bibliográficas")
    )
    annex_a_start = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip().startswith("Anexo A.")
    )
    annex_b_start = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip().startswith("Anexo B.")
    )
    ranges = ((chapter_5_start, references_start), (annex_a_start, annex_b_start))
    blocks: list[dict[str, object]] = []
    for index, paragraph in enumerate(document.paragraphs):
        in_rewritten_range = any(start <= index < end for start, end in ranges)
        if not in_rewritten_range or not paragraph._p.xpath(".//w:drawing | .//w:pict"):
            continue
        caption = document.paragraphs[index - 1].text.strip() if index > 0 else "Figura"
        title_parts = [part.strip() for part in caption.splitlines() if part.strip()]
        title = title_parts[-1] if title_parts else "Evidencia gráfica"
        blocks.append({"title": title, "xml": deepcopy(paragraph._p)})
    return blocks


def remove_range(start_paragraph, end_paragraph) -> None:
    """Remove body elements from start inclusive up to end exclusive."""
    node = start_paragraph._p
    end = end_paragraph._p
    while node is not end:
        nxt = node.getnext()
        node.getparent().remove(node)
        if nxt is None:
            raise RuntimeError("El rango XML no alcanzó el párrafo final.")
        node = nxt


def move_before_anchor(block, anchor) -> None:
    anchor._p.addprevious(block)


def add_paragraph_before(
    document: Document,
    anchor,
    text: str = "",
    style: str | None = None,
    bold_prefix: str | None = None,
):
    paragraph = document.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    move_before_anchor(paragraph._p, anchor)
    return paragraph


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(7.4)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_before(
    document: Document,
    anchor,
    headers: list[str],
    rows: list[list[str]],
    caption: str,
    note: str,
):
    cap = add_paragraph_before(document, anchor, caption, style="TablaIndice")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        shade_cell(table.rows[0].cells[idx], "1F4E78")
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            set_cell_text(cells[col_idx], value)
            if row_idx % 2:
                shade_cell(cells[col_idx], "EAF2F8")
    move_before_anchor(table._tbl, anchor)
    note_paragraph = add_paragraph_before(document, anchor, note)
    note_paragraph.paragraph_format.space_after = Pt(6)
    return table


def add_heading_before(document: Document, anchor, text: str, level: int):
    paragraph = add_paragraph_before(document, anchor, text, style=f"Heading {level}")
    if level == 1:
        paragraph.paragraph_format.page_break_before = True
    return paragraph


coverage = read_csv("run_coverage_quantitative.csv")
episodes = read_csv("episode_descriptive_50.csv")
primary = read_csv("primary_objective_values.csv")
effects = read_csv("objective_effect_descriptive.csv")
impact = read_csv("objective_impact_wilcoxon_holm.csv")
tests = read_csv("objective_friedman_posthoc_holm.csv")
ranking = read_csv("global_ranking_topsis.csv")
decisions = read_csv("hypothesis_decisions_quantitative.csv")
delta_audit = read_csv("stored_delta_consistency_audit.csv")

primary_index = {
    (row["axis"], row["algorithm"]): row
    for row in primary
}
effect_index = {
    (row["axis"], row["algorithm"]): row
    for row in effects
}
impact_index = {
    (row["axis"], row["algorithm"]): row
    for row in impact
}
friedman_index = {
    row["axis"]: row
    for row in tests
    if row["test"] == "Friedman"
}
decision_index = {row["scope"]: row for row in decisions}
delta_index = {
    (row["algorithm"], row["scenario"], row["family"]): row
    for row in delta_audit
}


document = Document(SOURCE_DOCX)
approved_before = exact_approved_texts(document)
approved_hash_before = approved_hash(approved_before)
preserved_figure_blocks = capture_removed_figure_blocks(document)
backup_dir = ROOT / "outputs" / "_word_backups"
backup_dir.mkdir(parents=True, exist_ok=True)
backup = backup_dir / f"{SOURCE_DOCX.stem}_antes_analisis_50ep_{datetime.now():%Y%m%d_%H%M%S}.docx"
shutil.copy2(SOURCE_DOCX, backup)


replace_paragraph(
    document,
    "La evidencia experimental corresponde exclusivamente",
    (
        "La evidencia se recalculó exclusivamente con los resultados reales de los 50 episodios "
        f"de cada uno de los 12 tratamientos de la corrida {RUN_ID}. Se auditaron 600 filas "
        "episódicas, 668 KPI distritales, 15.300 KPI por edificio y 204 resúmenes de comportamiento. "
        "En E1, MATD3 presentó el menor deterioro de flexibilidad (-0,0923 %), sin impacto ni "
        "diferencias significativas (Friedman p=0,614935). En E2, MATD3 presentó el menor aumento "
        "de CO₂ (-0,4193 % de efecto favorable, equivalente a +46.014,75 kgCO₂), con diferencias "
        "entre algoritmos (p=1,461×10⁻⁸). En E3, MAAC registró el menor incremento de costo "
        "(-0,2675 %, +18.978,30 EUR), sin diferencia global (p=0,109906). TOPSIS con pesos iguales "
        "identificó a MAAC (0,936627) como líder descriptivo global; Friedman global no alcanzó "
        "significancia (p=0,085801). Por ello, se cumplen los objetivos de cuantificación y "
        "selección descriptiva, pero no se demuestra superioridad estadística global. La inferencia "
        "es exploratoria porque existe una sola semilla independiente por tratamiento."
    ),
)
replace_paragraph(
    document,
    "The experimental evidence is restricted",
    (
        "All evidence was recalculated exclusively from the real 50-episode results for each of "
        f"the 12 treatments in run {RUN_ID}: 600 episode rows, 668 district KPI values, 15,300 "
        "building KPI values and 204 building-behaviour summaries. MATD3 showed the smallest "
        "flexibility deterioration in E1 (-0.0923%; Friedman p=0.614935) and the smallest CO2 "
        "increase in E2 (-0.4193% favourable effect, +46,014.75 kgCO2; p=1.461×10^-8). MAAC "
        "showed the smallest E3 cost increase (-0.2675%, +EUR 18,978.30; p=0.109906). Equal-weight "
        "TOPSIS ranked MAAC first globally (0.936627), while the global Friedman test was not "
        "significant (p=0.085801). The objectives are therefore fulfilled quantitatively and a "
        "descriptive leader is identified, but global statistical superiority is not demonstrated. "
        "Inference remains exploratory because only one independent seed is available per treatment."
    ),
)

replace_paragraph(
    document,
    "La evidencia experimental conservada se limita",
    (
        f"La evidencia experimental conservada se limita a la corrida canónica {RUN_ID}. "
        "Se dispone de 600 registros episódicos: 50 episodios para cada combinación de los cuatro "
        "algoritmos y los tres escenarios. En HAPPO, los episodios 0–48 proceden del archivo "
        "histórico y el episodio 49 del resultado post-resume; la unión se realizó por índice, "
        "sin inventar ni imputar valores. La inferencia se interpreta como exploratoria porque "
        "solo existe una semilla independiente por tratamiento."
    ),
)
replace_paragraph(
    document,
    "La evidencia empírica se restringe a los artefactos",
    (
        f"La evidencia empírica se restringe a los artefactos materializados de la corrida canónica "
        f"{RUN_ID}. La base auditada contiene 600 registros episódicos, equivalentes a 50 episodios "
        "por cada uno de los 12 tratamientos algoritmo–escenario."
    ),
)
replace_paragraph(
    document,
    "La cobertura desigual de HAPPO obliga",
    (
        "La serie completa de HAPPO fue reconstruida de manera trazable a partir de 49 episodios "
        "históricos y el episodio final post-resume. Por tanto, los cuatro algoritmos poseen 50 "
        "episodios descriptivos por escenario. Esta igualdad de longitud no convierte los episodios "
        "en réplicas independientes: las decisiones inferenciales principales se apoyan en bloques "
        "pareados de KPI o edificios y se reportan como evidencia intra-corrida exploratoria."
    ),
)
replace_paragraph(
    document,
    "La unidad experimental principal es el tratamiento algoritmo-escenario",
    (
        "La unidad experimental principal es la semilla de entrenamiento dentro de cada tratamiento "
        "algoritmo–escenario; en la corrida canónica solo se materializó seed=0. La unidad de "
        "observación episódica es el episodio y la unidad espacial es el edificio-agente. Los cuatro "
        "algoritmos conservan 50 episodios por escenario, pero estos episodios no se tratan como "
        "réplicas independientes. Para responder OE.1 se emplean los tres KPI primarios como bloques "
        "pareados; para OE.2 y OE.3, los 17 edificios. La inferencia es, por ello, intra-corrida y "
        "exploratoria."
    ),
)
replace_paragraph(
    document,
    "La decisión de usar estadística no paramétrica se debe",
    (
        "La estadística no paramétrica se selecciona porque los KPI de aprendizaje por refuerzo "
        "pueden presentar asimetría, valores atípicos y dependencia respecto de la política. La "
        "diferencia global entre los cuatro algoritmos se contrasta mediante Friedman sobre bloques "
        "pareados, con Kendall W como tamaño de efecto. Cuando corresponde, se aplican comparaciones "
        "post hoc de Wilcoxon pareado y corrección de Holm. El impacto frente al baseline también se "
        "evalúa con Wilcoxon-Holm. Esta estrategia es coherente con la recomendación de no asumir "
        "supuestos estadísticos no verificados en aprendizaje por refuerzo; sus p-valores se "
        "interpretan como evidencia intra-corrida exploratoria por existir una sola semilla."
    ),
)
replace_paragraph(
    document,
    "Como parte del control de calidad metodológico",
    (
        "Como parte del control de calidad metodológico, se adoptó la cobertura auditada de 600 filas "
        "episódicas: 50 por cada uno de los 12 tratamientos. HAPPO combina los episodios históricos "
        "0–48 con el episodio 49 post-resume, sin imputación. Las pruebas Friedman, Wilcoxon pareado "
        "con ajuste Holm y Kendall W se aplican a los KPI o edificios emparejados definidos para cada "
        "objetivo; los episodios se reservan para estadística descriptiva del aprendizaje. Cada tabla "
        "conserva una nota de fuente vinculada con results.json, training_summary.json, timeseries.csv, "
        "trace.csv, building_kpis.csv y checkpoint_manifest.json."
    ),
)
replace_paragraph(
    document,
    "La corrida canónica madrl_v3_20260627_164047 se ejecutó",
    (
        f"La corrida canónica {RUN_ID} se ejecutó en Google Colab sobre NVIDIA A100-SXM4-80GB "
        "mediante colab_a100_official_launcher.py y el protocolo two_phase_happo_masac_v3. La "
        "primera fase entrenó HAPPO y MASAC y la segunda MATD3 y MAAC. El monitor registró progreso, "
        "uso de GPU, recompensa y métricas energéticas; los archivos de progreso en vivo son "
        "transitorios y no sustituyen los artefactos finales. La cobertura auditada es de 50 "
        "episodios por escenario para los cuatro algoritmos; en HAPPO, el episodio final proviene "
        "del resultado post-resume."
    ),
)
# Replace Chapter 4.10 with the exact quantitative-analysis protocol.
chapter_4_10 = find_paragraph(document, "Respuesta operacional a las preguntas específicas")
chapter_5_anchor = find_paragraph(document, "Capítulo 5.")
remove_range(chapter_4_10, chapter_5_anchor)

add_heading_before(
    document,
    chapter_5_anchor,
    "4.10 Articulación de problemas, objetivos, hipótesis y análisis cuantitativo",
    2,
)
add_paragraph_before(
    document,
    chapter_5_anchor,
    (
        f"El corpus analítico corresponde únicamente a la corrida real {RUN_ID}, sincronizada con "
        f"la carpeta de Google Drive {DRIVE_URL}. El diseño contiene 4 algoritmos × 3 escenarios × "
        "50 episodios = 600 observaciones episódicas. No se generaron datos sintéticos, no se "
        "rellenaron KPI ausentes y no se utilizó recompensa publicada en otros artículos como "
        "sustituto de los resultados del proyecto."
    ),
)
add_heading_before(document, chapter_5_anchor, "4.10.1 Regla de impacto y sentido favorable", 3)
add_paragraph_before(
    document,
    chapter_5_anchor,
    (
        "Para flexibilidad (PE.1/OE.1/HE10–HE11), la operacionalización define tres indicadores "
        "primarios normalizados respecto del baseline: peak_average, ramping_average y "
        "one_minus_load_factor_average. El compuesto se calcula como F=(peak+ramping+OMLF)/3 y "
        "el efecto favorable como 100×(1−F). Para CO₂ y costo, el cambio absoluto se recalcula "
        "como Δ=control−baseline y el efecto favorable como 100×(baseline−control)/baseline. "
        "Así, un efecto favorable positivo representa reducción/mejora y uno negativo representa "
        "deterioro/aumento."
    ),
)
add_paragraph_before(
    document,
    chapter_5_anchor,
    (
        "La auditoría encontró 24 discrepancias entre los campos delta almacenados y la resta directa "
        "control−baseline (dos familias, carbono y costo, en los 12 tratamientos). En consecuencia, "
        "las decisiones usan la resta verificable de los totales; los delta originales se conservan "
        "solo para trazabilidad y no para decidir objetivos o hipótesis."
    ),
)
add_heading_before(document, chapter_5_anchor, "4.10.2 Unidad de análisis y pruebas", 3)
add_paragraph_before(
    document,
    chapter_5_anchor,
    (
        "OE.1 usa como bloques emparejados sus tres KPI primarios. OE.2 y OE.3 usan los 17 edificios "
        "como bloques emparejados. El impacto de cada algoritmo frente al baseline se contrasta con "
        "Wilcoxon y corrección de Holm; la diferencia entre cuatro algoritmos se contrasta con "
        "Friedman, Kendall W y post hoc Wilcoxon-Holm. La unidad independiente de entrenamiento es "
        "la semilla, y solo existe seed=0; por ello, los p-valores por KPI/edificio son evidencia "
        "intra-corrida exploratoria y no reemplazan un experimento multi-semilla."
    ),
)
add_heading_before(document, chapter_5_anchor, "4.10.3 Regla explícita de decisión", 3)
add_paragraph_before(
    document,
    chapter_5_anchor,
    (
        "Se fija α=0,05. Las hipótesis alternativas específicas requieren dos condiciones conjuntas: "
        "(a) impacto significativo frente al baseline y (b) diferencias significativas entre los "
        "algoritmos del escenario. H1G requiere que la condición se sostenga de manera integrada en "
        "las tres dimensiones y que exista diferencia global. Cuando una condición no se cumple, "
        "se informa «no rechazar la nula» o «no respaldar la alternativa»; no se confunde ausencia "
        "de evidencia con prueba de igualdad."
    ),
)


# Replace Chapter 5 completely so that no stale 49-episode or 3-algorithm claims remain.
chapter_5_start = find_paragraph(document, "Capítulo 5.")
chapter_6_anchor = find_paragraph(document, "Capítulo 6.")
remove_range(chapter_5_start, chapter_6_anchor)

add_heading_before(
    document,
    chapter_6_anchor,
    "Capítulo 5. Resultados y contrastación de hipótesis",
    1,
)
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        f"Los resultados de este capítulo proceden de {RUN_ID} y de sus archivos reales conservados "
        "en Drive. La ejecución del script reproducible materializó 600 filas episódicas, 668 valores "
        "KPI distritales, 648 KPI clasificados en los tres ejes, 216 KPI del escenario objetivo, "
        "104 ganancias directamente comparables, 15.300 valores KPI por edificio y 204 resúmenes "
        "edificio–tratamiento. Todos los valores presentados son cuantitativos y trazables."
    ),
)

add_heading_before(document, chapter_6_anchor, "5.1 Auditoría de cobertura de los 50 episodios", 2)
coverage_rows = []
for row in coverage:
    coverage_rows.append(
        [
            row["algorithm"],
            row["scenario"],
            row["episodes_recorded"],
            row["district_kpis"],
            row["buildings"],
            row["building_kpi_rows"],
            row["building_behavior_rows"],
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "Esc.", "Episodios", "KPI distrito", "Edificios", "KPI edificio", "Resúmenes"],
    coverage_rows,
    "Tabla 5.1. Cobertura cuantitativa real de los 12 tratamientos",
    (
        "Nota. Cada tratamiento contiene 50 episodios, 17 edificios y 1.275 valores KPI por edificio "
        "(75 tipos × 17). HAPPO conserva los episodios 0–48 en el archivo histórico y el episodio 49 "
        "en el resultado post-resume; se unieron por índice 0–49 sin imputación. Fuente: results.json, "
        "building_kpis.csv y building_behavior_summary.csv del espejo Drive."
    ),
)

add_heading_before(document, chapter_6_anchor, "5.2 Estadística descriptiva de los 600 episodios", 2)
episode_rows = []
for row in episodes:
    episode_rows.append(
        [
            row["algorithm"],
            row["scenario"],
            row["n"],
            fmt(row["mean"], 6),
            fmt(row["median"], 6),
            fmt(row["std"], 6),
            fmt(row["min"], 6),
            fmt(row["max"], 6),
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "Esc.", "n", "Media reward", "Mediana", "DE", "Mín.", "Máx."],
    episode_rows,
    "Tabla 5.2. Descriptivos de reward_mean_average en 50 episodios por tratamiento",
    (
        "Nota. La recompensa describe el aprendizaje intra-corrida y no sustituye los KPI físicos, "
        "ambientales y económicos usados para responder los objetivos. n=50 en las 12 combinaciones; "
        "total=600 episodios."
    ),
)
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        "Las medias episódicas varían entre -0,633044 (MATD3-E1) y -0,484054 (HAPPO-E2). "
        "El mayor cambio entre el primer y el último episodio fue MATD3-E1 (+0,101981), seguido de "
        "HAPPO-E1 (+0,088147). Estas cifras evidencian evolución de la recompensa, pero el cumplimiento "
        "de OE.1–OE.3 se determina con los indicadores de la operacionalización."
    ),
)

add_heading_before(
    document,
    chapter_6_anchor,
    "5.3 Respuesta a PE.1, cumplimiento de OE.1 y contrastación de HE10/HE11: flexibilidad energética en E1",
    2,
)
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        "PE.1 pregunta en qué medida los algoritmos impactan la flexibilidad y cuál presenta el mejor "
        "desempeño en E1. OE.1 exige cuantificar ese impacto e identificar al líder. Se evaluaron los "
        "tres indicadores primarios definidos en la operacionalización; baseline=1, por lo que un "
        "ratio inferior a 1 representa mejora y uno superior a 1 deterioro."
    ),
)
oe1_rows = []
for algorithm in ["HAPPO", "MAAC", "MASAC", "MATD3"]:
    row = primary_index[("OE1", algorithm)]
    oe1_rows.append(
        [
            algorithm,
            fmt(row["peak_average"], 6),
            fmt(row["ramping_average"], 6),
            fmt(row["one_minus_load_factor_average"], 6),
            fmt(row["control"], 6),
            fmt(row["favorable_effect_percent"], 4) + " %",
            row["rank_within_objective"],
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "Pico", "Ramping", "1−factor carga", "Compuesto F", "Efecto favorable", "Rango"],
    oe1_rows,
    "Tabla 5.3. Valores primarios de flexibilidad y efecto frente al baseline en E1",
    (
        "Nota. F=(pico+ramping+1−factor de carga)/3; efecto=100×(1−F). Los cuatro efectos compuestos "
        "son negativos: no hubo mejora integral de flexibilidad. MATD3 ocupa el primer lugar porque "
        "su deterioro fue el menor (-0,0923 %)."
    ),
)
oe1_test = friedman_index["OE1"]
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        "Impacto numérico. HAPPO: -11,0541 % (componentes pico -39,0323 %, ramping -8,3909 % y "
        "1−factor de carga +14,2609 %); MAAC: -1,2426 % (-3,0702 %, -1,0994 %, +0,4418 %); "
        "MASAC: -2,8553 % (-8,5063 %, -2,1969 %, +2,1374 %); MATD3: -0,0923 % (-0,8124 %, "
        "-0,1176 %, +0,6531 %)."
    ),
)
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        f"Contrastación. Friedman: χ²={fmt(oe1_test['statistic'], 4)}, gl=3, "
        f"p={fmt_p(oe1_test['p_value'])}, Kendall W={fmt(oe1_test['kendall_w'], 4)}. "
        "Ningún Wilcoxon contra baseline fue significativo tras Holm (p-Holm=1,0000 para los "
        "cuatro algoritmos) y ningún par fue significativo. Decisión: no se rechaza HE10 y no se "
        "respalda HE11. OE.1 se cumple porque se cuantificaron los cuatro impactos y se identificó "
        "a MATD3 como líder relativo, sin afirmar superioridad estadística."
    ),
)

add_heading_before(
    document,
    chapter_6_anchor,
    "5.4 Respuesta a PE.2, cumplimiento de OE.2 y contrastación de HE20/HE21: emisiones de CO₂ en E2",
    2,
)
oe2_rows = []
for algorithm in ["HAPPO", "MAAC", "MASAC", "MATD3"]:
    row = primary_index[("OE2", algorithm)]
    oe2_rows.append(
        [
            algorithm,
            fmt(row["baseline"], 2),
            fmt(row["control"], 2),
            fmt(row["control_minus_baseline"], 2),
            fmt(row["favorable_effect_percent"], 4) + " %",
            row["rank_within_objective"],
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "Baseline kgCO₂", "Control kgCO₂", "Δ control−base", "Efecto favorable", "Rango"],
    oe2_rows,
    "Tabla 5.4. Emisiones totales verificadas y efecto en E2",
    (
        "Nota. Un Δ positivo significa más emisiones; el efecto favorable es negativo cuando control "
        "supera baseline. MATD3 es el líder relativo porque presentó el menor incremento absoluto "
        "(46.014,75 kgCO₂) y porcentual (-0,4193 %)."
    ),
)
oe2_build_rows = []
for algorithm in ["HAPPO", "MAAC", "MASAC", "MATD3"]:
    desc = effect_index[("OE2", algorithm)]
    imp = impact_index[("OE2", algorithm)]
    oe2_build_rows.append(
        [
            algorithm,
            desc["n"],
            fmt(desc["mean"], 4) + " %",
            fmt(desc["median"], 4) + " %",
            f"{desc['favorable_units']}/{desc['unfavorable_units']}",
            fmt_p(imp["p_holm"]),
            "Sí" if imp["significant_holm"] == "True" else "No",
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "n edif.", "Media efecto", "Mediana", "Fav./desfav.", "p-Holm", "Impacto sig."],
    oe2_build_rows,
    "Tabla 5.5. Efecto de CO₂ por edificio y Wilcoxon-Holm en E2",
    (
        "Nota. Efecto favorable porcentual por edificio. HAPPO, MAAC y MASAC aumentaron emisiones en "
        "los 17 edificios; MATD3 mejoró 5 y empeoró 12."
    ),
)
oe2_test = friedman_index["OE2"]
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        f"Friedman detectó diferencias: χ²={fmt(oe2_test['statistic'], 4)}, gl=3, "
        f"p={fmt_p(oe2_test['p_value'])}, W={fmt(oe2_test['kendall_w'], 4)}. Tras Holm fueron "
        "significativos HAPPO–MAAC, HAPPO–MASAC, HAPPO–MATD3 y MASAC–MATD3; MAAC–MASAC y "
        "MAAC–MATD3 no lo fueron. Decisión: se rechaza HE20 y se respalda HE21 dentro de la "
        "inferencia exploratoria por edificios. El impacto significativo fue desfavorable, porque "
        "los cuatro totales de control superaron al baseline. OE.2 se cumple y MATD3 es el mejor "
        "desempeño relativo de E2, no una reducción neta de CO₂."
    ),
)

add_heading_before(
    document,
    chapter_6_anchor,
    "5.5 Respuesta a PE.3, cumplimiento de OE.3 y contrastación de HE30/HE31: costos energéticos en E3",
    2,
)
oe3_rows = []
for algorithm in ["HAPPO", "MAAC", "MASAC", "MATD3"]:
    row = primary_index[("OE3", algorithm)]
    oe3_rows.append(
        [
            algorithm,
            fmt(row["baseline"], 2),
            fmt(row["control"], 2),
            fmt(row["control_minus_baseline"], 2),
            fmt(row["favorable_effect_percent"], 4) + " %",
            row["rank_within_objective"],
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "Baseline EUR", "Control EUR", "Δ control−base", "Efecto favorable", "Rango"],
    oe3_rows,
    "Tabla 5.6. Costos totales verificados y efecto en E3",
    (
        "Nota. Todos los Δ son positivos. MAAC presenta el menor aumento absoluto (+18.978,30 EUR) "
        "y porcentual (-0,2675 %), seguido de MASAC, MATD3 y HAPPO."
    ),
)
oe3_build_rows = []
for algorithm in ["HAPPO", "MAAC", "MASAC", "MATD3"]:
    desc = effect_index[("OE3", algorithm)]
    imp = impact_index[("OE3", algorithm)]
    oe3_build_rows.append(
        [
            algorithm,
            desc["n"],
            fmt(desc["mean"], 4) + " %",
            fmt(desc["median"], 4) + " %",
            f"{desc['favorable_units']}/{desc['unfavorable_units']}",
            fmt_p(imp["p_holm"]),
            "Sí" if imp["significant_holm"] == "True" else "No",
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "n edif.", "Media efecto", "Mediana", "Fav./desfav.", "p-Holm", "Impacto sig."],
    oe3_build_rows,
    "Tabla 5.7. Efecto de costo por edificio y Wilcoxon-Holm en E3",
    (
        "Nota. HAPPO presenta heterogeneidad: 12 edificios favorables y 5 desfavorables, aunque el "
        "total distrital aumenta. MATD3 muestra impacto adverso significativo tras Holm."
    ),
)
oe3_test = friedman_index["OE3"]
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        f"Friedman no detectó diferencia entre algoritmos: χ²={fmt(oe3_test['statistic'], 4)}, "
        f"gl=3, p={fmt_p(oe3_test['p_value'])}, W={fmt(oe3_test['kendall_w'], 4)}; ningún post hoc "
        "fue significativo. Wilcoxon-Holm detectó impacto solo en MATD3 "
        f"(p-Holm={fmt_p(impact_index[('OE3', 'MATD3')]['p_holm'])}), de sentido desfavorable. "
        "La condición conjuntiva de HE31 no se cumple porque no hay diferencia global ni post hoc. "
        "Decisión: no se respalda HE31; HE30 no se rechaza como proposición conjunta, aunque su cláusula "
        "de ausencia total de impacto falla para MATD3. OE.3 se cumple y MAAC es el líder descriptivo."
    ),
)

add_heading_before(
    document,
    chapter_6_anchor,
    "5.6 Problema general, objetivo general y contrastación H0G/H1G",
    2,
)
rank_rows = []
for row in sorted(ranking, key=lambda x: int(float(x["rank_topsis"]))):
    rank_rows.append(
        [
            row["algorithm"],
            fmt(row["effect_flexibility_percent"], 4) + " %",
            fmt(row["effect_co2_percent"], 4) + " %",
            fmt(row["effect_cost_percent"], 4) + " %",
            fmt(row["mean_raw_effect_percent"], 4) + " %",
            fmt(row["topsis_equal_weight"], 6),
            row["rank_topsis"],
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "Flex. E1", "CO₂ E2", "Costo E3", "Media cruda", "TOPSIS", "Rango"],
    rank_rows,
    "Tabla 5.8. Integración global de los tres efectos primarios",
    (
        "Nota. TOPSIS usa igual peso (1/3) para flexibilidad, carbono y costo después de normalizar "
        "cada dimensión. Todos los efectos primarios son negativos; el ranking identifica el menor "
        "deterioro relativo, no una mejora absoluta."
    ),
)
global_decision = decision_index["OG"]
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        "TOPSIS ubica a MAAC en primer lugar (0,936627), MATD3 segundo (0,866227), MASAC tercero "
        "(0,840362) y HAPPO cuarto (0,000000). La media cruda de porcentajes favorece a MATD3 "
        "(-0,4766 %) frente a MAAC (-0,9543 %), lo que evidencia sensibilidad a la regla de "
        "agregación. Friedman sobre las tres dimensiones produce χ²=6,6000, p=0,085801 y "
        "Kendall W=0,7333: no alcanza α=0,05."
    ),
)
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        "Respuesta general. El OG se cumple cuantitativamente porque se integraron flexibilidad, CO₂ "
        "y costo, se calcularon magnitudes y se identificó a MAAC como líder TOPSIS bajo la regla "
        "preespecificada de pesos iguales. Sin embargo, H1G exige impacto y diferencias en las tres "
        "dimensiones y una diferencia global; la condición no se cumple. Decisión: no se rechaza H0G "
        "y no se respalda H1G. No existe evidencia suficiente para declarar un ganador estadístico "
        "global único."
    ),
)

add_heading_before(document, chapter_6_anchor, "5.7 Auditoría de los campos delta", 2)
delta_co2_rows = []
for algorithm in ["HAPPO", "MAAC", "MASAC", "MATD3"]:
    row = delta_index[(algorithm, "E2", "carbon_emissions")]
    delta_co2_rows.append(
        [
            algorithm,
            fmt(row["baseline"], 2),
            fmt(row["control"], 2),
            fmt(row["stored_delta"], 2),
            fmt(row["recomputed_control_minus_baseline"], 2),
            fmt(row["relative_discrepancy_pct"], 4) + " %",
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "Baseline", "Control", "Delta guardado", "Delta recalculado", "Discrepancia"],
    delta_co2_rows,
    "Tabla 5.9. Verificación de delta de CO₂ en el escenario E2",
    (
        "Nota. La discrepancia relativa del delta guardado frente a control−baseline es aproximadamente "
        "-49,8630 % en los cuatro algoritmos. La contrastación usa el delta recalculado."
    ),
)
delta_cost_rows = []
for algorithm in ["HAPPO", "MAAC", "MASAC", "MATD3"]:
    row = delta_index[(algorithm, "E3", "electricity_cost")]
    delta_cost_rows.append(
        [
            algorithm,
            fmt(row["baseline"], 2),
            fmt(row["control"], 2),
            fmt(row["stored_delta"], 2),
            fmt(row["recomputed_control_minus_baseline"], 2),
            fmt(row["relative_discrepancy_pct"], 4) + " %",
        ]
    )
add_table_before(
    document,
    chapter_6_anchor,
    ["Algoritmo", "Baseline", "Control", "Delta guardado", "Delta recalculado", "Discrepancia"],
    delta_cost_rows,
    "Tabla 5.10. Verificación de delta de costo en el escenario E3",
    (
        "Nota. Los cuatro delta guardados no coinciden con control−baseline. El mismo patrón aparece "
        "en carbono y costo de E1–E3: 24 inconsistencias auditadas. No se alteran los archivos fuente; "
        "se documenta la corrección analítica."
    ),
)

add_heading_before(document, chapter_6_anchor, "5.8 Síntesis explícita de objetivos e hipótesis", 2)
hyp_rows = [
    [
        "OE.1 / HE10–HE11",
        "MATD3; -0,0923 %",
        "χ²=1,8000; p=0,614935; W=0,2000",
        "No rechazar HE10; no respaldar HE11",
        "Cumplido",
    ],
    [
        "OE.2 / HE20–HE21",
        "MATD3; -0,4193 %",
        "χ²=39,3529; p=1,461×10⁻⁸; W=0,7716",
        "Rechazar HE20; respaldar HE21 (impacto adverso)",
        "Cumplido",
    ],
    [
        "OE.3 / HE30–HE31",
        "MAAC; -0,2675 %",
        "χ²=6,0353; p=0,109906; W=0,1183",
        "No respaldar HE31; HE30 no se rechaza conjuntamente",
        "Cumplido",
    ],
    [
        "OG / H0G–H1G",
        "MAAC; TOPSIS=0,936627",
        "χ²=6,6000; p=0,085801; W=0,7333",
        "No rechazar H0G; no respaldar H1G",
        "Cumplido cuantitativamente",
    ],
]
add_table_before(
    document,
    chapter_6_anchor,
    ["Alcance", "Líder y magnitud", "Prueba global", "Decisión", "Objetivo"],
    hyp_rows,
    "Tabla 5.11. Respuesta cuantitativa final a objetivos e hipótesis",
    (
        "Nota. «Cumplido» significa que se cuantificó el impacto y se identificó el mejor desempeño "
        "relativo. No implica que la hipótesis alternativa haya sido respaldada."
    ),
)

add_heading_before(document, chapter_6_anchor, "5.9 Discusión triangulada", 2)
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        "La triangulación entre teoría, operacionalización y evidencia conduce a tres hallazgos. "
        "Primero, el control continuo no garantiza reducción simultánea de pico, ramping y carga: "
        "MATD3 fue el más cercano al baseline de flexibilidad, pero su compuesto aún se deterioró "
        "0,0923 %. Segundo, la coordinación ambiental produjo diferencias fuertes entre algoritmos "
        "(W=0,7716), aunque todos aumentaron el CO₂ total; por tanto, «mejor» significa menor daño "
        "relativo y no éxito ambiental absoluto. Tercero, MAAC minimizó el incremento económico, "
        "pero la diferencia entre algoritmos de E3 no fue significativa."
    ),
)
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        "El cambio de líder entre dimensiones coincide con los antecedentes que muestran dependencia "
        "del entorno, del espacio de acciones y del KPI. Marchesini et al. (2026) observaron cambios "
        "de ranking al escalar tareas de red, y Wu et al. (2026) documentaron compromisos entre costo, "
        "desprendimiento de carga y tiempo computacional. En esta tesis, TOPSIS ofrece una regla "
        "transparente de selección global, pero la falta de significancia global y la sensibilidad "
        "frente a la media cruda obligan a presentar a MAAC como líder descriptivo, no como ganador "
        "universal."
    ),
)
add_paragraph_before(
    document,
    chapter_6_anchor,
    (
        "La evidencia por edificio revela heterogeneidad que el total distrital puede ocultar: HAPPO "
        "mejoró el costo en 12 de 17 edificios de E3, pero el agregado aumentó 213.072,16 EUR; MATD3 "
        "mejoró CO₂ en 5 de 17 edificios de E2, pero el agregado aumentó 46.014,75 kgCO₂. Este resultado "
        "justifica conservar simultáneamente totales distritales, efectos relativos por edificio y "
        "pruebas pareadas."
    ),
)


# Replace Chapter 6 with conclusions consistent with the recalculation.
chapter_6_start = find_paragraph(document, "Capítulo 6.")
references_anchor = find_paragraph(document, "Referencias bibliográficas")
remove_range(chapter_6_start, references_anchor)

add_heading_before(document, references_anchor, "Capítulo 6. Conclusiones", 1)
add_heading_before(document, references_anchor, "6.1 Conclusiones por objetivo", 2)
add_paragraph_before(
    document,
    references_anchor,
    (
        "OE.1. Los cuatro MADRL deterioraron el compuesto primario de flexibilidad en E1. MATD3 fue "
        "el mejor desempeño relativo con -0,0923 %, seguido de MAAC (-1,2426 %), MASAC (-2,8553 %) "
        "y HAPPO (-11,0541 %). Friedman p=0,614935 y ningún contraste Holm fue significativo. OE.1 "
        "se cumplió; no se rechazó HE10 ni se respaldó HE11."
    ),
)
add_paragraph_before(
    document,
    references_anchor,
    (
        "OE.2. Las emisiones de control superaron al baseline en los cuatro algoritmos. MATD3 presentó "
        "el menor incremento: +46.014,75 kgCO₂ (-0,4193 % de efecto favorable); MAAC +140.922,25 "
        "kgCO₂ (-1,3527 %), MASAC +154.873,62 kgCO₂ (-1,7041 %) y HAPPO +2.854.859,87 kgCO₂ "
        "(-31,2181 %). Friedman p=1,461×10⁻⁸, W=0,7716. OE.2 se cumplió; se rechazó HE20 y se "
        "respaldó HE21, precisando que el impacto diferenciado fue predominantemente adverso."
    ),
)
add_paragraph_before(
    document,
    references_anchor,
    (
        "OE.3. Todos los costos totales aumentaron en E3. MAAC fue el mejor desempeño relativo con "
        "+18.978,30 EUR (-0,2675 %), seguido de MASAC +39.477,42 EUR (-0,5555 %), MATD3 +88.556,08 "
        "EUR (-0,9181 %) y HAPPO +213.072,16 EUR (-2,9778 %). Friedman p=0,109906; solo MATD3 "
        "presentó impacto por edificio significativo tras Holm (p=0,026611), pero no hubo diferencia "
        "global ni post hoc. OE.3 se cumplió; HE31 no fue respaldada."
    ),
)
add_paragraph_before(
    document,
    references_anchor,
    (
        "OG. Con pesos iguales, TOPSIS identificó a MAAC (0,936627) como mejor desempeño global, "
        "seguido de MATD3 (0,866227), MASAC (0,840362) y HAPPO (0,000000). No obstante, Friedman "
        "global p=0,085801 no demostró diferencias al 5 %, y la media cruda favoreció a MATD3. El OG "
        "se cumplió al cuantificar e integrar las tres dimensiones y seleccionar un líder bajo una "
        "regla explícita; no se rechazó H0G ni se respaldó H1G, y no se declaró ganador estadístico "
        "global único."
    ),
)
add_heading_before(document, references_anchor, "6.2 Veredicto de hipótesis", 2)
add_table_before(
    document,
    references_anchor,
    ["Hipótesis", "Decisión", "Fundamento cuantitativo"],
    [
        [
            "H0G / H1G",
            "No rechazar H0G; no respaldar H1G",
            "Friedman global χ²=6,6000; p=0,085801. No se cumplen conjuntamente las tres dimensiones.",
        ],
        [
            "HE10 / HE11",
            "No rechazar HE10; no respaldar HE11",
            "Friedman E1 p=0,614935; ningún impacto ni par significativo tras Holm.",
        ],
        [
            "HE20 / HE21",
            "Rechazar HE20; respaldar HE21",
            "Friedman E2 p=1,461×10⁻⁸; W=0,7716; cuatro pares significativos; efecto adverso.",
        ],
        [
            "HE30 / HE31",
            "No respaldar HE31",
            "Friedman E3 p=0,109906; ningún post hoc; MATD3 sí impacta (p-Holm=0,026611).",
        ],
    ],
    "Tabla 6.1. Decisión final de las hipótesis",
    (
        "Nota. Las decisiones son exploratorias por semilla única. La nula no se «acepta»; se informa "
        "si se rechaza o no con la evidencia disponible."
    ),
)
add_heading_before(document, references_anchor, "6.3 Limitaciones y alcance de la evidencia", 2)
add_paragraph_before(
    document,
    references_anchor,
    (
        "La corrida canónica contiene 50 episodios por tratamiento, pero solo una semilla independiente "
        "(seed=0). Los 17 edificios y los tres KPI de flexibilidad permiten contrastes pareados "
        "intra-corrida, no estiman variabilidad entre reentrenamientos. Para una afirmación confirmatoria "
        "se requieren, como mínimo, semillas comunes adicionales, política bloqueada, idéntico presupuesto "
        "de interacción e intervalos de confianza entre semillas."
    ),
)
add_paragraph_before(
    document,
    references_anchor,
    (
        "Se detectaron 24 inconsistencias en los delta guardados de carbono y costo. Las conclusiones "
        "usan control−baseline recalculado y verificable. HAPPO aporta KPI finales comparables, pero su "
        "serie episódica está distribuida entre 49 filas históricas y el episodio final post-resume; "
        "la unión por índice no añade ni estima observaciones."
    ),
)
add_heading_before(document, references_anchor, "6.4 Recomendaciones", 2)
add_paragraph_before(
    document,
    references_anchor,
    (
        "Se recomienda ejecutar al menos 12 semillas por combinación algoritmo–escenario (144 corridas), "
        "mantener 2023 para entrenamiento, 2024 para validación y 2025 para evaluación bloqueada, y "
        "repetir Friedman/post hoc Holm sobre semillas independientes. El ranking final debe conservar "
        "TOPSIS con pesos iguales, frente de Pareto y sensibilidad de pesos; si las reglas no producen "
        "separación estadística y práctica, deberán declararse co-líderes o ausencia de ganador único."
    ),
)


# Replace Annex A with a compact quantitative, reproducible annex.
annex_a_start = find_paragraph(document, "Anexo A.")
annex_b_anchor = find_paragraph(document, "Anexo B.")
remove_range(annex_a_start, annex_b_anchor)

add_heading_before(
    document,
    annex_b_anchor,
    "Anexo A. Auditoría cuantitativa de los resultados reales de 50 episodios",
    1,
)
add_paragraph_before(
    document,
    annex_b_anchor,
    (
        f"Fuente única: {DRIVE_URL}, corrida {RUN_ID}. La auditoría contabiliza 12 tratamientos, "
        "600 episodios, 668 KPI distritales, 15.300 KPI por edificio y 204 resúmenes de comportamiento. "
        "No se inventaron valores ni se imputaron episodios."
    ),
)
add_heading_before(document, annex_b_anchor, "A.1 Matriz de magnitudes primarias", 2)
annex_primary_rows = []
for axis in ["OE1", "OE2", "OE3"]:
    for algorithm in ["HAPPO", "MAAC", "MASAC", "MATD3"]:
        row = primary_index[(axis, algorithm)]
        annex_primary_rows.append(
            [
                axis,
                row["scenario"],
                algorithm,
                row["primary_metric"],
                fmt(row["baseline"], 4),
                fmt(row["control"], 4),
                fmt(row["control_minus_baseline"], 4),
                fmt(row["favorable_effect_percent"], 4) + " %",
            ]
        )
add_table_before(
    document,
    annex_b_anchor,
    ["Eje", "Esc.", "Algoritmo", "Métrica", "Baseline", "Control", "Δ", "Efecto"],
    annex_primary_rows,
    "Tabla A.1. Valores primarios completos por objetivo y algoritmo",
    "Nota. En OE.1 el baseline es 1 y el control es el compuesto de los tres ratios de flexibilidad.",
)
add_heading_before(document, annex_b_anchor, "A.2 Cobertura de matrices completas", 2)
add_table_before(
    document,
    annex_b_anchor,
    ["Matriz", "Filas/valores", "Contenido"],
    [
        ["episode_metrics_50.csv", "600 filas", "Episodios 0–49 de 12 tratamientos"],
        ["district_all_kpis_long.csv", "668 valores", "KPI distritales disponibles"],
        ["objective_all_kpis_numeric.csv", "648 valores", "KPI clasificados OE1–OE3"],
        ["objective_target_scenario_kpis_numeric.csv", "216 valores", "E1/OE1, E2/OE2 y E3/OE3"],
        ["objective_comparable_kpi_gains.csv", "104 valores", "Ganancias con baseline comparable"],
        ["building_kpis_raw.csv", "15.300 valores", "75 KPI × 17 edificios × 12 tratamientos"],
        ["building_behavior_raw.csv", "204 filas", "17 edificios × 12 tratamientos"],
    ],
    "Tabla A.2. Inventario numérico completo generado",
    (
        "Nota. Las matrices se encuentran en outputs/madrl_v3_20260627_164047/"
        "resumen_comparativo/estadistica/analisis_cuantitativo_completo_50_episodios."
    ),
)
add_heading_before(document, annex_b_anchor, "A.3 Archivos de reproducibilidad estadística", 2)
for name in [
    "complete_quantitative_audit.json",
    "complete_quantitative_report.md",
    "run_coverage_quantitative.csv",
    "episode_descriptive_50.csv",
    "primary_objective_values.csv",
    "objective_effect_descriptive.csv",
    "objective_impact_wilcoxon_holm.csv",
    "objective_friedman_posthoc_holm.csv",
    "global_ranking_topsis.csv",
    "hypothesis_decisions_quantitative.csv",
    "stored_delta_consistency_audit.csv",
]:
    add_paragraph_before(
        document,
        annex_b_anchor,
        str((ANALYSIS / name).relative_to(ROOT)).replace("/", "\\"),
    )
add_paragraph_before(
    document,
    annex_b_anchor,
    (
        "El script de cálculo es tools\\run_complete_drive_kpi_objective_analysis.py y el script de "
        "integración documental es tools\\update_word_quantitative_50episodes.py."
    ),
)
add_heading_before(
    document,
    annex_b_anchor,
    "A.4 Evidencia gráfica conservada para trazabilidad",
    2,
)
add_paragraph_before(
    document,
    annex_b_anchor,
    (
        "Las figuras siguientes proceden del documento maestro y se conservan para no perder la "
        "trazabilidad visual de la ejecución, los equipos y los resultados. Tienen carácter "
        "descriptivo o histórico. Cuando una figura muestre la cobertura previa de HAPPO, campos "
        "delta almacenados o pruebas inferenciales anteriores, prevalecen los valores recalculados "
        "de las Tablas 5.1–5.11 y A.1–A.2."
    ),
)
for figure_number, figure_block in enumerate(preserved_figure_blocks, start=1):
    caption = add_paragraph_before(
        document,
        annex_b_anchor,
        f"Figura A.{figure_number}\n{figure_block['title']}",
        style="FiguraIndice",
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    annex_b_anchor._p.addprevious(deepcopy(figure_block["xml"]))
    add_paragraph_before(
        document,
        annex_b_anchor,
        (
            "Nota. Evidencia gráfica conservada del documento maestro para trazabilidad. No sustituye "
            "los valores cuantitativos recalculados ni la decisión estadística final. Fuente: "
            "artefactos originales de la corrida canónica."
        ),
    )


# Correct remaining captions outside the replaced chapters/annex.
for paragraph in document.paragraphs:
    text = paragraph.text
    if "597 filas episodicas" in text:
        paragraph.text = text.replace("597 filas episodicas", "600 filas episódicas")
    elif "597 filas" in text:
        paragraph.text = text.replace("597 filas", "600 filas")
    if "12 tratamientos, 597" in paragraph.text:
        paragraph.text = paragraph.text.replace("12 tratamientos, 597", "12 tratamientos, 600")
    if "105051 filas de series finales" in paragraph.text:
        paragraph.text = paragraph.text.replace(
            "105051 filas de series finales", "3.967.944 filas de timeseries"
        )
    if "351 checkpoints" in paragraph.text:
        paragraph.text = paragraph.text.replace("351 checkpoints", "294 archivos de checkpoint listados")

replace_paragraph(
    document,
    "Interpretación de la figura. El diagrama representa el pipeline de evaluación y selección",
    (
        "Interpretación de la figura. El diagrama representa el pipeline de evaluación y selección "
        "por objetivo. La decisión integra estadística descriptiva, Friedman, Kendall W, Wilcoxon "
        "pareado con corrección de Holm, KPI CityLearn v2 y resultados por edificio. MATD3 presenta "
        "el menor deterioro relativo en OE.1 y OE.2; MAAC, el menor deterioro en OE.3 y el primer "
        "lugar TOPSIS global. La prueba global no demuestra un ganador estadístico único."
    ),
)

table_method_replacements = {
    (
        "Estadística no paramétrica"
    ): "Estadística no paramétrica pareada",
    (
        "p-valores, epsilon2 y decisiones HE."
    ): "p-valores, Kendall W, Holm y decisiones HE.",
    (
        "Shapiro-Wilk, Kruskal-Wallis, Mann-Whitney-Holm y epsilon2."
    ): "Friedman, Kendall W y Wilcoxon pareado con corrección de Holm.",
    (
        "50 episodios por tratamiento disponible; HAPPO conserva 49 filas episódicas materializadas por escenario en el CSV final y se interpreta con esa limitación."
    ): (
        "50 episodios por tratamiento; HAPPO integra 49 episodios históricos y el episodio 49 "
        "post-resume, sin imputación."
    ),
    (
        "Shapiro-Wilk, Kruskal-Wallis y Mann-Whitney U con ajuste posterior."
    ): "Friedman, Kendall W y Wilcoxon pareado con corrección de Holm.",
}
table_method_replacement_count = 0
for old, new in table_method_replacements.items():
    table_method_replacement_count += replace_table_cell_text(document, old, new)


# Normalize document body typography without changing headings/captions semantically.
for paragraph in document.paragraphs:
    if paragraph.style.name.startswith("Heading") or paragraph.style.name in {
        "TablaIndice",
        "FiguraIndice",
    }:
        paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        if not run.font.name:
            run.font.name = "Arial"
        if not run.font.size and paragraph.style.name == "Normal":
            run.font.size = Pt(10)

for table in document.tables:
    if table.rows:
        repeat_table_header(table.rows[0])
    for row in table.rows:
        prevent_row_split(row)

request_field_update_on_open(document)
numbering_changes = normalize_explicit_numbering(document)
approved_after = exact_approved_texts(document)
approved_hash_after = approved_hash(approved_after)
if approved_before != approved_after or approved_hash_before != approved_hash_after:
    raise RuntimeError(
        "La protección falló: cambió al menos una formulación aprobada de problema, objetivo o hipótesis."
    )

document.core_properties.modified = datetime.now()
document.save(OUTPUT_DOCX)

verification_document = Document(OUTPUT_DOCX)
approved_reopened = exact_approved_texts(verification_document)
approved_hash_reopened = approved_hash(approved_reopened)
if approved_reopened != approved_before:
    raise RuntimeError("La reapertura del Word no conservó literalmente los textos aprobados.")

audit = {
    "source": str(SOURCE_DOCX),
    "output": str(OUTPUT_DOCX),
    "backup": str(backup),
    "approved_statement_count": len(approved_before),
    "approved_sha256_before": approved_hash_before,
    "approved_sha256_after": approved_hash_after,
    "approved_sha256_reopened": approved_hash_reopened,
    "approved_texts_unchanged": approved_hash_before == approved_hash_reopened,
    "explicit_numbering_changes": numbering_changes,
    "method_table_cell_replacements": table_method_replacement_count,
    "preserved_figure_blocks": len(preserved_figure_blocks),
    "paragraphs": len(verification_document.paragraphs),
    "tables": len(verification_document.tables),
    "source_master_overwritten": False,
    "quantitative_source": str(ANALYSIS),
    "episode_coverage": "4 algoritmos × 3 escenarios × 50 episodios = 600 episodios",
}
AUDIT_JSON.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")

print(f"UPDATED={OUTPUT_DOCX}")
print(f"BACKUP={backup}")
print(f"AUDIT={AUDIT_JSON}")
print(f"APPROVED_SHA256={approved_hash_reopened}")
print(f"PARAGRAPHS={len(verification_document.paragraphs)}")
print(f"TABLES={len(verification_document.tables)}")
