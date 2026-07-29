#!/usr/bin/env python3
"""Cierre: regenerar multicriterio 100% Drive 50 ep y reemplazar PNG en Word.

1) Regenera outputs/madrl_multicriteria_selection/ con --real-only --plots
2) Sustituye en los 2 Word canónicos cualquier media cuyo SHA coincida con
   las figuras MC previas (o con las nuevas si ya estaban) por los PNG nuevos.
3) Actualiza notas hybrid/ilustrativo → real Drive 50 ep.
4) Aplica anclas evaluate_v2 0,8805 / 0,8679.

Uso:
  py -3.11 scripts/regenerate_multicriteria_drive_50ep_closure.py
"""
from __future__ import annotations

import hashlib
import json
import re
import shutil
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
if str(REPO) not in sys.path:
    sys.path.insert(0, str(REPO))
THESIS = REPO / "tools" / "thesis"
if str(THESIS) not in sys.path:
    sys.path.insert(0, str(THESIS))

from docx import Document  # noqa: E402
from thesis_word_canons import CANONS, RUN_ID  # noqa: E402

OUT_DIR = REPO / "outputs" / "madrl_multicriteria_selection"
FIG_DIR = OUT_DIR / "figures"
BACKUP_DIR = REPO / "outputs" / "_word_backups"
REPORT = REPO / "docs" / "CIERRE_MULTICRITERIA_REAL_DRIVE_50EP_2026-07-29.json"
STAMP = datetime.now().strftime("%Y%m%d_%H%M%S")

MC_FIGS = (
    "pareto_cost_co2_flex.png",
    "learning_curves.png",
    "degradation_bars.png",
)

NOTE_REPLACEMENTS = [
    (
        "hybrid_real_c1c3_plus_illustrative",
        "real_drive_50ep_c1c6",
    ),
    (
        "Complemento visual de la Tabla 5.4.1 TOPSIS. Pareto, curvas de aprendizaje y "
        "gap early–late (C6) regenerados 100% desde Drive 50 ep "
        f"({RUN_ID}: district_objectives + episode_summary); sin relleno ilustrativo.",
        "Complemento visual de la Tabla 5.4.1 TOPSIS. Pareto, curvas de aprendizaje y "
        "gap early–late (C6) regenerados 100% desde Drive 50 ep "
        f"({RUN_ID}: district_objectives + episode_summary); "
        "exclusivamente con evidencia Drive 50 ep.",
    ),
    (
        "Complemento visual de la Tabla 5.4.1 TOPSIS. Pareto, curvas de aprendizaje y "
        "barras de degradación desde outputs/madrl_multicriteria_selection/figures/.",
        "Complemento visual de la Tabla 5.4.1 TOPSIS. Pareto, curvas de aprendizaje y "
        "gap early–late (C6) regenerados 100% desde Drive 50 ep "
        f"({RUN_ID}: district_objectives + episode_summary); "
        "exclusivamente con evidencia Drive 50 ep.",
    ),
    (
        "Nota. Figuras descriptivas de outputs/madrl_multicriteria_selection/figures/; "
        "no constituyen evidencia de HE10–HE31 (ver §5.3 y §5.5).",
        "Nota. Figuras descriptivas 100% reales (Drive 50 ep, "
        f"{RUN_ID}); no constituyen evidencia de HE10–HE31 (ver §5.3 y §5.5).",
    ),
]

EVAL_V2_REPLACEMENTS = [
    (
        "ranking evaluate_v2 4/4 lidera MAAC 0,9538).",
        "ranking evaluate_v2 4/4: MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000).",
    ),
    (
        "ranking evaluate_v2 4/4 lidera MAAC 0,9538 con HAPPO en score 0,0000)",
        "ranking evaluate_v2 4/4: MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000)",
    ),
    (
        "y el ranking evaluate_v2 4/4 (0,9538).",
        "y el ranking evaluate_v2 4/4 (MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000).",
    ),
    (
        "El ranking evaluate_v2 4/4 sitúa a MAAC primero (0,9538).",
        "El ranking evaluate_v2 4/4 sitúa a MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000.",
    ),
]


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_file(path: Path) -> str:
    return sha256_bytes(path.read_bytes())


def regenerate_artefacts() -> dict:
    from uc3m.multicriteria.pipeline import run_selection_pipeline

    # Snapshot old figure hashes before overwrite (for Word media match).
    old_hashes = {}
    for name in MC_FIGS:
        p = FIG_DIR / name
        if p.is_file():
            old_hashes[name] = sha256_file(p)

    result = run_selection_pipeline(
        repo=REPO,
        run_dir=REPO / "outputs" / RUN_ID,
        scenario="E1",
        prefer_real=True,
        allow_illustrative_fill=False,
        output_dir=OUT_DIR,
        make_plots=True,
        sensitivity_samples=48,
    )
    new_hashes = {name: sha256_file(FIG_DIR / name) for name in MC_FIGS}
    for name in MC_FIGS:
        if not (FIG_DIR / name).is_file():
            raise FileNotFoundError(FIG_DIR / name)
    if result["source"] != "real_drive_50ep_c1c6":
        raise RuntimeError(f"unexpected source={result['source']}")
    for algo, prov in result["provenance"].items():
        for cid, tag in prov.items():
            if "illustrative" in str(tag).lower() or "synthetic" in str(tag).lower():
                raise RuntimeError(f"illustrative leak {algo}/{cid}={tag}")
    return {
        "source": result["source"],
        "algorithms": list(result["decision_matrix"].keys()),
        "topsis_winner": result["ranking_consistency"]["topsis_winner"],
        "old_figure_sha256": old_hashes,
        "new_figure_sha256": new_hashes,
        "provenance_sample": {
            a: result["provenance"][a] for a in list(result["provenance"])[:1]
        },
        "decision_matrix": result["decision_matrix"],
    }


def replace_media_in_docx(docx_path: Path, old_to_new: dict[str, Path]) -> dict:
    """Replace embedded media whose SHA matches keys of old_to_new."""

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{docx_path.stem}_antes_mc_real_drive_{STAMP}{docx_path.suffix}"
    shutil.copy2(docx_path, backup)

    replaced = []
    tmp = docx_path.with_suffix(docx_path.suffix + f".tmp_{STAMP}")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
        tmp, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename.startswith("word/media/"):
                h = sha256_bytes(data)
                if h in old_to_new:
                    new_path = old_to_new[h]
                    data = new_path.read_bytes()
                    replaced.append(
                        {
                            "media": info.filename,
                            "old_sha256": h,
                            "new_file": str(new_path.relative_to(REPO)).replace("\\", "/"),
                            "new_sha256": sha256_bytes(data),
                        }
                    )
            zout.writestr(info, data)
    tmp.replace(docx_path)
    return {
        "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        "replaced": replaced,
        "n_replaced": len(replaced),
    }


def patch_text(docx_path: Path) -> list[dict]:
    doc = Document(str(docx_path))
    changes: list[dict] = []
    pairs = NOTE_REPLACEMENTS + EVAL_V2_REPLACEMENTS
    for old, new in pairs:
        for i, p in enumerate(doc.paragraphs):
            text = p.text or ""
            if old in text:
                p.text = text.replace(old, new)
                changes.append({"paragraph": i, "old": old[:80], "new": new[:80]})
    if changes:
        doc.save(str(docx_path))
    return changes


def verify_no_illustrative(docx_path: Path) -> dict:
    doc = Document(str(docx_path))
    blob = "\n".join(p.text or "" for p in doc.paragraphs)
    return {
        "has_hybrid_illustrative": bool(
            re.search(r"hybrid_real_c1c3_plus_illustrative", blob, flags=re.I)
        ),
        "has_08805": bool(re.search(r"0[,.]8805", blob)),
        "has_08679": bool(re.search(r"0[,.]8679", blob)),
        "has_real_drive_note": "real_drive_50ep" in blob or "100% reales" in blob or "100% desde Drive" in blob,
    }


def main() -> int:
    print("[1/4] Regenerando artefactos multicriterio real-only…")
    regen = regenerate_artefacts()
    print(f"  source={regen['source']} algos={regen['algorithms']} winner={regen['topsis_winner']}")

    # Map both old and (if needed) current on-disk hashes → new files.
    old_to_new: dict[str, Path] = {}
    for name in MC_FIGS:
        new_p = FIG_DIR / name
        old_h = regen["old_figure_sha256"].get(name)
        new_h = regen["new_figure_sha256"][name]
        if old_h:
            old_to_new[old_h] = new_p
        old_to_new[new_h] = new_p  # idempotent if already new

    # Also index any currently embedded hashes that match previous backup figs
    # by scanning Word after building old_to_new from file snapshots.

    print("[2/4] Reemplazando media MC en Word canónicos…")
    report: dict = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "stamp": STAMP,
        "run_id": RUN_ID,
        "regeneration": regen,
        "docs": {},
    }
    for path in CANONS:
        media_info = replace_media_in_docx(path, old_to_new)
        text_changes = patch_text(path)
        checks = verify_no_illustrative(path)
        # If no media replaced, try matching by scanning current media against
        # figure filenames via second pass using ANY previous known hashes from
        # selection_report if present — already covered by old_hashes.
        entry = {
            "path": str(path.relative_to(REPO)).replace("\\", "/"),
            "media": media_info,
            "text_changes": text_changes,
            "checks": checks,
            "ok": media_info["n_replaced"] >= 3
            and checks["has_08805"]
            and checks["has_08679"]
            and not checks["has_hybrid_illustrative"],
        }
        report["docs"][path.name] = entry
        print(
            f"  {path.name}: replaced={media_info['n_replaced']} "
            f"text={len(text_changes)} ok={entry['ok']}"
        )

    # If replacement failed (hashes already different), force-replace by
    # locating MC captions' nearby images is harder; fallback: re-insert via
    # hash of figures that may still be old from catalog copy.
    # Second pass: for docs with n_replaced < 3, search ALL media and replace
    # any that equal ANY file under figures/ before regen — already done.
    # Fallback: replace by size+order of MC captions using python-docx blips.
    for name, entry in report["docs"].items():
        if entry["media"]["n_replaced"] >= 3:
            continue
        print(f"  [fallback] forzando reescritura de media MC en {name}…")
        path = REPO / entry["path"]
        forced = force_replace_mc_images_by_caption(path)
        entry["media"]["fallback"] = forced
        entry["media"]["n_replaced"] = entry["media"]["n_replaced"] + forced.get("n_replaced", 0)
        entry["checks"] = verify_no_illustrative(path)
        entry["ok"] = (
            entry["media"]["n_replaced"] >= 3
            and entry["checks"]["has_08805"]
            and entry["checks"]["has_08679"]
            and not entry["checks"]["has_hybrid_illustrative"]
        )
        print(f"  {name} after fallback: replaced_total={entry['media']['n_replaced']} ok={entry['ok']}")

    report["ok"] = all(d.get("ok") for d in report["docs"].values())
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print("[3/4] Report:", REPORT)
    print("[4/4] VEREDICTO:", "PASS" if report["ok"] else "REVIEW")
    return 0 if report["ok"] else 2


def force_replace_mc_images_by_caption(docx_path: Path) -> dict:
    """Replace image parts referenced near MC / 5.3b-d / MC.1-3 captions."""

    doc = Document(str(docx_path))
    targets = {
        "pareto": FIG_DIR / "pareto_cost_co2_flex.png",
        "learning": FIG_DIR / "learning_curves.png",
        "degradation": FIG_DIR / "degradation_bars.png",
    }
    caption_map = [
        (re.compile(r"Figura (?:5\.3b|MC\.1).*Pareto", re.I), "pareto"),
        (re.compile(r"Figura (?:5\.3c|MC\.2).*aprendizaje", re.I), "learning"),
        (re.compile(r"Figura (?:5\.3d|MC\.3).*degrad", re.I), "degradation"),
    ]
    replaced = []
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        text = p.text or ""
        key = None
        for cre, k in caption_map:
            if cre.search(text):
                key = k
                break
        if key is None:
            continue
        # Image usually in previous non-empty paragraph with drawing.
        img_para = None
        for j in range(i - 1, max(-1, i - 4), -1):
            if paras[j]._element.xpath(".//a:blip"):
                img_para = paras[j]
                break
        if img_para is None:
            continue
        blips = img_para._element.xpath(".//a:blip")
        if not blips:
            continue
        embed = blips[0].get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if not embed:
            continue
        rel = doc.part.rels[embed]
        if "image" not in (rel.reltype or ""):
            continue
        new_bytes = targets[key].read_bytes()
        rel.target_part._blob = new_bytes
        replaced.append({"caption": text[:60], "key": key, "rId": embed})
    if replaced:
        doc.save(str(docx_path))
    return {"n_replaced": len(replaced), "items": replaced}


if __name__ == "__main__":
    raise SystemExit(main())
