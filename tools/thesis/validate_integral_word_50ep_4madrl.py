#!/usr/bin/env python3
"""Validación integral de los 2 Word canónicos vs ground truth 50 ep / 4 MADRL.

Compara cifras y claims cuantitativos de:
  - docs/Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx
  - docs/Inforne_tesisV4_FINAL_REVISADO_50_EPISODIOS.docx
contra artefactos reales del repo (best_madrl, ranking evaluate_v2, CANON MD).

No inventa números. Emite:
  - docs/VALIDACION_INTEGRAL_WORD_50EP_4MADRL_2026-07-29.json
  - docs/VALIDACION_INTEGRAL_WORD_50EP_4MADRL_2026-07-29.md

Uso:
  python tools/thesis/validate_integral_word_50ep_4madrl.py
"""

from __future__ import annotations

import csv
import json
import re
import sys
from collections import Counter
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from docx import Document  # noqa: E402
from thesis_word_canons import CANONS, DRIVE_FOLDER_URL, INFORME, RUN_ID, TESIS  # noqa: E402

DATE_TAG = "2026-07-29"
OUT_JSON = REPO / "docs" / f"VALIDACION_INTEGRAL_WORD_50EP_4MADRL_{DATE_TAG}.json"
OUT_MD = REPO / "docs" / f"VALIDACION_INTEGRAL_WORD_50EP_4MADRL_{DATE_TAG}.md"

BEST_MADRL = (
    REPO
    / "outputs"
    / RUN_ID
    / "resumen_comparativo"
    / "best_madrl_report.json"
)
RANKING_ALL = (
    REPO
    / "outputs"
    / "_drive_madrl"
    / "kpi_recalc_20260728"
    / "tables"
    / "ranking_oe_scores_all_values.csv"
)
CANON_MD = REPO / "docs" / "CANON_WORD_Y_VALIDEZ_50EP_DRIVE_2026-07-29.md"
POINTER = REPO / "outputs" / "latest_colab_output_root.txt"
HE_CSV = (
    REPO
    / "outputs"
    / RUN_ID
    / "resumen_comparativo"
    / "estadistica"
    / "hipotesis_estadisticas_madrl.csv"
)
MULTICRITERIA = REPO / "outputs" / "madrl_multicriteria_selection" / "selection_report.json"

ALGORITHMS = ("HAPPO", "MASAC", "MATD3", "MAAC")
TARGET_EPISODES = 50

# Anchors from CANON + ranking CSV (rounded forms accepted in Spanish/EN prose)
ANCHOR_SPECS: list[dict] = [
    {
        "id": "score_global_3x3_matd3",
        "label": "MATD3 score global 3×3",
        "value": 0.6667,
        "patterns": [r"0[,.]6667"],
        "source": "best_madrl_report.json / ranking canonical3_no_happo",
        "required": True,
    },
    {
        "id": "flex_matd3",
        "label": "flex_composite E1 MATD3",
        "value": 1.0009,
        "patterns": [r"1[,.]0009", r"1[,.]000922"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "flex_maac",
        "label": "flex_composite E1 MAAC",
        "value": 1.0124,
        "patterns": [r"1[,.]0124", r"1[,.]012425"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "flex_masac",
        "label": "flex_composite E1 MASAC",
        "value": 1.0286,
        "patterns": [r"1[,.]0286", r"1[,.]02855"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "flex_happo",
        "label": "flex_composite E1 HAPPO",
        "value": 1.1105,
        "patterns": [r"1[,.]1105", r"1[,.]11054"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "co2_matd3",
        "label": "ΔCO₂ E2 MATD3 kg",
        "value": 23070,
        "patterns": [r"23[\s.]?070", r"23070"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "co2_maac",
        "label": "ΔCO₂ E2 MAAC kg",
        "value": 70654,
        "patterns": [r"70[\s.]?654", r"70654"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "co2_masac",
        "label": "ΔCO₂ E2 MASAC kg",
        "value": 77649,
        "patterns": [r"77[\s.]?649", r"77[\s.]?648", r"77648", r"77649"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "co2_happo",
        "label": "ΔCO₂ E2 HAPPO kg",
        "value": 1431341,
        "patterns": [r"1[\s.]?431[\s.]?341", r"1431341"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "cost_maac",
        "label": "Δcosto E3 MAAC EUR",
        "value": 9515,
        "patterns": [r"9[\s.]?515", r"9515"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "cost_masac",
        "label": "Δcosto E3 MASAC EUR",
        "value": 19793,
        "patterns": [r"19[\s.]?793", r"19[\s.]?792", r"19792", r"19793"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "cost_matd3",
        "label": "Δcosto E3 MATD3 EUR",
        "value": 44399,
        "patterns": [r"44[\s.]?399", r"44399"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "cost_happo",
        "label": "Δcosto E3 HAPPO EUR",
        "value": 106828,
        "patterns": [r"106[\s.]?828", r"106[\s.]?827", r"106828", r"106827"],
        "source": "ranking_oe_scores_all_values.csv",
        "required": True,
    },
    {
        "id": "eval_v2_maac",
        "label": "evaluate_v2 4/4 MAAC score",
        "value": 0.9538,
        "patterns": [r"0[,.]9538"],
        "source": "ranking_oe_scores_all_values.csv all4",
        "required": True,
    },
    {
        "id": "eval_v2_matd3",
        "label": "evaluate_v2 4/4 MATD3 score",
        "value": 0.8805,
        "patterns": [r"0[,.]8805"],
        "source": "ranking_oe_scores_all_values.csv all4",
        "required": True,
    },
    {
        "id": "eval_v2_masac",
        "label": "evaluate_v2 4/4 MASAC score",
        "value": 0.8679,
        "patterns": [r"0[,.]8679"],
        "source": "ranking_oe_scores_all_values.csv all4",
        "required": True,
    },
    {
        "id": "run_id",
        "label": "run_id canónico",
        "value": RUN_ID,
        "patterns": [re.escape(RUN_ID)],
        "source": "best_madrl_report.json",
        "required": True,
    },
]

# Episode counts that contradict the 50-ep canonical campaign if asserted as primary
FORBIDDEN_PRIMARY_EPISODE_CLAIMS = [
    (r"(?i)(?:entrenamiento|corrida|campaña|experimento)\s+(?:de\s+)?10\s+episodios", "10 episodios como campaña primaria"),
    (r"(?i)(?:entrenamiento|corrida|campaña|experimento)\s+(?:de\s+)?20\s+episodios", "20 episodios como campaña primaria"),
    (r"(?i)(?:entrenamiento|corrida|campaña|experimento)\s+(?:de\s+)?5\s+episodios(?!\s+de\s+calentamiento)", "5 episodios como campaña primaria"),
    (r"(?i)target_episodes\s*=\s*(?!50)\d+", "target_episodes ≠ 50"),
]

THESIS_OBJECTIVE_MARKERS = [
    ("OG", r"OG\.\s*-"),
    ("OE.1", r"OE\.1\s*:"),
    ("OE.2", r"OE\.2\s*:"),
    ("OE.3", r"OE\.3\s*:"),
    ("flexibilidad", r"(?i)flexibilidad\s+energ"),
    ("CO2", r"(?i)(?:emisiones?\s+de\s+)?CO[\s₂2]"),
    ("costos", r"(?i)costos?\s+energ"),
    ("cuasiexperimental", r"(?i)cuasi[\s-]?experimental"),
    ("CityLearn", r"(?i)CityLearn"),
    ("Iquitos", r"(?i)Iquitos"),
    ("MADRL", r"\bMADRL\b"),
    ("Dec-POMDP", r"(?i)Dec[\s-]?POMDP"),
    ("H0G/H1G", r"H0G\.|H1G\."),
    ("HE10/HE11", r"HE10\.|HE11\."),
    ("HE20/HE21", r"HE20\.|HE21\."),
    ("HE30/HE31", r"HE30\.|HE31\."),
    ("Shapiro", r"(?i)Shapiro"),
    ("Kruskal", r"(?i)Kruskal"),
]


@dataclass
class CheckResult:
    id: str
    category: str
    status: str  # PASS | FAIL | WARN | GAP
    message: str
    word: str | None = None
    expected: str | None = None
    found: str | None = None
    source: str | None = None


@dataclass
class WordExtract:
    name: str
    path: str
    chars: int
    paragraphs: int
    tables: int
    text: str
    algo_counts: dict[str, int] = field(default_factory=dict)
    episode_50_mentions: int = 0
    episode_other_mentions: dict[str, int] = field(default_factory=dict)


def load_ground_truth() -> dict:
    gt: dict = {
        "run_id": RUN_ID,
        "drive_folder": DRIVE_FOLDER_URL,
        "target_episodes": TARGET_EPISODES,
        "algorithms": list(ALGORITHMS),
        "sources": {},
        "anchors": {},
        "gaps": [],
    }

    if POINTER.is_file():
        ptr = POINTER.read_text(encoding="utf-8").strip().replace("\\", "/")
        gt["sources"]["pointer"] = str(POINTER.relative_to(REPO)).replace("\\", "/")
        gt["pointer_value"] = ptr
        if RUN_ID not in ptr:
            gt["gaps"].append(f"Puntero latest no apunta a {RUN_ID}: {ptr}")
    else:
        gt["gaps"].append(f"Falta puntero {POINTER}")

    if not BEST_MADRL.is_file():
        gt["gaps"].append(f"Falta {BEST_MADRL}")
        return gt

    best = json.loads(BEST_MADRL.read_text(encoding="utf-8"))
    gt["sources"]["best_madrl"] = str(BEST_MADRL.relative_to(REPO)).replace("\\", "/")
    gt["best_madrl"] = {
        "mejor_madrl": best.get("mejor_madrl"),
        "target_episodes": best.get("target_episodes"),
        "algorithms_in_study": best.get("algorithms_in_study"),
        "nota_episodios": best.get("nota_episodios"),
        "ranking_with_kpis": best.get("ranking_with_kpis"),
        "kpis_primarios": best.get("kpis_primarios"),
    }
    if best.get("target_episodes") != 50:
        gt["gaps"].append(f"best_madrl target_episodes={best.get('target_episodes')} ≠ 50")

    ranking_rows: list[dict] = []
    if RANKING_ALL.is_file():
        with RANKING_ALL.open("r", encoding="utf-8-sig", newline="") as fh:
            ranking_rows = list(csv.DictReader(fh))
        gt["sources"]["ranking_all"] = str(RANKING_ALL.relative_to(REPO)).replace("\\", "/")
        gt["ranking_all4"] = [
            r for r in ranking_rows if r.get("ranking_set") == "all4_including_happo"
        ]
        gt["ranking_3x3"] = [
            r for r in ranking_rows if r.get("ranking_set") == "canonical3_no_happo"
        ]
    else:
        gt["gaps"].append(f"Falta {RANKING_ALL}")

    # Consolidate rounded anchors used in prose
    def round_flex(x: float) -> float:
        return round(x, 4)

    all4 = {r["algorithm"]: r for r in gt.get("ranking_all4", [])}
    gt["anchors"] = {
        "score_global_3x3_matd3": 0.6667,
        "mejor_madrl_3x3": best.get("mejor_madrl"),
        "flex_composite_e1": {
            algo: round_flex(float(all4[algo]["flex_composite_e1"]))
            for algo in ALGORITHMS
            if algo in all4
        },
        "co2_delta_kg_e2": {
            algo: round(float(all4[algo]["co2_delta_kg_e2"]))
            for algo in ALGORITHMS
            if algo in all4
        },
        "cost_delta_eur_e3": {
            algo: round(float(all4[algo]["cost_delta_eur_e3"]))
            for algo in ALGORITHMS
            if algo in all4
        },
        "eval_v2_scores": {
            algo: round(float(all4[algo]["score_global"]), 4)
            for algo in ALGORITHMS
            if algo in all4
        },
        "episodes": {
            "MATD3": "50/50",
            "MAAC": "50/50",
            "MASAC": "50/50",
            "HAPPO": "49/50",
        },
    }

    if HE_CSV.is_file():
        gt["sources"]["he_csv"] = str(HE_CSV.relative_to(REPO)).replace("\\", "/")
    else:
        gt["gaps"].append(f"Falta HE CSV {HE_CSV}")

    if MULTICRITERIA.is_file():
        mc = json.loads(MULTICRITERIA.read_text(encoding="utf-8"))
        gt["sources"]["multicriteria"] = str(MULTICRITERIA.relative_to(REPO)).replace("\\", "/")
        gt["multicriteria_source"] = mc.get("source")
        gt["sources"]["multicriteria_source"] = mc.get("source")
        gt["multicriteria_note"] = {
            "source_field": mc.get("source"),
            "role": (
                "TOPSIS = medida multicriterio formal adicional para OG/OE/HE "
                "(complementa evaluate_v2 y KPI-gains; no los sustituye)."
            ),
            "warning": (
                "Fuente multicriterio illustrative: no presentar TOPSIS como única evidencia OG "
                "ni inventar scores; citar valores del JSON."
                if "illustrative" in str(mc.get("source", "")).lower()
                else (
                    "Capa multicriterio real_drive_50ep_c1c6: medida formal adicional en Caps. V–VII "
                    "para objetivos/hipótesis; complementa best_madrl 3×3 y evaluate_v2 4/4."
                )
            ),
        }

    gt["sources"]["canon_md"] = str(CANON_MD.relative_to(REPO)).replace("\\", "/")
    gt["objectives"] = {
        "OG": "Determinar impacto de MADRLs en flexibilidad, CO₂ y costos en Iquitos e identificar mejor algoritmo global",
        "OE.1": "Flexibilidad energética / escenario E1",
        "OE.2": "Emisiones CO₂ / escenario E2",
        "OE.3": "Costos energéticos / escenario E3",
        "source": "agent-skills/madrl-citylearn-thesis-integrated/references/module-b-thesis-report.md",
    }
    return gt


def extract_word(path: Path) -> WordExtract:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    n_tables = 0
    for table in doc.tables:
        n_tables += 1
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(cells)
            if line.strip(" |"):
                parts.append(line)
    text = "\n".join(parts)
    algo_counts = {a: len(re.findall(rf"\b{a}\b", text)) for a in ALGORITHMS}
    ep50 = len(re.findall(r"(?i)50\s+episodios?", text))
    other = {
        "5": len(re.findall(r"(?i)(?<!1)(?<!2)(?<!3)(?<!4)(?<!5)\b5\s+episodios?", text)),
        "10": len(re.findall(r"(?i)\b10\s+episodios?", text)),
        "20": len(re.findall(r"(?i)\b20\s+episodios?", text)),
        "49": len(re.findall(r"(?i)\b49\s+episodios?", text)),
        "100": len(re.findall(r"(?i)\b100\s+episodios?", text)),
    }
    return WordExtract(
        name=path.name,
        path=str(path.relative_to(REPO)).replace("\\", "/"),
        chars=len(text),
        paragraphs=len(doc.paragraphs),
        tables=n_tables,
        text=text,
        algo_counts=algo_counts,
        episode_50_mentions=ep50,
        episode_other_mentions=other,
    )


def find_context(text: str, pattern: str, window: int = 80) -> str | None:
    m = re.search(pattern, text)
    if not m:
        return None
    a = max(0, m.start() - window)
    b = min(len(text), m.end() + window)
    return text[a:b].replace("\n", " ")


def run_checks(gt: dict, extracts: dict[str, WordExtract]) -> list[CheckResult]:
    checks: list[CheckResult] = []

    # Canon inventory
    docs_root = REPO / "docs"
    docx_root = sorted(p.name for p in docs_root.glob("*.docx"))
    if docx_root == [INFORME.name, TESIS.name] or set(docx_root) == {
        INFORME.name,
        TESIS.name,
    }:
        checks.append(
            CheckResult(
                id="canon_exactly_2",
                category="inventory",
                status="PASS",
                message=f"Exactamente 2 Word en docs/: {docx_root}",
            )
        )
    else:
        checks.append(
            CheckResult(
                id="canon_exactly_2",
                category="inventory",
                status="FAIL",
                message=f"docs/ tiene {len(docx_root)} .docx (esperado 2 canónicos): {docx_root}",
                found=str(docx_root),
                expected="Tesis + Informe",
            )
        )

    for path in CANONS:
        if path.is_file() and path.stat().st_size > 0:
            checks.append(
                CheckResult(
                    id=f"exists_{path.stem[:24]}",
                    category="inventory",
                    status="PASS",
                    message=f"Existe {path.name} ({path.stat().st_size} bytes)",
                    word=path.name,
                )
            )
        else:
            checks.append(
                CheckResult(
                    id=f"exists_{path.stem[:24]}",
                    category="inventory",
                    status="FAIL",
                    message=f"Falta Word canónico: {path}",
                    word=path.name,
                )
            )

    if gt.get("best_madrl", {}).get("target_episodes") == 50:
        checks.append(
            CheckResult(
                id="gt_target_50",
                category="ground_truth",
                status="PASS",
                message="best_madrl.target_episodes=50",
                source=gt["sources"].get("best_madrl"),
            )
        )
    else:
        checks.append(
            CheckResult(
                id="gt_target_50",
                category="ground_truth",
                status="FAIL",
                message="Ground truth no confirma 50 episodios",
                found=str(gt.get("best_madrl", {}).get("target_episodes")),
                expected="50",
            )
        )

    algos = set(gt.get("best_madrl", {}).get("algorithms_in_study") or [])
    if algos == set(ALGORITHMS):
        checks.append(
            CheckResult(
                id="gt_4_madrl",
                category="ground_truth",
                status="PASS",
                message=f"4 MADRL en estudio: {sorted(algos)}",
                source=gt["sources"].get("best_madrl"),
            )
        )
    else:
        checks.append(
            CheckResult(
                id="gt_4_madrl",
                category="ground_truth",
                status="FAIL",
                message="algorithms_in_study incompleto",
                found=str(sorted(algos)),
                expected=str(list(ALGORITHMS)),
            )
        )

    for name, ex in extracts.items():
        # 50 episodes presence
        if ex.episode_50_mentions >= 3:
            checks.append(
                CheckResult(
                    id=f"ep50_{name}",
                    category="episodes",
                    status="PASS",
                    message=f"{ex.episode_50_mentions} menciones de '50 episodios'",
                    word=ex.name,
                    expected="≥3 menciones de 50 episodios",
                    found=str(ex.episode_50_mentions),
                )
            )
        else:
            checks.append(
                CheckResult(
                    id=f"ep50_{name}",
                    category="episodes",
                    status="FAIL",
                    message=f"Pocas menciones de 50 episodios ({ex.episode_50_mentions})",
                    word=ex.name,
                    expected="≥3",
                    found=str(ex.episode_50_mentions),
                )
            )

        # Forbidden primary campaign claims
        for pat, label in FORBIDDEN_PRIMARY_EPISODE_CLAIMS:
            ctx = find_context(ex.text, pat)
            if ctx:
                checks.append(
                    CheckResult(
                        id=f"forbidden_ep_{name}_{label[:20]}",
                        category="episodes",
                        status="FAIL",
                        message=f"Claim de campaña no canónica: {label}",
                        word=ex.name,
                        found=ctx,
                        expected="campaña primaria = 50 episodios",
                    )
                )

        # HAPPO 49 note (49/50 or "49 episodios") — trazable a best_madrl
        happo_49 = bool(
            re.search(r"(?i)HAPPO.{0,60}49\s*/\s*50", ex.text)
            or re.search(r"(?i)49\s*/\s*50.{0,40}HAPPO", ex.text)
            or ex.episode_other_mentions.get("49", 0) > 0
        )
        if happo_49:
            checks.append(
                CheckResult(
                    id=f"happo_49_{name}",
                    category="episodes",
                    status="PASS",
                    message="Documenta HAPPO 49/50 (trazable a best_madrl)",
                    word=ex.name,
                    source=gt["sources"].get("best_madrl"),
                    found=find_context(ex.text, r"(?i)HAPPO.{0,40}49") or "49/50",
                )
            )
        else:
            checks.append(
                CheckResult(
                    id=f"happo_49_{name}",
                    category="episodes",
                    status="WARN",
                    message="No menciona explícitamente HAPPO 49/50",
                    word=ex.name,
                    expected="nota HAPPO 49/50",
                )
            )

        # 4 algorithms present with meaningful counts
        missing = [a for a in ALGORITHMS if ex.algo_counts.get(a, 0) < 5]
        if not missing:
            checks.append(
                CheckResult(
                    id=f"algos_{name}",
                    category="algorithms",
                    status="PASS",
                    message=f"4 MADRL presentes: {ex.algo_counts}",
                    word=ex.name,
                )
            )
        else:
            checks.append(
                CheckResult(
                    id=f"algos_{name}",
                    category="algorithms",
                    status="FAIL",
                    message=f"Algoritmos con presencia insuficiente: {missing}",
                    word=ex.name,
                    found=str(ex.algo_counts),
                    expected="≥5 menciones por algoritmo",
                )
            )

        # Numeric anchors
        for spec in ANCHOR_SPECS:
            hit = any(re.search(p, ex.text) for p in spec["patterns"])
            if hit:
                checks.append(
                    CheckResult(
                        id=f"anchor_{spec['id']}_{name}",
                        category="anchors",
                        status="PASS",
                        message=f"Ancla presente: {spec['label']}={spec['value']}",
                        word=ex.name,
                        expected=str(spec["value"]),
                        source=spec["source"],
                        found=find_context(ex.text, spec["patterns"][0]) or "match",
                    )
                )
            elif spec["required"]:
                checks.append(
                    CheckResult(
                        id=f"anchor_{spec['id']}_{name}",
                        category="anchors",
                        status="FAIL",
                        message=f"Ancla ausente o no trazable: {spec['label']}",
                        word=ex.name,
                        expected=str(spec["value"]),
                        source=spec["source"],
                    )
                )

        # Thesis objective / framing markers
        missing_obj = []
        for oid, pat in THESIS_OBJECTIVE_MARKERS:
            if not re.search(pat, ex.text):
                missing_obj.append(oid)
        if not missing_obj:
            checks.append(
                CheckResult(
                    id=f"objectives_{name}",
                    category="redaccion",
                    status="PASS",
                    message="Marcadores OG/OE/HE/metodología/caso presentes",
                    word=ex.name,
                )
            )
        else:
            # Cap. objectives may be abbreviated in Informe — WARN if few, FAIL if many critical
            critical = {"OG", "OE.1", "OE.2", "OE.3", "MADRL", "Iquitos", "flexibilidad", "CO2", "costos"}
            crit_miss = [m for m in missing_obj if m in critical]
            status = "FAIL" if len(crit_miss) >= 3 else "WARN"
            checks.append(
                CheckResult(
                    id=f"objectives_{name}",
                    category="redaccion",
                    status=status,
                    message=f"Marcadores de objetivo/tesis ausentes: {missing_obj}",
                    word=ex.name,
                    found=str(missing_obj),
                )
            )

        # Invented / stale campaign markers
        stale_patterns = [
            (r"(?i)\[Pendiente:\s*corrida can[oó]nica\s*50", "marcador pendiente 50 ep"),
            (r"(?i)\[REEMPLAZAR\]", "marcador REEMPLAZAR"),
            (r"(?i)resultado no verificado", "resultado no verificado"),
            (r"(?i)dato inventado", "dato inventado"),
            (r"citylearn_v3_madrl_full_20260615_074011_v4", "run legacy v4 5ep como canónico"),
            (r"(?i)hybrid_real_c1c3_plus_illustrative", "fuente multicriterio illustrative"),
        ]
        for pat, label in stale_patterns:
            ctx = find_context(ex.text, pat)
            if ctx:
                checks.append(
                    CheckResult(
                        id=f"stale_{name}_{label[:24]}",
                        category="trazabilidad",
                        status="FAIL",
                        message=f"Señal de contenido no final / no canónico: {label}",
                        word=ex.name,
                        found=ctx,
                    )
                )

        # TOPSIS as formal additional measure for OG/OE/HE (Caps. V–VII)
        has_topsis = bool(re.search(r"(?i)TOPSIS", ex.text))
        deny_topsis_role = bool(
            re.search(
                r"(?i)TOPSIS[^\n]{0,140}("
                r"solo\s+ilustrativ|no\s+decide\s+(hip[oó]tesis|HE)|"
                r"no\s+evidencia\s+de\s+HE|es\s+descriptiv[oa]\s+y\s+no\s+"
                r"|descriptivo\s*\(madrl_multicriteria|"
                r"not\s+hypothesis\s+evidence|does\s+not\s+decide\s+hypothes"
                r")",
                ex.text,
            )
        )
        formal_topsis = bool(
            re.search(
                r"(?i)TOPSIS[^\n]{0,160}("
                r"medida\s+multicriterio\s+formal\s+adicional|"
                r"additional\s+formal\s+multicriteria\s+measure|"
                r"refuerza[^\n]{0,40}(OG|OE|HE|objetivo|hip[oó]tes)"
                r")",
                ex.text,
            )
        )
        topsis_obj_he = bool(
            re.search(
                r"(?i)TOPSIS[^\n]{0,160}(OG|OE\.[123]|HE\d*|hip[oó]tesis|objetivo)",
                ex.text,
            )
        )
        stale_hybrid = bool(
            re.search(r"(?i)hybrid_real_c1c3_plus_illustrative|illustrative_methodology", ex.text)
        )

        if stale_hybrid and has_topsis:
            checks.append(
                CheckResult(
                    id=f"layer_topsis_hybrid_{name}",
                    category="trazabilidad",
                    status="FAIL",
                    message="TOPSIS citado junto a hybrid/illustrative — no usar como fuente canónica",
                    word=ex.name,
                    source=gt.get("sources", {}).get("multicriteria"),
                )
            )
        if has_topsis and deny_topsis_role:
            checks.append(
                CheckResult(
                    id=f"layer_topsis_deny_{name}",
                    category="redaccion",
                    status="FAIL",
                    message=(
                        "Disclaimer que niega el rol formal de TOPSIS "
                        "(ilustrativo/no decide HE) — reescribir como medida adicional OG/OE/HE"
                    ),
                    word=ex.name,
                    source=gt.get("sources", {}).get("multicriteria"),
                )
            )
        elif has_topsis and formal_topsis and topsis_obj_he:
            checks.append(
                CheckResult(
                    id=f"layer_topsis_formal_{name}",
                    category="redaccion",
                    status="PASS",
                    message=(
                        "TOPSIS tratado como medida multicriterio formal adicional "
                        "para objetivos/hipótesis (Caps. V–VII)"
                    ),
                    word=ex.name,
                    source=gt.get("sources", {}).get("multicriteria"),
                )
            )
        elif has_topsis:
            checks.append(
                CheckResult(
                    id=f"layer_topsis_role_{name}",
                    category="redaccion",
                    status="WARN",
                    message=(
                        "TOPSIS presente pero sin ancla clara de medida formal para OG/OE/HE"
                    ),
                    word=ex.name,
                    source=gt.get("sources", {}).get("multicriteria"),
                )
            )

        if has_topsis and (
            "real_drive_50ep" in (gt.get("multicriteria_source") or "")
            or "real_drive_50ep_c1c6"
            in str((gt.get("sources") or {}).get("multicriteria_source", ""))
        ):
            checks.append(
                CheckResult(
                    id=f"layer_topsis_real_{name}",
                    category="trazabilidad",
                    status="PASS",
                    message="TOPSIS anclado a multicriterio real_drive_50ep_c1c6",
                    word=ex.name,
                    source=gt.get("sources", {}).get("multicriteria"),
                )
            )

    # Cross-word consistency on anchors
    if len(extracts) == 2:
        names = list(extracts.keys())
        a, b = extracts[names[0]], extracts[names[1]]
        for spec in ANCHOR_SPECS:
            ha = any(re.search(p, a.text) for p in spec["patterns"])
            hb = any(re.search(p, b.text) for p in spec["patterns"])
            if ha and hb:
                checks.append(
                    CheckResult(
                        id=f"cross_{spec['id']}",
                        category="cross_word",
                        status="PASS",
                        message=f"Ambos Word contienen ancla {spec['label']}",
                        expected=str(spec["value"]),
                    )
                )
            elif ha != hb:
                checks.append(
                    CheckResult(
                        id=f"cross_{spec['id']}",
                        category="cross_word",
                        status="WARN",
                        message=f"Ancla {spec['label']} solo en uno de los Word",
                        found=f"{a.name}={ha}, {b.name}={hb}",
                        expected="presente en ambos (o documentar asimetría editorial)",
                        source=spec["source"],
                    )
                )

        # Cap5 score_06667 relative density
        c_a = len(re.findall(r"0[,.]6667", a.text))
        c_b = len(re.findall(r"0[,.]6667", b.text))
        if c_a > 0 and c_b > 0:
            checks.append(
                CheckResult(
                    id="cross_score_density",
                    category="cross_word",
                    status="PASS",
                    message=f"0,6667 en ambos (Tesis={c_a}, Informe={c_b})",
                )
            )
        else:
            checks.append(
                CheckResult(
                    id="cross_score_density",
                    category="cross_word",
                    status="FAIL",
                    message="0,6667 ausente en al menos un Word",
                    found=f"{a.name}={c_a}, {b.name}={c_b}",
                )
            )

    for gap in gt.get("gaps", []):
        checks.append(
            CheckResult(
                id=f"gt_gap_{abs(hash(gap)) % 10_000}",
                category="ground_truth",
                status="GAP",
                message=gap,
            )
        )

    return checks


def redaccion_gaps(extracts: dict[str, WordExtract], checks: list[CheckResult]) -> list[dict]:
    gaps: list[dict] = []
    for name, ex in extracts.items():
        # Conclusions should mention 50 ep + MATD3 / 3 pillars
        concl_hit = re.search(
            r"(?i)cap[ií]tulo\s*6|conclusiones", ex.text
        )
        if concl_hit:
            tail = ex.text[concl_hit.start() : concl_hit.start() + 25000]
            if not re.search(r"(?i)50\s+episodios?", tail):
                gaps.append(
                    {
                        "word": ex.name,
                        "gap": "Conclusiones: no se detecta mención cercana de 50 episodios en ventana Cap.6",
                        "severity": "redaccion",
                    }
                )
            if not re.search(r"\bMATD3\b", tail):
                gaps.append(
                    {
                        "word": ex.name,
                        "gap": "Conclusiones: no se detecta MATD3 (ganador 3×3 canónico) en ventana Cap.6",
                        "severity": "redaccion",
                    }
                )
            for pillar in ("flexibilidad", "CO", "costo"):
                if not re.search(rf"(?i){pillar}", tail):
                    gaps.append(
                        {
                            "word": ex.name,
                            "gap": f"Conclusiones: posible ausencia del pilar '{pillar}' en ventana Cap.6",
                            "severity": "redaccion",
                        }
                    )

        # OG wording alignment (soft)
        if not re.search(r"OG\.\s*-", ex.text):
            gaps.append(
                {
                    "word": ex.name,
                    "gap": "No se encontró el prefijo canónico 'OG. -' (formulación exacta del autor)",
                    "severity": "objetivos",
                }
            )

    # Checks already flagged
    for c in checks:
        if c.category == "redaccion" and c.status in {"FAIL", "WARN"}:
            gaps.append(
                {
                    "word": c.word,
                    "gap": c.message,
                    "severity": "redaccion",
                    "status": c.status,
                }
            )
    return gaps


def verdict_from(checks: list[CheckResult], redaccion: list[dict]) -> dict:
    counts = Counter(c.status for c in checks)
    fails = [c for c in checks if c.status == "FAIL"]
    warns = [c for c in checks if c.status == "WARN"]
    gaps = [c for c in checks if c.status == "GAP"]

    critical_fail_cats = {"anchors", "episodes", "algorithms", "inventory", "ground_truth"}
    critical_fails = [c for c in fails if c.category in critical_fail_cats]

    if critical_fails:
        verdict = "NO"
        rationale = f"{len(critical_fails)} fallos críticos en anclas/episodios/algoritmos/inventario"
    elif fails:
        verdict = "CONDICIONAL"
        rationale = f"{len(fails)} fallos no críticos; revisar antes de declarar versión final"
    elif warns or gaps or any(g.get("severity") == "objetivos" for g in redaccion):
        verdict = "CONDICIONAL"
        rationale = f"{len(warns)} warnings / {len(gaps)} gaps GT / gaps de redacción — listo con salvedades documentadas"
    else:
        verdict = "SÍ"
        rationale = "Todas las pruebas críticas PASS; anclas 50 ep / 4 MADRL trazables"

    return {
        "veredicto": verdict,
        "rationale": rationale,
        "counts": dict(counts),
        "pass": counts.get("PASS", 0),
        "fail": counts.get("FAIL", 0),
        "warn": counts.get("WARN", 0),
        "gap": counts.get("GAP", 0),
        "critical_fails": [asdict(c) for c in critical_fails],
        "all_fails": [asdict(c) for c in fails],
        "all_warns": [asdict(c) for c in warns],
    }


def write_md(report: dict) -> None:
    v = report["veredicto"]
    lines = [
        f"# Validación integral Word 50 ep / 4 MADRL — {DATE_TAG}",
        "",
        f"**Repo:** `D:/MADRLCitytleranflexresdr`",
        f"**Run canónico:** `{RUN_ID}`",
        f"**Drive:** {DRIVE_FOLDER_URL}",
        f"**Generado:** {report['generated_at']}",
        f"**Script:** `tools/thesis/validate_integral_word_50ep_4madrl.py`",
        "",
        f"## Veredicto: **{v['veredicto']}**",
        "",
        v["rationale"],
        "",
        "## Conteos de pruebas",
        "",
        f"| PASS | FAIL | WARN | GAP |",
        f"|---:|---:|---:|---:|",
        f"| {v['pass']} | {v['fail']} | {v['warn']} | {v['gap']} |",
        "",
        "## Word canónicos",
        "",
        "| Archivo | Rol | chars | tablas | 50 ep | HAPPO/MASAC/MATD3/MAAC |",
        "|---|---|---:|---:|---:|---|",
    ]
    for key, w in report["words"].items():
        counts = w["algo_counts"]
        lines.append(
            f"| `{w['path']}` | {key} | {w['chars']} | {w['tables']} | {w['episode_50_mentions']} | "
            f"{counts.get('HAPPO', 0)}/{counts.get('MASAC', 0)}/{counts.get('MATD3', 0)}/{counts.get('MAAC', 0)} |"
        )

    lines += [
        "",
        "## Ground truth (anclas)",
        "",
        "```json",
        json.dumps(report["ground_truth"].get("anchors", {}), indent=2, ensure_ascii=False),
        "```",
        "",
        "### Fuentes",
        "",
    ]
    for k, p in report["ground_truth"].get("sources", {}).items():
        lines.append(f"- **{k}:** `{p}`")

    lines += ["", "## Inconsistencias / fallos", ""]
    fails = v.get("all_fails") or []
    if not fails:
        lines.append("_Ningún FAIL._")
    else:
        for c in fails:
            lines.append(
                f"- **[{c['status']}]** `{c['id']}` ({c['category']}) — {c['message']}"
                + (f" | Word=`{c['word']}`" if c.get("word") else "")
                + (f" | esperado=`{c['expected']}`" if c.get("expected") else "")
                + (f" | hallado=`{(c.get('found') or '')[:120]}`" if c.get("found") else "")
                + (f" | fuente=`{c['source']}`" if c.get("source") else "")
            )

    lines += ["", "## Warnings", ""]
    warns = v.get("all_warns") or []
    if not warns:
        lines.append("_Ningún WARN._")
    else:
        for c in warns:
            lines.append(
                f"- **WARN** `{c['id']}` — {c['message']}"
                + (f" | Word=`{c['word']}`" if c.get("word") else "")
            )

    lines += ["", "## Gaps de redacción vs objetivo doctoral", ""]
    rg = report.get("redaccion_gaps") or []
    if not rg:
        lines.append("_Sin gaps de redacción detectados por heurística._")
    else:
        for g in rg:
            lines.append(
                f"- [{g.get('severity')}] {g.get('word', '')}: {g.get('gap')}"
            )

    lines += [
        "",
        "## Cambios aplicados a Word",
        "",
        report.get("word_patches", "_Ninguno en este pase (solo validación)._"),
        "",
        "## Recomendaciones (sin inventar datos)",
        "",
    ]
    for r in report.get("recommendations", []):
        lines.append(f"- {r}")

    lines += [
        "",
        "## Checklist de pruebas (resumen por categoría)",
        "",
    ]
    by_cat: dict[str, Counter] = {}
    for c in report["checks"]:
        by_cat.setdefault(c["category"], Counter())[c["status"]] += 1
    lines.append("| Categoría | PASS | FAIL | WARN | GAP |")
    lines.append("|---|---:|---:|---:|---:|")
    for cat, ctr in sorted(by_cat.items()):
        lines.append(
            f"| {cat} | {ctr.get('PASS', 0)} | {ctr.get('FAIL', 0)} | {ctr.get('WARN', 0)} | {ctr.get('GAP', 0)} |"
        )

    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    print("[1/4] Cargando ground truth…")
    gt = load_ground_truth()
    print(f"      anclas={list(gt.get('anchors', {}).keys())} gaps={len(gt.get('gaps', []))}")

    print("[2/4] Extrayendo Word canónicos…")
    extracts: dict[str, WordExtract] = {}
    for label, path in (("tesis", TESIS), ("informe", INFORME)):
        if not path.is_file():
            print(f"      FALTA {path}")
            continue
        print(f"      {path.name}…")
        extracts[label] = extract_word(path)
        print(
            f"        chars={extracts[label].chars} tables={extracts[label].tables} "
            f"ep50={extracts[label].episode_50_mentions} algos={extracts[label].algo_counts}"
        )

    print("[3/4] Ejecutando checks…")
    checks = run_checks(gt, extracts)
    redaccion = redaccion_gaps(extracts, checks)
    verdict = verdict_from(checks, redaccion)

    recommendations: list[str] = []
    if verdict["veredicto"] != "SÍ":
        recommendations.append(
            "Corregir FAILs listados antes de declarar versión final cerrada."
        )
    if any(c.status == "WARN" and "49" in c.id for c in checks):
        recommendations.append(
            "Asegurar nota explícita HAPPO 49/50 en ambos Word (trazable a best_madrl)."
        )
    recommendations.append(
        "TOPSIS es medida multicriterio formal adicional para OG/OE/HE en Caps. V–VII; "
        "complementa evaluate_v2/KPI-gains y no debe quedar como disclaimer «solo ilustrativo/no decide»."
    )
    recommendations.append(
        "No mezclar p-valores de HE KPI-gains 3×3 con batería episódica ni presentar TOPSIS "
        "como único veredicto OG en sustitución del omnibus."
    )
    recommendations.append(
        "Tras ediciones: abrir Informe en Word y actualizar TOC (F9)."
    )
    if gt.get("multicriteria_note"):
        recommendations.append(gt["multicriteria_note"]["warning"])

    report = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "date_tag": DATE_TAG,
        "script": "tools/thesis/validate_integral_word_50ep_4madrl.py",
        "canon_policy": str(CANON_MD.relative_to(REPO)).replace("\\", "/"),
        "ground_truth": {
            "run_id": gt.get("run_id"),
            "drive_folder": gt.get("drive_folder"),
            "target_episodes": gt.get("target_episodes"),
            "algorithms": gt.get("algorithms"),
            "sources": gt.get("sources"),
            "anchors": gt.get("anchors"),
            "objectives": gt.get("objectives"),
            "gaps": gt.get("gaps"),
            "multicriteria_note": gt.get("multicriteria_note"),
            "best_madrl_summary": {
                "mejor_madrl": (gt.get("best_madrl") or {}).get("mejor_madrl"),
                "target_episodes": (gt.get("best_madrl") or {}).get("target_episodes"),
                "nota_episodios": (gt.get("best_madrl") or {}).get("nota_episodios"),
            },
        },
        "words": {
            k: {
                "name": v.name,
                "path": v.path,
                "chars": v.chars,
                "paragraphs": v.paragraphs,
                "tables": v.tables,
                "algo_counts": v.algo_counts,
                "episode_50_mentions": v.episode_50_mentions,
                "episode_other_mentions": v.episode_other_mentions,
            }
            for k, v in extracts.items()
        },
        "checks": [asdict(c) for c in checks],
        "veredicto": verdict,
        "redaccion_gaps": redaccion,
        "word_patches": "Ninguno en este pase (solo validación; no se modificaron .docx).",
        "recommendations": recommendations,
    }

    print("[4/4] Escribiendo reportes…")
    OUT_JSON.write_text(
        json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    write_md(report)
    print(f"JSON: {OUT_JSON}")
    print(f"MD:   {OUT_MD}")
    print(
        f"VEREDICTO={verdict['veredicto']} PASS={verdict['pass']} "
        f"FAIL={verdict['fail']} WARN={verdict['warn']} GAP={verdict['gap']}"
    )
    return 0 if verdict["veredicto"] != "NO" else 2


if __name__ == "__main__":
    raise SystemExit(main())
