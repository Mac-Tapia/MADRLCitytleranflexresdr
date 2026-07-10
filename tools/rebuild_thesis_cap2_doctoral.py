from __future__ import annotations

import json
import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
sys_path = str(REPO)
if sys_path not in __import__("sys").path:
    __import__("sys").path.insert(0, sys_path)

from tools.thesis_antecedents_data import (  # noqa: E402
    ANTECEDENTES_INTERNACIONALES,
    ANTECEDENTES_NACIONALES,
    TABLE_INTERNATIONAL,
    TABLE_NATIONAL,
    antecedent_summary_table_rows,
    audit_json_payload,
)
SRC = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_resultados_drive_integrados_ordenado_con_diagramas_estructura_skill_objetivos_operacionalizacion.docx"
OUT = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_resultados_drive_integrados_ordenado_con_diagramas_marco_teorico_doctoral_sustentado.docx"
METRICS = REPO / "outputs" / "_drive_madrl" / "full_data" / "analysis_real_drive" / "thesis_docx_marco_teorico_doctoral_metrics.json"
TRIANGULATION_AUDIT = (
    REPO
    / "outputs"
    / "_drive_madrl"
    / "full_data"
    / "analysis_real_drive"
    / "marco_teorico_variables_triangulation_audit.json"
)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)

TRIANGULATION_ROWS: list[list[str]] = [
    ["VI", "D-VI.1 — Algoritmo MADRL", "Capa cooperativa multiagente bajo Dec-POMDP/CTDE (HAPPO, MASAC, MATD3, MAAC).", "Sutton y Barto (2018); Oliehoek y Amato (2016)", "Factor experimental (4 niveles)"],
    ["VI", "D-VI.1 — HAPPO", "On-policy con region de confianza para agentes heterogeneos.", "Kuba et al. (2021); Zhong et al. (2023)", "T01-T03"],
    ["VI", "D-VI.1 — MASAC", "Off-policy de maxima entropia con exploracion robusta.", "Haarnoja et al. (2018); Gao et al. (2023)", "T04-T06"],
    ["VI", "D-VI.1 — MATD3", "Off-policy con doble critico centralizado y retardo de politica.", "Lowe et al. (2017); Hu et al. (2023)", "T07-T09"],
    ["VI", "D-VI.1 — MAAC", "Critico con atencion selectiva entre agentes.", "Iqbal y Sha (2019)", "T10-T12"],
    ["VI", "D-VI.2 — Escenario E1", "Ponderacion dominante flexibilidad [0,70; 0,15; 0,15].", "Vazquez-Canteli et al. (2020); Yao et al. (2023)", "OE.1 / D-VD.1"],
    ["VI", "D-VI.2 — Escenario E2", "Ponderacion dominante CO2 [0,15; 0,70; 0,15].", "Liu et al. (2022); Tranberg et al. (2020)", "OE.2 / D-VD.2"],
    ["VI", "D-VI.2 — Escenario E3", "Ponderacion dominante costos [0,25; 0,15; 0,60].", "Dang et al. (2023); OSINERGMIN (2024)", "OE.3 / D-VD.3"],
    ["VI", "D-VI.3 — Control dataset", "Dataset comun citylearn_iquitos_2023_2025 (17 edificios, 8 760 h).", "Nweye et al. (2024); Nweye et al. (2023a)", "Mismo schema y semilla"],
    ["VI", "D-VI.3 — Control protocolo", "Funcion de recompensa unified_comparable_v4 y protocolo CTDE.", "Nweye et al. (2022); Hu et al. (2023)", "12 tratamientos 4x3"],
    ["VD", "D-VD.1 — Dim.1 pico", "Reduccion de demanda/importacion maxima distrital.", "Vazquez-Canteli et al. (2020); Lund et al. (2017)", "peak_average"],
    ["VD", "D-VD.1 — Dim.2 rampa", "Suavizado de variaciones horarias de importacion.", "Nweye et al. (2024); Yao et al. (2023)", "ramping_average"],
    ["VD", "D-VD.1 — Dim.3 factor de carga", "Equilibrio entre pico y demanda media (load factor).", "Vazquez-Canteli y Nagy (2019b); Xie et al. (2023)", "one_minus_load_factor_average; flex_composite"],
    ["VD", "D-VD.2 — Dim.1 emisiones totales", "Emisiones anuales agregadas del distrito.", "Liu et al. (2022); MINAM (2019)", "carbon_emissions_total"],
    ["VD", "D-VD.2 — Dim.2 delta CO2", "Reduccion frente a baseline CityLearn v2.", "Tranberg et al. (2020); Sarkar et al. (2024)", "carbon_emissions_delta"],
    ["VD", "D-VD.2 — Dim.3 consumo ponderado CI", "Importacion ponderada por intensidad de carbono horaria.", "Cao et al. (2023); Ye et al. (2025)", "CI-weighted consumption"],
    ["VD", "D-VD.3 — Dim.1 costo total", "Costo electrico anual agregado del distrito.", "Fang et al. (2021); Gao et al. (2023)", "electricity_cost_total"],
    ["VD", "D-VD.3 — Dim.2 delta costo", "Reduccion de costo frente a baseline.", "Dang et al. (2023); Xiong et al. (2024)", "electricity_cost_delta"],
    ["VD", "D-VD.3 — Dim.3 pico facturable", "Reduccion de demanda maxima en ventana de facturacion.", "OSINERGMIN (2024); Shojaeighadikolaei et al. (2022)", "peak_billing; price_signal_deviation"],
]


def triangulation_audit_payload() -> dict:
    return {
        "capitulo": 2,
        "seccion_variables": "2.2",
        "tabla_triangulacion": "Tabla 2.1",
        "vi_dimensiones": {
            "D-VI.1": ["HAPPO", "MASAC", "MATD3", "MAAC"],
            "D-VI.2": ["E1", "E2", "E3"],
            "D-VI.3": ["dataset comun", "protocolo experimental"],
        },
        "vd_dimensiones": {
            "D-VD.1": ["peak_average", "ramping_average", "one_minus_load_factor_average / flex_composite"],
            "D-VD.2": ["carbon_emissions_total", "carbon_emissions_delta", "CI-weighted consumption"],
            "D-VD.3": ["electricity_cost_total", "electricity_cost_delta", "peak_billing / price_signal_deviation"],
        },
        "triangulation_rows": [
            {
                "variable": r[0],
                "dimension": r[1],
                "definicion": r[2],
                "autores_apa": r[3],
                "kpi_tesis": r[4],
            }
            for r in TRIANGULATION_ROWS
        ],
        "min_autores_por_dimension": 2,
        "estructura_cap2": [
            "2.1 Fundamentos teoricos y matematicos",
            "2.1.1 Aprendizaje por refuerzo y MADRL",
            "2.1.2 Formalizacion matematica Dec-POMDP",
            "2.1.3 CityLearn y simulacion multiobjetivo",
            "2.2 Variables de la investigacion",
            "2.2.1 Variable independiente (VI)",
            "2.2.2 Variable dependiente (VD)",
            "2.3 Bases teoricas por eje",
            "2.4 Antecedentes",
            "2.5 Definicion de terminos y posicion teorica",
        ],
        "tablas": [
            "Tabla 2.1 Triangulacion VI/VD",
            "Tabla 2.2 Notacion Dec-POMDP",
            "Tabla 2.3 Algoritmos MADRL",
            TABLE_INTERNATIONAL,
            TABLE_NATIONAL,
            "Tabla 2.6 Constructos teoricos",
            "Tabla 2.7 Antecedentes por eje teorico",
        ],
    }


def text_of(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def clear_body_keep_sectpr(document: Document) -> None:
    body = document.element.body
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def append_before_sectpr(document: Document, el) -> None:
    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(el)
    else:
        body.insert(body.index(sect_pr), el)


def style_doc(document: Document) -> None:
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        if name not in [s.name for s in document.styles]:
            continue
        st = document.styles[name]
        st.font.name = "Calibri"
        if name == "Normal":
            st.font.size = Pt(11)
            st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.line_spacing = 1.15
        else:
            st.font.color.rgb = ACCENT
            st.font.bold = True
            st.font.size = Pt(16 if name == "Heading 1" else 13 if name == "Heading 2" else 11.5)


def set_bg(cell, color: str = "1F4E79") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_chapter_2(doc: Document) -> None:
    def h(text: str, level: int):
        return doc.add_heading(text, level=level)

    def p(text: str):
        para = doc.add_paragraph()
        para.add_run(text)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15
        return para

    def eq(text: str):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.italic = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        return para

    def table(caption: str, headers: list[str], rows: list[list[str]], font_size: float = 7.4):
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GREY
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Light Grid Accent 1"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, head in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = ""
            rr = cell.paragraphs[0].add_run(head)
            rr.bold = True
            rr.font.size = Pt(font_size)
            rr.font.color.rgb = RGBColor(255, 255, 255)
            set_bg(cell)
        for row in rows:
            cells = tbl.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                rr = cells[i].paragraphs[0].add_run(str(val))
                rr.font.size = Pt(font_size)
        doc.add_paragraph()
        return tbl

    h("Capitulo 2. Marco teorico", 1)
    p(
        "El marco teorico de esta tesis se organiza en cinco bloques secuenciales: "
        "(1) fundamentos teoricos y matematicos del aprendizaje por refuerzo multiagente; "
        "(2) definicion operativa y triangulacion bibliografica de la variable independiente "
        "(VI) y la variable dependiente (VD); (3) bases teoricas por eje (flexibilidad, CO2, "
        "costos y algoritmos MADRL); (4) antecedentes internacionales y nacionales; y "
        "(5) definicion de terminos y posicion teorica. Esta secuencia responde a OE.1, OE.2, "
        "OE.3 y al objetivo general, y conecta directamente con la operacionalizacion del "
        "Capitulo 3 (Seccion 3.4) y la contrastacion del Capitulo 5."
    )
    p(
        "La literatura revisada confirma que la gestion energetica de edificios ha transitado "
        "desde control basado en reglas hacia control por aprendizaje, y desde agentes "
        "centralizados hacia coordinacion multiagente. Sin embargo, tambien muestra una brecha: "
        "los estudios suelen evaluar un algoritmo, una metrica dominante o un conjunto limitado "
        "de edificios, mientras que esta tesis compara HAPPO, MASAC, MATD3 y MAAC bajo el mismo "
        "dataset, los mismos escenarios E1/E2/E3, la misma funcion de recompensa y el mismo "
        "protocolo de evidencia. Esa brecha teorica justifica el diseno experimental del "
        "Capitulo 3 y la contrastacion del Capitulo 5."
    )

    # ----- 2.1 Fundamentos teoricos y matematicos -----
    h("2.1 Fundamentos teoricos y matematicos", 2)

    h("2.1.1 Aprendizaje por refuerzo y MADRL", 3)
    p(
        "El aprendizaje por refuerzo modela un problema de decision secuencial donde un agente "
        "observa un estado, ejecuta una accion, recibe una recompensa y modifica su politica para "
        "maximizar retorno acumulado. Sutton y Barto (2018) formalizan este marco mediante "
        "procesos de decision de Markov, lo que permite definir estado, accion, transicion, "
        "recompensa y factor de descuento. En la tesis, esta base se proyecta a un sistema "
        "energetico donde las acciones representan cargar o descargar BESS, regular cargadores "
        "EV y desplazar cargas controlables."
    )
    p(
        "El aprendizaje por refuerzo profundo incorpora aproximadores neuronales para politicas y "
        "funciones de valor, permitiendo espacios de estado y accion de alta dimension. Haarnoja "
        "et al. (2018) introducen Soft Actor-Critic como algoritmo off-policy de maxima entropia, "
        "base conceptual para MASAC y para enfoques donde la exploracion es crucial. Esta "
        "propiedad es importante en comunidades energeticas porque los agentes deben descubrir "
        "estrategias de coordinacion bajo incertidumbre temporal y recompensas multiobjetivo."
    )
    p(
        "En un sistema multiagente, cada politica modifica el entorno efectivo que observan las "
        "demas. Esto genera no estacionariedad desde la perspectiva individual: un edificio "
        "aprende mientras los demas tambien cambian su comportamiento. Oliehoek y Amato (2016) "
        "formalizan esta dificultad en el marco Dec-POMDP, mientras la literatura de CTDE "
        "responde permitiendo criticos centralizados durante el entrenamiento y politicas "
        "descentralizadas durante la ejecucion (Lowe et al., 2017). En la tesis, cada edificio "
        "debe operar con observacion local, pero el aprendizaje puede usar informacion del "
        "distrito para internalizar picos, rampas, emisiones y costos agregados."
    )
    p(
        "La cooperacion se implementa mediante una recompensa mixta que combina recompensa "
        "individual y recompensa de equipo. Esto evita dos extremos: politicas totalmente "
        "egoistas que reducen el costo de un edificio desplazando problemas al distrito, y "
        "politicas totalmente globales que ignoran la heterogeneidad operativa de cada edificio. "
        "La formulacion cooperativa se alinea con los desafios descritos por Nweye et al. (2022) "
        "y con la necesidad de coordinacion observada en demanda respuesta multiagente (Yao et al., "
        "2023; Xie et al., 2023)."
    )

    h("2.1.2 Formalizacion matematica Dec-POMDP", 3)
    p(
        "El Decentralized Partially Observable Markov Decision Process (Dec-POMDP) permite "
        "formalizar decision cooperativa con informacion local incompleta. Oliehoek y Amato (2016) "
        "lo definen como una estructura donde N agentes comparten un criterio cooperativo comun, "
        "pero cada uno recibe observaciones parciales del estado global. Esta tesis adopta esa "
        "formulacion porque ningun edificio del SEAI Iquitos observa completamente el estado "
        "interno de los demas durante la ejecucion: cada agente ve su demanda, PV, BESS, EV, "
        "precio, intensidad de carbono y variables locales, pero no accede a temperatura, SOC, "
        "demanda ni perfiles EV de los otros edificios."
    )
    p(
        "El problema doctoral se modela como el Dec-POMDP cooperativo M, definido por la tupla "
        "siguiente (Oliehoek y Amato, 2016; Sutton y Barto, 2018):"
    )
    eq("M = <S, {A_i}_{i=1}^N, T, R, {O_i}_{i=1}^N, Omega, gamma, T_hor>")
    p(
        "donde N = 17 edificios institucionales/comerciales del dataset citylearn_iquitos_2023_2025, "
        "gamma = 0.9999 (factor de descuento para episodios de 8 760 pasos horarios) y "
        "T_hor = 8 760 (un ano simulado). El objetivo cooperativo es maximizar el retorno "
        "esperado J(pi) = E[ sum_{t=0}^{T_hor-1} gamma^t R_t ], donde pi = (pi_1, ..., pi_N) "
        "denota el conjunto de politicas locales. La Tabla 2.2 resume la notacion formal; "
        "la operacionalizacion computacional se desarrolla en el Capitulo 4 sin alterar esta "
        "definicion teorica."
    )
    table(
        "Tabla 2.2. Notacion formal del Dec-POMDP cooperativo (SEAI Iquitos, N = 17).",
        ["Simbolo", "Definicion teorica", "Valor / rango en esta tesis"],
        [
            ["N", "Numero de agentes cooperativos (edificios)", "17"],
            ["S", "Espacio de estado global", "Concatenacion s = [o_1, ..., o_17]; dim global = 1 856"],
            ["O_i", "Espacio de observacion local del agente i", "Heterogeneo: 57-330 dimensiones segun flota EV"],
            ["A_i", "Espacio de accion local del agente i", "Heterogeneo: 5-44 acciones (BESS, EV, carga desplazable)"],
            ["T", "Funcion de transicion estocastica S x A -> Delta(S)", "Balance energetico, modelo RC, BESS eta_RT = 0.9025, EV estocastico"],
            ["Omega", "Funcion de observacion O_i = Omega_i(s, a)", "Proyeccion parcial del estado global a informacion local"],
            ["R", "Recompensa cooperativa escalar o vector mixto", "CityLearnV3MADRLRewardFunction; agregacion team_mean"],
            ["gamma", "Factor de descuento", "0.9999"],
            ["T_hor", "Horizonte temporal del episodio", "8 760 pasos (1 h/paso)"],
            ["pi_i", "Politica descentralizada del edificio i", "pi_i(a_i | o_i); sin comunicacion inter-edificio"],
        ],
        font_size=7.0,
    )
    p(
        "Recompensa cooperativa multiobjetivo. La funcion R materializa los tres ejes "
        "doctorales (flexibilidad, CO2, costos) mediante una recompensa escalar por edificio "
        "y paso, con agregacion cooperativa tipo media de equipo. A nivel teorico:"
    )
    eq(
        "reward_i(t) = reward_scale * [ w_flex * flex_i(t) + w_carbon * carbon_i(t) "
        "+ w_cost * cost_i(t) + w_ev * ev_i(t) ]"
    )
    p(
        "El componente flex_i(t) penaliza, a nivel distrital compartido, el pico y la rampa "
        "mediante peak_share(t) = district_import(t) / N y ramp_share(t) = "
        "|district_import(t) - district_import(t-1)| / N, con funciones de suavizado tanh; "
        "carbon_i(t) pondera la importacion por la intensidad de carbono CI(t); cost_i(t) "
        "refleja la tarifa TOU mediante price_norm(t); y ev_i(t) incorpora urgencia de SOC y "
        "salida de vehiculos. Los pesos w_flex, w_carbon y w_cost se condicionan por escenario "
        "experimental (E1/E2/E3) segun la Tabla 3.1 del Capitulo 3."
    )
    eq("team_reward(t) = (1/N) * sum_{i=1}^N reward_i(t)")
    eq("mixed_reward_i(t) = (1 - r) * reward_i(t) + r * team_reward(t),   con r = 0.70")
    p(
        "La mezcla cooperativa con team_reward_ratio r = 0.70 evita politicas puramente "
        "egoistas (que desplazan picos o costos al distrito) y politicas puramente globales "
        "(que ignoran heterogeneidad operativa). Este esquema se alinea con la literatura de "
        "recompensa hibrida en MADRL energetico (Yao et al., 2023; Liu et al., 2022) y con "
        "los desafios de coordinacion identificados por Nweye et al. (2022). Durante la "
        "ejecucion descentralizada, cada politica pi_i(a_i | o_i) actua solo con o_i; el estado "
        "global s solo es accesible durante el entrenamiento bajo CTDE (Seccion 2.3.4)."
    )

    h("2.1.3 CityLearn y simulacion multiobjetivo", 3)
    p(
        "CityLearn constituye la linea de simulacion mas directamente vinculada con esta tesis. "
        "Vazquez-Canteli y Nagy (2019a) introdujeron CityLearn v1.0 como entorno compatible con "
        "OpenAI Gym para respuesta a la demanda en edificios, mostrando que un agente de "
        "aprendizaje profundo podia superar politicas basadas en reglas en reduccion de pico. "
        "Posteriormente, Vazquez-Canteli y Nagy (2019b) sistematizaron algoritmos y tecnicas de "
        "modelado para respuesta a la demanda con aprendizaje por refuerzo, y Vazquez-Canteli "
        "et al. (2020) reforzaron la necesidad de estandarizar entornos, KPIs y comparaciones en "
        "gestion energetica urbana."
    )
    p(
        "La evolucion hacia CityLearn v2 es decisiva porque integra edificios grid-interactive, "
        "almacenamiento, PV, vehiculos electricos, senales de carbono y costos. Nweye et al. "
        "(2024) presentan CityLearn v2 como entorno para gestion energetica flexible, resiliente, "
        "centrada en ocupantes y consciente de carbono; Nweye et al. (2023c) describen su "
        "formalizacion como entorno Gym para benchmarking de respuesta a la demanda; y Nweye y "
        "Nagy (2024b) amplian el uso de CityLearn Gym a evaluacion multiobjetivo. Esta tesis toma "
        "CityLearn v2 como base validada, pero no afirma que CityLearn v3 exista oficialmente: "
        "CityLearn v3 propuesto es una extension experimental desarrollada para el Dec-POMDP, "
        "CTDE y backends MADRL del proyecto."
    )
    p(
        "El paso de entornos de un edificio a comunidades heterogeneas demanda tratar problemas "
        "de escala, observabilidad parcial y no estacionariedad. Nweye et al. (2022) identifican "
        "desafios reales del aprendizaje por refuerzo multiagente en edificios grid-interactive, "
        "incluyendo generalizacion, seguridad, heterogeneidad, interaccion con ocupantes y "
        "reproducibilidad. Nweye et al. (2023a) proponen MERLIN para aprendizaje offline y "
        "transferencia en comunidades de 17 edificios, escala comparable al caso SEAI Iquitos, "
        "mientras Nweye et al. (2023b) aplican aprendizaje por refuerzo heterogeneo multiagente "
        "en comunidades grid-interactive. Estos antecedentes justifican que la unidad de analisis "
        "no sea un edificio aislado, sino una comunidad de 17 agentes."
    )

    # ----- 2.2 Variables de la investigacion -----
    h("2.2 Variables de la investigacion", 2)
    p(
        "La variable independiente (VI) y la variable dependiente (VD) se definen operativamente "
        "en esta seccion y se triangulan con autores de la bibliografia (Tabla 2.1). La "
        "triangulacion bibliografica asegura que cada dimension tenga al menos dos fuentes "
        "independientes que sustenten su definicion, siguiendo normas APA 7 en citas narrativas "
        "y parenteticas. La operacionalizacion computacional completa se desarrolla en el "
        "Capitulo 3, Seccion 3.4, que remite a esta seccion como marco conceptual."
    )

    h("2.2.1 Variable independiente (VI)", 3)
    p(
        "La variable independiente es la capa MADRL cooperativa implementada sobre CityLearn v2 "
        "(CityLearn v3 propuesto), manipulada en tres dimensiones: D-VI.1 algoritmo, D-VI.2 "
        "escenario experimental y D-VI.3 control metodologico. Sutton y Barto (2018) y Oliehoek "
        "y Amato (2016) fundamentan la formulacion Dec-POMDP cooperativa que constituye el marco "
        "comun de la VI; Lowe et al. (2017) justifican el paradigma CTDE bajo el cual se "
        "entrenan las politicas."
    )
    p(
        "D-VI.1 — Algoritmo MADRL. Comprende cuatro niveles del factor experimental: HAPPO "
        "(on-policy con region de confianza para agentes heterogeneos; Kuba et al., 2021; "
        "Zhong et al., 2023), MASAC (off-policy de maxima entropia; Haarnoja et al., 2018; "
        "Gao et al., 2023), MATD3 (off-policy con doble critico centralizado; Lowe et al., 2017; "
        "Hu et al., 2023) y MAAC (critico con atencion selectiva; Iqbal y Sha, 2019). Cada "
        "algoritmo representa una familia teorica distinta de coordinacion multiagente bajo el "
        "mismo entorno y la misma funcion de recompensa."
    )
    p(
        "D-VI.2 — Escenario E1/E2/E3. Los escenarios condicionan la politica hacia un objetivo "
        "dominante mediante vectores de peso distintos en la funcion de recompensa multiobjetivo. "
        "E1 prioriza flexibilidad [0,70; 0,15; 0,15] y se sustenta en KPIs de pico y rampa "
        "(Vazquez-Canteli et al., 2020; Yao et al., 2023); E2 prioriza emisiones de CO2 "
        "[0,15; 0,70; 0,15] y se fundamenta en control consciente de intensidad de carbono "
        "(Liu et al., 2022; Tranberg et al., 2020); E3 prioriza costos [0,25; 0,15; 0,60] y se "
        "ancla en respuesta a precios y cargos por demanda (Dang et al., 2023; OSINERGMIN, 2024)."
    )
    p(
        "D-VI.3 — Control experimental. Para aislar el efecto de D-VI.1 y D-VI.2, la tesis "
        "mantiene constantes el dataset citylearn_iquitos_2023_2025 (Nweye et al., 2024; "
        "Nweye et al., 2023a), la funcion de recompensa unified_comparable_v4, la semilla de "
        "reproducibilidad y el protocolo CTDE (Nweye et al., 2022; Hu et al., 2023). Esta "
        "dimension de control asegura que las diferencias observadas en D-VD.1, D-VD.2 y D-VD.3 "
        "se atribuyan al algoritmo y al escenario, no a variaciones del entorno o del protocolo."
    )

    h("2.2.2 Variable dependiente (VD)", 3)
    p(
        "La variable dependiente es el desempeno coordinado de la comunidad energetica medido "
        "en tres ejes: D-VD.1 flexibilidad energetica, D-VD.2 emisiones de CO2 y D-VD.3 "
        "costos energeticos. Cada eje se descompone en tres dimensiones operativas, alineadas "
        "con los KPIs oficiales de CityLearn v2 (Vazquez-Canteli et al., 2020; Nweye et al., "
        "2024). La lectura correcta del Capitulo 5 combina las tres dimensiones de cada eje, "
        "no un KPI aislado."
    )
    p(
        "D-VD.1 — Flexibilidad energetica. Dimension 1 (pico): capacidad de reducir la "
        "importacion maxima distrital, medida con peak_average. Lund et al. (2017) ubican la "
        "reduccion de pico dentro de los sistemas energeticos inteligentes, mientras Vazquez-"
        "Canteli et al. (2020) estandarizan peak_average como KPI comparable en CityLearn. "
        "Dimension 2 (rampa): suavizado de variaciones horarias de importacion, medido con "
        "ramping_average. Nweye et al. (2024) incorporan ramping como KPI distrital en CityLearn "
        "v2, y Yao et al. (2023) reportan mejoras de coordinacion MADRL en perfiles de demanda. "
        "Dimension 3 (factor de carga): equilibrio entre pico y demanda media, medido con "
        "one_minus_load_factor_average y el indicador compuesto flex_composite. Vazquez-Canteli "
        "y Nagy (2019b) vinculan el factor de carga con respuesta a la demanda, y Xie et al. "
        "(2023) demuestran que mecanismos de atencion multiagente mejoran la suavidad del "
        "perfil agregado."
    )
    p(
        "D-VD.2 — Emisiones de CO2. Dimension 1 (emisiones totales): carbon_emissions_total "
        "como indicador agregado anual del distrito. Liu et al. (2022) muestran reducciones de "
        "emisiones con MADRL en sistemas con energia renovable, y MINAM (2019) proporciona el "
        "factor de emision base para contextualizar redes aisladas peruanas. Dimension 2 (delta "
        "CO2): carbon_emissions_delta como reduccion frente a baseline CityLearn v2. Tranberg "
        "et al. (2020) fundamentan la contabilidad temporal de carbono, y Sarkar et al. (2024) "
        "aplican desplazamiento temporal de carga hacia periodos de baja intensidad de carbono. "
        "Dimension 3 (consumo ponderado por CI): importacion electrica ponderada por CI(t). "
        "Cao et al. (2023) estudian pronostico de intensidad de carbono para gestion de "
        "edificios, y Ye et al. (2025) avanzan hacia MARL seguro para operacion baja en carbono."
    )
    p(
        "D-VD.3 — Costos energeticos. Dimension 1 (costo total): electricity_cost_total como "
        "indicador agregado anual. Fang et al. (2021) proponen gestion distribuida de energia "
        "en microredes mediante MADRL, y Gao et al. (2023) validan MASAC para programacion "
        "colaborativa con respuesta a precios. Dimension 2 (delta costo): electricity_cost_delta "
        "como reduccion frente a baseline. Dang et al. (2023) muestran que BESS puede reducir "
        "cargos de demanda bajo precios dinamicos, y Xiong et al. (2024) confirman que DRL "
        "puede responder a tarifas TOU y almacenamiento. Dimension 3 (pico facturable): "
        "reduccion de demanda maxima en ventana de facturacion y desviacion respecto a la senal "
        "de precio (price_signal_deviation). OSINERGMIN (2024) justifica el uso de maxima "
        "demanda y tarifa en el contexto regulatorio peruano, y Shojaeighadikolaei et al. (2022) "
        "plantean gestion energetica distribuida y respuesta a la demanda en smart grids."
    )

    table(
        "Tabla 2.1. Triangulacion bibliografica de variables independientes y dependientes (VI/VD).",
        ["Variable", "Dimension", "Definicion operativa", "Autor(es) APA", "KPI tesis"],
        TRIANGULATION_ROWS,
        font_size=6.8,
    )

    # ----- 2.3 Bases teoricas por eje -----
    h("2.3 Bases teoricas por eje", 2)

    h("2.3.1 Flexibilidad energetica", 3)
    p(
        "La flexibilidad energetica se entiende como la capacidad de modificar el perfil temporal "
        "de demanda, importacion, exportacion, carga de almacenamiento y consumo flexible para "
        "aportar servicios al sistema electrico. En comunidades con PV, BESS y EV, la "
        "flexibilidad no equivale solamente a reducir demanda pico; incluye reducir rampas, "
        "mejorar factor de carga, incrementar autoconsumo y desplazar consumo hacia periodos de "
        "mayor disponibilidad renovable. Este concepto se alinea con los KPIs peak_average, "
        "ramping_average, one_minus_load_factor_average, autoconsumo y autosuficiencia empleados "
        "en CityLearn (Vazquez-Canteli et al., 2020; Nweye et al., 2024)."
    )
    p(
        "Los antecedentes recientes muestran que la flexibilidad requiere coordinacion y no solo "
        "optimizacion local. Yao et al. (2023) proponen una estrategia MADRL para gestion "
        "energetica de comunidades con PV, BESS y EV, reportando mejoras en pico y costo; Xie "
        "et al. (2023) introducen mecanismos de atencion multiagente para respuesta a la demanda "
        "en edificios grid-responsive; y Hribar et al. (2025) demuestran mejoras de autonomia "
        "energetica en distritos de energia positiva mediante MADRL. En paralelo, Felicetti et al. "
        "(2024) combinan programacion entera y aprendizaje por refuerzo para maximizar "
        "autoconsumo y recorte de picos, mientras Li et al. (2024) estudian programacion online "
        "de PV+BESS con DRL. Estos trabajos sustentan el eje OE.1."
    )
    p(
        "La tesis se diferencia porque no evalua una politica unica de flexibilidad. Define E1 "
        "como escenario dominante de flexibilidad y compara cuatro familias MADRL bajo "
        "condiciones identicas. La literatura muestra que la flexibilidad depende de la "
        "interaccion entre edificios, almacenamiento y cargas flexibles; por ello, el Capitulo 5 "
        "no debe interpretar un KPI aislado como evidencia suficiente. La lectura correcta "
        "combina resultados distritales, resultados por edificio, trazas de accion y carga "
        "controlada/no controlada."
    )
    p(
        "Los recursos que producen flexibilidad en el SEAI Iquitos son BESS, EV, PV y cargas "
        "desplazables. La teoria de PV+BESS indica que el almacenamiento permite transferir "
        "energia solar a horas de demanda o costo mayor; los estudios de Li et al. (2024) y "
        "Felicetti et al. (2024) respaldan esa logica. En la tesis, dicha teoria se materializa "
        "con acciones electrical_storage, electric_vehicle_storage_charger_* y "
        "washing_machine_*, mientras la carga base no controlada se mantiene como referencia "
        "para evaluar el efecto real de la accion."
    )

    h("2.3.2 Emisiones de carbono y control consciente de intensidad de carbono", 3)
    p(
        "El segundo eje teorico se relaciona con la gestion energetica orientada a carbono. En "
        "sistemas donde la intensidad de carbono varia temporalmente, el control puede reducir "
        "emisiones desplazando consumo hacia periodos de baja CI o cargando almacenamiento "
        "cuando la energia renovable desplaza generacion fosil. Tranberg et al. (2020) discuten "
        "metodos de contabilidad de carbono en tiempo real; Cao et al. (2023) estudian pronostico "
        "de intensidad de carbono para gestion energetica de edificios; y el Ministerio del "
        "Ambiente del Peru (2019) proporciona la referencia de factor de emision usada para "
        "contextualizar redes aisladas peruanas."
    )
    p(
        "En aprendizaje por refuerzo aplicado a energia, Liu et al. (2022) muestran que MADRL "
        "puede reducir costos y emisiones en sistemas de edificios con energia renovable. Ye "
        "et al. (2025) y Ma et al. (2025) avanzan hacia MARL seguro para operacion baja en "
        "carbono en redes activas y microredes heterogeneas. Sarkar et al. (2024) plantean "
        "reduccion de huella de carbono mediante desplazamiento temporal de carga, lo que es "
        "conceptualmente transferible a la tesis porque E2 tambien depende de desplazar consumo "
        "respecto a una senal de CI. Ren et al. (2025) extienden la discusion a mercados P2P "
        "de baja emision con decisiones multiagente."
    )
    p(
        "El caso SEAI Iquitos requiere adaptar este eje a una red aislada diesel+PV. La "
        "intensidad CI(t) de la tesis se deriva de un factor base de 0,790 kgCO2/kWh y un "
        "desplazamiento asociado a irradiancia, con rango aproximado 0,672-0,790 kgCO2/kWh. "
        "Esta construccion se sostiene teoricamente en la necesidad de hacer carbon-aware "
        "control, pero se implementa como CarbonIntensityModel en el motor CityLearn v3 "
        "propuesto. Asi, la teoria de carbono no queda desconectada del codigo: se transforma "
        "en senal de observacion, componente de recompensa y KPI de evaluacion en D-VD.2."
    )

    h("2.3.3 Costos energeticos, precios dinamicos y respuesta economica", 3)
    p(
        "El tercer eje se refiere a costos energeticos y respuesta a precios. En edificios "
        "comerciales, los costos no dependen solo de energia total, sino de horarios de consumo, "
        "precios, cargos de demanda y comportamiento frente a picos. Dang et al. (2023) estudian "
        "reduccion de cargos por demanda mediante BESS bajo precios en tiempo real; Xiong et al. "
        "(2024) analizan estrategias DRL para sistemas con tarifa por uso horario y "
        "almacenamiento; y Chen et al. (2024) modelan EV como almacenamiento movil en sistemas "
        "energeticos integrados. Estas bases justifican que la tesis incorpore costo energetico "
        "como D-VD.3 y no como metrica secundaria."
    )
    p(
        "La literatura MADRL de costos refuerza la necesidad de coordinacion. Fang et al. (2021) "
        "proponen gestion distribuida de energia y estrategia de mercado en microredes mediante "
        "MADRL; Gao et al. (2023) desarrollan MASAC mejorado para programacion colaborativa "
        "multi-microgrid; Shojaeighadikolaei et al. (2022) plantean gestion energetica "
        "distribuida y respuesta a la demanda en smart grids; y Shojaeighadikolaei et al. (2024) "
        "comparan enfoques centralizados y descentralizados para control de redes de carga EV. "
        "Estos estudios muestran que el costo emerge de decisiones coordinadas de multiples "
        "recursos, no de la optimizacion aislada de un edificio."
    )
    p(
        "La tesis adopta una senal TOU propia del contexto Iquitos y la incorpora en E3. En "
        "consecuencia, el marco teorico de costos se vincula con decisiones de carga/descarga "
        "BESS, programacion EV, respuesta a precio y reduccion de picos facturables. La "
        "inclusion de OSINERGMIN (2024) permite conectar el KPI de pico y costo con reglas "
        "regulatorias peruanas, evitando que el analisis economico sea un ejercicio abstracto "
        "separado de la realidad institucional. La importancia doctoral de este eje es que "
        "costo, flexibilidad y carbono pueden entrar en conflicto; por eso, la recompensa "
        "multiobjetivo y los escenarios E1/E2/E3 son teoricamente necesarios."
    )

    h("2.3.4 Algoritmos MADRL evaluados", 3)
    p(
        "HAPPO se fundamenta en optimizacion de politica con restricciones de region de confianza "
        "para agentes heterogeneos. Kuba et al. (2021) desarrollan una formulacion de "
        "trust-region policy optimization multiagente, y Zhong et al. (2023) profundizan el "
        "aprendizaje por refuerzo de agentes heterogeneos. En la tesis, HAPPO es relevante "
        "porque los edificios difieren en area, cargas, PV, BESS y numero de EV."
    )
    p(
        "MASAC deriva conceptualmente de Soft Actor-Critic y su regularizacion de entropia "
        "(Haarnoja et al., 2018). Gao et al. (2023) muestran la aplicabilidad de un MASAC "
        "mejorado en programacion colaborativa multi-microgrid. En la tesis, MASAC representa "
        "un enfoque off-policy con exploracion robusta y adaptacion de acciones para el entorno "
        "CityLearn v3 propuesto."
    )
    p(
        "MATD3 se basa en la idea de reducir sobreestimacion mediante criticos dobles, retraso "
        "de actualizacion de politica y ruido objetivo en la familia actor-critic multiagente "
        "(Lowe et al., 2017). En la tesis, MATD3 se plantea en las hipotesis como candidato de "
        "mayor efecto coordinado porque su doble critico puede ser ventajoso en horizontes "
        "largos y acciones energeticas continuas."
    )
    p(
        "MAAC se sustenta en Actor-Attention-Critic. Iqbal y Sha (2019) proponen que el critico "
        "utilice atencion para seleccionar los agentes mas relevantes durante la evaluacion de "
        "acciones. En comunidades energeticas, esta idea es pertinente porque no todos los "
        "edificios interactuan con la misma intensidad en cada hora: un hospital, un mall y una "
        "universidad presentan perfiles de carga y EV diferentes."
    )
    p(
        "El paradigma CTDE separa dos fases: durante el entrenamiento, los criticos pueden "
        "acceder al estado global; durante la ejecucion, las politicas actuan solo con "
        "observaciones locales (Lowe et al., 2017; Iqbal y Sha, 2019). MARLlib se considera "
        "solo como nombre propio de una biblioteca de referencia (Hu et al., 2023), no como "
        "sustituto conceptual de MADRL."
    )
    table(
        "Tabla 2.3. Fundamento teorico de los algoritmos MADRL evaluados.",
        ["Algoritmo", "Base teorica", "Ventaja esperada", "Riesgo metodologico", "Relacion con VI"],
        [
            ["HAPPO", "Trust-region y agentes heterogeneos (Kuba et al., 2021; Zhong et al., 2023).", "Estabilidad on-policy y tratamiento de heterogeneidad entre edificios.", "Costo de muestreo y cobertura parcial de artefactos finales.", "Nivel D-VI.1 del factor algoritmo."],
            ["MASAC", "SAC y regularizacion de entropia (Haarnoja et al., 2018; Gao et al., 2023).", "Exploracion robusta y aprendizaje off-policy.", "Adaptacion de acciones continuas/discretas y sensibilidad a hiperparametros.", "Nivel D-VI.1 del factor algoritmo."],
            ["MATD3", "Criticos dobles, retardo de politica y control continuo off-policy (Lowe et al., 2017).", "Reduccion de sobreestimacion y estabilidad en acciones continuas.", "Puede optimizar algunos ejes mejor que otros; requiere lectura por escenario.", "Nivel D-VI.1 e hipotesis direccional HE.2."],
            ["MAAC", "Critico con atencion multiagente (Iqbal y Sha, 2019).", "Coordinacion selectiva entre edificios heterogeneos.", "Complejidad computacional y sensibilidad a estructura de interacciones.", "Nivel D-VI.1 del factor algoritmo."],
        ],
    )

    h("2.3.5 Aportes fisicos al motor como base teorica de CityLearn v3 propuesto", 3)
    p(
        "El marco teorico tambien debe sustentar los aportes del motor de simulacion. La "
        "degradacion BESS con C-rate y temperatura se justifica por literatura de envejecimiento "
        "de baterias LiFePO4 y modelado de degradacion (Naumann et al., 2021; Rajagopalan et al., "
        "2024; Reniers et al., 2022; Xu et al., 2021). Esto evita que el entorno trate toda "
        "accion de almacenamiento como equivalente, ignorando condiciones termicas propias de "
        "Iquitos."
    )
    p(
        "La correccion PV por temperatura tropical se sustenta en IEC 61215-1:2021, Tamoor et al. "
        "(2022) y Antonanzas et al. (2021). En Iquitos, alta temperatura y humedad pueden reducir "
        "produccion frente a condiciones STC, por lo que un modelo sin correccion termica sesgaria "
        "las politicas MADRL hacia una confianza excesiva en PV de mediodia. El KPI de pico con "
        "ventana de facturacion se fundamenta en OSINERGMIN (2024) y en estudios de reduccion de "
        "demanda maxima con BESS (Dang et al., 2023). Finalmente, CarbonIntensityModel se "
        "sostiene en MINAM (2019), Tranberg et al. (2020) y Cao et al. (2023)."
    )

    table(
        "Tabla 2.6. Constructos teoricos, indicadores y fuentes.",
        ["Constructo", "Indicadores en tesis", "Fuentes teoricas principales", "Uso en Cap. 5"],
        [
            ["Flexibilidad energetica", "peak_average, ramping_average, one_minus_load_factor_average, autoconsumo, autosuficiencia.", "Vazquez-Canteli et al. (2020); Nweye et al. (2024); Lund et al. (2017).", "Comparacion E1, figuras A.1-A.2/A.7-A.8 y KPIs edificio."],
            ["Emisiones CO2", "carbon_emissions_total, carbon_emissions_delta, CI-weighted consumption.", "Liu et al. (2022); Tranberg et al. (2020); Cao et al. (2023); MINAM (2019).", "Comparacion E2, figura A.4/A.6 y tablas por edificio."],
            ["Costos energeticos", "electricity_cost_total, electricity_cost_delta, price_signal_deviation.", "Dang et al. (2023); Xiong et al. (2024); OSINERGMIN (2024); Gao et al. (2023).", "Comparacion E3, figura A.3/A.5 y deltas por edificio."],
            ["Coordinacion MADRL", "reward, trace por agente, acciones controladas, checkpoints.", "Lowe et al. (2017); Oliehoek y Amato (2016); Iqbal y Sha (2019); Kuba et al. (2021).", "Interpretacion distrito-edificio y reproducibilidad de politica."],
            ["Fidelidad fisica del simulador", "degradacion BESS, PV tropical, CI dinamica, pico facturable.", "Naumann et al. (2021); Tamoor et al. (2022); IEC (2021); MINAM (2019).", "Justifica CityLearn v3 propuesto y anexos de arquitectura."],
        ],
    )

    # ----- 2.4 Antecedentes -----
    h("2.4 Antecedentes", 2)
    p(
        "Los antecedentes se organizan en internacionales (Tabla 2.4), nacionales/peruanos "
        "(Tabla 2.5) y una sintesis critica de brecha cientifica. La Tabla 2.7 resume la "
        "relacion entre antecedentes, ejes teoricos y aportes de esta tesis."
    )
    table(
        "Tabla 2.7. Antecedentes usados por eje teorico y relacion con la tesis.",
        ["Eje", "Autores fuente", "Aporte teorico usado", "Limitacion que cubre esta tesis"],
        [
            ["CityLearn y benchmarking", "Vazquez-Canteli y Nagy (2019a, 2019b); Vazquez-Canteli et al. (2020); Nweye et al. (2023c, 2024); Nweye y Nagy (2024b)", "Entorno estandarizado, KPIs, comunidades grid-interactive, carbon-aware control.", "Se extiende a CityLearn v3 propuesto con Dec-POMDP/CTDE y cuatro backends MADRL."],
            ["Escala multiagente realista", "Nweye et al. (2022, 2023a, 2023b)", "Desafios reales, MERLIN y comunidades heterogeneas de 17 edificios.", "Se aplica a SEAI Iquitos con 17 edificios reales y artefactos Drive."],
            ["Flexibilidad energetica", "Yao et al. (2023); Xie et al. (2023); Hribar et al. (2025); Felicetti et al. (2024); Li et al. (2024); Zhao et al. (2024); Wu et al. (2025)", "Peak shaving, ramping, autoconsumo, coordinacion DR, PV+BESS, control seguro.", "E1 compara HAPPO/MASAC/MATD3/MAAC bajo mismo dataset y KPIs."],
            ["Emisiones CO2", "Liu et al. (2022); Ye et al. (2025); Ma et al. (2025); Sarkar et al. (2024); Ren et al. (2025); Tranberg et al. (2020); Cao et al. (2023)", "Control bajo intensidad de carbono, operacion baja en carbono, desplazamiento temporal de carga.", "E2 usa CI de red aislada Iquitos y CarbonIntensityModel."],
            ["Costos energeticos", "Fang et al. (2021); Gao et al. (2023); Shojaeighadikolaei et al. (2022, 2024); Xiong et al. (2024); Chen et al. (2024); Dang et al. (2023)", "Respuesta a precios, costo electrico, microredes, EV, BESS y demanda maxima.", "E3 vincula TOU local, BESS, EV, carga desplazable y costo distrital/edificio."],
            ["Fundamentos MADRL", "Sutton y Barto (2018); Oliehoek y Amato (2016); Lowe et al. (2017); Haarnoja et al. (2018); Iqbal y Sha (2019); Kuba et al. (2021); Zhong et al. (2023); Hu et al. (2023)", "MDP, Dec-POMDP, CTDE, SAC, atencion, HAPPO, MARLlib.", "Se operacionaliza en wrappers reales y tratamientos experimentales."],
            ["Modelado fisico y regulatorio", "Naumann et al. (2021); Rajagopalan et al. (2024); Reniers et al. (2022); Xu et al. (2021); Tamoor et al. (2022); Antonanzas et al. (2021); IEC (2021); MINAM (2019); OSINERGMIN (2024)", "Degradacion BESS, correccion PV tropical, carbono y tarifa.", "Aportes A1-A4 del motor CityLearn v3 propuesto."],
        ],
    )

    h("2.4.1 Antecedentes internacionales", 3)
    p(
        f"La {TABLE_INTERNATIONAL} sistematiza cinco antecedentes internacionales directamente "
        "vinculados con CityLearn, MADRL (HAPPO/MAAC/MADDPG), flexibilidad, CO2 y costos. "
        "Estos estudios constituyen la linea de comparacion externa para PE.1-PE.3 y para la "
        "contrastacion de hipotesis en el Capitulo 5."
    )
    table(
        f"{TABLE_INTERNATIONAL}. Antecedentes internacionales seleccionados (Modulo A).",
        ["Autor(es)", "Titulo", "Variables VI/VD", "Aporte para esta tesis", "DOI / enlace"],
        antecedent_summary_table_rows(ANTECEDENTES_INTERNACIONALES),
        font_size=6.8,
    )
    for ant in ANTECEDENTES_INTERNACIONALES:
        h(f"{ant['cita']}", 4)
        table(
            f"Ficha antecedente internacional — {ant['cita']}.",
            ["Campo", "Contenido"],
            [
                ["Titulo de la tesis/estudio", ant["titulo"]],
                ["Objetivo general", ant["objetivo_general"]],
                ["Tipo de metodologia", ant["tipo_metodologia"]],
                ["Diseno de investigacion", ant["diseno_investigacion"]],
                ["Resultados cuantitativos", ant["resultados_cuantitativos"]],
                ["Conclusion general", ant["conclusion_general"]],
                ["Aporte para esta tesis", ant["aporte_tesis"]],
            ],
            font_size=7.0,
        )

    h("2.4.2 Antecedentes nacionales y peruanos", 3)
    p(
        f"La {TABLE_NATIONAL} integra antecedentes nacionales y del contexto SEAI Iquitos: tesis "
        "doctorales peruanas/latinoamericanas y marcos regulatorios MINAM/OSINERGMIN. "
        "Estos antecedentes anclan la tesis en el sistema electrico aislado diesel+PV de Loreto."
    )
    table(
        f"{TABLE_NATIONAL}. Antecedentes nacionales/peruanos seleccionados.",
        ["Autor(es)", "Titulo", "Variables VI/VD", "Aporte para esta tesis", "DOI / enlace"],
        antecedent_summary_table_rows(ANTECEDENTES_NACIONALES),
        font_size=6.8,
    )
    for ant in ANTECEDENTES_NACIONALES:
        h(f"{ant['cita']}", 4)
        table(
            f"Ficha antecedente nacional — {ant['cita']}.",
            ["Campo", "Contenido"],
            [
                ["Titulo de la tesis/estudio", ant["titulo"]],
                ["Objetivo general", ant["objetivo_general"]],
                ["Tipo de metodologia", ant["tipo_metodologia"]],
                ["Diseno de investigacion", ant["diseno_investigacion"]],
                ["Resultados cuantitativos", ant["resultados_cuantitativos"]],
                ["Conclusion general", ant["conclusion_general"]],
                ["Aporte para esta tesis", ant["aporte_tesis"]],
            ],
            font_size=7.0,
        )

    h("2.4.3 Sintesis critica de antecedentes y brecha cientifica", 3)
    p(
        "La revision muestra cinco patrones. Primero, CityLearn proporciona un entorno "
        "estandarizado, pero no resuelve por si mismo la comparacion entre HAPPO, MASAC, MATD3 "
        "y MAAC. Segundo, los estudios de flexibilidad muestran beneficios de control aprendido, "
        "pero frecuentemente se concentran en pico o autoconsumo. Tercero, los estudios de "
        "carbono muestran utilidad de senales CI, pero rara vez se integran con costo y "
        "flexibilidad en un unico diseno factorial. Cuarto, la literatura de costos evidencia la "
        "importancia de precios y BESS, pero no siempre considera emisiones. Quinto, los "
        "algoritmos MADRL tienen fundamentos distintos, por lo que comparar solo resultados "
        "agregados sin control metodologico puede inducir conclusiones debiles."
    )
    p(
        "La brecha que sostiene esta tesis es metodologica y experimental: falta una evaluacion "
        "comparativa, reproducible y multiobjetivo de HAPPO, MASAC, MATD3 y MAAC en una "
        "comunidad inteligente realista de 17 edificios, bajo Dec-POMDP, CTDE, dataset comun, "
        "recompensa unificada y resultados trazables por distrito, edificio, escenario, KPIs y "
        "checkpoints. Esta brecha justifica la variable independiente, la variable dependiente "
        "y la matriz de operacionalizacion (Tabla 2.1 y Capitulo 3, Seccion 3.4)."
    )

    # ----- 2.5 Definicion de terminos y posicion teorica -----
    h("2.5 Definicion de terminos y posicion teorica", 2)
    h("2.5.1 Definicion de terminos y delimitaciones conceptuales", 3)
    p(
        "MADRL: aprendizaje por refuerzo profundo multiagente, donde multiples agentes aprenden "
        "politicas mediante redes neuronales y senales de recompensa. En esta tesis se usa MADRL "
        "para referirse a HAPPO, MASAC, MATD3 y MAAC bajo formulacion cooperativa. No se "
        "sustituye por MARL salvo cuando aparece como parte del nombre propio de una referencia, "
        "repositorio o biblioteca."
    )
    p(
        "CityLearn v2: entorno base oficial de simulacion para comunidades energeticas "
        "grid-interactive, con KPIs de energia, carbono y costo. CityLearn v3 propuesto: "
        "extension experimental de esta tesis, implementada localmente para Dec-POMDP, CTDE, "
        "recompensa multiobjetivo, wrappers MADRL y artefactos reproducibles. Esta distincion es "
        "obligatoria para no atribuir al paquete oficial una funcionalidad propia del proyecto."
    )
    p(
        "Tratamiento experimental: combinacion de un nivel D-VI.1 algoritmo y un nivel D-VI.2 "
        "escenario. El diseno factorial completo tiene 12 tratamientos. Distrito: agregacion de "
        "los 17 edificios para medir efecto comunitario. Edificio: agente individual con "
        "observacion local y acciones propias. Equipo controlado: variable de accion sobre BESS, "
        "EV o carga desplazable. Carga no controlada: demanda base observada o de referencia "
        "que no es actuada directamente por el agente."
    )
    p(
        "KPI: indicador cuantitativo usado para medir la variable dependiente. Los KPIs no son "
        "adornos del resultado; son la operacionalizacion de D-VD.1, D-VD.2 y D-VD.3. Por ello, "
        "toda grafica o tabla del Capitulo 5 debe indicar si corresponde a flexibilidad, "
        "carbono, costo, distrito, edificio, escenario o reproducibilidad de modelo."
    )

    h("2.5.2 Posicion teorica de la tesis", 3)
    p(
        "La posicion teorica adoptada es que una comunidad inteligente con recursos DER "
        "heterogeneos debe modelarse como un sistema multiagente parcialmente observable y no "
        "como un problema de control centralizado unico. El criterio de calidad no es solo "
        "minimizar un KPI, sino producir desempeno coordinado medible en tres dimensiones por "
        "eje (Tabla 2.1). Esta posicion combina la teoria Dec-POMDP de Oliehoek y Amato (2016), "
        "el paradigma CTDE de Lowe et al. (2017), la estandarizacion de CityLearn de Vazquez-"
        "Canteli et al. (2020) y Nweye et al. (2024), y la literatura de gestion energetica "
        "multiobjetivo."
    )
    p(
        "Desde esta posicion, el Capitulo 5 debe interpretar los resultados con tres reglas. "
        "Primero, no hay superioridad universal sin declarar escenario, KPI y escala. Segundo, "
        "un resultado distrital no reemplaza la evidencia por edificio. Tercero, un algoritmo "
        "con artefactos incompletos no debe usarse para comparaciones donde falten archivos. "
        "Estas reglas derivan directamente del marco teorico y evitan conclusiones no sustentadas."
    )


def main() -> None:
    shutil.copyfile(SRC, OUT)
    doc = Document(OUT)
    style_doc(doc)

    children = list(doc.element.body)
    idx_cap2 = idx_cap3 = None
    for i, el in enumerate(children):
        txt = text_of(el)
        if idx_cap2 is None and txt.startswith("Capitulo 2. Marco teorico"):
            idx_cap2 = i
        if idx_cap3 is None and txt.startswith("Capitulo 3. Metodologia"):
            idx_cap3 = i
    if idx_cap2 is None or idx_cap3 is None:
        raise RuntimeError(f"No se ubicaron limites Cap2/Cap3: {idx_cap2}, {idx_cap3}")

    before = [deepcopy(el) for el in children[:idx_cap2]]
    after = [deepcopy(el) for el in children[idx_cap3:] if el.tag != qn("w:sectPr")]

    clear_body_keep_sectpr(doc)
    for el in before:
        append_before_sectpr(doc, el)
    add_chapter_2(doc)
    for el in after:
        append_before_sectpr(doc, el)

    doc.save(OUT)

    v = Document(OUT)
    paras = [p.text.strip() for p in v.paragraphs if p.text.strip()]
    full = "\n".join(paras)
    cap2_start = next(i for i, x in enumerate(paras) if x == "Capitulo 2. Marco teorico")
    cap3_start = next(i for i, x in enumerate(paras) if x == "Capitulo 3. Metodologia")
    cap2_text = "\n".join(paras[cap2_start:cap3_start])
    metrics = {
        "output": str(OUT),
        "size_bytes": OUT.stat().st_size,
        "paragraphs_non_empty": len(paras),
        "word_count_estimated": len(re.findall(r"\b[\wáéíóúÁÉÍÓÚñÑüÜ-]+\b", full, re.UNICODE)),
        "cap2_word_count_estimated": len(re.findall(r"\b[\wáéíóúÁÉÍÓÚñÑüÜ-]+\b", cap2_text, re.UNICODE)),
        "tables": len(v.tables),
        "inline_images": len(v.inline_shapes),
        "cap2_tables_expected": all(
            x in cap2_text
            for x in [
                "Tabla 2.1",
                "Tabla 2.2",
                "Tabla 2.3",
                "Tabla 2.4",
                "Tabla 2.5",
                "Tabla 2.6",
                "Tabla 2.7",
            ]
        ),
        "cap2_has_section_21_fundamentos": "2.1 Fundamentos teoricos y matematicos" in cap2_text,
        "cap2_has_section_22_variables": "2.2 Variables de la investigacion" in cap2_text,
        "cap2_has_triangulation_table": "Triangulacion bibliografica de variables" in cap2_text,
        "cap2_vi_dimensions": all(x in cap2_text for x in ["D-VI.1", "D-VI.2", "D-VI.3"]),
        "cap2_vd_dimensions": all(x in cap2_text for x in ["D-VD.1", "D-VD.2", "D-VD.3"]),
        "cap2_antecedents_international": len(ANTECEDENTES_INTERNACIONALES),
        "cap2_antecedents_national": len(ANTECEDENTES_NACIONALES),
        "cap2_has_dec_pomdp_ctde": "Dec-POMDP" in cap2_text and "CTDE" in cap2_text,
        "cap2_has_dec_pomdp_tuple": "M = <S, {A_i}_{i=1}^N" in cap2_text,
        "cap2_has_reward_equations": "reward_i(t)" in cap2_text and "team_reward(t)" in cap2_text,
        "cap2_has_n17_agents": "N = 17" in cap2_text,
        "cap2_has_citylearn_v3_propuesto": "CityLearn v3 propuesto" in cap2_text,
        "cap2_citation_year_markers": len(re.findall(r"\(\d{4}[a-z]?\)", cap2_text)),
        "wrong_redaccion_blocks": "Redaccion doctoral ampliada" in full,
        "figures_a_1_a_9": all(f"Figura A.{i}" in full for i in range(1, 10)),
        "figures_b_1_a_9": all(f"Figura B.{i}" in full for i in range(1, 10)),
    }
    METRICS.parent.mkdir(parents=True, exist_ok=True)
    METRICS.write_text(json.dumps(metrics, ensure_ascii=False, indent=2), encoding="utf-8")
    audit_path = (
        REPO
        / "outputs"
        / "madrl_v3_20260627_164047"
        / "resumen_comparativo"
        / "estadistica"
        / "antecedentes_tesis_audit.json"
    )
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    audit_path.write_text(json.dumps(audit_json_payload(), ensure_ascii=False, indent=2), encoding="utf-8")
    TRIANGULATION_AUDIT.parent.mkdir(parents=True, exist_ok=True)
    TRIANGULATION_AUDIT.write_text(
        json.dumps(triangulation_audit_payload(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    metrics["antecedentes_audit_json"] = str(audit_path)
    metrics["triangulation_audit_json"] = str(TRIANGULATION_AUDIT)
    print(json.dumps(metrics, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
