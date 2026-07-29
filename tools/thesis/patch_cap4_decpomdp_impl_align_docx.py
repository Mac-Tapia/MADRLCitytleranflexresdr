#!/usr/bin/env python3
"""Actualiza datos Dec-POMDP en Cap. IV de los 2 Word canónicos (sin reescribir el bloque).

Solo corrige/amplía párrafos existentes de §4.2 (dims, r_team, anclaje implementación).
Idempotente.
"""
from __future__ import annotations

import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, INFORME, TESIS  # noqa: E402

REPORT = DOCS / "CAP4_DECPOMDP_IMPL_ALIGN_PATCH_REPORT_2026-07-29.json"
BACKUP_DIR = REPO / "outputs" / "_word_backups"
MARKER = "d_ai = 2 + n_i^ch"


def norm(s: str) -> str:
    t = (s or "").strip().lower()
    for a, b in (("á", "a"), ("é", "e"), ("í", "i"), ("ó", "o"), ("ú", "u"), ("ñ", "n")):
        t = t.replace(a, b)
    return re.sub(r"\s+", " ", t)


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def in_cap4_bounds(doc: Document) -> tuple[int, int]:
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if start is None and (
            re.search(r"^Cap[ií]tulo\s*4\b", t, re.I)
            or t.startswith("4.2 Modelo de IA")
            or t.startswith("4.2 Modelo de IA")
        ):
            if re.search(r"^Cap[ií]tulo\s*4\b", t, re.I) or t.startswith("4.2"):
                start = i
        if start is not None and re.search(r"^Cap[ií]tulo\s*5\b", t, re.I):
            end = i
            break
    if start is None:
        # fallback: first 4.2 Dec-POMDP heading
        for i, p in enumerate(doc.paragraphs):
            if "4.2" in (p.text or "") and "Dec-POMDP" in (p.text or ""):
                start = i
                break
    return start or 0, end or len(doc.paragraphs)


def patch_paragraphs(doc: Document) -> list[str]:
    changes: list[str] = []
    start, end = in_cap4_bounds(doc)
    for i in range(start, end):
        p = doc.paragraphs[i]
        t = p.text or ""
        nt = norm(t)
        new = t

        # Estado global: anclar ctde + suma
        if nt.startswith("estado global s") and "1856" in t.replace(" ", ""):
            if "sum" not in nt and MARKER.split("=")[0].strip() not in t:
                new = (
                    "Estado global S: concatenacion de observaciones locales "
                    "s = [o_1, ..., o_17] (ctde_state / concatenated_local_observations_for_ctde "
                    "en citylearn/v3/backends.py y CityLearnDecPOMDPEnv.state); "
                    "dim global medida = 1 856 (= suma_i d_oi), accesible solo por el critico "
                    "centralizado durante el entrenamiento."
                )

        # Observacion: quitar "aproximado", fijar medido + Tabla 2.A
        if ("observacion local" in nt or "observación local" in t.lower()) and (
            "54" in t and "327" in t
        ):
            if "tabla 2.a" not in nt or "aproximado" in nt or "citylearnenv" not in nt:
                new = (
                    "Observacion local Oi: heterogenea por edificio (tiempo, fisica, meteo, "
                    "precio, CI, SOC BESS y 7 canales EV por cargador); dimension medida en "
                    "CityLearnEnv = 54-327 (B05/B11/B12=54; B07=327). Desglose: Tabla 2.A Cap. 2."
                )

        # Accion: ~5-~44 -> exacto + formula
        if ("accion local" in nt or "acción local" in t.lower()) and (
            "5" in t and "44" in t
        ):
            if MARKER not in t and "2 + n" not in t:
                new = (
                    "Accion local Ai: electrical_storage + electric_vehicle_storage x n_i^ch "
                    "+ washing_machine; d_ai = 2 + n_i^ch en [5, 44] (suma distrital = 219). "
                    "B06 (32 carg.) d_ai=34; B07 (42 carg.) d_ai=44 concentran la mayor dimension."
                )

        # R: añadir r_team=0.70 si falta
        if nt.startswith("r:") and "citylearnv3madrlrewardfunction" in nt.replace(" ", ""):
            if "0.70" not in t and "0,70" not in t and "team_reward_ratio" not in nt:
                new = (
                    "R: recompensa cooperativa CityLearnV3MADRLRewardFunction; "
                    "team_reward = media de reward_i; mixed_i = (1 - r_team)*reward_i + "
                    "r_team*team_reward con team_reward_ratio r_team = 0,70 "
                    "(perfil unified_comparable_v4; canal de las 12 corridas)."
                )

        # Observabilidad: una frase distrito
        if "observabilidad parcial" in nt and "distrito" not in nt:
            new = (
                t.rstrip().rstrip(".")
                + ". El distrito no es un agente central: coordina via critico CTDE (s) "
                "y team_reward (r_team = 0,70); en ejecucion solo persiste pi_i(a_i | o_i)."
            )

        if new != t:
            set_paragraph_text(p, new)
            changes.append(f"p{i}")
    return changes


def already_aligned(doc: Document) -> bool:
    start, end = in_cap4_bounds(doc)
    blob = "\n".join((doc.paragraphs[i].text or "") for i in range(start, end))
    return MARKER in blob and "team_reward_ratio" in blob and "Tabla 2.A" in blob


def patch_doc(path: Path) -> dict:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / (
        f"{path.stem}_before_cap4_decpomdp_{datetime.now().strftime('%Y%m%d_%H%M%S')}{path.suffix}"
    )
    shutil.copy2(path, backup)
    doc = Document(str(path))
    if already_aligned(doc):
        backup.unlink(missing_ok=True)
        return {"path": str(path), "ok": True, "skipped": True, "reason": "already aligned"}
    changes = patch_paragraphs(doc)
    doc.save(str(path))
    return {
        "path": str(path),
        "ok": True,
        "backup": str(backup),
        "changes": changes,
        "n_changes": len(changes),
    }


def main() -> int:
    results = [patch_doc(p) for p in (TESIS, INFORME) if p.is_file()]
    report = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "policy": "solo actualizar datos Cap.4 Dec-POMDP; sin reescritura amplia",
        "results": results,
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0 if results and all(r.get("ok") for r in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
