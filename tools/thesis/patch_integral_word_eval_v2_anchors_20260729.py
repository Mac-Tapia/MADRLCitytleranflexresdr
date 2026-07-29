#!/usr/bin/env python3
"""Parches mínimos de trazabilidad en los 2 Word canónicos (2026-07-29).

Correcciones (solo datos reales del repo; nada inventado):
1) Completar ranking evaluate_v2 4/4 con MATD3 0,8805 y MASAC 0,8679
   (fuente: outputs/_drive_madrl/kpi_recalc_20260728/tables/ranking_oe_scores_all_values.csv).
2) Corregir celda corrupta Δ MATD3 E2 en Informe: '23 07047' → '46.014,7548'
   (Control−Baseline de la misma fila de la tabla).

Backups: outputs/_word_backups/*.pre_integral_validation_20260729.bak
"""

from __future__ import annotations

import json
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from docx import Document  # noqa: E402
from thesis_word_canons import INFORME, TESIS  # noqa: E402

BACKUP_DIR = REPO / "outputs" / "_word_backups"
REPORT = REPO / "docs" / "VALIDACION_INTEGRAL_WORD_PATCHES_2026-07-29.json"

# evaluate_v2 4/4 from ranking_oe_scores_all_values.csv (all4_including_happo)
REPLACEMENTS_ES = [
    (
        "ranking evaluate_v2 4/4 lidera MAAC 0,9538).",
        "ranking evaluate_v2 4/4: MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000).",
    ),
    (
        "ranking evaluate_v2 4/4 lidera MAAC 0,9538 con HAPPO en score 0,0000)",
        "ranking evaluate_v2 4/4: MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000)",
    ),
    (
        "y el ranking evaluate_v2 4/4 (0,9538).",
        "y el ranking evaluate_v2 4/4 (MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000).",
    ),
    (
        "El ranking evaluate_v2 4/4 sitúa a MAAC primero (0,9538).",
        "El ranking evaluate_v2 4/4 sitúa a MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000.",
    ),
]

REPLACEMENTS_EN = [
    (
        "evaluate_v2 4/4 ranks MAAC first at 0.9538).",
        "evaluate_v2 4/4 ranks MAAC 0.9538 > MATD3 0.8805 > MASAC 0.8679 > HAPPO 0.0000).",
    ),
    (
        "and evaluate_v2 4/4 (0.9538).",
        "and evaluate_v2 4/4 (MAAC 0.9538 > MATD3 0.8805 > MASAC 0.8679 > HAPPO 0.0000).",
    ),
]


def replace_in_paragraphs(doc: Document, pairs: list[tuple[str, str]], tag: str) -> list[dict]:
    changes: list[dict] = []
    for old, new in pairs:
        for i, p in enumerate(doc.paragraphs):
            text = p.text or ""
            if old in text:
                p.text = text.replace(old, new)
                changes.append(
                    {
                        "where": f"paragraph[{i}]",
                        "tag": tag,
                        "old": old,
                        "new": new,
                    }
                )
    return changes


def fix_corrupt_delta_cell(doc: Document) -> list[dict]:
    """Fix Informe table cell '23 07047' → Control−Baseline consistent value."""
    changes: list[dict] = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                raw = (cell.text or "").strip()
                if raw == "23 07047":
                    # Same row: Control − Baseline (Spanish thousands/decimals)
                    # Verified: 11.019.732,7202 − 10.973.717,9654 = 46.014,7548
                    new_val = "46.014,7548"
                    # Clear and set cell text
                    for paragraph in cell.paragraphs:
                        paragraph.text = ""
                    cell.paragraphs[0].text = new_val
                    changes.append(
                        {
                            "where": f"table[{ti}].row[{ri}].cell[{ci}]",
                            "tag": "corrupt_delta_matd3_e2",
                            "old": raw,
                            "new": new_val,
                            "rationale": (
                                "Celda corrupta; Δ coherente con Control−Baseline de la misma fila "
                                "(no confundir con ΔCO₂ distrito canónico 23 070 kg de ranking_oe)."
                            ),
                            "source": "misma fila tabla Informe (carbon_emissions_total MATD3 E2)",
                        }
                    )
    return changes


def backup(path: Path) -> Path:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    dest = BACKUP_DIR / f"{path.name}.pre_integral_validation_20260729.bak"
    shutil.copy2(path, dest)
    return dest


def main() -> int:
    report = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "patches": {},
        "backups": {},
    }

    for label, path in (("tesis", TESIS), ("informe", INFORME)):
        if not path.is_file():
            raise FileNotFoundError(path)
        bak = backup(path)
        report["backups"][label] = str(bak.relative_to(REPO)).replace("\\", "/")
        doc = Document(str(path))
        changes: list[dict] = []
        changes.extend(replace_in_paragraphs(doc, REPLACEMENTS_ES, "eval_v2_ranking_es"))
        changes.extend(replace_in_paragraphs(doc, REPLACEMENTS_EN, "eval_v2_ranking_en"))
        if label == "informe":
            changes.extend(fix_corrupt_delta_cell(doc))
        doc.save(str(path))
        report["patches"][label] = {
            "file": str(path.relative_to(REPO)).replace("\\", "/"),
            "n_changes": len(changes),
            "changes": changes,
        }
        print(f"{label}: {len(changes)} cambios -> {path.name}")
        for c in changes:
            print(f"  - {c['where']}: {c['old'][:60]!r} -> {c['new'][:80]!r}")

    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(f"Report: {REPORT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
