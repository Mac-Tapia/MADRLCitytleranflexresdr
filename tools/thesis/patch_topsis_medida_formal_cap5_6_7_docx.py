#!/usr/bin/env python3
"""Integra TOPSIS como medida formal adicional en Caps V–VII (2 Word canónicos).

Cambio de postura: TOPSIS deja de ser solo capa «ilustrativa/descriptiva que no decide»
y pasa a ser medida multicriterio adicional para sustentar/contrastar OG, OE y HE
junto a evaluate_v2, métricas OE y el diseño 4×3 — sin inventar scores.

GT: outputs/madrl_multicriteria_selection/{selection_report.json,topsis_ranking.csv}
C* canónicos: MAAC 0,9827 > MASAC 0,5656 > MATD3 0,3074
evaluate_v2: MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000

Backups: outputs/_word_backups/*.pre_topsis_formal_*.bak
Reportes: docs/TOPSIS_MEDIDA_OBJETIVOS_HE_CAP5_6_7_2026-07-29.{md,json}
"""
from __future__ import annotations

import json
import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import INFORME, TESIS  # noqa: E402

BACKUP_DIR = REPO / "outputs" / "_word_backups"
STAMP = datetime.now().strftime("%Y%m%d_%H%M%S")
MC_DIR = REPO / "outputs" / "madrl_multicriteria_selection"
REPORT_JSON = REPO / "docs" / "TOPSIS_MEDIDA_OBJETIVOS_HE_CAP5_6_7_2026-07-29.json"
REPORT_MD = REPO / "docs" / "TOPSIS_MEDIDA_OBJETIVOS_HE_CAP5_6_7_2026-07-29.md"
MARKER = "medida multicriterio formal adicional"

# GT from topsis_ranking.csv / selection_report.json (real_drive_50ep_c1c6)
TOPSIS_GT = {
    "source": "real_drive_50ep_c1c6",
    "path": "outputs/madrl_multicriteria_selection/topsis_ranking.csv",
    "report": "outputs/madrl_multicriteria_selection/selection_report.json",
    "ranking": [
        {"rank": 1, "algorithm": "MAAC", "closeness": 0.9826857601271446},
        {"rank": 2, "algorithm": "MASAC", "closeness": 0.5656251417928156},
        {"rank": 3, "algorithm": "MATD3", "closeness": 0.30739015628706934},
    ],
    "c_star_fmt": {
        "MAAC": "0,9827",
        "MASAC": "0,5656",
        "MATD3": "0,3074",
    },
    "c_star_dot": {
        "MAAC": "0.9827",
        "MASAC": "0.5656",
        "MATD3": "0.3074",
    },
}
EVAL_V2 = {
    "MAAC": "0,9538",
    "MATD3": "0,8805",
    "MASAC": "0,8679",
    "HAPPO": "0,0000",
}

# Ordered substring replacements (old → new). Longer / more specific first.
TEXT_REPLACEMENTS: list[tuple[str, str]] = [
    # Resumen ES / EN — C* + rol
    (
        "TOPSIS/AHP (outputs/madrl_multicriteria_selection) es descriptivo (MAAC C* ≈ 0,7828) y no sustituye las HE.",
        "TOPSIS/AHP (outputs/madrl_multicriteria_selection; fuente real_drive_50ep_c1c6) es "
        "medida multicriterio formal adicional (MAAC C* ≈ 0,9827 > MASAC 0,5656 > MATD3 0,3074) "
        "que integra costo–CO₂–flexibilidad para apoyar la determinación de OG/OE.1–OE.3 y el "
        "contraste multiobjetivo de HE, complementando —sin sustituir— el omnibus KPI-gains y evaluate_v2 4/4.",
    ),
    (
        "TOPSIS/AHP (MAAC C*≈0.7828) is descriptive and does not decide hypotheses.",
        "TOPSIS/AHP (real_drive_50ep_c1c6; MAAC C*≈0.9827 > MASAC 0.5656 > MATD3 0.3074) is an "
        "additional formal multicriteria measure supporting OG/OE determination and multiobjective "
        "hypothesis contrast, complementing —not replacing— KPI-gains omnibus tests and evaluate_v2 4/4.",
    ),
    (
        "TOPSIS (C* MAAC ≈ 0,7828) es selección descriptiva multicriterio y no decide hipótesis.",
        "TOPSIS (outputs/madrl_multicriteria_selection; C* MAAC ≈ 0,9827 > MASAC 0,5656 > MATD3 0,3074) "
        "es medida multicriterio formal adicional que integra los tres pilares para apoyar la "
        "determinación de OG/OE y el contraste multiobjetivo de HE, junto a evaluate_v2 y KPI-gains.",
    ),
    (
        "TOPSIS (MAAC C*≈0.7828) is descriptive multicriteria selection, not hypothesis evidence.",
        "TOPSIS (real_drive_50ep_c1c6; MAAC C*≈0.9827 > MASAC 0.5656 > MATD3 0.3074) is an additional "
        "formal multicriteria measure supporting OG/OE/HE multiobjective judgment alongside "
        "evaluate_v2 and KPI-gains (it does not replace the omnibus).",
    ),
    (
        "TOPSIS con pesos iguales identificó a MAAC (0,7828) como líder descriptivo global;",
        "TOPSIS multicriterio (madrl_multicriteria_selection; C* MAAC ≈ 0,9827) identifica a MAAC "
        "como líder multiobjetivo formal adicional;",
    ),
    (
        "Equal-weight TOPSIS ranked MAAC first globally (0.7828),",
        "Multicriteria TOPSIS (madrl_multicriteria_selection) ranked MAAC first (C*≈0.9827),",
    ),
    (
        "Las capas episódicas/TOPSIS son complementarias y no sustituyen el veredicto HE de KPI-gains.",
        "Las capas episódicas, evaluate_v2 4/4 y TOPSIS (medida multicriterio formal adicional; "
        "MAAC C* ≈ 0,9827) complementan el veredicto HE de KPI-gains: convergen con evaluate_v2 "
        "en coronar a MAAC (0,9538), y divergen del orden 2.º/3.º (TOPSIS: MASAC > MATD3; "
        "evaluate_v2: MATD3 > MASAC), sin sustituir el omnibus.",
    ),
    (
        "TOPSIS y ranking evaluate_v2 4/4 = estadistica descriptiva; HE11/HE21/HE31 = KPI-gains (impacto significativo y diferencias) + C3–C5.",
        "TOPSIS (medida multicriterio formal adicional; C* MAAC ≈ 0,9827) y ranking evaluate_v2 4/4 "
        "(MAAC 0,9538) apoyan la determinación multiobjetivo de OG/OE; HE11/HE21/HE31 se contrastan "
        "con KPI-gains (impacto significativo y diferencias) + C3–C5, complementados por TOPSIS/evaluate_v2.",
    ),
    (
        "TOPSIS/4/4 no sustituyen C1–C5.",
        "TOPSIS y evaluate_v2 4/4 son medidas formales adicionales que refuerzan el juicio de OG/OE/HE "
        "junto a C1–C5, sin reemplazar el conjunto de criterios de impacto.",
    ),
    (
        "TOPSIS y evaluate_v2 4/4 son descriptivos; no deciden HE.",
        "TOPSIS (C* MAAC ≈ 0,9827) y evaluate_v2 4/4 (MAAC 0,9538) son medidas formales adicionales "
        "para la determinación multiobjetivo de OG/OE y el contraste de HE; complementan KPI-gains "
        "sin sustituir el omnibus.",
    ),
    (
        "No usa TOPSIS ni ranking 4/4 como prueba de hipotesis. Decision formal en §5.5.",
        "El omnibus de hipótesis usa KPI-gains (§5.5); TOPSIS y evaluate_v2 4/4 aportan la medida "
        "multicriterio/ranking adicional para contrastar OG/OE/HE en clave multiobjetivo.",
    ),
    (
        "baseline v2, TOPSIS descriptivo y control de recursos (C5).",
        "baseline v2, TOPSIS como medida multicriterio formal adicional y control de recursos (C5).",
    ),
    (
        "Complementos del OG (descriptivos de soporte; no deciden H0G/H1G): ranking global Drive, best/worst por escenario, KPIs multiobjetivo de distrito y TOPSIS/AHP (estadistica descriptiva multicriterio sobre 50 ep Drive).",
        "Complementos del OG (medidas de soporte multiobjetivo para H0G/H1G/OE): ranking global Drive, "
        "best/worst por escenario, KPIs multiobjetivo de distrito y TOPSIS/AHP como medida multicriterio "
        "formal adicional (C* MAAC ≈ 0,9827; 50 ep Drive, real_drive_50ep_c1c6).",
    ),
    (
        "Tabla 5.4.1. TOPSIS descriptivo (madrl_multicriteria_selection; no evidencia de HE).",
        "Tabla 5.4.1. TOPSIS — medida multicriterio formal adicional "
        "(madrl_multicriteria_selection; C* canónicos; complementa OG/OE/HE junto a evaluate_v2).",
    ),
    (
        "Naturaleza descriptiva; no sustituye §5.3 ni §5.5 (HE/H0G).",
        "Incluye TOPSIS como medida multicriterio formal adicional para OG/OE/HE; "
        "complementa §5.3 y §5.5 (KPI-gains) sin sustituir el omnibus.",
    ),
    (
        "(5) TOPSIS/4/4/best_madrl son descriptivos y no respaldan HE por si solos.",
        "(5) TOPSIS/4/4/best_madrl son medidas formales adicionales que refuerzan el juicio "
        "multiobjetivo de OG/OE/HE junto a KPI-gains (sin sustituir el omnibus).",
    ),
    (
        "3) Trade-off: MATD3 flex+CO2; MAAC costos/TOPSIS/4/4 (descriptivo). 4) H1G exploratoria no equivale a HE11/HE21/HE31. 5) TOPSIS/4/4 = descriptivo de 50 ep Drive; HE = KPI-gains (impacto + diferencias).",
        "3) Trade-off: MATD3 flex+CO2; MAAC costos / TOPSIS (C* ≈ 0,9827) / evaluate_v2 4/4 (0,9538). "
        "4) H1G exploratoria no equivale a HE11/HE21/HE31. "
        "5) TOPSIS es medida multicriterio formal adicional (50 ep Drive) que apoya OG/OE/HE; "
        "el omnibus HE permanece en KPI-gains (impacto + diferencias).",
    ),
    (
        "TOPSIS/AHP en outputs/madrl_multicriteria_selection es descriptivo (MAAC C* ≈ 0,7828) y no decide hipótesis.",
        "TOPSIS/AHP en outputs/madrl_multicriteria_selection es medida multicriterio formal adicional "
        "(MAAC C* ≈ 0,9827 > MASAC 0,5656 > MATD3 0,3074) que apoya la determinación de OG/OE y el "
        "contraste multiobjetivo de HE, en convergencia con evaluate_v2 4/4 (ambos coronan MAAC) "
        "y con tensión en el orden 2.º/3.º respecto de MATD3/MASAC.",
    ),
    (
        "Las decisiones de HE usan la capa KPI-gains (Cap. 5); TOPSIS/4/4 y medias episódicas son descriptivas.",
        "Las decisiones omnibus de HE usan KPI-gains (Cap. 5); TOPSIS (C* MAAC ≈ 0,9827) y evaluate_v2 4/4 "
        "son medidas formales adicionales que refuerzan/matizan el cumplimiento multiobjetivo de OG/OE/HE.",
    ),
    (
        "MAAC gana costos/TOPSIS/4/4 descriptivo; KPI-gains E3 sin omnibus ni impacto vs cero tras Holm",
        "KPI-gains E3 sin omnibus ni impacto vs cero tras Holm; TOPSIS (C*≈0,9827) y evaluate_v2 "
        "(0,9538) coronan MAAC en multiobjetivo/costos como medida adicional, sin sustituir el omnibus HE31",
    ),
    (
        "Kruskal–Wallis E3 p = 0,7357 (KPI-gains); TOPSIS/4/4 no sustituyen HE31",
        "Kruskal–Wallis E3 p = 0,7357 (KPI-gains); TOPSIS/evaluate_v2 refuerzan el juicio multiobjetivo "
        "hacia MAAC/OE.3 sin sustituir el omnibus de HE31",
    ),
    # Stale C* leftovers (after specific phrases)
    ("MAAC C* ≈ 0,7828", "MAAC C* ≈ 0,9827"),
    ("MAAC C*≈0.7828", "MAAC C*≈0.9827"),
    ("C* MAAC ≈ 0,7828", "C* MAAC ≈ 0,9827"),
    ("(0,7828)", "(0,9827)"),
    ("(0.7828)", "(0.9827)"),
]

# Paragraphs to insert if missing (anchor substring → new text).
INSERT_AFTER: list[tuple[str, str]] = [
    (
        "5.6 Discusion de resultados",
        "En la discusión multiobjetivo, TOPSIS actúa como medida multicriterio formal adicional "
        "sobre los tres pilares (costo, CO₂, flexibilidad) derivados de la campaña "
        "madrl_v3_20260627_164047 (outputs/madrl_multicriteria_selection/; fuente real_drive_50ep_c1c6). "
        "El ranking TOPSIS (MAAC C* ≈ 0,9827 > MASAC 0,5656 > MATD3 0,3074) converge con evaluate_v2 4/4 "
        "en identificar a MAAC como mejor compromiso global (0,9538 > 0,8805 > 0,8679 > 0,0000) y "
        "diverge en el orden intermedio (TOPSIS: MASAC delante de MATD3; evaluate_v2: MATD3 delante de MASAC). "
        "Esta convergencia/tensión informa la contrastación de OG/OE.1–OE.3 y de HE en Caps. V–VII: "
        "refuerza el liderazgo multiobjetivo de MAAC hacia OE.3/OG sin anular el veredicto omnibus de "
        "KPI-gains (HE10–HE31 no respaldadas; H0G solo exploratoria).",
    ),
    (
        "6.1 Principales hallazgos",
        "Como medida multicriterio formal adicional, TOPSIS (C* MAAC ≈ 0,9827; MASAC 0,5656; MATD3 0,3074) "
        "refuerza las conclusiones sobre OG/OE: el cumplimiento descriptivo por eje (MATD3 en OE.1/OE.2; "
        "MAAC en OE.3) se articula con un juicio integrado costo–CO₂–flexibilidad que coincide con "
        "evaluate_v2 en coronar a MAAC y matiza el score best_madrl 3×3 (MATD3 = 0,6667). "
        "Respecto de las hipótesis, TOPSIS no reemplaza KPI-gains, pero sí aporta evidencia multicriterio "
        "adicional para determinar/sustentar el alcance del cumplimiento de objetivos y el contraste de HE.",
    ),
    (
        "6.1 Conclusiones por objetivo",
        "Como medida multicriterio formal adicional, TOPSIS (C* MAAC ≈ 0,9827; MASAC 0,5656; MATD3 0,3074; "
        "outputs/madrl_multicriteria_selection/) refuerza las conclusiones sobre OG/OE.1–OE.3: integra "
        "los tres pilares y converge con evaluate_v2 4/4 (MAAC 0,9538) en el liderazgo multiobjetivo de MAAC, "
        "mientras el best_madrl 3×3 favorece a MATD3 (0,6667) por OE.1/OE.2. Esta capa matiza —sin sustituir— "
        "el contraste de HE basado en KPI-gains.",
    ),
]


def set_run_font(run, *, bold: bool = False, size: float = 12.0) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def replace_paragraph_text(p: Paragraph, text: str) -> None:
    p.clear()
    run = p.add_run(text)
    set_run_font(run)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def backup(path: Path) -> Path:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    dest = BACKUP_DIR / f"{path.name}.pre_topsis_formal_{STAMP}.bak"
    shutil.copy2(path, dest)
    return dest


def load_gt_verify() -> dict:
    ranking_path = MC_DIR / "topsis_ranking.csv"
    report_path = MC_DIR / "selection_report.json"
    if not ranking_path.is_file() or not report_path.is_file():
        raise FileNotFoundError("Falta GT multicriterio en outputs/madrl_multicriteria_selection/")
    report = json.loads(report_path.read_text(encoding="utf-8"))
    rows = []
    with ranking_path.open(encoding="utf-8") as fh:
        header = fh.readline().strip().split(",")
        for line in fh:
            parts = line.strip().split(",")
            row = dict(zip(header, parts))
            rows.append(
                {
                    "rank": int(float(row["rank"])),
                    "algorithm": row["algorithm"],
                    "closeness": float(row["closeness"]),
                }
            )
    return {
        "source": report.get("source"),
        "rows": rows,
        "winner": (report.get("ranking_consistency") or {}).get("topsis_winner"),
    }


def patch_topsis_table(doc: Document) -> list[dict]:
    changes = []
    target = {
        "MAAC": TOPSIS_GT["c_star_dot"]["MAAC"],
        "MASAC": TOPSIS_GT["c_star_dot"]["MASAC"],
        "MATD3": TOPSIS_GT["c_star_dot"]["MATD3"],
    }
    for ti, table in enumerate(doc.tables):
        rows = table.rows
        if len(rows) < 2:
            continue
        header = " | ".join(c.text.strip() for c in rows[0].cells)
        if "C*" not in header.upper() and "TOPSIS" not in header.upper():
            continue
        # Expect Rank | Algoritmo | C*
        for ri, row in enumerate(rows[1:], start=1):
            cells = row.cells
            if len(cells) < 3:
                continue
            algo = cells[1].text.strip()
            if algo not in target:
                continue
            old = cells[2].text.strip()
            new = target[algo]
            if old != new:
                # clear cell paragraphs
                cells[2].text = new
                changes.append(
                    {
                        "type": "table_cell",
                        "table": ti,
                        "row": ri,
                        "algorithm": algo,
                        "old": old,
                        "new": new,
                    }
                )
    return changes


def apply_text_replacements(doc: Document) -> list[dict]:
    changes = []
    for i, p in enumerate(doc.paragraphs):
        old = p.text or ""
        if not old.strip():
            continue
        new = old
        applied: list[str] = []
        for a, b in TEXT_REPLACEMENTS:
            if a in new:
                new = new.replace(a, b)
                applied.append(a[:80])
        if new != old:
            replace_paragraph_text(p, new)
            changes.append(
                {
                    "type": "paragraph_replace",
                    "index": i,
                    "n_patterns": len(applied),
                    "old_preview": old[:220],
                    "new_preview": new[:220],
                    "patterns": applied,
                }
            )
    return changes


def apply_insertions(doc: Document) -> list[dict]:
    changes = []
    blob = "\n".join(p.text or "" for p in doc.paragraphs)
    # Walk a snapshot of paragraphs; insert after matching headings
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        for anchor, text in INSERT_AFTER:
            if anchor not in t:
                continue
            # Already inserted nearby?
            window = "\n".join(
                (paras[j].text or "") for j in range(i, min(i + 4, len(paras)))
            )
            if MARKER in window and "0,9827" in window and "evaluate_v2" in window:
                continue
            if text[:60] in blob:
                continue
            insert_paragraph_after(p, text)
            changes.append(
                {
                    "type": "paragraph_insert",
                    "after_index": i,
                    "anchor": anchor,
                    "preview": text[:220],
                }
            )
            blob += "\n" + text
    return changes


def patch_he31_table_cells(doc: Document) -> list[dict]:
    """Rewrite remaining 'descriptivo' HE31 fundament cells."""
    changes = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                old = cell.text.strip()
                new = old
                for a, b in TEXT_REPLACEMENTS:
                    if a in new:
                        new = new.replace(a, b)
                # Soft fallbacks inside cells
                new = new.replace(
                    "TOPSIS/4/4 no sustituyen HE31",
                    "TOPSIS/evaluate_v2 refuerzan el juicio multiobjetivo hacia MAAC/OE.3 sin sustituir el omnibus de HE31",
                )
                new = new.replace(
                    "MAAC gana costos/TOPSIS/4/4 descriptivo;",
                    "TOPSIS/evaluate_v2 coronan MAAC en multiobjetivo/costos (medida adicional);",
                )
                if new != old:
                    cell.text = new
                    changes.append(
                        {
                            "type": "table_cell_text",
                            "table": ti,
                            "row": ri,
                            "col": ci,
                            "old_preview": old[:160],
                            "new_preview": new[:160],
                        }
                    )
    return changes


def post_checks(doc: Document) -> dict:
    blob = "\n".join(p.text or "" for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                blob += "\n" + (cell.text or "")
    return {
        "has_marker": MARKER in blob,
        "has_c9827": ("0,9827" in blob) or ("0.9827" in blob),
        "has_c5656": ("0,5656" in blob) or ("0.5656" in blob),
        "has_c3074": ("0,3074" in blob) or ("0.3074" in blob),
        "stale_07828": ("0,7828" in blob) or ("0.7828" in blob),
        "stale_descriptivo_topsis": bool(
            re.search(r"(?i)TOPSIS[^\n]{0,80}descriptiv|descriptiv[^\n]{0,80}TOPSIS", blob)
        ),
        "deny_role": bool(
            re.search(
                r"(?i)TOPSIS[^\n]{0,100}(no decide|no evidencia de HE|no respaldan HE por si solos|solo ilustrativ)",
                blob,
            )
        ),
        "has_oe_he_anchor": bool(
            re.search(r"(?i)TOPSIS[^\n]{0,120}(OG|OE\.|HE|hip[oó]tesis|objetivo)", blob)
        ),
        "has_eval_compare": "evaluate_v2" in blob and "0,9827" in blob.replace(".", ",").replace("0.9827", "0,9827")
        or ("evaluate_v2" in blob and ("0,9827" in blob or "0.9827" in blob)),
    }


def patch_doc(path: Path) -> dict:
    bak = backup(path)
    doc = Document(str(path))
    changes: list[dict] = []
    changes.extend(apply_text_replacements(doc))
    changes.extend(patch_topsis_table(doc))
    changes.extend(patch_he31_table_cells(doc))
    changes.extend(apply_insertions(doc))
    # Second pass for any leftover stale C* after inserts
    changes.extend(apply_text_replacements(doc))
    doc.save(str(path))
    # reload for checks
    doc2 = Document(str(path))
    checks = post_checks(doc2)
    return {
        "path": str(path.relative_to(REPO)).replace("\\", "/"),
        "backup": str(bak.relative_to(REPO)).replace("\\", "/"),
        "n_changes": len(changes),
        "changes": changes,
        "checks": checks,
        "ok": checks["has_marker"]
        and checks["has_c9827"]
        and not checks["stale_07828"]
        and not checks["deny_role"],
    }


def write_reports(gt_verify: dict, results: dict) -> None:
    payload = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "stamp": STAMP,
        "script": "tools/thesis/patch_topsis_medida_formal_cap5_6_7_docx.py",
        "posture": (
            "TOPSIS = medida multicriterio formal adicional para determinar/sustentar "
            "OG, OE.1–OE.3 y HE en Caps. V–VII; complementa evaluate_v2 y KPI-gains; "
            "no inventa scores."
        ),
        "ground_truth": {
            **TOPSIS_GT,
            "verified_source": gt_verify.get("source"),
            "verified_rows": gt_verify.get("rows"),
            "winner": gt_verify.get("winner"),
        },
        "evaluate_v2_ranking": EVAL_V2,
        "relation_topsis_vs_evaluate_v2": {
            "convergence": "Ambos coronan MAAC (#1).",
            "divergence": (
                "TOPSIS: MASAC (#2) > MATD3 (#3); evaluate_v2: MATD3 (#2) > MASAC (#3); "
                "best_madrl 3×3 favorece MATD3 (0,6667) por OE.1/OE.2."
            ),
        },
        "chapter_mapping": {
            "Cap_V": "Capítulo 5 (resultados + §5.4 TOPSIS + contrastación)",
            "Cap_VI": "§5.6 Discusión de resultados (rol de discusión; no hay Cap.6 discusión separado)",
            "Cap_VII": "Capítulo 6 Conclusiones (no existe Heading 1 «Capítulo 7» en los 2 Word)",
        },
        "docs": results,
        "ok": all(r.get("ok") for r in results.values()),
    }
    REPORT_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# TOPSIS como medida formal para OG/OE/HE (Caps. V–VII) — 2026-07-29",
        "",
        f"Generado: `{payload['generated_at']}`",
        f"Script: `{payload['script']}`",
        "",
        "## Veredicto",
        "",
        f"- **OK global:** `{payload['ok']}`",
        f"- Postura: {payload['posture']}",
        "",
        "## Ground truth TOPSIS",
        "",
        f"- Fuente: `{TOPSIS_GT['path']}` / `{TOPSIS_GT['report']}` (`{gt_verify.get('source')}`)",
        f"- Ranking: MAAC **0,9827** > MASAC **0,5656** > MATD3 **0,3074** (ganador={gt_verify.get('winner')})",
        f"- evaluate_v2 4/4: MAAC 0,9538 > MATD3 0,8805 > MASAC 0,8679 > HAPPO 0,0000",
        "",
        "## Relación TOPSIS ↔ evaluate_v2",
        "",
        f"- Convergencia: {payload['relation_topsis_vs_evaluate_v2']['convergence']}",
        f"- Divergencia: {payload['relation_topsis_vs_evaluate_v2']['divergence']}",
        "",
        "## Mapeo de capítulos",
        "",
    ]
    for k, v in payload["chapter_mapping"].items():
        lines.append(f"- **{k}:** {v}")
    lines += ["", "## Documentos", ""]
    for name, r in results.items():
        lines += [
            f"### {name}",
            "",
            f"- Path: `{r['path']}`",
            f"- Backup: `{r['backup']}`",
            f"- Cambios: {r['n_changes']}",
            f"- Checks: `{json.dumps(r['checks'], ensure_ascii=False)}`",
            f"- OK: `{r['ok']}`",
            "",
            "Disclaimers / párrafos tocados (preview):",
            "",
        ]
        for ch in r["changes"][:25]:
            if ch["type"] == "paragraph_replace":
                lines.append(f"- p[{ch['index']}] OLD: {ch['old_preview'][:160]}")
                lines.append(f"  NEW: {ch['new_preview'][:160]}")
            elif ch["type"] == "paragraph_insert":
                lines.append(f"- INSERT after [{ch['after_index']}] `{ch['anchor']}`: {ch['preview'][:160]}")
            elif ch["type"].startswith("table"):
                lines.append(f"- TABLE {ch}")
        lines.append("")
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    print("[0] Verificando GT TOPSIS…")
    gt_verify = load_gt_verify()
    print(f"    source={gt_verify['source']} winner={gt_verify['winner']} rows={gt_verify['rows']}")
    if gt_verify["source"] and "illustrative" in str(gt_verify["source"]).lower():
        print("    WARN: source illustrative — no inventar scores; usar valores del JSON tal cual.")
    # Verify closeness matches expected
    for row, exp in zip(gt_verify["rows"], TOPSIS_GT["ranking"]):
        if abs(row["closeness"] - exp["closeness"]) > 1e-9 or row["algorithm"] != exp["algorithm"]:
            raise SystemExit(f"GT mismatch: {row} vs {exp}")

    results = {}
    for path in (TESIS, INFORME):
        print(f"[patch] {path.name}")
        results[path.name] = patch_doc(path)
        print(f"    changes={results[path.name]['n_changes']} ok={results[path.name]['ok']}")
        print(f"    checks={results[path.name]['checks']}")
        print(f"    backup={results[path.name]['backup']}")

    write_reports(gt_verify, results)
    print(f"[report] {REPORT_MD.relative_to(REPO)}")
    print(f"[report] {REPORT_JSON.relative_to(REPO)}")
    return 0 if all(r["ok"] for r in results.values()) else 1


if __name__ == "__main__":
    raise SystemExit(main())
