"""Create a DOCX skeleton for the doctoral thesis report (6 chapters)."""

from __future__ import annotations

import argparse
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from docx import Document
except ModuleNotFoundError:  # pragma: no cover - exercised when python-docx is absent.
    Document = None


TITLE = (
    "MULTI-AGENTE DE APRENDIZAJE POR REFUERZO PROFUNDO PARA LA GESTIÓN "
    "COORDINADA DE FLEXIBILIDAD ENERGÉTICA, EMISIONES DE CARBONO Y COSTOS "
    "ENERGÉTICOS EN COMUNIDADES INTELIGENTES"
)

SECTIONS = [
    ("CARÁTULA — UNI Escuela de Posgrado — Doctorado en Ingeniería", 0),
    ("Caso de estudio: SEAI Iquitos — 17 edificios (2023–2025)", 1),
    ("Índice", 0),
    ("Resumen", 0),
    ("Abstract", 0),
    ("Capítulo 1. Introducción", 0),
    ("1.1 Planteamiento y formulación del problema (PG, PE.1–PE.3)", 1),
    ("1.2 Objetivos (OG, OE.1–OE.3)", 1),
    ("1.3 Hipótesis (HG, HE.1–HE.3)", 1),
    ("1.4 Matriz de consistencia y operacionalización (Tablas 1.1, 1.2)", 1),
    ("1.5 Justificación", 1),
    ("1.6 Alcances y limitaciones", 1),
    ("Capítulo 2. Marco teórico", 0),
    ("2.1 Antecedentes de la investigación", 1),
    ("2.1.1 Flexibilidad energética con MADRL (D-VD.1)", 2),
    ("2.1.2 Emisiones de carbono con MADRL (D-VD.2)", 2),
    ("2.1.3 Costos energéticos con MADRL (D-VD.3)", 2),
    ("2.1.4 Marco técnico MADRL y sistemas aislados", 2),
    ("2.2 Bases teóricas", 1),
    ("2.3 Definición de términos", 1),
    ("Capítulo 3. Metodología", 0),
    ("3.1 Tipo, enfoque y nivel de investigación", 1),
    ("3.2 Diseño experimental factorial 4×3 (12 tratamientos)", 1),
    ("3.3 Unidad de análisis, población y muestra", 1),
    ("3.4 Datos: dataset citylearn_iquitos_2023_2025", 1),
    ("3.5 Variables y operacionalización (54 KPI)", 1),
    ("3.6 Técnicas e instrumentos de recolección", 1),
    ("3.7 Técnicas de análisis estadístico", 1),
    ("Capítulo 4. Desarrollo de la propuesta", 0),
    ("4.1 Arquitectura del sistema experimental", 1),
    ("4.2 Formulación Dec-POMDP", 1),
    ("4.3 Esquema CTDE y recompensa unified_comparable_v4", 1),
    ("4.4 Algoritmos e hiperparámetros", 1),
    ("4.5 Aportes originales al motor (A1–A4)", 1),
    ("4.6 Implementación y entorno computacional", 1),
    ("Capítulo 5. Resultados y contrastación de hipótesis", 0),
    ("5.1 Experimentos realizados", 1),
    ("5.2 Análisis descriptivo del efecto sobre la VD", 1),
    ("5.3 Efecto coordinado: ranking ponderado por escenario", 1),
    ("5.4 Contrastación inferencial de las hipótesis", 1),
    ("5.5 Discusión de resultados", 1),
    ("Capítulo 6. Conclusiones y trabajo futuro", 0),
    ("6.1 Conclusiones", 1),
    ("6.2 Limitaciones", 1),
    ("6.3 Trabajo futuro", 1),
    ("6.4 Cronograma de culminación", 1),
    ("Referencias bibliográficas (APA 7.ª ed.)", 0),
    ("ANEXOS", 0),
    ("Anexo A. Matriz bibliográfica de 50 investigaciones", 1),
    ("Anexo B. Pipeline dataset Iquitos (10 etapas)", 1),
    ("Anexo C. Tabla 17 edificios SEAI", 1),
    ("Anexo D. Configuración hiperparámetros y recompensa v4", 1),
    ("Anexo E. Resultados preliminares 5 ep / pendiente 50 ep", 1),
]


def _xml_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _paragraph_xml(text: str, *, heading: bool = False) -> str:
    style = '<w:pStyle w:val="Heading1"/>' if heading else ""
    return (
        "<w:p><w:pPr>"
        f"{style}"
        "</w:pPr><w:r><w:t>"
        f"{_xml_escape(text)}"
        "</w:t></w:r></w:p>"
    )


def _build_minimal_docx(output: Path) -> None:
    body = [_paragraph_xml(TITLE, heading=True)]
    body.append(_paragraph_xml("Documento base generado para tesis doctoral UNI — 6 capítulos (docs/Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx)."))
    body.append(_paragraph_xml("Nota: completar solo con fuentes verificadas, citas APA vigentes y resultados reales observados; los resultados incompletos deben marcarse como pendientes o no verificados."))

    for heading, _level in SECTIONS:
        body.append(_paragraph_xml(heading, heading=True))
        if heading == "CARÁTULA":
            body.extend([
                _paragraph_xml("Universidad: [pendiente]"),
                _paragraph_xml("Escuela de posgrado: [pendiente]"),
                _paragraph_xml("Tesista: [pendiente]"),
                _paragraph_xml("Asesor: [pendiente]"),
                _paragraph_xml("Lima, Perú - 2026"),
            ])
        elif heading == "REFERENCIAS":
            body.append(_paragraph_xml("Insertar únicamente referencias citadas en el texto, en formato APA vigente."))
        else:
            body.append(_paragraph_xml("[Redactar con evidencia del Módulo A, citas APA y trazabilidad de fuente.]"))

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + "".join(body)
        + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>'
        + "</w:body></w:document>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        "</Relationships>"
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(output, "w", ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)


def build(output: Path) -> None:
    if Document is None:
        _build_minimal_docx(output)
        return

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph("Documento base generado para tesis doctoral UNI — 6 capítulos (docs/Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx).")
    doc.add_paragraph("Nota: completar solo con fuentes verificadas, citas APA vigentes y resultados reales observados; los resultados incompletos deben marcarse como pendientes o no verificados.")

    for heading, level in SECTIONS:
        doc.add_heading(heading, level=level if level else 1)
        if heading == "CARÁTULA":
            doc.add_paragraph("Universidad: [pendiente]")
            doc.add_paragraph("Escuela de posgrado: [pendiente]")
            doc.add_paragraph("Tesista: [pendiente]")
            doc.add_paragraph("Asesor: [pendiente]")
            doc.add_paragraph("Lima, Perú - 2026")
        elif heading == "REFERENCIAS":
            doc.add_paragraph("Insertar únicamente referencias citadas en el texto, en formato APA vigente.")
        else:
            doc.add_paragraph("[Redactar con evidencia del Módulo A, citas APA y trazabilidad de fuente.]")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", required=True, help="Path to the .docx file.")
    args = parser.parse_args()
    build(Path(args.output))
    print(Path(args.output).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
