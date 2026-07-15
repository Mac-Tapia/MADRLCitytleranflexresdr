# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

p = Path(r"D:\MADRLCitytleranflexresdr\docs\ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx")
doc = Document(str(p))
lines = []
for i, para in enumerate(doc.paragraphs):
    t = (para.text or "").strip()
    if 640 <= i <= 655 or i in (419, 422, 574, 576, 577, 815, 817, 818, 788, 789, 790):
        lines.append(f"P{i}|{(para.style.name if para.style else '')}|{t[:320]}")

body = list(doc.element.body)
ti = 0
for ci, child in enumerate(body):
    if child.tag != qn("w:tbl"):
        continue
    prev = ""
    for j in range(ci - 1, -1, -1):
        if body[j].tag == qn("w:p"):
            prev = "".join(body[j].itertext()).strip()
            break
    if prev.startswith("Tabla 6.1") or prev.startswith("Tabla 6.2") or "Tabla A.2" in prev:
        tbl = doc.tables[ti]
        lines.append(f"TABLE|{prev[:100]}|rows={len(tbl.rows)}|cols={len(tbl.columns)}")
        for r, row in enumerate(tbl.rows[:9]):
            cells = [c.text.replace("\n", " ")[:60] for c in row.cells]
            lines.append("  R" + str(r) + "|" + " || ".join(cells))
    ti += 1

Path(r"D:\MADRLCitytleranflexresdr\outputs\_tmp_post_patch_status.txt").write_text("\n".join(lines), encoding="utf-8")
print("ok", len(lines))
