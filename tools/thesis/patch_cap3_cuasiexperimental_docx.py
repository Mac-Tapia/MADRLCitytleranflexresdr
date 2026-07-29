#!/usr/bin/env python3
"""Sincroniza Cap. III / Resumen a diseño cuasiexperimental en el Word canónico.

Fuente de verdad previa:
- docs/tesis_capitulos/Capitulo_3_Metodologia.md
- docs/tesis_capitulos/Referencias_APA.md
- tools/thesis/thesis_doctoral_sections.py (Resumen)
- tools/thesis/generate_borrador_tesis_docx.py (Cap. 3)

Este script no regenera toda la tesis: aplica un parche quirúrgico de coherencia
metodológica sobre el Word vigente y restaura la ruta canónica si falta.
"""
from __future__ import annotations

import json
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

import sys

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, TESIS, existing_canons  # noqa: E402

SRC_CANDIDATES = [TESIS]
OUT_CANON = TESIS
REPORT = DOCS / "CAP3_CUASIEXPERIMENTAL_PATCH_REPORT_2026-07-29.json"


def set_run_font(run, bold: bool = False, size: float = 12.0, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def replace_paragraph_text(p: Paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p.clear()
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def resolve_source() -> Path:
    for path in SRC_CANDIDATES:
        if path.is_file() and path.stat().st_size > 0:
            return path
    raise FileNotFoundError("No se encontró un .docx base para parchear Cap. III.")


def chapter_bounds(doc: Document, chapter: int) -> tuple[int | None, int | None]:
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if re.search(rf"^Cap[ií]tulo\s*{chapter}\b", t, re.I):
            start = i
        elif start is not None and re.search(rf"^Cap[ií]tulo\s*{chapter + 1}\b", t, re.I):
            end = i
            break
    return start, end


def patch_doc(doc: Document) -> list[str]:
    actions: list[str] = []

    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        n = norm(t)

        if n.startswith("esta tesis doctoral determina") and "factorial" in n:
            nt = t.replace("diseño experimental factorial", "diseño cuasiexperimental factorial")
            nt = nt.replace("diseno experimental factorial", "diseño cuasiexperimental factorial")
            if nt != t:
                replace_paragraph_text(p, nt)
                actions.append(f"resumen_body@{i}")

        if n.startswith("palabras clave:") and ("diseño experimental" in n or "diseno experimental" in n):
            nt = t.replace("diseño experimental", "diseño cuasiexperimental")
            nt = nt.replace("diseno experimental", "diseño cuasiexperimental")
            replace_paragraph_text(p, nt, italic=True)
            actions.append(f"resumen_keywords@{i}")

        if "this doctoral thesis" in n and "factorial design 4" in n and "quasi-experimental" not in n:
            nt = t.replace("factorial design 4×3", "quasi-experimental factorial design 4×3")
            nt = nt.replace("factorial design 4x3", "quasi-experimental factorial design 4×3")
            if nt != t:
                replace_paragraph_text(p, nt, italic=True)
                actions.append(f"abstract_body@{i}")

        if "metodologico: estudio cuantitativo, aplicado, comparativo, no experimental" in n or (
            "metodológico: estudio cuantitativo, aplicado, comparativo, no experimental" in n
        ):
            replace_paragraph_text(p, re.sub(r"no experimental", "cuasiexperimental", t, flags=re.I))
            actions.append(f"cap1_alcance@{i}")

        if n.startswith("el estudio es cuantitativo, aplicado y explicativo, basado en simulacion experimental") or n.startswith(
            "el estudio es cuantitativo, aplicado y explicativo, basado en simulación experimental"
        ):
            replace_paragraph_text(
                p,
                "El estudio es cuantitativo, aplicado y cuasiexperimental factorial 4×3, basado en simulación. "
                "A diferencia de los objetivos, las hipótesis formulan contrastes estadísticos sobre el "
                "factor algoritmo (VI) frente a la hipótesis nula de igualdad de distribuciones de KPI-gains. "
                "El protocolo inferencial (alpha = 0,05) se detalla en el Capítulo 3 y se resuelve en la "
                "sección 5.9 (Colas et al., 2019; Agarwal et al., 2021; Hernández-Sampieri & Mendoza, 2018):",
            )
            actions.append(f"cap1_hip_note@{i}")

        if n.startswith("diseno: no experimental") or n.startswith("diseño: no experimental"):
            replace_paragraph_text(
                p,
                "Diseño: cuasiexperimental, factorial 4×3 (algoritmo MADRL × escenario E1/E2/E3), "
                "basado en simulación computacional (Campbell & Stanley, 1963; "
                "Hernández-Sampieri & Mendoza, 2018).",
            )
            actions.append(f"cap3_design_bullet@{i}")

        if n.startswith("enfoque: cuantitativo.") and len(n) < 40:
            replace_paragraph_text(p, "Enfoque: cuantitativo (Hernández-Sampieri & Mendoza, 2018).")
            actions.append(f"cap3_enfoque@{i}")

        if n.startswith("tipo: aplicada.") and len(n) < 30:
            replace_paragraph_text(p, "Tipo: aplicada (Arias, 2020; Tamayo y Tamayo, 2004).")
            actions.append(f"cap3_tipo@{i}")

        if "3.2 diseno experimental" in n or "3.2 diseño experimental" in n:
            replace_paragraph_text(p, "3.2 Diseño cuasiexperimental factorial 4×3", bold=True)
            actions.append(f"cap3_heading32@{i}")

        if "experimental-computacional" in n or "experimental computacional factorial" in n:
            nt = re.sub(r"experimental[- ]computacional", "cuasiexperimental", t, flags=re.I)
            nt = nt.replace("Diseño experimental", "Diseño cuasiexperimental")
            nt = nt.replace("diseno experimental", "diseño cuasiexperimental")
            replace_paragraph_text(p, nt, bold=bool(re.search(r"^3\.2\b", t)))
            actions.append(f"cap3_expcomp@{i}")

        if "no se considera no experimental" in n or (
            "no experimental" in n and "manipul" in n and "algoritmo" in n
        ):
            replace_paragraph_text(
                p,
                "El diseño no se considera no experimental, porque existe manipulación controlada de la VI "
                "(algoritmo MADRL y escenario de recompensa E1/E2/E3) bajo protocolo fijo. Tampoco es un "
                "experimento puro de campo: no hay aleatorización de unidades naturales ni sujetos humanos. "
                "Por ello se adopta un diseño cuasiexperimental factorial 4×3 basado en simulación "
                "(Campbell & Stanley, 1963; Hernández-Sampieri & Mendoza, 2018), coherente con el Resumen "
                "(HAPPO, MASAC, MATD3, MAAC; CityLearn; SEAI Iquitos; OE.1 flexibilidad, OE.2 CO2, OE.3 costos).",
            )
            actions.append(f"cap3_quasi_explain@{i}")

    cap3_idx, cap4_idx = chapter_bounds(doc, 3)
    if cap3_idx is not None and cap4_idx is not None:
        cap3_text = "\n".join((p.text or "") for p in doc.paragraphs[cap3_idx:cap4_idx])
        if "cuasiexperimental" not in cap3_text.lower():
            h31 = next(
                (
                    i
                    for i in range(cap3_idx, cap4_idx)
                    if re.search(r"^3\.1\b", (doc.paragraphs[i].text or "").strip())
                ),
                None,
            )
            if h31 is not None:
                anchor = h31
                for j in range(h31 + 1, min(h31 + 12, cap4_idx)):
                    tj = (doc.paragraphs[j].text or "").strip()
                    if re.search(r"^3\.2\b", tj):
                        break
                    if tj:
                        anchor = j
                insert_paragraph_after(
                    doc.paragraphs[anchor],
                    "En coherencia con el Resumen, el diseño de investigación es cuasiexperimental factorial 4×3 "
                    "(4 algoritmos MADRL —HAPPO, MASAC, MATD3, MAAC— × 3 escenarios E1/E2/E3 alineados a OE.1 "
                    "flexibilidad, OE.2 CO2 y OE.3 costos) sobre CityLearn en el SEAI Iquitos. Se manipula la VI "
                    "bajo protocolo fijo, sin aleatorización de unidades naturales (Campbell & Stanley, 1963; "
                    "Hernández-Sampieri & Mendoza, 2018; Nweye et al., 2024).",
                )
                actions.append(f"cap3_insert_cuasi_after@{anchor}")

        if "Agarwal" not in cap3_text:
            for i in range(cap3_idx, cap4_idx):
                t = doc.paragraphs[i].text or ""
                if "Kruskal" in t and ("Mann" in t or "Wilcoxon" in t or "no param" in t.lower()):
                    insert_paragraph_after(
                        doc.paragraphs[i],
                        "El uso de pruebas no paramétricas y el reporte de tamaño de efecto siguen "
                        "recomendaciones para la comparación robusta de algoritmos de aprendizaje por "
                        "refuerzo (Demšar, 2006; Agarwal et al., 2021).",
                    )
                    actions.append(f"cap3_stats_cite@{i}")
                    break

    return actions


def verify(path: Path) -> dict:
    doc = Document(str(path))
    paras = [p.text or "" for p in doc.paragraphs]
    resumen = next((t for t in paras if t.strip().lower().startswith("esta tesis doctoral determina")), "")
    c3, c4 = chapter_bounds(doc, 3)
    cap3 = "\n".join(paras[c3:c4]) if c3 is not None and c4 is not None else ""
    low = "\n".join(paras).lower()
    return {
        "resumen_cuasiexperimental": "cuasiexperimental" in resumen.lower(),
        "cap3_cuasiexperimental": "cuasiexperimental" in cap3.lower(),
        "cap3_algorithms": all(a in cap3 for a in ("HAPPO", "MASAC", "MATD3", "MAAC")),
        "cap3_oe_terms": any(x in cap3 for x in ("OE.1", "flexibilidad"))
        and any(x in cap3 for x in ("OE.2", "CO2", "CO₂"))
        and any(x in cap3 for x in ("OE.3", "costo", "costos")),
        "doc_still_has_no_experimental": "no experimental" in low,
        "canonical_exists": path.is_file(),
    }


def main() -> int:
    src = resolve_source()
    if src.resolve() != OUT_CANON.resolve():
        shutil.copy2(src, OUT_CANON)

    targets = existing_canons() or [OUT_CANON]
    ordered: list[Path] = []
    for p in [OUT_CANON, *targets]:
        rp = p.resolve()
        if rp not in {x.resolve() for x in ordered} and p.is_file():
            ordered.append(p)

    per_file: list[dict] = []
    overall_ok = True
    for path in ordered:
        doc = Document(str(path))
        actions = patch_doc(doc)
        try:
            doc.save(str(path))
            saved = path
        except (PermissionError, OSError) as exc:
            raise SystemExit(
                f"No se pudo guardar el Word canónico {path.name} (¿abierto en Word?): {exc}"
            ) from exc
        checks = verify(saved)
        ok = bool(checks.get("cap3_cuasiexperimental") and checks.get("resumen_cuasiexperimental"))
        overall_ok = overall_ok and ok
        per_file.append(
            {
                "path": str(saved.relative_to(REPO)),
                "actions": actions,
                "checks": checks,
                "ok": ok,
            }
        )

    report = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "source": str(src.relative_to(REPO)),
        "canons": [str(p.relative_to(REPO)) for p in ordered],
        "files": per_file,
        "ok": overall_ok,
    }
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if overall_ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
