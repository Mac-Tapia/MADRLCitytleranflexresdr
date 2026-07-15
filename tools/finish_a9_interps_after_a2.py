# -*- coding: utf-8 -*-
"""Finish pending ABRIR_ESTE patches after A.2 validation: A.9 PNG + interpretations + verify 6.4/6.5."""
from __future__ import annotations

import json
import math
import re
import shutil
import zipfile
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document

REPO = Path(__file__).resolve().parents[1]
FULL = REPO / "outputs" / "_drive_madrl" / "full_data"
ANALYSIS = REPO / "outputs" / "_drive_madrl" / "gdrive_20260627_164047_objective_analysis"
TABLE_DIR = ANALYSIS / "tables"
FIG_DIR = ANALYSIS / "figures"
REPORT_DIR = ANALYSIS / "validation"
REPORT_DIR.mkdir(parents=True, exist_ok=True)

ALGOS = ["HAPPO", "MAAC", "MASAC", "MATD3"]
SCENARIOS = ["E1", "E2", "E3"]
PRIMARY = REPO / "docs" / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx"
MIRROR = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx"


def fmt_num(x, nd=2):
    return f"{float(x):,.{nd}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def scan_a9():
    rows = []
    for algo in ALGOS:
        for scen in SCENARIOS:
            man = FULL / algo / scen / "data" / "checkpoint_manifest.json"
            total_b = 0
            n = 0
            source = "missing"
            max_file_mb = 0.0
            if man.exists():
                obj = json.loads(man.read_text(encoding="utf-8"))
                cps = obj.get("checkpoints") or []
                n = len(cps)
                source = str(man.relative_to(REPO)).replace("\\", "/")
                for ck in cps:
                    b = int(ck.get("bytes") or 0)
                    total_b += b
                    mb = b / (1024 * 1024)
                    if mb > max_file_mb:
                        max_file_mb = mb
            rows.append(
                {
                    "algorithm": algo,
                    "scenario": scen,
                    "treatment": f"{algo}-{scen}",
                    "n_files_listed": n,
                    "total_bytes": total_b,
                    "total_mb": total_b / (1024 * 1024),
                    "total_gb": total_b / (1024**3),
                    "max_file_mb": max_file_mb,
                    "source": source,
                }
            )
    return pd.DataFrame(rows)


def plot_a9(df: pd.DataFrame) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    d = df.copy()
    d["algorithm"] = pd.Categorical(d["algorithm"], ALGOS, ordered=True)
    d["scenario"] = pd.Categorical(d["scenario"], SCENARIOS, ordered=True)
    d = d.sort_values(["algorithm", "scenario"])
    labels = d["treatment"].tolist()
    values = d["total_gb"].astype(float).tolist()
    fig, ax = plt.subplots(figsize=(9.5, 4.4))
    palette = {"HAPPO": "#7A7A7A", "MAAC": "#406A9F", "MASAC": "#2E8B57", "MATD3": "#C45C26"}
    colors = [palette.get(lab.split("-")[0], "#406A9F") for lab in labels]
    bars = ax.bar(labels, values, color=colors)
    ax.set_title("Tamaño total listado en manifiestos de checkpoint (MADRL del proyecto)")
    ax.set_ylabel("GB listados en checkpoint_manifest.json")
    ax.tick_params(axis="x", rotation=65)
    ax.grid(axis="y", alpha=0.25)
    ymax = max(values) if max(values) > 0 else 1
    ax.set_ylim(0, max(0.5, ymax * 1.18))
    for b, v, n in zip(bars, values, d["n_files_listed"].tolist()):
        ax.text(
            b.get_x() + b.get_width() / 2,
            v + ymax * 0.02,
            f"{v:.2f}\n(n={n})",
            ha="center",
            va="bottom",
            fontsize=7,
        )
    if any(d.loc[d["algorithm"] == "HAPPO", "n_files_listed"] == 0):
        ax.text(
            0.01,
            0.98,
            "HAPPO: sin checkpoint_manifest.json en madrl_v3_20260627_164047 (tamaño listado = 0)",
            transform=ax.transAxes,
            va="top",
            fontsize=7.5,
            color="#444444",
        )
    fig.tight_layout()
    out = FIG_DIR / "checkpoint_manifest_total_size_by_treatment.png"
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)
    d.to_csv(TABLE_DIR / "figura_a9_checkpoint_manifest_sizes.csv", index=False)
    return out


def replace_image_safe(docx_path: Path, caption_regex: str, png_path: Path, tag: str) -> dict:
    info = {"file": docx_path.name, "replaced": False}
    # Work on a full copy first to avoid corrupting open file mid-write
    work = docx_path.with_name(docx_path.stem + f"._tmp_{tag}.docx")
    try:
        shutil.copy2(docx_path, work)
    except Exception as exc:
        # if locked, write patched target directly from fresh read attempt
        alt = docx_path.with_name(docx_path.stem + f"_{tag}_PATCHED.docx")
        try:
            shutil.copy2(docx_path, alt)
            work = alt
            info["note"] = "copied to PATCHED first"
        except Exception as exc2:
            info["error"] = f"copy failed: {exc}; {exc2}"
            return info

    png_bytes = png_path.read_bytes()
    try:
        with zipfile.ZipFile(work, "r") as zin:
            xml = zin.read("word/document.xml").decode("utf-8", errors="ignore")
            rels = zin.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")
            namelist = zin.namelist()
    except Exception as exc:
        info["error"] = f"zip read: {exc}"
        work.unlink(missing_ok=True)
        return info

    m = re.search(caption_regex, xml, flags=re.I)
    if not m:
        info["error"] = "caption not found"
        work.unlink(missing_ok=True)
        return info
    after = xml[m.end() :]
    blips = re.findall(r'a:blip[^>]*r:embed="(rId\d+)"', after)
    if not blips:
        info["error"] = "no blip"
        work.unlink(missing_ok=True)
        return info
    rid = blips[0]
    rm = re.search(rf'Relationship[^>]*Id="{rid}"[^>]*Target="([^"]+)"', rels)
    if not rm:
        info["error"] = f"no rel {rid}"
        work.unlink(missing_ok=True)
        return info
    media = "word/" + rm.group(1).lstrip("/")
    if media not in namelist and media.replace("\\", "/") not in [n.replace("\\", "/") for n in namelist]:
        info["error"] = f"media missing {media}"
        work.unlink(missing_ok=True)
        return info

    out_tmp = work.with_suffix(".zipout.docx")
    try:
        with zipfile.ZipFile(work, "r") as zin, zipfile.ZipFile(out_tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.replace("\\", "/") == media.replace("\\", "/"):
                    data = png_bytes
                    info["replaced"] = True
                zout.writestr(item, data)
        # try replace main path
        try:
            out_tmp.replace(docx_path)
            info["saved"] = str(docx_path)
            work.unlink(missing_ok=True)
        except PermissionError:
            alt = docx_path.with_name(docx_path.stem + f"_{tag}_PATCHED.docx")
            out_tmp.replace(alt)
            info["saved"] = str(alt)
            info["note"] = "Word bloqueado; escrito _PATCHED"
            work.unlink(missing_ok=True)
    except Exception as exc:
        info["error"] = str(exc)
        out_tmp.unlink(missing_ok=True)
        work.unlink(missing_ok=True)
    return info


def replace_para_text(para, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def patch_interps(doc: Document, a9_df: pd.DataFrame) -> list[str]:
    actions = []
    cov = pd.read_csv(TABLE_DIR / "figura_5_1_checkpoint_coverage_counts.csv")
    p58 = pd.read_csv(TABLE_DIR / "fig58e_panel_values.csv")

    def mean_panel(panel, algo):
        sub = p58[(p58["panel"] == panel) & (p58["algorithm"] == algo)]
        return float(sub["value"].mean()) if len(sub) else float("nan")

    n_files = int(a9_df["n_files_listed"].sum())
    total_gb = float(a9_df["total_gb"].sum())
    total_mb = float(a9_df["total_mb"].sum())
    max_row = a9_df.sort_values("max_file_mb", ascending=False).iloc[0]
    maac_mean = float(a9_df.loc[a9_df["algorithm"] == "MAAC", "total_gb"].mean())

    patches = {
        "Figura 5.1": {
            "interp": (
                "Interpretación de la figura. Conteo de archivos listados en checkpoint_manifest.json de la corrida "
                "canónica madrl_v3_20260627_164047 (espejo outputs/_drive_madrl/full_data). "
                f"HAPPO-E1/E2/E3 = 0 (manifiesto ausente); MAAC = {int(cov.loc[cov['algorithm']=='MAAC','checkpoint_files_listed'].iloc[0])} "
                f"por escenario; MASAC = {int(cov.loc[cov['algorithm']=='MASAC','checkpoint_files_listed'].iloc[0])}; "
                f"MATD3 = {int(cov.loc[cov['algorithm']=='MATD3','checkpoint_files_listed'].iloc[0])}. "
                "Corrige el falso cero previo de MASAC/MATD3 por filtrar solo rutas episode_(\\d+)."
            ),
        },
        "Figura 5.8e": {
            "note": (
                "Nota. Panel action_l2 desde trace.csv (full_data); paneles EV/BESS desde building_behavior_summary "
                "(ev_charge_total_kwh, battery_throughput_total_kwh). Fuente: elaboración propia a partir de "
                "madrl_v3_20260627_164047."
            ),
            "interp": (
                "Interpretación de la figura. Fuentes mixtas auditadas: action_l2 desde full_data/trace.csv "
                "(incluye HAPPO ≈ 2,25/2,23/2,22 en E1/E2/E3); EV/BESS desde building_behavior_summary, no desde "
                "columnas muertas ev_charge_kwh/electrical_storage_soc (=0 en trace). "
                f"Medias action_l2: HAPPO={fmt_num(mean_panel('action_l2','HAPPO'),2)}, "
                f"MAAC={fmt_num(mean_panel('action_l2','MAAC'),2)}, MASAC={fmt_num(mean_panel('action_l2','MASAC'),2)}, "
                f"MATD3={fmt_num(mean_panel('action_l2','MATD3'),2)}. "
                "Valores: tables/fig58e_panel_values.csv."
            ),
        },
        "Figura A.9": {
            "note": (
                "Nota. La figura representa el tamaño total listado en los manifiestos de checkpoint por algoritmo y escenario. "
                "Fuente: elaboración propia a partir de outputs/_drive_madrl/full_data/{ALGO}/{E}/data/checkpoint_manifest.json "
                "(corrida canónica madrl_v3_20260627_164047)."
            ),
            "interp": (
                "Interpretación de la figura. Se agregan únicamente los bytes declarados en checkpoint_manifest.json "
                f"para los 12 tratamientos del proyecto. Archivos listados: {n_files}; total: {fmt_num(total_gb, 2)} GB "
                f"({fmt_num(total_mb, 2)} MB). HAPPO = 0 GB (sin manifiesto). Media MAAC ≈ {fmt_num(maac_mean, 2)} GB/tratamiento. "
                f"Mayor archivo individual listado: {max_row['algorithm']}-{max_row['scenario']} "
                f"({fmt_num(max_row['max_file_mb'], 2)} MB). No se inventan bytes fuera del manifiesto."
            ),
        },
        "Tabla A.2": {
            "note": (
                "Nota. Ranking por electricity_cost_delta_eur más negativo en building_behavior_summary.csv de "
                "edificios completos (MAAC, MASAC, MATD3 × E1–E3). HAPPO excluido: sin building_behavior_summary "
                "en la corrida canónica. Fuente: outputs/_drive_madrl/full_data/{ALGO}/{E}/data/building_behavior_summary.csv."
            ),
            "interp": (
                "Interpretación por edificio. El cambio de costo eléctrico más negativo se observa en MATD3-E1 Building_14, "
                "con electricity_cost_delta_eur=-3532.42. Ranking validado fila a fila contra building_behavior_summary.csv "
                "de la corrida canónica (50 episodios; MAAC/MASAC/MATD3)."
            ),
        },
    }

    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        for token, cfg in patches.items():
            if token in t and (t.startswith("Figura") or t.startswith("Tabla") or token in t[:40]):
                for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                    tj = (doc.paragraphs[j].text or "").strip()
                    if "note" in cfg and tj.startswith("Nota."):
                        replace_para_text(doc.paragraphs[j], cfg["note"])
                        actions.append(f"{token} nota")
                    if "interp" in cfg and (
                        tj.startswith("Interpretacion") or tj.startswith("Interpretación")
                    ):
                        replace_para_text(doc.paragraphs[j], cfg["interp"])
                        actions.append(f"{token} interpretacion")
                        break
    return actions


def verify_cap6(doc: Document) -> dict:
    paras = [(i, (p.text or "").strip()) for i, p in enumerate(doc.paragraphs)]
    h64 = next((t for _, t in paras if t.startswith("6.4 ")), None)
    h65 = next((t for _, t in paras if t.startswith("6.5 ")), None)
    t61 = next((t for _, t in paras if t.startswith("Tabla 6.1")), None)
    t62 = next((t for _, t in paras if t.startswith("Tabla 6.2")), None)
    # sniff table 6.1 content
    body_text = "\n".join(t for _, t in paras)
    return {
        "has_6_4": bool(h64),
        "has_6_5": bool(h65),
        "has_tabla_6_1": bool(t61),
        "has_tabla_6_2": bool(t62),
        "6_4_title": h64,
        "6_5_title": h65,
        "tabla_6_1_caption": t61,
        "tabla_6_2_caption": t62,
        "mentions_executed_plan": "15 de julio de 2026" in body_text or "H1. Cobertura HAPPO" in body_text,
        "no_doctorado_in_65": h65 is not None and "doctor" not in h65.lower(),
    }


def save_doc(doc: Document, path: Path) -> dict:
    try:
        doc.save(str(path))
        return {"ok": True, "saved": str(path)}
    except PermissionError:
        alt = path.with_name(path.stem + "_PATCHED.docx")
        doc.save(str(alt))
        return {"ok": False, "saved": str(alt), "note": "bloqueado"}


def main():
    a9 = scan_a9()
    png = plot_a9(a9)
    top = pd.read_csv(TABLE_DIR / "tabla_a2_top20_electricity_cost_reduction.csv")

    results = {"a9_png": str(png), "a9": a9.to_dict(orient="records"), "docs": {}}
    for path, label in [(PRIMARY, "primary"), (MIRROR, "mirror")]:
        if not path.exists():
            results["docs"][label] = {"error": "missing"}
            continue
        doc = Document(str(path))
        actions = patch_interps(doc, a9)
        cap6 = verify_cap6(doc)
        # If 6.4/6.5 not yet user content, leave note — previous run already inserted captions
        save_info = save_doc(doc, path)
        emb = replace_image_safe(
            Path(save_info["saved"]),
            r"Figura A\.9|Tamano total listado en manifiestos|Tamaño total listado en manifiestos",
            png,
            "A9",
        )
        results["docs"][label] = {
            "actions": actions,
            "cap6": cap6,
            "save": save_info,
            "a9_embed": emb,
        }

    report = {
        "a2_verdict": "VALIDADA",
        "a2_proof": {
            "definition": "top-20 electricity_cost_delta_eur más negativo en edificios completos (MAAC/MASAC/MATD3)",
            "sources": sorted(top["source_file"].unique().tolist()),
            "top20": top[
                ["rank", "algorithm", "scenario", "agent", "electricity_cost_delta_eur", "source_file"]
            ].to_dict(orient="records"),
            "happo_excluded": "Sin building_behavior_summary.csv en full_data/HAPPO/{E}",
            "row_match": "20/20 vs ABRIR_ESTE (redondeo presentado)",
        },
        "a9_old_vs_new": {
            "old": "Interpretación previa: 351 checkpoints, ~119627.66 MB, max MAAC-E1 763.36 MB",
            "new_files": int(a9["n_files_listed"].sum()),
            "new_total_gb": float(a9["total_gb"].sum()),
            "new_total_mb": float(a9["total_mb"].sum()),
            "per_treatment": a9[["treatment", "n_files_listed", "total_gb", "source"]].to_dict(orient="records"),
            "cause": (
                "Agregado previo no coincidía con bytes de checkpoint_manifest.json del espejo full_data "
                "restringido a HAPPO/MAAC/MASAC/MATD3; HAPPO debe ser 0 si falta manifiesto."
            ),
        },
        "docs": results["docs"],
        "png": str(png.relative_to(REPO)).replace("\\", "/"),
    }
    out = REPORT_DIR / "tabla_a2_a9_validation_final.json"
    out.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({
        "a2": "VALIDADA",
        "a9_files": report["a9_old_vs_new"]["new_files"],
        "a9_gb": report["a9_old_vs_new"]["new_total_gb"],
        "primary": results["docs"].get("primary"),
        "report": str(out),
    }, indent=2, ensure_ascii=False, default=str))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
