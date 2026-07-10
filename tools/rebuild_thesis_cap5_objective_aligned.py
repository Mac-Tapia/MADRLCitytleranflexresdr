from __future__ import annotations

import csv
import json
import math
import re
import shutil
from copy import deepcopy
from pathlib import Path

import pandas as pd
from scipy import stats
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
SRC = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_50EP_ANTECEDENTES.docx"
OUT = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_OBJETIVOS_ESTADISTICA_ALINEADA.docx"
EPISODE_CSV = REPO / "outputs" / "_drive_madrl" / "full_data" / "analysis_real_drive" / "tables" / "district_episode_kpis.csv"
SUMMARY_CSV = REPO / "outputs" / "_drive_madrl" / "full_data" / "analysis_real_drive" / "tables" / "objective_aligned_statistical_results.csv"
METRICS = REPO / "outputs" / "_drive_madrl" / "full_data" / "analysis_real_drive" / "thesis_docx_objective_aligned_metrics.json"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)

OBJECTIVES = [
    {
        "objective": "OE.1",
        "scenario": "E1",
        "dimension": "Flexibilidad energetica",
        "metric": "reward_mean",
        "direction": "max",
        "interpretation": "Se usa reward_mean del escenario E1 porque E1 pondera flexibilidad con peso dominante [0,70; 0,15; 0,15]. Es una medida de efecto del algoritmo sobre la recompensa de flexibilidad, no una prueba aislada de energia neta.",
    },
    {
        "objective": "OE.2",
        "scenario": "E2",
        "dimension": "Emisiones de CO2",
        "metric": "district_emission",
        "direction": "min",
        "interpretation": "Se usa district_emission del escenario E2 porque E2 pondera emisiones con peso dominante [0,15; 0,70; 0,15]. Menor valor implica menor emision distrital observada.",
    },
    {
        "objective": "OE.3",
        "scenario": "E3",
        "dimension": "Costos energeticos",
        "metric": "district_cost",
        "direction": "min",
        "interpretation": "Se usa district_cost del escenario E3 porque E3 pondera costos con peso dominante [0,25; 0,15; 0,60]. Menor valor implica menor costo distrital observado.",
    },
]


def text_of(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def clear_body_keep_sectpr(document: Document) -> None:
    body = document.element.body
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def append_before_sectpr(document: Document, el) -> None:
    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(el)
    else:
        body.insert(body.index(sect_pr), el)


def style_doc(document: Document) -> None:
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        if name not in [s.name for s in document.styles]:
            continue
        st = document.styles[name]
        st.font.name = "Calibri"
        if name == "Normal":
            st.font.size = Pt(11)
            st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.line_spacing = 1.15
        else:
            st.font.bold = True
            st.font.color.rgb = ACCENT
            st.font.size = Pt(16 if name == "Heading 1" else 13 if name == "Heading 2" else 11.5)


def set_bg(cell, color: str = "1F4E79") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def p(doc: Document, text: str):
    para = doc.add_paragraph()
    para.add_run(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    return para


def table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], font_size: float = 7.2):
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GREY
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(head)
        rr.bold = True
        rr.font.size = Pt(font_size)
        rr.font.color.rgb = RGBColor(255, 255, 255)
        set_bg(cell)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(font_size)
    doc.add_paragraph()
    return tbl


def fmt(x, nd: int = 4) -> str:
    try:
        return f"{float(x):,.{nd}f}"
    except Exception:
        return str(x)


def holm_adjust(pairs: list[tuple[str, float]]) -> list[tuple[str, float, float]]:
    m = len(pairs)
    ordered = sorted(pairs, key=lambda x: x[1])
    adjusted = []
    running = 0.0
    for rank, (name, pv) in enumerate(ordered, start=1):
        adj = min(1.0, (m - rank + 1) * pv)
        running = max(running, adj)
        adjusted.append((name, pv, running))
    return sorted(adjusted, key=lambda x: x[0])


def analyze_objectives() -> tuple[pd.DataFrame, dict]:
    df = pd.read_csv(EPISODE_CSV)
    rows = []
    detail = {}
    for spec in OBJECTIVES:
        sub = df[df["scenario"] == spec["scenario"]].copy()
        metric = spec["metric"]
        direction = spec["direction"]
        summary = (
            sub.groupby("algorithm")[metric]
            .agg(["count", "mean", "median", "std", "min", "max"])
            .reset_index()
        )
        ascending = direction == "min"
        summary_sorted = summary.sort_values("median", ascending=ascending).reset_index(drop=True)
        best_observed = summary_sorted.iloc[0]["algorithm"]
        complete = summary[summary["count"] >= 50].copy()
        complete_sorted = complete.sort_values("median", ascending=ascending).reset_index(drop=True)
        best_50 = complete_sorted.iloc[0]["algorithm"] if not complete_sorted.empty else "NA"
        grouped = [g[metric].dropna().values for _, g in sub.groupby("algorithm")]
        kw = stats.kruskal(*grouped)
        n = sum(len(g) for g in grouped)
        k = len(grouped)
        epsilon2 = (kw.statistic - k + 1) / (n - k) if n > k else float("nan")
        pair_p = []
        algorithms = sorted(sub["algorithm"].unique())
        for i, a in enumerate(algorithms):
            for b in algorithms[i + 1 :]:
                av = sub[sub["algorithm"] == a][metric].dropna().values
                bv = sub[sub["algorithm"] == b][metric].dropna().values
                pv = stats.mannwhitneyu(av, bv, alternative="two-sided").pvalue
                pair_p.append((f"{a} vs {b}", pv))
        pair_adj = holm_adjust(pair_p)
        for _, r in summary.iterrows():
            rows.append({
                "objective": spec["objective"],
                "scenario": spec["scenario"],
                "dimension": spec["dimension"],
                "metric": metric,
                "direction": direction,
                "algorithm": r["algorithm"],
                "n_episodes": int(r["count"]),
                "mean": r["mean"],
                "median": r["median"],
                "std": r["std"],
                "best_observed_by_median": best_observed,
                "best_50ep_by_median": best_50,
                "kw_h": kw.statistic,
                "kw_p": kw.pvalue,
                "kw_epsilon2": epsilon2,
            })
        detail[spec["objective"]] = {
            "spec": spec,
            "summary": summary_sorted,
            "best_observed": best_observed,
            "best_50": best_50,
            "kw": kw,
            "epsilon2": epsilon2,
            "pair_adj": pair_adj,
        }
    out = pd.DataFrame(rows)
    SUMMARY_CSV.parent.mkdir(parents=True, exist_ok=True)
    out.to_csv(SUMMARY_CSV, index=False, encoding="utf-8")
    return out, detail


def add_cap5(doc: Document, detail: dict) -> None:
    doc.add_heading("Capitulo 5. Resultados y contrastacion de hipotesis", level=1)
    p(doc, "Este capitulo se reestructura para responder directamente los tres objetivos especificos. No se mezclan KPIs de flexibilidad, emisiones y costo en una sola prueba global sin interpretacion. Cada objetivo se contrasta con su escenario dominante y con una variable de resultado observable en los artefactos reales: OE.1 usa E1, OE.2 usa E2 y OE.3 usa E3. La estadistica se calcula sobre episodios reales de district_episode_kpis.csv; no se usan valores heredados de corridas antiguas ni resultados no observados.")
    p(doc, "La regla de evidencia es la siguiente: HAPPO contiene 49 episodios reales por escenario y se incluye en analisis distritales; MAAC, MASAC y MATD3 contienen 50 episodios reales por escenario. HAPPO no se usa para resultados por edificio ni checkpoints porque faltan building_kpis.csv, building_behavior_summary.csv y checkpoint_manifest.json. Esta separacion evita alucinacion de resultados y preserva la trazabilidad de la tesis.")

    table(doc, "Tabla 5.1. Vinculacion directa entre objetivos, escenarios y variables estadisticas.",
          ["Objetivo", "Dimension", "Escenario", "Variable estadistica", "Direccion", "Justificacion"],
          [[d["spec"]["objective"], d["spec"]["dimension"], d["spec"]["scenario"], d["spec"]["metric"], "mayor es mejor" if d["spec"]["direction"] == "max" else "menor es mejor", d["spec"]["interpretation"]] for d in detail.values()],
          6.8)

    for idx, oe in enumerate(["OE.1", "OE.2", "OE.3"], start=2):
        d = detail[oe]
        spec = d["spec"]
        doc.add_heading(f"5.{idx} {oe}: efecto sobre {spec['dimension']}", level=2)
        p(doc, spec["interpretation"])
        rows = []
        for _, r in d["summary"].iterrows():
            rows.append([
                r["algorithm"],
                int(r["count"]),
                fmt(r["mean"], 6 if spec["metric"] == "reward_mean" else 2),
                fmt(r["median"], 6 if spec["metric"] == "reward_mean" else 2),
                fmt(r["std"], 6 if spec["metric"] == "reward_mean" else 2),
                "completo 50 ep" if int(r["count"]) >= 50 else "49 ep: evidencia distrital parcial",
            ])
        table(doc, f"Tabla 5.{idx}. Resultado descriptivo de {oe} por algoritmo.", ["Algoritmo", "n", "Media", "Mediana", "Desv. est.", "Cobertura"], rows, 7.0)
        p(doc, f"Resultado descriptivo: el mejor algoritmo observado por mediana para {oe} es {d['best_observed']}. Si se restringe la decision a algoritmos con 50 episodios completos, el mejor algoritmo es {d['best_50']}. Esta distincion es obligatoria porque HAPPO tiene 49 episodios reales.")
        kw = d["kw"]
        p(doc, f"Contrastacion global Kruskal-Wallis para {oe}: H={kw.statistic:.4f}, p={kw.pvalue:.6g}, epsilon2={d['epsilon2']:.4f}. Con alpha=0,05, {'se rechaza H0 de igualdad distribucional entre algoritmos' if kw.pvalue < 0.05 else 'no se rechaza H0 de igualdad distribucional entre algoritmos'}. La prueba se usa porque las comparaciones de algoritmos de RL suelen violar supuestos de normalidad y requieren protocolos estadisticos robustos (Colas et al., 2019; Agarwal et al., 2021; Patterson et al., 2024).")
        pair_rows = [[name, fmt(pv, 6), fmt(adj, 6), "significativo" if adj < 0.05 else "no significativo"] for name, pv, adj in d["pair_adj"]]
        table(doc, f"Tabla 5.{idx}a. Mann-Whitney U por pares con ajuste Holm para {oe}.", ["Par", "p sin ajustar", "p Holm", "Decision"], pair_rows, 7.2)

    doc.add_heading("5.5 Sintesis de contrastacion de objetivos", level=2)
    synth_rows = []
    for oe in ["OE.1", "OE.2", "OE.3"]:
        d = detail[oe]
        synth_rows.append([
            oe,
            d["spec"]["dimension"],
            d["spec"]["scenario"],
            d["best_observed"],
            d["best_50"],
            f"H={d['kw'].statistic:.3f}; p={d['kw'].pvalue:.3g}",
            "diferencia global detectada" if d["kw"].pvalue < 0.05 else "sin diferencia global",
        ])
    table(doc, "Tabla 5.8. Respuesta directa a OE.1, OE.2 y OE.3.", ["Objetivo", "Dimension", "Escenario", "Mejor observado", "Mejor con 50 ep", "Kruskal-Wallis", "Decision"], synth_rows, 7.0)
    p(doc, "La respuesta a los objetivos cambia respecto a versiones previas: no debe afirmarse que MATD3 domina todas las dimensiones si los datos reales no lo muestran. Para OE.1, el mejor desempeño de la recompensa E1 corresponde a MAAC. Para OE.2, HAPPO muestra menor emision distrital observada, pero con 49 episodios y sin KPIs de edificio; entre algoritmos completos de 50 episodios, MAAC presenta la menor mediana/valor distrital de emisiones del escenario E2. Para OE.3, HAPPO muestra menor costo observado con 49 episodios, pero entre algoritmos completos de 50 episodios MATD3 presenta el menor costo medio/competitivo segun la tabla descriptiva.")
    p(doc, "Las figuras A.1-A.9 y las tablas del Anexo A se mantienen como evidencia grafica y tabular. La interpretacion final debe citar siempre el objetivo asociado: A.1-A.2 apoyan dinamica de entrenamiento y energia; A.3 apoya costos; A.4 y A.6 apoyan emisiones; A.5 apoya costo por edificio; A.7-A.8 apoyan equipamiento controlado/no controlado; A.9 apoya reproducibilidad por checkpoints.")
    doc.add_heading("5.6 Limitaciones estadisticas", level=2)
    p(doc, "Los episodios dentro de una corrida no sustituyen multiples semillas independientes. Por ello, las pruebas de este capitulo son evidencia inferencial intra-corrida sobre episodios reales, no una prueba definitiva de robustez multi-semilla. La literatura de evaluacion rigurosa de aprendizaje por refuerzo recomienda multiples semillas, tamanos de efecto e intervalos de confianza para conclusiones causales finales (Henderson et al., 2018; Colas et al., 2019; Agarwal et al., 2021; Patterson et al., 2024).")


def rebuild_doc(detail: dict) -> None:
    shutil.copyfile(SRC, OUT)
    doc = Document(OUT)
    style_doc(doc)
    children = list(doc.element.body)
    idx_cap5 = idx_cap6 = None
    for i, el in enumerate(children):
        txt = text_of(el)
        if idx_cap5 is None and txt.startswith("Capitulo 5. Resultados"):
            idx_cap5 = i
        if idx_cap6 is None and txt.startswith("Capitulo 6. Conclusiones"):
            idx_cap6 = i
    if idx_cap5 is None or idx_cap6 is None:
        raise RuntimeError(f"No se encontraron limites Cap5/Cap6: {idx_cap5}, {idx_cap6}")
    before = [deepcopy(el) for el in children[:idx_cap5]]
    after = [deepcopy(el) for el in children[idx_cap6:] if el.tag != qn("w:sectPr")]
    clear_body_keep_sectpr(doc)
    for el in before:
        append_before_sectpr(doc, el)
    add_cap5(doc, detail)
    for el in after:
        append_before_sectpr(doc, el)
    doc.save(OUT)


def main() -> None:
    _, detail = analyze_objectives()
    rebuild_doc(detail)
    v = Document(OUT)
    paras = [p.text.strip() for p in v.paragraphs if p.text.strip()]
    full = "\n".join(paras)
    metrics = {
        "output": str(OUT),
        "stats_csv": str(SUMMARY_CSV),
        "size_bytes": OUT.stat().st_size,
        "paragraphs_non_empty": len(paras),
        "word_count_estimated": len(re.findall(r"\b[\wáéíóúÁÉÍÓÚñÑüÜ-]+\b", full, re.UNICODE)),
        "tables": len(v.tables),
        "inline_images": len(v.inline_shapes),
        "has_oe1": "OE.1: efecto sobre Flexibilidad energetica" in full,
        "has_oe2": "OE.2: efecto sobre Emisiones de CO2" in full,
        "has_oe3": "OE.3: efecto sobre Costos energeticos" in full,
        "has_no_old_global_kw": "p = 0,0459" not in full and "p=0,0459" not in full,
        "happo_caveat": "HAPPO contiene 49 episodios reales" in full,
        "figures_a_1_a_9": all(f"Figura A.{i}" in full for i in range(1, 10)),
        "figures_b_1_a_9": all(f"Figura B.{i}" in full for i in range(1, 10)),
    }
    METRICS.parent.mkdir(parents=True, exist_ok=True)
    METRICS.write_text(json.dumps(metrics, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(metrics, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
