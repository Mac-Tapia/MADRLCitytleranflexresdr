#!/usr/bin/env python3
"""Auditoría profunda: Dec-POMDP dims/fórmulas en todos los capítulos de los 2 Word."""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[2]
DOCS = REPO / "docs"
REPORT = DOCS / "AUDITORIA_DECPOMDP_CROSSCHAPTER_WORD_2026-07-29.json"

PATTERNS_OBSOLETE = [
    ("obsolete_oi_57_330", re.compile(r"57\s*[–\-]\s*330|57\s*a\s*330", re.I)),
    ("approx_oi", re.compile(r"rango aproximado|aproximad[oa].{0,60}(54|57).{0,30}327", re.I)),
    ("approx_ai_tilde", re.compile(r"~\s*5\s*a\s*~\s*44|~\s*5\s*[–\-]\s*44|de\s*~\s*5\s*a\s*~\s*44", re.I)),
    ("empirieca_typo", re.compile(r"empirieca", re.I)),
    ("wrong_sum_218", re.compile(r"suma.{0,20}218|218\s*dimensiones\s*de\s*acci", re.I)),
]

PATTERNS_PRESENT = [
    ("decpomdp", re.compile(r"Dec-POMDP|DecPOMDP", re.I)),
    ("oi_54_327", re.compile(r"54\s*[–\-]\s*327")),
    ("ai_5_44", re.compile(r"5\s*[–\-]\s*44|5\s*a\s*44")),
    ("ds_1856", re.compile(r"1\s*856|1856")),
    ("dai_formula", re.compile(r"d_a[ií]\s*=\s*2\s*\+\s*n|2\s*\+\s*n_i", re.I)),
    ("r_team", re.compile(r"team_reward_ratio|r_team\s*=\s*0[,.]70", re.I)),
    ("tabla_2a", re.compile(r"Tabla\s*2\.A", re.I)),
    ("distrito_no_agente", re.compile(r"distrito no es un agente", re.I)),
    ("ctde_backend", re.compile(r"concatenated_local_observations", re.I)),
]


def chapter_tracker():
    chap = 0

    def update(text: str) -> int:
        nonlocal chap
        t = text.strip()
        m = re.match(r"^Cap[ií]tulo\s*(\d+)\b", t, re.I)
        if m:
            chap = int(m.group(1))
            return chap
        m2 = re.match(r"^([1-6])\.\d+", t)
        if m2 and len(t) < 120:
            chap = int(m2.group(1))
        return chap

    return update


def audit(path: Path) -> dict:
    doc = Document(str(path))
    upd = chapter_tracker()
    obsolete = []
    presence: dict[int, dict[str, int]] = {i: {} for i in range(0, 7)}
    mentions: dict[int, list[dict]] = {i: [] for i in range(0, 7)}

    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        if not t.strip():
            continue
        chap = upd(t)
        for name, rx in PATTERNS_OBSOLETE:
            for m in rx.finditer(t):
                obsolete.append(
                    {
                        "chap": chap,
                        "para": i,
                        "kind": name,
                        "snippet": t[max(0, m.start() - 40) : m.end() + 60].replace("\n", " ")[:160],
                    }
                )
        for name, rx in PATTERNS_PRESENT:
            n = len(rx.findall(t))
            if n:
                presence[chap][name] = presence[chap].get(name, 0) + n
                if name == "decpomdp" and n:
                    mentions[chap].append({"para": i, "snippet": t[:120]})

    table_hits = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            row_txt = " | ".join((c.text or "").replace("\n", " ") for c in row.cells)
            if re.search(r"57\s*[–\-]\s*330", row_txt) or re.search(r"~\s*5", row_txt):
                table_hits.append({"table": ti, "row": ri, "text": row_txt[:180]})

    # Recommendations per chapter
    recs = []
    for ch in range(1, 7):
        pr = presence.get(ch, {})
        has_dec = pr.get("decpomdp", 0) > 0
        if not has_dec and ch not in (2, 4):
            # Cap1/3/5/6 may mention lightly
            pass
        issues = []
        if has_dec:
            if pr.get("oi_54_327", 0) == 0 and ch in (1, 3, 4, 5):
                # Cap1/3/5 may not need full dims
                if ch in (3, 5) and pr.get("ds_1856", 0) == 0:
                    issues.append("menciona Dec-POMDP pero sin dims canónicas; valorar remisión a Cap.2/Tabla 2.A")
            if ch == 3 and pr.get("tabla_2a", 0) == 0:
                issues.append("Cap.3 (metodología/dataset) podría remitar Tabla 2.A al definir agentes")
            if ch == 1 and pr.get("decpomdp", 0) > 0 and pr.get("oi_54_327", 0) == 0:
                issues.append("Cap.1: mención conceptual OK; no requiere tabla dims (evitar sobrecarga)")
            if ch == 5 and pr.get("decpomdp", 0) > 0:
                issues.append("Cap.5: resultados; solo corregir si hay dims erróneas (no re-axiomatizar)")
            if ch == 6 and pr.get("decpomdp", 0) > 0:
                issues.append("Cap.6: conclusión; una frase de remisión a tupla Cap.2/4 es opcional")
        ch_obs = [o for o in obsolete if o["chap"] == ch]
        if ch_obs:
            issues.append(f"CORREGIR {len(ch_obs)} dato(s) obsoleto(s)")
        if issues:
            recs.append({"chapter": ch, "decpomdp_hits": pr.get("decpomdp", 0), "presence": pr, "actions": issues})

    return {
        "file": path.name,
        "obsolete": obsolete,
        "table_hits": table_hits,
        "presence_by_chapter": {str(k): v for k, v in presence.items() if v},
        "recommendations": recs,
        "decpomdp_mentions_sample": {
            str(k): v[:3] for k, v in mentions.items() if v
        },
    }


def main() -> int:
    results = []
    for name in (
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx",
        "Inforne_tesisV4_FINAL_REVISADO_50_EPISODIOS.docx",
    ):
        results.append(audit(DOCS / name))
    # Cross summary
    summary = {
        "obsolete_total": sum(len(r["obsolete"]) for r in results),
        "table_hits_total": sum(len(r["table_hits"]) for r in results),
        "must_fix": [],
        "should_integrate_light": [],
        "ok_no_change": [],
    }
    for r in results:
        for o in r["obsolete"]:
            summary["must_fix"].append({"file": r["file"], **o})
        for rec in r["recommendations"]:
            acts = rec["actions"]
            if any(a.startswith("CORREGIR") for a in acts):
                summary["must_fix"].append({"file": r["file"], "chapter": rec["chapter"], "actions": acts})
            elif any("remit" in a.lower() or "remisión" in a.lower() or "valorar" in a.lower() for a in acts):
                summary["should_integrate_light"].append({"file": r["file"], **rec})
            else:
                summary["ok_no_change"].append({"file": r["file"], **rec})

    report = {
        "timestamp": "2026-07-29",
        "policy": "profundidad cross-chapter; solo integrar donde aporta consistencia",
        "results": results,
        "summary": summary,
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
