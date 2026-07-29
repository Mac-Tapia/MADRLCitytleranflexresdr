#!/usr/bin/env python3
"""Parche Cap. IV: narrativa de implementación + validación UC3M/schema en Word canónico.

Fuente de verdad:
- docs/tesis_capitulos/Capitulo_4_Desarrollo_Propuesta.md (§§4.1.1, 4.6.1, 4.7.1, 4.10–4.12)

No regenera toda la tesis: inserta/actualiza secciones al final de Cap. IV
(antes de Cap. V) y corrige el tipográfico de §4.7 si aparece corrupto.
No embebe figuras de resultados multicriterio (pertenecen a Cap. V).
"""
from __future__ import annotations

import json
import re
import shutil
import unicodedata
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

import sys

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, TESIS, existing_canons  # noqa: E402

REPORT = DOCS / "CAP4_IMPLEMENTACION_PATCH_REPORT_2026-07-29.json"
MARKER = "4.11 mapa de contribucion por carpetas"
VALIDATION_MARKER = "4.12 validacion implementada"


def set_run_font(run, bold: bool = False, size: float = 12.0, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def replace_paragraph_text(p: Paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p.clear()
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = False, italic: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, bold=bold, italic=italic)
    return new_para


def set_heading_style(paragraph: Paragraph, level: int) -> None:
    style_name = f"Heading {level}"
    try:
        paragraph.style = style_name
    except KeyError:
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
        set_run_font(run, bold=True, size=14.0 if level == 2 else 12.0)


def norm(s: str) -> str:
    """Lowercase + collapse spaces + strip accents (contribución → contribucion)."""
    folded = unicodedata.normalize("NFKD", s or "")
    folded = "".join(ch for ch in folded if not unicodedata.combining(ch))
    return re.sub(r"\s+", " ", folded.strip().lower())


def chapter_bounds(doc: Document, chapter: int) -> tuple[int | None, int | None]:
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if re.search(rf"^Cap[ií]tulo\s*{chapter}\b", t, re.I):
            start = i
        elif start is not None and re.search(rf"^Cap[ií]tulo\s*{chapter + 1}\b", t, re.I):
            end = i
            break
    return start, end


def insert_table_after_paragraph(
    paragraph: Paragraph,
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    caption: str,
) -> Paragraph:
    cap = insert_paragraph_after(paragraph, caption, italic=True)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=9.0)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=8.5)
    tbl = table._tbl
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)
    cap._p.addnext(tbl)
    note_el = OxmlElement("w:p")
    tbl.addnext(note_el)
    note = Paragraph(note_el, paragraph._parent)
    run = note.add_run(
        "Nota. Fuente: elaboración propia a partir del repositorio MADRLCitytleranflexresdr y pytest "
        "(2026-07-29)."
    )
    set_run_font(run, italic=True, size=10.0)
    return note


FOLDER_ROWS = [
    ["CityLearn/", "Dependencia", "Simulador v2, capa v3, train_citylearn_v3_*.py, dataset Iquitos"],
    ["external/", "Dependencia", "Backends HARL, MARL/src, off-policy, MAAC"],
    ["uc3m/", "Propio (diseño)", "BACT, 7 ejes, HPHI, EmpiricalProtocol; no canal de las 12 corridas"],
    ["tools/", "Propio", "Dataset, eval, thesis Word, validators de training"],
    ["scripts/", "Propio", "Orquestación, batería no paramétrica, multicriterio, guards"],
    ["tests/", "Propio", "tests/uc3m/ + tests/citylearn_v3/ (schema smoke)"],
    ["outputs/", "Artefactos", "Corridas, estadística, multicriterio (consumo Cap. V)"],
    ["docs/", "Propio", "Arquitectura, capítulos markdown, canones Word"],
    ["examples_madrl_v3/", "Propio", "Notebook Colab / two_phase_happo_masac_v3"],
]

VALIDATION_ROWS = [
    ["tests/uc3m/", "BACT, reward, HPHI, KPI, factory, env, multicriteria, …", "180 passed"],
    ["tests/citylearn_v3/", "test_schema_smoke.py", "4 passed"],
    ["Total", "184 ítems", "184 passed (91,09 s)"],
]


def strip_existing_block(doc: Document, start_idx: int, end_idx: int, markers: tuple[str, ...]) -> int | None:
    """If markers already present in Cap. IV, delete from first marker heading to Cap. V.

    Removes body siblings (paragraphs AND tables) from the marker element up to,
    but not including, the Cap. V heading. Paragraph-only deletion left orphan
    tables and caused duplicate 4.11/4.12 blocks on re-run.

    Returns the former block_start index, or None if nothing removed.
    """
    paras = doc.paragraphs
    block_start = None
    for i in range(start_idx, end_idx):
        n = norm(paras[i].text or "")
        if any(m in n for m in markers):
            block_start = i
            break
    if block_start is None:
        return None
    start_el = paras[block_start]._element
    end_el = paras[end_idx]._element
    body = start_el.getparent()
    if body is None:
        return None
    to_remove: list = []
    seen_start = False
    for child in list(body):
        if child is start_el:
            seen_start = True
        if not seen_start:
            continue
        if child is end_el:
            break
        to_remove.append(child)
    for child in to_remove:
        body.remove(child)
    return block_start


def patch_doc(doc: Document) -> list[str]:
    actions: list[str] = []
    c4, c5 = chapter_bounds(doc, 4)
    if c4 is None or c5 is None:
        actions.append("ERROR: no se localizaron límites Cap. IV / Cap. V")
        return actions

    # Fix corrupted 4.7 heading
    for i in range(c4, c5):
        t = (doc.paragraphs[i].text or "").strip()
        n = norm(t)
        if n.startswith("4.7") and ("disecuasiexperimental" in n or "diseno experimental" in n or "diseño experimental" in n):
            replace_paragraph_text(
                doc.paragraphs[i],
                "4.7 Diseño cuasiexperimental: matriz de 12 corridas",
                bold=True,
            )
            set_heading_style(doc.paragraphs[i], 2)
            actions.append(f"fix_heading_4_7@{i}")
            break
        if n.startswith("4.7 diseno") or n.startswith("4.7 diseño"):
            if "cuasiexperimental" not in n and "matriz" in n:
                replace_paragraph_text(
                    doc.paragraphs[i],
                    "4.7 Diseño cuasiexperimental: matriz de 12 corridas",
                    bold=True,
                )
                set_heading_style(doc.paragraphs[i], 2)
                actions.append(f"fix_heading_4_7_alt@{i}")
                break

    # Refresh Cap. IV bounds after possible edits
    c4, c5 = chapter_bounds(doc, 4)
    assert c4 is not None and c5 is not None

    markers = (MARKER, VALIDATION_MARKER, "4.13 reproducibilidad")
    # Remove previous patch block if re-running
    strip_existing_block(doc, c4, c5, markers)
    c4, c5 = chapter_bounds(doc, 4)
    assert c4 is not None and c5 is not None

    # Anchor: last non-empty paragraph before Cap. V
    anchor_idx = c5 - 1
    while anchor_idx > c4 and not (doc.paragraphs[anchor_idx].text or "").strip():
        anchor_idx -= 1
    cursor = doc.paragraphs[anchor_idx]

    # --- 4.11 Folder map ---
    h411 = insert_paragraph_after(
        cursor,
        "4.11 Mapa de contribución por carpetas (propio vs externo)",
        bold=True,
    )
    set_heading_style(h411, 2)
    actions.append("insert_4_11_heading")
    p1 = insert_paragraph_after(
        h411,
        "La narrativa de implementación del repositorio MADRLCitytleranflexresdr se organiza por "
        "frontera de propiedad. Solo el código propio de la tesis se trata como aporte metodológico; "
        "CityLearn/ y external/ se citan como dependencias. El framework uc3m/ es capa de diseño "
        "(BACT, siete ejes, EmpiricalProtocol) y no sustituye el canal de las doce corridas "
        "(CityLearn/scripts/train_citylearn_v3_*.py + backends external/).",
    )
    note = insert_table_after_paragraph(
        p1,
        doc,
        ["Carpeta", "Frontera", "Rol en la propuesta (Cap. IV)"],
        FOLDER_ROWS,
        "Tabla 4.11. Mapa de contribución por carpetas del repositorio (propio vs externo).",
    )
    actions.append("insert_4_11_table")

    # --- 4.12 Validation ---
    h412 = insert_paragraph_after(
        note,
        "4.12 Validación implementada (tests UC3M y schema CityLearn v3)",
        bold=True,
    )
    set_heading_style(h412, 2)
    actions.append("insert_4_12_heading")
    p_cmd = insert_paragraph_after(
        h412,
        "El 2026-07-29 se ejecutó la validación de software de la fachada UC3M y del smoke de schema "
        "CityLearn v3 en el entorno .venv39-citylearn-v3, tras verificar el contexto del proyecto con "
        "scripts/verify_project_context.ps1. Comando:",
    )
    p_code = insert_paragraph_after(
        p_cmd,
        ".\\.venv39-citylearn-v3\\Scripts\\python.exe -B -m pytest tests/uc3m/ tests/citylearn_v3/ -v --tb=line",
        italic=True,
    )
    p_interp = insert_paragraph_after(
        p_code,
        "Resultado: 184 tests passed (1 warning no bloqueante de SciPy en "
        "test_full_methodology_battery_smoke), duración 91,09 s. Esta validación acredita "
        "reproducibilidad de la capa de diseño y del contrato de datos; no sustituye el "
        "entrenamiento 4×3 ni las pruebas estadísticas del Capítulo 5.",
    )
    note2 = insert_table_after_paragraph(
        p_interp,
        doc,
        ["Suite", "Alcance", "Resultado"],
        VALIDATION_ROWS,
        "Tabla 4.12. Resumen de validación implementada (pytest, 2026-07-29).",
    )
    p_launch = insert_paragraph_after(
        note2,
        "Validación ligera adicional: tools/training/validate_launch_config.py confirmó la construcción "
        "de 12 jobs (4 algoritmos × 3 escenarios), hiperparámetros por fase y dry-run de dynamic "
        "backfill (rc=0, 12 jobs, cap ≤ 6, primeros 6 = fase 1 HAPPO+MASAC). "
        "No se reentrenaron las 12 corridas MADRL en esta integración.",
    )
    actions.append("insert_4_12_body")

    # --- 4.13 Reproducibility ---
    h413 = insert_paragraph_after(
        p_launch,
        "4.13 Reproducibilidad operativa y puente a evidencia",
        bold=True,
    )
    set_heading_style(h413, 2)
    insert_paragraph_after(
        h413,
        "El entorno canónico de desarrollo es Python 3.9 (.venv39-citylearn-v3) con activación "
        "scripts/activate_citylearn_v3.ps1. La orquestación Colab usa el protocolo "
        "two_phase_happo_masac_v3; la vía local oficial es launch_citylearn_v3_official_training.ps1. "
        "El puente multicriterio (scripts/run_madrl_multicriteria_selection.py; artefactos en "
        "outputs/madrl_multicriteria_selection/) y la batería no paramétrica se documentan aquí como "
        "contrato de implementación; las tablas numéricas y figuras de Pareto/ranking se interpretan "
        "en el Capítulo 5. El pipeline Word tools/thesis/run_thesis_word_pipeline.py actualiza los "
        "dos canones sin regeneración destructiva por defecto.",
    )
    actions.append("insert_4_13_body")

    return actions


def verify(path: Path) -> dict:
    doc = Document(str(path))
    c4, c5 = chapter_bounds(doc, 4)
    if c4 is None or c5 is None:
        return {"ok": False, "reason": "missing_cap4_bounds"}
    text = "\n".join((p.text or "") for p in doc.paragraphs[c4:c5])
    low = text.lower()
    return {
        "has_4_11_folders": "4.11" in text and "carpetas" in low,
        "has_4_12_validation": "4.12" in text and ("184" in text or "validación" in low or "validacion" in low),
        "has_4_13_repro": "4.13" in text,
        "has_uc3m": "uc3m" in low,
        "heading_4_7_ok": not ("disecuasiexperimental" in low),
        "ok": (
            "4.11" in text
            and "4.12" in text
            and "184" in text
            and "disecuasiexperimental" not in low
        ),
    }


def main() -> int:
    if not TESIS.is_file():
        raise SystemExit(f"Falta Word canónico: {TESIS}")

    # Política 2-Word: editar solo canónicos; backups fuera de docs/ si hace falta.
    backup_dir = REPO / "outputs" / "_word_backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    backup = backup_dir / f"Tesis_before_cap4_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    shutil.copy2(TESIS, backup)

    targets = [TESIS]
    per_file: list[dict] = []
    overall_ok = True
    for path in targets:
        doc = Document(str(path))
        actions = patch_doc(doc)
        try:
            doc.save(str(path))
            saved = path
        except (PermissionError, OSError) as exc:
            raise SystemExit(
                f"No se pudo guardar el Word canónico {path.name} (¿abierto en Word?): {exc}"
            ) from exc
        checks = verify(saved)
        ok = bool(checks.get("ok"))
        overall_ok = overall_ok and ok
        per_file.append(
            {
                "path": str(saved.relative_to(REPO)),
                "backup": str(backup.relative_to(REPO)),
                "actions": actions,
                "checks": checks,
                "ok": ok,
            }
        )

    report = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "source_md": "docs/tesis_capitulos/Capitulo_4_Desarrollo_Propuesta.md",
        "pytest_summary": {
            "command": r".\.venv39-citylearn-v3\Scripts\python.exe -B -m pytest tests/uc3m/ tests/citylearn_v3/ -v --tb=line",
            "passed": 184,
            "failed": 0,
            "warnings": 1,
            "duration_s": 91.09,
            "date": "2026-07-29",
        },
        "files": per_file,
        "ok": overall_ok,
    }
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if overall_ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
