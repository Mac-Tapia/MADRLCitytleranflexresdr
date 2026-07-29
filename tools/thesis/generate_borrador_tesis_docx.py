# -*- coding: utf-8 -*-
r"""
Generador del Borrador de Tesis (.docx) — MADRL CityLearn v3 Iquitos
====================================================================
Construye el documento Word `docs/Borrador_Tesis_MADRL_CityLearn_Iquitos.docx`
siguiendo la estructura de `docs/informedetesis.txt` (6 capitulos + referencias APA),
con portada, indice (campo TOC de Word), encabezados jerarquicos, tablas reales y
referencias en formato APA.

El contenido se redacta a partir de la documentacion real del proyecto
(docs/, configs, dataset, codigo fuente y resultados en outputs/).

Uso:
    .\.venv39-citylearn-v3\Scripts\python.exe -B tools\thesis\generate_borrador_tesis_docx.py
"""

from __future__ import annotations

import datetime as _dt
import sys
from pathlib import Path

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))
if str(REPO) not in sys.path:
    sys.path.insert(0, str(REPO))

from thesis_references_apa import load_all_thesis_references, reference_stats  # noqa: E402

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Rutas
# ---------------------------------------------------------------------------
OUT_PATH = REPO / "outputs" / "_word_backups" / "Borrador_Tesis_MADRL_CityLearn_Iquitos.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # azul oscuro
GREY = RGBColor(0x59, 0x59, 0x59)


# ---------------------------------------------------------------------------
# Helpers de estilo
# ---------------------------------------------------------------------------
def set_cell_background(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_toc(document) -> None:
    """Inserta un campo TOC de Word (se actualiza con F9 / al abrir)."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldText = OxmlElement("w:t")
    fldText.text = "Actualice este indice con clic derecho > Actualizar campos (o F9)."
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(fldText)
    run._r.append(fldChar_end)


def style_base(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 11.5, ACCENT),
    ):
        st = document.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True


def p(document, text: str, *, italic=False, bold=False, align=None, size=None, color=None, space_after=None):
    para = document.add_paragraph()
    run = para.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    para.alignment = align if align else WD_ALIGN_PARAGRAPH.JUSTIFY
    return para


def bullet(document, text: str):
    para = document.add_paragraph(style="List Bullet")
    para.add_run(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para


def numbered(document, text: str):
    para = document.add_paragraph(style="List Number")
    para.add_run(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para


def heading(document, text: str, level: int):
    h = document.add_heading(text, level=level)
    return h


def add_table(document, headers, rows, *, caption=None, col_widths=None, font_size=9):
    if caption:
        cap = document.add_paragraph()
        r = cap.add_run(caption)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = GREY
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(str(htext))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr[i], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    document.add_paragraph()
    return table


def status_note(document, text: str):
    """Marca un placeholder de informacion pendiente."""
    para = document.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return para


# ---------------------------------------------------------------------------
# Construccion del documento
# ---------------------------------------------------------------------------
def build(*, max_chapter: int | None = None):
    doc = Document()
    style_base(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    fecha = _dt.date.today().strftime("%d de %B de %Y")

    # ===================== PORTADA =====================
    for _ in range(2):
        doc.add_paragraph()
    p(doc, "UNIVERSIDAD NACIONAL DE INGENIERIA (UNI)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, color=ACCENT)
    p(doc, "Unidad de Posgrado [por confirmar: Seccion de Posgrado FIEE / FISI]", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=GREY)
    for _ in range(2):
        doc.add_paragraph()
    p(doc,
      "MULTI-AGENTE DE APRENDIZAJE POR REFUERZO PROFUNDO PARA LA GESTION "
      "COORDINADA DE FLEXIBILIDAD ENERGETICA, EMISIONES DE CARBONO Y COSTOS "
      "ENERGETICOS EN COMUNIDADES INTELIGENTES",
      align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=17, color=ACCENT)
    doc.add_paragraph()
    p(doc, "Caso de estudio: Sistema Electrico Aislado de Iquitos (SEAI) — "
           "17 edificios institucionales/comerciales reales, Loreto, Peru (2023-2025)",
      align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11.5)
    for _ in range(2):
        doc.add_paragraph()
    p(doc, "BORRADOR DE TESIS DE MAESTRIA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, color=RGBColor(0xC0, 0x00, 0x00))
    p(doc, "(Documento de avance — resultados preliminares)", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10.5, color=GREY)
    for _ in range(3):
        doc.add_paragraph()
    p(doc, "Autor: Mac Tapia (mac.tapia.c@uni.pe)", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    p(doc, "Asesor: [por definir]", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    p(doc, "Modalidad: Tesis de Maestria Profesionalizante", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    p(doc, "Linea: Inteligencia Artificial aplicada a Ingenieria Energetica — Sistemas Electricos Inteligentes",
      align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    for _ in range(2):
        doc.add_paragraph()
    p(doc, f"Lima / Iquitos, Peru — {fecha}", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=GREY)

    doc.add_page_break()

    # ===================== RESUMEN =====================
    heading(doc, "Resumen", 1)
    p(doc,
      "Esta tesis determina, mediante simulacion computacional, cual algoritmo "
      "Multi-Agente de Aprendizaje por Refuerzo Profundo (MADRL) gestiona de forma "
      "coordinada y simultanea la flexibilidad energetica, las emisiones de CO2 y los "
      "costos energeticos en una comunidad inteligente. El estudio compara cuatro "
      "algoritmos cooperativos —HAPPO, MASAC, MATD3 y MAAC— bajo una formulacion "
      "Descentralizada Parcialmente Observable (Dec-POMDP) y un esquema de Entrenamiento "
      "Centralizado con Ejecucion Descentralizada (CTDE), implementados sobre una capa "
      "experimental propia (CityLearn v3 propuesto) construida encima del simulador "
      "CityLearn v2. El caso de estudio es el Sistema Electrico Aislado de Iquitos "
      "(SEAI), con 17 edificios institucionales y comerciales reales y un dataset horario "
      "de 26 304 pasos (2023-2025). Se definen tres escenarios de recompensa multiobjetivo "
      "(E1 flexibilidad, E2 CO2, E3 costos) y una matriz de 12 corridas oficiales "
      "(4 algoritmos x 3 escenarios). En la corrida de referencia v4, MATD3 obtuvo el mejor "
      "desempeno global (score 0.7445), con diferencias estadisticamente significativas "
      "(Kruskal-Wallis p = 0.0459). Los resultados son preliminares y se reportan junto con "
      "su comparacion frente a la linea base de CityLearn v2.")
    p(doc, "Palabras clave: aprendizaje por refuerzo multiagente, Dec-POMDP, CTDE, "
           "CityLearn, flexibilidad energetica, emisiones de CO2, costos energeticos, "
           "microrred aislada, Iquitos.", italic=True)

    doc.add_page_break()

    # ===================== INDICE =====================
    heading(doc, "Indice", 1)
    add_toc(doc)
    doc.add_page_break()

    # ===================== CAP 1 =====================
    heading(doc, "Capitulo 1. Introduccion", 1)

    heading(doc, "1.1 Problema de investigacion", 2)
    p(doc,
      "Las comunidades inteligentes (smart communities) integran recursos de energia "
      "distribuida (DER): generacion solar fotovoltaica (PV), sistemas de almacenamiento "
      "en baterias (BESS) y estaciones de carga de vehiculos electricos (EV). La "
      "coordinacion multiagente de estos recursos bajo observabilidad parcial es un "
      "problema de decision secuencial no resuelto que afecta simultaneamente tres "
      "dimensiones criticas: la flexibilidad energetica, las emisiones de carbono y los "
      "costos energeticos.")
    p(doc,
      "En el Sistema Electrico Aislado de Iquitos (SEAI), operado por Electro Oriente "
      "S.A., el suministro depende mayoritariamente de generacion diesel con penetracion "
      "solar creciente (~15%). El factor de emision base es 0.790 kgCO2/kWh (MINAM RAGEI "
      "2019) y la tarifa por uso horario (TOU) distingue punta (18:00-22:59, 0.38 USD/kWh) "
      "de fuera de punta (0.26 USD/kWh). La ausencia de control coordinado de los DER "
      "deriva en picos de demanda elevados, consumo en horas de alta intensidad de carbono "
      "y costos energeticos innecesarios.")
    p(doc,
      "El estado del arte reporta evaluaciones aisladas de algoritmos individuales sobre "
      "dimensiones unicas. No existe un marco comparativo unificado que cubra HAPPO, MASAC, "
      "MATD3 y MAAC bajo formulacion Dec-POMDP y esquema CTDE para determinar el mejor "
      "agente MADRL en la gestion coordinada y simultanea de las tres dimensiones.")
    p(doc, "Problema general (PG):", bold=True)
    p(doc,
      "¿En que medida el algoritmo Multi-Agente de Aprendizaje por Refuerzo Profundo aplicado "
      "a una comunidad inteligente (variable independiente) produce un efecto diferenciado "
      "sobre la gestion coordinada de la flexibilidad energetica, las emisiones de CO2 y los "
      "costos energeticos (variable dependiente), y cual de los algoritmos comparados genera "
      "el mayor efecto?", italic=True)
    p(doc, "Problemas especificos:", bold=True)
    bullet(doc,
           "PE.1: ¿En que medida el algoritmo MADRL (VI) produce un efecto sobre la dimension "
           "de flexibilidad energetica de la comunidad (D-VD.1), y cual algoritmo genera el mayor efecto?")
    bullet(doc,
           "PE.2: ¿En que medida el algoritmo MADRL (VI) produce un efecto sobre la dimension "
           "de emisiones de CO2 de la comunidad (D-VD.2), y cual algoritmo genera el mayor efecto?")
    bullet(doc,
           "PE.3: ¿En que medida el algoritmo MADRL (VI) produce un efecto sobre la dimension "
           "de costos energeticos de la comunidad (D-VD.3), y cual algoritmo genera el mayor efecto?")
    p(doc,
      "Cada problema especifico PE.i se responde en el Capitulo 5 (secciones 5.3.1, 5.4.1 y 5.5.1) "
      "mediante analisis descriptivo e inferencial sobre la corrida canonica de 50 episodios, y se "
      "alinea uno a uno con OE.i y HE.i segun la Tabla 1.1 (PE.1→OE.1→HE.1; PE.2→OE.2→HE.2; "
      "PE.3→OE.3→HE.3).")

    heading(doc, "1.2 Objetivos", 2)
    p(doc, "Objetivo general (OG):", bold=True)
    p(doc,
      "Determinar el efecto del algoritmo MADRL aplicado a una comunidad inteligente (VI) "
      "sobre la gestion coordinada de la flexibilidad energetica, las emisiones de CO2 y los "
      "costos energeticos (VD), e identificar el algoritmo que produce el mayor efecto "
      "coordinado.", italic=True)
    p(doc, "Objetivos especificos:", bold=True)
    bullet(doc,
           "OE.1: Determinar el efecto del algoritmo MADRL (VI) sobre la flexibilidad energetica "
           "(D-VD.1) e identificar el algoritmo de mayor efecto en esta dimension.")
    bullet(doc,
           "OE.2: Determinar el efecto del algoritmo MADRL (VI) sobre las emisiones de CO2 "
           "(D-VD.2) e identificar el algoritmo de mayor efecto en esta dimension.")
    bullet(doc,
           "OE.3: Determinar el efecto del algoritmo MADRL (VI) sobre los costos energeticos "
           "(D-VD.3) e identificar el algoritmo de mayor efecto en esta dimension.")
    p(doc,
      "Los objetivos (OG, OE.1–OE.3) determinan el efecto del MADRL (VI) sobre cada dimension "
      "de la VD e identifican el algoritmo de mayor efecto mediante KPIs y ranking descriptivo "
      "(Capitulos 5 y 6). Las hipotesis (HG, HE.1–HE.3) son enunciados de contraste estadistico "
      "evaluados en la seccion 5.9; no deben confundirse con la redaccion de los objetivos. "
      "La cadena PG→OG→OE y PE.i→OE.i→HE.i se operacionaliza con el diseño factorial 4×3 "
      "(12 tratamientos), los 54 KPI oficiales de CityLearn v2 y la corrida canonica "
      "Colab/Drive madrl_v3_20260627_164047 (50 episodios).")

    heading(doc, "1.3 Hipotesis", 2)
    p(doc,
      "El estudio es cuantitativo, aplicado y explicativo, basado en simulacion experimental. "
      "A diferencia de los objetivos, las hipotesis formulan contrastes estadisticos sobre el "
      "factor algoritmo (VI) frente a la hipotesis nula de igualdad de distribuciones de "
      "KPI-gains. El protocolo inferencial (alpha = 0,05) se detalla en el Capitulo 3 y se "
      "resuelve en la seccion 5.9 (Colas et al., 2019; Agarwal et al., 2021):")
    p(doc, "Hipotesis general (HG):", bold=True)
    p(doc,
      "La aplicacion del algoritmo MADRL a la comunidad inteligente (VI) produce un efecto "
      "estadisticamente significativo y diferenciado sobre la gestion coordinada de la "
      "flexibilidad energetica, las emisiones de CO2 y los costos energeticos (VD).", italic=True)
    p(doc, "Hipotesis especificas:", bold=True)
    bullet(doc,
           "HE.1: El algoritmo MADRL (VI) produce un efecto significativo sobre la flexibilidad "
           "energetica (D-VD.1); el mayor efecto corresponde al algoritmo con menor variabilidad "
           "en los KPI de pico y rampa.")
    bullet(doc,
           "HE.2: El algoritmo MADRL (VI) produce un efecto significativo sobre las emisiones "
           "de CO2 (D-VD.2); el mayor efecto corresponde a MATD3.")
    bullet(doc,
           "HE.3: El algoritmo MADRL (VI) produce un efecto significativo sobre los costos "
           "energeticos (D-VD.3); el mayor efecto corresponde al algoritmo con menor delta "
           "de costo electrico.")
    p(doc,
      "Cada hipotesis especifica tiene una hipotesis nula asociada (sin diferencias significativas "
      "entre niveles del factor algoritmo). La decision inferencial se reporta en la seccion 5.9.5; "
      "el cumplimiento de OG y OE.1–OE.3 se reporta por separado en la seccion 5.11.")

    heading(doc, "1.4 Matriz de consistencia y operacionalizacion", 2)
    add_table(
        doc,
        ["Elemento", "PG / PE", "OG / OE", "Hipotesis", "Dimension VD"],
        [
            ["General", "PG: efecto MADRL coordinado", "OG: determinar efecto e identificar mayor efecto", "HG", "Flexibilidad + CO2 + costos"],
            ["Flexibilidad", "PE.1", "OE.1", "HE.1", "D-VD.1: peak, ramp, load factor"],
            ["Emisiones", "PE.2", "OE.2", "HE.2", "D-VD.2: carbon_emissions, delta CO2"],
            ["Costos", "PE.3", "OE.3", "HE.3", "D-VD.3: electricity_cost, delta costo"],
        ],
        caption="Tabla 1.1. Matriz de consistencia problema-objetivo-hipotesis.",
        col_widths=[2.5, 4.5, 5.0, 2.0, 4.0],
    )
    add_table(
        doc,
        ["Variable", "Codigo", "Niveles / medicion", "Control"],
        [
            ["VI tipo algoritmo", "D-VI.1", "HAPPO, MASAC, MATD3, MAAC", "Misma recompensa unified_comparable_v4"],
            ["VI escenario", "D-VI.2", "E1, E2, E3 (pesos multiobjetivo)", "Mismo dataset y semilla"],
            ["VD flexibilidad", "D-VD.1", "18 KPI OE1 (peak, ramp, etc.)", "Baseline CityLearn v2"],
            ["VD emisiones", "D-VD.2", "18 KPI OE2 (carbon_emissions, etc.)", "CI SEAI 0,672-0,790"],
            ["VD costos", "D-VD.3", "18 KPI OE3 (electricity_cost, etc.)", "TOU 0,26/0,38 USD/kWh"],
            ["Control", "-", "Dataset, clima, DER, hiperparametros", "seed=0; 50 ep Colab canonico"],
        ],
        caption="Tabla 1.2. Operacionalizacion de variables independientes y dependientes.",
        col_widths=[3.5, 2.0, 5.5, 5.0],
    )
    heading(doc, "1.4.1 Criterios de determinacion del impacto", 3)
    p(
        doc,
        "Para cumplir el OG y los OE.1–OE.3, y para demostrar las hipotesis, se exige el conjunto "
        "completo de criterios C1–C5 (sin parciales). C5 (control de recursos) es obligatorio. "
        "Cada eje se reporta a nivel distrito y a nivel edificio (evidencia Cap. 5).",
    )
    add_table(
        doc,
        ["Id", "Criterio", "Medida / prueba", "Rol"],
        [
            ["C1", "Impacto vs baseline", "Wilcoxon KPI-gains vs cero + Holm", "Inferencial HE"],
            ["C2", "Diferencias entre algoritmos", "Kruskal-Wallis / Friedman + Holm", "Inferencial HE"],
            ["C3", "KPIs fisicos de distrito por eje", "flex_composite / delta CO2 / delta costo", "Descriptivo distrito"],
            ["C4", "KPIs desagregados por edificio por eje", "17 edificios x E1/E2/E3", "Descriptivo edificio"],
            ["C5", "Control de recursos", "BESS, EV/V2G, carga desplazable", "Obligatorio (atribuibilidad)"],
        ],
        caption="Tabla 1.3. Criterios completos de determinacion del impacto (C5 = control de recursos).",
        col_widths=[1.5, 4.0, 6.0, 4.0],
    )

    heading(doc, "1.5 Justificacion", 2)
    bullet(doc, "Tecnica: aporta una evaluacion unificada de HAPPO, MASAC, MATD3 y MAAC bajo "
                "Dec-POMDP y CTDE, avanzando el estado del arte en control energetico cooperativo.")
    bullet(doc, "Ambiental: identificar el mejor MADRL para reduccion de CO2 contribuye a la "
                "descarbonizacion de comunidades grid-interactive, con aplicabilidad al SEAI Iquitos "
                "(factor de emision 0.790 kgCO2/kWh).")
    bullet(doc, "Economica: establecer el mejor MADRL para optimizacion de costos orienta la "
                "reduccion del gasto electrico bajo tarifas TOU.")
    bullet(doc, "Metodologica: la formulacion Dec-POMDP, el esquema CTDE y el benchmark unificado "
                "con CityLearn v3 propuesto, MARLlib y Optuna constituyen una contribucion reproducible.")
    bullet(doc, "Cientifica y social: llena una laguna en la literatura comparativa de MADRL y "
                "beneficia a usuarios institucionales y a la transicion energetica comunitaria.")

    heading(doc, "1.6 Alcances y limitaciones", 2)
    p(doc, "Alcances:", bold=True)
    bullet(doc, "Tematico: comparacion de HAPPO, MASAC, MATD3 y MAAC en KPIs de flexibilidad, CO2 y costos.")
    bullet(doc, "Espacial: comunidades inteligentes simuladas en CityLearn v2 / CityLearn v3 propuesto, "
                "con aplicabilidad al SEAI Iquitos (17 edificios reales).")
    bullet(doc, "Temporal: dataset horario 2023-2025 (26 304 pasos) y literatura 2015-2026.")
    bullet(doc, "Metodologico: estudio cuantitativo, aplicado, comparativo y cuasiexperimental, basado en simulacion.")
    bullet(doc, "Computacional: Python 3.9, PyTorch 2.8.0+cu126, CityLearn v2, MARLlib, Optuna, "
                "Gymnasium, PettingZoo; hardware local NVIDIA RTX 4060 Laptop 8 GB y Colab A100.")
    p(doc, "Limitaciones:", bold=True)
    bullet(doc, "No se modela una red electrica fisica; los resultados de simulacion no constituyen "
                "validacion de despliegue real.")
    bullet(doc, "CityLearn v3 propuesto es una extension experimental de tesis, no una version oficial de CityLearn.")
    bullet(doc, "Se excluyen el despliegue en campo en tiempo real, la investigacion con sujetos humanos, "
                "el despacho economico de unidades fisicas y el analisis de estabilidad de red.")

    # ===================== CAP 2 =====================
    doc.add_page_break()
    heading(doc, "Capitulo 2. Marco teorico", 1)

    heading(doc, "2.1 Estado del arte actualizado", 2)
    p(doc,
      "La revision sistematica (Modulo A) comprende 50 investigaciones verificadas, "
      "organizadas en cuatro ejes alineados con los objetivos de la tesis: (1) flexibilidad "
      "energetica con MADRL, (2) reduccion de emisiones de CO2 con MADRL, (3) optimizacion "
      "de costos energeticos con MADRL y (4) marco tecnico MADRL transversal.")
    p(doc, "Eje 1 — Flexibilidad energetica con MADRL:", bold=True)
    p(doc,
      "Vazquez-Canteli y Nagy (2019a) introducen CityLearn como entorno OpenAI Gym para "
      "respuesta a la demanda multiedificio; Vazquez-Canteli et al. (2020) estandarizan "
      "KPIs comparables (peak_average, ramping_average, one_minus_load_factor_average). "
      "Nweye et al. (2024) desarrollan CityLearn v2 con EV/V2G, intensidad de carbono "
      "dinamica, BESS, PV y KPIs de flexibilidad, CO2 y costos. Nweye et al. (2022) "
      "identifican nueve desafios reales del MARL en edificios grid-interactive; Nweye et "
      "al. (2023a) presentan MERLIN para comunidades de 17 edificios reales, escala "
      "equivalente al SEAI Iquitos. Yao et al. (2023) y Xie et al. (2023) demuestran "
      "mejoras de coordinacion con MARL cooperativo y mecanismos de atencion; Hribar et al. "
      "(2025) reportan +20% de autonomia energetica en distritos de energia positiva.")
    p(doc, "Eje 2 — Reduccion de emisiones de CO2 con MADRL:", bold=True)
    p(doc,
      "Liu et al. (2022) proponen MADDPG multiobjetivo con reducciones de ~15% de CO2 y "
      "~20% de costo; Ye et al. (2025) y Ma et al. (2025) introducen MARL seguro con "
      "restricciones de carbono en redes de distribucion y multi-microgrids; Sarkar et al. "
      "(2024) aplican desplazamiento temporal de carga hacia periodos de baja intensidad de "
      "carbono, directamente transferible al OE.2.")
    p(doc, "Eje 3 — Optimizacion de costos energeticos con MADRL:", bold=True)
    p(doc,
      "Yao et al. (2023) logran ~18% de reduccion de costo con LSD-MADDPG; Shojaeighadikolaei "
      "et al. (2022) ~22% frente a configuraciones no cooperativas; Gao et al. (2023) validan "
      "MASAC para programacion colaborativa multi-microgrid con respuesta a precios; Xiong et "
      "al. (2024) y Kim et al. (2025) abordan tarifas TOU y control de BESS en tiempo real.")
    p(doc, "Eje transversal — Marco tecnico MADRL:", bold=True)
    p(doc,
      "Lowe et al. (2017) establecen el paradigma CTDE (MADDPG); Oliehoek y Amato (2016) "
      "formalizan el Dec-POMDP; Kuba et al. (2021) presentan HAPPO con garantias de mejora "
      "monotona para agentes heterogeneos; Iqbal y Sha (2019) introducen MAAC con atencion "
      "multi-cabeza; Hu et al. (2023) presentan MARLlib; Akiba et al. (2019) presentan "
      "Optuna; Haarnoja et al. (2018) presentan SAC, base de MASAC y MAAC.")
    p(doc, "Antecedentes nacionales y de sistemas aislados:", bold=True)
    p(doc,
      "En el contexto peruano, el SEAI Iquitos concentra desafios de redes aisladas diesel-PV "
      "con intensidad de carbono y tarifas reguladas (MINAM, 2019; OSINERGMIN, 2024). "
      "Chevarria Moscoso (2024), en tesis doctoral de la UNI, aborda optimizacion de generacion "
      "hidroelectrica con metodos estocasticos; Peñalva Sanchez (2024), tambien en la UNI, "
      "combina sistemas fotovoltaicos hibridos con inteligencia artificial para prediccion de "
      "demanda. En el ambito latinoamericano, Rosero Bernal (2024) propone administracion de "
      "energia autonoma en microredes desde la nube. A nivel internacional, Domínguez Barbero "
      "(2026) demuestra DRL para microrredes aisladas con TD3, antecedente directo para el "
      "control de DER bajo incertidumbre operativa sin depender de pronosticos perfectos.")

    heading(doc, "2.2 Bases teoricas", 2)
    heading(doc, "2.2.1 Flexibilidad energetica en comunidades inteligentes", 3)
    p(doc,
      "La flexibilidad energetica es la capacidad del sistema para modificar su curva de "
      "demanda mediante desplazamiento de cargas, almacenamiento en BESS, autoconsumo solar "
      "y carga/descarga de EV. KPIs relevantes: peak_average (reduccion de pico), "
      "ramping_average (suavizado), one_minus_load_factor_average (factor de carga), "
      "autoconsumo, autosuficiencia, reduccion de importacion de red y utilizacion de renovables.")
    heading(doc, "2.2.2 Emisiones de carbono en comunidades inteligentes", 3)
    p(doc,
      "La intensidad de carbono (CI) varia con la mezcla de generacion horaria. En el SEAI "
      "Iquitos CI(t) = 0.790 x (1 - 0.15 x GHI(t)/1000) kgCO2/kWh, en el rango "
      "[0.672, 0.790]. La gestion consciente de la CI desplaza el consumo hacia periodos de "
      "menor emision. KPIs: emisiones totales, reduccion frente a baseline, consumo "
      "ponderado por CI y emisiones evitadas.")
    heading(doc, "2.2.3 Costos energeticos en comunidades inteligentes", 3)
    p(doc,
      "Las tarifas TOU y los precios en tiempo real crean incentivos para la flexibilidad "
      "economica. El arbitraje tarifario mediante BESS/EV reduce el cargo por demanda. KPIs: "
      "costo total de electricidad, reduccion frente a baseline, reduccion del cargo por "
      "demanda y desviacion respecto a la senal de precio.")
    heading(doc, "2.2.4 Marco tecnico MADRL", 3)
    add_table(
        doc,
        ["Concepto", "Definicion sintetica"],
        [
            ["Dec-POMDP", "Tupla <S, {Ai}, T, R, {Oi}, Z, gamma, T> para decision cooperativa bajo observabilidad parcial (Oliehoek y Amato, 2016)."],
            ["CTDE", "Entrenamiento centralizado (critico accede al estado global) con ejecucion descentralizada (politica usa solo la observacion local) (Lowe et al., 2017)."],
            ["HAPPO", "On-policy basado en PPO, actualizacion secuencial por agente, critico centralizado; mejora monotona para agentes heterogeneos (Kuba et al., 2021)."],
            ["MASAC", "Off-policy SAC + mezcla cooperativa tipo QMIX con redes RNN para observaciones parciales (Haarnoja et al., 2018; Gao et al., 2023)."],
            ["MATD3", "Off-policy con doble critico centralizado (anti sobreestimacion), policy delay y target noise (Fujimoto et al., 2018)."],
            ["MAAC", "Off-policy con critico de atencion multi-cabeza para coordinacion selectiva entre agentes (Iqbal y Sha, 2019)."],
            ["MARLlib", "Biblioteca unificada de algoritmos MADRL compatible con Ray/RLlib, Gymnasium y PettingZoo (Hu et al., 2023)."],
            ["CityLearn v2", "Entorno de simulacion multiagente para comunidades grid-interactive con DER, EV, CI y precios (Nweye et al., 2024)."],
            ["Optuna", "Framework de optimizacion automatica de hiperparametros basado en TPE (Akiba et al., 2019)."],
        ],
        caption="Tabla 2.1. Conceptos del marco tecnico MADRL.",
        col_widths=[3.0, 13.0],
    )

    heading(doc, "2.2.5 CityLearn v3 propuesto — capa experimental MADRL", 3)
    p(doc,
      "CityLearn v3 propuesto es la extension experimental de esta tesis sobre CityLearn v2 "
      "(Nweye et al., 2024). No existe como version oficial externa: conserva el simulador "
      "fisico y los KPI evaluate_v2, y agrega el contrato Dec-POMDP/CTDE, objetivos OE.1–OE.3, "
      "recompensa multiobjetivo y adaptadores a backends MADRL oficiales en external/. "
      "Fuentes: CityLearn/CITYLEARN_V3_MADRL.md, CityLearn/citylearn/v3/, "
      "docs/architecture/ARQUITECTURA_Y_FLUJO_TRABAJO_CITYLEARN_V3_MADRL.md y Módulo D del "
      "skill agent-skills/madrl-citylearn-thesis-plan.")
    p(doc,
      "Distincion UC3M (7 ejes) vs capa v3 ejecutada (3 ejes). El sustento formal "
      "agent-skills/madrl-sustento-doc-capa v3/madrl-modeladomatematico.md y el paquete "
      "uc3m/ (RewardAxes, HPHI, BACT, EmpiricalProtocol) se axiomatizan en §§2.2.6–2.2.9 "
      "como operador holistico de siete ejes (CO2, costo, flexibilidad, confort, "
      "degradacion BESS, resiliencia, ACS). La implementacion empirica —y las 12 corridas "
      "canónicas— operacionaliza solo tres ejes OE.1/OE.2/OE.3 via "
      "CityLearnV3MADRLRewardFunction (perfil unified_comparable_v4). Los ejes 4–7 se "
      "leen como sustento arquitectural/matematico, no como resultados ejecutados.")
    p(doc,
      "Modulos reales de la capa: v3/environment.py (entorno/fabrica), v3/objectives.py "
      "(OE.1 flexibilidad, OE.2 CO2, OE.3 costos), v3/config.py, v3/backends.py, "
      "v3/marllib_env.py (referencia MARLlib; Hu et al., 2023), y "
      "CityLearnV3MADRLRewardFunction en reward_function.py con perfil unified_comparable_v4. "
      "Wrappers: CityLearnHARLEnv (HAPPO), CityLearnSMACDiscreteEnv (MASAC), "
      "CityLearnOffPolicyVecEnv (MATD3), CityLearnMAACVecEnv (MAAC), orquestados por "
      "citylearn_v3_training_common.py.")
    p(doc,
      "Formalizacion: N=17 agentes-edificio; estado global CTDE por concatenacion de "
      "observaciones locales (≈1 856 dimensiones en entorno cargado); gamma=0,9999; "
      "T=8 760 pasos/episodio; team_reward_ratio=0,70. La operacionalizacion E1/E2/E3 "
      "(pesos flex/carbon/cost) se desarrolla en Caps. 3–4; la evidencia empirica de las "
      "12 corridas canónicas (madrl_v3_20260627_164047) se reporta en Cap. 5 §5.0–5.7.")
    add_table(
        doc,
        ["Concepto teorico (Cap. 2)", "Desarrollo (Cap. 4)", "Evidencia (Cap. 5)"],
        [
            ["Dec-POMDP / CTDE", "§§4.3–4.5", "§5.1 cobertura 4×3; HAPPO KPI-gains 4/4"],
            ["Capa v3 3 ejes + UC3M 7 ejes (sustento)", "Pasos 5–7; §§4.1/4.3–4.4", "§5.1.3 OE→E; no HPHI 7-D ejecutado"],
            ["Cuatro backends MADRL", "Pasos 8–9; §§4.5–4.6", "§§5.3–5.5 por OE"],
            ["Evaluacion v2 + multi-semilla (diseño)", "Pasos 10–13; §4.8", "§5.2 Shapiro→no parametrico; smoke n=3"],
        ],
        caption="Tabla 2.1b. Lectura vertical Cap. 2 → Cap. 4 → Cap. 5 (capa v3).",
        col_widths=[5.0, 5.5, 5.5],
    )

    heading(doc, "2.2.6 Formalizacion matematica del Meta-Dec-POMDP UC3M", 3)
    p(doc,
      "Fuente: agent-skills/madrl-sustento-doc-capa v3/madrl-modeladomatematico.md "
      "(adaptado a Cap. 2). El Meta-Dec-POMDP UC3M es la tupla 11-aria "
      "M_UC3M = <I, S, A, O, T, R, Z, gamma, H, b0, Lambda>, que extiende Oliehoek y "
      "Amato (2016) con vector de recompensas R=(r^(1),...,r^(7)) y simplex Lambda. "
      "En Iquitos N=17, gamma=0,9999, H=8 760. La ejecucion restringe Lambda al "
      "subsimplex OE.1/OE.2/OE.3 con r_team=0,70 (perfil unified_comparable_v4).")
    p(doc,
      "Definicion 2.1 (BACT): B en R^{N x Ka x Kc x Kb} codifica activos, clima y "
      "constructivo (uc3m/env/bact.py). Observaciones parciales incluyen carga, termica, "
      "PV, SoC, disponibilidad EV, tarifas, CI y meteorologia; acciones continuas "
      "acotadas en [-1,1]^{d_ai} con d_ai heterogeneo. Prop. 2.1: S compacto. Lema 2.2: "
      "R escalarizada Borel-medible y acotada. Nucleo T factorizado exogeno/endogeno.")

    heading(doc, "2.2.7 Operador de recompensa holistico, Pareto y convergencia", 3)
    p(doc,
      "Definicion 2.2: R_i = -sum_k lambda_k * r_i~(k) con normalizacion a base RBC. "
      "Teorema 2.3: V_pi existe y |V|<=M/(1-gamma). Prop. 2.4: Lipschitz con "
      "L=sum lambda_k L_k. Teorema 2.5: existencia de frontera de Pareto si Pi compacto "
      "(Roijers et al., 2013). Definicion 2.4 (HPHI): hipervolumen 7-D normalizado "
      "(uc3m/reward/hphi.py); esta tesis no ejecuta HPHI 7-D en Cap. 5. Prop. 2.6: CTDE "
      "mitiga no-estacionariedad (Lowe et al., 2017); HAPPO/HATRPO bajo Zhong et al. (2023).")

    heading(doc, "2.2.8 Modelado fisico-matematico de los ejes operacionales", 3)
    p(doc,
      "Ejes ejecutados: UC3M-1/OE.2 (CO2 con factor marginal y desplazamiento EV), "
      "UC3M-2/OE.3 (costo TOU + potencia), UC3M-3/OE.1 (ramping y picos vs umbral DSO; "
      "KPIs peak/ramping/load_factor). Ejes de sustento (no Cap. 5): 4 confort "
      "adaptativo De Dear-Brager, 5 degradacion Arrhenius-SEI/Peukert, 6 resiliencia "
      "isla (CCI/LOLP), 7 ACS con perdidas UA. Balance Kirchhoff: "
      "P_net = P_load + P_HP + P_EH + P_BESS + P_EV - P_PV - P_wind.")

    heading(doc, "2.2.9 Arquitectura MARLlib-CTDE y universalidad algoritmica (sustento)", 3)
    p(doc,
      "MARLlib (Hu et al., 2023) y el plugin algoritmico UC3M "
      "P=<Theta, Phi, L_actor, L_critic, U_step, B> sustentan universalidad via "
      "AlgorithmFactory. CTDE: pi_i(a_i|o_i) descentralizado y Q_centr(s,a) centralizado. "
      "HAPPO/MASAC/MATD3/MAAC (§2.2.4) son instancias; el launcher de 12 corridas usa "
      "wrappers CityLearn/scripts/ + external/, no UC3MEnv/MARLlib.")

    heading(doc, "2.3 Trabajos relacionados", 2)
    p(doc,
      "El antecedente directo mas cercano es Nweye et al. (2023b), que evalua HAPPO (HARL) "
      "en comunidades grid-interactive heterogeneas con CityLearn a escala de 17 edificios. "
      "Esta tesis se diferencia por: (i) comparar simultaneamente cuatro algoritmos (HAPPO, "
      "MASAC, MATD3, MAAC) bajo la misma funcion de recompensa unificada; (ii) usar un "
      "dataset real propio del SEAI Iquitos (red aislada diesel+PV, clima ecuatorial, "
      "tarifas locales); y (iii) introducir cuatro aportes originales al motor de simulacion "
      "(degradacion BESS con C-rate y Arrhenius, correccion PV tropical IEC 61215, KPI de "
      "pico con ventana de facturacion OSINERGMIN, y clase CarbonIntensityModel).")
    add_table(
        doc,
        ["Trabajo", "Enfoque", "Resultado / aporte"],
        [
            ["Lowe et al. (2017)", "MADDPG, CTDE", "Paradigma base de critico centralizado y politica descentralizada."],
            ["Kuba et al. (2021)", "HAPPO (HARL)", "Mejora monotona para agentes heterogeneos; supera a MAPPO/IPPO."],
            ["Iqbal y Sha (2019)", "MAAC (atencion)", "Coordinacion selectiva; mejoras de ~15-30% sobre MADDPG/COMA."],
            ["Nweye et al. (2024)", "CityLearn v2", "Entorno base con EV/V2G, CI dinamica y KPIs multiobjetivo."],
            ["Nweye et al. (2023a)", "MERLIN", "MARL en comunidad real de 17 edificios (escala SEAI Iquitos)."],
            ["Yao et al. (2023)", "LSD-MADDPG", "~15% pico y ~18% costo frente a agentes no cooperativos."],
            ["Liu et al. (2022)", "MADDPG multiobjetivo", "~15% CO2 y ~20% costo; recompensa con termino de exportacion."],
            ["Gao et al. (2023)", "MASAC multi-microgrid", "Arbitraje BESS en mercados TOU con respuesta a precios."],
        ],
        caption="Tabla 2.2. Sintesis de trabajos relacionados.",
        col_widths=[3.6, 4.4, 8.0],
    )

    # ===================== CAP 3 =====================
    doc.add_page_break()
    heading(doc, "Capitulo 3. Metodologia", 1)

    heading(doc, "3.1 Tipo de investigacion", 2)
    bullet(doc, "Enfoque: cuantitativo (Hernández-Sampieri & Mendoza, 2018).")
    bullet(doc, "Tipo: aplicada (Arias, 2020; Tamayo y Tamayo, 2004).")
    bullet(doc, "Nivel: descriptivo, comparativo y propositivo.")
    bullet(doc, "Diseno: cuasiexperimental, factorial 4×3 (algoritmo MADRL × escenario E1/E2/E3), basado en simulacion computacional (Campbell & Stanley, 1963; Hernández-Sampieri & Mendoza, 2018).")
    p(doc,
      "El nivel comparativo es esencial: el estudio determina el mejor MADRL comparando "
      "HAPPO, MASAC, MATD3 y MAAC en tres ejes (flexibilidad, CO2, costos). El nivel "
      "propositivo radica en que CityLearn v3 propuesto es una extension arquitectonica "
      "original sobre CityLearn v2.")

    heading(doc, "3.2 Diseno cuasiexperimental factorial 4×3", 2)
    p(doc,
      "El experimento manipula la variable independiente definida en el Capitulo 2, Seccion 2.2.1: "
      "D-VI.1 (tipo de algoritmo: HAPPO, MASAC, MATD3, MAAC), D-VI.2 (escenario de ponderacion: "
      "E1, E2, E3) y D-VI.3 (controles metodologicos). La variable dependiente, definida en "
      "el Capitulo 2, Seccion 2.2.2 (Tabla 2.1), se observa en tres ejes con tres dimensiones "
      "cada uno: D-VD.1 flexibilidad, D-VD.2 emisiones de CO2 y D-VD.3 costos energeticos "
      "(54 KPI oficiales CityLearn v2). El "
      "diseno factorial completo comprende 12 tratamientos (4×3), ejecutados con la misma "
      "funcion de recompensa multiobjetivo. Cada escenario condiciona la politica hacia un "
      "objetivo dominante mediante un vector de pesos distinto, respondiendo a OE.1, OE.2 "
      "y OE.3 por separado. La escalarizacion lineal de objetivos (Roijers et al., 2013) "
      "y el entrenamiento de politicas separadas por vector de pesos (Abels et al., 2019; "
      "Felten et al., 2024) sustentan este diseno.")
    add_table(
        doc,
        ["Escenario", "Objetivo", "Pesos [flex, CO2, costo]", "Eje dominante"],
        [
            ["E1", "OE.1 Flexibilidad", "[0.70, 0.15, 0.15]", "Flexibilidad"],
            ["E2", "OE.2 Emisiones CO2", "[0.15, 0.70, 0.15]", "Carbono"],
            ["E3", "OE.3 Costos", "[0.25, 0.15, 0.60]", "Costos"],
            ["Global", "O.G. Coordinado", "[0.50, 0.25, 0.25]", "Coordinado"],
        ],
        caption="Tabla 3.1. Escenarios de entrenamiento y pesos de recompensa por eje.",
        col_widths=[2.5, 4.5, 5.0, 4.0],
    )
    add_table(
        doc,
        ["Tratamiento", "D-VI.1", "D-VI.2", "D-VD medida", "Corrida canonica"],
        [
            ["T01-T03", "HAPPO", "E1/E2/E3", "54 KPI por eje", "49 ep (sin KPI finales)"],
            ["T04-T06", "MASAC", "E1/E2/E3", "54 KPI por eje", "50/50/50 ep"],
            ["T07-T09", "MATD3", "E1/E2/E3", "54 KPI por eje", "50/50/50 ep"],
            ["T10-T12", "MAAC", "E1/E2/E3", "54 KPI por eje", "50/50/50 ep"],
        ],
        caption="Tabla 3.1b. Diseño factorial 4×3: 12 tratamientos (algoritmo × escenario).",
        col_widths=[2.5, 2.5, 2.5, 4.0, 4.5],
    )

    heading(doc, "3.3 Datos utilizados — Dataset citylearn_iquitos_2023_2025", 2)
    p(doc,
      "El dataset de tesis se construye integramente desde datos primarios mediante la "
      "orquestacion tools/dataset/orchestrate_citylearn_dataset.py (que invoca generate_iquitos_dataset.py "
      "y herramientas de dimensionamiento/auditoria). No se adoptan datasets preexistentes de "
      "CityLearn porque el SEAI Iquitos posee condiciones irrepresentables: sistema aislado, "
      "generacion diesel dominante, clima ecuatorial sin calefaccion y tarifas TOU locales.")
    add_table(
        doc,
        ["Componente", "Valor"],
        [
            ["Edificios (agentes MADRL)", "17 (institucionales/comerciales reales)"],
            ["Periodo / pasos horarios", "2023-2025 / 26 304 h (episodio = 8 760 h)"],
            ["Archivos CSV auditados", "222 (0 NaN, 0 Inf)"],
            ["Cargadores EV controlables", "185 tomas Mode 3 / 96 unidades fisicas / 192 sockets"],
            ["Potencia EV nominal", "749.4 kW (31 con V2G; pool de 1 850 EV)"],
            ["BESS total", "26 266 kWh / 6 648 kW"],
            ["PV total", "48 790.9 kWp"],
            ["Maquinas controladas", "17 lavadoras (~876.6 MWh/ano)"],
            ["Intensidad de carbono", "0.672-0.790 kgCO2/kWh (MINAM RAGEI 2019 + modulacion PV)"],
            ["Precios TOU", "0.26 USD/kWh fuera punta; 0.38 USD/kWh punta (18-22 h)"],
            ["Ubicacion", "lat -3.7491, lon -73.2538, alt 106 m, tz America/Lima"],
            ["State dim global (cargado)", "1856 (E1/E2/E3); 17 agentes con EV en obs/accion"],
        ],
        caption="Tabla 3.2. Caracteristicas del dataset citylearn_iquitos_2023_2025.",
        col_widths=[6.0, 10.0],
    )
    p(doc,
      "Fuentes de entrada: (a) meteorologia horaria PVGIS-ERA5 (2023, via pvlib) y NASA "
      "POWER (2024-2025); (b) consumo real destilado desde facturacion mensual de Electro "
      "Oriente S.A. (tools/dataset/distill_building_loads.py); (c) senales regulatorias y de mercado "
      "(CI dinamica y tarifas TOU). El pipeline ejecuta 10 etapas reproducibles con control "
      "de semilla y validacion en cada etapa (descarga meteorologica, seleccion de modulo PV "
      "Sandia, generacion PV con pvlib.ModelChain, dimensionamiento BESS por balance, "
      "generacion de Building_X.csv, charger_X_Y.csv, Washing_Machine_X.csv, señales de red, "
      "schema.json y validacion final con CityLearnEnv).")

    p(doc, "Los 17 edificios del SEAI Iquitos y su dimensionamiento DER vigente:", bold=True)
    add_table(
        doc,
        ["ID", "Edificio", "Tipo", "Area m2", "PV kWp", "BESS kWh", "EV kW", "Cargadores"],
        [
            ["B01", "Electro Oriente S.A.", "Office", "14 000", "3 360.2", "6 747", "21.8", "4"],
            ["B02", "Municipalidad San Juan Bautista", "Office", "8 000", "1 920.0", "244", "24.4", "6"],
            ["B03", "Aeropuerto Internacional", "Assembly", "6 000", "1 440.2", "2 363", "37.8", "8"],
            ["B04", "Hipermercados Tottus Oriente", "Retail", "2 500", "600.2", "454", "24.4", "6"],
            ["B05", "Hotel Plaza S.A.", "Hotel", "1 142", "274.1", "234", "14.4", "3"],
            ["B06", "Mall Aventura S.A.", "Mall", "20 637", "4 952.9", "2 541", "119.6", "32"],
            ["B07", "UNAP Facultad de Biologia", "Education", "8 103", "1 944.9", "984", "153.2", "42"],
            ["B08", "PNP Escuela Tecnica Superior", "Military", "21 000", "5 040.2", "601", "73.6", "17"],
            ["B09", "GORE Loreto - COER", "Office_Critical", "4 480", "1 075.3", "138", "37.4", "10"],
            ["B10", "Gobierno Regional de Loreto", "Office", "14 296", "3 431.1", "2 353", "36.6", "6"],
            ["B11", "Hospital Regional de Loreto", "Hospital", "42 649", "10 236.1", "1 901", "14.4", "3"],
            ["B12", "EsSalud", "Healthcare", "18 197", "4 367.5", "4 346", "14.4", "3"],
            ["B13", "UNAP Ciencias Economicas", "Education", "2 723", "653.8", "272", "41.4", "11"],
            ["B14", "Autoridad Portuaria Nacional", "Port", "17 761", "4 262.9", "229", "21.8", "4"],
            ["B15", "DREL Colegio Nacional", "Education", "9 890", "2 373.8", "500", "31.4", "8"],
            ["B16", "SIMA Iquitos S.R.Ltda", "Industrial", "10 294", "2 470.8", "1 622", "41.4", "11"],
            ["B17", "Asociacion Civil Selva Amazonica", "Laboratory", "1 611", "386.9", "737", "41.4", "11"],
            ["Tot", "17 edificios", "-", "-", "48 790.9", "26 266", "749.4", "185"],
        ],
        caption="Tabla 3.3. Edificios reales del SEAI Iquitos y dimensionamiento DER (auditado).",
        col_widths=[1.0, 4.6, 2.2, 1.8, 1.8, 1.8, 1.4, 1.6],
        font_size=8,
    )

    heading(doc, "3.4 Variables", 2)
    p(doc,
      "La definicion conceptual, la triangulacion bibliografica y las tres dimensiones de cada "
      "variable se establecen en el Capitulo 2, Seccion 2.2 (Tabla 2.1). Esta seccion "
      "operacionaliza computacionalmente esa definicion para el diseno experimental factorial "
      "4x3 descrito en la Seccion 3.2.")
    p(doc, "Variable independiente (VI):", bold=True)
    p(doc,
      "Capa MADRL cooperativa implementada sobre CityLearn v2 (CityLearn v3 propuesto), "
      "manipulada en tres dimensiones: D-VI.1 algoritmo (HAPPO, MASAC, MATD3, MAAC), "
      "D-VI.2 escenario (E1/E2/E3) y D-VI.3 control experimental (dataset comun, recompensa "
      "unified_comparable_v4, semilla y protocolo CTDE). Ver definicion teorica en "
      "Capitulo 2, Secciones 2.2.1 y 2.1.2.")
    p(doc, "Variable dependiente (VD):", bold=True)
    p(doc,
      "Desempeno coordinado en tres ejes con tres dimensiones cada uno: D-VD.1 flexibilidad "
      "(pico, rampa, factor de carga/flex_composite), D-VD.2 emisiones de CO2 (totales, "
      "delta, consumo ponderado por CI) y D-VD.3 costos energeticos (total, delta, pico "
      "facturable/price_signal_deviation), medidos con los KPIs oficiales de CityLearn v2. "
      "Ver triangulacion en Capitulo 2, Seccion 2.2.2 y Tabla 2.1.")
    add_table(
        doc,
        ["Eje (OE)", "KPIs principales"],
        [
            ["OE.1 Flexibilidad", "peak_average, ramping_average, one_minus_load_factor_average, autoconsumo, autosuficiencia."],
            ["OE.2 CO2", "carbon_emissions, carbon_emissions_delta, consumo ponderado por CI, emisiones evitadas."],
            ["OE.3 Costos", "electricity_cost, electricity_cost_delta, reduccion de cargo por demanda, price_signal_deviation."],
        ],
        caption="Tabla 3.4. Operacionalizacion de la variable dependiente por eje.",
        col_widths=[3.5, 12.5],
    )
    p(doc,
      "Variables de control: dataset climatico (PVGIS-ERA5/NASA POWER), perfil de demanda "
      "destilado de mediciones reales, CI (0.672-0.790 kgCO2/kWh), precios TOU (0.26/0.38 "
      "USD/kWh), capacidad BESS, penetracion PV, escenario de carga EV e hiperparametros.")

    heading(doc, "3.5 Tecnicas, herramientas e instrumentos", 2)
    p(doc, "Tecnicas:", bold=True)
    bullet(doc, "Revision bibliografica sistematica (matriz de 50 antecedentes, Modulo A).")
    bullet(doc, "Extraccion y preprocesamiento del dataset; registro de metricas de entrenamiento "
                "(reward acumulada y media, pesos por eje, CI, precio, carga neta).")
    bullet(doc, "Evaluacion de KPIs por eje mediante env.evaluate_v2() de CityLearn v2.")
    bullet(doc, "Comparacion de algoritmos: tablas por KPI y eje, ranking integrado y score global.")
    bullet(doc, "Pruebas estadisticas no parametricas: Shapiro-Wilk (normalidad), Kruskal-Wallis "
                "(diferencias globales), Mann-Whitney U (pares, con tamanos de efecto Cliff's delta, "
                "Vargha-Delaney A12, Cohen d, Hedges g) y Wilcoxon signed-rank (pares).")
    p(doc, "Herramientas e instrumentos:", bold=True)
    bullet(doc, "Simulacion: CityLearn v2 y CityLearn v3 propuesto (schema Iquitos).")
    bullet(doc, "Backends MADRL: HAPPO (HARL), MASAC, MATD3 y MAAC; MARLlib como referencia tecnica.")
    bullet(doc, "Stack computacional: Python 3.9, PyTorch/CUDA, Optuna (TPE), Gymnasium/PettingZoo.")
    bullet(doc, "Dataset e infra: citylearn_iquitos_2023_2025; ejecucion local (RTX 4060) y Colab "
                "(corrida canonica madrl_v3_20260627_164047).")
    p(doc,
      "Instrumentos de evidencia: scripts de entrenamiento/evaluacion del repositorio, "
      "best_madrl_report.json, pruebas no parametricas y figuras Drive auditadas.")

    heading(doc, "3.6 Procedimiento experimental", 2)
    numbered(doc, "Verificacion de contexto del proyecto (scripts/verify_project_context.ps1).")
    numbered(doc, "Construccion y auditoria del dataset (orchestrate_citylearn_dataset.py) y gates de validacion.")
    numbered(doc, "Verificacion de integridad del flujo (verify_workflow_integrity.py).")
    numbered(doc, "Entrenamiento de las 12 corridas (4 algoritmos x 3 escenarios) con monitor visible y "
                  "concurrencia controlada (HAPPO/MATD3 hasta 2; MASAC/MAAC 1).")
    numbered(doc, "Benchmark de la linea base CityLearn v2 (baseline, hour_rbc) y comparaciones SB3 (PPO/SAC/A2C) "
                  "sobre el mismo schema de Iquitos.")
    numbered(doc, "Comparacion v2 vs v3 (compare_citylearn_v2_vs_v3_madrl.py) y generacion de evidencia de "
                  "tesis (generate_thesis_objective_evidence.py): KPIs, pruebas estadisticas, figuras y tablas.")

    # ===================== CAP 4 =====================
    doc.add_page_break()
    heading(doc, "Capitulo 4. Desarrollo de la propuesta", 1)

    heading(doc, "4.1 Desarrollo del sistema y arquitectura de software", 2)
    p(doc,
      "El sistema se organiza en seis capas de software, desde el simulador base hasta la "
      "generacion de evidencia para la tesis:")
    add_table(
        doc,
        ["Capa", "Componente", "Funcion"],
        [
            ["1", "CityLearn v2 (CityLearn/citylearn/*.py)", "Fisica de edificios, BESS, PV, EV y KPIs oficiales."],
            ["2", "CityLearn v3 propuesto (CityLearn/citylearn/v3/)", "Entorno Dec-POMDP, objetivos, configuracion y recompensa multiobjetivo."],
            ["3", "Framework UC3M (uc3m/)", "Wrapper universal (UC3MEnv), 7 ejes de recompensa, KPIEvaluator, AlgorithmFactory; capa auxiliar/generica."],
            ["4", "Backends MADRL (external/)", "HARL (HAPPO), MARL/src (MASAC), off-policy (MATD3 activo), MAAC."],
            ["5", "Launchers y orquestacion", "train_citylearn_v3_*.py, launchers PowerShell y Colab A100, monitoreo y checkpointing."],
            ["6", "Evaluacion y evidencia", "benchmark_citylearn_v2_agents.py, compare_citylearn_v2_vs_v3_madrl.py, generate_thesis_objective_evidence.py."],
        ],
        caption="Tabla 4.1. Capas de software del sistema MADRL CityLearn v3.",
        col_widths=[1.2, 6.3, 8.5],
    )
    p(doc,
      "Nota de implementacion: las 12 corridas oficiales y los resultados reportados se "
      "ejecutan mediante los scripts CityLearn/scripts/train_citylearn_v3_<algoritmo>.py con "
      "la funcion de recompensa CityLearnV3MADRLRewardFunction (tres ejes flex/CO2/costo mas "
      "termino EV). El framework UC3M (capa 3, con 7 ejes de recompensa en uc3m/reward/axes.py "
      "y el sustento matematico en agent-skills/madrl-sustento-doc-capa v3/) es "
      "infraestructura generica/auxiliar: provee EmpiricalProtocol (n_seeds=12), HPHI, BACT "
      "y la batería no parametrica, pero no es el canal de entrenamiento de las 12 corridas "
      "reportadas en el Capitulo 5. En Cap. 2 se formaliza el mapeo 7→3; aqui solo se "
      "implementa el contrato de tres ejes.")
    p(doc,
      "Diseño multi-semilla (implementado). EmpiricalProtocol.n_seeds=12 en "
      "uc3m/multicriteria/scenarios.py; el runner scripts/run_madrl_nonparametric_battery.py "
      "acepta --run-root con jobs {ALGO}/{E*}_seed_{k}/ o CSV con columna seed. Estado "
      "ejecutado: (i) campaña canónica seed=0 (12 jobs); (ii) smoke ilustrativo n=3 en "
      "outputs/madrl_nonparametric_battery_smoke_n3/; (iii) batería episódica complementaria "
      "con HAPPO en outputs/madrl_nonparametric_battery/. No existen 12 semillas entrenadas "
      "reales en outputs/_drive_madrl; no se inventan.")

    heading(doc, "4.2 Modelo de IA: formulacion Dec-POMDP", 2)
    p(doc,
      "El problema se formaliza como el Dec-POMDP M = <S, {Ai}, T, R, {Oi}, Omega, gamma, T>:")
    bullet(doc, "N = 17 agentes (edificios del SEAI Iquitos).")
    bullet(doc, "Estado global S: concatenacion de las observaciones locales (state dim global = 1856 "
                "en el entorno cargado), accesible solo por el critico centralizado durante el entrenamiento.")
    bullet(doc, "Observacion local Oi: heterogenea por edificio (tiempo, fisica del edificio, "
                "subset meteorologico, precio, CI, estado SOC del BESS y estado de cada cargador EV); "
                "su dimension depende del numero de cargadores EV, en un rango aproximado de 57 a 330 "
                "dimensiones por agente (notebook madrl_citylearn_v3_tutorial.ipynb, Seccion 4).")
    bullet(doc, "Accion local Ai: heterogenea por edificio (potencia de BESS carga/descarga, potencia "
                "de carga EV por cargador y control de carga desplazable); la dimensionalidad varia de "
                "~5 a ~44 acciones segun la flota EV del edificio (B06 con 32 y B07 con 42 cargadores "
                "concentran las acciones de mayor dimension).")
    bullet(doc, "T: dinamica del balance energetico, modelo RC de temperatura, BESS con eficiencia "
                "round-trip 0.9025 y perfiles EV estocasticos.")
    bullet(doc, "R: recompensa cooperativa CityLearnV3MADRLRewardFunction; team_reward = media de reward_i.")
    bullet(doc, "gamma = 0.9999; horizonte T = 8 760 pasos (1 ano horario).")
    p(doc,
      "La condicion de observabilidad parcial estricta se satisface: cada edificio solo "
      "observa su propio estado local y no accede a la temperatura, demanda, SOC ni perfil EV "
      "de los demas durante la ejecucion.")

    heading(doc, "4.3 Esquema CTDE", 2)
    p(doc,
      "Entrenamiento centralizado: el critico Qi(s, a1,...,aN) o V(s) accede al estado global "
      "s = [o1,...,o17] durante el entrenamiento, corrigiendo la no-estacionariedad del "
      "aprendizaje independiente. Ejecucion descentralizada: la politica pi_i(ai|oi) usa solo "
      "la observacion local, sin comunicacion entre edificios. Tras el entrenamiento, el "
      "critico centralizado se descarta y solo se conservan las politicas locales.")

    heading(doc, "4.4 Algoritmos", 2)
    add_table(
        doc,
        ["Algoritmo", "Tipo", "Backend / wrapper", "Caracteristica clave"],
        [
            ["HAPPO", "On-policy", "external/HARL · CityLearnHARLEnv", "Actualizacion secuencial con trust region; mejora monotona."],
            ["MASAC", "Off-policy", "external/MARL/src · CityLearnSMACDiscreteEnv", "SAC + QMIX, acciones discretizadas por eje, RNN."],
            ["MATD3", "Off-policy", "external/off-policy · CityLearnOffPolicyVecEnv", "Doble critico TD3, policy delay, target noise (backend activo)."],
            ["MAAC", "Off-policy", "external/MAAC · CityLearnMAACVecEnv", "Critico de atencion multi-cabeza (4 heads)."],
        ],
        caption="Tabla 4.2. Los cuatro algoritmos MADRL y sus backends.",
        col_widths=[2.0, 2.0, 6.0, 6.0],
    )
    p(doc, "Hiperparametros: corrida canonica Colab vs referencia local v4:", bold=True)
    add_table(
        doc,
        ["Parametro", "Colab canonico (50 ep)", "Local v4 (5 ep)", "Notas"],
        [
            ["Episodios × pasos", "50 × 8 760", "5 × 8 760", "438 000 vs 43 800 pasos/corrida"],
            ["gamma", "0.9999", "0.9999", "Identico"],
            ["hidden / actor", "[512,512] HAPPO; 768 MATD3/MAAC", "[256,256] todos", "Mayor capacidad en Colab"],
            ["replay_buffer MATD3", "2 000 000", "4 096", "Restriccion VRAM 8 GB local"],
            ["batch MATD3", "1 024", "256", "Colab A100"],
            ["MASAC replay", "2 ep CPU", "on-GPU reducido", "Protocolo two_phase_happo_masac"],
            ["Semillas", "seed 0", "seed 0", "Limitacion inferencial"],
            ["GPU", "A100-SXM4-80GB", "RTX 4060 8 GB", "madrl_v3_20260627_164047"],
        ],
        caption="Tabla 4.3. Hiperparametros canonicos Colab vs referencia local v4.",
        col_widths=[4.0, 4.5, 4.0, 3.5],
    )
    p(doc,
      "La optimizacion automatica de hiperparametros con Optuna (TPE) queda como mejora "
      "experimental posterior; la evidencia cuantitativa se toma de la corrida canonica Colab.")

    heading(doc, "4.5 Funcion de recompensa multiobjetivo", 2)
    p(doc,
      "La clase CityLearnV3MADRLRewardFunction (hereda de Electric_Vehicles_Reward_Function) "
      "calcula la recompensa de cada edificio como:")
    p(doc,
      "reward_i(t) = reward_scale x [ w_flex·flex_i + w_carbon·carbon_i + w_cost·cost_i + "
      "ev_weight·ev_i ]", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc,
      "El componente de flexibilidad opera a nivel de distrito (peak_share y ramp_share sobre "
      "la carga neta agregada, con incentivos de exportacion/headroom BESS); el de carbono "
      "pondera la importacion por CI(t); y el de costo refleja la tarifa TOU. La agregacion "
      "cooperativa usa mixed_reward_i = (1 - r)·reward_i + r·team_reward, con team_reward_ratio "
      "r = 0.70. El perfil activo es unified_comparable (v4) e identico para los cuatro "
      "algoritmos, garantizando comparabilidad estadistica.")
    add_table(
        doc,
        ["Parametro", "Valor", "Justificacion sintetica"],
        [
            ["w_flex / w_carbon / w_cost (E1)", "0.70 / 0.15 / 0.15", "Escenario lexicografico de flexibilidad."],
            ["w_flex / w_carbon / w_cost (E2)", "0.15 / 0.70 / 0.15", "Escenario lexicografico de CO2."],
            ["w_flex / w_carbon / w_cost (E3)", "0.25 / 0.15 / 0.60", "Costo-pico correlacionados en TOU."],
            ["team_reward_ratio (r)", "0.70", "Coordinacion cooperativa Dec-POMDP (media de perfiles del plan)."],
            ["peak_weight / ramp_weight", "0.45 / 0.35", "KPI primario y secundario de flexibilidad CityLearn."],
            ["ev_weight", "0.25", "Termino EV/SOC reforzado para evitar dejar EV sin cargar."],
            ["reward_scale", "1.00", "Escala uniforme para gradientes comparables."],
            ["carbon_reference / price_reference", "0.35 / 0.20", "Referencia global IEA y spot competitivo."],
        ],
        caption="Tabla 4.4. Parametros de la funcion de recompensa multiobjetivo (perfil unificado v4).",
        col_widths=[5.2, 3.8, 7.0],
    )

    heading(doc, "4.6 Aportes originales al motor de simulacion", 2)
    p(doc,
      "Como contribucion metodologica diferencial, se realizaron cuatro extensiones "
      "originales y retrocompatibles al fork CityLearn (commit 54b1938e):")
    add_table(
        doc,
        ["#", "Aporte", "Clase/Metodo", "Modelo incorporado"],
        [
            ["A1", "Degradacion BESS C-rate + Arrhenius", "Battery.degrade(temperature_celsius)", "ΔC = base·(C_rate)^0.55·exp[Ea/R·(1/T_ref - 1/T)], Ea=24500 J/mol."],
            ["A2", "Correccion PV tropical (IEC 61215)", "PV.get_generation(dry_bulb_temperature, ghi)", "T_cell = T_amb + (NOCT-20)/800·G; P(T)=P_STC·[1+γ(T_cell-25)], γ=-0.0035/°C."],
            ["A3", "KPI pico con ventana de facturacion", "CostFunction.peak(billing_window_steps)", "Maximo rodante sub-horario (OSINERGMIN MT-3/MT-4, 15 min)."],
            ["A4", "Modelo CI dinamico diesel+PV", "CarbonIntensityModel(base_ci, pv_displacement_factor)", "CI(t)=0.790·(1-0.15·min(GHI/1000,1))."],
        ],
        caption="Tabla 4.5. Aportes originales al motor de simulacion CityLearn.",
        col_widths=[0.8, 4.4, 5.0, 5.8],
        font_size=8,
    )

    heading(doc, "4.7 Diseno experimental: matriz de 12 corridas", 2)
    p(doc,
      "El experimento oficial ejecuta 12 corridas (4 algoritmos x 3 escenarios, seed = 0). "
      "La configuracion canonica objetivo es de 50 episodios x 8 760 "
      "pasos (438 000 pasos/corrida); la corrida local de referencia v4 completo 5 episodios "
      "x 8 760 pasos (43 800 pasos/corrida) en GPU RTX 4060 Laptop 8 GB. La paralelizacion "
      "es valida porque las corridas son experimentos independientes (Roijers et al., 2013; "
      "MALib, Zhou et al., 2021).")
    add_table(
        doc,
        ["Escenario", "HAPPO", "MASAC", "MATD3", "MAAC"],
        [
            ["E1 Flexibilidad", "happo/E1_s0", "masac/E1_s0", "matd3/E1_s0", "maac/E1_s0"],
            ["E2 CO2", "happo/E2_s0", "masac/E2_s0", "matd3/E2_s0", "maac/E2_s0"],
            ["E3 Costos", "happo/E3_s0", "masac/E3_s0", "matd3/E3_s0", "maac/E3_s0"],
        ],
        caption="Tabla 4.6. Matriz experimental de 12 corridas oficiales.",
        col_widths=[3.5, 3.1, 3.1, 3.1, 3.1],
    )

    heading(doc, "4.8 Implementacion", 2)
    p(doc,
      "Entorno: Python 3.9 (.venv39-citylearn-v3), PyTorch 2.8.0+cu126, CUDA 12.6. Los "
      "scripts de entrenamiento comparten el adaptador citylearn_v3_training_common.py "
      "(normalizacion, KPIs, trazas, figuras, tablas y metadatos). Cada corrida produce "
      "results.json, training_summary.json, timeseries.csv, trace.csv, checkpoint_manifest.json "
      "y figures_manifest.json. El despliegue admite ejecucion local (PowerShell) y en la nube "
      "(Colab A100; Docker/AWS EC2 opcional).")
    p(doc,
      "La corrida canonica de 50 episodios se ejecuta en Colab sobre una NVIDIA A100-SXM4-80GB "
      "(Pro+ High-RAM, ~167 GiB de RAM) mediante colab_a100_official_launcher.py en modo "
      "two_phase_happo_masac (protocolo two_phase_happo_masac_v3): la Fase 1 entrena HAPPO y "
      "MASAC y la Fase 2 MATD3 y MAAC, con seis trabajos en paralelo por fase, replay de MASAC "
      "en CPU, monitoreo con colab_a100_live_monitor.py y reanudacion intra-job mediante "
      "live_progress.json y --skip-completed. Los artefactos de esa corrida se integraran desde "
      "outputs/colab_50ep/ o desde la carpeta de Google Drive del entrenamiento. Fuente: "
      "notebook examples_madrl_v3/madrl_citylearn_v3_tutorial.ipynb.")
    p(doc,
      "Validacion de software (Cap. IV): pytest tests/uc3m/ + tests/citylearn_v3/ en el venv "
      "del proyecto (184 passed el 2026-07-29; ver seccion 4.12 del Word canónico vía "
      "tools/thesis/patch_cap4_implementacion_docx.py y docs/tesis_capitulos/"
      "Capitulo_4_Desarrollo_Propuesta.md). La fachada UC3M y el schema smoke no sustituyen "
      "las 12 corridas ni la batería estadística del Capitulo 5.")

    if max_chapter is not None and max_chapter < 5:
        doc.add_page_break()
        heading(doc, "Referencias bibliograficas", 1)
        stats = reference_stats()
        p(
            doc,
            f"Formato APA (7.a ed.). Fuente consolidada del skill: docs/tesis_capitulos/Referencias_APA.md "
            f"({stats['from_apa_md']} entradas) + citas en texto y protocolo estadistico RL "
            f"({stats['total_unique']} referencias unicas"
            + (f"; {stats['pv_marked']} con metadatos bibliograficos incompletos)." if int(stats['pv_marked']) else ")."),
            italic=True,
            size=10,
            color=GREY,
        )
        for ref in load_all_thesis_references():
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.first_line_indent = Cm(-1.0)
            para.paragraph_format.space_after = Pt(6)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(ref).font.size = Pt(10)
        OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(OUT_PATH))
        print(f"OK -> {OUT_PATH}")
        return OUT_PATH

    # ===================== CAP 5 =====================
    doc.add_page_break()
    heading(doc, "Capitulo 5. Resultados", 1)

    status_note(doc,
      "AVISO - RESULTADOS PRELIMINARES. Todas las cifras de este capitulo son PRELIMINARES y "
      "provienen de la corrida local de 5 episodios (citylearn_v3_madrl_full_20260615_074011_v4). "
      "La corrida canonica de 50 episodios se esta ejecutando en Google Colab (A100, modo "
      "two_phase_happo_masac). Al finalizar, estos valores se reemplazaran por los de 50 "
      "episodios, recalculando KPIs normalizados, pruebas estadisticas y % de mejora vs baseline, "
      "e insertando las figuras .png definitivas; los artefactos se integraran desde "
      "outputs/colab_50ep/ o la carpeta de Google Drive de la corrida. "
      "[REEMPLAZAR con resultados de la corrida canonica de 50 episodios en Colab]")

    heading(doc, "5.1 Experimentos realizados", 2)
    p(doc,
      "Se completaron las 12 corridas oficiales de la corrida de referencia v4 "
      "(outputs/citylearn_v3_madrl_full_20260615_074011_v4), todas con exit_code = 0. El "
      "entrenamiento total tomo del 2026-06-15 al 2026-06-16 (~39 h de reloj) en GPU RTX 4060 "
      "Laptop. La duracion por corrida vario por algoritmo: HAPPO ~57-67 min, MASAC ~126-148 "
      "min, MATD3 ~376-451 min y MAAC ~323-332 min por escenario.")

    heading(doc, "5.2 Metricas utilizadas", 2)
    p(doc,
      "Se emplean los KPIs oficiales de CityLearn v2, normalizados respecto a la linea base "
      "(valor 1.0 = baseline; valores < 1.0 indican mejora). Por eje: OE.1 peak_average, "
      "ramping_average, one_minus_load_factor_average; OE.2 carbon_emissions; OE.3 "
      "electricity_cost. El ranking integrado usa un score compuesto por reward promedio, "
      "reduccion de picos, gestion del SOC del BESS, reduccion de CO2, cumplimiento de "
      "restricciones y robustez.")

    heading(doc, "5.3 Resultados obtenidos", 2)
    p(doc,
      "El ranking global de la corrida v4 identifica a MATD3 como el mejor algoritmo MADRL, "
      "seguido de MASAC, MAAC y HAPPO. La prueba de Kruskal-Wallis sobre el score global "
      "arroja p = 0.0459 (significativo a alpha = 0.05).")
    add_table(
        doc,
        ["Rango", "Algoritmo", "OE.1 Flex", "OE.2 CO2", "OE.3 Costo", "Score global"],
        [
            ["1", "MATD3", "0.7486", "0.7515", "0.7333", "0.7445"],
            ["2", "MASAC", "0.74", "0.74", "0.72", "~0.73"],
            ["3", "MAAC", "0.72", "0.72", "0.73", "~0.72"],
            ["4", "HAPPO", "0.70", "0.70", "0.70", "~0.70"],
        ],
        caption="Tabla 5.1. Scores por eje y ranking global (corrida v4).",
        col_widths=[1.6, 3.0, 2.8, 2.8, 2.8, 3.0],
    )
    p(doc,
      "KPIs normalizados de evaluacion (ejemplo MATD3 en E3, respecto a baseline): "
      "peak_average = 1.0112, ramping_average = 1.0009, carbon_emissions = 1.0847 y "
      "electricity_cost = 1.0092. Estos valores cercanos o ligeramente superiores a 1.0 "
      "indican que, con el presupuesto de entrenamiento local (5 episodios), las politicas "
      "MADRL aun no superan de forma consistente a la linea base en los KPIs crudos, lo que "
      "es coherente con el caracter preliminar de los resultados.")
    add_table(
        doc,
        ["KPI (normalizado a baseline)", "MATD3 / E3", "Interpretacion"],
        [
            ["peak_average", "1.0112", "~1% por encima del baseline (pico)."],
            ["ramping_average", "1.0009", "Practicamente igual al baseline."],
            ["carbon_emissions", "1.0847", "~8% por encima del baseline (CO2)."],
            ["electricity_cost", "1.0092", "~1% por encima del baseline (costo)."],
            ["ev_departure_success_rate", "0.4749", "47.5% de salidas EV con SOC requerido."],
        ],
        caption="Tabla 5.2. KPIs normalizados de la corrida MATD3/E3 (data/training_summary.json).",
        col_widths=[6.5, 3.0, 6.5],
    )

    status_note(doc,
      "[REEMPLAZAR con resultados de la corrida canonica de 50 episodios en Colab - "
      "scores y KPIs de las Tablas 5.1 y 5.2 son preliminares (5 episodios).]")

    heading(doc, "5.4 Comparacion con baseline / trabajos relacionados", 2)
    p(doc,
      "La comparacion entre familias (CityLearn v2 original vs CityLearn v3 MADRL) se calcula "
      "con compare_citylearn_v2_vs_v3_madrl.py usando pesos iguales por eje (OE1 0.34, OE2 "
      "0.33, OE3 0.33). En el escenario E1, el mejor por eje en flexibilidad (OE1) fue HAPPO "
      "(score 0.5679), superando a la linea base; en CO2 (OE2) y costos (OE3) la linea base "
      "(baseline / hour_rbc) mantuvo ventaja, reflejando que las politicas MADRL aun "
      "requieren mas episodios de entrenamiento para dominar todos los ejes simultaneamente.")
    add_table(
        doc,
        ["Escenario", "Mejor OE1 (flex)", "Mejor OE2 (CO2)", "Mejor OE3 (costo)", "Mejor global"],
        [
            ["E1", "HAPPO (0.5679)", "baseline v2 (~1.000)", "hour_rbc v2 (0.7474)", "baseline v2 (0.7254)"],
            ["E2", "HAPPO (0.6769)", "MATD3 (0.9858)", "MATD3 (0.8401)", "MATD3 (0.7515)"],
            ["E3", "HAPPO (0.6806)", "MATD3 (0.9811)", "MAAC (0.7879)", "MATD3 (0.7333)"],
        ],
        caption="Tabla 5.3. Comparacion v2 vs v3 por eje y escenario (best_by_axis).",
        col_widths=[2.2, 3.6, 3.6, 3.6, 3.0],
    )
    p(doc,
      "Frente a los trabajos relacionados (Yao et al., 2023: ~18% de reduccion de costo; Liu "
      "et al., 2022: ~15% CO2 y ~20% costo), los resultados preliminares de esta tesis aun no "
      "alcanzan esas magnitudes de mejora porque el presupuesto de entrenamiento local es "
      "reducido. La contribucion vigente es metodologica: un benchmark unificado, reproducible "
      "y estadisticamente fundamentado de cuatro algoritmos sobre un dataset real.")

    heading(doc, "5.5 Pruebas estadisticas", 2)
    add_table(
        doc,
        ["Prueba", "Resultado", "p-valor", "Conclusion"],
        [
            ["Shapiro-Wilk", "Algunos grupos no normales", "-", "Justifica el uso de pruebas no parametricas."],
            ["Kruskal-Wallis", "Diferencia entre los 4 algoritmos", "0.0459", "Significativo (alpha = 0.05)."],
            ["Mann-Whitney U (MATD3 vs HAPPO)", "MATD3 superior", "0.0182", "Significativo."],
            ["Wilcoxon SR (MATD3 vs HAPPO)", "Diferencia sistematica", "2.62e-6", "Muy significativo."],
        ],
        caption="Tabla 5.4. Pruebas estadisticas sobre el score global (corrida v4 preliminar, 5 episodios).",
        col_widths=[5.0, 4.5, 2.5, 4.0],
    )

    status_note(doc,
      "[REEMPLAZAR con resultados de la corrida canonica de 50 episodios en Colab - "
      "comparacion vs baseline y pruebas estadisticas (Tablas 5.3 y 5.4) son preliminares.]")

    heading(doc, "5.6 Figuras", 2)
    p(doc,
      "Cada corrida genera figuras estandarizadas (figures/, 13 por corrida): "
      "reward_timeseries.png, convergence_returns.png, episode_reward_summary.png, "
      "learning_efficiency.png, citylearn_v2_district_timeseries.png, "
      "axis_baseline_comparison.png, baseline_gain_by_kpi.png, core_kpis.png, "
      "OE1_flexibility_kpis.png, OE2_co2_kpis.png y OE3_cost_kpis.png. La comparacion v2 vs "
      "v3 genera ademas OE1/OE2/OE3_comparison.png y baseline_gain_heatmap.png por escenario.")
    status_note(doc,
      "[Pendiente: insertar en esta seccion las figuras .png seleccionadas de "
      "outputs/citylearn_v3_madrl_full_20260615_074011_v4/<algoritmo>/<escenario>/figures/ "
      "y de outputs/comparison_citylearn_v2_vs_v3_madrl/ una vez elegidas las definitivas.]")

    heading(doc, "5.7 Discusion de resultados", 2)
    p(doc,
      "Los resultados preliminares permiten tres lecturas. Primero, el ranking inter-algoritmo "
      "es estadisticamente significativo (Kruskal-Wallis p = 0.0459) y consistente: MATD3 "
      "—off-policy con doble critico— alcanza el mejor score global, lo que sugiere que la "
      "estabilidad del critico TD3 es ventajosa en este entorno heterogeneo de 17 agentes. "
      "Segundo, en flexibilidad (OE.1) HAPPO destaca por eje, coherente con su diseno para "
      "agentes heterogeneos y con el antecedente de Nweye et al. (2023b). Tercero, los KPIs "
      "crudos normalizados aun se situan cerca o por encima del baseline, evidenciando que el "
      "presupuesto de entrenamiento local (5 episodios) es insuficiente para que las politicas "
      "MADRL superen consistentemente a un control basado en reglas bien sintonizado. La "
      "configuracion canonica (50 episodios) y la HPO con Optuna son las palancas previstas "
      "para cerrar esta brecha en la version final de la tesis.")

    # ===================== CAP 6 =====================
    doc.add_page_break()
    heading(doc, "Capitulo 6. Conclusiones preliminares", 1)

    heading(doc, "6.1 Principales hallazgos", 2)
    bullet(doc, "Se construyo un benchmark unificado, reproducible y estadisticamente fundamentado "
                "de cuatro algoritmos MADRL (HAPPO, MASAC, MATD3, MAAC) bajo Dec-POMDP y CTDE sobre "
                "un dataset real del SEAI Iquitos (17 edificios, 26 304 h, 222 CSV auditados).")
    bullet(doc, "En la corrida de referencia v4, MATD3 es el mejor MADRL global (score 0.7445) con "
                "diferencias significativas (Kruskal-Wallis p = 0.0459; Mann-Whitney MATD3 vs HAPPO "
                "p = 0.0182; Wilcoxon p = 2.62e-6).")
    bullet(doc, "Por eje, HAPPO destaca en flexibilidad (OE.1) y MATD3 en CO2 (OE.2); MAAC lidera "
                "costos (OE.3) en la corrida canonica Colab/Drive, segun la comparacion v2 vs v3.")
    bullet(doc, "Se aportaron cuatro extensiones originales y retrocompatibles al motor CityLearn "
                "(degradacion BESS C-rate+Arrhenius, correccion PV tropical IEC 61215, KPI de pico "
                "con ventana de facturacion OSINERGMIN y clase CarbonIntensityModel).")

    heading(doc, "6.2 Limitaciones encontradas", 2)
    bullet(doc, "Presupuesto de entrenamiento local reducido (5 episodios en RTX 4060 8 GB), por lo "
                "que los KPIs crudos normalizados aun no superan consistentemente al baseline.")
    bullet(doc, "Restricciones de VRAM que obligan a limitar buffers y concurrencia (MASAC/MAAC a 1 corrida).")
    bullet(doc, "Una sola semilla (seed = 0) en la corrida de referencia, lo que limita el analisis de robustez.")
    bullet(doc, "Algunas referencias bibliograficas estan marcadas como pendientes de verificacion de datos secundarios.")

    heading(doc, "6.3 Trabajo pendiente", 2)
    bullet(doc, "Ejecutar la configuracion canonica de 50 episodios x 8 760 pasos (438 000 pasos/corrida), "
                "preferentemente en Colab A100, para las 12 corridas.")
    bullet(doc, "Repetir con multiples semillas y aplicar HPO con Optuna (TPE) a cada backend.")
    bullet(doc, "Completar el benchmark de comparacion con SB3 (PPO/SAC/A2C) sobre el mismo schema.")
    bullet(doc, "Insertar las figuras definitivas y consolidar las tablas de KPIs por edificio.")
    bullet(doc, "Verificar y completar las referencias marcadas como pendientes.")

    heading(doc, "6.4 Plan para culminar la tesis", 2)
    add_table(
        doc,
        ["Fase", "Actividades", "Meses"],
        [
            ["Preparatoria", "Revision bibliografica, Modulo A (50 antecedentes), diagnostico, dataset y KPIs.", "1-3"],
            ["Diseno tecnico", "Arquitectura CityLearn v3, Dec-POMDP, CTDE, integracion de backends, Optuna.", "4-8"],
            ["Evaluacion por eje", "Entrenamiento E1/E2/E3 x 4 algoritmos (50 ep), evaluacion de KPIs OE.1/OE.2/OE.3.", "9-18"],
            ["Determinacion y cierre", "Comparacion, ranking, analisis estadistico, discusion SEAI Iquitos, redaccion y sustentacion.", "19-24"],
        ],
        caption="Tabla 6.1. Cronograma para la culminacion de la tesis.",
        col_widths=[3.5, 9.5, 3.0],
    )

    # ===================== REFERENCIAS =====================
    doc.add_page_break()
    heading(doc, "Referencias bibliograficas", 1)
    stats = reference_stats()
    p(
        doc,
        f"Formato APA (7.a ed.). Fuente consolidada del skill: docs/tesis_capitulos/Referencias_APA.md "
        f"({stats['from_apa_md']} entradas) + citas en texto y protocolo estadistico RL "
        f"({stats['total_unique']} referencias unicas"
        + (f"; {stats['pv_marked']} con metadatos bibliograficos incompletos)." if int(stats['pv_marked']) else ")."),
        italic=True,
        size=10,
        color=GREY,
    )

    references = load_all_thesis_references()
    for ref in references:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.0)
        para.paragraph_format.first_line_indent = Cm(-1.0)
        para.paragraph_format.space_after = Pt(6)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(ref).font.size = Pt(10)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"OK -> {OUT_PATH}")
    print(f"size_bytes={OUT_PATH.stat().st_size}")
    rs = reference_stats()
    print(f"referencias_apa={rs['total_unique']} (md={rs['from_apa_md']}, pv={rs['pv_marked']})")


if __name__ == "__main__":
    build()
