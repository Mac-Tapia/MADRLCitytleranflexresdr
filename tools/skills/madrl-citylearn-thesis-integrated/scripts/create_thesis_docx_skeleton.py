"""Create a DOCX skeleton for the professional master's thesis report."""

from __future__ import annotations

import argparse
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from docx import Document
except ModuleNotFoundError:  # pragma: no cover - exercised when python-docx is absent.
    Document = None


TITLE = (
    "MULTI-AGENTE DE APRENDIZAJE POR REFUERZO PROFUNDO PARA LA GESTIÓN "
    "COORDINADA DE FLEXIBILIDAD ENERGÉTICA, EMISIONES DE CARBONO Y COSTOS "
    "ENERGÉTICOS EN COMUNIDADES INTELIGENTES"
)

SECTIONS = [
    ("CARÁTULA", 0),
    ("DATOS GENERALES", 0),
    ("Dedicatoria", 1),
    ("Agradecimientos", 1),
    ("Copia de documentos", 1),
    ("Índice de contenidos", 1),
    ("Lista de tablas, ilustraciones y cuadros", 1),
    ("Resumen - Abstract", 1),
    ("Introducción", 1),
    ("CAPÍTULO I. PLANTEAMIENTO DEL PROBLEMA", 0),
    ("1.1 Diagnóstico", 1),
    ("1.2 Identificación y descripción del problema de estudio", 1),
    ("1.3 Formulación del problema", 1),
    ("1.3.1 Formulación del problema general", 2),
    ("1.3.2 Formulación de los problemas específicos", 2),
    ("1.4 Objetivos", 1),
    ("1.4.1 Objetivo general", 2),
    ("1.4.2 Objetivos específicos", 2),
    ("1.5 Justificación del estudio", 1),
    ("1.6 Alcance del estudio", 1),
    ("CAPÍTULO II. MARCO TEÓRICO", 0),
    ("2.1 Antecedentes", 1),
    ("2.2 Bases teóricas", 1),
    ("2.3 Definición de términos", 1),
    ("CAPÍTULO III. DESARROLLO DEL TRABAJO DE TESIS", 0),
    ("3.1 Presentación de la propuesta de solución", 1),
    ("3.2 Desarrollo de la propuesta de solución", 1),
    ("3.3 Análisis de los datos y resultados", 1),
    ("3.4 Discusión e interpretación de los resultados", 1),
    ("3.5 Estimación del impacto de la solución", 1),
    ("CAPÍTULO IV. CONCLUSIONES Y RECOMENDACIONES", 0),
    ("4.1 Conclusiones", 1),
    ("4.2 Recomendaciones", 1),
    ("REFERENCIAS", 0),
    ("ANEXOS", 0),
    ("Anexo 1. Matriz de consistencia", 1),
    ("Anexo 2. Matriz de operacionalización de variables", 1),
    ("Anexo 3. Matriz de antecedentes", 1),
    ("Anexo 4. Matriz de KPIs", 1),
    ("Anexo 5. Arquitectura CityLearn v3 propuesta", 1),
    ("Anexo 6. Comparación de backends MADRL", 1),
    ("Anexo 7. Datasets y fuentes", 1),
    ("Anexo 8. Configuración de hiperparámetros", 1),
    ("Anexo 9. Recompensa multiobjetivo", 1),
    ("Anexo 10. Resultados de simulación vigentes o pendientes", 1),
    ("Anexo 11. Evidencias de GitHub", 1),
    ("Anexo 12. Glosario MADRL", 1),
    ("Anexo 13. Cadenas de búsqueda", 1),
    ("Anexo 14. Matriz bibliográfica de 50 investigaciones", 1),
]


def _xml_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _paragraph_xml(text: str, *, heading: bool = False) -> str:
    style = '<w:pStyle w:val="Heading1"/>' if heading else ""
    return (
        "<w:p><w:pPr>"
        f"{style}"
        "</w:pPr><w:r><w:t>"
        f"{_xml_escape(text)}"
        "</w:t></w:r></w:p>"
    )


def _build_minimal_docx(output: Path) -> None:
    body = [_paragraph_xml(TITLE, heading=True)]
    body.append(_paragraph_xml("Documento base generado para tesis de Maestría de Especialización o Profesionalizante."))
    body.append(_paragraph_xml("Nota: completar solo con fuentes verificadas, citas APA vigentes y resultados reales observados; los resultados incompletos deben marcarse como pendientes o no verificados."))

    for heading, _level in SECTIONS:
        body.append(_paragraph_xml(heading, heading=True))
        if heading == "CARÁTULA":
            body.extend([
                _paragraph_xml("Universidad: [pendiente]"),
                _paragraph_xml("Escuela de posgrado: [pendiente]"),
                _paragraph_xml("Tesista: [pendiente]"),
                _paragraph_xml("Asesor: [pendiente]"),
                _paragraph_xml("Lima, Perú - 2026"),
            ])
        elif heading == "REFERENCIAS":
            body.append(_paragraph_xml("Insertar únicamente referencias citadas en el texto, en formato APA vigente."))
        else:
            body.append(_paragraph_xml("[Redactar con evidencia del Módulo A, citas APA y trazabilidad de fuente.]"))

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + "".join(body)
        + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>'
        + "</w:body></w:document>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        "</Relationships>"
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(output, "w", ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)


def build(output: Path) -> None:
    if Document is None:
        _build_minimal_docx(output)
        return

    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph("Documento base generado para tesis de Maestría de Especialización o Profesionalizante.")
    doc.add_paragraph("Nota: completar solo con fuentes verificadas, citas APA vigentes y resultados reales observados; los resultados incompletos deben marcarse como pendientes o no verificados.")

    for heading, level in SECTIONS:
        doc.add_heading(heading, level=level if level else 1)
        if heading == "CARÁTULA":
            doc.add_paragraph("Universidad: [pendiente]")
            doc.add_paragraph("Escuela de posgrado: [pendiente]")
            doc.add_paragraph("Tesista: [pendiente]")
            doc.add_paragraph("Asesor: [pendiente]")
            doc.add_paragraph("Lima, Perú - 2026")
        elif heading == "REFERENCIAS":
            doc.add_paragraph("Insertar únicamente referencias citadas en el texto, en formato APA vigente.")
        else:
            doc.add_paragraph("[Redactar con evidencia del Módulo A, citas APA y trazabilidad de fuente.]")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", required=True, help="Path to the .docx file.")
    args = parser.parse_args()
    build(Path(args.output))
    print(Path(args.output).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
