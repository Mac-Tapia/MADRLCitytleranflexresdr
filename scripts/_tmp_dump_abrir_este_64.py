# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

OUT = Path(r"D:\MADRLCitytleranflexresdr\outputs\_tmp_abrir_este_dump.txt")
DOC = Path(
    r"D:\MADRLCitytleranflexresdr\docs\ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx"
)


def main() -> None:
    doc = Document(str(DOC))
    lines: list[str] = []
    for i, para in enumerate(doc.paragraphs):
        if 636 <= i <= 655 or i in (419, 420, 421, 422, 574, 575, 576, 577):
            style = para.style.name if para.style else ""
            t = (para.text or "").replace("\n", " | ")
            lines.append(f"P{i}|{style}|{t}")

    body = list(doc.element.body)
    ti = 0
    for ci, child in enumerate(body):
        if child.tag != qn("w:tbl"):
            continue
        prev = ""
        for j in range(ci - 1, -1, -1):
            if body[j].tag == qn("w:p"):
                prev = "".join(body[j].itertext()).strip()
                break
        if prev.startswith("Tabla 6.1") or prev.startswith("Tabla 6.2"):
            tbl = doc.tables[ti]
            lines.append(f"TABLE after={prev[:100]} idx={ti} rows={len(tbl.rows)}")
            for r, row in enumerate(tbl.rows):
                cells = [c.text.replace("\n", " ") for c in row.cells]
                lines.append("  R" + str(r) + "|" + " || ".join(cells))
        ti += 1

    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"wrote {OUT} lines={len(lines)}")


if __name__ == "__main__":
    main()
