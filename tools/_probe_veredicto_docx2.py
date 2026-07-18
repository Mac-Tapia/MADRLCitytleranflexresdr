#!/usr/bin/env python3
"""Probe exact Cap1 PG/OG/HG paragraphs in PATCHED Word."""
from pathlib import Path
from docx import Document

p = Path(r"D:\MADRLCitytleranflexresdr\docs\ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS_PATCHED.docx")
doc = Document(str(p))
for i in range(50, 100):
    t = (doc.paragraphs[i].text or "").strip()
    if t:
        print(f"--- {i} ---")
        print(t[:500])
print("==== CAP3 280-295 ====")
for i in range(280, 300):
    t = (doc.paragraphs[i].text or "").strip()
    if t:
        print(f"--- {i} ---")
        print(t[:400])
print("==== CAP6 620-640 ====")
for i in range(620, 640):
    t = (doc.paragraphs[i].text or "").strip()
    if t:
        print(f"--- {i} ---")
        print(t[:500])
