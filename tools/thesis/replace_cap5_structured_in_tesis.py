#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Rebuild Cap. 5 into Tesis (and optionally Informe) without regenerating Caps. 1–4/6.

Preserves the existing fat Tesis body outside Cap. 5.
"""

from __future__ import annotations

import argparse
import json
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))
if str(REPO) not in sys.path:
    sys.path.insert(0, str(REPO))

from generate_borrador_tesis_docx import (  # noqa: E402
    ACCENT,
    GREY,
    add_table,
    heading,
    p,
    status_note,
    style_base,
)
from sync_cap5_to_canon_words import extract_cap5, replace_cap5  # noqa: E402
from thesis_cap5_structured import add_chapter_5_structured  # noqa: E402
from thesis_word_canons import INFORME, TESIS, require_tesis  # noqa: E402


def _build_cap5_elements() -> list:
    tmp = Document()
    style_base(tmp)
    section = tmp.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    add_chapter_5_structured(tmp, p, heading, add_table, status_note)
    # Cap 5 starts at first body child (no front matter)
    return [deepcopy(child) for child in tmp.element.body]


def main() -> int:
    parser = argparse.ArgumentParser(description="Replace Cap.5 in Tesis from structured builder")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--informe-also", action="store_true", help="Also sync Cap.5 into Informe")
    args = parser.parse_args()

    require_tesis()
    elements = _build_cap5_elements()
    # Validate built Cap5 has expected independent headings in text
    blob = "\n".join("".join(t.text or "" for t in el.iter(qn("w:t"))) for el in elements)
    expected = [
        "5.2 Resultados descriptivos",
        "5.2.1 OG",
        "5.2.2 OE.1",
        "5.2.3 OE.2",
        "5.2.4 OE.3",
        "5.3 Resultados inferenciales",
        "5.3.1 OG",
        "5.3.2 OE.1",
        "5.3.3 OE.2",
        "5.3.4 OE.3",
        "5.4 Otros resultados",
        "5.4.1 OG",
        "5.4.2 OE.1",
        "5.4.3 OE.2",
        "5.4.4 OE.3",
        "5.4.5 Control de recursos",
        "5.1.1 Criterios de determinacion del impacto",
        "5.5 Contrastacion de hipotesis",
        "5.5.1 Hipotesis general",
        "5.5.2 Hipotesis especificas OE.1",
        "5.5.3 Hipotesis especificas OE.2",
        "5.5.4 Hipotesis especificas OE.3",
        "5.6 Discusion de resultados",
    ]
    missing = [e for e in expected if e not in blob]
    report = {
        "built_blocks": len(elements),
        "missing_headings": missing,
        "ok_structure": not missing,
    }
    if missing:
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return 1

    tesis_info = replace_cap5(TESIS, elements, toc_safe=False, dry_run=args.dry_run)
    report["tesis"] = tesis_info
    if args.informe_also and INFORME.is_file():
        # Re-extract from Tesis after write (or use same elements)
        if args.dry_run:
            report["informe"] = {"dry_run": True}
        else:
            cap5_from_tesis, _meta = extract_cap5(TESIS)
            report["informe"] = replace_cap5(
                INFORME, cap5_from_tesis, toc_safe=False, dry_run=False
            )

    out = REPO / "docs" / "_cap5_restructure_report.json"
    if not args.dry_run:
        out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        report["report_path"] = str(out)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if tesis_info.get("ok") and not missing else 1


if __name__ == "__main__":
    raise SystemExit(main())
