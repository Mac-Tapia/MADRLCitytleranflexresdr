#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Inserta §2.2.10 Adecuaciones de los 4 MADRL al dominio eléctrico en Cap. 2.

Edita solo los 2 Word canónicos. Idempotente (marcador MARKER).
Sustento: CityLearn/scripts/citylearn_v3_training_common.py + train_*.py +
docs/tesis_capitulos/Capitulo_2_Marco_Teorico.md §§2.2.4.1–2.2.4.4.
"""
from __future__ import annotations

import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, INFORME, TESIS  # noqa: E402

REPORT = DOCS / "CAP2_MADRL_ADECUACIONES_PATCH_REPORT_2026-07-29.json"
BACKUP_DIR = REPO / "outputs" / "_word_backups"
MARKER = "2.2.10 Adecuaciones de los cuatro backends MADRL al dominio electrico"
MARKER_NORM = "2.2.10 adecuaciones de los cuatro backends madrl al dominio electrico"

HEADING = (
    "2.2.10 Adecuaciones de los cuatro backends MADRL al dominio electrico "
    "(OE.1 flexibilidad, OE.2 CO2, OE.3 costos)"
)

PARAS: list[str] = [
    (
        "Premisa. Los backends oficiales pinados en external/ (HARL/HAPPO, MARL/src-MASAC, "
        "off-policy/MATD3 y MAAC) fueron disenados, validados y entrenados de forma "
        "predominante en dominios de referencia multiagente (p. ej. SMAC/StarCraft, MPE u "
        "otros benchmarks MARL). En su forma original no estan preparados, ni entrenados, "
        "ni aplicados out-of-the-box al control de una comunidad electrica con DER, EV/V2G "
        "y senales horarias de precio e intensidad de carbono. Por tanto, no resuelven por "
        "si solos el problema de esta tesis: flexibilidad energetica (OE.1), emisiones de "
        "CO2 (OE.2) y costos economico-energeticos (OE.3) sobre el SEAI Iquitos."
    ),
    (
        "La contribucion teorica no consiste en reinventar HAPPO, MASAC, MATD3 o MAAC, sino "
        "en adecuarlos —via la capa CityLearn v3 propuesto, el adaptador comun "
        "citylearn_v3_training_common.py y wrappers tipados— a un Dec-POMDP cooperativo de "
        "17 edificios heterogeneos, recompensa multiobjetivo unificada "
        "(CityLearnV3MADRLRewardFunction, perfil unified_comparable_v4) y evaluacion con "
        "KPIs oficiales evaluate_v2. El Cap. 4 documenta la implementacion; esta seccion "
        "fija el sustento teorico de cada adecuacion."
    ),
    (
        "Capa comun de adecuacion (los cuatro MADRL): (i) contrato Dec-POMDP/CTDE con "
        "estado global s en R^{1856}; (ii) padding / policy mapping ante heterogeneidad "
        "d_oi en [54, 327] y d_ai en [5, 44]; (iii) recompensa multiobjetivo con escenarios "
        "E1/E2/E3 y r_team = 0,70; (iv) dataset citylearn_iquitos_2023_2025 (BESS, PV, "
        "185 cargadores EV, TOU Electro Oriente, CI dinamico); (v) orquestacion "
        "train_citylearn_v3_{happo,masac,matd3,maac}.py bajo matriz 4x3 comparable."
    ),
    (
        "HAPPO — CityLearnHARLEnv + external/HARL. El HAPPO original (Kuba et al., 2021; "
        "Zhong et al., 2023) aporta actualizacion secuencial con trust region para agentes "
        "heterogeneos, pero HARL espera un ShareVecEnv de juegos multiagente, no setpoints "
        "BESS/EV. Adecuacion: wrapper CityLearnHARLEnv con observacion local continua y "
        "share_observation_space = estado CTDE; acciones continuas en [-1,1]^{d_ai} "
        "(BESS + EV + lavadora); recompensa mixta del perfil v4; script "
        "train_citylearn_v3_happo.py. Teoricamente es el candidato on-policy a la "
        "heterogeneidad tipologica del SEAI."
    ),
    (
        "MASAC — CityLearnSMACDiscreteEnv + external/MARL/src. El backend MASAC/mSAC asume "
        "API tipo SMAC (get_obs, get_state, acciones discretas). CityLearn expone control "
        "continuo multidimensional. Adecuacion: adaptador SMAC-like con get_state() = s "
        "CTDE; discretizacion action_bins=3 y discrete_action_mode=axis (base un-eje + "
        "no-op; no producto cartesiano 3^{d_ai}, inviable con d_ai hasta 44); mezcla "
        "cooperativa tipo QMIX sobre la recompensa energetica de equipo; script "
        "train_citylearn_v3_masac.py. Aporta exploracion entropica off-policy sobre un "
        "espacio de accion energetico tractable."
    ),
    (
        "MATD3 — CityLearnOffPolicyVecEnv + external/off-policy. El repositorio autor "
        "(external/MATD3implementation, TensorFlow 1.x / Python 3.6) no es el stack de "
        "entrenamiento: se adecua via backend PyTorch marlbenchmark/off-policy (clases "
        "MATD3 / R_MATD3). Adecuacion: wrapper vectorizado de un hilo; acciones continuas "
        "con doble critico, policy delay y target noise sobre actuadores BESS/EV; "
        "policy_mapping_fn y padding para 17 politicas heterogeneas; script "
        "train_citylearn_v3_matd3.py. Es el candidato deterministico off-policy al control "
        "continuo distrital."
    ),
    (
        "MAAC — CityLearnMAACVecEnv + external/MAAC. MAAC (Iqbal y Sha, 2019) introduce "
        "atencion multi-cabeza en el critico para dominios donde la relevancia entre "
        "agentes cambia; no trae de fabrica recompensa de pico/rampa/CO2/TOU. Adecuacion: "
        "wrapper que expone observaciones por agente al mecanismo de atencion; misma "
        "discretizacion eje-wise que MASAC (bins=3); atencion sobre los 17 edificios "
        "heterogeneos bajo recompensa v4; script train_citylearn_v3_maac.py. Operacionaliza "
        "coordinacion selectiva (p. ej. hospital vs mall vs campus) sin comunicacion en "
        "ejecucion."
    ),
]

TABLE_TITLE = (
    "Tabla 2.B — Adecuacion de los cuatro MADRL originales al problema electrico "
    "(OE.1 / OE.2 / OE.3)"
)

TABLE_NOTE = (
    "Nota. Los backends originales no fueron preparados ni entrenados ni aplicados "
    "out-of-the-box a flexibilidad, CO2 y costos energeticos. La adecuacion se sustenta "
    "en wrappers y adaptadores del proyecto (citylearn_v3_training_common.py; Cap. 4). "
    "Fuente: elaboracion propia a partir de external/HARL, external/MARL/src, "
    "external/off-policy, external/MAAC y la capa CityLearn v3 propuesto."
)

TABLE_ROWS: list[list[str]] = [
    [
        "Algoritmo",
        "Dominio tipico del original",
        "Wrapper / backend en el proyecto",
        "Adecuacion clave",
        "Capacidad resultante OE.1–OE.3",
    ],
    [
        "HAPPO",
        "MARL heterogeneo generico (HARL)",
        "CityLearnHARLEnv / external/HARL",
        "Share-obs CTDE; accion continua BESS/EV",
        "On-policy cooperativo sobre DER/EV heterogeneos",
    ],
    [
        "MASAC",
        "SMAC-like / accion discreta",
        "CityLearnSMACDiscreteEnv / MARL/src",
        "API SMAC; bins=3 axis; Q-mix energetico",
        "Off-policy entropico sobre control electrico discretizado",
    ],
    [
        "MATD3",
        "Paper TF1; activo: off-policy PyTorch",
        "CityLearnOffPolicyVecEnv / off-policy",
        "Backend PyTorch 3.9; accion continua",
        "Off-policy deterministico continuo BESS/EV",
    ],
    [
        "MAAC",
        "Atencion multiagente generica",
        "CityLearnMAACVecEnv / external/MAAC",
        "Atencion 17 agentes + bins=3",
        "Coordinacion atencional bajo flex/CO2/costo",
    ],
]

CLOSING = (
    "En sintesis: los cuatro MADRL originales no resuelven por si solos el problema "
    "electrico multiobjetivo; las adecuaciones del proyecto (wrappers + Dec-POMDP/CTDE + "
    "recompensa unified_comparable_v4 + dataset Iquitos + evaluate_v2) son la condicion "
    "de posibilidad teorica y operativa del cuasiexperimento 4x3 (Caps. 3–5)."
)


def set_run_font(run, bold: bool = False, size: float = 12.0) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, bold=bold)
    return new_para


def norm(s: str) -> str:
    t = (s or "").strip().lower()
    for a, b in (
        ("á", "a"),
        ("é", "e"),
        ("í", "i"),
        ("ó", "o"),
        ("ú", "u"),
        ("ñ", "n"),
    ):
        t = t.replace(a, b)
    return re.sub(r"\s+", " ", t)


def already_done(doc: Document) -> bool:
    for p in doc.paragraphs:
        if MARKER_NORM in norm(p.text):
            return True
        if "tabla 2.b" in norm(p.text) and "adecuacion de los cuatro madrl" in norm(p.text):
            return True
    return False


def find_anchor_after_229(doc: Document) -> Paragraph | None:
    """Return the last paragraph of §2.2.9 body (insert after it)."""
    body_229: Paragraph | None = None
    seen_heading = False
    for p in doc.paragraphs:
        t = norm(p.text)
        if not t:
            continue
        if t.startswith("2.2.9 ") or (
            "arquitectura marllib-ctde" in t and "universalidad" in t
        ):
            seen_heading = True
            continue
        if seen_heading:
            if t.startswith("2.2 ") or t.startswith("2.3") or t.startswith("2.2.10"):
                break
            if t.startswith("metodologia de seleccion") or t.startswith("variables de la investigacion"):
                break
            if t.startswith("capitulo 3") or t.startswith("3."):
                break
            body_229 = p
    return body_229


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 10.0) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=size)


def insert_table_after(paragraph: Paragraph, rows: list[list[str]]) -> Table:
    tbl = OxmlElement("w:tbl")
    tbl_pr = OxmlElement("w:tblPr")
    tbl.append(tbl_pr)
    tbl_grid = OxmlElement("w:tblGrid")
    for _ in rows[0]:
        tbl_grid.append(OxmlElement("w:gridCol"))
    tbl.append(tbl_grid)

    for r_i, row_data in enumerate(rows):
        tr = OxmlElement("w:tr")
        for cell_text in row_data:
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            tc.append(tc_pr)
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            if r_i == 0:
                b = OxmlElement("w:b")
                r_pr.append(b)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "20")
            r_pr.append(sz)
            r_fonts = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                r_fonts.set(qn(attr), "Times New Roman")
            r_pr.append(r_fonts)
            r.append(r_pr)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = cell_text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    paragraph._p.addnext(tbl)
    return Table(tbl, paragraph._parent)


def patch_doc(path: Path) -> dict:
    result: dict = {"path": str(path), "status": "pending", "changes": []}
    if not path.is_file() or path.stat().st_size <= 0:
        result["status"] = "missing"
        return result

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    bak = BACKUP_DIR / f"{path.name}.pre_madrl_adecuaciones_{stamp}.bak"
    shutil.copy2(path, bak)
    result["backup"] = str(bak)

    doc = Document(str(path))
    if already_done(doc):
        result["status"] = "already_done"
        return result

    anchor = find_anchor_after_229(doc)
    if anchor is None:
        result["status"] = "anchor_not_found"
        return result

    cursor = insert_paragraph_after(anchor, HEADING, bold=True)
    result["changes"].append("heading_2.2.10")

    for para in PARAS:
        cursor = insert_paragraph_after(cursor, para)
        result["changes"].append("para")

    cursor = insert_paragraph_after(cursor, TABLE_TITLE, bold=True)
    result["changes"].append("table_title")

    insert_table_after(cursor, TABLE_ROWS)
    # find the newly inserted table's following point: next sibling after title is table;
    # insert note + closing after the table by walking XML next.
    # Re-locate title paragraph and append after table element.
    title_para = None
    for p in doc.paragraphs:
        if TABLE_TITLE in (p.text or ""):
            title_para = p
            break
    if title_para is not None:
        # table is next sibling of title
        nxt = title_para._p.getnext()
        # create note paragraph after table
        note_p = OxmlElement("w:p")
        if nxt is not None:
            nxt.addnext(note_p)
        else:
            title_para._p.addnext(note_p)
        note_para = Paragraph(note_p, title_para._parent)
        run = note_para.add_run(TABLE_NOTE)
        set_run_font(run, bold=False, size=10.0)
        result["changes"].append("table_note")

        close_p = OxmlElement("w:p")
        note_p.addnext(close_p)
        close_para = Paragraph(close_p, title_para._parent)
        run2 = close_para.add_run(CLOSING)
        set_run_font(run2)
        result["changes"].append("closing")
        result["changes"].append("table")

    doc.save(str(path))
    result["status"] = "patched"
    result["size"] = path.stat().st_size
    return result


def main() -> int:
    results = []
    for path in (TESIS, INFORME):
        results.append(patch_doc(path))

    report = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "marker": MARKER,
        "results": results,
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    ok = all(r["status"] in {"patched", "already_done"} for r in results)
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
