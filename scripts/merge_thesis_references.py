#!/usr/bin/env python
"""Merge split reference lists into one APA section in thesis Word docs."""
from __future__ import annotations

import re
import shutil
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "scripts"))

from thesis_references_apa import _normalize_ref, _ref_key, load_all_thesis_references  # noqa: E402

PRIMARY = REPO / "docs" / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx"
MIRROR = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx"

REF_MAIN = "Referencias bibliograficas"
REF_COMP_PATTERNS = (
    "Referencias complementarias verificadas",
    "Referencias complementarias incorporadas",
    "Referencias complementarias",
)


def is_ref_heading(text: str) -> bool:
    t = text.strip().lower()
    return t == REF_MAIN.lower() or any(t.startswith(p.lower()) for p in REF_COMP_PATTERNS)


def is_apa_ref(text: str) -> bool:
    t = text.strip()
    return bool(re.match(r"^[A-Za-z]", t) and re.search(r"\(\d{4}", t))


def para_text(para: Paragraph) -> str:
    return (para.text or "").strip()


def heading_level(para: Paragraph) -> int | None:
    style = para.style.name if para.style else ""
    m = re.match(r"Heading\s+(\d+)", style, re.I)
    return int(m.group(1)) if m else None


def find_ref_regions(doc: Document) -> list[dict]:
    regions: list[dict] = []
    for i, para in enumerate(doc.paragraphs):
        t = para_text(para)
        if is_ref_heading(t):
            regions.append({"start_idx": i, "heading": t, "level": heading_level(para)})
    for j, region in enumerate(regions):
        start = region["start_idx"]
        end = regions[j + 1]["start_idx"] if j + 1 < len(regions) else len(doc.paragraphs)
        refs = []
        for k in range(start + 1, end):
            t = para_text(doc.paragraphs[k])
            if is_apa_ref(t):
                refs.append(_normalize_ref(t))
        region["end_idx"] = end
        region["refs"] = refs
        region["count"] = len(refs)
    return regions


def pick_best(existing: str, candidate: str) -> str:
    if len(candidate) > len(existing):
        return candidate
    if "http" in candidate.lower() and "http" not in existing.lower():
        return candidate
    if "[PV]" in existing.upper() and "[PV]" not in candidate.upper():
        return candidate
    return existing


def merge_all_refs(regions: list[dict], canonical: list[str]) -> list[str]:
    merged: dict[str, str] = {}
    for ref in canonical:
        merged[_ref_key(ref)] = ref
    for region in regions:
        for ref in region["refs"]:
            key = _ref_key(ref)
            if key in merged:
                merged[key] = pick_best(merged[key], ref)
            else:
                merged[key] = ref
    return sorted(merged.values(), key=lambda r: r.lower())


def apply_hanging_indent(para: Paragraph) -> None:
    pf = para.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def remove_paragraph(para: Paragraph) -> None:
    el = para._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def clone_paragraph_after(template: Paragraph, text: str) -> Paragraph:
    new_el = deepcopy(template._element)
    template._element.addnext(new_el)
    new_para = Paragraph(new_el, template._parent)
    for run in list(new_para.runs):
        run.text = ""
    if new_para.runs:
        new_para.runs[0].text = text
    else:
        new_para.add_run(text)
    apply_hanging_indent(new_para)
    return new_para


def merge_doc(path: Path, dry_run: bool = False) -> dict:
    doc = Document(str(path))
    regions = find_ref_regions(doc)
    before_lists = len(regions)
    before_total = sum(r["count"] for r in regions)

    canonical = load_all_thesis_references()
    unified = merge_all_refs(regions, canonical)
    final_count = len(unified)

    if dry_run:
        return {
            "file": str(path),
            "before_lists": before_lists,
            "before_total_raw": before_total,
            "after_count": final_count,
            "regions": [(r["heading"], r["count"]) for r in regions],
        }

    if not regions:
        raise RuntimeError(f"No reference sections found in {path}")

    main = regions[0]
    main_para = doc.paragraphs[main["start_idx"]]

    # Remove later reference-region paragraphs (headings + entries), keep anexos.
    for region in reversed(regions[1:]):
        for idx in range(region["end_idx"] - 1, region["start_idx"] - 1, -1):
            remove_paragraph(doc.paragraphs[idx])

    # Refresh indices after removals
    regions = find_ref_regions(doc)
    main = regions[0]
    main_para = doc.paragraphs[main["start_idx"]]

    # Remove old entries under main heading
    end = main["end_idx"]
    for idx in range(end - 1, main["start_idx"], -1):
        p = doc.paragraphs[idx]
        t = para_text(p)
        if is_apa_ref(t) or (not t and not heading_level(p)):
            remove_paragraph(p)

    # Re-find anchor after cleanup
    main_para = next(p for p in doc.paragraphs if para_text(p) == REF_MAIN)
    template = None
    anchor = main_para
    for para in doc.paragraphs:
        if para._element.getparent() is None:
            continue
        if para_text(para) == REF_MAIN:
            anchor = para
            continue
        if para_text(para) == REF_MAIN:
            break
        if is_apa_ref(para_text(para)):
            template = para
            break

    if template is None:
        template = anchor

    # Remove remaining body under main heading until next H1 (intro note + old refs).
    to_remove = []
    seen_main = False
    for para in doc.paragraphs:
        t = para_text(para)
        if t == REF_MAIN:
            seen_main = True
            continue
        if not seen_main:
            continue
        lvl = heading_level(para)
        if lvl == 1:
            break
        if t:
            to_remove.append(para)
    for para in reversed(to_remove):
        remove_paragraph(para)

    anchor = next(p for p in doc.paragraphs if para_text(p) == REF_MAIN)
    for ref in unified:
        anchor = clone_paragraph_after(anchor, ref)

    tmp = path.with_suffix(".docx.tmp_merged")
    doc.save(str(tmp))
    try:
        tmp.replace(path)
    except PermissionError:
        alt = path.with_name(path.stem + "_REFERENCIAS_UNIFICADAS.docx")
        tmp.replace(alt)
        raise PermissionError(
            f"Could not overwrite locked file {path.name}. Saved as {alt.name} — close Word and rerun."
        ) from None
    after_regions = find_ref_regions(Document(str(path)))
    return {
        "file": str(path),
        "before_lists": before_lists,
        "before_total_raw": before_total,
        "after_lists": len(after_regions),
        "after_count": final_count,
        "regions_before": [(r["heading"], r["count"]) for r in regions],
    }


def main() -> int:
    dry = "--dry-run" in sys.argv
    paths = [PRIMARY, MIRROR]
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if not dry:
        for p in paths:
            if p.is_file():
                bak = p.with_suffix(f".docx.bak_refs_merge_{stamp}")
                shutil.copyfile(p, bak)
                print(f"backup: {bak.name}")

    # merge primary first
    report_primary = merge_doc(PRIMARY, dry_run=dry)
    print("PRIMARY", report_primary)

    if not dry:
        shutil.copyfile(PRIMARY, MIRROR)
        report_mirror = merge_doc(MIRROR, dry_run=True)
        report_mirror["mirrored_from"] = PRIMARY.name
        print("MIRROR verified", report_mirror)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
