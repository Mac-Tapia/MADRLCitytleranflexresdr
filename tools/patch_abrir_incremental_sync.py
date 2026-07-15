#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Incremental patch for ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx.

Preserves existing body. Only:
  - regenerate Figura 5.8e PNG and swap binary near its caption
  - fill Tabla A.4 cells from real checkpoint_summary.csv if zeros/empty
  - update 5.8e caption if still claim-only-trace with dead columns
  - fix residual language artifacts (tesises, doctorales, doctorado)
  - leave 6.4/6.5 alone if already filled
"""
from __future__ import annotations

import io
import json
import re
import zipfile
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.oxml.ns import qn

REPO = Path(__file__).resolve().parents[1]
ABRIR = REPO / "docs" / "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx"
FIG_DIR = REPO / "outputs" / "_drive_madrl" / "gdrive_20260627_164047_objective_analysis" / "figures"
TABLE_DIR = REPO / "outputs" / "_drive_madrl" / "gdrive_20260627_164047_objective_analysis" / "tables"
FULL_DATA = REPO / "outputs" / "_drive_madrl" / "full_data"
CK_CSV = FULL_DATA / "analysis_real_drive" / "tables" / "checkpoint_summary.csv"
REPORT = FIG_DIR.parent / "abrir_incremental_sync_report.json"

ALGOS = ["HAPPO", "MAAC", "MASAC", "MATD3"]
SCENARIOS = ["E1", "E2", "E3"]

LANG_FIXES = [
    (re.compile(r"\bejes de tesises\b", re.I), "ejes de la tesis"),
    (re.compile(r"\btesises\b", re.I), "tesis"),
    (re.compile(r"\bdoctorales\b", re.I), "de la tesis"),
    (re.compile(r"\bformación doctoral\b", re.I), "formación de la tesis"),
    (re.compile(r"\bformacion doctoral\b", re.I), "formación de la tesis"),
    (re.compile(r"\bestudio doctoral\b", re.I), "estudio de tesis"),
    (re.compile(r"\btesis doctoral\b", re.I), "tesis"),
    (re.compile(r"\bcierre doctoral\b", re.I), "cierre de la tesis"),
    (re.compile(r"\bdoctorado\b", re.I), "posgrado"),
    (re.compile(r"\bdoctoral\b", re.I), "de tesis"),
    (re.compile(r"Politicas y acciones medias desde trace\.csv", re.I),
     "Politicas y acciones medias (action_l2 desde trace.csv; EV/BESS desde building_behavior_summary)"),
    (re.compile(r"Políticas y acciones medias desde trace\.csv", re.I),
     "Políticas y acciones medias (action_l2 desde trace.csv; EV/BESS desde building_behavior_summary)"),
]


def set_cell_text(cell, text: str) -> None:
    """Replace cell text preserving first paragraph formatting lightly."""
    paras = cell.paragraphs
    if not paras:
        cell.text = text
        return
    # Clear all runs in first para; clear extra paras
    p0 = paras[0]
    if p0.runs:
        p0.runs[0].text = text
        for r in p0.runs[1:]:
            r.text = ""
    else:
        p0.add_run(text)
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def patch_paragraph_text(para, patterns) -> int:
    full = para.text or ""
    if not full:
        return 0
    new = full
    n = 0
    for rx, repl in patterns:
        new2, k = rx.subn(repl, new)
        if k:
            n += k
            new = new2
    if n and new != full:
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.add_run(new)
    return n


def load_traces_action_l2() -> pd.DataFrame:
    frames = []
    for algo in ALGOS:
        for scen in SCENARIOS:
            # Prefer full_data (has HAPPO); fallback local madrl mirror
            candidates = [
                FULL_DATA / algo / scen / "data" / "trace.csv",
                REPO / "outputs" / "madrl_v3_20260627_164047" / algo / scen / "data" / "trace.csv",
            ]
            path = next((p for p in candidates if p.exists() and p.stat().st_size > 100), None)
            if path is None:
                continue
            try:
                df = pd.read_csv(path, usecols=lambda c: c in {"action_l2", "action_mean"})
            except Exception:
                continue
            if df.empty or "action_l2" not in df.columns:
                continue
            frames.append(
                pd.DataFrame(
                    {
                        "algorithm": [algo],
                        "scenario": [scen],
                        "action_l2": [pd.to_numeric(df["action_l2"], errors="coerce").mean()],
                    }
                )
            )
    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()


def load_behavior_ev_bess() -> pd.DataFrame:
    """Aggregate per algorithm/scenario EV and BESS metrics from behavior summaries."""
    rows = []
    for algo in ALGOS:
        for scen in SCENARIOS:
            path = FULL_DATA / algo / scen / "data" / "building_behavior_summary.csv"
            if not path.exists():
                # try consolidated
                continue
            df = pd.read_csv(path)
            # Prefer real non-dead columns
            ev_cols = [c for c in df.columns if c in {
                "ev_charge_total_kwh", "ev_charger_energy_total", "electrical_vehicle_charger_energy",
                "ev_charge_kwh", "ev_charged_kwh",
            } or ("ev" in c.lower() and "charge" in c.lower() and "kwh" in c.lower())]
            bess_cols = [c for c in df.columns if c in {
                "battery_throughput_total_kwh", "electrical_storage_throughput_kwh",
                "battery_throughput", "electrical_storage_energy_balance_kwh",
            } or ("throughput" in c.lower() and ("batt" in c.lower() or "storage" in c.lower()))]
            # SOC mean if available as secondary
            soc_cols = [c for c in df.columns if "soc" in c.lower() and "storage" in c.lower()]

            def pick_mean(cols):
                for c in cols:
                    s = pd.to_numeric(df[c], errors="coerce")
                    if s.notna().any() and (s.abs() > 1e-12).any():
                        return float(s.mean()), c
                return float("nan"), None

            ev_mean, ev_src = pick_mean(ev_cols)
            bess_mean, bess_src = pick_mean(bess_cols)
            if np.isnan(bess_mean):
                bess_mean, bess_src = pick_mean(soc_cols)
            rows.append(
                {
                    "algorithm": algo,
                    "scenario": scen,
                    "ev_metric": ev_mean,
                    "ev_src": ev_src,
                    "bess_metric": bess_mean,
                    "bess_src": bess_src,
                }
            )
    out = pd.DataFrame(rows)
    # Also try consolidated file
    cons = TABLE_DIR / "gdrive_building_behavior_summary_all.csv"
    if (out.empty or out["ev_metric"].isna().all()) and cons.exists():
        df = pd.read_csv(cons)
        rows = []
        for algo in ALGOS:
            for scen in SCENARIOS:
                sub = df[(df.get("algorithm") == algo) & (df.get("scenario") == scen)] if "algorithm" in df.columns else df.iloc[0:0]
                if sub.empty:
                    continue
                ev_cols = [c for c in sub.columns if "ev" in c.lower() and "charge" in c.lower()]
                bess_cols = [c for c in sub.columns if "throughput" in c.lower() or ("battery" in c.lower() and "kwh" in c.lower())]
                def mean_first(cols):
                    for c in cols:
                        s = pd.to_numeric(sub[c], errors="coerce")
                        if s.notna().any() and (s.abs() > 1e-12).any():
                            return float(s.mean()), c
                    return float("nan"), None
                ev_m, ev_s = mean_first(ev_cols)
                be_m, be_s = mean_first(bess_cols)
                rows.append({"algorithm": algo, "scenario": scen, "ev_metric": ev_m, "ev_src": ev_s, "bess_metric": be_m, "bess_src": be_s})
        if rows:
            out = pd.DataFrame(rows)
    return out


def regenerate_fig_58e() -> tuple[Path, dict]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    action = load_traces_action_l2()
    behav = load_behavior_ev_bess()
    out = FIG_DIR / "trace_policy_action_heatmaps.png"
    verify = {"action_l2": {}, "ev": {}, "bess": {}, "sources": {}}

    panels = [
        ("action_l2", "Intensidad de accion (mean action_l2)", action, "action_l2"),
        ("ev_metric", "Carga EV (behavior summary)", behav, "ev_metric"),
        ("bess_metric", "Throughput/SOC BESS (behavior summary)", behav, "bess_metric"),
    ]

    fig, axes = plt.subplots(1, 3, figsize=(13.2, 4.2))
    for ax, (key, title, df, col) in zip(axes, panels):
        pivot = pd.DataFrame(index=ALGOS, columns=SCENARIOS, dtype=float)
        if not df.empty and col in df.columns:
            for _, r in df.iterrows():
                pivot.loc[r["algorithm"], r["scenario"]] = r[col]
        vals = pivot.values.astype(float)
        verify[key if key != "ev_metric" else "ev" if key.startswith("ev") else key] = {
            f"{a}-{s}": (None if pd.isna(pivot.loc[a, s]) else float(pivot.loc[a, s]))
            for a in ALGOS
            for s in SCENARIOS
        }
        # store under clear names
        if col == "action_l2":
            verify["action_l2"] = {f"{a}-{s}": (None if pd.isna(pivot.loc[a, s]) else float(pivot.loc[a, s])) for a in ALGOS for s in SCENARIOS}
        elif col == "ev_metric":
            verify["ev"] = {f"{a}-{s}": (None if pd.isna(pivot.loc[a, s]) else float(pivot.loc[a, s])) for a in ALGOS for s in SCENARIOS}
            if not behav.empty and "ev_src" in behav.columns:
                verify["sources"]["ev"] = sorted({x for x in behav["ev_src"].dropna().unique()})
        else:
            verify["bess"] = {f"{a}-{s}": (None if pd.isna(pivot.loc[a, s]) else float(pivot.loc[a, s])) for a in ALGOS for s in SCENARIOS}
            if not behav.empty and "bess_src" in behav.columns:
                verify["sources"]["bess"] = sorted({x for x in behav["bess_src"].dropna().unique()})

        finite = vals[np.isfinite(vals)]
        if finite.size:
            vmin, vmax = float(np.nanmin(finite)), float(np.nanmax(finite))
            if abs(vmax - vmin) < 1e-12:
                vmax = vmin + 1e-6
        else:
            vmin, vmax = 0.0, 1.0
        img = ax.imshow(vals, cmap="plasma", vmin=vmin, vmax=vmax, aspect="auto")
        ax.set_xticks(range(len(SCENARIOS)), SCENARIOS)
        ax.set_yticks(range(len(ALGOS)), ALGOS)
        ax.set_title(title, fontsize=9)
        for i in range(len(ALGOS)):
            for j in range(len(SCENARIOS)):
                v = vals[i, j]
                label = "NA" if not np.isfinite(v) else f"{v:.2f}"
                ax.text(j, i, label, ha="center", va="center", color="white", fontsize=7)
        fig.colorbar(img, ax=ax, fraction=0.046, pad=0.04)
    fig.suptitle("Figura 5.8e — politicas/acciones (fuentes mixtas auditadas)", fontsize=11)
    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(
        [
            {"panel": "action_l2", "algorithm": a, "scenario": s, "value": verify["action_l2"].get(f"{a}-{s}")}
            for a in ALGOS
            for s in SCENARIOS
        ]
        + [
            {"panel": "ev", "algorithm": a, "scenario": s, "value": verify["ev"].get(f"{a}-{s}")}
            for a in ALGOS
            for s in SCENARIOS
        ]
        + [
            {"panel": "bess", "algorithm": a, "scenario": s, "value": verify["bess"].get(f"{a}-{s}")}
            for a in ALGOS
            for s in SCENARIOS
        ]
    ).to_csv(TABLE_DIR / "fig58e_panel_values.csv", index=False)
    return out, verify


def find_table_after_caption(doc: Document, caption_substr: str):
    body = list(doc.element.body)
    para_idxs = {}
    for i, p in enumerate(doc.paragraphs):
        para_idxs[id(p._element)] = i
    # Walk body sequentially
    captions = []
    for el in body:
        if el.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in el.iter(qn("w:t")))
        if caption_substr in text:
            captions.append(el)
    for cap_el in captions:
        # next table sibling
        nxt = cap_el.getnext()
        while nxt is not None:
            if nxt.tag == qn("w:tbl"):
                # map to doc.tables
                for ti, tbl in enumerate(doc.tables):
                    if tbl._tbl is nxt:
                        return ti, tbl, "".join(t.text or "" for t in cap_el.iter(qn("w:t")))
            if nxt.tag == qn("w:p"):
                t = "".join(x.text or "" for x in nxt.iter(qn("w:t"))).strip()
                if t.startswith("Tabla ") or t.startswith("Figura ") or t.startswith("A."):
                    break
            nxt = nxt.getnext()
    return None, None, None


def table_looks_all_zero(tbl) -> bool:
    nums = []
    for ri, row in enumerate(tbl.rows):
        if ri == 0:
            continue
        for cell in row.cells:
            t = (cell.text or "").strip().replace(",", "")
            try:
                nums.append(float(t))
            except ValueError:
                continue
    return bool(nums) and all(abs(x) < 1e-12 for x in nums)


def fill_tabla_a4(doc: Document) -> dict:
    info = {"found": False, "updated": False, "reason": ""}
    ti, tbl, cap = find_table_after_caption(doc, "Tabla A.4")
    if tbl is None:
        # fallback: search headers
        for ti2, t2 in enumerate(doc.tables):
            hdr = " ".join(c.text.strip().lower() for c in t2.rows[0].cells)
            if "checkpoint" in hdr and ("algoritmo" in hdr or "algorithm" in hdr):
                # Prefer summary-like headers
                if any(k in hdr for k in ["archivos", "bytes", "listed", "declared", "count", "manifiest"]):
                    ti, tbl, cap = ti2, t2, hdr
                    break
        if tbl is None:
            # last resort: table nearest to A.4 caption paragraph index
            for i, p in enumerate(doc.paragraphs):
                if "Tabla A.4" in (p.text or ""):
                    # find first table element after this para in body order — handled above; if miss, scan by content zeros
                    pass
            info["reason"] = "Tabla A.4 no localizada"
            return info

    info["found"] = True
    info["table_index"] = ti
    info["caption"] = (cap or "")[:120]
    info["zeros_before"] = table_looks_all_zero(tbl)

    if not CK_CSV.exists():
        info["reason"] = f"Falta {CK_CSV}"
        return info
    ck = pd.read_csv(CK_CSV)

    # Decide columns based on existing header if compatible; else rewrite header+rows keeping style
    desired_headers = [
        "Algoritmo",
        "Escenario",
        "Backend",
        "Checkpoints declarados",
        "Archivos listados",
        "Tipos",
        "Bytes totales",
    ]
    # HapPO missing from ck file historically — note N/A rows
    hdr_now = [c.text.strip() for c in tbl.rows[0].cells]
    # Resize: if column count differs, rewrite using existing ncol if possible
    ncol = len(tbl.rows[0].cells)
    if ncol >= 5:
        # Map into available columns
        headers = desired_headers[:ncol]
        # Ensure enough rows: 1 header + len(ck) (+ HAPPO placeholders if missing)
        algorithms_in_ck = set(ck["algorithm"])
        extra = []
        for algo in ALGOS:
            if algo not in algorithms_in_ck:
                for scen in SCENARIOS:
                    extra.append(
                        {
                            "algorithm": algo,
                            "scenario": scen,
                            "backend": "N/D (sin checkpoint_manifest.json en Drive)",
                            "checkpoint_count_declared": 0,
                            "checkpoint_files_listed": 0,
                            "checkpoint_file_types": "—",
                            "checkpoint_bytes_total": 0,
                        }
                    )
        ck_all = pd.concat([ck, pd.DataFrame(extra)], ignore_index=True)
        ck_all["algorithm"] = pd.Categorical(ck_all["algorithm"], ALGOS, ordered=True)
        ck_all["scenario"] = pd.Categorical(ck_all["scenario"], SCENARIOS, ordered=True)
        ck_all = ck_all.sort_values(["algorithm", "scenario"])

        needed = 1 + len(ck_all)
        # Add rows if short
        while len(tbl.rows) < needed:
            tbl.add_row()
        # Clear surplus text but keep extra rows empty
        for j, h in enumerate(headers):
            if j < ncol:
                set_cell_text(tbl.rows[0].cells[j], h)
        for ri, (_, r) in enumerate(ck_all.iterrows(), start=1):
            vals = [
                str(r["algorithm"]),
                str(r["scenario"]),
                str(r.get("backend", "")),
                str(int(r["checkpoint_count_declared"]) if pd.notna(r["checkpoint_count_declared"]) else 0),
                str(int(r["checkpoint_files_listed"]) if pd.notna(r["checkpoint_files_listed"]) else 0),
                str(r.get("checkpoint_file_types", "")),
                f"{int(r['checkpoint_bytes_total']):,}".replace(",", " "),
            ]
            for j in range(min(ncol, len(vals))):
                set_cell_text(tbl.rows[ri].cells[j], vals[j])
        # Blank leftover rows
        for ri in range(needed, len(tbl.rows)):
            for j in range(ncol):
                set_cell_text(tbl.rows[ri].cells[j], "")
        info["updated"] = True
        info["rows_written"] = len(ck_all)
        info["zeros_after"] = table_looks_all_zero(tbl)
        info["sample_rows"] = [
            [c.text.strip() for c in tbl.rows[i].cells[:ncol]] for i in range(min(5, len(tbl.rows)))
        ]
    else:
        info["reason"] = f"Tabla A.4 con ncol inesperado={ncol}"
    return info


def replace_image_near_caption(docx_path: Path, caption_key: str, png_path: Path) -> dict:
    """Replace the first image relationship that appears shortly after caption in document.xml order."""
    info = {"replaced": False}
    # Open as zip; locate document.xml caption, then next a:blip r:embed
    with zipfile.ZipFile(docx_path, "r") as zin:
        xml = zin.read("word/document.xml").decode("utf-8")
        rels = zin.read("word/_rels/document.xml.rels").decode("utf-8")

    # Find caption position (strip tags roughly by searching plain text after removing tags for locate)
    text_only = re.sub(r"<[^>]+>", "", xml)
    # Caption in ABRIR is often split across lines: "Figura 5.8e\nPoliticas..."
    idx = text_only.find("Figura 5.8e")
    if idx < 0:
        info["reason"] = "caption Figura 5.8e no encontrada en document.xml"
        return info

    # Map approximate: find blip after caption occurrence in XML
    # Search for 'Figura 5.8e' in xml (may be split across w:t)
    # Strategy: find all drawings with rId, pick the drawing that occurs after first w:t containing 5.8e
    m = re.search(r"5\.8e", xml)
    if not m:
        info["reason"] = "5.8e no encontrado en XML"
        return info
    after = xml[m.end() : m.end() + 25000]
    blips = re.findall(r'a:blip[^>]+r:embed="(rId\d+)"', after)
    if not blips:
        # broader search: next few blips in remaining doc
        blips = re.findall(r'a:blip[^>]+r:embed="(rId\d+)"', xml[m.end() :])
    if not blips:
        info["reason"] = "No se halló a:blip después de 5.8e"
        return info
    rid = blips[0]
    rm = re.search(rf'Target="([^"]+)"[^>]*Id="{rid}"|Id="{rid}"[^>]*Target="([^"]+)"', rels)
    if not rm:
        # relationships usually Id then Target
        rm = re.search(rf'Relationship[^>]*Id="{rid}"[^>]*Target="([^"]+)"', rels)
    if not rm:
        info["reason"] = f"No Target para {rid}"
        return info
    target = rm.group(1) if rm.lastindex and rm.group(1) else rm.group(rm.lastindex)
    media_path = "word/" + target.lstrip("/")
    if not media_path.startswith("word/"):
        media_path = "word/" + target

    png_bytes = png_path.read_bytes()
    # Rewrite zip replacing that media part
    tmp = docx_path.with_suffix(".docx.tmp_58e")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        replaced = False
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.replace("\\", "/") == media_path.replace("\\", "/"):
                data = png_bytes
                replaced = True
            zout.writestr(item, data)
        info["replaced"] = replaced
        info["rId"] = rid
        info["media"] = media_path
        info["png_bytes"] = len(png_bytes)
    if not info["replaced"]:
        tmp.unlink(missing_ok=True)
        info["reason"] = f"Parte {media_path} no estaba en el zip"
        return info
    tmp.replace(docx_path)
    return info


def incremental_language_cleanup(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        n += patch_paragraph_text(p, LANG_FIXES)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n += patch_paragraph_text(p, LANG_FIXES)
    return n


def verify_abrir(doc: Document, png_verify: dict) -> dict:
    text = "\n".join(p.text or "" for p in doc.paragraphs)
    return {
        "has_6_4": "6.4 Plan para culminar la tesis" in text,
        "has_6_5": "6.5 Criterios de cierre" in text,
        "has_5_8e_caption": "Figura 5.8e" in text,
        "has_a4_caption": "Tabla A.4" in text,
        "no_doctorado": not bool(re.search(r"doctorado|doctoral", text, re.I)),
        "has_tesises_artifact": "tesises" in text.lower(),
        "happo_action_present": any(
            (png_verify.get("action_l2") or {}).get(f"HAPPO-{s}") not in (None, 0) for s in SCENARIOS
        ),
        "words_approx_6_4": len(re.findall(r"\w+", text.split("6.4 Plan")[1].split("6.5 ")[0])) if "6.4 Plan" in text and "6.5 " in text else None,
    }


def main() -> int:
    actions = []
    report: dict = {"target": str(ABRIR), "actions": actions}

    # 1) Regenerate figure
    png_path, png_verify = regenerate_fig_58e()
    actions.append({"regen_58e": str(png_path), "verify": png_verify})

    # 2) Open ABRIR and patch tables/language (preserve rest)
    doc = Document(str(ABRIR))
    a4 = fill_tabla_a4(doc)
    actions.append({"tabla_a4": a4})

    # Caption update for 5.8e if needed
    cap_n = 0
    for p in doc.paragraphs:
        t = p.text or ""
        if "Figura 5.8e" in t or ("Politicas y acciones" in t and "trace.csv" in t) or ("Políticas y acciones" in t and "trace.csv" in t):
            cap_n += patch_paragraph_text(p, LANG_FIXES)
    # also adjacent paragraph after heading-like caption
    for i, p in enumerate(doc.paragraphs):
        if "Figura 5.8e" in (p.text or ""):
            # next non-empty
            for q in doc.paragraphs[i : i + 3]:
                cap_n += patch_paragraph_text(q, LANG_FIXES)
    actions.append({"caption_58e_patches": cap_n})

    lang_n = incremental_language_cleanup(doc)
    actions.append({"language_fixes": lang_n})

    # Do NOT rebuild 6.4/6.5 — already present
    actions.append({"cap6": "preserved_existing_filled_sections"})

    # Save docx (table/language) then swap image binary
    try:
        doc.save(str(ABRIR))
        actions.append({"saved_docx": True})
    except PermissionError:
        alt = ABRIR.with_name(ABRIR.stem + "_INCREMENTAL_PATCH.docx")
        doc.save(str(alt))
        actions.append({"saved_docx": False, "saved_alt": str(alt), "note": "Archivo abierto/bloqueado; cierre Word y renombre el _INCREMENTAL_PATCH"})
        report["blocked"] = True
        report["alt"] = str(alt)
        # continue image replace on alt
        target = alt
    else:
        target = ABRIR

    img = replace_image_near_caption(target, "5.8e", png_path)
    actions.append({"image_58e": img})

    # Re-open verify
    doc2 = Document(str(target))
    v = verify_abrir(doc2, png_verify)
    # re-check A.4 zeros
    _, tbl, _ = find_table_after_caption(doc2, "Tabla A.4")
    v["a4_zeros"] = table_looks_all_zero(tbl) if tbl is not None else None
    if tbl is not None:
        v["a4_preview"] = [[c.text.strip() for c in row.cells[:7]] for row in tbl.rows[:4]]
    report["verification"] = v
    report["png_path"] = str(png_path)
    report["png_size"] = png_path.stat().st_size
    report["target_final"] = str(target)

    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
