# -*- coding: utf-8 -*-
"""Generate the thesis-plan DOCX for the CityLearn v3 MADRL project."""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path
from typing import Iterable, Mapping, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
CITYLEARN_ROOT = ROOT / "CityLearn"
DOCS_DIR = ROOT / "docs"
OUTPUT_PATH = DOCS_DIR / "PLAN_TESIS_MADRL_CITYLEARN_V3.docx"
PDF_PATH = ROOT / "Plan_Tesis_MADRL_Diagnostico_v17.pdf"
LOCK_PATH = ROOT / "external" / "backends.lock.json"
STATUS_PATH = ROOT / "outputs" / "citylearn_v3_madrl_official_full_cuda_v2" / "official_full_status.json"
SCHEMA_PATH = CITYLEARN_ROOT / "data" / "datasets" / "citylearn_challenge_2022_phase_all_plus_evs" / "schema.json"

sys.path.insert(0, str(CITYLEARN_ROOT))
from citylearn.v3.objectives import objective_manifest  # noqa: E402


TITLE = (
    "MULTI-AGENTE DE APRENDIZAJE POR REFUERZO PROFUNDO PARA GESTIÓN "
    "COORDINADA DE FLEXIBILIDAD ENERGÉTICA, EMISIONES DE CARBONO Y "
    "EFICIENCIA ECONÓMICA EN COMUNIDADES INTELIGENTES"
)


def _read_json(path: Path) -> dict:
    if not path.exists():
        return {}

    return json.loads(path.read_text(encoding="utf-8-sig"))


def _schema_summary() -> dict:
    schema = _read_json(SCHEMA_PATH)
    buildings = [
        name
        for name, payload in schema.get("buildings", {}).items()
        if payload.get("include", True)
    ]
    observations = [
        name
        for name, payload in schema.get("observations", {}).items()
        if payload.get("active", False)
    ]
    actions = [
        name
        for name, payload in schema.get("actions", {}).items()
        if payload.get("active", False)
    ]

    return {
        "dataset": SCHEMA_PATH.parent.name,
        "schema_path": str(SCHEMA_PATH.relative_to(ROOT)),
        "buildings": buildings,
        "building_count": len(buildings),
        "simulation_start": schema.get("simulation_start_time_step"),
        "simulation_end": schema.get("simulation_end_time_step"),
        "time_steps": (
            schema.get("simulation_end_time_step", 0)
            - schema.get("simulation_start_time_step", 0)
            + 1
        ),
        "observations": observations,
        "observation_count": len(observations),
        "actions": actions,
        "action_count": len(actions),
        "has_ev": any("electric_vehicle" in item for item in observations + actions),
    }


def _pdf_page_count() -> int | str:
    if not PDF_PATH.exists():
        return "no disponible"

    return len(PdfReader(str(PDF_PATH)).pages)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    *,
    style: str = "Table Grid",
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = style
    header_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        _set_cell_text(header_cells[i], header, bold=True)
        _set_cell_shading(header_cells[i], "D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            _set_cell_text(cells[i], str(value))

    doc.add_paragraph()


def _add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _add_modification_registry(doc: Document) -> None:
    doc.add_heading(
        "4.8.1 Registro de modificaciones, mejoras y adecuaciones CityLearn v2 -> CityLearn v3",
        level=3,
    )
    _add_paragraphs(
        doc,
        [
            (
                "Este numeral documenta las intervenciones realizadas sobre el proyecto para dejar funcional "
                "el entrenamiento MADRL en CityLearn v3, manteniendo CityLearn v2 como base oficial de datos, "
                "simulación física, baseline y KPIs. La intención metodológica es que cada mejora sea "
                "auditada como parte de la solución y no como una reimplementación informal de los algoritmos."
            ),
            (
                "La adecuación respeta el principio central del proyecto: CityLearn v3 es una capa experimental "
                "MADRL sobre CityLearn v2. Por tanto, los datasets, el `schema.json`, los espacios de acción y "
                "observación, los edificios, PV, baterías, EVs, precios, intensidad de carbono y evaluación "
                "`evaluate_v2` siguen siendo la fuente técnica de referencia."
            ),
        ],
    )
    _add_table(
        doc,
        ["Área intervenida", "Adecuación implementada", "Evidencia en el proyecto"],
        [
            (
                "Arquitectura v3",
                "Creación de la capa `citylearn.v3` para configuración, entorno, objetivos, backends y adaptadores.",
                "`CityLearn/citylearn/v3/config.py`, `environment.py`, `objectives.py`, `backends.py`.",
            ),
            (
                "Dec-POMDP",
                "Cada edificio opera como agente descentralizado con observación local, acción local y recompensa colaborativa.",
                "`build_citylearn_v3_env`, `build_default_17_building_ev_env`, `CityLearnV3MARLlibEnv`.",
            ),
            (
                "CTDE",
                "Exposición de estado global para entrenamiento centralizado y ejecución con políticas locales por edificio.",
                "Uso de `state()`, `state_space`, `share_observation_space` y adaptadores por backend.",
            ),
            (
                "Dataset de tesis",
                "Definición del dataset por defecto con 17 edificios, EVs, 8760 pasos horarios y soporte para otros schemas CityLearn v2.",
                "`citylearn_challenge_2022_phase_all_plus_evs/schema.json`.",
            ),
            (
                "Backends oficiales",
                "Registro de HAPPO, MASAC/mSAC, MATD3, MATD3 PyTorch, MAAC y MARLlib, evitando duplicar implementaciones MADRL dentro de `citylearn.agents`.",
                "`CityLearn/citylearn/official_madrl.py` y `external/backends.lock.json`.",
            ),
            (
                "Scripts de entrenamiento",
                "Creación de scripts independientes para cada MADRL con parámetros, CUDA, semillas, escenario, episodios y salida reproducible.",
                "`train_citylearn_v3_happo.py`, `train_citylearn_v3_masac.py`, `train_citylearn_v3_matd3.py`, `train_citylearn_v3_maac.py`.",
            ),
            (
                "Lanzamiento oficial",
                "Automatización secuencial de entrenamiento oficial para reducir conflictos de memoria GPU y registrar estado global.",
                "`launch_citylearn_v3_official_training.ps1`, `official_full_status.json`, `official_full_manifest.json`.",
            ),
            (
                "CUDA y Python 3.9",
                "Preparación del entorno `.venv39-citylearn-v3` con PyTorch CUDA para entrenamiento acelerado.",
                "Estado reportado: `torch 2.8.0+cu126`, `cuda=true`.",
            ),
            (
                "MATD3 funcional",
                "Ajustes para backend PyTorch compatible con Python 3.9, manejo correcto de dispositivo/dtype y módulos del crítico.",
                "`external/off-policy` como backend PyTorch source-backed; MATD3 original queda como referencia del paper.",
            ),
            (
                "Objetivos científicos",
                "Reformulación del proyecto en tres ejes: OE1 flexibilidad, OE2 emisiones de CO2 y OE3 costos energéticos.",
                "`CityLearn/citylearn/v3/objectives.py` y `ESTRATEGIA_3PILARES_MADRL.md`.",
            ),
            (
                "KPIs por eje",
                "Exposición de 36 KPIs para OE1, 7 para OE2 y 11 para OE3, con fuente, dirección de mejora y relación con baseline.",
                "`objective_manifest()`, `objective_kpis.csv`, notebook MADRL CityLearn v3.",
            ),
            (
                "KPI derivado",
                "Inclusión de `price_signal_deviation` como KPI derivado desde importación neta distrital y tarifa dinámica.",
                "`price_signal_deviation()` en `CityLearn/citylearn/v3/objectives.py`.",
            ),
            (
                "Artefactos reproducibles",
                "Estandarización de carpetas por algoritmo con resultados, series, trazas, checkpoints, figuras y tablas.",
                "`data/results.json`, `data/timeseries.csv`, `data/trace.csv`, `data/checkpoint_manifest.json`, `checkpoints/`, `figures/`.",
            ),
            (
                "Figuras y cuadros",
                "Generación de recompensas, convergencia, returns, eficiencia, comparación baseline, perfiles por eje y tablas CSV/Markdown.",
                "`regenerate_citylearn_v3_figures.py` y `figures/figures_manifest.json`.",
            ),
            (
                "Tutorial académico",
                "Creación de notebook CityLearn v3 MADRL con diagnóstico, problemática, marco teórico, backends, papers, GitHub y KPIs.",
                "`CityLearn/examples/madrl_citylearn_v3_tutorial.ipynb`.",
            ),
            (
                "Validaciones",
                "Pruebas de construcción del entorno, KPIs, imports de backends, smoke tests y entrenamiento de 5 episodios.",
                "`validate_citylearn_v3_objectives.py`, logs smoke y salidas en `outputs/`.",
            ),
            (
                "Repositorio público",
                "Versionamiento del proyecto raíz y fork CityLearn con commits trazables para reproducibilidad.",
                "`Mac-Tapia/MADRLCitytleranflexresdr` y `Mac-Tapia/CityLearn`.",
            ),
        ],
    )


def _add_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


def _set_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    for style_name, size in [
        ("Title", 16),
        ("Heading 1", 14),
        ("Heading 2", 13),
        ("Heading 3", 12),
        ("Heading 4", 11),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True


def _setup_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)


def _add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Plan de tesis - CityLearn v3 MADRL"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = "Documento generado desde el proyecto MADRLCitytleranflexresdr"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_cover(doc: Document, schema: Mapping[str, object], status: Mapping[str, object]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD NACIONAL DE INGENIERÍA\nESCUELA DE POSGRADO")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PLAN DE TESIS DE MAESTRÍA DE ESPECIALIZACIÓN")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"“{TITLE}”")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PARA OPTAR EL GRADO ACADÉMICO DE\n").bold = True
    p.add_run("MAESTRO EN CIENCIAS CON MENCIÓN EN INGENIERÍA DE SISTEMAS")

    doc.add_paragraph()
    _add_table(
        doc,
        ["Campo", "Dato"],
        [
            ("Tesista", "[Nombre del tesista]"),
            ("Asesor", "[Nombre del asesor]"),
            ("Línea de investigación", "Inteligencia artificial aplicada a sistemas energéticos urbanos"),
            ("Proyecto computacional", "CityLearn v3 MADRL sobre CityLearn v2"),
            ("Dataset de tesis", schema["dataset"]),
            ("Caso experimental", f"{schema['building_count']} edificios + EV, {schema['time_steps']} pasos horarios"),
            ("Backends MADRL", "HAPPO, MASAC/mSAC, MATD3 y MAAC"),
            ("Estado de entrenamiento oficial", status.get("status", "no disponible")),
            ("Lugar y fecha", f"Lima, Perú - {date.today().isoformat()}"),
        ],
    )
    doc.add_page_break()


def _add_preliminary(doc: Document, schema: Mapping[str, object], status: Mapping[str, object]) -> None:
    doc.add_heading("I. CUERPO PRELIMINAR", level=1)
    doc.add_heading("Datos Generales", level=2)
    _add_table(
        doc,
        ["Elemento", "Descripción"],
        [
            ("Título", TITLE),
            ("Tipo de documento", "Plan de tesis de maestría de especialización"),
            ("Problema real", "Coordinación energética multiobjetivo en comunidades de edificios con recursos distribuidos, EVs y señales dinámicas."),
            ("Solución propuesta", "Capa CityLearn v3 MADRL con Dec-POMDP, CTDE, backends oficiales y evaluación con KPIs CityLearn v2."),
            ("Dataset", schema["dataset"]),
            ("Esquema", schema["schema_path"]),
            ("Edificios", schema["building_count"]),
            ("Horizonte", f"{schema['simulation_start']} a {schema['simulation_end']} ({schema['time_steps']} pasos)"),
            ("Observaciones activas", schema["observation_count"]),
            ("Acciones activas", f"{schema['action_count']}: {', '.join(schema['actions'])}"),
            ("EV incluido", "Sí, por observaciones/acciones `electric_vehicle_*`" if schema["has_ev"] else "No detectado"),
            ("Documento PDF base revisado", f"{PDF_PATH.name}, { _pdf_page_count() } páginas"),
            ("Entrenamiento oficial actual", f"{status.get('status', 'no disponible')} en {status.get('output_root', 'no disponible')}"),
        ],
    )

    doc.add_heading("Resumen Ejecutivo", level=2)
    _add_paragraphs(
        doc,
        [
            (
                "El presente plan de tesis propone una solución computacional para la gestión coordinada de "
                "flexibilidad energética, emisiones de carbono y eficiencia económica en comunidades "
                "inteligentes. La propuesta conserva CityLearn v2 como simulador, fuente de datos, física "
                "de edificios, baterías, PV, EVs y KPIs, y agrega una capa CityLearn v3 orientada a "
                "aprendizaje por refuerzo profundo multiagente."
            ),
            (
                "El estudio no plantea contrastación de hipótesis. Por tratarse de una tesis de especialización "
                "orientada a solución, su demostración se apoya en simulación reproducible, comparación contra "
                "línea base, estadística descriptiva, análisis multicriterio y trazabilidad de artefactos "
                "técnicos: checkpoints, resultados JSON, series CSV, trazas por agente, figuras y tablas."
            ),
        ],
    )

    doc.add_heading("Índice de Contenidos", level=2)
    _add_numbered(
        doc,
        [
            "CAPÍTULO I. Planteamiento del Problema.",
            "CAPÍTULO II. Objetivos.",
            "CAPÍTULO III. Marco Teórico.",
            "CAPÍTULO IV. Metodología.",
            "II. Propuesta y Desarrollo de la Solución.",
            "Referencias Bibliográficas.",
            "Anexos: matriz de consistencia y operacionalización de KPIs.",
        ],
    )
    doc.add_page_break()


def _add_chapter_1(doc: Document) -> None:
    doc.add_heading("CAPÍTULO I. PLANTEAMIENTO DEL PROBLEMA", level=1)

    doc.add_heading("1.1 Diagnóstico", level=2)
    _add_paragraphs(
        doc,
        [
            (
                "El sector de edificios constituye un frente crítico para la transición energética. UNEP y "
                "GlobalABC reportan que, en 2022, los edificios y la construcción representaron una fracción "
                "significativa de la demanda energética y de las emisiones globales de CO2 relacionadas con "
                "energía y procesos. La IEA también identifica a la operación de edificios como un consumidor "
                "relevante de energía final y como fuente importante de emisiones energéticas, especialmente "
                "por el uso de electricidad y calor."
            ),
            (
                "La penetración de recursos energéticos distribuidos, baterías, generación fotovoltaica, "
                "cargas flexibles, bombas de calor y vehículos eléctricos crea oportunidades de flexibilidad, "
                "pero introduce decisiones operativas acopladas. Una acción local que reduce el costo de un "
                "edificio puede desplazar picos al distrito; una acción que carga un EV puede coincidir con "
                "horas de alta intensidad de carbono; y una política que reduce emisiones puede incrementar "
                "costos si ignora tarifas dinámicas."
            ),
            (
                "CityLearn v2 ofrece un entorno estandarizado para comunidades grid-interactive, con evaluación "
                "de flexibilidad, resiliencia, confort, emisiones y costos. Sin embargo, el proyecto actual "
                "requiere extender su soporte MARL hacia MADRL con Dec-POMDP, CTDE y backends oficiales, sin "
                "reimplementar algoritmos de forma local ni perder compatibilidad con los KPIs de CityLearn v2."
            ),
        ],
    )

    _add_table(
        doc,
        ["Síntoma", "Causa probable", "Consecuencia técnica"],
        [
            ("Picos y rampas distritales", "Control local no coordinado de baterías, PV y EVs", "Mayor estrés de red y peor flexibilidad."),
            ("Importación en horas con alta intensidad de carbono", "Políticas que optimizan energía o costo sin señal ambiental", "Aumento de huella ambiental."),
            ("Costos altos bajo tarifas dinámicas", "Carga flexible desplazada a periodos caros", "Menor eficiencia económica."),
            ("Dificultad de comparación entre algoritmos", "Entrenamientos con backends, KPIs u horizontes distintos", "Resultados no reproducibles o no defendibles."),
            ("Escalabilidad limitada", "Control centralizado puro o agentes independientes sin coordinación", "Pérdida de privacidad, estabilidad o asignación de crédito."),
        ],
    )

    doc.add_heading("1.2 Identificación y Descripción del Problema de Estudio", level=2)
    _add_paragraphs(
        doc,
        [
            (
                "El problema de estudio es la falta de una capa experimental integrada, reproducible y "
                "científicamente trazable que permita entrenar y comparar algoritmos MADRL colaborativos "
                "sobre CityLearn v2 para optimizar simultáneamente flexibilidad energética, emisiones de CO2 "
                "y costos energéticos en una comunidad de 17 edificios con EVs."
            ),
            (
                "La situación problemática contiene variables técnicas y metodológicas. En lo técnico, cada "
                "edificio observa parcialmente el sistema y ejecuta acciones locales sobre almacenamiento, EVs "
                "y cargas; en lo metodológico, la comparación requiere mantener constante el dataset, horizonte, "
                "semillas, baseline y KPIs. Por ello, el problema no se resuelve con un único agente ni con "
                "algoritmos prototipo, sino con Dec-POMDP, CTDE y backends oficiales acoplados a CityLearn v2."
            ),
        ],
    )

    doc.add_heading("1.2.1 Antecedentes Bibliográficos", level=3)
    _add_paragraphs(
        doc,
        [
            (
                "CityLearn fue propuesto para estandarizar investigación en aprendizaje por refuerzo aplicado "
                "a respuesta de demanda y gestión energética urbana. CityLearn v2 amplía ese enfoque hacia "
                "comunidades grid-interactive, DERs, V2G, resiliencia, confort y control carbon-aware."
            ),
            (
                "En coordinación multiagente, el Dec-POMDP formaliza toma de decisiones descentralizada bajo "
                "observabilidad parcial; CTDE permite entrenar con información global y ejecutar con políticas "
                "locales. MADDPG y QMIX son antecedentes clave de esta lógica, mientras que HAPPO, MASAC, "
                "MATD3 y MAAC representan familias MADRL avanzadas relevantes para el proyecto."
            ),
            (
                "Los antecedentes aplicados incluyen tesis sobre MARL para HVAC, comunidades energéticas con "
                "EV/V2G, respuesta de demanda con DERs y pricing dinámico multi-energía. Estos trabajos "
                "muestran la pertinencia de MADRL cuando existen recursos distribuidos, privacidad local, "
                "objetivos conflictivos y necesidad de evaluación cuantitativa."
            ),
        ],
    )

    _add_table(
        doc,
        ["Línea", "Antecedente", "Aporte al proyecto"],
        [
            ("CityLearn", "Vázquez-Canteli et al. (2020); Nweye et al. (2025)", "Entorno, datasets, evaluación y KPIs de comunidades energéticas."),
            ("Dec-POMDP/CTDE", "Bernstein et al. (2002); Lowe et al. (2017); Rashid et al. (2018)", "Formulación descentralizada y entrenamiento centralizado."),
            ("HAPPO/HARL", "Zhong et al. (2024)", "Políticas heterogéneas con mejora conjunta y crítico centralizado."),
            ("MASAC/mSAC", "Pu et al. (2021)", "Soft actor-critic multiagente cooperativo con descomposición."),
            ("MATD3", "Ackermann et al. (2019)", "Críticos centralizados dobles para reducir sobreestimación."),
            ("MAAC", "Iqbal y Sha (2019)", "Crítico de atención para seleccionar información relevante entre agentes."),
            ("Tesis energéticas", "González Rotger (2021); Fonseca (2023); Dong (2022); Almannouny (2025)", "Aplicaciones de RL/MARL a HVAC, EVs, DERs y respuesta de demanda."),
        ],
    )

    doc.add_heading("1.2.2 Formulación del Problema", level=3)
    doc.add_heading("1.2.2.1 Formulación del Problema General", level=4)
    doc.add_paragraph(
        "¿Cómo diseñar, implementar y evaluar una capa CityLearn v3 MADRL, basada en Dec-POMDP, CTDE y "
        "backends oficiales, que permita optimizar de forma coordinada la flexibilidad energética, las "
        "emisiones de CO2 y los costos energéticos de una comunidad inteligente de 17 edificios con EVs, "
        "manteniendo los KPIs y la línea base de CityLearn v2?"
    )

    doc.add_heading("1.2.2.2 Formulación de los Problemas Específicos", level=4)
    _add_bullets(
        doc,
        [
            "PE1: ¿Cómo aumentar la capacidad de desplazar cargas y aprovechar baterías, EVs/V2G, PV y autoconsumo usando MADRL colaborativo?",
            "PE2: ¿Cómo reducir emisiones de CO2 distritales minimizando importaciones en horas de alta intensidad de carbono?",
            "PE3: ¿Cómo optimizar costos energéticos reduciendo picos de demanda y respondiendo a tarifas dinámicas?",
            "PE4: ¿Cómo garantizar que HAPPO, MASAC, MATD3 y MAAC sean comparables usando el mismo dataset, baseline, KPIs y artefactos reproducibles?",
        ],
    )

    doc.add_heading("1.2.3 Justificación y Alcances", level=3)
    doc.add_heading("1.2.3.1 Justificación", level=4)
    _add_paragraphs(
        doc,
        [
            (
                "La justificación científica radica en articular CityLearn v2 con MADRL profundo sin duplicar "
                "implementaciones ni alterar los KPIs oficiales. Esto habilita comparaciones reproducibles y "
                "aporta una base extensible para futuros estudios en comunidades energéticas inteligentes."
            ),
            (
                "La justificación social y ambiental se relaciona con la necesidad de reducir emisiones y "
                "mejorar el uso de recursos energéticos distribuidos. La justificación económica está asociada "
                "a la reducción de costos bajo tarifas dinámicas y al potencial de disminuir picos de demanda."
            ),
            (
                "La justificación metodológica reside en que el planteamiento no contrasta hipótesis, sino que "
                "desarrolla y evalúa una solución usando simulación, estadística descriptiva, KPIs, análisis de "
                "deltas contra baseline y comparación multicriterio."
            ),
        ],
    )

    doc.add_heading("1.2.3.2 Alcances", level=4)
    _add_bullets(
        doc,
        [
            "El entorno base es CityLearn v2; la propuesta agrega una capa CityLearn v3 MADRL.",
            "El caso de tesis usa el dataset `citylearn_challenge_2022_phase_all_plus_evs`, 17 edificios, EVs y horizonte horario completo.",
            "Los algoritmos incluidos son HAPPO, MASAC/mSAC, MATD3 y MAAC con backends fuente-oficiales o source-backed.",
            "La evaluación se limita a simulación computacional; no incluye despliegue físico en red eléctrica real.",
            "La comparación usa línea base CityLearn v2, KPIs por eje, estadísticas descriptivas y artefactos reproducibles.",
        ],
    )
    doc.add_page_break()


def _kpi_rows(manifest: Mapping[str, object]) -> list[tuple[str, str, str]]:
    rows = []
    for axis, payload in manifest["axes"].items():
        rows.append((axis, payload["name"], ", ".join(payload["kpis"])))
    return rows


def _add_chapter_2(doc: Document, manifest: Mapping[str, object]) -> None:
    doc.add_heading("CAPÍTULO II. OBJETIVOS", level=1)
    doc.add_heading("2.1 Objetivo General", level=2)
    doc.add_paragraph(
        "Diseñar, implementar y evaluar una capa CityLearn v3 MADRL sobre CityLearn v2 para la gestión "
        "coordinada de flexibilidad energética, emisiones de CO2 y eficiencia económica en comunidades "
        "inteligentes, mediante Dec-POMDP, CTDE, backends oficiales y comparación contra línea base con KPIs "
        "CityLearn v2."
    )

    doc.add_heading("2.2 Objetivos Específicos", level=2)
    _add_bullets(
        doc,
        [
            "OE1: Aumentar la flexibilidad energética del distrito mediante desplazamiento de carga, baterías, EVs/V2G, PV, autoconsumo e intercambio comunitario.",
            "OE2: Reducir las emisiones de CO2 del distrito evitando importaciones en horas de alta intensidad de carbono.",
            "OE3: Optimizar costos energéticos reduciendo picos de demanda y aprovechando tarifas dinámicas.",
            "OE4 metodológico: comparar HAPPO, MASAC, MATD3 y MAAC bajo el mismo dataset, baseline, KPIs, horizonte y protocolo de artefactos.",
        ],
    )

    doc.add_heading("KPIs por eje", level=2)
    _add_table(doc, ["Eje", "Nombre", "KPIs"], _kpi_rows(manifest))
    doc.add_page_break()


def _add_chapter_3(doc: Document, lock: Mapping[str, object]) -> None:
    doc.add_heading("CAPÍTULO III. MARCO TEÓRICO", level=1)
    doc.add_heading("3.1 Bases Teóricas", level=2)

    theoretical_blocks = [
        (
            "Comunidades inteligentes e interacción con la red",
            "Una comunidad inteligente integra edificios, almacenamiento, generación distribuida, EVs, precios y señales de carbono. Su control requiere coordinar decisiones locales para mejorar resultados distritales.",
        ),
        (
            "CityLearn v2 como entorno de simulación",
            "CityLearn v2 proporciona datasets, dinámica física, API de simulación, línea base y KPIs para evaluar comunidades grid-interactive. En este proyecto no se reemplaza; se usa como entorno oficial de entrenamiento y evaluación.",
        ),
        (
            "Dec-POMDP",
            "El problema se modela como un Dec-POMDP: cada edificio es un agente, observa solo información local, actúa sobre sus recursos disponibles y comparte un objetivo colaborativo a nivel de distrito.",
        ),
        (
            "CTDE",
            "Centralized Training, Decentralized Execution permite entrenar críticos o estados globales usando información del distrito, mientras que en ejecución cada política decide con observaciones locales.",
        ),
        (
            "Multiobjetivo",
            "Flexibilidad, emisiones y costos no siempre mejoran simultáneamente. Por ello se reportan KPIs por eje y se habilita análisis multicriterio posterior, por ejemplo TOPSIS o ranking ponderado.",
        ),
    ]

    for title, text in theoretical_blocks:
        doc.add_heading(title, level=3)
        doc.add_paragraph(text)

    doc.add_heading("Backends MADRL oficiales", level=3)
    backends = lock.get("backends", {})
    rows = []
    for name in ["HAPPO", "MASAC", "MATD3", "MATD3_PYTORCH", "MAAC", "MARLlib"]:
        data = backends.get(name, {})
        rows.append(
            (
                name,
                data.get("path", ""),
                data.get("repository", ""),
                data.get("commit", data.get("branch", "")),
            )
        )
    _add_table(doc, ["Backend", "Ruta local", "Repositorio", "Commit/branch"], rows)

    doc.add_heading("3.2 Definición de Términos", level=2)
    _add_table(
        doc,
        ["Término", "Definición operativa"],
        [
            ("MADRL", "Aprendizaje por refuerzo profundo multiagente para políticas coordinadas."),
            ("Dec-POMDP", "Modelo de decisión secuencial descentralizado con observabilidad parcial."),
            ("CTDE", "Entrenamiento centralizado y ejecución descentralizada."),
            ("Baseline", "Referencia de comparación provista por CityLearn v2/evaluate_v2."),
            ("KPI", "Indicador cuantitativo usado para medir un eje de investigación."),
            ("OE1", "Eje de flexibilidad energética."),
            ("OE2", "Eje de emisiones de CO2."),
            ("OE3", "Eje de costos energéticos."),
            ("V2G", "Vehicle-to-grid: exportación de energía desde EV hacia edificio/red."),
            ("TOPSIS", "Método multicriterio para ordenar alternativas por cercanía a solución ideal."),
            ("Checkpoint", "Archivo de modelo entrenado que permite trazabilidad y reproducción."),
        ],
    )
    doc.add_page_break()


def _add_chapter_4(
    doc: Document,
    schema: Mapping[str, object],
    manifest: Mapping[str, object],
    status: Mapping[str, object],
) -> None:
    doc.add_heading("CAPÍTULO IV. METODOLOGÍA", level=1)

    doc.add_heading("4.1 Tipo de Investigación", level=2)
    doc.add_paragraph(
        "La investigación es aplicada, cuantitativa y tecnológica-computacional. Está orientada a desarrollar "
        "una solución experimental reproducible para un problema real de gestión energética urbana."
    )

    doc.add_heading("4.2 Nivel de Investigación", level=2)
    doc.add_paragraph(
        "El nivel es descriptivo, comparativo y evaluativo. Describe el comportamiento de cada algoritmo, "
        "compara resultados contra baseline y evalúa desempeño por KPIs de los tres ejes."
    )

    doc.add_heading("4.3 Métodos de Trabajo", level=2)
    _add_bullets(
        doc,
        [
            "Modelamiento del entorno como Dec-POMDP colaborativo.",
            "Entrenamiento CTDE con críticos/estados globales según la naturaleza de cada backend.",
            "Ejecución descentralizada con actores por edificio.",
            "Simulación horaria en CityLearn v2 usando el mismo schema para todos los algoritmos.",
            "Evaluación con KPIs CityLearn v2 y KPIs derivados documentados.",
            "Comparación contra baseline usando estadística descriptiva y análisis multicriterio.",
        ],
    )

    doc.add_heading("4.4 Población y Muestra", level=2)
    _add_paragraphs(
        doc,
        [
            (
                "La población conceptual está conformada por comunidades de edificios interactivos con la red "
                "eléctrica que integran recursos flexibles, almacenamiento, PV, EVs y señales de precio/carbono."
            ),
            (
                f"La muestra computacional del caso de tesis es el dataset `{schema['dataset']}`, con "
                f"{schema['building_count']} edificios, {schema['time_steps']} pasos horarios, "
                f"{schema['observation_count']} observaciones activas y {schema['action_count']} acciones activas."
            ),
        ],
    )

    doc.add_heading("4.5 Tipo de Diseño", level=2)
    doc.add_paragraph(
        "El diseño es cuasi-experimental computacional, longitudinal y comparativo. Las unidades de análisis "
        "son las corridas algoritmo-escenario-semilla. El diseño completo previsto por configuración contempla "
        "4 algoritmos, 3 escenarios y 10 semillas; las corridas oficiales pueden ejecutarse progresivamente "
        "según disponibilidad computacional."
    )
    _add_table(
        doc,
        ["Escenario", "Prioridad", "Propósito"],
        [
            ("E1", "OE1", "Entrenamiento y evaluación orientados a flexibilidad energética."),
            ("E2", "OE2", "Entrenamiento y evaluación orientados a emisiones de CO2."),
            ("E3", "OE3", "Entrenamiento y evaluación orientados a costos energéticos."),
        ],
    )

    doc.add_heading("4.6 Técnicas e Instrumentos de Recolección de Datos", level=2)
    _add_table(
        doc,
        ["Instrumento", "Dato recolectado"],
        [
            ("CityLearn v2 schema", "Edificios, observaciones, acciones, precios, carbono, PV, baterías y EVs."),
            ("Scripts de entrenamiento", "Configuración, episodios, semillas, hiperparámetros y backend."),
            ("results.json", "Reporte técnico, KPIs, baseline, hiperparámetros y resumen de ejecución."),
            ("timeseries.csv", "Series temporales distritales por paso/episodio."),
            ("trace.csv", "Trazas por agente, acción, recompensa y paso."),
            ("checkpoint_manifest.json", "Inventario de checkpoints generados."),
            ("figures/", "Gráficas de recompensa, convergencia, eficiencia, baseline y KPIs por eje."),
        ],
    )

    doc.add_heading("4.7 Técnicas e Instrumentos de Análisis y Procesamiento de Datos", level=2)
    _add_bullets(
        doc,
        [
            "Estadística descriptiva: media, suma, mínimo, máximo, desviación estándar y evolución temporal.",
            "Comparación contra baseline: ratios, diferencias absolutas, deltas y porcentaje de mejora.",
            "Análisis por eje: OE1, OE2 y OE3 se reportan por separado para evitar mezclar objetivos conflictivos.",
            "Análisis de aprendizaje: rewards, returns, convergencia, exploración y estabilidad.",
            "Análisis multicriterio: TOPSIS o ranking ponderado con pesos explícitos después de normalizar KPIs.",
            "Trazabilidad: verificación de scripts, commits de backends, checkpoints, JSON, CSV y figuras.",
        ],
    )

    doc.add_heading("4.8 Etapas de Intervención del Estudio", level=2)
    _add_table(
        doc,
        ["Etapa", "Actividad", "Producto verificable"],
        [
            ("1", "Diagnóstico y revisión bibliográfica", "Capítulo I, marco teórico y referencias APA 7."),
            ("2", "Adecuación CityLearn v2 -> CityLearn v3", "Módulos `citylearn.v3`, Dec-POMDP y MARLlib env."),
            ("3", "Integración de backends oficiales", "HARL, MARL, MATD3/off-policy, MAAC y lockfile."),
            ("4", "Definición de ejes y KPIs", "Manifest OE1/OE2/OE3 y `objective_kpis.csv`."),
            ("5", "Entrenamiento MADRL", "Checkpoints, `results.json`, `timeseries.csv`, `trace.csv`."),
            ("6", "Generación de figuras y tablas", "`figures/`, PNG, CSV y Markdown por corrida."),
            ("7", "Comparación contra baseline", "Tablas por eje, deltas, mejoras y ranking multicriterio."),
            ("8", "Redacción final", "Plan de tesis, anexos, matriz de consistencia y repositorio reproducible."),
        ],
    )

    _add_modification_registry(doc)

    doc.add_heading("Estado experimental actual", level=2)
    jobs = status.get("jobs", [])
    rows = [
        (
            job.get("name", ""),
            job.get("started_at", ""),
            job.get("completed_at", "en ejecución"),
            job.get("exit_code", "en ejecución"),
            job.get("output_dir", ""),
        )
        for job in jobs
    ]
    if rows:
        _add_table(doc, ["MADRL", "Inicio", "Fin", "Exit code", "Salida"], rows)
    else:
        doc.add_paragraph("No se encontró manifest de entrenamiento oficial en la ruta esperada.")

    doc.add_page_break()


def _add_solution_section(doc: Document) -> None:
    doc.add_heading("II. PROPUESTA Y DESARROLLO DE LA SOLUCIÓN", level=1)
    _add_paragraphs(
        doc,
        [
            (
                "La solución propuesta es una arquitectura CityLearn v3 MADRL que conserva CityLearn v2 como "
                "núcleo de simulación y evaluación, y agrega una capa multiagente profunda basada en backends "
                "oficiales. El objetivo no es reemplazar CityLearn v2, sino extenderlo hacia un entorno de "
                "entrenamiento MADRL colaborativo y multiobjetivo."
            ),
            (
                "La viabilidad se demuestra mediante ejecución local reproducible, disponibilidad de backends "
                "clonados, entrenamiento con CUDA, generación de checkpoints y evaluación con KPIs. La "
                "sostenibilidad científica se apoya en trazabilidad de fuentes, scripts reutilizables y "
                "compatibilidad con otros datasets CityLearn v2."
            ),
        ],
    )

    _add_table(
        doc,
        ["Componente", "Descripción"],
        [
            ("Datos", "Dataset CityLearn v2 `citylearn_challenge_2022_phase_all_plus_evs`."),
            ("Entorno", "Wrapper Dec-POMDP con 17 agentes-edificio y estado global CTDE."),
            ("Algoritmos", "HAPPO, MASAC, MATD3 y MAAC con backends oficiales/source-backed."),
            ("Evaluación", "KPIs CityLearn v2 y derivados por OE1, OE2 y OE3."),
            ("Artefactos", "Checkpoints, JSON, CSV, logs, figuras, tablas y manifiestos."),
            ("Comparación", "Baseline CityLearn v2, deltas, ratios, mejora por KPI y análisis multicriterio."),
        ],
    )

    doc.add_heading("Demostración de viabilidad y sostenibilidad", level=2)
    _add_bullets(
        doc,
        [
            "Viabilidad técnica: el entorno Python 3.9 usa PyTorch CUDA y scripts por algoritmo.",
            "Viabilidad metodológica: todos los algoritmos comparten dataset, horizonte, semilla, baseline y KPIs.",
            "Viabilidad científica: los algoritmos provienen de papers y repositorios oficiales o source-backed.",
            "Sostenibilidad: la capa v3 es reutilizable para otros datasets CityLearn v2.",
            "Evaluación cuantitativa: se aplican estadísticas descriptivas, deltas, ratios y ranking multicriterio.",
            "Evaluación económica de la solución: en la etapa final se puede estimar beneficio-costo computacional, ahorro energético, reducción de emisiones y reducción de costo contra baseline.",
        ],
    )
    doc.add_page_break()


def _add_annexes(doc: Document, manifest: Mapping[str, object]) -> None:
    doc.add_heading("ANEXOS", level=1)
    doc.add_heading("Anexo I. Matriz de Consistencia", level=2)
    _add_table(
        doc,
        ["Problema específico", "Objetivo específico", "Eje", "Indicadores"],
        [
            ("PE1", "Aumentar flexibilidad energética", "OE1", "KPIs de importación, exportación, pico, rampa, load factor, PV, baterías y EVs."),
            ("PE2", "Reducir emisiones de CO2", "OE2", "KPIs de emisiones control, baseline, delta y promedios diarios."),
            ("PE3", "Optimizar costos energéticos", "OE3", "KPIs de costo, picos/rampas de costo y desviación respecto a tarifa."),
            ("PE4", "Comparar algoritmos MADRL", "Transversal", "Rewards, returns, checkpoints, trazas, figuras, baseline y ranking."),
        ],
    )

    doc.add_heading("Anexo II. Operacionalización de KPIs", level=2)
    rows = []
    axis_kpis = manifest["axis_kpis"]
    for axis, payload in manifest["axes"].items():
        for kpi in payload["kpis"]:
            trace = axis_kpis.get(kpi, {})
            rows.append(
                (
                    axis,
                    kpi,
                    trace.get("source", ""),
                    "Menor" if trace.get("lower_is_better", True) else "Mayor",
                    trace.get("note", ""),
                )
            )
    _add_table(doc, ["Eje", "KPI", "Fuente", "Mejor dirección", "Nota"], rows)
    doc.add_page_break()


def _add_references(doc: Document) -> None:
    doc.add_heading("REFERENCIAS BIBLIOGRÁFICAS", level=1)
    references = [
        "Ackermann, J., Gabler, V., Osa, T., & Sugiyama, M. (2019). Reducing overestimation bias in multi-agent domains using double centralized critics. arXiv. https://doi.org/10.48550/arXiv.1910.01465",
        "Almannouny, G. A. (2025). Intelligent dynamic pricing and integrated demand response for multi-energy systems using deep reinforcement learning [Doctoral dissertation, University of Glasgow]. Enlighten Theses. https://doi.org/10.5525/gla.thesis.85367",
        "Bernstein, D. S., Givan, R., Immerman, N., & Zilberstein, S. (2002). The complexity of decentralized control of Markov decision processes. Mathematics of Operations Research, 27(4), 819-840. https://doi.org/10.1287/moor.27.4.819.297",
        "Dong, J. (2022). Peak load ensemble prediction and multi-agent reinforcement learning for DER demand response management in smart grids [Master's thesis, Lakehead University]. Knowledge Commons. https://knowledgecommons.lakeheadu.ca/handle/2453/4944",
        "Fonseca, T. C. C. (2023). A multi-agent reinforcement learning approach to integrate flexible assets into energy communities [Master's thesis, Instituto Superior de Engenharia do Porto]. Repositório Científico do Instituto Politécnico do Porto. http://hdl.handle.net/10400.22/24068",
        "González Rotger, C. (2021). Multi-agent reinforcement learning applied to heating, ventilation, and air conditioning in a building energy management system [Master's thesis, Universitat de les Illes Balears]. http://hdl.handle.net/11201/158415",
        "Hu, S., Zhong, Y., Gao, M., Wang, W., Dong, H., Liang, X., Li, Z., Chang, X., & Yang, Y. (2023). MARLlib: A scalable and efficient multi-agent reinforcement learning library. arXiv. https://arxiv.org/abs/2210.13708",
        "International Energy Agency. (2023). Buildings. https://www.iea.org/energy-system/buildings",
        "Iqbal, S., & Sha, F. (2019). Actor-attention-critic for multi-agent reinforcement learning. arXiv. https://doi.org/10.48550/arXiv.1810.02912",
        "Lowe, R., Wu, Y., Tamar, A., Harb, J., Abbeel, P., & Mordatch, I. (2017). Multi-agent actor-critic for mixed cooperative-competitive environments. Advances in Neural Information Processing Systems, 30. https://papers.nips.cc/paper/7217-multi-agent-actor-critic-for-mixed-cooperative-competitive-environments",
        "Nweye, K., Kaspar, K., Buscemi, G., Fonseca, T., Pinto, G., Ghose, D., Duddukuru, S., Pratapa, P., Li, H., Mohammadi, J., Lino Ferreira, L., Hong, T., Ouf, M., Capozzoli, A., & Nagy, Z. (2025). CityLearn v2: Energy-flexible, resilient, occupant-centric, and carbon-aware management of grid-interactive communities. Journal of Building Performance Simulation, 18(1), 17-38. https://doi.org/10.1080/19401493.2024.2418813",
        "Pu, Y., Wang, S., Yang, R., Yao, X., & Li, B. (2021). Decomposed soft actor-critic method for cooperative multi-agent reinforcement learning. arXiv. https://doi.org/10.48550/arXiv.2104.06655",
        "Rashid, T., Samvelyan, M., Schroeder de Witt, C., Farquhar, G., Foerster, J., & Whiteson, S. (2018). QMIX: Monotonic value function factorisation for deep multi-agent reinforcement learning. arXiv. https://doi.org/10.48550/arXiv.1803.11485",
        "United Nations Environment Programme & Global Alliance for Buildings and Construction. (2024). Global status report for buildings and construction. https://www.unep.org/resources/report/global-status-report-buildings-and-construction",
        "Vázquez-Canteli, J. R., Dey, S., Henze, G., & Nagy, Z. (2020). CityLearn: Standardizing research in multi-agent reinforcement learning for demand response and urban energy management. arXiv. https://doi.org/10.48550/arXiv.2012.10504",
        "Zhong, Y., Kuba, J. G., Feng, X., Hu, S., Ji, J., & Yang, Y. (2024). Heterogeneous-agent reinforcement learning. Journal of Machine Learning Research, 25(32), 1-67. https://jmlr.org/papers/v25/23-0488.html",
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)


def build_docx() -> Path:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    schema = _schema_summary()
    status = _read_json(STATUS_PATH)
    lock = _read_json(LOCK_PATH)
    manifest = objective_manifest()

    doc = Document()
    _set_styles(doc)
    _setup_sections(doc)
    _add_header_footer(doc)

    _add_cover(doc, schema, status)
    _add_preliminary(doc, schema, status)
    doc.add_heading("II. CONTENIDOS", level=1)
    _add_chapter_1(doc)
    _add_chapter_2(doc, manifest)
    _add_chapter_3(doc, lock)
    _add_chapter_4(doc, schema, manifest, status)
    _add_solution_section(doc)
    _add_references(doc)
    doc.add_page_break()
    _add_annexes(doc, manifest)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_docx()
    print(path)
