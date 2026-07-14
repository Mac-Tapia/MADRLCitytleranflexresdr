#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Audit informe-final chapter/subsection coverage for docs/informedetesis.txt."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
DOCX = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx"
OUT = REPO / "outputs" / "_audit_tesis_structure_checklist.json"


def _heading_level(style_name: str) -> int | None:
    m = re.match(r"Heading\s+(\d+)", style_name or "", re.I)
    return int(m.group(1)) if m else None


def extract(doc: Document) -> dict:
    headings: list[dict] = []
    chapter_words: dict[str, int] = {}
    current = "front"
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        n = len(text.split())
        chapter_words[current] = chapter_words.get(current, 0) + n
        level = _heading_level(para.style.name if para.style else "")
        if level is None:
            continue
        headings.append({"level": level, "text": text})
        m = re.match(r"Capitulo\s+(\d+)", text, re.I)
        if m:
            current = f"cap{m.group(1)}"
        elif text.lower().startswith("referencias"):
            current = "refs"
        elif text.lower().startswith("anexo"):
            current = "anexo"
    return {"headings": headings, "chapter_words": chapter_words}


def section_range(headings: list[dict], cap: int) -> list[dict]:
    out: list[dict] = []
    capture = False
    for h in headings:
        if re.match(rf"Capitulo\s+{cap}\b", h["text"], re.I):
            capture = True
            out.append(h)
            continue
        if not capture:
            continue
        if re.match(r"Capitulo\s+\d+", h["text"], re.I):
            break
        if h["text"].lower().startswith("referencias") or h["text"].lower().startswith("anexo"):
            break
        out.append(h)
    return out


def find_like(texts: list[str], *patterns: str) -> bool:
    blob = " | ".join(texts).lower()
    for pat in patterns:
        if re.search(pat, blob, re.I):
            return True
    return False


def body_has(full: str, *patterns: str) -> bool:
    for pat in patterns:
        if re.search(pat, full, re.I):
            return True
    return False


def status(ok: bool, incomplete: bool = False) -> str:
    if ok and not incomplete:
        return "OK"
    if ok and incomplete:
        return "incomplete"
    return "missing"


def main() -> int:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else DOCX
    doc = Document(str(path))
    data = extract(doc)
    headings = data["headings"]
    full = "\n".join(p.text or "" for p in doc.paragraphs)
    n_tables = len(doc.tables)
    n_images = sum(1 for r in doc.part.rels.values() if "image" in r.reltype)

    caps = {i: section_range(headings, i) for i in range(1, 7)}
    h_texts = {i: [h["text"] for h in caps[i]] for i in range(1, 7)}

    checklist = {
        "Capítulo 1. Introducción": {
            "Problema de investigación": status(find_like(h_texts[1], r"problema")),
            "Objetivos": status(find_like(h_texts[1], r"objetivo")),
            "Hipótesis": status(find_like(h_texts[1], r"hip[oó]tes")),
            "Justificación": status(find_like(h_texts[1], r"justificaci")),
            "Alcances y limitaciones": status(find_like(h_texts[1], r"alcance")),
        },
        "Capítulo 2. Marco teórico": {
            "Estado del arte actualizado": status(
                find_like(h_texts[2], r"estado del arte")
                or body_has(full, r"2\.1.*estado del arte|Estado del arte actualizado"),
                incomplete=not find_like(h_texts[2], r"estado del arte"),
            ),
            "Bases teóricas": status(find_like(h_texts[2], r"bases? te")),
            "Trabajos relacionados": status(
                find_like(h_texts[2], r"trabajos relacion|antecedentes"),
                incomplete=not find_like(h_texts[2], r"trabajos relacion"),
            ),
        },
        "Capítulo 3. Metodología": {
            "Tipo de investigación": status(find_like(h_texts[3], r"tipo de investig")),
            "Diseño metodológico": status(find_like(h_texts[3], r"dise")),
            "Datos utilizados": status(find_like(h_texts[3], r"datos|dataset")),
            "Variables": status(find_like(h_texts[3], r"variable")),
            "Técnicas": status(find_like(h_texts[3], r"t[eé]cnica")),
            "Herramientas": status(
                find_like(h_texts[3], r"herramient|instrumento"),
                incomplete=not find_like(h_texts[3], r"herramient"),
            ),
            "Procedimiento experimental": status(find_like(h_texts[3], r"procedimiento")),
        },
        "Capítulo 4. Desarrollo de la propuesta": {
            "Desarrollo del sistema": status(find_like(h_texts[4], r"desarrollo")),
            "Arquitectura": status(find_like(h_texts[4], r"arquitectura")),
            "Modelo de IA": status(find_like(h_texts[4], r"modelo|Dec-POMDP|CTDE")),
            "Algoritmos": status(find_like(h_texts[4], r"algoritmo")),
            "Diseño experimental": status(find_like(h_texts[4], r"dise")),
            "Implementación": status(find_like(h_texts[4], r"implementaci")),
        },
        "Capítulo 5. Resultados": {
            "Experimentos realizados": status(
                find_like(h_texts[5], r"cobertura experimental|experimento|marco de contrast")
                or body_has(full, r"experimentos realizados|12 corridas|cobertura experimental")
            ),
            "Métricas utilizadas": status(
                find_like(h_texts[5], r"m[eé]trica|KPI|descriptivo")
                or body_has(full, r"m[eé]tricas utilizadas|KPI CityLearn")
            ),
            "Resultados obtenidos": status(
                find_like(h_texts[5], r"resultado|sintesis|respuesta a|OE\.")
            ),
            "Comparación baseline / trabajos relacionados": status(
                find_like(h_texts[5], r"comparaci|baseline|linea base|l[ií]nea base|antecedentes")
            ),
            "Tablas": status(n_tables >= 10),
            "Figuras": status(n_images >= 10),
            "Discusión de resultados": status(find_like(h_texts[5], r"discusi")),
        },
        "Capítulo 6. Conclusiones preliminares": {
            "Principales hallazgos": status(
                find_like(h_texts[6], r"conclusion|hallazgo")
                or body_has(full, r"principales hallazgos"),
                incomplete=not find_like(h_texts[6], r"hallazgo|principales"),
            ),
            "Limitaciones encontradas": status(find_like(h_texts[6], r"limitaci")),
            "Trabajo pendiente": status(
                find_like(h_texts[6], r"trabajo (pendiente|futuro)|pendiente"),
                incomplete=not find_like(h_texts[6], r"pendiente"),
            ),
            "Plan para culminar la tesis": status(
                find_like(h_texts[6], r"plan.*(culmin|cierre)|trabajo futuro|cronograma")
                or body_has(full, r"plan para culminar|plan de cierre|H1-H7|H1–H7"),
                incomplete=not find_like(h_texts[6], r"plan.*(culmin|cierre)"),
            ),
            "Referencias bibliográficas (APA)": status(
                body_has(full, r"Referencias bibliogr")
                or any(h["text"].lower().startswith("referencias") for h in headings)
            ),
        },
    }

    total_body = sum(v for k, v in data["chapter_words"].items() if k.startswith("cap"))
    cap5_share = (
        data["chapter_words"].get("cap5", 0) / total_body if total_body else 0.0
    )

    report = {
        "docx": str(path),
        "n_tables": n_tables,
        "n_images": n_images,
        "chapter_words": data["chapter_words"],
        "cap5_word_share_body": round(cap5_share, 3),
        "h1_h2": [h for h in headings if h["level"] <= 2],
        "checklist": checklist,
        "gaps": [
            f"{ch} / {sec}: {st}"
            for ch, secs in checklist.items()
            for sec, st in secs.items()
            if st != "OK"
        ],
    }
    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({"gaps": report["gaps"], "cap5_share": report["cap5_word_share_body"],
                      "words": data["chapter_words"], "tables": n_tables, "images": n_images},
                     indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
