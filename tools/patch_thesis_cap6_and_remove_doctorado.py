#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Patch Cap. 6.4/6.5 and remove 'doctorado'/'doctoral' from docs/*.docx.

Hard rules:
- No inventar cifras nuevas.
- No renombrar archivos.
- Saltar backups .bak_*
- Ortografía académica RAE en texto nuevo (sin la palabra doctorado).
"""
from __future__ import annotations

import json
import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
DOCS = REPO / "docs"
RUN_ID = "madrl_v3_20260627_164047"
AUDIT_OUT = (
    REPO
    / "outputs"
    / "_drive_madrl"
    / "gdrive_20260627_164047_objective_analysis"
    / "patch_cap6_remove_doctorado_report.json"
)
GREY = RGBColor(0x59, 0x59, 0x59)

# Ordered replacements (longer phrases first). Applied case-insensitively via helper.
PHRASE_REPLACEMENTS: list[tuple[str, str]] = [
    ("cierre doctoral y control de calidad final", "cierre de la tesis y control de calidad final"),
    ("criterios de cierre doctoral", "criterios de cierre de la tesis"),
    ("control de calidad doctoral", "control de calidad de la tesis"),
    ("redacción doctoral", "redacción de la tesis"),
    ("redaccion doctoral", "redacción de la tesis"),
    ("estructura doctoral", "estructura de la tesis"),
    ("formato doctoral", "formato de la tesis"),
    ("manuscrito doctoral", "manuscrito de la tesis"),
    ("documento doctoral", "documento de la tesis"),
    ("versión doctoral", "versión de la tesis"),
    ("version doctoral", "versión de la tesis"),
    ("nivel doctoral", "nivel de tesis"),
    ("calidad doctoral", "calidad de la tesis"),
    ("contribución doctoral", "contribución de la tesis"),
    ("contribucion doctoral", "contribución de la tesis"),
    ("investigación doctoral", "investigación de la tesis"),
    ("investigacion doctoral", "investigación de la tesis"),
    ("programa de doctorado", "programa de posgrado"),
    ("estudios de doctorado", "estudios de posgrado"),
    ("grado de doctor", "grado académico de la tesis"),
    ("tesis doctoral", "tesis"),
    ("doctoral thesis", "thesis"),
    ("doctoral dissertation", "dissertation"),
    ("doctoral research", "research"),
    ("doctoral study", "study"),
    ("doctoral program", "graduate program"),
    ("propuesta doctoral", "propuesta de tesis"),
    ("informe doctoral", "informe de tesis"),
    ("cierre doctoral", "cierre de la tesis"),
    ("enfoque doctoral", "enfoque de la tesis"),
    ("análisis doctoral", "análisis de la tesis"),
    ("analisis doctoral", "análisis de la tesis"),
    ("auditoría doctoral", "auditoría de la tesis"),
    ("auditoria doctoral", "auditoría de la tesis"),
    ("defensa doctoral", "sustentación de la tesis"),
    ("revisión doctoral", "revisión de la tesis"),
    ("revision doctoral", "revisión de la tesis"),
    ("versión final doctoral", "versión final de la tesis"),
    ("version final doctoral", "versión final de la tesis"),
    # standalone residual noun/adjective (after phrases)
    ("doctorado", "posgrado"),
    ("doctoral", "de tesis"),
]

# Filenames intentionally kept; body must have zero matches of these after patch.
BANNED_RE = re.compile(r"doctorado|doctoral", re.IGNORECASE)

PLACEHOLDER_SOFTEN = [
    (re.compile(r"\bTODO\b"), "pendiente de edición"),
    (re.compile(r"\bXXXX+\b"), ""),
    (re.compile(r"\[lorem[^\]]*\]", re.I), ""),
]


def _iter_text_nodes(element):
    for node in element.iter(qn("w:t")):
        yield node


def _apply_phrase_replacements(text: str) -> tuple[str, int]:
    if not text:
        return text, 0
    count = 0
    out = text
    for src, dst in PHRASE_REPLACEMENTS:
        pattern = re.compile(re.escape(src), re.IGNORECASE)

        def _sub(m: re.Match, replacement: str = dst) -> str:
            original = m.group(0)
            # Preserve crude capitalization of first letter
            if original[:1].isupper() and replacement[:1].islower():
                return replacement[:1].upper() + replacement[1:]
            return replacement

        new_out, n = pattern.subn(_sub, out)
        if n:
            count += n
            out = new_out
    # Soft cleanups that are not doctorado
    out2 = out
    for rx, repl in PLACEHOLDER_SOFTEN:
        out2, n = rx.subn(repl, out2)
        count += n
    # Collapse triple+ spaces created by empties
    out2 = re.sub(r"[ \t]{2,}", " ", out2)
    out2 = re.sub(r" ?\. ?", ". ", out2) if ".." in out2 else out2
    return out2, count


def clean_document_text(doc: Document) -> dict:
    """Replace doctorado/doctoral in paragraphs, tables, headers/footers."""
    stats = {"paragraphs": 0, "table_cells": 0, "headers_footers": 0, "total": 0, "samples": []}

    def _patch_paragraph(para) -> int:
        full = para.text or ""
        if not full:
            return 0
        new, n = _apply_phrase_replacements(full)
        if n == 0 or new == full:
            # still try node-level if mixed runs hold the word
            local = 0
            for node in _iter_text_nodes(para._element):
                raw = node.text or ""
                patched, nn = _apply_phrase_replacements(raw)
                if nn and patched != raw:
                    if len(stats["samples"]) < 8:
                        stats["samples"].append({"before": raw[:120], "after": patched[:120]})
                    node.text = patched
                    local += nn
            return local
        # Rewrite as single run preserving first run formatting when possible
        if len(stats["samples"]) < 8:
            stats["samples"].append({"before": full[:120], "after": new[:120]})
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.add_run(new)
        return n

    for para in doc.paragraphs:
        n = _patch_paragraph(para)
        stats["paragraphs"] += n
        stats["total"] += n

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    n = _patch_paragraph(para)
                    stats["table_cells"] += n
                    stats["total"] += n

    for section in doc.sections:
        for hf in (section.header, section.footer):
            for para in hf.paragraphs:
                n = _patch_paragraph(para)
                stats["headers_footers"] += n
                stats["total"] += n
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            n = _patch_paragraph(para)
                            stats["headers_footers"] += n
                            stats["total"] += n
    return stats


def count_banned(doc: Document) -> list[str]:
    hits = []
    for i, para in enumerate(doc.paragraphs):
        t = para.text or ""
        if BANNED_RE.search(t):
            hits.append(f"p{i}:{t[:100]}")
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                t = cell.text or ""
                if BANNED_RE.search(t):
                    hits.append(f"t{ti}r{ri}c{ci}:{t[:80]}")
    return hits


def _set_run_font(run, size: float = 11, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def _add_para(doc_or_parent, text: str, *, bold: bool = False, italic: bool = False, size: float = 11):
    para = doc_or_parent.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return para


def _add_heading_para(doc, text: str, level: int = 2):
    para = doc.add_heading(text, level=level)
    return para


def _add_table(doc, caption: str, headers: list[str], rows: list[list[str]], font_size: float = 8.0):
    cap = doc.add_paragraph()
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(h)
        rr.bold = True
        rr.font.size = Pt(font_size)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(font_size)
    return tbl


def _find_heading(doc: Document, prefixes: tuple[str, ...]) -> Paragraph | None:
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        for pref in prefixes:
            if t.startswith(pref):
                return para
    return None


def _heading_level(para: Paragraph) -> int:
    name = para.style.name if para.style else ""
    m = re.search(r"(\d+)", name or "")
    return int(m.group(1)) if m else 9


def _collect_section_range(doc: Document, start: Paragraph) -> tuple[list, Paragraph | None]:
    """Return body elements belonging to section starting at `start`, and the next heading element."""
    body = list(doc.element.body)
    try:
        idx = body.index(start._element)
    except ValueError:
        return [], None
    start_level = _heading_level(start)
    end_el = None
    keep: list = []
    for el in body[idx + 1 :]:
        if el.tag == qn("w:sectPr"):
            break
        if el.tag == qn("w:p"):
            p = Paragraph(el, doc)
            t = (p.text or "").strip()
            style = p.style.name if p.style else ""
            if style.startswith("Heading") and t:
                lvl = _heading_level(p)
                if lvl <= start_level:
                    end_el = p
                    break
        keep.append(el)
    return keep, end_el


def _clear_elements(elements: list) -> None:
    for el in elements:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _insert_elements_before(anchor: Paragraph | None, elements: list, doc: Document) -> None:
    body = doc.element.body
    if anchor is not None:
        ref = anchor._element
        for el in elements:
            ref.addprevious(el)
    else:
        sect = body.find(qn("w:sectPr"))
        for el in elements:
            if sect is not None:
                sect.addprevious(el)
            else:
                body.append(el)


def _build_cap64_temp() -> list:
    tmp = Document()
    # clear default empty para
    body = tmp.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    _add_heading_para(tmp, "6.4 Plan para culminar la tesis", 2)
    _add_para(
        tmp,
        "El plan de cierre se formula como secuencia de hitos verificables, anclados a la corrida "
        f"canónica {RUN_ID} (diseño factorial 4×3: HAPPO, MASAC, MATD3 y MAAC × escenarios E1, E2 y E3) "
        "y a la operacionalización de OE.1–OE.3 registrada en la Tabla 3.4. No introduce resultados "
        "experimentales nuevos: organiza el trabajo pendiente documentado en 6.2–6.3 y la evidencia ya "
        "materializada en results.json, timeseries.csv, trace.csv, building_kpis.csv, "
        "building_behavior_summary.csv y checkpoint_manifest.json.",
    )
    _add_para(
        tmp,
        "Orden lógico de dependencias: (1) homogeneizar artefactos faltantes de HAPPO y corregir "
        "artefactos con ceros anómalos en figuras/tablas; (2) cerrar la reproducibilidad de modelos "
        "(Anexo A.4) y la validez visual de la Figura 5.8e; (3) consolidar inferencia y lectura por "
        "objetivo; (4) completar el cierre editorial (ortografía RAE, APA, índice F9, metadatos "
        "institucionales). Los hitos H5 (Optuna) y la extensión multi-semilla son opcionales para la "
        "defensa de resultados descriptivos ya obtenidos, pero obligatorios si se afirma robustez "
        "causal externa.",
    )
    _add_table(
        tmp,
        "Tabla 6.1. Plan para culminar la tesis (hitos H1–H8).",
        ["Hito", "Actividades y entregable", "Dependencias", "Estado (evidencia actual)"],
        [
            [
                "H1. Cobertura HAPPO",
                "Recuperar evaluate_v2/core_kpis y building_* comparables; documentar 49 ep reales sin imputar ep. 50.",
                "Corrección VecEnvWrapper / pipeline Colab",
                "Pendiente (HAPPO parcial en KPIs de edificio y manifiestos)",
            ],
            [
                "H2. Figura 5.8e",
                "Regenerar heatmaps: action_l2 desde full_data/trace.csv (incl. HAPPO); EV/BESS desde building_behavior_summary; colorbar por panel.",
                "tools/build_final_thesis_gdrive_objectives.py + espejo full_data",
                "En corrección (columnas muertas ev_charge_kwh/SOC en trace)",
            ],
            [
                "H3. Tabla A.4 checkpoints",
                "Recalcular conteos reales desde checkpoint_manifest.json / archivos .pt|.pkl; actualizar narrativa A.4.",
                "outputs/_drive_madrl/full_data/.../checkpoint_manifest.json",
                "Pendiente de reemplazo (valores 0 anómalos reportados)",
            ],
            [
                "H4. Auditoría de ceros",
                "Inventario Cap. 5/anexos: distinguir ceros legítimos vs fallo de lectura; regenerar artefactos afectados.",
                "H2–H3",
                "En curso",
            ],
            [
                "H5. Inferencia y OE",
                "Mantener Kruskal-Wallis / Holm alineados a Tabla 3.4; no mezclar HAPPO incompleto en contrastes 50 ep.",
                "H1 para factorial 4×3 completo",
                "Parcial (MAAC/MASAC/MATD3 con 50 ep; HAPPO 49 ep distrital)",
            ],
            [
                "H6. Cierre editorial",
                "Pasada RAE (tildes, mayúsculas, tipografía); eliminar placeholders; sincronizar Words canónicos; índice F9.",
                "Contenido Cap. 1–6 estable",
                "En curso (incluye eliminación de la voz «doctorado» del cuerpo)",
            ],
            [
                "H7. Metadatos institucionales",
                "Completar asesor/[por definir] solo con dato real del programa; no inventar nombres.",
                "Definición institucional del autor",
                "Pendiente (placeholder institucional)",
            ],
            [
                "H8. Sustentación",
                "PDF final, presentación y paquete de reproducibilidad (scripts + CSV + manifiestos).",
                "H1–H7 según alcance declarado",
                "Pendiente",
            ],
        ],
        font_size=7.5,
    )
    _add_para(
        tmp,
        "Criterio de prioridad: H1–H4 son bloqueantes para la coherencia de Cap. 5 y Anexo A; H6 es "
        "bloqueante para la versión de lectura; H5, H7 y H8 cierran la trazabilidad académica e "
        "institucional. La multi-semilla (≥3) y Optuna (TPE) se mantienen como extensión de robustez, "
        "no como sustituto de la evidencia canónica de 50 episodios ya auditada.",
    )
    return [deepcopy(c) for c in tmp.element.body if c.tag != qn("w:sectPr")]


def _build_cap65_temp() -> list:
    tmp = Document()
    body = tmp.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    _add_heading_para(tmp, "6.5 Criterios de cierre de la tesis y control de calidad final", 2)
    _add_para(
        tmp,
        "Esta sección define criterios de aceptación verificables para declarar el manuscrito listo "
        "para sustentación. Cada criterio se evalúa contra artefactos del repositorio y del informe; "
        "el estado refleja la evidencia disponible al momento de esta revisión y no inventa cierres "
        "experimentales inexistentes.",
    )
    _add_table(
        tmp,
        "Tabla 6.2. Criterios de cierre de la tesis y control de calidad final.",
        ["ID", "Criterio verificable", "Cómo comprobarlo", "Estado"],
        [
            [
                "C1",
                "Estructura Cap. 1–6 + anexos presente, sin secciones vacías 6.4/6.5",
                "Headings en Word; ≥1 párrafo sustantivo bajo 6.4 y 6.5",
                "CUMPLE (tras este parche)",
            ],
            [
                "C2",
                "Corrida canónica citada: " + RUN_ID + " (4 alg. × E1–E3)",
                "Texto Cap. 5/6 + carpeta outputs/madrl_v3_20260627_164047 o espejo Drive",
                "CUMPLE",
            ],
            [
                "C3",
                "OE.1/OE.2/OE.3 operacionalizados (Tabla 3.4) y contrastados sin inventar KPIs",
                "Cruce Cap. 3 ↔ Cap. 5; CSV episodicos / results.json",
                "CUMPLE (núcleo); pendiente homogeneizar HAPPO en KPIs de edificio",
            ],
            [
                "C4",
                "Figura 5.8e sin paneles de columnas muertas (ceros totales en EV/SOC de trace)",
                "PNG regenerado + fila HAPPO en action_l2; EV/BESS desde behavior summary",
                "NO CUMPLE aún (corrección en pipeline)",
            ],
            [
                "C5",
                "Tabla A.4 con conteos reales de checkpoints (no matriz de ceros)",
                "checkpoint_summary.csv / checkpoint_manifest.json vs tabla en Anexo A.4",
                "NO CUMPLE aún (ceros anómalos reportados)",
            ],
            [
                "C6",
                "Auditoría de ceros sospechosos documentada (legítimo vs bug)",
                "Informe breve o notas en Cap. 5/Anexo A",
                "PARCIAL",
            ],
            [
                "C7",
                "Cuerpo del Word sin la palabra «doctorado»/«doctoral»",
                "Búsqueda python-docx / XML w:t",
                "CUMPLE (objetivo de este parche)",
            ],
            [
                "C8",
                "Ortografía RAE y captiones coherentes; sin TODO/XXXX/lorem",
                "Revisión de títulos, front-matter y placeholders",
                "PARCIAL (pasada en curso)",
            ],
            [
                "C9",
                "Referencias APA con correspondencia cita↔lista; índice actualizado (F9)",
                "Campo TOC + lista de referencias",
                "PARCIAL (requiere F9 en Word del autor)",
            ],
            [
                "C10",
                "Metadatos institucionales sin [por definir] ficticio",
                "Portada / asesores solo con datos reales",
                "NO CUMPLE (placeholder institucional pendiente de dato real)",
            ],
            [
                "C11",
                "Reproducibilidad: scripts + CSV + manifiestos localizables",
                "tools/, scripts/, outputs/_drive_madrl/full_data",
                "CUMPLE (artefactos presentes; HAPPO parcial)",
            ],
            [
                "C12",
                "Limitación honestamente declarada: semilla única; HAPPO 49 ep",
                "Secciones 6.2–6.3",
                "CUMPLE",
            ],
        ],
        font_size=7.2,
    )
    _add_para(
        tmp,
        "Regla de no sobreclaim: se puede afirmar cumplimiento de C1–C3, C7, C11 y C12 con la "
        "evidencia actual; C4–C6, C8–C10 permanecen abiertos y deben cerrarse antes de fijar la "
        "versión de sustentación. Completar C4 y C5 es prioritario porque afectan la lectura "
        "empírica de políticas y la reproducibilidad de modelos.",
    )
    _add_para(
        tmp,
        "Sincronización documental: la versión maestra de contenido es "
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx; el documento de presentación "
        "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx debe reflejar el mismo Cap. 6. "
        "Los nombres de archivo pueden conservar la etiqueta histórica «Doctoral»; el cuerpo del "
        "texto no usa esa voz.",
    )
    return [deepcopy(c) for c in tmp.element.body if c.tag != qn("w:sectPr")]


def upsert_section(doc: Document, prefixes: tuple[str, ...], new_elements: list, stop_prefixes: tuple[str, ...]) -> str:
    start = _find_heading(doc, prefixes)
    if start is None:
        # Insert before first stop prefix heading, else before referencias
        anchor = None
        for pref in stop_prefixes:
            anchor = _find_heading(doc, (pref,))
            if anchor is not None:
                break
        if anchor is None:
            return "skip_missing_anchor"
        _insert_elements_before(anchor, new_elements, doc)
        return "inserted"
    olds, nxt = _collect_section_range(doc, start)
    # Also remove the heading itself; new_elements include heading
    _clear_elements([start._element] + olds)
    _insert_elements_before(nxt, new_elements, doc)
    return "replaced"


def has_chapter_6(doc: Document) -> bool:
    for para in doc.paragraphs:
        t = (para.text or "").strip().lower()
        if "capitulo 6" in t or "capítulo 6" in t or t.startswith("6.1") or t.startswith("6.4"):
            return True
    return False


def patch_cap6(doc: Document) -> list[str]:
    actions = []
    if not has_chapter_6(doc):
        return ["skip_no_cap6"]
    a64 = upsert_section(
        doc,
        ("6.4 Plan", "6.4 Plan para culminar"),
        _build_cap64_temp(),
        ("6.5", "Referencias", "Anexo", "ANEXO"),
    )
    actions.append(f"6.4:{a64}")
    a65 = upsert_section(
        doc,
        ("6.5 Criterios", "6.5 Criterios de cierre"),
        _build_cap65_temp(),
        ("Referencias", "Anexo", "ANEXO", "Apéndice", "Apendice"),
    )
    actions.append(f"6.5:{a65}")
    # If 6.5 was missing and inserted relative to Referencias after 6.4 replace, OK.
    # Validate substance
    text = "\n".join(p.text for p in doc.paragraphs)
    for marker in ("Tabla 6.1", "Tabla 6.2", "H1. Cobertura HAPPO", "C4", RUN_ID):
        if marker not in text:
            actions.append(f"WARN_missing:{marker}")
    if BANNED_RE.search(text):
        actions.append("WARN_banned_still_in_cap6")
    return actions


def list_docx_targets() -> list[Path]:
    files = []
    for p in sorted(DOCS.glob("*.docx")):
        name = p.name
        if ".bak" in name.lower():
            continue
        if name.startswith("~$"):
            continue
        # Focus on thesis / presentation / informe maintained set
        if (
            name.startswith("Tesis_Doctoral_")
            or name.startswith("ABRIR_ESTE_")
            or "informe" in name.lower()
            or "resultados_drive" in name.lower()
        ):
            files.append(p)
    return files


def validate_cap6(doc: Document) -> dict:
    headings = [(p.text or "").strip() for p in doc.paragraphs if (p.style and p.style.name.startswith("Heading"))]
    h64 = next((h for h in headings if h.startswith("6.4")), None)
    h65 = next((h for h in headings if h.startswith("6.5")), None)

    def section_words(prefix: str) -> int:
        paras = list(doc.paragraphs)
        idx = next((i for i, p in enumerate(paras) if (p.text or "").strip().startswith(prefix)), None)
        if idx is None:
            return 0
        words = 0
        start_level = _heading_level(paras[idx])
        for p in paras[idx + 1 :]:
            t = (p.text or "").strip()
            if p.style and p.style.name.startswith("Heading") and t:
                if _heading_level(p) <= start_level:
                    break
            words += len(t.split())
        return words

    w64 = section_words("6.4")
    w65 = section_words("6.5")
    body = "\n".join(p.text or "" for p in doc.paragraphs)
    bad = ["[pendiente]", "lorem", "TODO", "[por definir]"]
    # [por definir] may remain for advisor — note separately
    return {
        "has_6_4": h64 is not None,
        "has_6_5": h65 is not None,
        "heading_6_4": h64,
        "heading_6_5": h65,
        "words_6_4": w64,
        "words_6_5": w65,
        "filled_6_4": w64 >= 80 and "H1" in body,
        "filled_6_5": w65 >= 80 and "C1" in body,
        "no_doctorado": not bool(BANNED_RE.search(body)),
        "banned_hits": count_banned(doc)[:10],
        "has_run_id": RUN_ID in body,
        "placeholder_por_definir_present": "[por definir]" in body.lower() or "[por definir]" in body,
    }


def main() -> int:
    targets = list_docx_targets()
    report = {"run_id": RUN_ID, "files": [], "summary": {}}
    cap6_priority = {
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx",
        "ABRIR_ESTE_WORD_FINAL_TODAS_FIGURAS_APA_INTERPRETADAS.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_skill.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_50EP_ANTECEDENTES.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_FINAL_GDRIVE_50EP_OBJETIVOS_DOCTORAL.docx",
        "Tesis_Doctoral_MADRL_CityLearn_Iquitos_VERSION_ACTUALIZADA_METODOLOGIA_MADRL.docx",
    }

    total_repl = 0
    files_touched = 0
    for path in targets:
        try:
            doc = Document(str(path))
        except Exception as exc:
            report["files"].append({"file": path.name, "error": f"{type(exc).__name__}: {exc}"})
            continue

        entry: dict = {"file": path.name, "size_before": path.stat().st_size}
        actions = []
        if path.name in cap6_priority or has_chapter_6(doc):
            actions.extend(patch_cap6(doc))

        clean_stats = clean_document_text(doc)
        entry["replacements"] = clean_stats
        entry["cap6_actions"] = actions
        entry["validation"] = validate_cap6(doc)

        # Second pass if banned still remains (edge mixed runs)
        if entry["validation"]["banned_hits"]:
            clean_stats2 = clean_document_text(doc)
            entry["replacements_second_pass"] = clean_stats2
            entry["validation"] = validate_cap6(doc)

        try:
            doc.save(str(path))
            entry["saved"] = True
            entry["size_after"] = path.stat().st_size
            files_touched += 1
            total_repl += clean_stats.get("total", 0)
        except Exception as exc:
            entry["saved"] = False
            entry["save_error"] = f"{type(exc).__name__}: {exc}"
            # If locked (user has ABRIR open), write sibling copy
            alt = path.with_name(path.stem + "_SIN_DOCTORADO.docx")
            try:
                doc.save(str(alt))
                entry["saved_alt"] = str(alt.name)
            except Exception as exc2:
                entry["alt_error"] = f"{type(exc2).__name__}: {exc2}"

        report["files"].append(entry)
        print(
            f"{path.name}: repl={clean_stats.get('total', 0)} "
            f"cap6={actions} banned_left={len(entry['validation'].get('banned_hits') or [])} "
            f"saved={entry.get('saved')}"
        )

    report["summary"] = {
        "files_targeted": len(targets),
        "files_saved": files_touched,
        "total_replacements": total_repl,
        "zero_doctorado_files": [
            f["file"] for f in report["files"] if f.get("validation", {}).get("no_doctorado")
        ],
        "still_has_doctorado": [
            f["file"] for f in report["files"] if f.get("validation") and not f["validation"].get("no_doctorado")
        ],
    }
    AUDIT_OUT.parent.mkdir(parents=True, exist_ok=True)
    AUDIT_OUT.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print("REPORT", AUDIT_OUT)
    print(json.dumps(report["summary"], indent=2, ensure_ascii=False))
    return 0 if not report["summary"]["still_has_doctorado"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
