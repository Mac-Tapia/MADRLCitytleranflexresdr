#!/usr/bin/env python3
"""Integra KPIs distrito/edificio (rankings Drive 50 ep) en los 2 Word canónicos.

NO crea Word nuevos en docs/. Solo edita:
  - docs/Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx
  - docs/Inforne_tesisV4_FINAL_REVISADO_50_EPISODIOS.docx

Backups: outputs/_word_backups/
Idempotente: marker §5.4.5 KPIs distrito y edificio.
"""
from __future__ import annotations

import csv
import json
import re
import shutil
import sys
import unicodedata
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import CANONS, DOCS, RUN_ID  # noqa: E402

BACKUP_DIR = REPO / "outputs" / "_word_backups"
KPI_DIR = (
    REPO
    / "outputs"
    / RUN_ID
    / "resumen_comparativo"
    / "multiobjetivo"
    / "kpi_rankings_drive50"
)
REPORT = DOCS / "CAP5_DISTRICT_BUILDING_KPI_RANKINGS_PATCH_REPORT_2026-07-29.json"
STAMP = datetime.now().strftime("%Y%m%d_%H%M%S")
GREY = RGBColor(0x55, 0x55, 0x55)

MARKER = "5.4.5 KPIs distrito y edificio"
ANCHOR_55 = re.compile(r"(?i)^\s*5\.5\s+Contrastaci[oó]n\s+de\s+hip[oó]tesis")

FIGURES = (
    (
        KPI_DIR / "district_kpis_4madrl.png",
        "Figura 5.4.5a. KPIs de distrito (4 MADRL × 50 episodios Drive).",
    ),
    (
        KPI_DIR / "building_rank_MATD3_E1_flex.png",
        "Figura 5.4.5b. Ranking por edificio OE.1 flexibilidad — MATD3/E1.",
    ),
    (
        KPI_DIR / "building_rank_MATD3_E2_co2.png",
        "Figura 5.4.5c. Ranking por edificio OE.2 reducción CO₂ — MATD3/E2.",
    ),
    (
        KPI_DIR / "building_rank_MATD3_E3_cost.png",
        "Figura 5.4.5d. Ranking por edificio OE.3 reducción costo — MATD3/E3.",
    ),
    (
        KPI_DIR / "building_resource_control_matd3_e1.png",
        "Figura 5.4.5e. Control de recursos por edificio (BESS/EV) — MATD3/E1.",
    ),
)


def norm(s: str) -> str:
    folded = unicodedata.normalize("NFKD", s or "")
    folded = "".join(ch for ch in folded if not unicodedata.combining(ch))
    return re.sub(r"\s+", " ", folded.strip().lower())


def set_run_font(
    run,
    *,
    bold: bool = False,
    size: float = 12.0,
    italic: bool = False,
    grey: bool = False,
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if grey:
        run.font.color.rgb = GREY
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), "Times New Roman")


def insert_paragraph_after(
    anchor: Paragraph,
    text: str = "",
    *,
    bold: bool = False,
    italic: bool = False,
    size: float = 12.0,
    grey: bool = False,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    if text:
        run = para.add_run(text)
        set_run_font(run, bold=bold, italic=italic, size=size, grey=grey)
    return para


def set_heading_style(paragraph: Paragraph, level: int) -> None:
    try:
        paragraph.style = f"Heading {level}"
    except KeyError:
        if paragraph.runs:
            set_run_font(paragraph.runs[0], bold=True, size=12.0)


def insert_figure_after(
    anchor: Paragraph, path: Path, caption: str, width_cm: float = 14.0
) -> Paragraph:
    pic_para = insert_paragraph_after(anchor)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = insert_paragraph_after(pic_para, caption, italic=True, size=9.0, grey=True)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap


def insert_table_after(
    paragraph: Paragraph,
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    caption: str,
    note: str,
    *,
    font_size: float = 8.0,
) -> Paragraph:
    cap = insert_paragraph_after(paragraph, caption, italic=True, size=11.0)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=font_size)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=font_size)
    tbl = table._tbl
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)
    cap._p.addnext(tbl)
    note_el = OxmlElement("w:p")
    tbl.addnext(note_el)
    note_p = Paragraph(note_el, paragraph._parent)
    run = note_p.add_run(note)
    set_run_font(run, italic=True, size=10.0)
    return note_p


def already_patched(doc: Document) -> bool:
    return any(MARKER in (p.text or "") for p in doc.paragraphs)


def find_paragraph_before_55(doc: Document) -> Paragraph | None:
    """Return the paragraph immediately before §5.5 (insert after this)."""
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if ANCHOR_55.match((p.text or "").strip()):
            if i == 0:
                return None
            return paras[i - 1]
    return None


def load_district_rank_rows() -> list[list[str]]:
    path = KPI_DIR / "district_ranking_by_scenario.csv"
    by_scen: dict[str, list[dict[str, str]]] = {"E1": [], "E2": [], "E3": []}
    with path.open(encoding="utf-8", newline="") as f:
        for r in csv.DictReader(f):
            by_scen.setdefault(r["scenario"], []).append(r)
    rows: list[list[str]] = []
    labels = {"E1": "OE.1 flex", "E2": "OE.2 CO2", "E3": "OE.3 costo"}
    for scen in ("E1", "E2", "E3"):
        ranked = sorted(by_scen[scen], key=lambda x: int(x["rank"]))
        cells = [labels[scen]]
        for r in ranked:
            val = float(r["value"])
            vtxt = f"{val:.4f}" if abs(val) < 100 else f"{val:,.0f}"
            cells.append(f"#{r['rank']} {r['algorithm']} ({vtxt})")
        while len(cells) < 5:
            cells.append("—")
        rows.append(cells)
    return rows


def load_best_building_rows() -> list[list[str]]:
    path = KPI_DIR / "building_best_per_algo_scenario.csv"
    rows: list[list[str]] = []
    with path.open(encoding="utf-8", newline="") as f:
        for r in csv.DictReader(f):
            val = float(r["best_value"])
            vtxt = f"{val:.4f}" if abs(val) < 100 else f"{val:,.0f}"
            bess = r.get("battery_throughput_kwh") or ""
            evs = r.get("ev_departure_success_ratio") or ""
            try:
                bess_txt = f"{float(bess):,.0f}" if bess != "" else "—"
            except ValueError:
                bess_txt = "—"
            try:
                evs_txt = f"{float(evs):.3f}" if evs != "" else "—"
            except ValueError:
                evs_txt = "—"
            rows.append(
                [
                    r["algorithm"],
                    r["scenario"],
                    r["oe"],
                    f"B{int(r['best_building_id']):02d}",
                    (r["best_nombre"] or "")[:36],
                    vtxt,
                    bess_txt,
                    evs_txt,
                ]
            )
    return rows


def insert_block(doc: Document, anchor: Paragraph) -> list[str]:
    actions: list[str] = []
    cursor = insert_paragraph_after(
        anchor,
        "5.4.5 KPIs distrito y edificio (rankings Drive 50 episodios)",
        bold=True,
    )
    set_heading_style(cursor, 3)
    actions.append("heading_5.4.5")

    cursor = insert_paragraph_after(
        cursor,
        "Cálculo descriptivo sobre artefactos reales de la corrida canónica "
        f"{RUN_ID} (4 MADRL × 3 escenarios × 50 episodios). Fuente: "
        "kpi_recalc_20260728 (distrito y building_kpis) y "
        "building_behavior_summary.csv (comportamiento BESS/EV por edificio). "
        "No es caja negra: cada edificio reporta ganancia de flexibilidad "
        "(1 − ratio de importación vs baseline), reducción de CO₂ (−Δ emisiones) "
        "y reducción de costo (−Δ costo), junto con throughput BESS, carga EV, "
        "éxito de salida EV y rol de red.",
    )
    actions.append("intro")

    cursor = insert_table_after(
        cursor,
        doc,
        ["Escenario", "1.º", "2.º", "3.º", "4.º"],
        load_district_rank_rows(),
        "Tabla 5.4.5a. Ranking distrital por escenario (menor KPI primario = mejor).",
        "Nota. E1: flex_composite; E2: carbon_emissions_delta (kg); "
        "E3: electricity_cost_delta (EUR). Fuente: all_core_kpis_wide.csv.",
    )
    actions.append("table_district")

    cursor = insert_table_after(
        cursor,
        doc,
        ["Algo", "Esc.", "OE", "Mejor edif.", "Nombre", "Valor", "BESS thr.", "EV éxito"],
        load_best_building_rows(),
        "Tabla 5.4.5b. Mejor edificio por algoritmo × escenario (17 edificios).",
        "Nota. OE.1: flex_gain_vs_baseline; OE.2: co2_reduction_kgco2; "
        "OE.3: cost_reduction_eur. Valores negativos en reducción indican "
        "menor empeoramiento (sin reducción neta). Control de recursos: "
        "BESS throughput y éxito EV desde evaluate_v2 / behavior summary.",
        font_size=7.5,
    )
    actions.append("table_buildings")

    cursor = insert_paragraph_after(
        cursor,
        "En la referencia MATD3, el mejor edificio en flexibilidad (E1) es B14 "
        "(Autoridad Portuaria Nacional); en CO₂ (E2) y costo (E3) es B12 "
        "(EsSalud). A nivel distrito, MATD3 lidera E1 y E2; MAAC lidera E3. "
        "El comportamiento no se resume solo en un score: las figuras siguientes "
        "exponen el ranking completo de los 17 edificios y el control de recursos "
        "(BESS/EV).",
    )

    for path, caption in FIGURES:
        if not path.is_file():
            raise FileNotFoundError(path)
        cursor = insert_figure_after(cursor, path, caption)
        actions.append(caption)

    cursor = insert_paragraph_after(
        cursor,
        "Nota. Sección descriptiva; no sustituye §5.3 ni §5.5. Artefactos: "
        "outputs/.../multiobjetivo/kpi_rankings_drive50/.",
        italic=True,
        size=9.0,
        grey=True,
    )
    actions.append("closing_note")
    return actions


def patch_doc(path: Path) -> dict:
    for fig_path, _ in FIGURES:
        if not fig_path.is_file():
            raise FileNotFoundError(fig_path)
    for req in (
        KPI_DIR / "district_ranking_by_scenario.csv",
        KPI_DIR / "building_best_per_algo_scenario.csv",
    ):
        if not req.is_file():
            raise FileNotFoundError(req)

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{path.stem}_antes_kpi_rankings_{STAMP}{path.suffix}"
    shutil.copy2(path, backup)

    doc = Document(str(path))
    if already_patched(doc):
        return {
            "path": str(path.relative_to(REPO)).replace("\\", "/"),
            "ok": True,
            "skipped": True,
            "reason": "already_has_5.4.5_kpi_rankings",
            "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        }

    anchor = find_paragraph_before_55(doc)
    if anchor is None:
        return {
            "path": str(path.relative_to(REPO)).replace("\\", "/"),
            "ok": False,
            "error": "anchor before 5.5 not found",
            "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        }

    actions = insert_block(doc, anchor)
    doc.save(str(path))

    verify = Document(str(path))
    text = "\n".join(p.text or "" for p in verify.paragraphs)
    return {
        "path": str(path.relative_to(REPO)).replace("\\", "/"),
        "ok": True,
        "skipped": False,
        "backup": str(backup.relative_to(REPO)).replace("\\", "/"),
        "actions": actions,
        "checks": {
            "has_marker": MARKER in text,
            "has_table_a": "Tabla 5.4.5a" in text,
            "has_table_b": "Tabla 5.4.5b" in text,
            "has_fig_a": "Figura 5.4.5a" in text,
            "has_fig_e": "Figura 5.4.5e" in text,
            "has_b12": "EsSalud" in text or "ESSALUD" in text.upper(),
            "has_b14": "PORTUARIA" in text.upper(),
        },
    }


def main() -> int:
    report: dict = {
        "stamp": STAMP,
        "run_id": RUN_ID,
        "kpi_dir": str(KPI_DIR.relative_to(REPO)).replace("\\", "/"),
        "files": {},
    }
    ok_all = True
    for path in CANONS:
        if not (path.is_file() and path.stat().st_size > 0):
            report["files"][path.name] = {"ok": False, "error": "missing"}
            ok_all = False
            continue
        result = patch_doc(path)
        report["files"][path.name] = result
        ok_all = ok_all and bool(result.get("ok"))
        print(json.dumps(result, ensure_ascii=False))

    # Política: no crear Word nuevos; solo report JSON de auditoría.
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print("WROTE", REPORT)
    return 0 if ok_all else 1


if __name__ == "__main__":
    raise SystemExit(main())
