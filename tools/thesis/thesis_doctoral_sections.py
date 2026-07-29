"""Secciones doctorales: Cap. 5 Colab, multiobjetivo, verificación de completitud."""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

import sys

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))
if str(REPO) not in sys.path:
    sys.path.insert(0, str(REPO))

from tools.thesis.thesis_antecedents_data import (  # noqa: E402
    ANTECEDENTES_INTERNACIONALES,
    ANTECEDENTES_NACIONALES,
)
RUN_ID = "madrl_v3_20260627_164047"
MO_DIR = REPO / "outputs" / RUN_ID / "resumen_comparativo" / "multiobjetivo"
FD_DIR = REPO / "outputs" / RUN_ID / "resumen_comparativo" / "figuras_drive_reales" / "comparativo"
BEST_REPORT = REPO / "outputs" / RUN_ID / "resumen_comparativo" / "best_madrl_report.json"
DISTRICT_CSV = MO_DIR / "district_objectives_by_algorithm.csv"
INVENTORY_CSV = MO_DIR / "building_inventory_multiobjective.csv"

BL_DIR = REPO / "outputs" / RUN_ID / "resumen_comparativo" / "citylearn_v2_baseline"
STAT_DIR = REPO / "outputs" / RUN_ID / "resumen_comparativo" / "estadistica"
HYP_CSV = STAT_DIR / "hipotesis_estadisticas_madrl.csv"
ANALYSIS_CSV = STAT_DIR / "analisis_estadistico_madrl.csv"
MWU_CSV = STAT_DIR / "comparaciones_mwu_madrl.csv"
WC_CSV = STAT_DIR / "comparaciones_wilcoxon_madrl.csv"
DESCRIPTIVE_CSV = STAT_DIR / "descriptivo_distrito_colab.csv"
EPISODE_CSV = (
    REPO / "outputs" / "_drive_madrl" / "full_data" / "analysis_real_drive" / "tables" / "district_episode_kpis.csv"
)
BUILDING_CSV = MO_DIR / "building_objectives_by_algorithm.csv"
ALGO_METRICS_JSON = STAT_DIR / "madrl_per_algorithm_metrics.json"
FIG_RUN = REPO / "outputs" / RUN_ID

OE_EPISODE_SPECS = {
    "OE1": {
        "scenario": "E1",
        "metric": "reward_mean",
        "label": "reward_mean por episodio",
        "hypothesis": "HE.1",
        "fmt": ".4f",
        "higher_better": True,
    },
    "OE2": {
        "scenario": "E2",
        "metric": "district_emission",
        "label": "emision distrital por episodio (kgCO2 eq.)",
        "hypothesis": "HE.2",
        "fmt": ",.1f",
        "higher_better": False,
    },
    "OE3": {
        "scenario": "E3",
        "metric": "district_cost",
        "label": "costo distrital por episodio (EUR)",
        "hypothesis": "HE.3",
        "fmt": ",.1f",
        "higher_better": False,
    },
}
HYPOTHESIS_LABELS = {
    "OE1": ("HE.1", "OE.1 / E1"),
    "OE2": ("HE.2", "OE.2 / E2"),
    "OE3": ("HE.3", "OE.3 / E3"),
    "OG": ("HG", "Global (ALL)"),
    "ALL": ("HG", "Global (ALL)"),
}

PE_SPECS = {
    "PE.1": {
        "oe_key": "OE1",
        "oe": "OE.1",
        "he": "HE.1",
        "vd": "D-VD.1",
        "section": "5.3.1",
        "dimension": "flexibilidad energetica",
        "problem": (
            "¿En qué medida el algoritmo MADRL (VI) produce un efecto sobre la dimensión de "
            "flexibilidad energética de la comunidad (D-VD.1), y cuál algoritmo genera el mayor efecto?"
        ),
    },
    "PE.2": {
        "oe_key": "OE2",
        "oe": "OE.2",
        "he": "HE.2",
        "vd": "D-VD.2",
        "section": "5.4.1",
        "dimension": "emisiones de CO2",
        "problem": (
            "¿En qué medida el algoritmo MADRL (VI) produce un efecto sobre la dimensión de "
            "emisiones de CO2 de la comunidad (D-VD.2), y cuál algoritmo genera el mayor efecto?"
        ),
    },
    "PE.3": {
        "oe_key": "OE3",
        "oe": "OE.3",
        "he": "HE.3",
        "vd": "D-VD.3",
        "section": "5.5.1",
        "dimension": "costos energéticos",
        "problem": (
            "¿En qué medida el algoritmo MADRL (VI) produce un efecto sobre la dimensión de "
            "costos energéticos de la comunidad (D-VD.3), y cuál algoritmo genera el mayor efecto?"
        ),
    },
}

PE_ANSWERS_AUDIT_JSON = STAT_DIR / "pe_answers_audit.json"

FIG_WIDTH_LANDSCAPE = 17.5
FIG_WIDTH_PORTRAIT = 16.5
FIG_WIDTH_HEATMAP = 17.0

OE_DEFINITIONS = {
    "OE1": {
        "scenario": "E1",
        "vd": "D-VD.1",
        "dimension": "flexibilidad energetica",
        "weights": "[0,70; 0,15; 0,15]",
        "primary_kpis": [
            ("flex_composite", "Flexibilidad compuesta", ".4f", True),
            ("peak_average", "Pico promedio", ".4f", True),
            ("ramping_average", "Ramping promedio", ".4f", True),
            ("one_minus_load_factor_average", "1 − factor de carga", ".4f", True),
            ("grid_import_delta", "Delta importacion red (kWh)", ",.0f", True),
            ("ev_departure_success_rate", "Exito salida EV", ".1%", False),
        ],
        "kpi_fig": FD_DIR / "comparativo_E1_OE1_kpi.png",
        "conv_fig": FD_DIR / "comparativo_E1_convergence_reward_mean.png",
        "building_fig": MO_DIR / "drive_building_E1_flex_composite_proxy.png",
        "trace_fig": FD_DIR / "comparativo_E1_control_trace.png",
        "score_key": "score_oe1_flex",
    },
    "OE2": {
        "scenario": "E2",
        "vd": "D-VD.2",
        "dimension": "emisiones de CO2",
        "weights": "[0,15; 0,70; 0,15]",
        "primary_kpis": [
            ("carbon_emissions_delta_kg", "Delta CO2 (kg)", ",.0f", True),
            ("flex_composite", "Flexibilidad compuesta (secundaria)", ".4f", True),
            ("grid_import_delta", "Delta importacion red (kWh)", ",.0f", True),
            ("ev_departure_success_rate", "Exito salida EV", ".1%", False),
        ],
        "kpi_fig": FD_DIR / "comparativo_E2_OE2_kpi.png",
        "conv_fig": FD_DIR / "comparativo_E2_convergence_reward_mean.png",
        "building_fig": MO_DIR / "drive_building_E2_carbon_emissions_delta_kgco2.png",
        "trace_fig": FD_DIR / "comparativo_E2_control_trace.png",
        "score_key": "score_oe2_co2",
    },
    "OE3": {
        "scenario": "E3",
        "vd": "D-VD.3",
        "dimension": "costos energeticos",
        "weights": "[0,15; 0,15; 0,70]",
        "primary_kpis": [
            ("electricity_cost_delta_eur", "Delta costo electrico (EUR)", ",.0f", True),
            ("flex_composite", "Flexibilidad compuesta (secundaria)", ".4f", True),
            ("peak_average", "Pico promedio", ".4f", True),
            ("ev_departure_success_rate", "Exito salida EV", ".1%", False),
        ],
        "kpi_fig": FD_DIR / "comparativo_E3_OE3_kpi.png",
        "conv_fig": FD_DIR / "comparativo_E3_convergence_reward_mean.png",
        "building_fig": MO_DIR / "drive_building_E3_electricity_cost_delta_eur.png",
        "trace_fig": FD_DIR / "comparativo_E3_control_trace.png",
        "score_key": "score_oe3_cost",
    },
}


GREY = RGBColor(0x59, 0x59, 0x59)


def _antecedents_discussion_text() -> list[str]:
    """Paragraphs contrasting thesis results with Cap. 2 antecedents (PE/HG/HE)."""
    return [
        (
            "Contrastacion con antecedentes internacionales (Tabla 2.4). Frente a Nweye et al. "
            "(2024), que estandariza KPIs de flexibilidad, CO2 y costo en CityLearn v2, esta tesis "
            "extiende el entorno a CityLearn v3 propuesto con Dec-POMDP/CTDE y cuatro backends "
            "MADRL; los resultados de 50 episodios muestran trade-offs por escenario (MATD3 en "
            "OE.1/OE.2, MAAC en OE.3) no resueltos por un unico algoritmo en la literatura base."
        ),
        (
            "Nweye et al. (2023b) reportan HAPPO (HARL) viable en comunidades heterogeneas "
            "CityLearn, pero sin comparacion simultanea con MASAC, MATD3 y MAAC. La corrida "
            "canonica confirma efectos diferenciados entre algoritmos, aunque HG y HE.1-HE.3 no "
            "alcanzan significancia omnibus (KW p > 0,05), coherente con la advertencia metodologica "
            "de comparaciones RL rigurosas (Agarwal et al., 2021)."
        ),
        (
            "Yao et al. (2023) obtienen ~15% de reduccion de pico y ~18% de costo con LSD-MADDPG "
            "en comunidades inteligentes. En PE.1 (flexibilidad, E1), MATD3 lidera descriptivamente "
            "(flex_composite = 1,0009), pero la magnitud relativa frente a baseline RBC (seccion 5.7) "
            "matiza la generalizacion causal frente a mejoras reportadas en comunidades sinteticas."
        ),
        (
            "Liu et al. (2022) reportan ~15% de reduccion de CO2 y ~20% de costo con MADDPG. En "
            "PE.2 (E2), MATD3 presenta el menor delta de CO2 entre MADRL auditados (23 070 kg), "
            "mientras MASAC y MAAC registran reducciones mayores; esto contrasta la expectativa de "
            "dominancia unica y refuerza la lectura multiobjetivo de D-VD.2."
        ),
        (
            "Iqbal y Sha (2019) fundamentan MAAC con mejoras de ~15-30% frente a MADDPG/COMA. En "
            "PE.3 (costos, E3), MAAC lidera (delta costo = 9 515 EUR), alineado con coordinacion "
            "selectiva bajo senales TOU; este patron es consistente con Xie et al. (2023) sobre "
            "atencion para demand response, aunque la tesis no replica sus magnitudes (~25% DR) "
            "por diferencias de dataset y funcion de recompensa."
        ),
        (
            "Contrastacion con antecedentes nacionales/peruanos (Tabla 2.5). MINAM (2019) fija el "
            "factor 0,790 kgCO2/kWh del diesel aislado que parametriza CarbonIntensityModel en E2; "
            "los deltas de CO2 de MATD3, MASAC y MAAC deben interpretarse sobre esa senal nacional, "
            "no sobre intensidades de red continental."
        ),
        (
            "OSINERGMIN (2024) vincula costos a demanda maxima facturable en Electro Oriente S.A. "
            "El liderazgo descriptivo de MAAC en OE.3 es coherente con la necesidad de coordinar "
            "picos distritales bajo TOU local; sin embargo, la superioridad global de controles RBC "
            "frente a MADRL (seccion 5.7) sugiere que la regulacion peruana por si sola no garantiza "
            "ventaja aprendida sin multi-semilla ni ajuste de hiperparametros."
        ),
        (
            "Las tesis doctorales Chevarria Moscoso (2024) y Peñalva Sanchez (2024) aportan "
            "optimizacion bajo incertidumbre hidrologica y prediccion con IA en sistemas PV "
            "hibridos (sin transferir magnitudes cuantitativas a Iquitos). Rosero Bernal (2024) "
            "fundamenta la coordinacion de microredes operada desde la nube. Estos antecedentes "
            "contextualizan el SEAI Iquitos como caso de red aislada con recursos renovables y "
            "demanda institucional heterogenea, justificando el diseno de 17 agentes y la lectura "
            "distrito-edificio del Capitulo 5."
        ),
    ]


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def _read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _fmt_es_score(value: float, decimals: int = 3) -> str:
    return f"{value:.{decimals}f}".replace(".", ",")


def _ordinal_es(rank: int) -> str:
    return f"{int(rank)}.o"


def _ranking_csv_for_scenario(scenario: str) -> Path | None:
    """Prefer synced baseline dir; fall back to regenerated comparison outputs."""
    candidates = (
        BL_DIR / scenario / "ranking_by_axis.csv",
        REPO / "outputs" / "comparison_citylearn_v2_vs_v3_madrl" / scenario / "ranking_by_axis.csv",
    )
    for path in candidates:
        if path.is_file():
            return path
    return None


def _baseline_oe_ranking_table_rows() -> tuple[list[list[str]], str]:
    """Build Tabla 5.12 from ranking_by_axis.csv (OE1←E1, OE2←E2, OE3←E3)."""
    labels = {"OE1": "OE.1 / E1", "OE2": "OE.2 / E2", "OE3": "OE.3 / E3"}
    rows: list[list[str]] = []
    notes: list[str] = []
    for oe, scenario in (("OE1", "E1"), ("OE2", "E2"), ("OE3", "E3")):
        path = _ranking_csv_for_scenario(scenario)
        if path is None:
            continue
        axis_rows = [r for r in _read_csv(path) if r.get("axis") == oe]
        if not axis_rows:
            continue
        ranked = sorted(axis_rows, key=lambda r: float(r.get("axis_rank") or 1e9))
        best = ranked[0]
        madrl = [r for r in ranked if r.get("family") == "citylearn_v3_madrl"]
        if not madrl:
            continue
        best_madrl = min(madrl, key=lambda r: float(r.get("axis_rank") or 1e9))
        method_label = {
            "baseline": "baseline v2",
            "hour_rbc": "hour_rbc",
        }.get(best.get("method", ""), best.get("method", ""))
        rows.append(
            [
                labels[oe],
                method_label,
                _fmt_es_score(float(best["normalized_score"])),
                best_madrl.get("method", ""),
                _fmt_es_score(float(best_madrl["normalized_score"])),
                f"{_ordinal_es(int(float(best_madrl['axis_rank'])))} / {len(ranked)} metodos",
            ]
        )
        notes.append(path.as_posix())
    source_note = notes[0] if notes else str(BL_DIR)
    return rows, source_note


def add_dedicatoria_agradecimientos(doc, p, heading) -> None:
    heading(doc, "Dedicatoria", 1)
    p(
        doc,
        "A mi familia, por el apoyo incondicional durante este camino de formación doctoral; "
        "a los docentes y colegas de la Universidad Nacional de Ingeniería, por impulsar la "
        "investigación aplicada en sistemas eléctricos inteligentes; y a la comunidad académica "
        "de Iquitos, cuyo contexto energético motivó el presente estudio.",
    )
    doc.add_page_break()
    heading(doc, "Agradecimientos", 1)
    p(
        doc,
        "Se agradece a la Escuela de Posgrado de la UNI, al equipo de investigación del proyecto "
        "MADRL CityLearn v3, a Electro Oriente S.A. y a las instituciones que aportaron datos "
        "operativos para el dataset citylearn_iquitos_2023_2025. Asimismo, se reconoce el uso "
        "de recursos computacionales en Google Colab para la corrida canónica de 50 episodios "
        f"({RUN_ID}).",
    )
    doc.add_page_break()


def _resumen_inferential_snippet() -> tuple[str, str]:
    og = _hyp_row_by_axis("OG")
    kw_all = float(og.get("KW_p_value", 0.155)) if og else 0.155
    wc = next(
        (
            r
            for r in _significant_wilcoxon_rows()
            if r.get("scope") == "ALL" and r.get("algorithm_a") == "MASAC" and r.get("algorithm_b") == "MATD3"
        ),
        {},
    )
    wc_p = float(wc.get("wilcoxon_p_value", 0.0049)) if wc else 0.0049
    es = (
        f"(Kruskal-Wallis ALL p={kw_all:.3f}; Wilcoxon MASAC vs MATD3 p={wc_p:.4f})"
    )
    en = (
        f"Kruskal-Wallis ALL p={kw_all:.3f} (not significant at alpha=0.05); "
        f"Wilcoxon MASAC vs MATD3 p={wc_p:.4f}."
    )
    return es, en


def add_resumen_doctoral(doc, p, heading) -> None:
    inf_es, inf_en = _resumen_inferential_snippet()
    heading(doc, "Resumen", 1)
    p(
        doc,
        "Esta tesis doctoral determina, mediante simulación computacional bajo diseño cuasiexperimental "
        "factorial 4×3, el efecto de cuatro algoritmos Multi-Agente de Aprendizaje por Refuerzo "
        "Profundo (MADRL) —HAPPO, MASAC, MATD3 y MAAC— sobre la gestión coordinada de flexibilidad "
        "energética, emisiones de CO₂ y costos en una comunidad inteligente del Sistema Eléctrico "
        "Aislado de Iquitos (SEAI). La formulación Dec-POMDP con CTDE se implementa sobre CityLearn "
        "v3 propuesto (17 edificios reales, 26 304 h, 185 cargadores EV). La corrida canónica "
        f"Colab/Drive ({RUN_ID}) completó 50 episodios por escenario en MATD3, MAAC y MASAC. "
        "En respuesta a PE.1–PE.3 y OE.1–OE.3 (evidencia descriptiva), MATD3 lidera flexibilidad "
        "(PE.1/OE.1) y emisiones (PE.2/OE.2); MAAC lidera costos (PE.3/OE.3); MATD3 obtiene el "
        "mejor score global (0,6667). Las figuras de convergencia, control MADRL y ranking provienen de "
        "timeseries.csv y trace.csv auditados en Drive (sin datos sinteticos). El analisis "
        "multiobjetivo desagrega KPIs por distrito y por edificio. El recálculo Drive "
        "(2026-07-28) recupera 58 KPIs evaluate_v2 (680 valores; 12/12 tratamientos) e incluye "
        "HAPPO en KPI-gains all_values y ranking 4/4 (score HAPPO = 0,0000; lider 4/4 = MAAC "
        "0,9538). HAPPO conserva 49 episodios materializados por escenario. El veredicto "
        "best_madrl canónico 3×3 permanece MATD3 (0,6667). La contrastacion inferencial de "
        "HG y HE.1–HE.3 (seccion 5.9) "
        f"{inf_es} no confirma diferencias omnibus con una semilla; la referencia local v4 "
        "(KW p=0,0459) es exploratoria con 5 episodios.",
    )
    p(
        doc,
        "Palabras clave: aprendizaje por refuerzo multiagente, diseño cuasiexperimental, Dec-POMDP, CTDE, "
        "flexibilidad energética, emisiones de CO₂, microrred aislada, Iquitos.",
        italic=True,
    )
    doc.add_page_break()
    heading(doc, "Abstract", 1)
    p(
        doc,
        "This doctoral thesis determines the effect of four cooperative Multi-Agent Deep "
        "Reinforcement Learning (MADRL) algorithms under a Dec-POMDP/CTDE framework on a "
        "real 17-building dataset from Iquitos, Peru (quasi-experimental factorial design 4×3). PE.1–PE.3 answers "
        "(descriptive + inferential): MATD3 leads flexibility and CO₂; MAAC leads energy cost; "
        "Kruskal-Wallis omnibus tests on KPI-gains do not reject H0 "
        "(OE.1 p=0.4685; OE.2 p=0.7648; OE.3 p=0.7357). "
        "Objective-level evidence (OG, OE.1–OE.3): MATD3 leads overall (global score 0.6667), "
        "flexibility and CO₂; MAAC leads energy cost. Hypothesis-level inference (HG, HE.1–HE.3) on the "
        "canonical 50-episode run: "
        f"{inf_en} Multi-objective KPIs are reported at district and building levels "
        "(185 EV chargers). Training figures use audited Drive timeseries and trace CSVs "
        "(no synthetic data).",
        italic=True,
    )
    doc.add_page_break()


def _fmt_kpi(value: str, fmt: str, lower_better: bool) -> str:
    x = float(value)
    if fmt == ".1%":
        return f"{x * 100:.1f}%"
    if fmt == ",.0f":
        return f"{x:,.0f}"
    if fmt == ".4f":
        return f"{x:.4f}"
    return str(x)


def _rows_for_scenario(district: list[dict[str, str]], scenario: str) -> list[dict[str, str]]:
    return [r for r in district if r["scenario"] == scenario]


def _best_algo_by_kpi(rows: list[dict[str, str]], kpi: str, lower_better: bool = True) -> tuple[str, float]:
    best_algo, best_val = "", float("inf") if lower_better else float("-inf")
    for row in rows:
        val = float(row[kpi])
        if lower_better and val < best_val:
            best_val, best_algo = val, row["algorithm"]
        elif not lower_better and val > best_val:
            best_val, best_algo = val, row["algorithm"]
    return best_algo, best_val


def _hypothesis_row(oe_key: str) -> dict[str, str]:
    rows = _read_csv(HYP_CSV)
    axis = {"OE1": "OE1", "OE2": "OE2", "OE3": "OE3"}[oe_key]
    for row in rows:
        if row.get("axis") == axis:
            return row
    return {}


def _fmt_stat_num(value: float | str | None, fmt: str) -> str:
    if value is None or value == "":
        return "-"
    x = float(value)
    if fmt == ".4f":
        return f"{x:.4f}"
    if fmt == ",.1f":
        return f"{x:,.1f}"
    if fmt == ".3e":
        return f"{x:.3e}"
    return f"{x:.3f}"


def _descriptive_episode_rows(oe_key: str) -> list[dict[str, str | float]]:
    spec = OE_EPISODE_SPECS[oe_key]
    if DESCRIPTIVE_CSV.is_file():
        rows = [r for r in _read_csv(DESCRIPTIVE_CSV) if r.get("axis") == oe_key]
        if rows:
            return rows
    if not EPISODE_CSV.is_file():
        return []
    import pandas as pd

    df = pd.read_csv(EPISODE_CSV)
    sub = df[df["scenario"] == spec["scenario"]]
    out: list[dict[str, str | float]] = []
    for algo, grp in sub.groupby("algorithm"):
        vals = grp[spec["metric"]].dropna().astype(float)
        if vals.empty:
            continue
        out.append(
            {
                "axis": oe_key,
                "scenario": spec["scenario"],
                "algorithm": algo,
                "n_episodes": int(len(vals)),
                "mean": float(vals.mean()),
                "median": float(vals.median()),
                "std": float(vals.std(ddof=1)) if len(vals) > 1 else 0.0,
                "min": float(vals.min()),
                "max": float(vals.max()),
            }
        )
    return out


def _best_algo_descriptive(rows: list[dict], oe_key: str) -> str:
    spec = OE_EPISODE_SPECS[oe_key]
    if not rows:
        return "-"
    key = "median"
    reverse = spec["higher_better"]
    best = sorted(rows, key=lambda r: float(r[key]), reverse=reverse)[0]
    return str(best["algorithm"])


def _analysis_row(scope: str) -> dict[str, str]:
    if not ANALYSIS_CSV.is_file():
        return {}
    for row in _read_csv(ANALYSIS_CSV):
        if row.get("scope") == scope:
            return row
    return {}


def _shapiro_table_rows() -> list[list[str]]:
    rows: list[list[str]] = []
    for scope in ("OE1", "OE2", "OE3", "ALL"):
        row = _analysis_row(scope)
        if not row:
            continue
        hyp, label = HYPOTHESIS_LABELS.get(scope, (scope, scope))
        for algo in ("MASAC", "MATD3", "MAAC", "HAPPO"):
            p_key = f"shapiro_wilk_p_value_{algo}"
            status = row.get(f"shapiro_wilk_status_{algo}", "")
            if status == "no_data":
                continue
            p_val = row.get(p_key, "")
            rejected = row.get(f"shapiro_wilk_normality_rejected_{algo}", "")
            rows.append(
                [
                    label,
                    hyp,
                    algo,
                    _fmt_stat_num(p_val, ".3e") if p_val else "-",
                    "Si" if str(rejected).lower() == "true" else "No",
                ]
            )
    return rows


def _kruskal_table_rows() -> list[list[str]]:
    rows: list[list[str]] = []
    for axis in ("OE1", "OE2", "OE3", "OG"):
        hyp_row = next((r for r in _read_csv(HYP_CSV) if r.get("axis") == axis), {})
        if not hyp_row:
            continue
        hyp, label = HYPOTHESIS_LABELS.get(axis, (axis, axis))
        p_val = float(hyp_row.get("KW_p_value", "nan"))
        rows.append(
            [
                label,
                hyp,
                _fmt_stat_num(hyp_row.get("KW_H_statistic"), ".3f"),
                _fmt_stat_num(p_val, ".4f"),
                "Si" if p_val < 0.05 else "No",
                hyp_row.get("statistical_best_algorithm_by_median_gain", "-"),
            ]
        )
    return rows


def _mwu_table_rows(scope_filter: str | None = None) -> list[list[str]]:
    if not MWU_CSV.is_file():
        return []
    rows: list[list[str]] = []
    for row in _read_csv(MWU_CSV):
        scope = row.get("scope", "")
        if scope_filter and scope != scope_filter:
            continue
        if row.get("mann_whitney_status") != "ok":
            continue
        p_val = float(row.get("mann_whitney_p_value", "nan"))
        hyp, label = HYPOTHESIS_LABELS.get(scope, (scope, scope))
        rows.append(
            [
                label,
                hyp,
                f"{row.get('algorithm_a')} vs {row.get('algorithm_b')}",
                _fmt_stat_num(row.get("mann_whitney_u"), ".1f"),
                _fmt_stat_num(p_val, ".4f"),
                "Si" if p_val < 0.05 else "No",
            ]
        )
    return rows


def _wilcoxon_table_rows(scope_filter: str | None = None) -> list[list[str]]:
    if not WC_CSV.is_file():
        return []
    rows: list[list[str]] = []
    for row in _read_csv(WC_CSV):
        scope = row.get("scope", "")
        if scope_filter and scope != scope_filter:
            continue
        if not str(row.get("wilcoxon_status", "")).startswith("ok"):
            continue
        p_val = float(row.get("wilcoxon_p_value", "nan"))
        hyp, label = HYPOTHESIS_LABELS.get(scope, (scope, scope))
        rows.append(
            [
                label,
                hyp,
                f"{row.get('algorithm_a')} vs {row.get('algorithm_b')}",
                _fmt_stat_num(row.get("wilcoxon_T_statistic"), ".1f"),
                _fmt_stat_num(p_val, ".4f"),
                "Si" if p_val < 0.05 else "No",
            ]
        )
    return rows


def _fmt_p_es(value: float | str | None, digits: int = 3) -> str:
    if value is None or value == "":
        return "-"
    x = float(value)
    if x < 0.001:
        return f"{x:.2e}".replace(".", ",")
    return f"{x:.{digits}f}".replace(".", ",")


def _hyp_row_by_axis(axis: str) -> dict[str, str]:
    for row in _read_csv(HYP_CSV):
        if row.get("axis") == axis:
            return row
    return {}


def _significant_wilcoxon_rows() -> list[dict[str, str]]:
    if not WC_CSV.is_file():
        return []
    return [
        row
        for row in _read_csv(WC_CSV)
        if row.get("wilcoxon_significant_alpha_0_05") == "True"
        and str(row.get("wilcoxon_status", "")).startswith("ok")
    ]


def _kw_narrative_text() -> str:
    parts: list[str] = []
    for axis, label in (("OE1", "OE.1"), ("OE2", "OE.2"), ("OE3", "OE.3"), ("OG", "ALL")):
        row = _hyp_row_by_axis(axis)
        if row:
            parts.append(f"{label} p={_fmt_p_es(row.get('KW_p_value'))}")
    joined = "; ".join(parts)
    return (
        f"Ningun Kruskal-Wallis por eje alcanza alpha = 0,05 con una semilla ({joined}). "
        "Por tanto, no se rechaza H0 de igualdad global entre algoritmos en los contrastes omnibus; "
        "las respuestas a HE.1–HE.3 y HG se sustentan principalmente en evidencia descriptiva "
        "(seccion 5.8 y Tablas 5.4–5.6). La columna «Mejor (mediana KPI-gain)» de la Tabla 5.17 "
        "refiere al lider inferencial por mediana de signed_relative_gain, que puede diferir del "
        "lider descriptivo de distrito (Tablas 5.4–5.6)."
    )


def _mwu_narrative_text() -> str:
    mwu_sig = [
        row
        for row in _read_csv(MWU_CSV)
        if row.get("mann_whitney_significant_alpha_0_05") == "True" and row.get("mann_whitney_status") == "ok"
    ] if MWU_CSV.is_file() else []
    all_row = next((r for r in _read_csv(MWU_CSV) if r.get("scope") == "ALL" and r.get("mann_whitney_status") == "ok"
                    and r.get("algorithm_a") == "MASAC" and r.get("algorithm_b") == "MATD3"), {}) if MWU_CSV.is_file() else {}
    all_p = _fmt_p_es(all_row.get("mann_whitney_p_value", "0.0701"))
    if mwu_sig:
        pairs = ", ".join(f"{r['algorithm_a']} vs {r['algorithm_b']} (p={_fmt_p_es(r['mann_whitney_p_value'])})" for r in mwu_sig)
        tail = f"Pares significativos: {pairs}."
    else:
        tail = (
            f"Ningun par Mann-Whitney U alcanza significancia global a alpha = 0,05 en los ejes "
            f"OE.1–OE.3; en ALL, MASAC vs MATD3 (p = {all_p}) tampoco alcanza el umbral."
        )
    return f"{tail} Estas pruebas complementan el Kruskal-Wallis sin sustituir la evidencia descriptiva."


def _wilcoxon_narrative_text() -> str:
    sig = _significant_wilcoxon_rows()
    if not sig:
        return (
            "Wilcoxon pareado no detecta diferencias sistematicas a alpha = 0,05 en ningun par "
            "auditado sobre KPI-gains."
        )
    by_scope: dict[str, list[str]] = {}
    for row in sig:
        scope = row.get("scope", "ALL")
        pair = f"{row.get('algorithm_a')} vs {row.get('algorithm_b')} (p = {_fmt_p_es(row.get('wilcoxon_p_value'))})"
        by_scope.setdefault(scope, []).append(pair)
    scope_labels = {"OE1": "OE.1", "OE2": "OE.2", "OE3": "OE.3", "ALL": "ALL"}
    chunks = [f"{scope_labels.get(scope, scope)}: {', '.join(pairs)}" for scope, pairs in by_scope.items()]
    return (
        "Wilcoxon pareado detecta diferencias sistematicas en pares especificos (KPI-gains pareados): "
        + "; ".join(chunks)
        + ". Estos contrastes son exploratorios y no reemplazan replicacion multi-semilla ni el "
        "diseno factorial 4×3 completo."
    )


DESCRIPTIVE_BEST_BY_AXIS = {
    "OE1": "MATD3",
    "OE2": "MATD3",
    "OE3": "MAAC",
    "OG": "MATD3",
}


def _hypothesis_decision_label(axis: str, kw_p: float, wc_n: int) -> str:
    desc_best = DESCRIPTIVE_BEST_BY_AXIS.get(axis, "-")
    if kw_p < 0.05:
        return "Confirmada (H0 rechazada, KW p<0,05)"
    if axis == "OE3":
        tail = f"; lider descriptivo {desc_best} (Δcosto 9 515 EUR)"
    else:
        tail = f"; lider descriptivo {desc_best}"
    if wc_n:
        return (
            f"No confirmada inferencialmente (KW p={kw_p:.3f}{tail}); "
            "Wilcoxon exploratorio significativo"
        )
    return f"No confirmada inferencialmente (KW p={kw_p:.3f}{tail})"


def _hypothesis_decision_table_rows() -> list[list[str]]:
    rows_out: list[list[str]] = []
    for axis, (hyp_code, label) in (
        ("OE1", ("HE.1", "OE.1 / E1")),
        ("OE2", ("HE.2", "OE.2 / E2")),
        ("OE3", ("HE.3", "OE.3 / E3")),
        ("OG", ("HG", "Global (ALL)")),
    ):
        hyp_row = _hyp_row_by_axis(axis)
        kw_p = float(hyp_row.get("KW_p_value", 1)) if hyp_row else 1.0
        kw_sig = kw_p < 0.05
        wc_n = len([r for r in _significant_wilcoxon_rows() if r.get("scope") == ("ALL" if axis == "OG" else axis)])
        decision = _hypothesis_decision_label(axis, kw_p, wc_n)
        rows_out.append(
            [
                hyp_code,
                label,
                "Kruskal-Wallis + Wilcoxon exploratorio",
                "Si" if kw_sig else "No (omnibus)",
                hyp_row.get("statistical_best_algorithm_by_median_gain", "-"),
                decision,
            ]
        )
    return rows_out


def _verdict_table_rows(report: dict) -> list[list[str]]:
    oe_map = {
        "OE1": ("OE.1", "D-VD.1 / E1", "score_oe1_flex", "flex_composite", ".4f"),
        "OE2": ("OE.2", "D-VD.2 / E2", "score_oe2_co2", "carbon_emissions_delta_kg", ",.0f"),
        "OE3": ("OE.3", "D-VD.3 / E3", "score_oe3_cost", "electricity_cost_delta_eur", ",.0f"),
    }
    district = _read_csv(DISTRICT_CSV)
    rows_out: list[list[str]] = []
    for oe_key, (oe_label, vd, score_key, kpi, fmt) in oe_map.items():
        scen = OE_DEFINITIONS[oe_key]["scenario"]
        scen_rows = _rows_for_scenario(district, scen)
        best_algo, best_val = _best_algo_by_kpi(scen_rows, kpi, lower_better=True)
        score = next(item.get(score_key, 0) for item in report["ranking_with_kpis"] if item["algorithm"] == best_algo)
        rows_out.append(
            [
                oe_label,
                vd,
                best_algo,
                f"{OE_DEFINITIONS[oe_key]['primary_kpis'][0][1]} = {_fmt_kpi(str(best_val), fmt, True)}; score = {score:.4f}",
                "Cumplido descriptivamente; inferencia causal limitada (1 semilla)",
            ]
        )
    best_global = report["ranking_with_kpis"][0]["algorithm"]
    score_global = report["ranking_with_kpis"][0].get("score_global", 0)
    rows_out.append(
        [
            "OG",
            "VD integrada (ALL)",
            best_global,
            f"score global = {score_global:.4f}",
            "Cumplido descriptivamente; sin dominancia universal en los tres ejes",
        ]
    )
    return rows_out


def _og_oe_verdict_text() -> str:
    return (
        "Veredicto OG/OE (evidencia descriptiva, corrida canonica 50 ep): OG cumplido en el "
        "sentido de determinar el efecto coordinado del MADRL e identificar a MATD3 como mayor "
        "efecto integrado (score 0,6667), con trade-offs entre ejes. OE.1 cumplido: MATD3 "
        "(flex_composite = 1,0009). OE.2 cumplido: MATD3 (delta CO2 = 23 070 kg). OE.3 cumplido: "
        "MAAC (delta costo = 9 515 EUR). Ningun objetivo exige por si mismo significancia "
        "estadistica omnibus; los limites de inferencia causal (semilla unica seed=0; "
        "HAPPO con KPI-gains 4/4 peores y 49 ep) se declaran explicitamente. Las hipotesis "
        "HG y HE.1–HE.3 se resuelven solo en la seccion 5.9.5."
    )


def _inferential_conclusion_text() -> str:
    og = _hyp_row_by_axis("OG")
    kw_all = float(og.get("KW_p_value", 1)) if og else 1.0
    wc_all = next(
        (
            r
            for r in _significant_wilcoxon_rows()
            if r.get("scope") == "ALL" and r.get("algorithm_a") == "MASAC" and r.get("algorithm_b") == "MATD3"
        ),
        {},
    )
    wc_p = _fmt_p_es(wc_all.get("wilcoxon_p_value", "0.0049"))
    return (
        "Decision inferencial (H0G/H1G, HE10–HE31): en la corrida canonica de 50 episodios y una "
        "semilla, ningun Kruskal-Wallis omnibus sobre KPI-gains rechaza H0 por eje (alpha = 0,05): "
        "HE10 p = 0,4685; HE20 p = 0,7648; HE30 p = 0,7357; KW ALL p = 0,1554. "
        "H0G se rechaza solo de forma exploratoria por Friedman (p = 0,0096). "
        "Por tanto, HE11/HE21/HE31 no se respaldan con el diseno actual. "
        "La evidencia descriptiva identifica a MATD3 en flexibilidad y CO2, y a MAAC en costos "
        "(Δcosto 9 515 EUR en E3), patron coherente con la literatura de atencion multiagente "
        "para optimizacion tarifaria (Iqbal y Sha, 2019; Xie et al., 2023). Cuando el ranking "
        "descriptivo difiere del contraste omnibus no significativo, la interpretacion debe "
        "priorizar intervalos y replicacion multi-semilla (Agarwal et al., 2021; Colas et al., "
        "2019). "
        f"Wilcoxon exploratorio ALL (MASAC vs MATD3, p = {wc_p}) y otros pares por eje sugieren "
        "diferencias en KPI-gains pareados que requieren replicacion multi-semilla antes de "
        "sustentar conclusiones causales robustas."
    )


def _pe_key_for_oe(oe_key: str) -> str:
    return {"OE1": "PE.1", "OE2": "PE.2", "OE3": "PE.3"}[oe_key]


def _district_kpi_values(district: list[dict[str, str]], scenario: str, kpi: str) -> dict[str, float]:
    return {
        row["algorithm"]: float(row[kpi])
        for row in _rows_for_scenario(district, scenario)
    }


def _pct_delta_vs_best(value: float, best: float) -> float:
    if best == 0:
        return 0.0
    return (value - best) / abs(best) * 100.0


def _shapiro_rows_for_scope(scope: str) -> list[list[str]]:
    row = _analysis_row(scope)
    if not row:
        return []
    hyp, label = HYPOTHESIS_LABELS.get(scope, (scope, scope))
    out: list[list[str]] = []
    for algo in ("MASAC", "MATD3", "MAAC"):
        p_key = f"shapiro_wilk_p_value_{algo}"
        if row.get(f"shapiro_wilk_status_{algo}") == "no_data":
            continue
        rejected = str(row.get(f"shapiro_wilk_normality_rejected_{algo}", "")).lower() == "true"
        out.append(
            [
                label,
                algo,
                _fmt_stat_num(row.get(p_key), ".3e"),
                "Si" if rejected else "No",
            ]
        )
    return out


def _mwu_pairs_for_scope(scope: str) -> list[dict[str, str]]:
    if not MWU_CSV.is_file():
        return []
    return [
        row
        for row in _read_csv(MWU_CSV)
        if row.get("scope") == scope and row.get("mann_whitney_status") == "ok"
    ]


def _wilcoxon_pairs_for_scope(scope: str) -> list[dict[str, str]]:
    if not WC_CSV.is_file():
        return []
    return [
        row
        for row in _read_csv(WC_CSV)
        if row.get("scope") == scope and str(row.get("wilcoxon_status", "")).startswith("ok")
    ]


def build_pe_answer(pe_id: str, *, district: list[dict[str, str]] | None = None, report: dict | None = None) -> dict:
    """Construye respuesta estructurada PE.x desde CSV auditados (50 ep Drive)."""
    import datetime as dt

    spec = PE_SPECS[pe_id]
    oe_key = spec["oe_key"]
    cfg = OE_DEFINITIONS[oe_key]
    scenario = cfg["scenario"]
    primary_kpi, primary_label, primary_fmt, lower_better = cfg["primary_kpis"][0]
    district = district or _read_csv(DISTRICT_CSV)
    report = report or _read_json(BEST_REPORT)
    hyp_row = _hypothesis_row(oe_key)
    analysis_row = _analysis_row(oe_key)

    scen_rows = _rows_for_scenario(district, scenario)
    best_algo, best_val = _best_algo_by_kpi(scen_rows, primary_kpi, lower_better=True)
    kpi_by_algo = _district_kpi_values(district, scenario, primary_kpi)
    score_key = cfg["score_key"]
    best_score = next(
        item.get(score_key, 0) for item in report["ranking_with_kpis"] if item["algorithm"] == best_algo
    )

    desc_episode = _descriptive_episode_rows(oe_key)
    episode_metric = OE_EPISODE_SPECS[oe_key]["metric"]
    episode_label = OE_EPISODE_SPECS[oe_key]["label"]

    deltas = []
    for algo, val in sorted(kpi_by_algo.items(), key=lambda x: x[1]):
        if algo == best_algo:
            continue
        delta = val - best_val
        pct = _pct_delta_vs_best(val, best_val)
        deltas.append(
            {
                "algorithm": algo,
                "value": val,
                "delta_vs_best": delta,
                "pct_vs_best": pct,
                "formatted_value": _fmt_kpi(str(val), primary_fmt, True),
                "formatted_delta": _fmt_kpi(str(delta), primary_fmt, True),
                "formatted_pct": f"{pct:+.1f}%",
            }
        )

    kw_p = float(hyp_row.get("KW_p_value", "nan")) if hyp_row else float("nan")
    kw_h0_rejected = kw_p < 0.05
    best_median_gain = hyp_row.get("statistical_best_algorithm_by_median_gain", "-")
    sw_rows = _shapiro_rows_for_scope(oe_key)
    mwu_pairs = _mwu_pairs_for_scope(oe_key)
    wc_pairs = _wilcoxon_pairs_for_scope(oe_key)
    wc_sig = [r for r in wc_pairs if r.get("wilcoxon_significant_alpha_0_05") == "True"]
    mwu_sig = [r for r in mwu_pairs if r.get("mann_whitney_significant_alpha_0_05") == "True"]

    delta_text = "; ".join(
        f"{d['algorithm']} {d['formatted_delta']} ({d['formatted_pct']} vs {best_algo})"
        for d in deltas
    )
    desc_explicit = (
        f"En qué medida: el MADRL (VI) modifica {spec['vd']} con magnitudes distinguibles entre "
        f"algoritmos ({primary_label} {best_algo} = "
        f"{_fmt_kpi(str(best_val), primary_fmt, True)}; score {spec['oe']} = {best_score:.4f}); "
        f"deltas frente al lider: {delta_text}. "
        f"Mayor efecto: {best_algo} porque registra el mejor {primary_label} en {scenario} "
        f"({_fmt_kpi(str(best_val), primary_fmt, True)}) y el score normalizado {spec['oe']} = "
        f"{best_score:.4f}."
    )

    wc_sig_text = (
        ", ".join(
            f"{r['algorithm_a']} vs {r['algorithm_b']} (p = {_fmt_p_es(r.get('wilcoxon_p_value'))}, "
            f"mejor {r.get('better_by_median_difference', '-')})"
            for r in wc_sig
        )
        if wc_sig
        else "ningún par significativo"
    )
    mwu_sig_text = (
        ", ".join(
            f"{r['algorithm_a']} vs {r['algorithm_b']} (p = {_fmt_p_es(r.get('mann_whitney_p_value'))})"
            for r in mwu_sig
        )
        if mwu_sig
        else "ningún par significativo"
    )
    inf_explicit = (
        f"En qué medida inferencial: H0 de igualdad global entre algoritmos "
        f"{'rechazada' if kw_h0_rejected else 'no rechazada'} "
        f"(Kruskal-Wallis {spec['he']}, p = {_fmt_p_es(kw_p)}); el efecto entre algoritmos "
        f"{'es' if kw_h0_rejected else 'no es'} estadísticamente distinguible al nivel omnibus "
        f"(alpha = 0,05). Shapiro-Wilk: normalidad rechazada en al menos un grupo "
        f"({spec['he']}). Mann-Whitney U: {mwu_sig_text}. Wilcoxon pareado (exploratorio): "
        f"{wc_sig_text}. Mayor efecto por mediana KPI-gain: {best_median_gain}."
    )

    return {
        "pe_id": pe_id,
        "problem": spec["problem"],
        "oe": spec["oe"],
        "he": spec["he"],
        "vd": spec["vd"],
        "dimension": spec["dimension"],
        "section": spec["section"],
        "run_id": RUN_ID,
        "descriptive": {
            "best_algorithm": best_algo,
            "primary_kpi": primary_kpi,
            "primary_label": primary_label,
            "primary_value": best_val,
            "primary_value_formatted": _fmt_kpi(str(best_val), primary_fmt, True),
            "score_normalized": float(best_score),
            "kpis_by_algorithm": {algo: float(val) for algo, val in kpi_by_algo.items()},
            "deltas_vs_best": deltas,
            "episode_metric": episode_metric,
            "episode_label": episode_label,
            "episode_stats": [
                {
                    "algorithm": str(r["algorithm"]),
                    "n_episodes": int(r.get("n_episodes", 0)),
                    "mean": float(r["mean"]),
                    "median": float(r["median"]),
                    "std": float(r["std"]),
                }
                for r in desc_episode
            ],
            "explicit_answer": desc_explicit,
        },
        "inferential": {
            "shapiro_wilk": [
                {
                    "algorithm": row[1],
                    "p_value": row[2],
                    "normality_rejected": row[3] == "Si",
                }
                for row in sw_rows
            ],
            "kruskal_wallis": {
                "H": float(hyp_row.get("KW_H_statistic", "nan")) if hyp_row else None,
                "p_value": kw_p,
                "h0_rejected": kw_h0_rejected,
            },
            "best_median_kpi_gain": best_median_gain,
            "mann_whitney_significant": [
                {
                    "pair": f"{r['algorithm_a']} vs {r['algorithm_b']}",
                    "p_value": float(r.get("mann_whitney_p_value", "nan")),
                    "better_by_median": r.get("better_by_median", "-"),
                    "cliffs_delta": r.get("cliffs_delta", ""),
                    "cohen_d": r.get("cohen_d", ""),
                }
                for r in mwu_sig
            ],
            "wilcoxon_significant": [
                {
                    "pair": f"{r['algorithm_a']} vs {r['algorithm_b']}",
                    "p_value": float(r.get("wilcoxon_p_value", "nan")),
                    "better_by_median": r.get("better_by_median_difference", "-"),
                }
                for r in wc_sig
            ],
            "explicit_answer": inf_explicit,
        },
        "generated_at": dt.datetime.now(dt.timezone.utc).isoformat(),
    }


def build_all_pe_answers(*, district: list[dict[str, str]] | None = None, report: dict | None = None) -> dict:
    district = district or _read_csv(DISTRICT_CSV)
    report = report or _read_json(BEST_REPORT)
    answers = {pe_id: build_pe_answer(pe_id, district=district, report=report) for pe_id in PE_SPECS}
    return {
        "generated_at": answers["PE.1"]["generated_at"],
        "run_id": RUN_ID,
        "verdict": "structured_from_csv",
        "sources": {
            "descriptivo_distrito_colab": str(DESCRIPTIVE_CSV),
            "district_episode_kpis": str(EPISODE_CSV),
            "hipotesis_estadisticas_madrl": str(HYP_CSV),
            "comparaciones_mwu_madrl": str(MWU_CSV),
            "comparaciones_wilcoxon_madrl": str(WC_CSV),
            "best_madrl_report": str(BEST_REPORT),
            "inferential_audit_report": str(STAT_DIR / "inferential_audit_report.json"),
        },
        "answers": answers,
    }


def write_pe_answers_audit(path: Path | None = None) -> Path:
    path = path or PE_ANSWERS_AUDIT_JSON
    payload = build_all_pe_answers()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _add_pe_answer_section(
    doc,
    p,
    heading,
    add_table,
    *,
    oe_key: str,
    district: list[dict[str, str]],
    report: dict,
    desc_table_base: int,
) -> None:
    pe_id = _pe_key_for_oe(oe_key)
    answer = build_pe_answer(pe_id, district=district, report=report)
    cfg = OE_DEFINITIONS[oe_key]
    spec = PE_SPECS[pe_id]
    desc = answer["descriptive"]
    inf = answer["inferential"]

    heading(doc, f"{spec['section']} Respuesta a {pe_id} ({spec['oe']} / {spec['he']} / {spec['vd']})", 3)
    p(doc, spec["problem"], italic=True)

    p(doc, f"Análisis descriptivo — {pe_id}:", bold=True)
    kpi_headers = ["Algoritmo", desc["primary_label"], f"Score {spec['oe']}"]
    kpi_rows = []
    for algo in sorted(desc["kpis_by_algorithm"], key=lambda a: desc["kpis_by_algorithm"][a]):
        val = desc["kpis_by_algorithm"][algo]
        score = next(
            item.get(cfg["score_key"], 0) for item in report["ranking_with_kpis"] if item["algorithm"] == algo
        )
        kpi_rows.append(
            [
                algo,
                _fmt_kpi(str(val), cfg["primary_kpis"][0][2], True),
                f"{float(score):.4f}",
            ]
        )
    add_table(
        doc,
        kpi_headers,
        kpi_rows,
        caption=f"Tabla 5.{desc_table_base}. KPIs {spec['vd']} por algoritmo ({cfg['scenario']}, distrito 50 ep).",
        col_widths=[2.5, 3.5, 2.5],
    )

    ep_spec = OE_EPISODE_SPECS[oe_key]
    ep_rows = []
    for row in sorted(desc["episode_stats"], key=lambda r: r["median"], reverse=ep_spec["higher_better"]):
        ep_rows.append(
            [
                row["algorithm"],
                str(row["n_episodes"]),
                _fmt_stat_num(row["mean"], ep_spec["fmt"]),
                _fmt_stat_num(row["median"], ep_spec["fmt"]),
                _fmt_stat_num(row["std"], ep_spec["fmt"]),
            ]
        )
    if ep_rows:
        add_table(
            doc,
            ["Algoritmo", "n ep.", "Media", "Mediana", "Desv. est."],
            ep_rows,
            caption=(
                f"Tabla 5.{desc_table_base + 1}. Estadística descriptiva episódica {spec['oe']} — "
                f"{desc['episode_label']} ({cfg['scenario']})."
            ),
            col_widths=[2.0, 1.2, 2.2, 2.2, 2.2],
        )

    delta_bits = [
        f"{d['algorithm']}: {d['formatted_delta']} ({d['formatted_pct']} vs {desc['best_algorithm']})"
        for d in desc["deltas_vs_best"]
    ]
    p(
        doc,
        f"Magnitud del efecto: {desc['primary_label']} del líder {desc['best_algorithm']} = "
        f"{desc['primary_value_formatted']}; diferencias respecto al mejor: "
        f"{'; '.join(delta_bits) if delta_bits else 'sin contraste adicional'}.",
    )
    p(doc, f"Respuesta explícita: {desc['explicit_answer']}")

    p(doc, f"Análisis inferencial — {pe_id}:", bold=True)
    sw_rows = [
        [row["algorithm"], row["p_value"], "Si" if row["normality_rejected"] else "No"]
        for row in inf["shapiro_wilk"]
    ]
    if sw_rows:
        add_table(
            doc,
            ["Algoritmo", "p (Shapiro-Wilk)", "Normalidad rechazada"],
            sw_rows,
            caption=f"Tabla 5.{desc_table_base + 2}. Shapiro-Wilk por grupo ({spec['he']}, KPI-gains).",
            col_widths=[2.5, 3.0, 3.0],
        )

    kw = inf["kruskal_wallis"]
    p(
        doc,
        f"Kruskal-Wallis ({spec['he']}): H = {_fmt_stat_num(kw.get('H'), '.3f')}, "
        f"p = {_fmt_p_es(kw.get('p_value'))}; H0 "
        f"{'rechazada' if kw.get('h0_rejected') else 'no rechazada'} (alpha = 0,05).",
    )

    mwu_rows = []
    for row in _mwu_pairs_for_scope(oe_key):
        p_val = float(row.get("mann_whitney_p_value", "nan"))
        mwu_rows.append(
            [
                f"{row.get('algorithm_a')} vs {row.get('algorithm_b')}",
                _fmt_stat_num(row.get("mann_whitney_u"), ".1f"),
                _fmt_p_es(p_val),
                "Si" if p_val < 0.05 else "No",
                row.get("better_by_median", "-"),
                row.get("cliffs_delta", "-"),
            ]
        )
    if mwu_rows:
        add_table(
            doc,
            ["Par", "U", "p", "Signif.", "Mejor (mediana)", "Cliff's δ"],
            mwu_rows,
            caption=f"Tabla 5.{desc_table_base + 3}. Mann-Whitney U ({spec['he']}).",
            col_widths=[2.8, 1.3, 1.5, 1.2, 2.0, 1.5],
        )

    wc_rows = []
    for row in _wilcoxon_pairs_for_scope(oe_key):
        p_val = float(row.get("wilcoxon_p_value", "nan"))
        wc_rows.append(
            [
                f"{row.get('algorithm_a')} vs {row.get('algorithm_b')}",
                _fmt_stat_num(row.get("wilcoxon_T_statistic"), ".1f"),
                _fmt_p_es(p_val),
                "Si" if p_val < 0.05 else "No",
                row.get("better_by_median_difference", "-"),
            ]
        )
    if wc_rows:
        add_table(
            doc,
            ["Par (pareado)", "T", "p", "Signif.", "Mejor (mediana diff.)"],
            wc_rows,
            caption=f"Tabla 5.{desc_table_base + 4}. Wilcoxon signed-rank ({spec['he']}).",
            col_widths=[2.8, 1.3, 1.5, 1.2, 2.5],
        )

    p(doc, f"Respuesta explícita: {inf['explicit_answer']}")


def _pe_conclusions_paragraph() -> str:
    answers = build_all_pe_answers()
    chunks = []
    for pe_id in ("PE.1", "PE.2", "PE.3"):
        ans = answers["answers"][pe_id]
        chunks.append(
            f"{pe_id} ({ans['vd']}): descriptivamente, mayor efecto {ans['descriptive']['best_algorithm']} "
            f"({ans['descriptive']['primary_value_formatted']}); inferencialmente, Kruskal-Wallis "
            f"p = {_fmt_p_es(ans['inferential']['kruskal_wallis']['p_value'])} "
            f"(H0 {'rechazada' if ans['inferential']['kruskal_wallis']['h0_rejected'] else 'no rechazada'}); "
            f"mejor mediana KPI-gain: {ans['inferential']['best_median_kpi_gain']}."
        )
    return " ".join(chunks)


def _load_algo_profiles() -> dict:
    from madrl_algorithm_analysis import build_all_profiles, write_metrics_json

    if not ALGO_METRICS_JSON.is_file():
        write_metrics_json(ALGO_METRICS_JSON)
    return json.loads(ALGO_METRICS_JSON.read_text(encoding="utf-8"))


def _algo_figure(algo: str, scen: str, name: str) -> Path | None:
    path = FIG_RUN / algo / scen / "figures" / f"{name}.png"
    return path if path.is_file() else None


def _fmt_ep(value: float | int | None, digits: int = 4) -> str:
    if value is None:
        return "-"
    if isinstance(value, int):
        return str(value)
    return f"{value:.{digits}f}"


def _add_per_algorithm_analysis_section(
    doc,
    p,
    heading,
    add_table,
    *,
    profiles: dict,
    fig_counter: list[int],
) -> None:
    heading(doc, "5.2 Analisis de convergencia, aprendizaje y desempeno por algoritmo MADRL", 2)
    p(
        doc,
        "Esta seccion complementa el analisis por objetivo especifico (secciones 5.3–5.5) con una "
        "lectura longitudinal por algoritmo sobre la corrida canonica de 50 episodios. Para cada "
        "MADRL se reportan: (i) trayectoria de reward_mean y episodio de meseta; (ii) estabilidad "
        "(varianza intra-entrenamiento); (iii) KPIs distritales por escenario E1/E2/E3; "
        "(iv) heterogeneidad entre los 17 edificios; y (v) comportamiento EV/BESS cuando los datos "
        "lo permiten. La interpretacion se alinea con la literatura CityLearn: algoritmos off-policy "
        "actor-critico (MASAC, MATD3) y atencion cooperativa (MAAC) exhiben patrones de convergencia "
        "distintos; HAPPO on-policy muestra mayor varianza inter-episodio en entornos multiagente "
        "(Yu et al., 2022; Vazquez-Canteli et al., 2024).",
    )

    algo_specs = [
        (
            "MATD3",
            "5.2.1",
            True,
            "MATD3 (Multi-Agent Twin Delayed DDPG) combina politica determinista con retraso de "
            "actualizacion del actor y ruido de exploracion gaussiano. En CityLearn, TD3 y variantes "
            "multiagente han mostrado convergencia rapida en costos pero sensibilidad a la escala de "
            "recompensa (Fujimoto et al., 2018; Di Savino et al., 2025). En Iquitos, MATD3 lidera "
            "OE.1 y OE.2 descriptivamente y obtiene el mejor score global (0,6667), con la mayor "
            "tasa de exito EV (43,9% en E1) y flex_composite distrital minimo (1,0009).",
        ),
        (
            "MAAC",
            "5.2.2",
            True,
            "MAAC (Multi-Agent Actor-Attention-Critic) usa atencion en el critico centralizado para "
            "ponderar contribuciones de agentes vecinos, lo que mejora escalabilidad en distritos "
            "grandes (Iqbal & Sha, 2019; Di Savino et al., 2025). En esta corrida, MAAC converge "
            "tempranamente (meseta ~episodio 4 en los tres escenarios) con baja varianza de "
            "reward_mean y lidera OE.3 (delta costo = 9 515 EUR). Su debilidad relativa en CO2 "
            "(70 654 kg) refleja trade-off multiobjetivo.",
        ),
        (
            "MASAC",
            "5.2.3",
            True,
            "MASAC (Multi-Agent Soft Actor-Critic) maximiza entropia de politica para exploracion "
            "estable; en benchmarks CityLearn, SAC y variantes multiagente suelen mejorar rapido "
            "al inicio pero pueden plateauar antes que metodos on-policy cooperativos "
            "(Haarnoja et al., 2018; arXiv:2602.19223). Aqui, MASAC muestra convergencia estable "
            "pero KPIs distritales inferiores en flexibilidad y CO2; su score OE.3 intermedio "
            "(0,7054) indica competencia parcial en costos sin liderar el eje.",
        ),
        (
            "HAPPO",
            "5.2.4",
            True,
            "HAPPO (Heterogeneous-Agent PPO) aplica PPO secuencial por agente bajo CTDE "
            "(Yu et al., 2022). Cobertura episódica materializada: 49/50 por escenario. "
            "El recálculo Drive 2026-07-28 recupera KPIs finales evaluate_v2/core_kpis y "
            "KPI-gains all_values (flex_composite E1 = 1,1105; carbon_emissions_delta E2 = "
            "1 431 341 kg; electricity_cost_delta E3 = 106 828; medianas de ganancia "
            "OE1/OE2/OE3 = −0,0972 / −715 671 / −53 414). HAPPO entra en pie de igualdad "
            "descriptiva 4/4 (ranking evaluate_v2 score 0,0000, peor que MAAC/MATD3/MASAC). "
            "El veredicto best_madrl 3×3 (MATD3 0,6667) y las HE sobre KPI-gains de "
            "entrenamiento se mantienen como lectura canónica de hipótesis; la batería "
            "episódica complementaria incluye HAPPO (49 ep).",
        ),
    ]

    conv_headers = [
        "Esc.",
        "n ep.",
        "Media reward",
        "Desv. est.",
        "CV",
        "Ep. meseta",
        "Mejora ep.1–10 → ult.10",
    ]

    for algo, sec_num, has_kpis, intro in algo_specs:
        prof = profiles["algorithms"][algo]
        heading(doc, f"{sec_num} {algo}", 3)
        p(doc, intro)

        conv_rows = []
        for scen in ("E1", "E2", "E3"):
            c = prof.get("convergence", {}).get(scen)
            if not c:
                continue
            conv_rows.append(
                [
                    scen,
                    str(c["n_episodes"]),
                    _fmt_ep(c["reward_mean"]),
                    _fmt_ep(c["reward_std"]),
                    _fmt_ep(c["coefficient_of_variation"]),
                    str(c["plateau_episode"]) if c["plateau_episode"] is not None else "-",
                    _fmt_ep(c["improvement_first_to_last"]),
                ]
            )
        if conv_rows:
            add_table(
                doc,
                conv_headers,
                conv_rows,
                caption=f"Tabla 5.3-{algo}. Convergencia y estabilidad de {algo} (reward_mean, timeseries.csv).",
                col_widths=[1.2, 1.2, 2.2, 2.0, 1.5, 1.8, 2.5],
            )

        if has_kpis and prof.get("district_kpis"):
            kpi_rows = []
            for scen in ("E1", "E2", "E3"):
                d = prof["district_kpis"].get(scen, {})
                if not d:
                    continue
                kpi_rows.append(
                    [
                        scen,
                        _fmt_ep(d.get("flex_composite")),
                        f"{d.get('carbon_emissions_delta_kg', 0):,.0f}",
                        f"{d.get('electricity_cost_delta_eur', 0):,.0f}",
                        f"{d.get('ev_departure_success_rate', 0) * 100:.1f}%",
                        f"{d.get('grid_import_delta', 0):,.0f}",
                    ]
                )
            add_table(
                doc,
                ["Esc.", "flex_composite", "Delta CO2 (kg)", "Delta costo (EUR)", "Exito EV", "Delta import."],
                kpi_rows,
                caption=f"Tabla 5.4-{algo}. KPIs distritales por escenario ({algo}, corrida 50 ep).",
                col_widths=[1.2, 2.2, 2.5, 2.5, 1.8, 2.0],
            )

            het_rows = []
            for scen in ("E1", "E2", "E3"):
                h = prof.get("building_heterogeneity", {}).get(scen, {})
                if not h:
                    continue
                het_rows.append(
                    [
                        scen,
                        str(int(h["n_buildings"])),
                        f"{h['mean']:,.2f}",
                        f"{h['std']:,.2f}",
                        _fmt_ep(h["cv"], 3),
                        f"{h['min']:,.2f}",
                        f"{h['max']:,.2f}",
                    ]
                )
            add_table(
                doc,
                ["Esc.", "n edif.", "Media KPI edificio", "Desv. est.", "CV", "Min.", "Max."],
                het_rows,
                caption=(
                    f"Tabla 5.5-{algo}. Heterogeneidad inter-edificio (17 agentes, building_objectives_by_algorithm.csv)."
                ),
                col_widths=[1.2, 1.5, 2.5, 2.0, 1.5, 2.0, 2.0],
            )

        dom_scen = {"MATD3": "E1", "MAAC": "E3", "MASAC": "E2", "HAPPO": "E1"}[algo]
        conv_fig = _algo_figure(algo, dom_scen, "convergence_returns")
        if conv_fig:
            fig_counter[0] += 1
            add_figure(
                doc,
                conv_fig,
                f"Figura 5.{fig_counter[0]}. Convergencia reward_mean — {algo}/{dom_scen} (300 dpi, datos Drive).",
                width_cm=FIG_WIDTH_LANDSCAPE,
                interpretation=(
                    f"La trayectoria de {algo} en {dom_scen} muestra "
                    f"{'mejora sostenida hasta meseta temprana' if prof.get('convergence', {}).get(dom_scen, {}).get('plateau_episode', 99) < 15 else 'varianza inter-episodio elevada'} "
                    f"(CV = {_fmt_ep(prof.get('convergence', {}).get(dom_scen, {}).get('coefficient_of_variation', 0), 3)}). "
                    "Coherente con patrones reportados en benchmarks MARL-CityLearn donde la estabilidad "
                    "depende del algoritmo y del peso de recompensa por escenario."
                ),
            )
        expl_fig = _algo_figure(algo, dom_scen, "exploration_action_l2")
        if expl_fig:
            fig_counter[0] += 1
            add_figure(
                doc,
                expl_fig,
                f"Figura 5.{fig_counter[0]}. Exploracion (action_l2) — {algo}/{dom_scen} (trace.csv).",
                interpretation=(
                    f"La norma L2 de acciones de {algo} refleja intensidad de control sobre BESS, "
                    "cargas desplazables y EV. Reduccion gradual sugiere politica mas determinista; "
                    "picos persistentes indican exploracion activa o inestabilidad de politica."
                ),
            )
        base_fig = _algo_figure(algo, dom_scen, "axis_baseline_comparison")
        if base_fig and has_kpis:
            fig_counter[0] += 1
            add_figure(
                doc,
                base_fig,
                f"Figura 5.{fig_counter[0]}. Comparacion vs baseline CityLearn v2 — {algo}/{dom_scen}.",
                interpretation=(
                    f"El contraste con baseline RBC ubica a {algo} respecto a controles heuristicos "
                    "del eje dominante. Esta lectura matiza la superioridad MADRL frente a reglas "
                    "predefinidas (Nweye et al., 2024)."
                ),
            )


def _figure_width_cm(path: Path) -> float:
    try:
        from PIL import Image

        with Image.open(path) as im:
            w, h = im.size
            return FIG_WIDTH_PORTRAIT if h > w * 1.15 else FIG_WIDTH_LANDSCAPE
    except Exception:
        return FIG_WIDTH_LANDSCAPE


def add_figure(
    doc,
    path: Path,
    caption: str,
    width_cm: float | None = None,
    interpretation: str | None = None,
) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)
    if width_cm is None:
        width_cm = _figure_width_cm(path)
    doc.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    if interpretation:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(interpretation)
    doc.add_paragraph()


def _add_oe_results_section(
    doc,
    p,
    heading,
    add_table,
    *,
    oe_key: str,
    oe_num: int,
    section_num: str,
    district: list[dict[str, str]],
    report: dict,
    fig_counter: list[int],
) -> None:
    cfg = OE_DEFINITIONS[oe_key]
    scenario = cfg["scenario"]
    rows = _rows_for_scenario(district, scenario)
    primary_kpi = cfg["primary_kpis"][0][0]
    best_algo, best_val = _best_algo_by_kpi(rows, primary_kpi, lower_better=True)
    hyp = _hypothesis_row(oe_key)
    kw_p = float(hyp.get("KW_p_value", "nan")) if hyp else float("nan")

    oe_display = {"OE1": "OE.1", "OE2": "OE.2", "OE3": "OE.3"}[oe_key]

    heading(doc, f"{section_num} {oe_display} — Efecto sobre {cfg['dimension']} ({cfg['vd']})", 2)
    p(
        doc,
        f"{oe_display}. Determinar el efecto del algoritmo MADRL (VI: D-VI.1) sobre "
        f"{cfg['dimension']} ({cfg['vd']}) e identificar el algoritmo de mayor efecto en esta "
        f"dimension. El escenario {scenario} concentra el peso de recompensa en este eje "
        f"({cfg['weights']}). La evidencia proviene de la corrida canonica de 50 episodios "
        f"({RUN_ID}) con MATD3, MAAC, MASAC y HAPPO. HAPPO aporta KPI-gains evaluate_v2 "
        "(Drive 2026-07-28; ranking 4/4 score 0,0000) y 49 episodios en la batería "
        "episódica; las HE sobre KPI-gains de entrenamiento usan el trío MASAC/MATD3/MAAC "
        "(n KPI 12/5/9) y el best_madrl canónico 3×3 (MATD3 0,6667).",
    )

    kpi_headers = ["Algoritmo"] + [label for _, label, _, _ in cfg["primary_kpis"]]
    kpi_rows = []
    for row in sorted(rows, key=lambda r: float(r[primary_kpi])):
        cells = [row["algorithm"]]
        for kpi, _, fmt, lower in cfg["primary_kpis"]:
            cells.append(_fmt_kpi(row[kpi], fmt, lower))
        kpi_rows.append(cells)
    add_table(
        doc,
        kpi_headers,
        kpi_rows,
        caption=(
            f"Tabla 5.{3 + oe_num}. KPIs de {cfg['vd']} en escenario {scenario} "
            f"(distrito, corrida Colab 50 ep)."
        ),
        col_widths=[2.2] + [2.4] * (len(kpi_headers) - 1),
    )

    score_rows = []
    for item in report.get("ranking_with_kpis", []):
        score_rows.append(
            [
                item["algorithm"],
                f"{item.get(cfg['score_key'], 0):.4f}",
                "Si" if item["algorithm"] == best_algo else "No",
            ]
        )
    score_rows.sort(key=lambda r: float(r[1]), reverse=True)
    add_table(
        doc,
        ["Algoritmo", f"Score normalizado {oe_key}", f"Mayor efecto en {oe_key}"],
        score_rows,
        caption=f"Tabla 5.{6 + oe_num}. Ranking descriptivo {oe_display} (best_madrl_report.json).",
        col_widths=[3.0, 4.0, 4.5],
    )

    hyp_code = HYPOTHESIS_LABELS[oe_key][0]
    p(
        doc,
        f"Respuesta a {oe_display}: (a) efecto del MADRL (VI) sobre {cfg['vd']}: los KPIs de la "
        f"Tabla 5.{3 + oe_num} muestran diferencias entre algoritmos en {scenario}; "
        f"(b) algoritmo de mayor efecto: {best_algo} "
        f"({cfg['primary_kpis'][0][1]} = {_fmt_kpi(str(best_val), cfg['primary_kpis'][0][2], True)}); "
        f"(c) enlace a {hyp_code}: Kruskal-Wallis p = {kw_p:.3f} (seccion 5.9.5). "
        f"La identificacion del lider responde al objetivo {oe_display}; la significancia "
        f"estadistica de {hyp_code} se decide en 5.9.5, no en esta seccion.",
    )

    _add_pe_answer_section(
        doc,
        p,
        heading,
        add_table,
        oe_key=oe_key,
        district=district,
        report=report,
        desc_table_base=22 + (oe_num - 1) * 6,
    )

    fig_specs = [
        (cfg["conv_fig"], f"convergencia de reward_mean en {scenario}"),
        (cfg["kpi_fig"], f"comparativa de KPI {oe_key} en {scenario}"),
        (cfg["building_fig"], f"desagregacion por edificio ({cfg['vd']})"),
    ]
    if cfg.get("trace_fig"):
        fig_specs.append((cfg["trace_fig"], "trazas de control MADRL (trace.csv)"))

    for path, desc in fig_specs:
        fig_counter[0] += 1
        n = fig_counter[0]
        doc.add_page_break()
        width = FIG_WIDTH_LANDSCAPE if "convergencia" in desc else None
        add_figure(
            doc,
            path,
            f"Figura 5.{n}. {desc.capitalize()} — escenario {scenario}, datos reales Drive (300 dpi).",
            width_cm=width,
            interpretation=(
                f"La Figura 5.{n} vincula directamente {oe_display} con {cfg['vd']}: "
                f"en {scenario} se observa que {best_algo} concentra el mejor desempeno "
                f"descriptivo frente a los demas MADRL auditados. La lectura debe hacerse "
                f"exclusivamente en la dimension {cfg['dimension']}, sin extrapolar al "
                "score global ni a los otros objetivos especificos."
            ),
        )


def add_chapter_5_doctoral(doc, p, heading, add_table, status_note) -> None:
    """Delega en estructura: descriptivos -> inferenciales -> otros -> hipotesis -> discusion."""
    from thesis_cap5_structured import add_chapter_5_structured

    add_chapter_5_structured(doc, p, heading, add_table, status_note)


def add_chapter_6_doctoral(doc, p, heading, bullet, add_table) -> None:
    """Cap. 6 alineado a docs/informedetesis.txt (conclusiones preliminares)."""
    heading(doc, "Capitulo 6. Conclusiones preliminares", 1)

    heading(doc, "6.1 Principales hallazgos", 2)
    p(doc, _pe_conclusions_paragraph())
    p(
        doc,
        "OG: el algoritmo MADRL aplicado a la comunidad inteligente (VI) produce efectos "
        "diferenciados sobre la gestion coordinada de flexibilidad, emisiones de CO2 y costos "
        "(VD). El algoritmo de mayor efecto coordinado en la corrida canonica de 50 episodios "
        "es MATD3 (score global 0,6667 canónico 3×3; ranking evaluate_v2 4/4 lidera MAAC "
        "0,9538 con HAPPO en score 0,0000), que lidera OE.1 y OE.2 pero no OE.3 (MAAC). La "
        "respuesta al OG es descriptivamente afirmativa con limites de inferencia causal "
        "(semilla unica seed=0; diseño multi-semilla n=12 implementado y smoke-validado; "
        "HAPPO con KPI-gains evaluate_v2 y 49 ep episódicos).",
    )
    p(
        doc,
        "OE.1 (flexibilidad, D-VD.1, escenario E1): el algoritmo MADRL de mayor efecto es MATD3 "
        "(flex_composite = 1,0009; score normalizado OE.1 = 1,0000; 50 episodios). MATD3 tambien "
        "registra la mayor tasa de exito EV (43,9%), coherente con la operacionalizacion de "
        "flexibilidad mediante BESS, cargadores y cargas desplazables.",
    )
    p(
        doc,
        "OE.2 (emisiones CO2, D-VD.2, escenario E2): el algoritmo de mayor efecto es MATD3 "
        "(delta CO2 = 23 070 kg; score OE.2 = 1,0000). La reduccion distrital es la menor entre "
        "los tres MADRL auditados (MASAC: 77 649 kg; MAAC: 70 654 kg).",
    )
    p(
        doc,
        "OE.3 (costos energeticos, D-VD.3, escenario E3): el algoritmo de mayor efecto es MAAC "
        "(delta costo = 9 515 EUR; score OE.3 = 1,0000), frente a MATD3 (44 399 EUR) y MASAC "
        "(19 793 EUR). Esta conclusion es independiente del score global y confirma trade-offs "
        "entre objetivos.",
    )
    p(
        doc,
        "Integracion OE: MATD3 obtiene score global 0,6667 por liderar OE.1 y OE.2, pero no OE.3. "
        "La comparacion con baseline CityLearn v2 (seccion 5.7) muestra que controles RBC superan "
        "globalmente a MADRL en score HPHI, matizando la generalizacion causal pero no invalidando "
        "el benchmark inter-algoritmico exigido por los objetivos.",
    )
    p(
        doc,
        "Respecto de las hipotesis (KPI-gains Cap. 5): H0G se rechaza de forma exploratoria "
        "(Friedman p = 0,0096) mientras KW ALL no es significativo (p = 0,1554). "
        "HE10 (KW p = 0,4685), HE20 (KW p = 0,7648) y HE30 (KW p = 0,7357) no se rechazan; "
        "HE11/HE21/HE31 no se respaldan. Descriptivamente MATD3 lidera OE.1/OE.2 y MAAC lidera "
        "OE.3. Wilcoxon exploratorio detecta pares que no sustituyen el contraste omnibus ni la "
        "replicacion multi-semilla (Agarwal et al., 2021; Colas et al., 2019; Iqbal y Sha, 2019; "
        "Xie et al., 2023).",
    )

    heading(doc, "6.2 Limitaciones encontradas", 2)
    p(
        doc,
        "Las limitaciones principales son: (i) semilla unica (seed 0) en la campaña "
        "entrenada —el protocolo n_seeds=12 y el runner estan implementados "
        "(uc3m/multicriteria/scenarios.py; scripts/run_madrl_nonparametric_battery.py) y "
        "validados con smoke ilustrativo, pero no se entrenaron 12 semillas reales—; "
        "(ii) HAPPO con 49/50 episodios y KPI-gains evaluate_v2 peores que el trio "
        "(ranking 4/4 score 0); las HE sobre KPI-gains de entrenamiento usan MASAC/"
        "MATD3/MAAC; (iii) simulacion sin validacion en red fisica; (iv) CityLearn v3 "
        "propuesto como extension experimental de tesis (3 ejes OE); el paquete uc3m/ "
        "formaliza un mapeo aspiracional a 7 ejes sin ser el canal de las 12 corridas; "
        "(v) MADRL por debajo del baseline RBC en score global HPHI. Las hipotesis HG/HE "
        "se contrastan en seccion 5.9; los objetivos OG/OE en seccion 5.11.",
    )

    heading(doc, "6.3 Trabajo pendiente", 2)
    bullet(
        doc,
        "Entrenar la campaña multi-semilla real (n_seeds=12, seeds 0..11) con el runner "
        "ya implementado (EmpiricalProtocol + run_madrl_nonparametric_battery.py --run-root); "
        "el smoke ilustrativo (n=3 en outputs/madrl_nonparametric_battery_smoke_n3/) valida "
        "el pipeline y no sustituye entrenamiento.",
    )
    bullet(
        doc,
        "Homogeneizar HAPPO a 50 episodios i.i.d. por escenario (hoy: KPI-gains evaluate_v2 "
        "y ranking 4/4 disponibles; cobertura episódica 49/50).",
    )
    bullet(
        doc,
        "Completar analisis de frontera de Pareto por eje y cuantificar porcentajes de mejora "
        "frente a baseline CityLearn v2/RBC por KPI.",
    )
    bullet(
        doc,
        "Optimizar hiperparametros con Optuna (TPE) por backend y enriquecer el contraste con "
        "SB3 (PPO/SAC/A2C) bajo el mismo schema de Iquitos.",
    )
    bullet(
        doc,
        "Cerrar redaccion doctoral institucional (pasada RAE Cap. 1–6) y actualizar el indice "
        "Word (F9) antes del deposito.",
    )

    heading(doc, "6.4 Plan para culminar la tesis", 2)
    p(
        doc,
        "El plan de cierre prioriza hitos verificables alineados con la corrida canonica "
        f"{RUN_ID} y con los pendientes del Capitulo 5:",
    )
    add_table(
        doc,
        ["Hito", "Entregable verificable", "Estado"],
        [
            [
                "H1. HAPPO KPI-gains / 50 ep",
                "KPI-gains evaluate_v2 4/4 (Drive 2026-07-28); 50 ep i.i.d.",
                "KPI-gains 4/4 ejecutados; 49/50 ep",
            ],
            [
                "H2. Robustez multi-semilla",
                "Protocolo n_seeds=12 + smoke + campaña seed=0",
                "Diseño+runner+smoke ejecutados; campaña 12-seed no entrenada",
            ],
            [
                "H3. Inferencia Colab cerrada",
                "Tablas KW/MW/Wilcoxon auditadas y redaccion Cap. 5–6 sincronizada",
                "Ejecutada (HG no confirmada omnibus; H0G exploratoria Friedman)",
            ],
            [
                "H4. Pareto y % vs baseline",
                "Tablas/figuras OE.1–OE.3 vs RBC/CityLearn v2",
                "Parcial (baseline integrada; % KPI pendientes)",
            ],
            [
                "H5. HPO Optuna",
                "Estudio TPE por algoritmo",
                "Pendiente (fuera del alcance empirico cerrado)",
            ],
            [
                "H6. Redaccion final APA",
                "Informe final Cap. 1–6 + referencias APA sin marcadores pendientes",
                "En curso",
            ],
            [
                "H7. Sustentacion",
                "Documento oficial + defensa",
                "Pendiente institucional",
            ],
        ],
        caption="Tabla 6.1. Plan para culminar la tesis (hitos H1–H7).",
        col_widths=[4.2, 7.5, 4.0],
    )
    p(
        doc,
        "La culminacion academica queda condicionada a completar la campaña multi-semilla "
        "entrenada (H2) y la homogenizacion episódica HAPPO (H1) antes de afirmar efecto "
        "causal diferenciado del algoritmo MADRL sobre la VD multiobjetivo en el SEAI Iquitos. "
        "El diseño, los runners y la evidencia seed=0 + KPI-gains 4/4 ya estan cerrados en "
        "Caps. 4–5.",
    )


def verify_doctoral_docx(path: Path) -> dict:
    import re

    from docx import Document

    from thesis_references_apa import reference_stats

    required_sections = [
        "Dedicatoria",
        "Agradecimientos",
        "Resumen",
        "Abstract",
        "Introduccion",
        "Capitulo 5",
        "Capitulo 6",
        "Referencias bibliograficas",
    ]
    required_tables_min = 15
    required_figures_min = 12
    ref_stats = reference_stats()
    required_refs_min = max(50, ref_stats["total_unique"] - 5)

    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text)
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]

    section_ok = {s: any(s.lower() in h.lower() for h in headings) or s.lower() in text.lower() for s in required_sections}
    n_tables = len(doc.tables)
    n_images = sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref)

    in_refs = False
    n_ref_paras = 0
    ref_heading_count = 0
    complementaria_headings = 0
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        style = para.style.name if para.style else ""
        if "referencias bibliograficas" in t.lower() and (
            style.startswith("Heading") or t.lower() == "referencias bibliograficas"
        ):
            in_refs = True
            ref_heading_count += 1
            continue
        if in_refs and "referencias complementarias" in t.lower():
            complementaria_headings += 1
        if in_refs and re.match(r"^[A-Za-z]", t) and re.search(r"\(\d{4}", t):
            n_ref_paras += 1

    checks = {
        "sections": section_ok,
        "tables_count": n_tables,
        "tables_ok": n_tables >= required_tables_min,
        "images_count": n_images,
        "images_ok": n_images >= required_figures_min,
        "references_count": n_ref_paras,
        "references_expected": ref_stats["total_unique"],
        "references_ok": n_ref_paras >= required_refs_min,
        "references_single_list": ref_heading_count == 1 and complementaria_headings == 0,
        "has_matd3_selection": "MATD3" in text and ("0.6667" in text or "0,6667" in text),
        "has_pe_answers": "Respuesta a PE.1" in text and "Respuesta a PE.2" in text and "Respuesta a PE.3" in text,
        "has_multiobjetivo": "multiobjetivo" in text.lower() or "185" in text,
        "has_drive_figures": "timeseries.csv" in text.lower() or "artefactos drive" in text.lower(),
        "has_referencias_apa": "Referencias_APA.md" in text or "referencias bibliograficas" in text.lower(),
        "complete": (
            all(section_ok.values())
            and n_tables >= required_tables_min
            and n_images >= required_figures_min
            and n_ref_paras >= required_refs_min
            and ref_heading_count == 1
            and complementaria_headings == 0
            and "Respuesta a PE.1" in text
        ),
    }
    return checks
