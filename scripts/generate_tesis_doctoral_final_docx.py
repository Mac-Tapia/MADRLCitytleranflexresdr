#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Genera la version final de tesis doctoral en Word con resultados Colab y multiobjetivo."""

from __future__ import annotations

import sys
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "scripts"))

from generate_borrador_tesis_docx import (  # noqa: E402
    ACCENT,
    GREY,
    add_table,
    add_toc,
    build as build_borrador,
    bullet,
    heading,
    p,
    status_note,
    style_base,
)
from thesis_doctoral_sections import (  # noqa: E402
    add_chapter_5_doctoral,
    add_chapter_6_doctoral,
    add_dedicatoria_agradecimientos,
    add_resumen_doctoral,
    verify_doctoral_docx,
)

OUT_PATH = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos.docx"
SKILL_COPY = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_skill.docx"


def build_doctoral() -> Path:
    """Genera borrador base y compone version doctoral final."""
    import datetime as dt
    from copy import deepcopy

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm

    temp_borrador = REPO / "docs" / "_tmp_borrador_base.docx"
    import generate_borrador_tesis_docx as borrador_mod

    original_out = borrador_mod.OUT_PATH
    borrador_mod.OUT_PATH = temp_borrador
    try:
        build_borrador(max_chapter=4)
    finally:
        borrador_mod.OUT_PATH = original_out

    base = Document(str(temp_borrador))

    def extract_body_elements(source_doc, start_marker: str, end_marker: str | None = None) -> list:
        copying_el = False
        copied = []
        for child in list(source_doc.element.body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            line = ""
            if tag == "p":
                texts = [t.text or "" for t in child.iter(qn("w:t"))]
                line = "".join(texts).strip()
                if line.startswith(start_marker):
                    copying_el = True
                if end_marker and line.startswith(end_marker):
                    break
            elif tag == "tbl" and not copying_el:
                continue
            if copying_el:
                copied.append(deepcopy(child))
        return copied

    def insert_body_elements(target_doc, elements: list, insert_at: int = 0) -> None:
        target_body = target_doc.element.body
        for offset, child in enumerate(elements):
            target_body.insert(insert_at + offset, child)

    # 1) Cuerpo en orden 1→6→Referencias (documento intermedio con API consistente)
    doc = Document()
    style_base(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    insert_body_elements(doc, extract_body_elements(base, "Capitulo 1", "Referencias bibliograficas"), insert_at=0)
    doc.add_page_break()
    add_chapter_5_doctoral(doc, p, heading, add_table, status_note)
    doc.add_page_break()
    add_chapter_6_doctoral(doc, p, heading, bullet, add_table)
    doc.add_page_break()
    insert_body_elements(doc, extract_body_elements(base, "Referencias bibliograficas"), insert_at=len(doc.element.body))

    # 2) Materiales previos insertados al inicio (portada, resumen, indice)
    front = Document()
    style_base(front)
    fecha = dt.date.today().strftime("%d de %B de %Y")
    for _ in range(2):
        front.add_paragraph()
    p(front, "UNIVERSIDAD NACIONAL DE INGENIERIA (UNI)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, color=ACCENT)
    p(front, "Escuela de Posgrado — Doctorado en Ingenieria", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=GREY)
    p(front, "Inteligencia Artificial aplicada a Sistemas Electricos Inteligentes", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    for _ in range(2):
        front.add_paragraph()
    p(
        front,
        "MULTI-AGENTE DE APRENDIZAJE POR REFUERZO PROFUNDO PARA LA GESTION "
        "COORDINADA DE FLEXIBILIDAD ENERGETICA, EMISIONES DE CARBONO Y COSTOS "
        "ENERGETICOS EN COMUNIDADES INTELIGENTES",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=16,
        color=ACCENT,
    )
    p(
        front,
        "Caso de estudio: SEAI Iquitos — 17 edificios reales (2023-2025)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        size=11,
    )
    for _ in range(3):
        front.add_paragraph()
    p(front, "TESIS PARA OPTAR EL GRADO ACADEMICO DE DOCTOR EN INGENIERIA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
    p(front, "Autor: Mac Tapia (mac.tapia.c@uni.pe)", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    p(front, "Asesor: [por definir]", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    p(front, f"Lima / Iquitos, Peru — {fecha}", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=GREY)
    front.add_page_break()
    add_dedicatoria_agradecimientos(front, p, heading)
    add_resumen_doctoral(front, p, heading)
    heading(front, "Indice", 1)
    add_toc(front)
    front.add_page_break()

    front_elements = [deepcopy(child) for child in front.element.body]
    for offset, child in enumerate(reversed(front_elements)):
        doc.element.body.insert(offset, child)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    SKILL_COPY.write_bytes(OUT_PATH.read_bytes())

    if temp_borrador.exists():
        temp_borrador.unlink()

    checks = verify_doctoral_docx(OUT_PATH)
    print(f"OK -> {OUT_PATH}")
    print(f"OK -> {SKILL_COPY}")
    print(f"Verificacion: complete={checks['complete']} tables={checks['tables_count']} images={checks['images_count']}")
    for sec, ok in checks["sections"].items():
        print(f"  seccion {sec}: {'OK' if ok else 'FALTA'}")
    if not checks["complete"]:
        print("AVISO: revisar secciones faltantes antes de sustentacion.")
    return OUT_PATH


if __name__ == "__main__":
    build_doctoral()
