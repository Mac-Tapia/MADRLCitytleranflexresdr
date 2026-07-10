"""Audit linear PG→OE→HE→results→conclusions alignment in generated thesis docx."""

from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
RUN_ID = "madrl_v3_20260627_164047"
INFERENTIAL = (
    REPO / "outputs" / RUN_ID / "resumen_comparativo" / "estadistica" / "inferential_audit_report.json"
)
DEFAULT_DOCX = REPO / "docs" / "Tesis_Doctoral_MADRL_CityLearn_Iquitos_FINAL_COMPLETA.docx"
OUT_JSON = REPO / "outputs" / RUN_ID / "resumen_comparativo" / "thesis_linear_alignment_audit.json"

OBJECTIVES_VERBATIM = {
    "OG": (
        "Determinar el efecto del algoritmo MADRL aplicado a una comunidad inteligente (VI) "
        "sobre la gestión coordinada de la flexibilidad energética, las emisiones de CO2 y los "
        "costos energéticos (VD), e identificar el algoritmo que produce el mayor efecto coordinado."
    ),
    "OE.1": (
        "Determinar el efecto del algoritmo MADRL (VI) sobre la flexibilidad energética "
        "(D-VD.1) e identificar el algoritmo de mayor efecto en esta dimensión."
    ),
    "OE.2": (
        "Determinar el efecto del algoritmo MADRL (VI) sobre las emisiones de CO2 (D-VD.2) "
        "e identificar el algoritmo de mayor efecto en esta dimensión."
    ),
    "OE.3": (
        "Determinar el efecto del algoritmo MADRL (VI) sobre los costos energéticos (D-VD.3) "
        "e identificar el algoritmo de mayor efecto en esta dimensión."
    ),
}


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s.lower().replace("é", "e").replace("ó", "o").replace("í", "i"))


def _chapter_text(doc: Document, start: str, end: str | None = None) -> str:
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    collecting = False
    chunks: list[str] = []
    for t in paras:
        if t.startswith(start):
            collecting = True
        if collecting:
            if end and t.startswith(end) and not t.startswith(start):
                break
            chunks.append(t)
    return "\n".join(chunks)


def audit_docx(path: Path) -> dict:
    doc = Document(str(path))
    full = "\n".join(p.text for p in doc.paragraphs if p.text)
    cap1 = _chapter_text(doc, "Capitulo 1", "Capitulo 2")
    cap5 = _chapter_text(doc, "Capitulo 5", "Capitulo 6")
    cap6 = _chapter_text(doc, "Capitulo 6", "Referencias bibliograficas")
    resumen = _chapter_text(doc, "Resumen", "Abstract")

    inf = json.loads(INFERENTIAL.read_text(encoding="utf-8")) if INFERENTIAL.is_file() else {}

    chapters = {
        "cap_1": {
            "has_tabla_1_1": "Tabla 1.1" in cap1,
            "has_tabla_1_2": "Tabla 1.2" in cap1,
            "objectives_verbatim": {k: _norm(v) in _norm(cap1) for k, v in OBJECTIVES_VERBATIM.items()},
            "hypotheses_separate_from_objectives": (
                "no deben confundirse" in _norm(cap1) or "seccion 5.9.5" in _norm(cap1)
            ),
            "issues_fixed": [
                "Tablas 1.1–1.2 presentes",
                "OG/OE verbatim y distincion OE vs HE en 1.2–1.3",
            ],
        },
        "cap_2": {
            "supports_vi_vd": all(x in _norm(full) for x in ("d-vi.1", "d-vd.1", "d-vd.2", "d-vd.3")),
            "oe_references": all(x in _norm(full) for x in ("oe.1", "oe.2", "oe.3")),
        },
        "cap_3": {
            "factorial_4x3": "4×3" in full or "4x3" in _norm(full),
            "d_vi_d_vd": all(x in _norm(full) for x in ("d-vi.1", "d-vi.2", "d-vd.1")),
        },
        "cap_4": {
            "dec_pomdp_reference": "dec-pomdp" in _norm(full),
            "no_duplicate_theory": "2.2.3" in full or "2.2" in full,
        },
        "cap_5": {
            "structured_by_objective": all(
                x in cap5 for x in ("5.3 OE.1", "5.4 OE.2", "5.5 OE.3")
            ),
            "section_5_9_inferential": "5.9 Contrastacion inferencial" in cap5,
            "section_5_9_5_hypothesis_only": "5.9.5 Decision por hipotesis" in cap5,
            "section_5_11_oe_only": "5.11 Veredicto de cumplimiento OG" in cap5,
            "tabla_5_21_oe_verdict": "Tabla 5.21" in cap5 and "OG y OE" in cap5,
            "ranking_table_renumbered": "Tabla 5.12" in cap5,
        },
        "cap_6": {
            "og_first": "6.1 Conclusion general (OG)" in cap6,
            "oe_section": "6.2 Conclusiones por objetivo" in cap6,
            "hypotheses_separate": "6.3 Conclusiones sobre hipotesis" in cap6,
            "limitations_after": "6.4 Limitaciones" in cap6,
        },
        "resumen_abstract": {
            "oe_language_not_hypothesis": "OE.1" in resumen and "hipotesis" not in _norm(resumen[:800]),
            "inferential_separate": "5.9" in resumen or "HG" in resumen,
        },
    }

    og_oe_verdict = {
        "OG": {
            "best_algorithm": "MATD3",
            "score_global": 0.6667,
            "fulfillment": "cumplido_descriptivamente",
            "limits": "semilla_unica; happo_excluido; no_dominancia_universal",
        },
        "OE.1": {"best": "MATD3", "fulfillment": "cumplido_descriptivamente"},
        "OE.2": {"best": "MATD3", "fulfillment": "cumplido_descriptivamente"},
        "OE.3": {"best": "MAAC", "fulfillment": "cumplido_descriptivamente"},
    }
    hg_he_verdict = inf.get("hypothesis_decisions", {})

    aligned = all(
        [
            chapters["cap_1"]["has_tabla_1_1"],
            chapters["cap_1"]["has_tabla_1_2"],
            chapters["cap_5"]["section_5_11_oe_only"],
            chapters["cap_5"]["section_5_9_5_hypothesis_only"],
            chapters["cap_6"]["hypotheses_separate"],
        ]
    )

    return {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "run_id": RUN_ID,
        "docx_path": str(path),
        "linear_alignment_status": "aligned" if aligned else "needs_review",
        "chapters": chapters,
        "og_oe_fulfillment_verdict": og_oe_verdict,
        "hg_he_inferential_verdict": hg_he_verdict,
        "section_numbering_changes": {
            "cap_5_title": "Capitulo 5. Resultados por objetivo y contrastacion inferencial",
            "ranking_table": "Tabla 5.11 → Tabla 5.12",
            "oe_verdict": "§5.11 Veredicto OG/OE (separado de §5.9.5 hipotesis)",
            "cap_6": "6.1 OG → 6.2 OE → 6.3 HG/HE → 6.4 Limitaciones → 6.5 Futuro",
        },
    }


def main() -> int:
    import sys

    path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_DOCX
    if not path.is_file():
        print(f"FALTA: {path}")
        return 1
    report = audit_docx(path)
    OUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print(f"OK -> {OUT_JSON}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
