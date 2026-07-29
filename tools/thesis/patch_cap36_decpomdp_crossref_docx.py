#!/usr/bin/env python3
"""Remisiones ligeras Dec-POMDP Cap.3/Cap.6 en los 2 Word canónicos.

No re-axiomatiza Cap.1/5. Idempotente.
"""
from __future__ import annotations

import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

_THESIS_DIR = Path(__file__).resolve().parent
REPO = _THESIS_DIR.parents[1]
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))

from thesis_word_canons import DOCS, INFORME, TESIS  # noqa: E402

REPORT = DOCS / "CAP36_DECPOMDP_CROSSREF_PATCH_REPORT_2026-07-29.json"
BACKUP_DIR = REPO / "outputs" / "_word_backups"

CAP3_NEEDLE = "permite modelar el problema como Dec-POMDP"
CAP3_ADD = (
    " Las dimensiones locales medidas (d_oi en [54, 327], d_ai = 2 + n_i^ch en [5, 44], "
    "estado CTDE d_s = 1 856) constan en Cap. 2 Tabla 2.A y se operacionalizan en Cap. 4."
)

CAP6_MARKERS = (
    "benchmark reproducible Dec-POMDP",
    "Contribucion metodologica",
    "Contribución metodológica",
)
CAP6_REPLACEMENT_SNIPPET = (
    "benchmark reproducible Dec-POMDP/CTDE sobre 17 edificios del SEAI Iquitos "
    "(d_s = 1 856; d_oi en [54, 327]; d_ai en [5, 44]; r_team = 0,70) con cuatro "
    "algoritmos MADRL bajo CityLearn v3 (extension experimental de tesis; Caps. 2 y 4)"
)


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def patch_doc(path: Path) -> dict:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / (
        f"{path.stem}_before_cap36_xref_{datetime.now().strftime('%Y%m%d_%H%M%S')}{path.suffix}"
    )
    shutil.copy2(path, backup)
    doc = Document(str(path))
    changes: list[str] = []

    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        # Cap 3 dataset paragraph
        if CAP3_NEEDLE in t and "Tabla 2.A" not in t:
            set_paragraph_text(p, t.rstrip().rstrip(".") + "." + CAP3_ADD)
            changes.append(f"cap3_p{i}")

        # Cap 6 methodological contribution
        low = t.lower()
        if "dec-pomdp/ctde sobre 17" in low or (
            "benchmark reproducible" in low and "dec-pomdp" in low and "17 edificios" in low
        ):
            if "tabla 2.a" in low or "d_s = 1 856" in t or "d_s=1 856" in t.replace(" ", ""):
                continue
            if "d_s = 1 856" in t or "54, 327" in t or "[54, 327]" in t:
                continue
            # Replace the contribution clause in-place when present
            new = t
            # Spanish variants without accents in Word
            pat = re.compile(
                r"benchmark reproducible Dec-POMDP/CTDE sobre 17 edificios del SEAI Iquitos"
                r"(?: con cuatro algoritmos MADRL bajo CityLearn v3"
                r"(?: \(extensi[oó]n experimental de tesis\))?)?",
                re.I,
            )
            if pat.search(new):
                new = pat.sub(CAP6_REPLACEMENT_SNIPPET, new, count=1)
                if new != t:
                    set_paragraph_text(p, new)
                    changes.append(f"cap6_p{i}")

        # Appendix diagram interpretation: ensure formula mention if dims present but no Tabla 2.A
        if "Interpretación de la figura" in t or "Interpretacion de la figura" in t:
            if "54" in t and "327" in t and "1 856" in t.replace(" ", " ") and "Tabla 2.A" not in t:
                if "Caps. 2 y 4" not in t and "Cap. 2" not in t:
                    set_paragraph_text(
                        p,
                        t.rstrip().rstrip(".")
                        + " (valores alineados a Cap. 2 Tabla 2.A y Cap. 4 §4.2).",
                    )
                    changes.append(f"fig_p{i}")

    if not changes:
        backup.unlink(missing_ok=True)
        return {"path": str(path), "ok": True, "skipped": True, "reason": "nothing to patch or already xref"}

    doc.save(str(path))
    return {"path": str(path), "ok": True, "backup": str(backup), "changes": changes}


def main() -> int:
    results = []
    for path in (TESIS, INFORME):
        if path.is_file():
            results.append(patch_doc(path))
    report = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "policy": "Cap.3/6 remision ligera; Cap.1/5 sin re-axiomatizar; Cap.2/4 ya cerrados",
        "results": results,
    }
    REPORT.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0 if results and all(r.get("ok") for r in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
